"""Customer Requirements Document (CRD) — the Plan-phase requirements-capture instrument (.docx).

NEW-V3.23.156: the missing first link of the AS document chain (CRD → HLD → LLD → MOP → NRFU → PIR).
A CRD captures the customer's business / technical / operational requirements as testable, owned,
traceable statements (REQ-IDs), gathered through a requirements workshop, key-personnel interviews
and a questionnaire — and it is a LIVING document, revisited when the detailed design surfaces new
requirements.

Evidence discipline is the whole point of this generator: requirements are the CUSTOMER'S statements,
so the toolkit cannot invent them. What it CAN do is prime the instrument with the assessment's
evidence — the current-environment summary is real, the technical-requirement sections are gated by
what was actually observed (no wireless section for a fleet with no wireless evidence), and each
seeded technical requirement is a PROPOSAL derived from observed state ("preserve the N production
VLANs") that the customer must confirm, amend, or strike. Unconfirmed requirements are UNKNOWN, not
assumed — no confidence theater.

python-docx is OPTIONAL: imported inside the function so the package imports without it; a missing
library is a warning + skip, never a crash. Every snapshot read is defensive. Deterministic; no network.
"""
import logging
from datetime import datetime

from cisco_toolkit.analyze import vlan_inventory
from cisco_toolkit.docmeta import SEV_RANK as _SEV_RANK
from cisco_toolkit.docmeta import add_acceptance, add_document_control, add_excellence_front, add_glossary, add_inputs_required, add_table, add_toc
from cisco_toolkit.docmeta import as_dict as _as_dict, as_list as _as_list   # coerce truthy non-dict/list sections
from cisco_toolkit.textutils import _as_num, xml_safe, xml_safe_deep   # entry deep-sanitize of device text (audit-5) + fail-soft numeric coercion
from cisco_toolkit import ssot as _ssot_mod   # Law 1 accessors (canonical facts + segmentation posture)

logger = logging.getLogger(__name__)


def _evidence_facts(snap: dict) -> dict:
    """Pull the evidence the CRD primes its sections with — all defensive reads of known shapes."""
    devices = _as_dict(snap.get("devices"))
    ifaces = _as_dict(snap.get("interfaces"))
    endpoints = 0
    for host, ports in ifaces.items():
        for p, d in _as_dict(ports).items():
            # PER-PORT coercion, not `d or {}`: the container is guarded but a single truthy non-dict
            # PORT value (interfaces={"acc1": {"Gi1/0/5": 5}} in an uploaded snapshot) survives `or {}`
            # and crashes d.get("switchport_mode") below -> stored 500. Same class as the section guards.
            d = _as_dict(d)
            # str()-coerce every device string-field before .strip()/.lower() (audit-6 finding-#3 class): a
            # wrong-typed leaf (a truthy dict/list/number in an uploaded snapshot) would otherwise .strip() -> 500.
            if str(d.get("switchport_mode") or "").lower() == "access" and str(d.get("end_host_mac") or "").strip():
                endpoints += 1
    # SSOT (Law 1): the segmentation posture has ONE owner (analyze.compute_segmentation ->
    # snap['segmentation']), read via ssot.segmentation_facts. The inline recount this replaced asked a
    # different question under the same label — it treated EVERY non-default VRF on the box (mgmt0's
    # management VRF, a vPC keepalive VRF) as observed segmentation, so REQ-T-SEC-001 told the customer
    # to "preserve the observed segmentation: VRF(s) Mgmt-vrf, VPC, management, mgmtVrf" on a fabric the
    # owner reports FLAT (1 VRF bucket across 231 gateway SVIs) — a requirement written off a fact that
    # segments no user traffic. It also counted an ACL on ANY port carrying an svi_ip, not just a VlanN
    # SVI, so its numerator could exceed the design doc's and the archreview's for the same estate.
    _segf = _ssot_mod.segmentation_facts(snap)
    gw_vrfs = list(_segf.get("gateway_vrfs") or [])
    other_vrfs = list(_segf.get("other_vrfs") or [])
    n_acl_svis = _as_num(_segf.get("n_with_acl"))
    n_seg_gateways = _as_num(_segf.get("n_gateways"))
    # Dual-homed endpoints: the CANONICAL redundancy-bearing count is the engine's
    # endpoint_dependencies.dual_homed (host MAC observed on two switches) — the SAME source the
    # design blueprint's preserve-dual-homed-endpoints decision reads — so the CRD and the HLD agree
    # on one number instead of the CRD reporting a looser per-port dual_connection tally (A1 SSOT fix).
    dual = len(_as_list(_as_dict(snap.get("endpoint_dependencies")).get("dual_homed")))
    rn = _as_dict(snap.get("routing_neighbors"))
    # str() the protocol KEY before .upper(): JSON object keys are always strings, so this is latent on
    # the upload path, but the CLI hands write_crd_docx the IN-PROCESS snapshot dict (never round-tripped
    # through JSON), where a non-str key would AttributeError. Identical output for every string key.
    protos = sorted({str(p).upper() for host in rn for p, nbrs in _as_dict(rn.get(host)).items() if nbrs})
    l3f = _as_list(snap.get("l3_forwarding"))
    # Real FHRP only. The parser writes the literal string "none" when no HSRP/VRRP/GLBP is present,
    # and "none".strip() is truthy — so a bare truthiness test mislabels EVERY gateway VLAN as
    # FHRP-protected (a false-redundancy claim). Mirror the engine's canonical gate
    # `(fhrp or "none") != "none"` (analyze.py) so the CRD agrees with the FHRP Consistency sheet.
    fhrp_vlans = sorted({str(r.get("vlan")) for r in l3f
                         if isinstance(r, dict) and (r.get("fhrp", "none") or "none") != "none"})
    svc = _as_dict(snap.get("service_map"))
    services = [s for s in _as_list(svc.get("services")) if isinstance(s, dict)]
    mc = _as_dict(svc.get("multicast"))
    mcast_active = any((
        _as_num(mc.get("active_interfaces")), _as_num(mc.get("active_switch_count")),
        # coerce, don't `or []`/`or {}`: a truthy non-list/non-dict inside the (already guarded) multicast
        # block reaches len() / .values() / .get() and aborts the whole CRD (classified_groups=5 -> len(5),
        # ptp=5 -> 5.values(), ptp={"x": 5} -> 5.get("operational")).
        len(_as_list(mc.get("classified_groups"))), len(_as_list(mc.get("igmp_queriers"))),
        any(_as_dict(v).get("operational") for v in _as_dict(mc.get("ptp")).values()),
    ))
    lc = _as_dict(_as_dict(snap.get("lifecycle_risk")).get("summary"))
    coll = _as_dict(_as_dict(snap.get("collection_completeness")).get("summary"))
    # str()-coerce the RANK LOOKUP KEY: an UNHASHABLE severity (punchlist=[{"severity": ["x"]}] in an
    # uploaded snapshot) makes dict.get itself raise TypeError: unhashable type -> the whole CRD aborts.
    # Behaviour-identical for real input: every _SEV_RANK key is a string, so str(x) hits the same entry
    # for a string severity and misses (-> the 5 default) for everything else, exactly as before.
    punch = sorted([i for i in _as_list(snap.get("punchlist")) if isinstance(i, dict)],
                   key=lambda i: _SEV_RANK.get(str(i.get("severity")), 5))
    # isinstance-guard the canonical scale once: a TRUTHY non-dict executive_brief (malformed/slimmed snapshot)
    # slips through `or {}` and crashes .get('scale') -> the whole CRD silently aborts (audit L4).
    _eb = snap.get("executive_brief"); _eb_scale = _eb.get("scale") if isinstance(_eb, dict) else None
    _eb_scale = _eb_scale if isinstance(_eb_scale, dict) else {}
    return {
        # canonical-first (SSOT) like the n_vlans / n_endpoints reads below -- not a raw len(devices) recount,
        # which is the device-count drift seam ssot.py was created to eliminate (design.py/engagement.py do this)
        "devices": devices,
        "n_devices": _eb_scale.get("n_devices") or len(devices),
        # SSOT: the published canonical VLAN count first (the one source design/explorer/webapp/workbook
        # reconcile to), with the local vlan_inventory recount only as the pre-brief fallback (mirrors
        # n_endpoints below + design.py's A5 canonical-first read — closes the recompute drift seam).
        "n_vlans": _eb_scale.get("n_vlans") or len(vlan_inventory(snap)),
        "endpoints": endpoints, "dual": dual, "protos": protos, "fhrp_vlans": fhrp_vlans,
        # gateway-carrying VRFs vs configured-elsewhere VRFs: two DIFFERENT facts, kept apart (see above)
        "gw_vrfs": gw_vrfs, "other_vrfs": other_vrfs,
        "n_acl_svis": n_acl_svis, "n_seg_gateways": n_seg_gateways, "services": services,
        "mcast": mc, "mcast_active": mcast_active, "lifecycle": lc, "coll": coll, "punch": punch,
        "n_l3": len(l3f),
        # canonical endpoint scale (the published single source) with the access-port tally as fallback
        "n_endpoints": _eb_scale.get("n_endpoints") or endpoints,
    }


def _requirements_overlay(bp: dict) -> dict:
    """Read the customer's supplied requirements register OFF the canonical blueprint (SSOT).

    `design_blueprint.requirements_model` is the ONE place the `--requirements` register is echoed
    into the snapshot (design_advisor._requirements_model): `provided` (was a register supplied?),
    `fields` ([{key, label, value}] — availability_tier / constraints / fabric_operating_model / …),
    and `open_questions`. The CRD reads that model rather than re-loading the register file, so the
    Constraints / Out-of-Scope sections agree with the HLD/deck/explorer instead of drifting. When no
    register was supplied every value is None and `provided` is False — the caller surfaces the
    engagement constraints as OPEN QUESTIONS, never inventing them (coverage-honesty).
    """
    rm = _as_dict(bp.get("requirements_model"))
    # str() the comprehension KEY: an unhashable `key` (fields=[{"key": ["a"]}] in an uploaded snapshot)
    # makes building this dict raise TypeError: unhashable type -> the whole CRD aborts. Identical for a
    # real register, whose keys are the string field names the three vals.get() lookups below use.
    vals = {str(f.get("key")): f.get("value")
            for f in _as_list(rm.get("fields")) if isinstance(f, dict) and f.get("key")}
    return {
        "provided": bool(rm.get("provided")),
        "constraints": [str(c).strip() for c in _as_list(vals.get("constraints")) if str(c).strip()],
        "availability_tier": (str(vals.get("availability_tier") or "").strip() or None),
        "fabric_operating_model": (str(vals.get("fabric_operating_model") or "").strip() or None),
    }


def write_crd_docx(output_path: str, snap_dict: dict, label: str) -> None:
    """Emit the Customer Requirements Document (.docx) to `output_path`. Fail-soft: a missing
    python-docx is a warning + skip, and every snapshot read is defensive so malformed input degrades
    to placeholders. A genuine render error still propagates to the caller — the CLI phase wrapper
    logs-and-continues, AssessHub turns it into a clean 500 (the old 'logged, never raised' claim was
    false: the only `try` here is the optional-import guard, and swallowing errors would be
    coverage-dishonest; the crash class it hid is now guarded at the read sites instead)."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor
    except ImportError:
        logger.warning("  CRD (DOCX) skipped: python-docx not installed "
                       "(pip install python-docx to enable the requirements-capture deliverable).")
        return

    snap = xml_safe_deep(snap_dict if isinstance(snap_dict, dict) else {})   # one bad device byte (noncharacter/surrogate) must not abort the docx
    label = xml_safe(label) if isinstance(label, str) else (str(label) if label is not None else "")
    from cisco_toolkit.brand_tokens import DOC_NAVY_RGB
    NAVY = RGBColor(*DOC_NAVY_RGB)
    GREY = RGBColor(0x59, 0x59, 0x59)
    RED = RGBColor(0xB0, 0x2A, 0x1E)
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    def _label_run(p, label_text, value, color=NAVY):
        r = p.add_run(label_text); r.bold = True; r.font.color.rgb = color
        p.add_run(" " + (str(value) if value not in (None, "") else "—"))

    def table(headers, rows, widths=None):
        # Delegates to the family's single table builder (docmeta.add_table).
        return add_table(doc, headers, rows, widths, fixed=False)

    ev = _evidence_facts(snap)
    _bp = snap.get("design_blueprint")        # isinstance-guard, not `or {}`: a truthy non-dict crashes .get below
    bp = _bp if isinstance(_bp, dict) else {}   # (audit-2 L5 -- design.py guarded this; crd.py was the gap)
    reg = _requirements_overlay(bp)             # the customer's WHY register, read off the canonical blueprint (SSOT)
    req_ids: list = []   # (req_id, origin) — feeds the traceability skeleton

    def _verify_method(rid):
        # how the requirement is PROVEN — Source is implicit in the REQ-ID class (B/T/O/D), so the
        # higher-value column to add is the verification method, kept consistent with the RTM (N5).
        rid = str(rid)
        if rid.startswith("REQ-D"):
            return "Design-driven NRFU (§10 + HLD §4)"
        if rid.startswith("REQ-T"):
            return "NRFU technical acceptance test"
        if rid.startswith("REQ-B"):
            return "Sponsor acceptance / sign-off"
        if rid.startswith("REQ-O"):
            return "Operational acceptance (Day-2 handover)"
        return "<NRFU test — detail in HLD/NRFU>"

    def req_table(rows):
        """A requirement-capture table; registers every REQ-ID for traceability, has each requirement
        DECLARE how it is proven (Verification), and classifies its strength with a BCP 14 / RFC 2119
        normative keyword (Class) rather than an ad-hoc H/M/L."""
        out = []
        for r in rows:
            req_ids.append(r[0])
            cls = "<MUST/SHOULD/MAY>" if str(r[3]).strip() in ("<H/M/L>", "") else r[3]
            out.append((r[0], r[1], r[2], cls, _verify_method(r[0]), r[4]))
        table(["REQ-ID", "Requirement (testable statement)", "Owner", "Class (RFC 2119)", "Verification",
               "Confirmed?"], out, widths=[0.85, 2.6, 0.75, 0.95, 1.4, 0.65])

    # ---- title page ----
    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Customer Requirements Document (CRD)"); tr.bold = True
    tr.font.size = Pt(26); tr.font.color.rgb = NAVY
    sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s2 = sub2.add_run("Plan-phase requirements capture — workshop / interview / questionnaire instrument")
    s2.font.size = Pt(14); s2.font.color.rgb = GREY
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(label); sr.font.size = Pt(13); sr.font.color.rgb = GREY
    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')}  ·  "
                 f"{ev['coll'].get('complete', ev['n_devices'])} of {ev['n_devices']} devices collected  ·  script {snap.get('script_version', '')}"
                 ).font.color.rgb = GREY
    status = doc.add_paragraph(); status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st = status.add_run("DRAFT TEMPLATE — requirements are the customer's statements; confirm, amend or "
                        "strike every seeded row in the workshop.")
    st.italic = True; st.font.color.rgb = RED
    doc.add_paragraph()
    note = doc.add_paragraph()
    _label_run(note, "How to use this CRD:",
               "Run one requirements workshop, interview the key personnel, and leave the questionnaire "
               "columns with the customer. Every requirement needs a REQ-ID, an owner, and a TESTABLE "
               "statement ('the network should be fast' is not a requirement; 'voice VLANs tolerate at "
               "most one 30-minute outage window' is). Seeded technical rows are PROPOSALS derived from "
               "the observed network — they are not requirements until the customer confirms them. This "
               "is a living document: revisit it when the detailed design surfaces new requirements.", GREY)
    doc.add_page_break()

    # ---- document control (AS-style front matter) ----
    add_document_control(
        doc, document="Customer Requirements Document (CRD)", label=label,
        engine_version=str(snap.get("script_version", "")), generated_at=snap.get("generated_at"),
        collected_at=snap.get("collected_at"),
        audience="The customer's project sponsor, network owner and operations lead (requirement "
                 "owners), and the delivery engineer facilitating the requirements workshop.",
        exclude=("crd",),
        extra_assumptions=(
            "Requirements are the customer's statements. The evidence-primed rows below are proposals "
            "derived from observed state and must be confirmed or corrected by the customer; an "
            "unconfirmed requirement is recorded as UNKNOWN, never assumed.",))
    doc.add_page_break()

    # ---- table of contents (shared field-code helper, V3.23.171) ----
    add_toc(doc)


    # Deliverable Excellence (DE-01): answer-first at-a-glance register + single-source-of-truth signal.
    add_excellence_front(doc, snap_dict)
    # Deliverable Excellence (DE-01): consolidated Inputs-Required register (Law 9). The CRD's whole
    # purpose is capturing what the customer must confirm, so it names that ask explicitly up front.
    add_inputs_required(doc, snap_dict, extra_rows=[
        ("Confirm every seeded requirement",
         "Every REQ-ID below is a PROPOSAL derived from the observed network — confirm, amend, or strike "
         "each in the requirements workshop; a seeded row is not a requirement until the customer confirms it.",
         "Customer (requirement owner)"),
        ("Complete the questionnaire fields",
         "Fill every <…> owner / class / confirmation field the questionnaire leaves with the customer.",
         "Customer / delivery"),
    ])
    # ===== 1. Engagement context =====
    doc.add_heading("1. Engagement Context", level=1)
    doc.add_paragraph(
        "Drivers, scope and decision owners — completed with the customer in the workshop. The "
        "evidence scope below is what the assessment actually observed; the engagement scope may be "
        "wider (name the gap explicitly if so).")
    table(["Field", "Value"], [
        ("Business drivers for the change", "<why now — e.g. EoL exposure, capacity, site move>"),
        ("Engagement scope (sites / domains)", "<sites and network domains in scope>"),
        ("Evidence scope (this assessment)", f"{ev['coll'].get('complete', ev['n_devices'])} of "
                                             f"{ev['n_devices']} devices collected, {ev['n_vlans']} VLANs, "
                                             f"{ev['n_endpoints']} evidenced endpoints"),
        ("Project sponsor", "<name, role>"),
        ("Network owner (requirement owner)", "<name, role>"),
        ("Operations lead", "<name, role>"),
        ("Decision process", "<who approves the design, the windows, and the acceptance>"),
    ], widths=[2.4, 4.3])

    # ===== 2. Current environment summary (evidence) =====
    doc.add_heading("2. Current Environment Summary (evidence)", level=1)
    doc.add_paragraph(
        "What the assessment observed — the factual baseline the requirements are written against. "
        "Full evidence is in the assessment workbook; every number here reconciles to it.")
    lc, coll = ev["lifecycle"], ev["coll"]
    table(["Fact", "Observed"], [
        ("Devices", ev["n_devices"]),
        ("VLANs in use", ev["n_vlans"]),
        ("Evidenced endpoints (access ports with a host MAC)", ev["endpoints"]),
        ("Dual-homed endpoints (host MAC on two switches)", ev["dual"]),
        ("Routing protocols observed", ", ".join(ev["protos"]) or "none (pure L2 fleet)"),
        ("Gateway SVIs / FHRP-protected VLANs", f"{ev['n_l3']} / {len(ev['fhrp_vlans'])}"),
        ("Non-default VRFs carrying a gateway SVI", ", ".join(ev["gw_vrfs"]) or "none (flat global table)"),
        ("Other non-default VRFs (no gateway SVI — segment no user traffic)",
         ", ".join(ev["other_vrfs"]) or "none"),
        ("Hardware past last-day-of-support (LDoS)", lc.get("n_past_ldos", "—")),
        ("Collection completeness", f"{coll.get('complete', '—')} complete / "
                                    f"{coll.get('partial', '—')} partial / "
                                    f"{coll.get('not_collected', '—')} not collected"),
    ], widths=[3.4, 3.3])
    if ev["punch"]:
        doc.add_paragraph("Known issues the requirements must take a position on (top of the "
                          "assessment punch-list):")
        table(["Severity", "Category", "Issue"],
              [(i.get("severity"), i.get("category") or "—", i.get("title") or "—")
               for i in ev["punch"][:8]], widths=[0.9, 1.4, 4.3])
        # Disclose the cut, and say so in the sentence that assigns the work. "Take a position on"
        # is the CRD's whole function: the workshop writes requirements against the issues printed
        # here, and §2 states no punch-list total anywhere else, so 8 severity-ranked rows out of
        # 1,805 (the [HISTORY-REDACTED] fleet; 216 Critical + 916 High) read as the complete known-issue set and the
        # requirements register silently inherits the other 1,797 as unaddressed.
        if len(ev["punch"]) > 8:
            doc.add_paragraph(
                f"…and {len(ev['punch']) - 8} further punch-list item(s) NOT listed above "
                f"({len(ev['punch'])} in total) — the table shows the 8 highest-severity items only. "
                "The requirements must take a position on the full set (carry forward, remediate, or "
                "explicitly risk-accept); work it from the Migration Punch-List workbook sheet.")

    # requirements-classification legend (BCP 14 / RFC 2119 + RFC 8174) — establishes the normative-keyword
    # convention before the first requirement table (N1/N8)
    doc.add_paragraph(
        "Requirements classification (BCP 14). Each requirement's strength is stated with an RFC 2119 / "
        "RFC 8174 normative keyword in the Class column — MUST / MUST NOT (an absolute requirement), "
        "SHOULD / SHOULD NOT (recommended; deviate only with a documented reason), MAY / OPTIONAL (truly "
        "discretionary) — and the keyword carries this meaning only when written in all capitals (RFC 8174). "
        "Confirm the class of each row in the workshop.")

    # ===== 3. Business requirements =====
    doc.add_heading("3. Business Requirements", level=1)
    doc.add_paragraph(
        "Success criteria and constraints, as testable statements with owners. These are pure "
        "customer inputs — the rows are prompts, not proposals.")
    req_table([
        ("REQ-B-001", "Outage tolerance per service class: <e.g. voice VLANs tolerate at most one "
                      "30-minute window; CCTV tolerates none>", "<owner>", "<H/M/L>", "<YES/AMEND>"),
        ("REQ-B-002", "Timeline and window constraints: <deadline, change-freeze periods, allowed "
                      "maintenance windows>", "<owner>", "<H/M/L>", "<YES/AMEND>"),
        ("REQ-B-003", "Compliance / policy constraints: <regulatory, security policy, vendor-support "
                      "requirements>", "<owner>", "<H/M/L>", "<YES/AMEND>"),
        ("REQ-B-004", "Definition of success: <the measurable end-state the sponsor will accept>",
         "<owner>", "<H/M/L>", "<YES/AMEND>"),
    ])

    # ===== 4. Technical requirements (per observed discipline) =====
    doc.add_heading("4. Technical Requirements", level=1)
    doc.add_paragraph(
        "One subsection per discipline the evidence shows in use — a discipline with no observed "
        "footprint gets no section (add one in the workshop if the TARGET state introduces it). "
        "Seeded rows are evidence-derived proposals.")

    doc.add_heading("4.1 Campus LAN / Layer 2", level=2)
    req_table([
        ("REQ-T-LAN-001", f"Preserve the {ev['n_vlans']} production VLAN IDs and names through the "
                          "migration (observed as-built).", "<owner>", "<H/M/L>", "<YES/AMEND>"),
        ("REQ-T-LAN-002", f"Maintain dual-homing for the {ev['dual']} dual-homed endpoint(s) (host "
                          "MAC present on two switches); no dual-homed endpoint may lose both legs "
                          "together.", "<owner>", "<H/M/L>", "<YES/AMEND>"),
    ])

    if ev["protos"] or ev["n_l3"]:
        doc.add_heading("4.2 Layer 3 & routing", level=2)
        req_table([
            ("REQ-T-L3-001",
             (f"Preserve gateway redundancy on the {len(ev['fhrp_vlans'])} FHRP-protected VLAN(s); "
              "single-gateway VLANs are remediated, not carried forward." if ev["fhrp_vlans"]
              else f"No first-hop redundancy (HSRP/VRRP/GLBP) was observed on any of the {ev['n_l3']} "
                   "gateway SVI instance(s) — every gateway is single-homed; evaluate introducing FHRP "
                   "so VLANs are not carried forward without gateway redundancy."),
             "<owner>", "<H/M/L>", "<YES/AMEND>"),
            ("REQ-T-L3-002", "Maintain the observed routing adjacencies ("
                             + (", ".join(ev["protos"]) or "static/connected only")
                             + ") and their policy through the migration.",
             "<owner>", "<H/M/L>", "<YES/AMEND>"),
        ])

    if ev["mcast_active"]:
        doc.add_heading("4.3 Multicast & timing", level=2)
        # `classified_groups` is coerced AGAIN here, per ELEMENT: [5] passes the mcast_active length gate
        # above (len == 1), so this section renders and only then does groups[0].get(...) crash — the
        # activity check is not a type check, and the crash is two reads downstream of it.
        groups = _as_list(ev["mcast"].get("classified_groups"))
        g0 = _as_dict(groups[0]) if groups else {}
        gname = g0.get("name") or g0.get("group") or "the observed groups"
        req_table([
            ("REQ-T-MC-001", f"Preserve multicast delivery for the {len(groups)} classified group(s) "
                             f"(e.g. {gname}) including snooping/querier behaviour per VLAN.",
             "<owner>", "<H/M/L>", "<YES/AMEND>"),
        ])

    if ev["gw_vrfs"] or ev["n_acl_svis"] or ev["other_vrfs"]:
        doc.add_heading("4.4 Segmentation & security", level=2)
        # "Preserve the observed segmentation" must name what actually segments USER traffic. A
        # management/keepalive VRF carries no gateway SVI, so listing it here wrote a carry-forward
        # requirement off a fabric the engine reports FLAT.
        req_table([
            ("REQ-T-SEC-001",
             ("Preserve the observed segmentation: VRF(s) " + ", ".join(ev["gw_vrfs"])
              if ev["gw_vrfs"] else
              f"NO gateway SVI sits in a non-default VRF (all {ev['n_seg_gateways']} are in the global "
              "table) — there is no VRF segmentation to preserve; the target must introduce it"
              if ev["n_seg_gateways"] else
              # coverage-honest: no gateway tier was observed, so neither verdict is assertable
              "[NOT OBSERVED] — no gateway SVI was collected, so the L3 segmentation posture could "
              "not be assessed (this is a blind spot, not a flat fabric)")
             + f"; {ev['n_acl_svis']} gateway SVI(s) carry ACLs that must be "
               "carried forward or consciously redesigned."
             + (f" ({len(ev['other_vrfs'])} non-default VRF(s) — {', '.join(ev['other_vrfs'][:6])} — "
                "exist but carry no gateway SVI; they are management/infrastructure separation, not "
                "user-traffic segmentation.)" if ev["other_vrfs"] else ""),
             "<owner>", "<H/M/L>", "<YES/AMEND>"),
        ])

    if ev["services"]:
        doc.add_heading("4.5 Services & applications", level=2)
        cats = sorted({str(s.get("category") or "").strip() for s in ev["services"] if s.get("category")})
        req_table([
            ("REQ-T-SVC-001", "End-to-end continuity for the detected service categories ("
                              + (", ".join(cats[:6]) or "see service map")
                              + ") — each gets an NRFU end-to-end test.",
             "<owner>", "<H/M/L>", "<YES/AMEND>"),
        ])

    # ===== 5. Operational requirements =====
    doc.add_heading("5. Operational Requirements", level=1)
    req_table([
        ("REQ-O-001", "Management & monitoring continuity: <SNMP/syslog/flow reach the same "
                      "collectors throughout; no monitoring blackout longer than X>",
         "<owner>", "<H/M/L>", "<YES/AMEND>"),
        ("REQ-O-002", "Change windows & freeze calendar: <when work may happen>",
         "<owner>", "<H/M/L>", "<YES/AMEND>"),
        ("REQ-O-003", "Support model & knowledge transfer: <who operates the target network; what "
                      "handover/training is required before acceptance>",
         "<owner>", "<H/M/L>", "<YES/AMEND>"),
    ])

    # ===== 6. Future plans & growth =====
    doc.add_heading("6. Future Plans & Growth", level=1)
    # evidence-anchored scale baseline — growth targets are measured against observed reality (N3)
    table(["Scale dimension", "Today (observed)", "Target horizon", "Headroom / note"], [
        ("Switches in scope", ev["n_devices"], "<confirm>",
         "replacement / refresh tracked in the lifecycle plan"),
        ("Production VLANs", ev["n_vlans"], "<confirm>", "new segments add to the L2/L3 + addressing plan"),
        ("Evidenced endpoints", ev["n_endpoints"], "<confirm>",
         "PoE / port / multicast capacity must absorb the growth"),
    ], widths=[1.7, 1.3, 1.2, 2.5])
    doc.add_paragraph(
        "Planned initiatives the design must not block — capacity horizon, new sites or services, "
        "technology directions. Captured in the workshop; each becomes a requirement row or an "
        "explicit exclusion.")
    table(["Horizon", "Plan", "Design implication"],
          [("<12 months>", "<e.g. new building / IP camera expansion>", "<ports, PoE, multicast>"),
           ("<36 months>", "<e.g. platform refresh, segmentation program>", "<to be assessed>")],
          widths=[1.1, 3.0, 2.6])

    # ===== 7. Constraints & assumptions =====
    # The fixed boundaries the design must respect. Register-supplied constraints (the customer's WHY,
    # read off the canonical blueprint) are CONFIRMED; evidence-derived constraints are proposals gated on
    # observed state; when NO register was supplied the standard engagement constraints are surfaced as
    # OPEN QUESTIONS (coverage-honesty — a constraint the customer never stated is never invented).
    doc.add_heading("7. Constraints & Assumptions", level=1)
    doc.add_paragraph(
        "The fixed boundaries the design MUST respect — the non-negotiables the target architecture is "
        "built inside. A constraint is stronger than a requirement: it is not up for trade-off. Rows "
        "marked CONFIRMED come from the supplied requirements register; EVIDENCE rows are proposals "
        "derived from the assessment; OPEN QUESTION rows are constraints the engagement typically carries "
        "but that this assessment cannot observe — answer them in the workshop (the engine never invents "
        "a constraint the customer has not stated).")
    con_rows = []   # (CON-ID, constraint, source)
    _ci = 0
    if reg["provided"] and reg["constraints"]:
        for c in reg["constraints"]:
            _ci += 1
            con_rows.append((f"CON-{_ci:03d}", c, "CONFIRMED (requirements register)"))
    if reg["fabric_operating_model"]:
        _ci += 1
        _fm = {"nxos-evpn": "Standalone NX-OS VXLAN BGP-EVPN fabric (Nexus Dashboard / NDFC-managed), "
                            "not a Cisco ACI / APIC policy fabric.",
               "aci": "Cisco ACI (APIC-controlled application-centric policy fabric), not a standalone "
                      "NX-OS VXLAN-EVPN fabric."}.get(reg["fabric_operating_model"],
                                                      f"Target fabric operating model: {reg['fabric_operating_model']}.")
        con_rows.append((f"CON-{_ci:03d}", _fm, "CONFIRMED (requirements register)"))
    if reg["availability_tier"]:
        _ci += 1
        con_rows.append((f"CON-{_ci:03d}",
                         f"Target availability tier is '{reg['availability_tier']}' — the design is "
                         "right-sized to this tier, not 5x9 everywhere.",
                         "CONFIRMED (requirements register)"))
    # Evidence-derived constraint: EoL / past-LDoS hardware forces replacement (from lifecycle_risk).
    _ldos = ev["lifecycle"].get("n_past_ldos")
    _eos = ev["lifecycle"].get("n_past_eos")
    if isinstance(_ldos, int) and _ldos > 0:
        _ci += 1
        con_rows.append((f"CON-{_ci:03d}",
                         f"{_ldos} device(s) are past last-day-of-support (LDoS) — these are forced "
                         "replacements the migration cannot carry forward on the current hardware.",
                         "EVIDENCE (lifecycle risk)"))
    elif isinstance(_eos, int) and _eos > 0:
        _ci += 1
        con_rows.append((f"CON-{_ci:03d}",
                         f"{_eos} device(s) are past end-of-sale (EoS) — plan the hardware refresh within "
                         "the vendor-support horizon.", "EVIDENCE (lifecycle risk)"))
    if not (reg["provided"] and reg["constraints"]):
        # No register (or an empty constraints list): surface the standard engagement constraints as OPEN
        # QUESTIONS rather than asserting them — the coverage-honest default (mirrors §4's open questions).
        for q in (
            "Fabric operating model — standalone NX-OS VXLAN BGP-EVPN (NDFC-managed) vs Cisco ACI "
            "(APIC-managed)? The whole target design pivots on this; confirm it in the workshop.",
            "Out-of-band management must NOT transit the fabric it manages — confirm the OOB gear is "
            "physically separate.",
            "Overlapping legacy VLAN IDs across zones — confirm whether port-VLAN translation to a new "
            "fabric VLAN/VNI is required at the migration handoff.",
            "Installed-base reuse — which existing hardware (if any) must be carried forward vs replaced?",
            "Cutover model — multi-window, NRFU-gated (fabric fully deployed and tested before any "
            "workload migration)?",
        ):
            _ci += 1
            con_rows.append((f"CON-{_ci:03d}", q, "OPEN QUESTION — no requirements register supplied"))
    if not con_rows:   # register provided-but-empty AND no lifecycle evidence: still ask, never blank
        con_rows.append(("CON-001", "<capture the engagement's fixed constraints in the workshop>",
                         "OPEN QUESTION"))
    table(["CON-ID", "Constraint / assumption", "Source"], con_rows, widths=[0.9, 4.4, 1.4])

    # ===== 8. Out of scope =====
    # The explicit boundary statement — the standard CRD guardrail. Register-seeded when supplied; the
    # default list is marked "confirm with customer" so a boundary is never silently assumed.
    doc.add_heading("8. Out of Scope", level=1)
    doc.add_paragraph(
        "What this engagement does NOT cover — the explicit boundary that stops scope creep and sets the "
        "customer's expectations. Each line is a proposed exclusion to CONFIRM, AMEND or STRIKE in the "
        "workshop; anything the customer moves in-scope becomes a requirement row above.")
    oos_default = [
        ("Application / server / storage changes",
         "The assessment and design are network-layer (L1–L4); application, compute and storage "
         "remediation are out of scope unless explicitly added."),
        ("Physical cabling & structured-cabling works",
         "New fibre/copper runs, patching and facilities work are assumed pre-existing or handled by a "
         "separate work-order."),
        ("Endpoint / client-side configuration",
         "Host, camera and appliance configuration is the customer's responsibility beyond the switchport."),
        ("Security policy authorship",
         "Firewall rule-set and security-policy design beyond preserving the observed segmentation."),
        ("WAN / Internet-edge & provider circuits",
         "Carrier services and the Internet edge are out of scope unless named in the engagement scope."),
        ("Not-collected devices",
         "Any device not in the collected evidence set is outside this assessment's coverage — its role "
         "and dependencies are an explicit unknown, not assumed."),
    ]
    table(["Excluded area", "Boundary statement (confirm with customer)"],
          oos_default, widths=[2.1, 4.6])

    # Design-driven requirements (REQ-D) are derived from the canonical blueprint; register their IDs in
    # the traceability set BEFORE the RTM (§9) renders so they are NOT orphaned from the matrix (a
    # requirement that traces to nothing is the exact review defect the RTM forbids). Their detail is §10.
    # `bp` is isinstance-guarded above, but its CHILDREN are not: `or []` lets a truthy non-list through
    # to `for d in 5` -> TypeError. Coerce, then keep the existing per-element isinstance filter.
    bp_decisions = [d for d in _as_list(bp.get("decisions")) if isinstance(d, dict)]
    rec = [d for d in bp_decisions if d.get("status") == "recommended"]
    req_d = [(f"REQ-D-{i:03d}", d) for i, d in enumerate(rec, 1)]
    req_ids.extend(rid for rid, _ in req_d)

    # ===== 9. Requirements Traceability Matrix (RTM) =====
    # The Cisco Advanced-Services instrument: each captured REQ-ID traced FORWARD through the delivery
    # chain — HLD section → LLD object → MOP step → NRFU test case → evidence. The design/MOP/NRFU are
    # authored downstream, so every forward cell is an HONEST PLACEHOLDER ("to be traced in <deliverable>")
    # rather than a fabricated section number — this closes the requirement→design→MOP→NRFU→evidence chain
    # the AS standard mandates without inventing references that do not yet exist.
    doc.add_heading("9. Requirements Traceability Matrix (RTM)", level=1)
    doc.add_paragraph(
        "The forward-traceability instrument the Cisco Advanced-Services standard mandates: every captured "
        "requirement is traced forward to the design object that satisfies it, the MOP step that "
        "implements it, and the NRFU test case that proves it — so no requirement is orphaned (traced to "
        "nothing) and no design choice is unjustified (tracing to no requirement). The design, MOP and "
        "NRFU are authored downstream; each forward cell below is a PLACEHOLDER to complete when that "
        "deliverable is produced, never a fabricated reference. This is the spine of the "
        "requirement→design→MOP→NRFU→evidence chain.")

    def _rtm_row(rid):
        # REQ-D requirements already trace to the KNOWN HLD §4 design blueprint (the same design_blueprint
        # this CRD read); the forward LLD/MOP/NRFU cells are honest placeholders naming the deliverable
        # that will author them. Non-REQ-D rows are placeholders across all four forward columns.
        if rid.startswith("REQ-D"):
            return (rid, "HLD §4 (design blueprint)", "to be traced in LLD",
                    "to be traced in MOP", "design-driven NRFU (to be traced in NRFU)")
        return (rid, "to be traced in HLD", "to be traced in LLD",
                "to be traced in MOP", "to be traced in NRFU")
    table(["REQ-ID", "HLD section", "LLD object", "MOP step", "NRFU test case"],
          [_rtm_row(rid) for rid in req_ids], widths=[1.0, 1.35, 1.4, 1.4, 1.55])

    # ===== 10. Design-driven requirements (target-state blueprint) — evidence-gated =====
    if bp_decisions:
        doc.add_heading("10. Design-Driven Requirements (Target-State Blueprint)", level=1)
        doc.add_paragraph(
            "Requirements the assessment itself derives: the CCDE-grounded target-state design blueprint "
            "(the SAME design_blueprint behind the HLD/LLD §4 and the explorer Design mode), read as "
            "requirement candidates. Each is gated on observed evidence and cites a network-design "
            "principle — confirm, amend or strike it like any workshop requirement; each traces forward to "
            "an HLD §4 design decision (and appears in the §9 RTM).")
        if rec:
            table(["REQ-D", "Design requirement (recommended target pattern)", "Driver / evidence",
                   "CCDE basis", "Confirmed?"],
                  [(rid,
                    f"{d.get('title')}: {d.get('recommended_action') or ''}".strip().rstrip(":"),
                    # per-decision children coerced: a decision row is isinstance-filtered, but its
                    # evidence / principle sub-object can still be a truthy scalar -> 5.get(...) -> 500
                    _as_dict(d.get("evidence")).get("summary") or d.get("driver") or "—",
                    _as_dict(d.get("principle")).get("citation") or "—", "<YES/AMEND>")
                   for rid, d in req_d],
                  widths=[0.8, 2.7, 2.0, 1.4, 0.8])
        needs = [d for d in bp_decisions if d.get("status") == "needs-requirement"]
        if needs:
            doc.add_paragraph(
                "Open design questions — these decisions depend on requirements the assessment cannot "
                "observe; answer them in the workshop and the target design right-sizes (the engine never "
                "assumes an answer):")
            for d in needs:
                # str()-coerce each element too: `", ".join([7])` raises TypeError even when the list
                # itself is well-formed (the same wrong-typed-leaf class the device-string reads guard).
                _label_run(doc.add_paragraph(), f"{d.get('title')} —",
                           ", ".join(str(r) for r in _as_list(d.get("requirements_needed")))
                           or "requirement to confirm", GREY)
        cov = _as_dict(bp.get("coverage"))
        if cov.get("caveat"):
            _label_run(doc.add_paragraph(), "Coverage:", cov.get("caveat"), GREY)

    # ---- closing acceptance gate ----
    # Deliverable Excellence (DE-01): shared glossary before the sign-off gate.
    add_glossary(doc)

    add_acceptance(
        doc, scope_note="Acceptance of this CRD baselines the requirements for the high-level design; "
                        "changes after acceptance go through change control as requirement amendments.")

    n_req = len(req_ids)
    doc.save(output_path)
    logger.info(f"[Phase 35] CRD (DOCX) written: {output_path} ({n_req} requirement rows seeded)")
