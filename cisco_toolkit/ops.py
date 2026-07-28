"""Operations Handbook — the PPDIOO Operate-phase deliverable (.docx).

NEW-V3.23.168: the document family covered Prepare→Plan→Design→Implement (CRD, design,
MOP, cutover, NRFU) and the close-out (PIR), but nothing for the OPERATE phase — the
handbook the customer's NOC actually runs the network with after the migration. This
generator turns the assessment's own evidence into that handbook: the monitoring
baseline comes from what the syslog axis actually saw fire, the capacity baseline from
the platform-health sample, the drift-control standard from the golden-config baseline,
and the patch-governance section from the software-risk screening.

Evidence discipline: every axis-backed section is GATED on its snapshot key — an older
snapshot (or a collection without that evidence) gets an honest one-line declaration and
a re-collection prompt, never an invented baseline. Placeholders mark everything only the
customer can decide (contacts, escalation tiers, maintenance windows).

python-docx is OPTIONAL: imported inside the function so the package imports without it;
a missing library is a warning + skip, never a crash. Deterministic; no network.
"""
import logging
from datetime import datetime

from cisco_toolkit.docmeta import add_acceptance, add_document_control, add_excellence_front, add_glossary, add_inputs_required, add_table, add_toc
from cisco_toolkit.docmeta import as_dict as _as_dict
from cisco_toolkit.docmeta import as_list as _as_list
from cisco_toolkit.textutils import _as_num, xml_safe, xml_safe_deep   # entry deep-sanitize of device text (audit-5) + fail-soft numeric coercion

logger = logging.getLogger(__name__)


def _facts(snap: dict) -> dict:
    """Defensive reads of the snapshot sections the handbook is grounded in."""
    devices = _as_dict(snap.get("devices"))
    fi = _as_list(snap.get("failure_impact"))
    keystones = [r for r in fi if isinstance(r, dict)]

    def _stranded(r):
        return _as_num(r.get("stranded"))
    keystones.sort(key=lambda r: (-_stranded(r), str(r.get("host") or "")))
    keystones = [k for k in keystones if _stranded(k) > 0]
    si = _as_dict(snap.get("syslog_intelligence"))
    ph = _as_dict(snap.get("platform_health"))
    gd = _as_dict(snap.get("golden_drift"))
    qa = _as_dict(snap.get("qos_audit"))
    sr = _as_dict(snap.get("software_risk"))
    lc = _as_dict(snap.get("lifecycle_risk"))
    sec = _as_dict(snap.get("security"))
    sec_fail = 0
    for host, blk in sec.items():
        s = _as_dict(_as_dict(blk).get("summary"))
        sec_fail += _as_num(s.get("fail"))   # fail-soft: a bad 'fail' count contributes 0, never aborts the handbook
    # routing-adjacency + first-hop-redundancy Day-2 health (N36)
    ph2 = _as_list(snap.get("protocol_health"))
    routing_protos = sorted({str(r.get("protocol", "")).upper() for r in ph2 if isinstance(r, dict)
                             and str(r.get("protocol", "")).upper() in ("OSPF", "BGP", "EIGRP", "ISIS", "IS-IS")})
    n_proto_high = sum(1 for r in ph2 if isinstance(r, dict) and r.get("severity") in ("High", "Critical"))
    l3f = _as_list(snap.get("l3_forwarding"))
    n_gw = len(l3f)
    # Real FHRP only, and the fallback must run BEFORE str(): `str(None)` is the truthy string
    # "None", which never equals "none", so a present-but-null fhrp counted as FHRP-PROTECTED and
    # 3.3 rendered "N of N gateway SVI(s) FHRP-protected" off gateways with no first-hop redundancy
    # at all. Mirrors the engine's canonical gate `(fhrp or "none") != "none"` (analyze.py) exactly
    # as crd.py:68 does — this was the outlier.
    n_fhrp = sum(1 for r in l3f if isinstance(r, dict) and (r.get("fhrp", "none") or "none") != "none")
    return {"devices": devices, "keystones": keystones[:5], "si": si, "ph": ph,
            "gd": gd, "qa": qa, "sr": sr, "lc": lc, "sec": sec, "n_sec_fail": sec_fail,
            "routing_protos": routing_protos, "n_proto_high": n_proto_high, "n_gw": n_gw, "n_fhrp": n_fhrp}


def _hosts_from(rows, key="host", cap=6):
    """De-duplicated, ordered list of affected device names from a list of finding dicts, capped
    (with an ellipsis when truncated) so a known-issue cites WHICH devices, not just a count."""
    seen = []
    for r in rows:
        if not isinstance(r, dict):
            continue
        h = r.get(key)
        if h and h not in seen:
            seen.append(str(h))
    txt = ", ".join(seen[:cap])
    if len(seen) > cap:
        txt += f", +{len(seen) - cap} more"
    return txt or "—"


def _known_issues(ev: dict) -> tuple:
    """Synthesize the operate-phase Known-Issues register from the assessment's OWN finding axes.

    Returns (issues, absent) where `issues` is a list of (axis, issue, affected, day2-note) rows —
    each grounded in a collected axis and citing the source axis + affected devices — and `absent`
    is a list of (axis, how-to-collect) rows for finding-bearing axes that were NOT collected. An
    uncollected axis is DECLARED not-assessable (in `absent`), never silently dropped — a silent drop
    would read as 'no known issues on that axis', the false-health class this whole toolkit guards
    against. Pure/deterministic; no docx dependency (unit-testable in isolation)."""
    issues, absent = [], []

    # -- syslog_intelligence: recurring signatures already seen on this fleet --
    si_sum = _as_dict(ev["si"].get("summary"))
    if si_sum.get("n_collected"):
        dets = [d for d in _as_list(ev["si"].get("detections")) if isinstance(d, dict)]
        if dets:
            # Aggregate by distinct signature CLASS (label) BEFORE the top-N cut — otherwise one
            # signature firing on many devices produces many rows and crowds every other class off
            # a real fleet's list (adversarial-review finding, 2026-07-05).
            _SEV = {"Critical": 3, "High": 2, "Medium": 1, "Low": 0}
            by_label: dict = {}
            for d in dets:
                a = by_label.setdefault(str(d.get("label") or "?"), {"count": 0, "sev": "—", "devs": 0})
                a["count"] += _as_num(d.get("count"))   # fail-soft: a malformed detection count contributes 0
                a["devs"] += 1
                sev = str(d.get("severity") or "—")
                if _SEV.get(sev, -1) > _SEV.get(a["sev"], -1):
                    a["sev"] = sev
            top = sorted(by_label.items(), key=lambda kv: (-kv[1]["count"], kv[0]))[:6]
            sigs = "; ".join(f"{lbl} (×{a['count']} on {a['devs']} device(s), {a['sev']})" for lbl, a in top)
            issues.append((
                "Syslog Intelligence",
                f"Recurring log signatures already fired on this fleet: {sigs}.",
                _hosts_from(dets),
                "These are live day-2 caveats, not historical noise — alert on each class (§3.1) and "
                "root-cause before it recurs post-cutover."))
        else:
            issues.append((
                "Syslog Intelligence", "Log evidence collected; no recurring-signature detections in the "
                "assessed window.", "(fleet)",
                "Absence is bounded by the buffer window — keep central syslog and the §3.1 alert list live."))
    elif si_sum.get("n_devices"):
        absent.append(("Syslog Intelligence (recurring signatures: MAC-flap / err-disable / link-flap)",
                       "re-run the collection ('show logging' is in the standard set) to derive the "
                       "recurring-signature caveats from this fleet's own history."))
    else:
        absent.append(("Syslog Intelligence (recurring signatures: MAC-flap / err-disable / link-flap)",
                       "re-run the assessment with a current engine to populate the syslog axis."))

    # -- software_risk: open advisory / PSIRT surface --
    sr_sum = _as_dict(ev["sr"].get("summary"))
    if sr_sum.get("n_devices"):
        srf = [f for f in _as_list(ev["sr"].get("findings")) if isinstance(f, dict)]
        n_find = sr_sum.get("n_findings", 0)
        if srf:
            surfaces = "; ".join(sorted({str(f.get("label") or f.get("kind") or "?") for f in srf}))
            issues.append((
                "Software Risk",
                f"{n_find} exposed advisory / hardening surface(s) open at assessment: {surfaces}.",
                _hosts_from(srf),
                "Each is a standing exposure until closed — validate every running release against the "
                "Cisco PSIRT Software Checker (§5) and remediate the surface."))
        else:
            issues.append((
                "Software Risk", f"{n_find} exposed advisory surface(s) flagged; see the Software Risk sheet.",
                "(see sheet)", "Screen every release against the Cisco PSIRT Software Checker on the §8 cadence."))
    else:
        absent.append(("Software Risk (open PSIRT / advisory surface)",
                       "re-run the assessment with a current engine to screen the running releases."))

    # -- platform_health: hot control planes --
    ph_sum = _as_dict(ev["ph"].get("summary"))
    if ph_sum.get("n_collected"):
        phf = [f for f in _as_list(ev["ph"].get("findings")) if isinstance(f, dict)]
        if phf:
            labels = "; ".join(sorted({str(f.get("label") or f.get("kind") or "?") for f in phf}))
            issues.append((
                "Platform Health",
                f"{len(phf)} control-plane capacity finding(s): {labels}.",
                _hosts_from(phf),
                "A hot control plane converges slowly and drops adjacencies under load — treat as a day-2 "
                "caveat on any change to the affected device (§3.2) and trend it."))
        else:
            issues.append((
                "Platform Health", "Capacity sampled; no device outside its band at collection time.",
                "(fleet)", "Single-point sample — re-sample on the §8 cadence and alert on departure from baseline."))
    else:
        absent.append(("Platform Health (hot control planes / CPU-memory pressure)",
                       "re-run the collection to capture the control-plane capacity baseline."))

    # -- lifecycle_risk: past-LDoS / EoS platforms as a standing day-2 caveat --
    lc_sum = _as_dict(ev["lc"].get("summary"))
    if lc_sum:
        n_ldos = lc_sum.get("n_past_ldos")
        n_eos = lc_sum.get("n_past_eos")
        if n_ldos or n_eos:
            issues.append((
                "Lifecycle Risk",
                f"{n_ldos or 0} device(s) past last-date-of-support / {n_eos or 0} past end-of-sale — "
                "no vendor bug/PSIRT remediation path until migrated.",
                "(see Lifecycle Risk sheet)",
                "Standing caveat until the fleet is migrated: an unfixable defect on a past-LDoS platform is "
                "a when-not-if incident. Feed into budget and the migration wave order."))
        else:
            issues.append((
                "Lifecycle Risk", "No past-LDoS / past-EoS platform flagged at assessment.",
                "(fleet)", "Re-check EoX bands quarterly (§8); lifecycle status drifts as Cisco publishes notices."))
    else:
        absent.append(("Lifecycle Risk (past-LDoS / past-EoS platforms)",
                       "re-run the assessment with a current engine to populate the hardware-lifecycle bands."))

    # -- qos_audit: doctrine gaps (a best-effort media fabric is a caveat) --
    qa_sum = _as_dict(ev["qa"].get("summary"))
    if qa_sum.get("n_assessable"):
        qaf = [f for f in _as_list(ev["qa"].get("findings")) if isinstance(f, dict)]
        if qaf:
            labels = "; ".join(sorted({str(f.get("label") or f.get("kind") or "?") for f in qaf}))
            issues.append((
                "QoS Audit",
                f"{len(qaf)} QoS-doctrine gap(s): {labels}.",
                _hosts_from(qaf),
                "On a broadcast/media estate an un-marked, un-queued path is a caveat for real-time feeds — "
                "hold the role's trust + queuing template (§4) on every switch carrying media."))
        else:      # collected AND clean — represent it, don't drop it (a silent drop reads as 'not assessed')
            issues.append((
                "QoS Audit", "QoS posture screened; no marking/queuing doctrine gap flagged at assessment.",
                "(fleet)", "Hold the role's trust + queuing template (§4) as QoS is extended for real-time media."))
    else:
        absent.append(("QoS Audit (marking / queuing doctrine gaps)",
                       "re-run the assessment to screen the fleet's QoS posture."))

    # -- security: CIS hardening failures open at assessment --
    # Coverage-honest, mirroring every other axis (adversarial-review finding, 2026-07-05): declare
    # the axis not-assessable when it was never collected; when collected-and-clean, say so; and when
    # there are failures, name ONLY the devices that actually failed (never every device that merely
    # carries a security block — that would assert open failures on clean boxes to a change board).
    if ev["sec"]:                                          # the CIS/hardening axis WAS collected
        def _fail_of(blk):                                 # total on a malformed 'fail' (e.g. 'many' / a JSON Infinity)
            return _as_num(_as_dict(_as_dict(blk).get("summary")).get("fail"))
        fail_hosts = sorted(h for h, blk in ev["sec"].items() if _fail_of(blk) > 0)
        if ev["n_sec_fail"] and fail_hosts:
            issues.append((
                "Security Posture",
                f"{ev['n_sec_fail']} CIS hardening check failure(s) open at assessment.",
                _hosts_from([{"host": h} for h in fail_hosts]),
                "Each is an operations-owned remediation with a named owner and date — track to zero (§4)."))
        else:
            issues.append((
                "Security Posture", "CIS / hardening posture screened; no failing check at assessment.",
                "(fleet)", "Re-screen on the §8 cadence — hardening posture drifts as configs change."))
    else:
        absent.append(("Security Posture (CIS / hardening failures)",
                       "re-run the assessment to screen the fleet's hardening posture (per-device CIS checks)."))

    return issues, absent


def write_ops_handbook_docx(output_path: str, snap_dict: dict, label: str) -> None:
    """Emit the Operations Handbook (.docx) to `output_path`. Fail-soft: a missing python-docx is a
    warning + skip, and every snapshot read is defensive so malformed input degrades to an honest
    absence declaration. A genuine render error still propagates to the caller — the CLI phase
    wrapper logs-and-continues, AssessHub turns it into a clean 500. (Docstring corrected: the old
    'any unexpected render error is logged, never raised' claim was false — the only `try` here is
    the optional-import guard below — and it hid a crash class now guarded at the read sites. The
    fix is the guards, NOT a broad `except`: swallowing a render error would ship a silently short
    document, which the coverage-honesty doctrine forbids.)"""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor
    except ImportError:
        logger.warning("  Operations handbook (DOCX) skipped: python-docx not installed "
                       "(pip install python-docx to enable the Operate-phase deliverable).")
        return

    snap = xml_safe_deep(snap_dict if isinstance(snap_dict, dict) else {})   # one bad device byte (noncharacter/surrogate) must not abort the docx
    label = xml_safe(label) if isinstance(label, str) else (str(label) if label is not None else "")
    from cisco_toolkit.brand_tokens import DOC_NAVY_RGB
    NAVY = RGBColor(*DOC_NAVY_RGB)
    GREY = RGBColor(0x59, 0x59, 0x59)
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    def _label_run(p, label_text, value, color=NAVY):
        r = p.add_run(label_text); r.bold = True; r.font.color.rgb = color
        p.add_run(" " + (str(value) if value not in (None, "") else "—"))

    def table(headers, rows, widths=None):
        return add_table(doc, headers, rows, widths, fixed=False)

    def absent(what, how):
        p = doc.add_paragraph()
        _label_run(p, "Not in this snapshot:", f"{what} — {how}", GREY)

    ev = _facts(snap)
    devices = ev["devices"]

    # ---- title page ----
    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Operations Handbook"); tr.bold = True
    tr.font.size = Pt(26); tr.font.color.rgb = NAVY
    sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s2 = sub2.add_run("Operate-phase handbook — monitoring baseline, drift control, "
                      "software governance and escalation readiness")
    s2.font.size = Pt(14); s2.font.color.rgb = GREY
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr2 = sub.add_run(label); sr2.font.size = Pt(13); sr2.font.color.rgb = GREY
    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _scale = _as_dict(_as_dict(snap.get("executive_brief")).get("scale"))   # coerce: a truthy non-dict eb/scale must not crash (audit-4 #10)
    _ncoll = _scale.get("n_collected")
    _ninv = _scale.get("n_devices")                                         # canonical inventoried (SSOT owner: executive_brief.scale.n_devices); len() only pre-brief (CROSS-03)
    meta.add_run(f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')}  ·  "
                 f"{_ncoll if isinstance(_ncoll, int) else len(devices)} collected of "
                 f"{_ninv if isinstance(_ninv, int) else len(devices)} inventoried  ·  script {snap.get('script_version', '')}"
                 ).font.color.rgb = GREY
    doc.add_paragraph()
    note = doc.add_paragraph()
    _label_run(note, "How to use this handbook:",
               "This is the Day-2 companion to the migration documents: the monitoring thresholds, "
               "baselines and cadences below are derived from this fleet's own assessed evidence, "
               "not generic defaults. Adopt it as the NOC's working document — complete the "
               "<placeholder> contact and window fields, then re-generate from a fresh assessment "
               "whenever the network changes materially (the baselines move with the evidence).", GREY)
    doc.add_page_break()

    # ---- document control ----
    add_document_control(
        doc, document="Operations Handbook", label=label,
        engine_version=str(snap.get("script_version", "")), generated_at=snap.get("generated_at"),
        collected_at=snap.get("collected_at"),
        audience="The customer's NOC / operations team (primary users), the network owner "
                 "(standards and escalation authority) and the delivery engineer handing over.",
        exclude=("opshandbook",),
        extra_assumptions=(
            "Baselines in this handbook are derived from the named assessment snapshot; they are "
            "the network AS ASSESSED, and drift with every change. Re-generate after material "
            "changes rather than hand-editing the numbers.",))
    doc.add_page_break()

    # ---- table of contents (shared field-code helper, V3.23.171) ----
    add_toc(doc)


    # Deliverable Excellence (DE-01): answer-first at-a-glance register + single-source-of-truth signal.
    add_excellence_front(doc, snap_dict)
    # Deliverable Excellence (DE-01): consolidated Inputs-Required register (Law 9).
    add_inputs_required(doc, snap_dict)
    # ===== 1. Purpose & audience =====
    doc.add_heading("1. Purpose & Audience", level=1)
    doc.add_paragraph(
        "The document family covers planning (CRD), design (HLD/LLD), execution (MOP, cutover, "
        "NRFU) and close-out (PIR); this handbook covers what comes after — running the network. "
        "It gives the operations team four things grounded in this fleet's own evidence: a "
        "monitoring baseline (what this network's logs and control planes actually look like), a "
        "drift-control standard, a software-governance cadence, and an escalation pack. It is a "
        "living document owned by operations after handover.")

    # ===== 2. Network quick reference =====
    doc.add_heading("2. Network Quick Reference", level=1)
    if devices:
        rows = [[h, _as_dict(devices[h]).get("model") or "—",
                 _as_dict(devices[h]).get("platform") or "—",
                 _as_dict(devices[h]).get("sw_version") or "—",
                 _as_dict(devices[h]).get("serial_number") or "—"]
                for h in sorted(devices)]
        table(["Device", "Model", "Platform", "Software", "Serial"],
              rows[:60], widths=[1.6, 1.7, 0.8, 1.4, 1.5])
        if len(rows) > 60:
            doc.add_paragraph(f"… and {len(rows) - 60} more — full inventory in the workbook.")
    else:
        absent("device inventory", "re-run the collection to populate the quick reference.")
    if ev["keystones"]:
        doc.add_heading("2.1 Keystone devices (handle with care)", level=2)
        doc.add_paragraph(
            "Failure-impact analysis ranks these as the fleet's keystones — the devices whose loss "
            "strands the most endpoints. Treat any work on them as high-risk change control.")
        table(["Device", "Severity", "Endpoints stranded if it fails", "VLANs impacted"],
              [[k.get("host"), k.get("severity", "—"), k.get("stranded", 0),
                k.get("vlans_impacted", 0)]
               for k in ev["keystones"]], widths=[2.0, 1.2, 2.0, 1.6])

    # ===== 3. Monitoring & alerting baseline =====
    doc.add_heading("3. Monitoring & Alerting Baseline", level=1)
    doc.add_heading("3.1 Log signatures to alert on (from this fleet's own logs)", level=2)
    si_sum = _as_dict(ev["si"].get("summary"))
    if si_sum.get("n_devices"):
        if si_sum.get("n_collected"):
            doc.add_paragraph(
                f"The assessment parsed {si_sum.get('total_events', 0)} buffered log event(s) on "
                f"{si_sum.get('n_collected', 0)} device(s) and raised "
                f"{si_sum.get('n_detections', 0)} detection(s). Alert on these signature classes "
                "FIRST — they have already occurred on this network:")
            det_rows = [[d.get("host"), d.get("label"), d.get("severity"), d.get("count", 0)]
                        for d in _as_list(ev["si"].get("detections"))[:12]
                        if isinstance(d, dict)]
            if det_rows:
                table(["Device", "Signature (already seen here)", "Severity", "Count"],
                      det_rows, widths=[1.5, 3.0, 0.9, 0.7])
            else:
                doc.add_paragraph("No detections in the assessed window — start from the standard "
                                  "set: MAC flaps, err-disable, link flaps, environmental alarms, "
                                  "duplex mismatches, login failures.")
            doc.add_paragraph(
                "Forward syslog to a central collector (the CIS check in the workbook flags devices "
                "without 'logging host') and page on severities 0–2; review severity 3 daily.")
        else:
            absent("log evidence ('show logging')",
                   "re-run the collection (the command is in the standard set) to derive the alert "
                   "baseline from this fleet's own history.")
    else:
        absent("the syslog-intelligence axis", "re-run the assessment with a current engine.")

    doc.add_heading("3.2 Control-plane capacity baseline", level=2)
    ph_sum = _as_dict(ev["ph"].get("summary"))
    if ph_sum.get("n_collected"):
        pb = _as_dict(ph_sum.get("bands"))   # `or {}` caught only a FALSY bands; a truthy scalar crashed .items()
        btxt = ", ".join(f"{v}× {k}" for k, v in pb.items())
        doc.add_paragraph(
            f"Capacity sample at collection time ({btxt}). These figures are the NORMAL for this "
            "fleet — alert when a device departs its own baseline, not just on absolute thresholds. "
            "The sample is a single point in time: re-sample on the §8 cadence.")
        prow = [[d.get("host"), d.get("cpu_5min") if d.get("cpu_5min") is not None else "—",
                 d.get("mem_free_pct") if d.get("mem_free_pct") is not None else "—",
                 d.get("band")]
                for d in _as_list(ev["ph"].get("per_device"))
                if isinstance(d, dict) and d.get("collected")][:30]
        if prow:
            table(["Device", "CPU 5-min % (baseline)", "Memory free % (baseline)", "Band"],
                  prow, widths=[1.9, 1.7, 1.7, 1.2])
    else:
        absent("CPU/memory capacity output",
               "re-run the collection to capture the control-plane baseline.")

    # ---- 3.3 routing-adjacency & first-hop-redundancy health (Day-2 monitored items, N36) ----
    doc.add_heading("3.3 Routing-adjacency & first-hop-redundancy health", level=2)
    doc.add_paragraph(
        "Beyond CIS / syslog / capacity, the operate phase continuously watches the control plane's "
        "routing adjacencies and gateway redundancy — a silent adjacency drop or a lost FHRP peer is an "
        "outage waiting for the next failure.")
    r36 = []
    if ev["routing_protos"]:
        r36.append(("Routing adjacencies", ", ".join(ev["routing_protos"]),
                    "Baseline the neighbour list per device; alert on any adjacency leaving Full/Up and on "
                    "flap counts."))
    else:
        r36.append(("Routing adjacencies", "none observed (L2-forwarded fleet)",
                    "If routing is introduced in the target design, add adjacency-state + flap monitoring."))
    if ev["n_gw"]:
        if ev["n_fhrp"]:
            r36.append(("First-hop redundancy", f"{ev['n_fhrp']} of {ev['n_gw']} gateway SVI(s) FHRP-protected",
                        "Alert on HSRP/VRRP/GLBP active↔standby transitions and track-decrement events."))
        else:
            r36.append(("First-hop redundancy", f"0 of {ev['n_gw']} gateway SVI(s) — none observed",
                        "No FHRP state to monitor today; introducing gateway redundancy is a design "
                        "recommendation — then alarm on active/standby transitions."))
    if ev["n_proto_high"]:
        r36.append(("Protocol-health findings", f"{ev['n_proto_high']} flagged High at assessment",
                    "Re-verify post-cutover and monitor for recurrence (see the punch-list)."))
    table(["Monitored item", "Observed", "Day-2 alerting"], r36, widths=[1.5, 2.2, 2.8])

    # ===== 4. Operational standards & drift control =====
    doc.add_heading("4. Operational Standards & Drift Control", level=1)
    gd_sum = _as_dict(ev["gd"].get("summary"))
    if gd_sum.get("n_baseline"):
        gmode = gd_sum.get("mode") or "majority"
        gsrc = ("the supplied golden-config standard" if gmode == "golden-file"
                else "the fleet's de-facto majority baseline (auto-derived)")
        doc.add_paragraph(
            f"Configuration standard: {gsrc}, {gd_sum.get('n_baseline', 0)} required directive(s). "
            f"At assessment time {gd_sum.get('n_drifting', 0)} of {gd_sum.get('n_devices', 0)} "
            f"device(s) drifted (average compliance {gd_sum.get('avg_compliance_pct', 0)}%). "
            "Operations owns keeping this at zero: re-run the drift check on the §8 cadence and "
            "treat every new MISSING directive as an unauthorized change until explained.")
    else:
        absent("a configuration baseline",
               "supply --golden-config (or collect ≥3 full running-configs for the majority "
               "baseline) and re-run.")
    qa_sum = _as_dict(ev["qa"].get("summary"))
    if qa_sum.get("n_assessable"):
        qmodes = _as_dict(qa_sum.get("modes"))
        doc.add_paragraph(
            f"QoS posture to hold: {', '.join(f'{v}× {k}' for k, v in qmodes.items())} "
            f"({qa_sum.get('n_findings', 0)} finding(s) open in the QoS Audit sheet). Any new "
            "switch joining the fleet must match its role's trust + queuing template before "
            "carrying traffic.")
    if ev["n_sec_fail"]:
        doc.add_paragraph(
            f"Security hardening: {ev['n_sec_fail']} CIS check failure(s) were open at assessment "
            "time (Security Posture sheet). Each is an operations-owned remediation with a named "
            "owner and date — track them to zero.")

    # ===== 5. Software & lifecycle governance =====
    doc.add_heading("5. Software & Lifecycle Governance", level=1)
    sr_sum = _as_dict(ev["sr"].get("summary"))
    if sr_sum.get("n_devices"):
        tb = _as_dict(sr_sum.get("train_bands"))
        doc.add_paragraph(
            f"Software trains at assessment: {', '.join(f'{v}× {k}' for k, v in tb.items())}. "
            f"{sr_sum.get('n_findings', 0)} exposed advisory surface(s) were open (Software Risk "
            "sheet). Governance: validate every running release with the Cisco PSIRT Software "
            "Checker on the §8 cadence, close the exposed surfaces, and plan upgrades for every "
            "Replace/Upgrade and Verify-EoL train against Cisco's published notices.")
        worst = [d for d in _as_list(ev["sr"].get("per_device"))
                 if isinstance(d, dict)
                 and d.get("train_band") in ("Replace/Upgrade", "Verify EoL")][:12]
        if worst:
            table(["Device", "Version", "Train", "Band"],
                  [[d.get("host"), d.get("sw_version"), d.get("train"), d.get("train_band")]
                   for d in worst], widths=[1.7, 1.5, 1.8, 1.6])
    else:
        absent("the software-risk screening", "re-run the assessment with a current engine.")
    lc_sum = _as_dict(ev["lc"].get("summary"))
    if lc_sum:
        doc.add_paragraph(
            "Hardware lifecycle: the Lifecycle Risk sheet carries the per-device EoX bands — "
            "review quarterly and feed Past-EoS / Near-LDoS devices into budget planning.")

    # ===== 6. Backup & recovery =====
    # Config-backup strategy + cadence, the restore procedure, and — load-bearing — the restore-TEST
    # discipline. Evidence-gated: the assessment collects read-only SHOW output, so it does NOT directly
    # evidence backup state; that is DECLARED as a blind spot (coverage-honesty), never asserted healthy.
    doc.add_heading("6. Backup & Recovery", level=1)
    doc.add_paragraph(
        "A config that cannot be restored is an outage waiting for its trigger. Operations owns a "
        "backup regime that is proven by restore, not assumed by the presence of a backup file.")
    # count reconciled to the canonical scale (SSOT: executive_brief.scale.n_devices), len() only pre-brief
    _n_scope = _ninv if isinstance(_ninv, int) else len(devices)
    doc.add_heading("6.1 Strategy, cadence & retention", level=2)
    table(["Control", "Standard to hold", "Evidence / owner"], [
        ("What is backed up",
         "The full running- AND startup-config of every managed device, plus any off-box state "
         "(licences, certificates, NDFC/controller policy where applicable).",
         f"All {_n_scope} device(s) in scope (§2)"),
        ("Cadence",
         "Automated nightly config backup, and an explicit on-change backup captured immediately "
         "before and after every MOP (the pre-change capture is the rollback baseline).",
         "Scheduler / §8 calendar"),
        ("Storage & retention",
         "Version-controlled off-box store, access-controlled, ≥90 days of daily versions plus a "
         "monthly archive; the store must survive loss of the device it backs up (off-site copy).",
         "Customer backup platform"),
        ("Integrity",
         "Each backup diffed against the prior version so an empty or truncated capture is caught the "
         "morning after, not at restore time.",
         "Backup platform / §4 drift check"),
    ], widths=[1.5, 3.6, 1.6])
    doc.add_heading("6.2 Restore procedure & the restore-test discipline", level=2)
    doc.add_paragraph(
        "A backup that has never been restore-tested is not a backup — it is an untested assumption. "
        "The restore procedure and its proving discipline:")
    table(["Step", "Action"], [
        ("1. Retrieve", "Pull the last-known-good version for the device from the off-box store "
                        "(identify it by the pre-change backup, not by date alone)."),
        ("2. Stage & review", "Diff the backup against current running-config; a peer reviews the delta "
                             "before any load — a restore is a change and follows change control."),
        ("3. Restore", "Apply in a maintenance window with console/OOB access proven FIRST (never rely on "
                       "the path you are restoring); capture a post-restore backup."),
        ("4. Validate", "Run the NRFU/post-change checks (§8, MOP/NRFU documents) to confirm the device "
                        "returned to its baseline before handing back."),
        ("Restore-test drill", "On the §8 cadence, restore a representative device's backup to a lab or "
                              "spare and prove it boots and forwards. An untested restore path is a "
                              "latent outage — this drill is what turns a backup file into a recovery "
                              "capability."),
    ], widths=[1.6, 5.1])
    doc.add_heading("6.3 Backup coverage (evidence-gated)", level=2)
    doc.add_paragraph(
        "Coverage-honesty: this assessment collects read-only SHOW output and does not directly "
        "evidence whether each device is under a working backup regime — that is a blind spot the "
        "customer closes, not a state this document can assert. What the evidence CAN anchor:")
    _bk_rows = []
    gd_sum_bk = _as_dict(ev["gd"].get("summary"))
    if gd_sum_bk.get("n_baseline"):
        _bk_rows.append(("Config baseline captured",
                         f"A configuration baseline of {gd_sum_bk.get('n_baseline', 0)} directive(s) exists "
                         f"across {gd_sum_bk.get('n_devices', 0)} device(s) — the drift check (§4) is the "
                         "day-2 signal that a config changed outside change control."))
    else:
        _bk_rows.append(("Config baseline captured",
                         "No configuration baseline was captured — supply --golden-config or collect "
                         "≥3 full running-configs so drift (an unexpected change to a backed-up config) "
                         "is detectable. Until then, backup completeness is [NOT OBSERVED]."))
    _bk_rows.append(("Per-device backup status",
                     f"Not directly assessed for any of the {_n_scope} in-scope device(s) — verify in the "
                     "customer's backup platform that every device has a recent, restore-tested backup; "
                     "any device missing one is a blind spot, not assumed backed-up."))
    table(["What the evidence anchors", "Status (coverage-honest)"], _bk_rows, widths=[2.2, 4.5])

    # NDFC / EVPN-fabric backup discipline — gated on the SAME design_blueprint.evpn_migration.applicable
    # flag the MOP §2.2 uses, so it is silent on a non-EVPN engagement (evidence-gated, never assumed).
    _evpn = _as_dict(_as_dict(snap.get("design_blueprint")).get("evpn_migration"))
    if _evpn.get("applicable"):
        # Match the CRD/MOP confidence (adversarial-review finding, 2026-07-05): assert the EVPN/NDFC target
        # only when the requirements register confirms it; otherwise it is an engine assessment to confirm.
        _rm = _as_dict(_as_dict(snap.get("design_blueprint")).get("requirements_model"))
        _fom = next((str(f.get("value") or "") for f in _as_list(_rm.get("fields"))
                     if isinstance(f, dict) and f.get("key") == "fabric_operating_model"), "")
        _confirmed = bool(_rm.get("provided")) and "evpn" in _fom.lower()
        _lead = (f"The target is an NX-OS VXLAN BGP-EVPN fabric ({_evpn.get('model_basis', '')}), "
                 "customer-confirmed via the requirements register."
                 if _confirmed else
                 f"IF the target is an NX-OS VXLAN BGP-EVPN fabric ({_evpn.get('model_basis', '')}) — an engine "
                 "assessment to confirm with the customer, not a settled plan-of-record — the day-2 backup model "
                 "changes as follows.")
        doc.add_heading("6.4 Fabric backup discipline (NDFC-managed target)", level=2)
        doc.add_paragraph(
            _lead + " On the "
            "NDFC-managed fabric the backup model changes: NDFC (Nexus Dashboard Fabric Controller) owns "
            "scheduled and on-demand configuration backups for every managed switch, and the source of "
            "truth is NDFC's intent, not each device. Two disciplines follow:")
        for _b in (
            "NDFC config-backup is enabled per fabric (scheduled + pre/post-change), and an NDFC "
            "backup/restore is exercised on the §8 restore-test cadence like any other device — the same "
            "'never restore-tested = not a backup' rule applies to the controller.",
            "No out-of-band CLI once managed: once a switch is under NDFC, configuration changes go "
            "through NDFC so intent and device stay reconciled. An out-of-band CLI edit creates drift NDFC "
            "will flag (or overwrite) — treat any OOB change as an incident, matching the §4 drift standard."):
            doc.add_paragraph(_b, style="List Bullet")

    # ===== 7. Known issues & operating caveats =====
    # Synthesized from the assessment's OWN finding axes (syslog / software-risk / platform-health /
    # lifecycle / qos / security). Each collected axis cites its source + affected devices; an axis that
    # was NOT collected is DECLARED not-assessable — a silent drop would read as "no known issues" on that
    # axis, the false-health class the toolkit exists to prevent.
    doc.add_heading("7. Known Issues & Operating Caveats", level=1)
    doc.add_paragraph(
        "The day-2 caveats the network hands over WITH — synthesized from this assessment's own findings, "
        "not a generic list. Each issue names the assessment axis it comes from and the affected "
        "device(s); an axis that was not collected is declared not-assessable below, never silently "
        "treated as clean.")
    _issues, _absent_axes = _known_issues(ev)
    if _issues:
        table(["Source axis", "Known issue (from this fleet's evidence)", "Affected", "Day-2 caveat / action"],
              [[a, i, aff, note] for (a, i, aff, note) in _issues], widths=[1.2, 2.7, 1.2, 1.6])
    else:
        doc.add_paragraph(
            "No finding-bearing axis was collected, so NO known issue can be asserted either way — see "
            "the not-assessable axes below. This is a coverage gap, not an all-clear.")
    if _absent_axes:
        doc.add_heading("7.1 Not-assessable axes (declared, not assumed clean)", level=2)
        doc.add_paragraph(
            "These finding axes were NOT collected. Per coverage-honesty each is a blind spot, NOT a "
            "'no issues' result — collect the evidence before treating the axis as clean:")
        for _axis, _how in _absent_axes:
            p = doc.add_paragraph(style="List Bullet")
            _label_run(p, f"{_axis}:", f"not-assessable — {_how}", GREY)

    # ===== 8. Routine operations calendar =====
    doc.add_heading("8. Routine Operations Calendar", level=1)
    doc.add_paragraph(
        "The cadence that keeps the baselines above honest. Each row names its evidence source so "
        "the activity stays mechanical, not aspirational.")
    table(["Cadence", "Activity", "Evidence / tool"], [
        ("Daily", "Review severity 0–3 syslog events against the §3.1 alert list; clear or ticket "
                  "every new detection-class event.", "Central syslog / §3.1"),
        ("Weekly", "Sample control-plane CPU/memory and compare to the §3.2 baseline; investigate "
                   "any device that departed its own normal.", "show processes cpu / §3.2"),
        ("Monthly", "Re-run the configuration drift check; every new missing directive is an "
                    "unauthorized change until explained.", "Toolkit --no-collect re-run / §4"),
        ("Monthly", "Validate running releases against new advisories.",
         "Cisco PSIRT Software Checker / §5"),
        ("Quarterly", "Review hardware EoX bands and software-train lifecycle; feed into budget "
                      "and upgrade planning.", "Lifecycle Risk + Software Risk sheets"),
        ("Quarterly", "Re-run the full assessment and trend it against prior snapshots; the "
                      "baselines in this handbook move with the evidence.",
         "Toolkit (--trend OLD NEW)"),
        ("Quarterly", "Restore-test drill: restore a representative device's backup to a lab/spare and "
                      "prove it boots and forwards — an untested restore is not a recovery capability.",
         "Backup platform / §6.2"),
        ("Per change", "MOP with per-step success criteria + post-change validation; capture a pre- and "
                       "post-change backup and update the drift baseline when the standard itself changes.",
         "MOP / NRFU documents / §6"),
    ], widths=[1.0, 4.0, 1.9])

    # ===== 9. Escalation & TAC readiness =====
    doc.add_heading("9. Escalation & TAC Readiness", level=1)
    doc.add_paragraph(
        "Before any vendor case: capture the evidence pack below — it answers the first round of "
        "TAC questions in one attachment and anchors the case to facts.")
    table(["Item", "Where it comes from"], [
        ("Device model, serial and running release", "§2 quick reference (verify against the device)"),
        ("The failing device's full command capture", "Re-run the toolkit collection for that device "
                                                      "(read-only show commands)"),
        ("'show tech-support' from the affected device", "Collected at incident time (large; collect "
                                                         "to file)"),
        ("Recent log excerpt around the event", "Central syslog / 'show logging'"),
        ("Relevant advisory IDs already screened", "Software Risk sheet (§5)"),
        ("The assessment workbook + this handbook", "Engagement deliverables"),
    ], widths=[3.0, 3.9])
    doc.add_paragraph(
        "Escalation tiers, contacts and response targets are the customer's to define:")
    table(["Tier", "Condition", "Contact", "Response target"], [
        ("P1", "<service-down condition>", "<name / desk>", "<minutes>"),
        ("P2", "<degraded / redundancy-lost condition>", "<name / desk>", "<hours>"),
        ("P3", "<single-user / non-urgent condition>", "<queue>", "<business days>"),
    ], widths=[0.6, 2.7, 1.9, 1.7])

    # ---- acceptance ----
    # Deliverable Excellence (DE-01): shared glossary before the sign-off gate.
    add_glossary(doc)

    add_acceptance(
        doc, scope_note="Acceptance of this handbook transfers Day-2 operational ownership to the "
                        "customer's operations team; the baselines inside it are superseded by each "
                        "newer accepted re-generation.")

    doc.save(output_path)
    logger.info(f"[Phase 38] Operations handbook (DOCX) written: {output_path} "
                f"({_ninv if isinstance(_ninv, int) else len(devices)} devices in scope)")   # canonical n_devices (reuse title's read); len() only pre-brief (CROSS-03)
