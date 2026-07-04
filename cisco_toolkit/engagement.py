"""Engagement Workflow & Plan of Record — the engagement-management instrument (.docx).

NEW-V3.23.157: the document chain (CRD → HLD/LLD → MOP → cutover plan → NRFU → PIR) says WHAT each
phase produces; nothing said HOW THE ENGAGEMENT RUNS — the judgment layer a senior AS engineer adds
between the documents. This deliverable is that layer: an evidence-led engagement verdict (proceed /
proceed-with-conditions / hold), a phase tracker with entry/exit gates across the whole document set,
the ordered next-action queue, a T-minus gate calendar per migration wave (commit → checkpoint →
go/no-go → window → hypercare exit, the cadence large-migration governance playbooks converge on),
and a seeded RAID log (Risks / Assumptions / Issues / Dependencies + a decision log) with the
assessment's own findings as the opening entries.

Evidence discipline: every verdict, gate criterion and RAID row is traced to a computed block of the
snapshot (punch-list, migration readiness, wave sequencing, scenarios, validation plan, collection
completeness, lifecycle risk, executive brief). The toolkit never invents project state — documents
it generated are DRAFTS until signed, unconfirmed assumptions stay marked to-confirm, and a fleet
with no derivable waves gets the unscheduled fallback, not a fabricated schedule.

python-docx is OPTIONAL: imported inside the function so the package imports without it; a missing
library is a warning + skip, never a crash. Every snapshot read is defensive. Deterministic; no network.
"""
import logging
import re
from datetime import datetime

from cisco_toolkit.docmeta import SEV_RANK as _SEV_RANK
from cisco_toolkit.docmeta import add_acceptance, add_document_control, add_excellence_front, add_glossary, add_inputs_required, add_table, add_toc
from cisco_toolkit.docmeta import as_dict as _docmeta_as_dict
from cisco_toolkit.docmeta import as_list as _docmeta_as_list
from cisco_toolkit.textutils import xml_safe, xml_safe_deep   # entry deep-sanitize of device text (audit-5)

logger = logging.getLogger(__name__)

# The per-wave T-minus gate cadence — ONE source of truth: §4.1 renders from it, AssessHub's gate
# board offers exactly these gates (webapp/backend/gates.py imports it), and the recorded
# dispositions land back in §4.3 keyed by these keys. (key, label, when, purpose, go_criteria).
# CONTRACT (V3.23.159): the keys are persisted into AssessHub's `gates.gate` column — they are a
# storage schema. Keys are APPEND-ONLY: never rename or remove one without a webapp data migration,
# or existing sign-offs orphan (invisible on the board, rendered with a raw key in §4.3).
GATE_SEQUENCE = (
    ("commit", "Commit", "T-28 days", "Lock scope, owners and window request",
     "Wave scope frozen; change request submitted; rollback owner named"),
    ("checkpoint", "Checkpoint", "T-14 days", "Verify preparation is on track",
     "Blockers for this wave closed; MOP peer-reviewed; spares staged"),
    ("readiness", "Readiness review", "T-7 days", "Confirm the window can be used",
     "Change approved; comms sent; baseline snapshot captured (--compare ready)"),
    ("go_no_go", "Go / No-Go", "T-1 day", "The decision meeting",
     "All readiness checks pass; on-call roster confirmed; rollback rehearsed"),
    ("window", "Window", "T-0", "Execute the MOP",
     "Run-of-show followed; every validation step at or above its expected baseline"),
    ("hypercare_exit", "Hypercare exit", "T+5 days", "Stand down elevated support",
     "No Critical/High incident attributable to the wave; NRFU signed"),
)


# V3.23.171: the truthy-non-dict coercers are shared family furniture now (docmeta.as_dict /
# as_list) -- local aliases keep the call sites unchanged.
_as_dict = _docmeta_as_dict
_as_list = _docmeta_as_list


def _as_int(v, default: int = 0) -> int:
    """Coerce snapshot counts defensively: None / '' / 'unknown' all degrade to `default`."""
    try:
        return int(v)
    except (TypeError, ValueError):
        return default


def _missing_text(d: dict) -> str:
    """Render a blind-spot row's `missing` commands: a list joins normally, a bare string is one
    item (never char-joined), anything else degrades to str() — no input shape can crash or
    garble the issues table."""
    m = d.get("missing")
    items = [str(x) for x in m] if isinstance(m, list) else ([str(m)] if m else [])
    return ", ".join(items)[:80]


def _workflow_facts(snap: dict) -> dict:
    """Pull the evidence the workflow synthesizes — all defensive reads of known shapes."""
    punch = sorted((i for i in _as_list(snap.get("punchlist")) if isinstance(i, dict)),
                   key=lambda i: _SEV_RANK.get(i.get("severity"), 5))
    sev_count = {}
    for i in punch:
        s = i.get("severity") or "Info"
        sev_count[s] = sev_count.get(s, 0) + 1
    mr = [r for r in _as_list(snap.get("migration_readiness")) if isinstance(r, dict)]
    ws_by = {w.get("group"): w for w in _as_list(snap.get("wave_sequencing")) if isinstance(w, dict)}
    sc = _as_dict(snap.get("migration_scenarios"))
    sc_by = {g.get("group"): g for g in _as_list(sc.get("per_group")) if isinstance(g, dict)}
    by_wave = _as_dict(_as_dict(snap.get("validation_plan")).get("by_wave"))
    coll = _as_dict(snap.get("collection_completeness"))
    csum = _as_dict(coll.get("summary"))
    # Tri-state honesty: an ABSENT completeness block means completeness is UNKNOWN — it must
    # never render as "collection complete" (absence of evidence is not evidence of completeness).
    coll_present = bool(csum)
    blind = [d for d in _as_list(coll.get("devices")) if isinstance(d, dict)]
    brief = _as_dict(snap.get("executive_brief"))
    lc = _as_dict(_as_dict(snap.get("lifecycle_risk")).get("summary"))
    rem = _as_dict(_as_dict(snap.get("remediation_plan")).get("summary"))
    notready = [r.get("group") or "?" for r in mr if r.get("readiness") == "NOT READY"]

    # Pilot-first doctrine: the recommended pilot is the READY group with the SMALLEST blast radius
    # (fewest endpoints); a CAUTION group only if nothing is READY; never a NOT READY group.
    # A group with UNKNOWN endpoints (null / non-numeric) sorts LAST — missing data must never
    # win the pilot slot by masquerading as zero blast radius.
    def _pilot_key(r):
        ep = _as_int(r.get("endpoints"), default=-1)
        return (ep < 0, ep if ep >= 0 else 0, str(r.get("group") or ""))

    pilot = None
    for want in ("READY", "CAUTION"):
        cands = sorted((r for r in mr if r.get("readiness") == want), key=_pilot_key)
        if cands:
            pilot = cands[0]
            break

    return {"punch": punch, "sev_count": sev_count, "mr": mr, "ws_by": ws_by,
            "sc": sc, "sc_by": sc_by, "by_wave": by_wave, "csum": csum,
            "coll_present": coll_present, "blind": blind,
            "brief": brief, "lc": lc, "rem": rem, "notready": notready, "pilot": pilot,
            # SSOT: read the canonical inventory count first (the same source n_collected reads), with
            # the local len(devices) only as a pre-brief fallback — else "N collected of M inventoried"
            # mixes a canonical n_collected with a recount and can render "253 collected of 3 inventoried".
            "n_devices": (((_as_dict(snap.get("executive_brief")).get("scale")) or {}).get("n_devices")
                          or len(_as_dict(snap.get("devices")))),
            "n_collected": ((_as_dict(snap.get("executive_brief")).get("scale")) or {}).get("n_collected")}


def _verdict(f: dict) -> tuple:
    """The senior engineer's call: (verdict, [conditions]) — every condition cites its evidence.
    An unconditional PROCEED returns an EMPTY conditions list (no placeholder dressed as a
    condition); the writer renders the no-conditions statement itself."""
    conds = []
    n_crit = f["sev_count"].get("Critical", 0)
    n_high = f["sev_count"].get("High", 0)
    n_blind = _as_int(f["csum"].get("partial")) + _as_int(f["csum"].get("not_collected"))
    if n_crit:
        conds.append(f"{n_crit} Critical punch-list item(s) must be closed before any wave is scheduled.")
    if f["notready"]:
        conds.append("Readiness gating fails for " + ", ".join(str(g) for g in f["notready"][:4])
                     + " — resolve the failing checks, then re-assess.")
    if not f["coll_present"]:
        conds.append("Collection-completeness evidence is absent from this snapshot — completeness "
                     "is UNKNOWN; re-assess with a current engine collection before relying on the "
                     "verdicts.")
    elif n_blind:
        conds.append(f"{n_blind} device(s) have partial or missing collection — every verdict on them "
                     "is provisional until re-collection closes the blind spots.")
    if n_high:
        conds.append(f"{n_high} High punch-list item(s) should be remediated or explicitly "
                     "risk-accepted before the go/no-go gate.")
    if _as_int(f["lc"].get("n_past_ldos")):
        conds.append(f"{f['lc']['n_past_ldos']} device(s) are past last-day-of-support (LDoS) — no "
                     "TAC escalation path during the windows; stage spares or replace first.")
    if _as_int(f["lc"].get("n_past_eos")):
        conds.append(f"{f['lc']['n_past_eos']} device(s) are past end-of-sale (EoS), support window "
                     "closing — plan a refresh before they reach last-day-of-support.")
    if n_crit or f["notready"]:
        return "HOLD", conds
    if conds:
        return "PROCEED WITH CONDITIONS", conds
    return "PROCEED", []


def write_engagement_docx(output_path: str, snap_dict: dict, label: str,
                          gate_record: dict | None = None) -> None:
    """Emit the Engagement Workflow & Plan of Record (.docx) to `output_path`. Fail-soft: a missing
    python-docx is a warning + skip, and every snapshot read is defensive so malformed input degrades
    to placeholders. A genuine render error still propagates to the caller — the CLI phase wrapper
    logs-and-continues, AssessHub turns it into a clean 500 (V3.23.159: docstring corrected; the old
    'never raised' claim was false and the crash class it hid is now guarded at the read sites).

    `gate_record` (optional) is engagement state fed back from AssessHub's gate board:
    {wave: {gate_key: {"decision", "signed_by", "note", "decided_at"}}} keyed by GATE_SEQUENCE keys.
    When present, §4.3 renders the as-signed trail; the CLI passes nothing (no project state offline)."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor
    except ImportError:
        logger.warning("  Engagement workflow (DOCX) skipped: python-docx not installed "
                       "(pip install python-docx to enable the plan-of-record deliverable).")
        return

    snap = xml_safe_deep(snap_dict if isinstance(snap_dict, dict) else {})   # one bad device byte (noncharacter/surrogate) must not abort the docx
    label = xml_safe(label) if isinstance(label, str) else (str(label) if label is not None else "")
    NAVY = RGBColor(0x1F, 0x38, 0x64)
    GREY = RGBColor(0x59, 0x59, 0x59)
    RED = RGBColor(0xB0, 0x2A, 0x1E)
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    def _label_run(p, label_text, value, color=NAVY):
        r = p.add_run(label_text); r.bold = True; r.font.color.rgb = color
        p.add_run(" " + (str(value) if value not in (None, "") else "—"))

    def table(headers, rows, widths=None):
        # Delegates to the family's single table builder (docmeta.add_table).
        return add_table(doc, headers, rows, widths, fixed=False)

    f = _workflow_facts(snap)
    verdict, conditions = _verdict(f)

    # ---- title page ----
    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Engagement Workflow & Plan of Record"); tr.bold = True
    tr.font.size = Pt(26); tr.font.color.rgb = NAVY
    sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s2 = sub2.add_run("Phase gates, next actions, wave calendar and RAID log — synthesized from the assessment evidence")
    s2.font.size = Pt(14); s2.font.color.rgb = GREY
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(label); sr.font.size = Pt(13); sr.font.color.rgb = GREY
    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')}  ·  "
                 f"{f.get('n_collected') if f.get('n_collected') is not None else f['n_devices']} collected of "
                 f"{f['n_devices']} inventoried  ·  script {snap.get('script_version', '')}"
                 ).font.color.rgb = GREY
    status = doc.add_paragraph(); status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st = status.add_run("DRAFT PLAN OF RECORD — the verdict and every gate below are evidence-led "
                        "proposals; the engagement adopts them by signing the acceptance block.")
    st.italic = True; st.font.color.rgb = RED
    doc.add_paragraph()
    note = doc.add_paragraph()
    _label_run(note, "How to use this document:",
               "This is the working plan of record for the engagement — review it at the weekly "
               "cadence, keep the RAID log current (a RAID log nobody reviews is documentation "
               "theater, not management), and re-generate it from a fresh snapshot after each wave "
               "so the verdict tracks reality. Each phase ends at a signed gate; no further work "
               "until the gate's exit criteria are met.", GREY)
    doc.add_page_break()

    # ---- document control (AS-style front matter) ----
    add_document_control(
        doc, document="Engagement Workflow & Plan of Record", label=label,
        engine_version=str(snap.get("script_version", "")), generated_at=snap.get("generated_at"),
        collected_at=snap.get("collected_at"),
        audience="The delivery engineer running the engagement, the customer's project sponsor and "
                 "network owner (gate signatories), and the project / change manager who owns the "
                 "calendar and the RAID cadence.",
        exclude=("engagement",),
        extra_assumptions=(
            "Project state (gate sign-offs, window approvals, RAID dispositions) lives with the "
            "engagement, not the snapshot — generated documents are DRAFTS until signed, and the "
            "phase tracker describes each phase's pending gate, never generation or approval "
            "status it cannot evidence.",))
    doc.add_page_break()

    # ---- table of contents (shared field-code helper, V3.23.171) ----
    add_toc(doc)


    # Deliverable Excellence (DE-01): answer-first at-a-glance register + single-source-of-truth signal.
    add_excellence_front(doc, snap_dict)
    # Deliverable Excellence (DE-01): consolidated Inputs-Required register (Law 9).
    add_inputs_required(doc, snap_dict)
    # ===== 1. Engagement verdict =====
    doc.add_heading("1. Engagement Verdict", level=1)
    vp_ = doc.add_paragraph()
    vr = vp_.add_run(verdict); vr.bold = True; vr.font.size = Pt(16)
    vr.font.color.rgb = RED if verdict == "HOLD" else NAVY
    brief = f["brief"]
    if brief.get("posture_statement"):
        doc.add_paragraph(str(brief["posture_statement"]))
    if conditions:
        doc.add_paragraph("Conditions attached to this verdict (each cites its evidence):")
        for c in conditions:
            doc.add_paragraph(c, style="List Bullet")
    else:
        doc.add_paragraph("No conditions attached — the evidence shows no gating findings; "
                          "schedule the pilot wave.")
    gating = _as_list(brief.get("top_gating"))
    if gating:
        doc.add_paragraph("Cross-axis gating headlines (from the executive brief — full detail in "
                          "the workbook and runbook):")
        for g in gating[:6]:
            doc.add_paragraph(str(g), style="List Bullet")

    # ===== 2. Engagement phase tracker =====
    doc.add_heading("2. Engagement Phase Tracker", level=1)
    doc.add_paragraph(
        "One row per engagement phase, mapped to the instrument that records it and the gate that "
        "closes it. Status reflects only what this snapshot evidences — whether a document was "
        "generated, reviewed or signed lives with the engagement (and, when recorded in AssessHub, "
        "in the gate record at §4.3); the gate column says what approval looks like.")
    # Assess status is tri-state: complete / has blind spots / completeness UNKNOWN (no block).
    if not f["coll_present"]:
        assess_status = "Evidence collected; completeness UNKNOWN (no collection-completeness block)"
    elif f["blind"]:
        assess_status = f"Evidence collected; {len(f['blind'])} device(s) with blind spots"
    else:
        assess_status = "Evidence collected (this snapshot)"
    has_waves = bool(f["mr"])
    table(["Phase", "Instrument", "Entry criteria", "Exit gate", "Status"], [
        ("Assess", "Assessment workbook · explorer · runbook",
         "Collection captured for the inventory",
         "Findings reviewed with the network owner", assess_status),
        ("Plan", "Customer Requirements Document (CRD)",
         "Assessment reviewed", "Requirements workshop held; CRD signed",
         "Pending — requirements workshop not yet held"),
        ("Design", "As-Built Network Design Document (HLD/LLD)",
         "CRD signed", "Design review held; document accepted",
         "Pending — design review not yet held"),
        ("Implement", "MOP (per wave) · Cutover Plan (run-of-show)",
         "Design accepted; blockers closed", "Per-wave go/no-go passed; window executed",
         ("Waves derived — gate calendar pending" if has_waves
          else "No waves derivable from this evidence yet")),
        ("Validate", "NRFU / Acceptance Test Plan",
         "Wave executed", "NRFU results at or above baseline; acceptance signed",
         "Runs after each wave"),
        ("Optimize", "Post-Implementation Review (PIR)",
         "All waves executed; hypercare complete", "PIR held; lessons recorded; project closed",
         "Not started — post-cutover"),
    ], widths=[0.9, 1.9, 1.4, 1.6, 1.5])

    # ===== 3. Next actions =====
    doc.add_heading("3. Next Actions", level=1)
    doc.add_paragraph(
        "The ordered work queue — what a senior engineer would do next, each action traced to the "
        "evidence that demands it and to the gate it unblocks.")
    actions = []
    if f["blind"]:
        hosts = ", ".join(str(d.get("host")) for d in f["blind"][:6]) \
                + (" …" if len(f["blind"]) > 6 else "")
        actions.append(("Close collection blind spots: re-collect " + hosts,
                        f"collection completeness: {f['csum'].get('partial', 0)} partial / "
                        f"{f['csum'].get('not_collected', 0)} not collected",
                        "Delivery engineer", "Assess"))
    elif not f["coll_present"]:
        actions.append(("Re-collect with a current engine — completeness evidence is absent from "
                        "this snapshot, so collection state is UNKNOWN",
                        "collection completeness (block missing)", "Delivery engineer", "Assess"))
    actions.append(("Hold the requirements workshop; confirm, amend or strike every seeded CRD row",
                    "CRD draft (evidence-primed proposals)", "Customer + delivery engineer", "Plan"))
    n_blockers = f["sev_count"].get("Critical", 0) + f["sev_count"].get("High", 0)
    if n_blockers:
        rem_n = f["rem"].get("n_items") or ""
        actions.append((f"Remediate the {n_blockers} Critical/High punch-list item(s)"
                        + (f" — {rem_n} generated config snippet(s) await engineering review"
                           if rem_n else ""),
                        "punch-list + remediation plan", "Customer ops + delivery engineer",
                        "Implement"))
    actions.append(("Hold the design review; accept or annotate the HLD/LLD",
                    "design document draft", "Customer network owner", "Design"))
    if f["pilot"]:
        p = f["pilot"]
        # endpoints can be present-but-null/non-numeric in uploaded snapshots: say UNKNOWN, never
        # render 'None endpoint(s)' (and _pilot_key already keeps unknown groups out of the slot
        # unless nothing better exists).
        ep_txt = (f"{_as_int(p.get('endpoints'))} endpoint(s)"
                  if _as_int(p.get("endpoints"), -1) >= 0 else "endpoint count unknown")
        actions.append((f"Schedule the pilot wave: {p.get('group')} "
                        f"({len(_as_list(p.get('switches')))} switch(es), "
                        f"{ep_txt}) — smallest blast radius first",
                        f"migration readiness: {p.get('readiness')}", "Project manager", "Implement"))
    elif has_waves:
        actions.append(("Re-assess wave readiness after blocker remediation; no group currently "
                        "qualifies as a pilot", "migration readiness (all gated)",
                        "Delivery engineer", "Implement"))
    else:
        actions.append(("Derive migration waves once collection covers the fleet — none are "
                        "derivable from the current evidence", "move groups (empty)",
                        "Delivery engineer", "Implement"))
    actions.append(("Execute remaining waves per their recommended scenario; NRFU after each",
                    "migration scenarios + validation plan", "All parties", "Validate"))
    actions.append(("Run the PIR and hypercare exit review; close the project",
                    "PIR (post-cutover)", "Project manager", "Optimize"))
    table(["#", "Action", "Evidence basis", "Owner (role)", "Unblocks"],
          [(i, a, e, o, g) for i, (a, e, o, g) in enumerate(actions, 1)],
          widths=[0.4, 2.9, 1.6, 1.3, 0.8])

    # ===== 4. Wave gate calendar =====
    doc.add_heading("4. Wave Gate Calendar", level=1)
    doc.add_paragraph(
        "Migration governance runs each wave through a T-minus gate cadence; a gate that cannot "
        "answer its criteria does not pass — the wave slips, the calendar does not bend.")
    doc.add_heading("4.1 Gate cadence (per wave)", level=2)
    table(["Gate", "When", "Purpose", "Go criteria"],
          [(g_label, when, purpose, criteria)
           for _key, g_label, when, purpose, criteria in GATE_SEQUENCE],
          widths=[1.1, 0.8, 1.7, 3.1])

    doc.add_heading("4.2 Wave schedule (from the evidence)", level=2)
    if has_waves:
        rows = []
        for r in f["mr"]:
            g = r.get("group") or "?"
            w = f["ws_by"].get(g, {})
            s = f["sc_by"].get(g, {})
            hard = len(_as_list(w.get("hard_cutover")))
            pilot_mark = " (PILOT)" if f["pilot"] and f["pilot"].get("group") == g else ""
            rows.append((str(g) + pilot_mark, len(_as_list(r.get("switches"))),
                         _as_int(r.get("endpoints")), r.get("readiness") or "—",
                         # the engine emits 'recommended_scenario' (every other consumer reads it); 'scenario' is a
                         # key it never emits, so the §4.2 Scenario column was '—' for every wave (audit-4 #9)
                         s.get("recommended_scenario") or "—",
                         (f"{hard} switch(es) / {_as_int(w.get('hard_cutover_endpoints'))} endpoint(s)"
                          if hard else "none"),
                         len(_as_list(f["by_wave"].get(g)))))
        table(["Wave", "Switches", "Endpoints", "Readiness", "Scenario",
               "Hard-cutover exposure", "Post-checks"], rows,
              widths=[1.1, 0.7, 0.8, 0.9, 0.9, 1.6, 0.7])
        if f["sc"].get("fleet_recommendation"):
            p2 = doc.add_paragraph()
            _label_run(p2, "Fleet recommendation:", f["sc"]["fleet_recommendation"])
        doc.add_paragraph(
            "Order of execution: pilot first, then ascending blast radius. A NOT READY wave is not "
            "scheduled — it re-enters the calendar when its gating checks pass. Per-wave procedure "
            "is the MOP; per-wave validation is the cutover-validation checklist.")
    else:
        doc.add_paragraph(
            "No migration waves are derivable from the current evidence (no multi-switch move "
            "groups). The calendar starts when the Assess phase closes its blind spots or the "
            "design defines the move boundaries explicitly.")

    # §4.3 — the feedback loop: gate dispositions recorded against the campaign (AssessHub gate
    # board) land back in the plan of record as the as-signed trail. Only rendered when something
    # was actually recorded — an empty record is project state the document must not invent.
    gr = {str(w): g for w, g in _as_dict(gate_record).items() if isinstance(g, dict) and g}
    if gr:
        gate_label = {k: lbl for k, lbl, *_rest in GATE_SEQUENCE}
        gate_order = {k: i for i, (k, *_r) in enumerate(GATE_SEQUENCE)}
        mr_order = {str(r.get("group")): i for i, r in enumerate(f["mr"])}

        def _wave_key(w):
            # Calendar order (the §4.2 row order) first; anything else falls back to NATURAL
            # numeric order — 'Group 2' before 'Group 10', never lexicographic.
            m = re.search(r"(\d+)", w)
            return (mr_order.get(w, len(mr_order)), int(m.group(1)) if m else 10 ** 9, w)

        def _stamp(ts) -> str:
            # storage._now() emits ISO-8601 UTC ('…T14:23:01+00:00'); render it labeled, never a
            # bare wall-clock time that silently dropped its offset.
            s = str(ts or "")
            if not s:
                return "—"
            base = s[:16].replace("T", " ")
            if s.endswith("+00:00") or s.endswith("Z"):
                return base + " UTC"
            if len(s) > 19 and s[19] in "+-":
                return base + " UTC" + s[19:]
            return base

        def _gate_rows(waves):
            out = []
            for wave in sorted(waves, key=_wave_key):
                for key in sorted(gr[wave], key=lambda k: gate_order.get(k, 99)):
                    d = gr[wave][key] if isinstance(gr[wave][key], dict) else {}
                    out.append((wave, gate_label.get(key, key),
                                str(d.get("decision") or "").upper().replace("_", "-") or "—",
                                d.get("signed_by") or "—", _stamp(d.get("decided_at")),
                                d.get("note") or "—"))
            return out

        doc.add_heading("4.3 Gate record (as signed)", level=2)
        doc.add_paragraph(
            "Dispositions recorded against this campaign in AssessHub — the as-run trail of the "
            "calendar above (times are UTC). A NO-GO or SLIPPED row keeps its history here; the "
            "calendar shows where the wave re-enters.")
        # Gate records are CAMPAIGN state; this document is a SNAPSHOT view. Rows whose wave label
        # exists in this snapshot's calendar render as the trail; rows signed against a different
        # collection's waves are kept (never silently dropped) but clearly fenced off.
        known = [w for w in gr if w in mr_order] if has_waves else list(gr)
        stray = [w for w in gr if w not in mr_order] if has_waves else []
        if known:
            table(["Wave", "Gate", "Decision", "Signed by", "When", "Note"], _gate_rows(known),
                  widths=[0.9, 1.1, 0.9, 1.1, 1.1, 1.6])
        if stray:
            doc.add_paragraph(
                "Recorded against waves NOT in this snapshot's calendar — signed against a "
                "different collection of this campaign; reconcile before relying on them:")
            table(["Wave", "Gate", "Decision", "Signed by", "When", "Note"], _gate_rows(stray),
                  widths=[0.9, 1.1, 0.9, 1.1, 1.1, 1.6])

    # ===== 5. RAID log (seeded) =====
    doc.add_heading("5. RAID Log (seeded from the assessment)", level=1)
    doc.add_paragraph(
        "Risks, Assumptions, Issues, Dependencies — plus the decision log. Seeded rows are the "
        "assessment's own findings; the engagement keeps this log current at the weekly review. "
        "Every row needs a named owner — an unowned RAID row is a wish, not a control.")

    doc.add_heading("5.1 Risks", level=2)
    risks = []
    for i in f["punch"][:8]:
        if i.get("severity") in ("Critical", "High"):
            risks.append((i.get("severity"), i.get("title") or "—",
                          "punch-list: " + (i.get("category") or "—"),
                          "Remediate before the owning wave's T-14 checkpoint"))
    if _as_int(f["lc"].get("n_past_ldos")):
        risks.append(("High", f"{f['lc']['n_past_ldos']} device(s) past last-day-of-support (LDoS) "
                              "in the migration path", "lifecycle risk",
                      "Stage spares / replace before their wave; no TAC path otherwise"))
    hard_total = sum(_as_int(w.get("hard_cutover_endpoints")) for w in f["ws_by"].values())
    if hard_total:
        risks.append(("Medium", f"{hard_total} endpoint(s) take a hard outage during cutover "
                                "(single-homed switches)", "wave sequencing",
                      "Schedule those switches inside agreed maintenance windows; dual-home first "
                      "where feasible"))
    if not risks:
        risks.append(("—", "No Critical/High findings recorded by the assessment", "punch-list",
                      "Keep the log open — new risks land here as the engagement runs"))
    table(["Severity", "Risk", "Evidence", "Mitigation  (owner: <name>)"],
          [("RSK-%03d · %s" % (i, r[0]), r[1], r[2], r[3]) for i, r in enumerate(risks, 1)],
          widths=[1.1, 2.4, 1.2, 2.3])

    doc.add_heading("5.2 Assumptions (confirm in the workshop)", level=2)
    table(["ID", "Assumption", "If wrong"], [
        ("ASM-001", "Maintenance windows will be granted on the cadence the calendar needs",
         "The schedule stretches; re-baseline the plan of record"),
        ("ASM-002", "The snapshot reflects the production network (no change since collection)",
         "Re-collect before the readiness review of the next wave"),
        ("ASM-003", "Spares, licences and support contracts cover the target-state hardware",
         "Procurement lead time enters the critical path"),
    ], widths=[0.8, 3.4, 2.5])

    doc.add_heading("5.3 Issues (open now)", level=2)
    if f["blind"]:
        table(["ID", "Issue", "Evidence", "Action"],
              [(f"ISS-{i:03d}",
                f"{d.get('host')}: collection {d.get('status')} ({d.get('data_quality')}% of "
                "essential commands)",
                "collection completeness",
                "Re-collect; missing: " + (_missing_text(d) or "—"))
               for i, d in enumerate(f["blind"][:8], 1)], widths=[0.7, 2.4, 1.3, 2.6])
    elif not f["coll_present"]:
        doc.add_paragraph("Collection-completeness evidence is absent from this snapshot — "
                          "collection state is UNKNOWN. Treat closing that gap as the first open "
                          "issue; new issues land here as the engagement runs.")
    else:
        doc.add_paragraph("No open issues seeded — collection is complete for the inventory. New "
                          "issues land here as the engagement runs.")

    doc.add_heading("5.4 Dependencies", level=2)
    table(["ID", "Dependency", "Depends on"], [
        ("DEP-001", "Wave scheduling", "CRD signed (windows + outage tolerance are customer requirements)"),
        ("DEP-002", "MOP execution", "Design acceptance + blocker remediation for the wave's switches"),
        ("DEP-003", "Bulk waves", "Pilot wave NRFU signed (pilot-first: lessons feed the bulk MOPs)"),
        ("DEP-004", "Project closure", "PIR held and hypercare exit criteria met on the final wave"),
    ], widths=[0.8, 2.0, 3.9])

    doc.add_heading("5.5 Decision log", level=2)
    # DE-01/Law 7: a decision record must show the ALTERNATIVES considered + the rationale, not just a
    # position — otherwise the reader cannot tell whether the choice was reasoned or defaulted.
    decisions = [("DEC-001", "Fleet cutover scenario",
                  str(f["sc"].get("fleet_recommendation") or "<to be decided at design review>"),
                  "Big-bang (all at once) vs phased-by-domain vs pilot-first; recommended for the lowest "
                  "blast radius and reversibility — per-group basis in the Migration Scenarios sheet.",
                  "PROPOSED")]
    if f["pilot"]:
        decisions.append(("DEC-002", "Pilot wave selection",
                          f"{f['pilot'].get('group')} — smallest qualifying blast radius",
                          "Other candidate groups carried larger blast radius / readiness gaps; the smallest "
                          "qualifying group de-risks the bulk waves that inherit its lessons.",
                          "PROPOSED"))
    decisions.append(("DEC-%03d" % (len(decisions) + 1), "Window calendar",
                      "<dates per wave — set at the commit gate>",
                      "Driven by the customer's approved maintenance windows + outage tolerance (a CRD input).",
                      "OPEN"))
    table(["ID", "Decision", "Position", "Alternatives / rationale", "Status"], decisions,
          widths=[0.6, 1.3, 1.9, 2.5, 0.7])

    # ===== 6. Operating rhythm =====
    doc.add_heading("6. Operating Rhythm", level=1)
    doc.add_paragraph(
        "The cadence that keeps the plan of record honest. Re-generate this document from a fresh "
        "snapshot after every wave — the verdict, the wave table and the RAID seeds then track the "
        "network as it actually is, not as it was at kickoff.")
    table(["Ritual", "Cadence", "Owner", "Output"], [
        ("RAID + next-actions review", "Weekly", "Project / change manager",
         "Updated log; aged rows escalated"),
        ("Gate reviews", "Per the T-minus calendar", "Delivery engineer",
         "Signed gate record (or a slip decision)"),
        ("Snapshot refresh + re-assessment", "After each wave; before each readiness review",
         "Delivery engineer", "Fresh deliverable set; --compare diff vs baseline"),
        ("Sponsor checkpoint", "Per wave", "Project sponsor", "Verdict acknowledged; conditions accepted"),
    ], widths=[2.0, 1.6, 1.5, 1.6])

    # ---- closing acceptance gate ----
    # Deliverable Excellence (DE-01): shared glossary before the sign-off gate.
    add_glossary(doc)

    add_acceptance(
        doc, scope_note="Acceptance of this document adopts the verdict, the gate calendar and the "
                        "seeded RAID log as the engagement's working plan of record; subsequent "
                        "changes are recorded in the decision log.")

    doc.save(output_path)
    logger.info(f"[Phase 36] Engagement workflow (DOCX) written: {output_path} "
                f"(verdict: {verdict}; {len(f['mr'])} wave(s) on the calendar)")
