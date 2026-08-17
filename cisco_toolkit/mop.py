"""Method of Procedure (MOP) — per-wave migration cutover runbook (.docx).

NEW-V3.23.149: the Implement-phase deliverable. A Cisco Advanced Services migration ends each wave in
a maintenance window driven by a MOP: scope, pre-cutover baseline capture, the ordered procedure,
post-cutover validation, and rollback. This module emits one MOP document with a section per migration
wave (move group), assembled from the SAME snapshot the other deliverables read so it stays consistent
with the workbook / explorer / runbook / design doc.

It is a TEMPLATE primed with the assessment's evidence — never an auto-executed change. The procedure
steps carry <placeholder> markers for window-specific facts (date, owner, approver), and the cutover
strategy (make-before-break vs hard cutover) and the rollback come from the per-group migration
scenario the engine already recommended. The post-cutover checks are the EXISTING validation plan
(one source of truth: the same per-wave checks the workbook's Cutover Validation sheet carries).

python-docx is OPTIONAL: imported inside the function so the package imports without it; a missing
library is a warning + skip, never a crash. Every snapshot read is defensive. Deterministic; no network.
"""
import logging
from datetime import datetime

from cisco_toolkit.docmeta import add_acceptance, add_document_control, add_excellence_front, add_glossary, add_inputs_required, add_table, add_toc
from cisco_toolkit.docmeta import as_dict as _as_dict, as_list as _as_list
from cisco_toolkit.textutils import _as_num, xml_safe, xml_safe_deep   # entry deep-sanitize of device text (audit-5) + fail-soft numeric coercion

logger = logging.getLogger(__name__)

_SEV_RANK = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3, "Info": 4}


def _as_strs(x) -> list:
    """Coerce a snapshot list the document renders as TEXT into a list of strings.

    `as_list` guards the CONTAINER's type but is structurally blind to what is inside it, so a
    perfectly well-formed list holding one truthy scalar (`["distA", 5]`, the shape an attacker-supplied
    snapshot can carry through the webapp's `isinstance(snap, dict) and "devices" in snap` upload check)
    still detonates at the first `", ".join(...)` / `doc.add_paragraph(...)` / `"x" in <str>` of the
    element. This is the ELEMENT-granular sibling of docmeta's as_dict/as_list, and it mirrors the
    `str(dd.get(...) or "")` coercion already applied to wrong-typed device leaves below: identity on
    every string element, so a well-formed snapshot renders byte-identically."""
    return [str(v or "") for v in _as_list(x)]


def _as_dicts(x) -> list:
    """Coerce a snapshot list-of-records to the records only. Same element-granular blind spot as
    `_as_strs`, for the lists whose elements are dereferenced with `.get(...)` rather than rendered."""
    return [v for v in _as_list(x) if isinstance(v, dict)]


def _waves(snap: dict):
    """Resolve the ordered list of migration waves as (name, switches[]) pairs. Prefer the
    migration_readiness verdict rows (they carry the group name + switch set the validation plan keys
    on); fall back to synthesising 'Group N' from move_groups so the MOP still renders on an older or
    minimal snapshot."""
    mr = [_as_dict(r) for r in _as_list(snap.get("migration_readiness"))]
    if mr:
        return [(r.get("group") or f"Group {i + 1}", _as_strs(r.get("switches")))
                for i, r in enumerate(mr)]
    return [(f"Group {i + 1}", _as_strs(_as_dict(g).get("switches")))
            for i, g in enumerate(_as_list(snap.get("move_groups")))]


def _wave_sections(snap: dict):
    """Per-WAVE section descriptors keyed on the canonical design_blueprint.target_state.wave_plan (one
    MOP section per candidate wave). Each wave resolves its per-group records by SWITCH-MEMBERSHIP, never
    by index (wave_plan.source_groups is a filtered/0-based index while the readiness groups are raw/
    1-based -- joining by index would misattribute checks). Falls back to _waves() (one section per
    move-group) when wave_plan is absent, so older/minimal snapshots still render unchanged.
    Returns (name, switches, kind, group_names)."""
    wp = _as_dict(_as_dict(_as_dict(snap.get("design_blueprint")).get("target_state")).get("wave_plan"))
    waves = _as_list(wp.get("waves"))
    if not waves:
        return [(name, switches, "move-group", [name]) for name, switches in _waves(snap)]
    sw2grp = {}
    for r in _as_list(snap.get("migration_readiness")):
        r = _as_dict(r)
        g = r.get("group")
        for s in _as_strs(r.get("switches")):
            sw2grp.setdefault(s, g)
    out = []
    for w in waves:
        w = _as_dict(w)
        switches = _as_strs(w.get("switches"))
        kind = w.get("kind") or "wave"
        gnames = []
        for s in switches:
            g = sw2grp.get(s)
            if g and g not in gnames:
                gnames.append(g)
        out.append((f"Wave {w.get('wave')} ({kind})", switches, kind, gnames))
    return out


def _join_group_records(gnames, wave_switches, readiness_by_group, seq_by_group, scen_by_group, val_by_wave):
    """Union the per-group records across the groups a wave's switches belong to (a coupled-subwave maps to
    one group; an independent-batch maps to many): worst readiness verdict, fail/warn checks, first
    non-empty sequencing/scenario, deduped validation checks. Endpoints are APPORTIONED to this wave's
    switch-share of each group -- a sub-wave is a SLICE, so it must not inherit the parent group's full
    endpoint total. Single source of truth -- reads the engine's records, never recomputes them."""
    # Rank the engine's ACTUAL readiness vocabulary -- analyze.compute_migration_readiness emits exactly
    # {"NOT READY","CAUTION","READY"}. Lower = worse, so the reduce below keeps the worst verdict; an
    # unknown value defaults high (least-worst) and never masks a real NOT READY.
    rank = {"not ready": 0, "caution": 1, "ready": 2}

    def _i(x):
        # Delegate to the shared fail-soft coercer: it rejects the JSON Infinity/NaN that a plain
        # int() would let through (inf -> OverflowError, missed by (TypeError, ValueError)).
        return int(_as_num(x))

    r = {"endpoints": 0, "n_fail": 0, "n_warn": 0, "readiness": None}
    seq, scen, val_items, worst, seen = {}, {}, [], None, set()
    mbb_all, hard_all = [], []   # cutover host split UNIONED across every group in the wave (not just the first)
    members_all = set()          # full membership of the joined groups — detects a sub-wave SLICE below
    wsw = set(wave_switches or [])
    for g in gnames:
        rr = readiness_by_group.get(g) or {}
        # element-coerced: a scalar member would otherwise break `s in note` (TypeError: 'in <string>'
        # requires string) and `sorted(gsw)` (int vs str) in the check-attribution below, and every
        # downstream render of `switches` (the §x.1 join, the port map, the OOB evidence lookup).
        gsw = _as_strs(rr.get("switches"))
        members_all.update(gsw)
        share = (len(wsw & set(gsw)) / len(gsw)) if gsw else 1.0      # a sub-wave is a SLICE -> apportion its switch-share
        r["endpoints"] += round(_i(rr.get("endpoints")) * share)   # endpoints ARE divisible -> apportion by share
        # n_fail/n_warn: ATTRIBUTE each fail/warn check to the ONE sub-wave that owns its canonical (lowest-sorted)
        # member switch, so the per-sub-wave counts SUM to the group SSOT. The old share-with-floor reprinted a
        # 'max(.,1)' on every slice, so an N-way split of a positive count summed toward N (Meridian Group 1's true 3/4
        # rendered 7/7 across 7 sub-waves) -- a one-source-of-truth break vs migration_readiness (audit-2 #3).
        checks = [c for c in _as_list(rr.get("checks"))
                  if isinstance(c, dict) and str(c.get("status", "")).lower() in ("fail", "warn")]
        if checks:
            gsw_set = set(gsw)
            for chk in checks:
                note = str(chk.get("note", ""))
                members = sorted(s for s in gsw_set if s and s in note)   # group switches named in this check's note
                owner = members[0] if members else (sorted(gsw)[0] if gsw else None)   # canonical single owner
                if owner is not None and owner in wsw:                    # this sub-wave owns the check
                    r["n_fail" if str(chk.get("status")).lower() == "fail" else "n_warn"] += 1
        else:
            # older snapshot with no per-check attribution -> apportion by switch-share, never rounding a lone
            # blocker DOWN to 0 (that would mask a NOT-READY wave's gate). Single-slice correctness preserved.
            _ef, _ew = _i(rr.get("n_fail")) * share, _i(rr.get("n_warn")) * share
            r["n_fail"] += max(round(_ef), 1) if _ef > 0 else 0
            r["n_warn"] += max(round(_ew), 1) if _ew > 0 else 0
        rd = rr.get("readiness")
        if rd and (worst is None or rank.get(str(rd).lower(), 9) < rank.get(str(worst).lower(), 9)):
            worst = rd
        gseq = seq_by_group.get(g) or {}
        if not seq:
            seq = gseq
        for _h in _as_strs(gseq.get("make_before_break")):
            if _h not in mbb_all:
                mbb_all.append(_h)
        for _h in _as_strs(gseq.get("hard_cutover")):
            if _h not in hard_all:
                hard_all.append(_h)
        if not scen:
            scen = scen_by_group.get(g) or {}
        # keep only RECORDS: every consumer of val_items dereferences it (`it.get("device")` for the
        # §x.3 capture table, `it.get("severity")` for the §x.6 go/no-go sort), so a scalar element --
        # a truthy int OR a plain string, both of which the old `isinstance` branch let through into
        # the list -- crashed the render two sections later.
        for v in _as_dicts(val_by_wave.get(g)):
            # A command is a capture action, not a validation-check identity.  FHRP deliberately emits
            # one semantic check per VLAN while reusing a single `show standby brief` capture; collapsing
            # on (device, command) silently removed every VLAN after the first from the §x.6 gate.  Keep
            # exact semantic duplicates out when groups are joined, but retain distinct VLAN/subtype/check
            # and acceptance rows.  repr() keeps malformed JSON-like leaves hashable so this join remains
            # total; all downstream rendering still applies the existing defensive coercion.
            key = tuple(repr(v.get(field)) for field in (
                "device", "category", "subtype", "vlan", "check", "command", "expect", "severity",
            ))
            if key not in seen:
                seen.add(key)
                val_items.append(v)
    r["readiness"] = worst or "—"
    # A wave that covers only a SLICE of its groups' membership (an oversized group split into coupled
    # sub-waves) must not inherit the parent's full host split — same doctrine as the endpoint
    # apportioning above (QA row-13: every ≤40-switch sub-wave row repeated the parent's full
    # '107 hard cutover + 144 make-before-break', and §wi.5 sequenced parent-scale counts). A wave
    # covering its WHOLE group keeps the raw union, so hosts the sequencing lists name beyond the
    # membership record are preserved rather than silently dropped (fail-soft on inconsistent uploads).
    if wsw and members_all and not members_all.issubset(wsw):
        mbb_all = [h for h in mbb_all if h in wsw]
        hard_all = [h for h in hard_all if h in wsw]
    # Overlay the SLICED cutover host split so the procedure branches on THIS wave's members, not just the
    # representative group (the free-text 'sequence' stays the representative's — descriptive only). Copy
    # so the engine's record is never mutated (this joiner reads the engine's records, never recomputes).
    if seq:
        seq = {**seq, "make_before_break": mbb_all, "hard_cutover": hard_all}
    return r, seq, scen, val_items


def _strategy_cell(seq, scen):
    """Count-honest cutover-strategy label derived from THIS wave's (sliced) host split. The engine's
    free-text 'sequence' describes the WHOLE parent group, so on a coupled sub-wave row its counts are
    impossible (QA row-13: a 40-switch row claiming '107 hard cutover + 144 make-before-break'). The
    free text / scenario is the fallback only when the split lists are absent."""
    seq = seq if isinstance(seq, dict) else {}
    scen = scen if isinstance(scen, dict) else {}
    mbb, hard = _as_list(seq.get("make_before_break")), _as_list(seq.get("hard_cutover"))
    if mbb and hard:
        return f"{len(hard)} hard cutover (single-homed) + {len(mbb)} make-before-break (dual-homed)"
    if mbb:
        return f"make-before-break ({len(mbb)} dual-homed)"
    if hard:
        return f"hard cutover ({len(hard)} single-homed)"
    return seq.get("sequence") or scen.get("recommended_scenario") or "—"


# ---- readiness roll-up + window estimate (BLUF facts) -----------------------------------------
# The engine's readiness vocabulary is exactly {"NOT READY","CAUTION","READY"} (analyze.compute_
# migration_readiness). Lower rank = worse; an unknown value defaults least-worst so it never masks a
# real NOT READY. This mirrors _join_group_records' reducer — one shared ordering, not a second copy.
_READY_RANK = {"not ready": 0, "caution": 1, "ready": 2}


def _worst_readiness(waves, readiness_by_group, seq_by_group, scen_by_group, val_by_wave):
    """The go/no-go gate for the whole change: the WORST per-wave readiness verdict across every wave
    (each wave's verdict is itself the worst across its constituent groups, via _join_group_records — one
    source of truth). Returns the verdict string, or None when no wave publishes one (coverage-honest)."""
    worst = None
    for _name, switches, _kind, gnames in waves:
        r, *_ = _join_group_records(gnames, switches, readiness_by_group, seq_by_group,
                                    scen_by_group, val_by_wave)
        rd = r.get("readiness")
        if rd and rd != "—" and (worst is None
                                 or _READY_RANK.get(str(rd).lower(), 9) < _READY_RANK.get(str(worst).lower(), 9)):
            worst = rd
    return worst


def _current_baseline_context(validation_plan, waves):
    """Return the canonical current-baseline gate plus every unbound blocker.

    The gate owner first reconciles ``items`` / ``by_wave`` / ``summary``. Only after that
    integrity check succeeds may this document project rows from the plan. A row is bound to a
    scheduled wave by its declared validation-group label or by device membership; anything else
    remains an explicit fleet-level blocker instead of disappearing from a per-wave MOP.
    """
    from cisco_toolkit.analyze import classify_current_baseline_item, compute_current_baseline_gate

    gate = compute_current_baseline_gate(validation_plan)
    if _as_dict(gate.get("integrity")).get("valid") is not True:
        return gate, [], []

    blockers = []
    for raw in _as_list(_as_dict(validation_plan).get("items")):
        if not isinstance(raw, dict):
            continue
        state = classify_current_baseline_item(raw)
        if state not in {"degraded", "review", "not_verified"}:
            continue
        blockers.append({**raw, "_baseline_state": state})

    scheduled = [
        (set(_as_strs(switches)), set(_as_strs(group_names)))
        for _name, switches, _kind, group_names in waves
    ]
    unbound = [
        row for row in blockers
        if not any(
            (str(row.get("wave") or "") in groups)
            or (str(row.get("device") or "") in switches)
            for switches, groups in scheduled
        )
    ]
    return gate, blockers, unbound


def _window_estimate(n_waves: int) -> str:
    """A DERIVED planning figure for the per-window duration — NOT a collected fact, so it is always
    labelled an estimate to confirm with the change owner (coverage-honesty: a planning heuristic must
    never masquerade as an authoritative number). The engine publishes no per-window duration, so this
    is a standard AS default band, stated as a range to be sized against the wave's real port count."""
    return ("2–4 hours per maintenance window (standard planning estimate — includes pre-cutover "
            "baseline capture, the procedure, post-cutover validation and a rollback contingency; "
            "confirm with the change owner and size against each wave's port count and blast radius).")


def _blast_for(switches, fi_by_host):
    """Max endpoints stranded if one of THIS wave's devices is lost mid-move — or None when the
    failure-impact simulation carries no covering evidence for these switches (coverage-honesty).

    `max(..., default=0)` over `fi_by_host.get(s, {}).get("stranded")` printed a measured-looking
    **0** for an axis that was never collected, and that 0 is load-bearing TWICE: §x.1's scope row,
    and the quantified "roll back if ... more endpoints than the §x.1 max-blast-radius figure are
    affected" trigger — which a fabricated 0 silently rewrites to "roll back if ANY endpoint is
    affected", i.e. always true in a window whose whole purpose is moving endpoints. An observed
    `stranded: 0` is still a real 0; only an ABSENT figure abstains."""
    vals = [_as_num(v) for v in (_as_dict(fi_by_host.get(s)).get("stranded") for s in switches)
            if v is not None]
    return max(vals) if vals else None


def _oob_evidence_for(switches, ifaces_all):
    """Coverage-honest OOB precondition evidence: return the management-port identifier(s) observed for a
    wave's devices (mgmt0 / Management* / Gi0/0 / Fa0/0, or any interface carrying an mgmt_ip), or None
    when the snapshot has NO management-port evidence — in which case the precondition is [NOT OBSERVED]
    (not assumed reachable). Reads only the collected `interfaces` evidence; never infers reachability."""
    import re
    found = []
    for h in switches:
        for pname, dd in sorted((_as_dict(ifaces_all.get(h))).items()):
            dd = _as_dict(dd)
            if re.match(r"^(mgmt|management|gigabitethernet0/0|gi0/0|fastethernet0/0|fa0/0)",
                        str(pname), re.IGNORECASE) or str(dd.get("mgmt_ip") or "").strip():
                tag = f"{h} {pname}"
                ip = str(dd.get("mgmt_ip") or "").strip()
                found.append(tag + (f" ({ip})" if ip else ""))
    return found or None


def _write_bluf(doc, waves, readiness_by_group, seq_by_group, scen_by_group, val_by_wave,
                snap, NAVY, RED, _label_run, table, blockers_for=None,
                current_baseline=None, unbound_baseline_blockers=None):
    """Render the Bottom-Line-Up-Front executive summary (unnumbered H1 so the numbered sections are
    untouched). Answer-first decision facts, all snapshot-grounded and coverage-honest.

    `blockers_for(switches) -> (remediation_items, punchlist_items)` is the SAME callable §1's
    overview table and each wave's §x.2 use. It is required for the blocker row to cite the section
    it cross-references: "blockers" names two different sets in this document — the readiness FAIL
    checks (§x.1's 'Blocking / warning checks') and the Critical/High remediation + punch-list items
    (§1's Blockers column and §x.2's 'Blockers to clear before this window'). The BLUF used to print
    the first under the second's label while pointing the approver at §x.2, so the one-screen
    decision view contradicted the section it sent them to. Both are reported now, each named."""
    n_waves = len(waves)
    gate = _worst_readiness(waves, readiness_by_group, seq_by_group, scen_by_group, val_by_wave)
    current = _as_dict(current_baseline)
    current_verdict = str(current.get("verdict") or "NOT_ASSESSED").upper()
    current_summary = _as_dict(current.get("summary"))
    n_current_blockers = int(_as_num(current_summary.get("n_blockers")))
    unbound = _as_dicts(unbound_baseline_blockers)
    fleet_rec = _as_dict(snap.get("migration_scenarios")).get("fleet_recommendation")
    gating = _as_list(_as_dict(snap.get("executive_brief")).get("top_gating"))
    n_fail_checks = sum(1 for _n, sw, _k, gn in waves
                        for r in [_join_group_records(gn, sw, readiness_by_group, seq_by_group,
                                                      scen_by_group, val_by_wave)[0]]
                        for _ in range(int(r.get("n_fail", 0) or 0)))
    # The §x.2 / §1 set, summed the same way §1's column is read (per wave), so the BLUF total and
    # that column reconcile by addition. None -> the caller did not supply it; abstain rather than
    # recompute a second answer to a question §1 already owns.
    n_blockers = (sum(len(rem) + len(pl) for _n, sw, _k, _gn in waves
                      for rem, pl in [blockers_for(sw)])
                  if callable(blockers_for) else None)

    doc.add_heading("Executive Summary — Bottom Line Up Front (BLUF)", level=1)
    doc.add_paragraph(
        "The one-screen decision view for the change approver: what these windows do, the go/no-go "
        "gate, the rollback in one line, and the planning window. Every figure below is the assessment "
        "snapshot's own (single source of truth) — a fact the snapshot did not publish reads "
        "“[NOT OBSERVED]”, never an assumed value.")

    def _v(x):
        return x if x not in (None, "", "—") else "[NOT OBSERVED]"

    if current_verdict == "CLEAR":
        current_detail = (
            "CLEAR within the bounded validation-plan scope; this is not cutover authorization. "
            "Per-wave readiness still applies."
        )
        combined_gate = (
            f"{gate} — the current baseline is CLEAR within its bounded scope; "
            "do NOT open a window whose wave is NOT READY until its blockers are cleared or risk-accepted."
            if gate else
            "HOLD — no per-wave readiness verdict was published; treat every wave as no-go until readiness is assessed."
        )
    elif current_verdict == "BLOCKED":
        current_detail = (
            f"HOLD — current baseline BLOCKED: {n_current_blockers} degraded/review/not-verified "
            "validation row(s) require resolution or explicit disposition before any window. "
            f"{len(unbound)} blocker(s) are not bound to a scheduled wave"
            + ("; see Unscheduled current-baseline blockers." if unbound else ".")
        )
        combined_gate = current_detail
    else:
        current_detail = (
            f"HOLD — current baseline {current_verdict}; the validation-plan contract does not authorize "
            "an all-clear. Recollect or reconcile before any window."
        )
        combined_gate = current_detail

    rows = [
        ("What this MOP does",
         f"Cuts the migration over in {n_waves} candidate wave(s), each in its own maintenance window "
         "(clear blockers → baseline → procedure → validate → proceed or roll back)."),
        ("Per-wave worst readiness (scheduled waves)",
         (str(gate) if gate else
          "[NOT OBSERVED] — no per-wave readiness verdict was published; treat every wave as no-go "
          "until readiness is assessed.")),
        ("Current baseline gate (all validation groups)", current_detail),
        ("Go / no-go gate", combined_gate),
        ("Open blockers gating the change",
         ((f"{n_blockers} Critical/High blocker(s) (remediation + punch-list) attributed across the "
           "waves — the same set §1's Blockers column totals and each wave's Blockers-to-clear "
           "subsection lists."
           if n_blockers is not None else
           "[NOT OBSERVED] — the per-wave blocker attribution was not available to this summary; "
           "read each wave's Blockers-to-clear subsection directly.")
          + f" Separately, {n_fail_checks} migration-readiness check(s) FAIL across the waves "
            "(§x.1 'Blocking / warning checks') — a different count of a different thing; both "
            "must be cleared or risk-accepted.")),
        ("Rollback in one line",
         "If any wave's post-cutover validation fails and cannot be corrected inside the window, "
         "re-apply that wave's captured pre-change config and re-home to the legacy path — every wave "
         "carries an explicit quantified rollback trigger (in each wave's Rollback subsection)."),
        ("Maintenance window (estimate)", _window_estimate(n_waves)),
        ("Fleet-level recommendation", _v(fleet_rec)),
        ("Top gating item to clear first", _v(gating[0] if gating else None)),
    ]
    table(["Bottom line", "Detail"], rows, widths=[2.4, 4.3])
    # the go/no-go gate restated as a red one-liner so it cannot be missed
    _label_run(doc.add_paragraph(), "GO/NO-GO:",
               combined_gate
               + " No window opens without CAB approval, a tested rollback, and the wave's blockers "
                 "cleared or explicitly risk-accepted.", RED)

    if unbound:
        doc.add_heading("Unscheduled current-baseline blockers", level=2)
        doc.add_paragraph(
            f"{len(unbound)} validated current-baseline blocker(s) are not bound to any scheduled wave. "
            "They remain fleet-level HOLD conditions: re-plan them into scope, resolve them, or explicitly "
            "disposition them before opening any window. Every row is retained below.")
        table(
            ["State", "Validation group", "Device", "Check", "Command", "Acceptance / evidence"],
            [
                (
                    row.get("_baseline_state") or "review",
                    row.get("wave") or "(unscheduled)",
                    row.get("device") or "—",
                    row.get("check") or "Current baseline blocker",
                    row.get("command") or "—",
                    str(row.get("expect") or "")
                    + f"\nCustody: {row.get('projection_custody') or '—'}"
                    + f" | Source: {row.get('source_key') or '—'}",
                )
                for row in unbound
            ],
            widths=[0.7, 1.1, 1.0, 1.5, 1.3, 1.8],
        )


def write_mop_docx(output_path: str, snap_dict: dict, label: str) -> None:
    """Emit the per-wave Method of Procedure (.docx) to `output_path`. Fail-soft: a missing python-docx
    is a warning + skip, and every snapshot read is defensive so malformed input degrades to placeholders.
    A genuine render error still propagates to the caller — the CLI phase wrapper logs-and-continues,
    AssessHub turns it into a clean 500 (docstring corrected: the old 'logged, never raised' claim was
    false — the only `try` here is the optional-import guard — and the crash class it hid is now guarded
    at the read sites; matches engagement.write_engagement_docx)."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor
    except ImportError:
        logger.warning("  MOP (DOCX) skipped: python-docx not installed "
                       "(pip install python-docx to enable the Method-of-Procedure deliverable).")
        return

    snap = xml_safe_deep(snap_dict if isinstance(snap_dict, dict) else {})   # one bad device byte (noncharacter/surrogate) must not abort the docx
    label = xml_safe(label) if isinstance(label, str) else (str(label) if label is not None else "")
    from cisco_toolkit.brand_tokens import DOC_NAVY_RGB
    NAVY = RGBColor(*DOC_NAVY_RGB)
    GREY = RGBColor(0x59, 0x59, 0x59)
    RED = RGBColor(0xB0, 0x2A, 0x1E)
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    def _label_run(p, label_text, value, color=NAVY):
        r = p.add_run(label_text); r.bold = True; r.font.color.rgb = color
        p.add_run(" " + (str(value) if value not in (None, "") else "—"))

    def table(headers, rows, widths=None):
        # Delegates to the family's single table builder (docmeta.add_table) — V3.23.152 dedup.
        return add_table(doc, headers, rows, widths, fixed=False)

    def steps(items):
        # Manual numbering rather than Word's built-in "List Number" style: that style's single shared
        # numbering definition makes numbers CONTINUE across every list block in the document, so a wave's
        # procedure and rollback (and the next wave's procedure) would not restart at 1. Explicit prefixes
        # restart per call and render identically in every viewer.
        # An item may be a plain string, or an (action, success_criterion) tuple — the MOP definition in
        # the AS migration service is steps + precautions + SUCCESS CRITERIA per step (V3.23.155).
        for i, s in enumerate(items, 1):
            if isinstance(s, tuple):
                action, success = s
                p = doc.add_paragraph(f"{i}. {action}")
                if success:
                    r = p.add_run("  Success: " + success)
                    r.bold = True
                    r.font.color.rgb = NAVY
            else:
                doc.add_paragraph(f"{i}. {s}")

    # ---- snapshot-derived facts ----
    # Shared coercers (not `or {}`/`or []`, which do NOT guard a truthy non-dict/list) so a malformed
    # uploaded snapshot -- where a section that should be a list is an object, etc. -- degrades instead of
    # char-iterating or raising AttributeError mid-render (matching engagement/ops/archreview).
    devices = _as_dict(snap.get("devices"))
    waves = _wave_sections(snap)
    readiness_by_group = {d.get("group"): d for d in (_as_dict(r) for r in _as_list(snap.get("migration_readiness")))}
    seq_by_group = {d.get("group"): d for d in (_as_dict(r) for r in _as_list(snap.get("wave_sequencing")))}
    scen_by_group = {d.get("group"): d
                     for d in (_as_dict(r) for r in _as_list(_as_dict(snap.get("migration_scenarios")).get("per_group")))}
    val_by_wave = _as_dict(_as_dict(snap.get("validation_plan")).get("by_wave"))
    current_baseline, _all_baseline_blockers, unbound_baseline_blockers = _current_baseline_context(
        snap.get("validation_plan"), waves)
    current_baseline_verdict = str(current_baseline.get("verdict") or "NOT_ASSESSED").upper()
    fi_by_host = {d.get("host"): d for d in (_as_dict(r) for r in _as_list(snap.get("failure_impact")))}
    rem_items = [_as_dict(r) for r in _as_list(_as_dict(snap.get("remediation_plan")).get("items"))]
    punchlist = [_as_dict(r) for r in _as_list(snap.get("punchlist"))]

    def _validation_blocker(it):
        state = str(it.get("evidence_state") or "").strip().lower()
        expected = str(it.get("expect") or "").strip().upper()
        return state in {"degraded", "review", "not_verified"} or expected.startswith(
            ("PRE-CUTOVER DEGRADED — BLOCKER:", "PRE-CUTOVER REVIEW — BLOCKER:",
             "FHRP CONFIGURED GROUP NOT VERIFIED — BLOCKER:",
             "ROUTING BASELINE NOT VERIFIED — BLOCKER:",
             "ETHERCHANNEL BASELINE NOT VERIFIED — BLOCKER:")
        )

    def _validation_acceptance(it):
        acceptance = str(it.get("expect") or "")
        state = str(it.get("evidence_state") or "").strip()
        custody = str(it.get("projection_custody") or "").strip()
        source = str(it.get("source_key") or "").strip()
        if not (state or custody or source):
            return acceptance
        if custody == "embedded_unverified":
            custody = "embedded_unverified (embedded projection; integrity unverified)"
        evidence = (f"Evidence state: {state or '—'} | Projection custody: {custody or '—'} | "
                    f"Source: {source or '—'}")
        return acceptance + ("\n" if acceptance else "") + evidence

    def _blockers_for(switches):
        sw = set(switches)
        # high-severity remediation snippets that touch this wave's devices
        rem = [it for it in rem_items if it.get("device") in sw
               and _SEV_RANK.get(it.get("severity"), 4) <= 1]
        # critical/high punch-list items that touch this wave's devices
        pl = [i for i in punchlist if _SEV_RANK.get(i.get("severity"), 4) <= 1
              and (set(_as_list(i.get("devices"))) & sw)]
        return rem, pl

    # ---- title page ----
    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Migration Method of Procedure (MOP)"); tr.bold = True
    tr.font.size = Pt(26); tr.font.color.rgb = NAVY
    sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s2 = sub2.add_run("Per-wave maintenance-window cutover procedure")
    s2.font.size = Pt(14); s2.font.color.rgb = GREY
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(label); sr.font.size = Pt(13); sr.font.color.rgb = GREY
    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _ninv = _as_dict(_as_dict(snap.get("executive_brief")).get("scale")).get("n_devices")   # canonical inventoried (SSOT owner); len() only pre-brief (CROSS-04)
    meta.add_run(f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')}  ·  "
                 f"{len(waves)} wave(s)  ·  {_ninv if isinstance(_ninv, int) else len(devices)} devices  ·  script {snap.get('script_version', '')}"
                 ).font.color.rgb = GREY
    status = doc.add_paragraph(); status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st = status.add_run("DRAFT TEMPLATE — fill every <placeholder>, peer-review, and dry-run before any window.")
    st.italic = True; st.font.color.rgb = RED
    doc.add_paragraph()
    note = doc.add_paragraph()
    _label_run(note, "How to use this MOP:",
               "This is a change template primed with the assessment evidence, NOT an auto-executed "
               "change. Each wave below is one maintenance window. Work top to bottom: clear the wave's "
               "blockers, capture the pre-cutover baseline, run the procedure, then PROVE the wave with "
               "the post-cutover validation checks — and only proceed to the next wave once they pass. "
               "If any validation check fails and cannot be corrected within the window, execute the "
               "rollback. Every <placeholder> (date, owner, approver, exact uplinks) must be completed "
               "by the implementing engineer.", GREY)
    doc.add_page_break()

    # ---- Executive summary (BLUF — Bottom Line Up Front) ------------------------------------------
    # DE-01 / Cisco-AS: answer-first executive summary at the TOP of the MOP so a change approver reads
    # the window intent, the go/no-go gate, the one-line rollback and the window estimate WITHOUT paging
    # into the numbered procedure. Every fact is the snapshot's own (never invented): wave count from
    # `waves`, the go/no-go gate is the WORST readiness verdict across the waves (via _join_group_records,
    # one source of truth), the recommendation + top-gating are the engine's. A fact the snapshot did not
    # publish reads "[NOT OBSERVED]", never a fabricated value (coverage-honesty).
    _write_bluf(doc, waves, readiness_by_group, seq_by_group, scen_by_group, val_by_wave,
                snap, NAVY, RED, _label_run, table, blockers_for=_blockers_for,
                current_baseline=current_baseline,
                unbound_baseline_blockers=unbound_baseline_blockers)
    doc.add_page_break()

    # ---- document control (AS-style front matter; unnumbered so the wave sections are untouched) ----
    add_document_control(
        doc, document="Migration Method of Procedure (MOP)", label=label,
        engine_version=str(snap.get("script_version", "")), generated_at=snap.get("generated_at"),
        collected_at=snap.get("collected_at"),
        audience="The implementing engineer(s) executing each maintenance window, the validation / "
                 "war-room lead, and the change manager approving the windows.",
        exclude=("mop",),
        extra_assumptions=(
            "“Ready for service” for a wave means the environment functions as per the "
            "post-cutover validation checks documented in that wave's section — acceptance criteria "
            "live in this MOP, not in a separate test document.",))
    doc.add_page_break()

    # ---- table of contents (shared field-code helper, V3.23.171) ----
    add_toc(doc)


    # Deliverable Excellence (DE-01): answer-first at-a-glance register + single-source-of-truth signal.
    add_excellence_front(doc, snap_dict)
    # Deliverable Excellence (DE-01): consolidated Inputs-Required register (Law 9).
    add_inputs_required(doc, snap_dict)
    # ===== 1. Change overview =====
    doc.add_heading("1. Change Overview", level=1)
    doc.add_paragraph(
        f"This migration is sequenced into {len(waves)} candidate wave(s) (right-sized from the L2-coupling "
        "move-groups), each cut over in its own maintenance window. The recommended order is: clear all "
        "blocking items first, then cut the lower-risk waves to build confidence, protecting the "
        "highest-blast-radius (keystone) devices.")
    ov_rows = []
    for name, switches, _kind, _gnames in waves:
        r, seq, _scen, _vi = _join_group_records(_gnames, switches, readiness_by_group, seq_by_group,
                                                 scen_by_group, val_by_wave)
        blast = _blast_for(switches, fi_by_host)
        rem, pl = _blockers_for(switches)
        ov_rows.append((name, len(switches), r.get("endpoints", "—"),
                        r.get("readiness", "—"), _strategy_cell(seq, _scen),
                        blast if blast is not None else "[NOT OBSERVED]", len(rem) + len(pl)))
    table(["Wave", "Devices", "Endpoint MACs", "Readiness", "Strategy", "Max blast", "Blockers"],
          ov_rows, widths=[1.5, 0.8, 1.0, 1.2, 1.5, 0.9, 0.9])

    fleet_rec = _as_dict(snap.get("migration_scenarios")).get("fleet_recommendation")
    if fleet_rec:
        _label_run(doc.add_paragraph(), "Fleet-level recommendation:", fleet_rec)

    # ===== 2. Global prerequisites =====
    doc.add_heading("2. Global Prerequisites (before any window)", level=1)
    doc.add_paragraph(
        "Complete these once, before the first wave. They are the fleet-wide gating items the assessment "
        "flagged; cutting over before they are resolved or risk-accepted carries the documented risk.")
    # element-coerced: unlike a table cell (add_table str()s every value), add_paragraph() takes the raw
    # element and char-iterates it -- a scalar gating item raised TypeError mid-document.
    gating = _as_strs(_as_dict(snap.get("executive_brief")).get("top_gating"))
    if gating:
        for g in gating[:6]:
            doc.add_paragraph(g, style="List Bullet")
        # Disclose the cap (same class as §x.2's blocker lists). This section is a PRECONDITION for
        # the whole change -- "complete these once, before the first wave ... cutting over before
        # they are resolved or risk-accepted carries the documented risk" -- so a silent 6-of-9 list
        # makes an engineer who cleared everything printed here believe the fleet-wide gate is met.
        # The engine publishes the full ordered set; the deck already says "5 of 9 shown" for the
        # same list, so a silent MOP was also the set-wide outlier.
        if len(gating) > 6:
            doc.add_paragraph(
                f"…and {len(gating) - 6} further fleet-wide gating item(s) NOT listed above "
                f"({len(gating)} in total) — this prerequisite covers all of them, not only the 6 "
                "shown. Work the full set from the workbook's Executive Summary before the first "
                "wave.")
    steps([
        "Confirm the assessment workbook, runbook and this MOP are the agreed change record, and that "
        "the change ticket <CHG-NUMBER> references them.",
        "Obtain change-advisory-board approval and a signed maintenance window per wave "
        "(<DATE/TIME>, owner <NAME>, approver <NAME>).",
        "Verify out-of-band/console access to every device in scope, and that current configurations "
        "are backed up off-box (golden snapshot).",
        "Stage and peer-review the remediation/config snippets from the Remediation Plan for each wave's "
        "blockers (review-only — do not paste un-reviewed).",
        "Brief the war room: roles, the rollback decision-maker, and the go/no-go criteria (the "
        "post-cutover validation must pass).",
    ])

    # 2.1 responsibilities — the AS migration-service ownership split, stated up front
    doc.add_heading("2.1 Responsibilities (RACI)", level=2)
    doc.add_paragraph(
        "Ownership follows the standard AS migration-service split: the delivery engineer prepares "
        "the MOP, mappings and pre/post checks; the customer's team executes the steps in the window "
        "and owns the key in-window decisions. Adjust names per the engagement before approval.")
    table(["Activity", "Responsible", "Notes"], [
        ("MOP preparation, port mapping & staged configuration", "Delivery engineer (author)",
         "Reviewed and approved by the customer network owner"),
        ("Pre/post check definition & baseline capture", "Delivery engineer (author)",
         "From the assessment's validation plan — one source of truth"),
        ("Change approval & window scheduling", "Customer change manager",
         "Via the customer's change-advisory process"),
        ("In-window execution of each step", "Customer operations / implementing engineer",
         "Delivery engineer advises; steps are executed exactly as written"),
        ("Go/no-go and rollback decision", "Customer network owner (change owner)",
         "At the rollback decision time stated in each wave's scope table"),
        ("Validation evidence & sign-off records", "Validation / war-room lead",
         "Archived with the post-implementation review"),
    ], widths=[2.6, 2.0, 2.2])

    # Communications & escalation plan (Cisco-AS standard) — the roles, the T-minus comms cadence and
    # the escalation matrix (incl. Cisco TAC pre-open for high-risk windows). Unnumbered H2 so the
    # conditional §2.2/§2.3 EVPN/software numbering (pinned by tests) is untouched.
    doc.add_heading("Communications & escalation plan", level=2)
    doc.add_paragraph(
        "Who runs the window, who is told what and when, and how a problem escalates. Fill the "
        "<contact> placeholders per engagement before approval; the roles and cadence are the standard "
        "AS change-communications model.")
    doc.add_paragraph("War-room roles", style="List Bullet")
    table(["Role", "Who", "Responsibility"], [
        ("Change owner (go/no-go)", "<name / on-call>",
         "Owns the go/no-go and rollback decision; the only person who calls a rollback."),
        ("Executor (implementing engineer)", "<name>",
         "Executes each MOP step exactly as written; calls out each success criterion."),
        ("Verifier (validation / war-room lead)", "<name>",
         "Runs each wave's post-cutover validation independently of the executor; records evidence."),
        ("Escalation / bridge lead", "<name>",
         "Runs the conference bridge, drives escalation, owns stakeholder comms."),
    ], widths=[2.0, 1.6, 3.1])
    doc.add_paragraph("T-minus communications cadence", style="List Bullet")
    table(["When", "Message", "Audience"], [
        ("T-minus 1 week", "Window confirmed, CAB approved, MOP + rollback distributed.",
         "Stakeholders, NOC, change owner"),
        ("T-minus 1 day", "Go/no-go readiness re-confirmed; preconditions checklist reviewed.",
         "War-room roles, change owner"),
        ("T-minus 1 hour", "War room open, roll-call, OOB access + backups confirmed.", "War-room roles"),
        ("T-0 (start)", "Window start announced; outage clock started (if hard cutover).",
         "Stakeholders, NOC"),
        ("Per step / gate", "Each success criterion and each go/no-go gate called on the bridge.",
         "War-room roles"),
        ("T-plus (close)", "Window closed as PROCEEDED or ROLLED-BACK; validation evidence archived.",
         "Stakeholders, NOC, change owner"),
    ], widths=[1.3, 3.3, 2.1])
    doc.add_paragraph("Escalation matrix", style="List Bullet")
    table(["Tier", "Trigger", "Contact / action", "Response target"], [
        ("L1 — war room", "A step fails its success criterion.",
         "Executor → verifier → change owner on the bridge.", "Immediate"),
        ("L2 — engineering", "A trigger fires or a fault is not understood in one cycle.",
         "Escalation lead → senior network engineering / <SME>.", "<15 min>"),
        ("L3 — vendor (Cisco TAC)", "A device/software fault needs vendor help.",
         "Open/attach the Cisco TAC case (SEV per impact); reference the pre-opened SR.", "<per SLA>"),
        ("Rollback authority", "The rollback decision time passes or a trigger stands.",
         "Change owner declares rollback; executor runs that wave's Rollback procedure.", "Immediate"),
    ], widths=[1.4, 2.1, 2.4, 0.9])
    _label_run(doc.add_paragraph(), "High-risk windows:",
               "for any NOT-READY wave or a wave with Critical/High blockers, PRE-OPEN a Cisco TAC case "
               "before the window and hold the SR number in the war room, so L3 escalation is immediate "
               "rather than started mid-incident (see each wave's pre-implementation checklist).", RED)

    # §2.2 EVPN-migration guardrails — gated, evidence-grounded brownfield->NX-OS-VXLAN-EVPN cutover doctrine.
    # Silent unless a VXLAN-EVPN fabric is the target; renders the load-bearing, primary-source migration
    # gotchas (NX-OS 10.2(3) HSRP/DAG gate, the gateway-vMAC pre-step, single-active-L2-interconnect loop
    # safety, the vPC back-to-back method, rollback triggers) grounded in the fleet's OWN observed evidence.
    evpn = _as_dict(_as_dict(snap.get("design_blueprint")).get("evpn_migration"))
    evpn_on = bool(evpn.get("applicable"))
    # Match the CRD's confidence (adversarial-review finding, 2026-07-05): assert "the target IS EVPN"
    # only when the customer's requirements register confirms fabric_operating_model=nxos-evpn; otherwise
    # this is an ENGINE assessment from the fleet's evidence and must be framed as a to-confirm target-state,
    # never a settled plan-of-record.
    _rm = _as_dict(_as_dict(snap.get("design_blueprint")).get("requirements_model"))
    _fom = next((str(f.get("value") or "") for f in _as_list(_rm.get("fields"))
                 if isinstance(f, dict) and f.get("key") == "fabric_operating_model"), "")
    evpn_confirmed = bool(_rm.get("provided")) and "evpn" in _fom.lower()
    if evpn_on:
        doc.add_heading("2.2 EVPN-migration guardrails (brownfield → NX-OS VXLAN-EVPN)", level=2)
        _basis = (
            f"The target fabric is NX-OS VXLAN BGP-EVPN ({evpn.get('model_basis', '')}), customer-confirmed "
            "via the requirements register."
            if evpn_confirmed else
            f"The assessed target-state fabric is modeled as NX-OS VXLAN BGP-EVPN ({evpn.get('model_basis', '')}) "
            "— an engine assessment from this fleet's evidence, NOT a settled plan-of-record; confirm the fabric "
            "operating model with the customer before cutover (a requirements register with fabric_operating_model "
            "settles it).")
        doc.add_paragraph(
            _basis + " These migration guardrails are grounded in this fleet's own evidence and documented in "
            "primary Cisco sources; each is an easily-missed cutover failure mode that must be satisfied — or "
            "explicitly risk-accepted — before the relevant wave.")
        table(["Guardrail", "Phase", "Severity", "Basis (observed) & action — [source]"],
              [(g.get("title", ""), g.get("phase", ""), g.get("severity", ""),
                f"{g.get('basis', '')} — {g.get('detail', '')} [{g.get('source', '')}]")
               for g in _as_list(evpn.get("guardrails")) if isinstance(g, dict)],
              widths=[1.7, 0.9, 0.8, 3.4])

    # §2.3/§2.2 software / image standardization (per in-scope platform) — N24 (number follows the EVPN subsection)
    swrisk_pd = [d for d in _as_list(_as_dict(snap.get("software_risk")).get("per_device")) if isinstance(d, dict)]
    if swrisk_pd:
        doc.add_heading(f"{'2.3' if evpn_on else '2.2'} Software / image standardization", level=2)
        doc.add_paragraph(
            "Before any wave, agree the target NOS image per platform. The current software trains and their "
            "disposition (from the software-advisory screening) are below; the target image is selected at "
            "detailed design and applied + verified per wave in the staged config and post-cutover checks.")
        agg: dict = {}
        for d in swrisk_pd:
            key = (str(d.get("platform") or "—"), str(d.get("train") or "—"), str(d.get("train_band") or "—"))
            agg[key] = agg.get(key, 0) + 1
        rows = sorted(agg.items(), key=lambda kv: (-kv[1], kv[0]))
        table(["Platform", "Current train", "Disposition", "Devices", "Target image"],
              [(p, tr, band, n, "<confirm>") for (p, tr, band), n in rows][:30],
              widths=[1.0, 1.5, 1.6, 0.7, 1.4])

    # ===== per-wave sections =====
    for wi, (name, switches, _kind, _gnames) in enumerate(waves, start=3):
        r, seq, scen, _val_items = _join_group_records(_gnames, switches, readiness_by_group, seq_by_group,
                                                       scen_by_group, val_by_wave)
        playbook = _as_dict(scen.get("playbook"))
        rem, pl = _blockers_for(switches)
        blast = _blast_for(switches, fi_by_host)
        verdict = r.get("readiness", "—")

        # The constituent move-groups this wave covers (a coupled-subwave maps to ONE group; an
        # independent-batch may bundle MANY). Each keeps its own strategy/rationale/rollback -- surface
        # every one, never just the first (the joiner picks a representative seq/scen, but the document
        # must not silently drop the siblings').
        wave_groups = [(g, scen_by_group.get(g) or {}, seq_by_group.get(g) or {}) for g in _gnames]
        wave_groups = [(g, sc, sq) for g, sc, sq in wave_groups if sc or sq]
        multi = len(wave_groups) > 1
        if multi:
            strat_cell = "; ".join(dict.fromkeys(
                str(sq.get("sequence") or sc.get("recommended_scenario") or "").strip()
                for _g, sc, sq in wave_groups if (sq.get("sequence") or sc.get("recommended_scenario")))) or "—"
        else:
            strat_cell = _strategy_cell(seq, scen)

        doc.add_heading(f"{wi}. MOP — {name}", level=1)

        # 2.x.1 scope
        doc.add_heading(f"{wi}.1 Scope & risk", level=2)
        table(["Field", "Value"], [
            ("Devices in scope", ", ".join(switches) or "—"),
            ("Endpoint MACs (apportioned)", r.get("endpoints", "—")),
            ("Readiness verdict", verdict),
            ("Fleet current baseline verdict", current_baseline_verdict),
            ("Blocking / warning checks", f"{r.get('n_fail', 0)} / {r.get('n_warn', 0)}"),
            ("Cutover strategy", strat_cell),
            ("Max blast radius (endpoints stranded if a device is lost mid-move)",
             blast if blast is not None else
             "[NOT OBSERVED] — no failure-impact evidence for this wave's devices; do NOT read this "
             "as zero (it is the threshold the Blast-radius rollback trigger below would size on)"),
            ("Maintenance window", "<DATE> <START>–<END>  ·  owner <NAME>  ·  approver <NAME>"),
            ("Rollback decision time", "<HH:MM> — if validation has not passed by this time, roll back"),
        ], widths=[2.6, 4.2])
        if multi:
            doc.add_paragraph(
                "This wave bundles several INDEPENDENT move-groups (they share no VLANs, so they may be "
                "cut in any order within the window — but each must pass its own validation before the "
                "next). Each keeps its own strategy and rollback:")
            table(["Move-group", "Strategy", "Why"],
                  [(g, (sq.get("sequence") or sc.get("recommended_scenario") or "—"),
                    sc.get("rationale") or "—") for g, sc, sq in wave_groups], widths=[1.4, 1.8, 3.6])
        elif scen.get("rationale"):
            _label_run(doc.add_paragraph(), "Why this strategy:", scen.get("rationale"))
        if _kind == "coupled-subwave":
            _label_run(doc.add_paragraph(), "Shared-L2 coordination (coupled sub-wave):",
                       "this wave is one slice of an oversized L2-coupled move-group; the sub-waves SHARE "
                       "VLANs, so they are a SEQUENCE (order matters), not independent cutovers. Keep the "
                       "shared VLANs trunked/extended across the sibling sub-waves until the whole group "
                       "is migrated; STP root, FHRP active and trunk allowed-VLAN lists are NOT independent "
                       "of the siblings.", RED)

        # pre-implementation checklist — the precondition gate for THIS window (Cisco-AS standard).
        # Evidence-gated where the snapshot can genuinely check the precondition (OOB management-port
        # evidence; the wave's own readiness/blocker gate); human-attested where it cannot (backup taken,
        # CAB approval, rollback tested). Coverage-honesty: an OOB precondition with no mgmt-port evidence
        # is [NOT OBSERVED], never assumed reachable. Unnumbered H2 so the existing .1-.8 chain is intact.
        high_risk = (str(verdict).upper() == "NOT READY") or bool(pl) or bool(rem)
        oob_ev = _oob_evidence_for(switches, _as_dict(snap.get("interfaces")))
        doc.add_heading("Pre-implementation checklist (precondition gate)", level=2)
        doc.add_paragraph(
            "Do NOT open this window until every precondition below is met. Each row is either "
            "EVIDENCE-gated (the assessment snapshot can confirm it) or ATTESTED (the change owner must "
            "confirm it — it is not machine-verifiable). Coverage-honesty: an evidence-gated item with no "
            "collected evidence reads “[NOT OBSERVED]”, never assumed satisfied.")
        chk_rows = [
            ("Out-of-band / console access to every device in scope", "Evidence",
             # The bare "…" marked THAT something was dropped but never WHAT: this row is the
             # evidence gate for "OOB / console access to EVERY device in scope", so the engineer
             # must know the gate covers 35 management ports, not the 6 printed. Name the remainder
             # and the total (the house disclosure pattern), never an anonymous ellipsis.
             (f"observed management-port(s): {'; '.join(oob_ev[:6])}"
              + (f" …and {len(oob_ev) - 6} further management-port(s) not listed here "
                 f"({len(oob_ev)} observed in total)" if len(oob_ev) > 6 else "")
              + " — confirm each is reachable before the window")
             if oob_ev else
             "[NOT OBSERVED] — no management-port evidence for this wave's devices; confirm OOB "
             "reachability out-of-band before the window (do NOT assume it)."),
            ("Config backup taken off-box (golden snapshot)", "Attested",
             "Confirm a fresh off-box config backup of every in-scope device (captured in §baseline below)."),
            ("Maintenance window confirmed & communicated", "Attested",
             "Date/time agreed, owner + approver named, stakeholders notified per the §2 "
             "Communications & escalation plan."),
            ("Change-advisory-board (CAB) approval obtained", "Attested",
             "Signed change record referencing this MOP, the workbook and the runbook."),
            ("Rollback tested / dry-run walked", "Attested",
             "The §rollback for this wave has been walked or dry-run; the rollback owner is identified."),
            ("Wave blockers cleared or explicitly risk-accepted", "Evidence",
             (f"readiness {verdict}; {len(pl)} Critical/High punch-list + {len(rem)} remediation item(s) "
              "attributed (see §blockers) — clear or risk-accept before proceeding")
             if high_risk else
              f"readiness {verdict}; no Critical/High blocker attributed to this wave's devices."),
            ("Current baseline gate across every validation group", "Evidence",
             ("CLEAR within the bounded validation-plan scope; this is not cutover authorization. "
              "Per-wave readiness still applies."
              if current_baseline_verdict == "CLEAR" else
              f"HOLD — current baseline {current_baseline_verdict}; do not open this or any other window "
              "until the validation-plan gate is reconciled and every blocker is resolved or dispositioned.")),
        ]
        if high_risk:
            chk_rows.append((
                "Cisco TAC case pre-opened (high-risk window)", "Attested",
                "This wave is HIGH-RISK (NOT-READY and/or Critical/High blockers). Pre-open a Cisco TAC "
                "case and have the SR number in the war room before the window, so escalation is "
                "immediate — do not open the case mid-incident."))
        table(["Precondition", "Gate", "Status / evidence"], chk_rows, widths=[2.4, 0.9, 3.4])

        # 2.x.2 blockers to clear first
        doc.add_heading(f"{wi}.2 Blockers to clear before this window", level=2)
        if not rem and not pl:
            doc.add_paragraph("No Critical/High blockers attributed to this wave's devices. Proceed once "
                              "the global prerequisites are met.")
        else:
            # Both lists are DISPLAY-capped, and this is the one section the precondition gate points
            # at: that row states the FULL attributed counts and says "clear or risk-accept before
            # proceeding (see §blockers)". Silently rendering 10 of 80 punch-list items made an
            # engineer who cleared everything printed here believe the gate was met — the window opens
            # on 70 uncleared Critical/High blockers. Every other capped table in this writer discloses
            # its overflow (§x.3 baselines, §x.4 port map, §x.6 checks); these two did not.
            if pl:
                table(["Severity", "Category", "Item", "Devices"],
                      [(i.get("severity"), i.get("category") or "—", i.get("title") or "—",
                        ", ".join(str(x) for x in _as_list(i.get("devices")) if x in set(switches)))
                       for i in pl[:10]], widths=[0.9, 1.3, 3.0, 1.6])
                if len(pl) > 10:
                    doc.add_paragraph(
                        f"…and {len(pl) - 10} further Critical/High punch-list item(s) attributed to "
                        f"this wave and NOT listed above ({len(pl)} in total) — the precondition gate "
                        "covers all of them, not only the 10 shown. Work the full set from the "
                        "workbook's Punch List sheet before opening this window.")
            for it in rem[:8]:
                p = doc.add_paragraph()
                _label_run(p, f"{it.get('device')} — {it.get('title') or it.get('source') or 'fix'}"
                              f" ({it.get('severity', '—')}):",
                           it.get("why") or "review-only config change; see the Remediation Plan sheet.")
                for c in _as_strs(it.get("commands"))[:8]:      # same add_paragraph() element class
                    doc.add_paragraph(c, style="List Bullet")
            if len(rem) > 8:
                doc.add_paragraph(
                    f"…and {len(rem) - 8} further Critical/High remediation item(s) for this wave's "
                    f"devices and NOT listed above ({len(rem)} in total) — see the workbook's "
                    "Remediation Plan sheet; the precondition gate covers all of them.")

        # 2.x.3 pre-cutover baseline capture
        doc.add_heading(f"{wi}.3 Pre-cutover baseline capture", level=2)
        doc.add_paragraph(
            "Capture and SAVE the output of these commands before any change, so the pre-change state is "
            "recorded and each post-cutover check can be evaluated against its stated acceptance criterion. "
            "A captured degraded state is evidence to resolve or disposition, not a definition of 'good'.")
        # The CONFIGURATION capture is unconditional, and must stay that way: §x.7's rollback says
        # "re-apply the pre-change configuration captured in §x.3". Before this, that artifact only
        # existed on the FALLBACK branch below -- when the snapshot carried a validation plan (the
        # normal case) this section rendered only the plan's show commands, none of which is
        # 'show running-config'. So on the hard-cutover path, which tears the legacy config down, the
        # rollback instructed the engineer to restore a file nobody had been told to make -- mid
        # outage, inside the window, on production gear. A rollback that cites a non-existent artifact
        # is not a rollback, and CLAUDE.md guardrail 1 requires production change to carry a real one.
        steps([f"CONFIGURATION BACKUP (required for the §{wi}.7 rollback): on each of "
               f"{', '.join(switches) or 'the wave devices'}, capture 'show running-config' and save "
               f"it to the change record. §{wi}.7 restores from exactly this file; without it the "
               f"rollback cannot be executed."])
        val_items = _val_items
        cap_cmds = []
        seen_cap = set()
        for it in val_items:
            dev, cmd = it.get("device"), it.get("command")
            if cmd and (dev, cmd) not in seen_cap:
                seen_cap.add((dev, cmd))
                cap_cmds.append((dev, cmd))
        if cap_cmds:
            doc.add_paragraph("State baseline for the post-cutover comparison in "
                              f"§{wi}.6 (in addition to the configuration backup above):")
            blocker_capture_keys = {
                (it.get("device"), it.get("command")) for it in val_items
                if _validation_blocker(it) and it.get("command")
            }
            blocker_captures = [row for row in cap_cmds if row in blocker_capture_keys]
            ordinary_captures = [row for row in cap_cmds if row not in blocker_capture_keys]
            # An evidence-state blocker without its pre-change command capture cannot be dispositioned. Retain
            # every such capture outside the ordinary 20-command display allowance.
            shown = blocker_captures + ordinary_captures[:20]
            table(["Device", "Command to capture"], shown, widths=[1.8, 4.6])
            # Disclose the cap. §x.6 renders up to 40 checks and tells the engineer every one must
            # match "its captured baseline"; silently listing 20 baselines against 40 checks makes the
            # surplus either a guess or an unresolvable comparison. Every other capped table in this
            # writer discloses its overflow; this one did not.
            omitted_captures = len(ordinary_captures) - min(len(ordinary_captures), 20)
            if omitted_captures:
                doc.add_paragraph(
                    f"…and {omitted_captures} further ordinary device/command pair(s) not listed "
                    f"here — capture the full validation-plan set for this wave, not only the "
                    f"{len(shown)} shown, or the §{wi}.6 checks beyond that point have no baseline. "
                    "Every evidence-state blocker capture is retained above.")
        else:
            steps([f"On each of {', '.join(switches) or 'the wave devices'}: also capture "
                   "'show ip interface brief', 'show cdp neighbors', "
                   "'show spanning-tree summary', and (if L3) 'show ip route' / 'show standby brief'."])

        # 2.x.4 port / interface mapping & staged configuration (AS migration-service deliverable)
        doc.add_heading(f"{wi}.4 Port / interface mapping & staged configuration", level=2)
        doc.add_paragraph(
            "The MOP carries the physical port/interface mapping and the staged (converted) "
            "configuration — the legacy side is read from the collected evidence; every <target> "
            "column must be completed by the implementing engineer before the window.")
        map_rows = []
        ifaces_all = _as_dict(snap.get("interfaces"))
        for h in switches:
            for pname, dd in sorted(_as_dict(ifaces_all.get(h)).items()):
                dd = _as_dict(dd)
                mode = str(dd.get("switchport_mode") or "").lower()   # str()-coerce wrong-typed device leaves (audit-6 #3 class)
                if mode == "access" and (str(dd.get("end_host_mac") or "").strip()
                                         or str(dd.get("end_host_ip") or "").strip()):
                    kind = "Endpoint"
                elif mode == "trunk" and str(dd.get("cdp_neighbor") or "").strip():
                    po = str(dd.get("port_channel") or "").strip()
                    # surface bundle membership so redundant uplink pairs (two ports sharing a Po) are explicit
                    kind = f"Uplink → {dd.get('cdp_neighbor')}" + (f" [{po}]" if po else "")
                else:
                    continue
                map_rows.append((h, pname, kind, dd.get("vlan") or "—",
                                 (dd.get("description") or dd.get("endpoint_type") or "—"),
                                 "<target device>", "<target port>"))
        if map_rows:
            table(["Legacy device", "Port", "Role", "VLAN", "Description / endpoint",
                   "Target device", "Target port"],
                  map_rows[:30], widths=[1.0, 0.8, 1.2, 0.5, 1.4, 1.0, 0.9])
            if len(map_rows) > 30:
                doc.add_paragraph(f"… and {len(map_rows) - 30} more port(s); the full per-port "
                                  "inventory is in the workbook's interface sheets.")
            stub_by_host: dict = {}
            for h, pname, kind, v, desc, *_ in map_rows:
                if kind == "Endpoint" and v != "—":
                    stub_by_host.setdefault(h, []).append((pname, v, desc))
            if stub_by_host:
                doc.add_paragraph(
                    "Staged endpoint-port configuration, derived from the evidence (mode/VLAN/"
                    "description only — review, complete and peer-review before any use):")
                # BOTH caps here are display caps and BOTH must disclose. §x.5 tells the engineer to
                # "apply the staged target configuration ... per the §x.4 port mapping", so a wave of
                # 40 stub-bearing devices that silently stages 2 of them (6 of up to 95 ports each)
                # reads as the wave's COMPLETE staged configuration — the same defect class as §x.2's
                # blocker list. The port-map table above discloses ITS cap; these two did not.
                _stub_hosts = sorted(stub_by_host)
                _shown_hosts = _stub_hosts[:2]
                for h in _shown_hosts:
                    _ports = stub_by_host[h]
                    lines = [f"! {h} — staged endpoint-port config (evidence-derived)"]
                    for pname, v, desc in _ports[:6]:
                        lines.append(f"interface <target port for {pname}>")
                        if desc != "—":
                            lines.append(f" description {desc}")
                        lines += [" switchport mode access", f" switchport access vlan {v}", "!"]
                    if len(_ports) > 6:
                        lines.append(
                            f"! …and {len(_ports) - 6} further endpoint port(s) on {h} NOT staged "
                            f"above ({len(_ports)} in total) — derive them the same way from the "
                            "port mapping above / the workbook's interface sheets.")
                    cp = doc.add_paragraph()
                    cr = cp.add_run("\n".join(lines))
                    cr.font.name = "Consolas"; cr.font.size = Pt(9)
                if len(_stub_hosts) > len(_shown_hosts):
                    doc.add_paragraph(
                        f"…and {len(_stub_hosts) - len(_shown_hosts)} further device(s) in this wave "
                        f"carry endpoint ports needing the same staged configuration "
                        f"({len(_stub_hosts)} in total); only the first {len(_shown_hosts)} are "
                        "stubbed above. These stubs are a WORKED EXAMPLE of the conversion, NOT this "
                        "wave's complete staged configuration — build the remainder the same way "
                        "from the port mapping above before the window.")
        else:
            doc.add_paragraph("No endpoint or uplink port evidence for this wave's devices in the "
                              "snapshot — build the mapping from a fresh interface collection before "
                              "the window.")

        # 2.x.5 procedure (each step carries its success criterion — the MOP definition; and a
        # PRE/DURING/POST phase tag so the window structure is visible at a glance — roadmap C2).
        doc.add_heading(f"{wi}.5 Cutover procedure", level=2)
        proc = [("[PRE] Open the change window and announce start in the war room; confirm rollback owner "
                 "is present.",
                 "roles confirmed; the rollback decision time in §" + f"{wi}.1 is agreed.")]
        # str()-coerced: the bare truthiness gate admits ANY truthy value, and the next call is a STRING
        # method -- a scalar leaf inside a well-formed playbook dict raised AttributeError here.
        if playbook.get("pre"):
            proc.append(("[PRE] Preparation: " + str(playbook["pre"]).rstrip(". ") + ".",
                         "staging complete with no production-facing change."))
        # classify off the engine's STRUCTURED dual-homed/single-homed split, NOT the free-text 'sequence'
        # (which for a mixed wave reads "...hard cutover (single-homed) + ...make-before-break (dual-homed)" —
        # the old substring match flipped the WHOLE wave to make-before-break, applying "keep the legacy leg
        # as the live fallback" steps to single-homed switches that have NO second path). Make-before-break
        # is the procedure ONLY when every member is dual-homed; any single-homed member -> hard-cutover flow.
        mbb_hosts, hard_hosts = (seq.get("make_before_break") or []), (seq.get("hard_cutover") or [])
        pure_mbb = bool(mbb_hosts) and not hard_hosts
        if pure_mbb:
            proc += [
                ("[DURING] Build the new/target path BESIDE the existing one (do not remove the legacy "
                 "path yet): stage the target device config and bring up the new uplinks per the §"
                 f"{wi}.4 mapping.",
                 "target links up; no STP topology change on the legacy path."),
                ("[DURING] Verify the new path forwards in isolation (link up, STP/trunk consistent, "
                 "gateway reachable) before moving any production load.",
                 "isolation checks pass on the target path before any endpoint moves."),
                ("[DURING] Migrate endpoints leg-by-leg / VLAN-by-VLAN onto the new path, validating "
                 "after each increment; keep the legacy leg as the live fallback.",
                 "each increment re-homes and passes its checks before the next one moves."),
                # The decommission step USED TO SIT HERE — i.e. the legacy path was torn down BEFORE the
                # §x.6 go/no-go ran two steps later, while §x.7's rollback told the engineer "the legacy
                # path was never torn down, this is non-disruptive". A §x.6 failure at 2am therefore met
                # a rollback whose stated precondition the procedure itself had just destroyed, and this
                # branch carries NO restore-from-§x.3 step. It is now emitted AFTER the §x.6 gate below.
            ]
        else:  # any single-homed member present (mixed or pure single-homed) -> scheduled-outage flow
            if mbb_hosts and hard_hosts:   # MIXED: do the dual-homed make-before-break FIRST, then the outage
                proc.append((
                    f"[DURING] Sequence within the wave: make-before-break the {len(mbb_hosts)} dual-homed "
                    f"switch(es) first (build beside, keep the legacy leg), THEN take the scheduled outage "
                    f"for the {len(hard_hosts)} single-homed switch(es) that have no second path.",
                    "dual-homed members re-homed with no outage before the single-homed window opens."))
            proc += [
                ("[DURING] Announce the hard cutover; this wave has a brief outage for the moved "
                 "endpoints.",
                 "stakeholders acknowledged; the outage clock is started."),
                ("[DURING] Apply the staged target configuration and move the uplinks/endpoints per the §"
                 f"{wi}.4 port mapping from legacy to target.",
                 "every port in the §" + f"{wi}.4 mapping is re-terminated and up."),
                ("[DURING] Bring up the target links and confirm STP converges and trunks negotiate as "
                 "expected.",
                 "STP converged; trunk allowed-VLAN checks pass; no err-disabled ports."),
            ]
        if playbook.get("validate"):
            proc.append(("[DURING] In-window validation: " + str(playbook["validate"]).rstrip(". ") + ".",
                         "in-window checks match the §" + f"{wi}.3 baseline."))
        proc.append(("[POST] Run §" + f"{wi}.6 post-cutover validation in full before declaring the wave "
                     "complete.",
                     "every check meets its stated acceptance criterion; reproducing a degraded baseline "
                     "or leaving an evidence gap unresolved is not acceptance."))
        if pure_mbb:
            # Ordered strictly AFTER the §x.6 gate: until §x.6 passes, the legacy path IS the rollback,
            # and §x.7's non-disruptive fail-back is only true while it is still standing.
            proc.append((
                "[POST] ONLY once §" + f"{wi}.6 has passed IN FULL and the wave is accepted: decommission "
                "the legacy path. Do NOT tear it down earlier — until §" + f"{wi}.6 passes, the legacy "
                "path IS the §" + f"{wi}.7 rollback, and removing it converts a non-disruptive fail-back "
                "into a disruptive restore from the §" + f"{wi}.3 configuration backup.",
                "§" + f"{wi}.6 passed in full BEFORE this step ran; no endpoint remains on the legacy leg "
                "(MAC/ARP tables clean)."))
        steps(proc)

        # 2.x.6 post-cutover validation (reuse the existing validation plan)
        doc.add_heading(f"{wi}.6 Post-cutover validation (go/no-go)", level=2)
        if val_items:
            doc.add_paragraph("Every acceptance condition must pass. A PRE-CUTOVER DEGRADED row is a blocker "
                              "to resolve or explicitly risk-accept before the window; matching that degraded "
                              "state after cutover is not success. A PRE-CUTOVER REVIEW row is also a blocker "
                              "and requires live simultaneous verification. Any BASELINE NOT VERIFIED row is "
                              "a blocker, not an acceptance target: re-collect the protocol evidence or explicitly "
                              "disposition the gap before the window. Evidence-bearing rows disclose evidence state, "
                              "embedded-projection custody, and source. A failure that cannot be corrected "
                              "in-window triggers the rollback in §" + f"{wi}.7.")
            # Every FHRP row is a VLAN-specific cutover gate and must remain executable in the MOP;
            # a shared summary command is only a capture optimization, not permission to omit a VLAN.
            # Keep those rows outside the bounded allowance for the other categories, then order the
            # combined set High-first.  The rationale `why` and any omitted non-FHRP rows remain in the
            # Cutover Validation workbook sheet.
            _vr = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3, "Info": 4}
            _ranked = sorted(val_items, key=lambda it: _vr.get(it.get("severity"), 5))
            _fhrp_rows = [it for it in _ranked if str(it.get("category") or "").upper() == "FHRP"]
            _other_rows = [it for it in _ranked if str(it.get("category") or "").upper() != "FHRP"]
            _evidence_blockers = [it for it in _other_rows if _validation_blocker(it)]
            _generic_rows = [it for it in _other_rows if not _validation_blocker(it)]
            _non_fhrp_cap = 40
            vshow = sorted(
                _fhrp_rows + _evidence_blockers + _generic_rows[:_non_fhrp_cap],
                key=lambda it: _vr.get(it.get("severity"), 5),
            )
            table(["Sev", "Category", "Device", "Check", "Command", "Observed baseline / acceptance"],
                  [(it.get("severity") or "—", it.get("category"), it.get("device"), it.get("check"),
                    it.get("command"), _validation_acceptance(it)) for it in vshow],
                  widths=[0.5, 0.9, 0.9, 1.6, 1.4, 1.5])
            _n_omitted = len(_generic_rows) - min(len(_generic_rows), _non_fhrp_cap)
            if _n_omitted:
                doc.add_paragraph(f"… and {_n_omitted} more ordinary non-FHRP check(s); full set in the "
                                  "Cutover Validation workbook sheet. Every FHRP VLAN gate and evidence-state "
                                  "blocker is retained above.")
        else:
            doc.add_paragraph("No machine-generated validation checks for this wave. As a minimum, verify "
                              "gateway reachability (ping the SVI/HSRP vIP), FHRP roles, routing "
                              "adjacencies, STP root, and port-channel membership are unchanged from the "
                              "§" + f"{wi}.3 baseline.")

        # rollback triggers — the EXPLICIT, QUANTIFIED abort conditions for this wave (Cisco-AS standard:
        # not "if something breaks" but a boolean with a measurable threshold). Each is a "Roll back if:"
        # test the in-window engineer can evaluate unambiguously. Where the snapshot carries no device-
        # specific threshold, the AS standard DEFAULT is stated and flagged "confirm with the change
        # owner" (coverage-honesty: a default is not a measured value). Unnumbered H2 so the .1-.8 chain
        # is intact; sits directly ahead of the §x.7 rollback PROCEDURE.
        doc.add_heading("Rollback triggers (quantified go/no-go)", level=2)
        doc.add_paragraph(
            "Execute the §" + f"{wi}.7 rollback if ANY trigger below is true. These are explicit boolean "
            "conditions, not prose judgement — the standard defaults are flagged “confirm with the change "
            "owner” because the assessment snapshot carries no device-specific threshold for them.")
        trig_rows = [
            ("Validation failure", "Roll back if: any Critical or High §" + f"{wi}.6 validation check "
             "fails, OR > 0.1% of the checks fail and cannot be corrected in-window.",
             "standard AS default — confirm with the change owner"),
            ("No convergence in time", "Roll back if: STP / routing / control-plane has not converged "
             "within the agreed budget (default 5 minutes) after the cutover step, OR endpoints are not "
             "reachable by the §" + f"{wi}.1 rollback decision time.",
             "5 min default — confirm with the change owner"),
            # The endpoint half of this trigger is sized on §x.1's max-blast-radius figure. When that
            # figure is [NOT OBSERVED] the clause must be withdrawn, not evaluated against a
            # fabricated 0 — "more than 0 endpoints affected" is true the moment the cutover starts,
            # so the in-window engineer would be reading a trigger that is unconditionally satisfied.
            ("Blast-radius / outage overrun",
             ("Roll back if: the observed outage exceeds the approved window, OR more endpoints than "
              "the §" + f"{wi}.1 max-blast-radius figure are affected."
              if blast is not None else
              "Roll back if: the observed outage exceeds the approved window. The §"
              + f"{wi}.1 max-blast-radius figure is [NOT OBSERVED] (no failure-impact evidence for "
              "this wave's devices), so NO endpoint threshold is derivable from this assessment — "
              "agree an explicit one with the change owner before the window; an absent figure is "
              "NOT a threshold of zero."),
             ("sized from §" + f"{wi}.1 — confirm with the change owner" if blast is not None else
              "[NOT OBSERVED] — the change owner must set this threshold")),
            ("Unrecoverable error", "Roll back if: any device err-disables, drops OOB reachability, or "
             "enters a state not covered by this MOP and not resolved within one escalation cycle.",
             "standard AS default — confirm with the change owner"),
        ]
        # EVPN-target specificity (gated ONLY on an NX-OS VXLAN-EVPN target): surface the primary-source
        # EVPN abort conditions so the trigger set is grounded in the fleet's own guardrail evidence, not
        # generic prose. Prefer the engine's rollback guardrail detail verbatim; fall back to the canonical
        # documented failure modes when a slim snapshot carries no rollback-phase guardrail.
        if evpn_on:
            evpn_rb = next((g for g in _as_list(evpn.get("guardrails"))
                            if isinstance(g, dict) and g.get("phase") == "rollback"), None)
            if evpn_rb and "if ANY of:" in str(evpn_rb.get("detail", "")):
                cond = str(evpn_rb["detail"]).split("if ANY of:", 1)[-1].strip().rstrip(".")
                basis = f"[{evpn_rb.get('source', 'BRKDCN-2951')}]"
            else:
                cond = ("a gateway-MAC mismatch black-holes a migrated subnet (host ARP not refreshed); "
                        "a Layer-2 loop forms (the overlay will not break it); the EVPN control plane "
                        "(underlay IGP or iBGP-EVPN) is not Established on a migrated leg; or post-cutover "
                        "reachability / NRFU fails")
                basis = "BRKDCN-2951; Cisco migration white papers"
            trig_rows.append((
                "EVPN cutover abort (target = NX-OS VXLAN-EVPN)",
                "Roll back the affected subnet/wave across the still-present L2+L3 interconnect if: "
                + cond + ".", basis))
        table(["Trigger", "Roll back if (quantified condition)", "Threshold basis"],
              trig_rows, widths=[1.5, 3.9, 1.3])

        # 2.x.7 rollback
        doc.add_heading(f"{wi}.7 Rollback", level=2)
        rb = [("Declare rollback in the war room and record the trigger (which validation check "
               "failed).",
               "the trigger and time are recorded in the change log.")]
        if multi:
            for g, sc, _sq in wave_groups:
                rbk = _as_dict(sc.get("playbook")).get("rollback")
                if rbk:                                        # same truthy-scalar leaf class as §x.5
                    rb.append((f"{g}: " + str(rbk).rstrip(". ") + ".",
                               "this group's endpoints are back on their legacy path."))
        elif playbook.get("rollback"):
            rb.append("Strategy-specific: " + str(playbook["rollback"]).rstrip(". ") + ".")
        if _kind == "coupled-subwave":
            rb.append(("Coupled sub-wave: rolling back this sub-wave must NOT strand a sibling sub-wave's "
                       "shared VLAN — keep the shared VLANs extended on the legacy path and coordinate with "
                       "any sibling sub-waves already cut over before withdrawing anything.",
                       "no sibling sub-wave loses a shared VLAN as a result of this rollback."))
        if pure_mbb:   # pure make-before-break: the legacy path stands until §x.6 passes
            rb.append(("Move any migrated endpoints back onto the still-live legacy leg; remove the "
                       "new path. The §" + f"{wi}.5 procedure decommissions the legacy path only AFTER "
                       "§" + f"{wi}.6 has passed in full, so on any §" + f"{wi}.6-triggered rollback the "
                       "legacy leg is still up and this is non-disruptive.",
                       "all endpoints re-home onto the legacy leg with no loss."))
            # …and the contingency the old text asserted out of existence: if the trigger fires after
            # the final §x.5 step, the fail-back above is NOT available and this branch must still name
            # a restore path — otherwise the pure-make-before-break rollback never references the
            # §x.3 configuration backup that §x.3 declares mandatory "for the §x.7 rollback".
            rb.append(("If the legacy path has ALREADY been decommissioned (the trigger fired after the "
                       "final §" + f"{wi}.5 step), the non-disruptive fail-back above is NOT available: "
                       "re-apply the pre-change configuration captured in §" + f"{wi}.3 and re-cable to "
                       "the legacy ports. This restore is DISRUPTIVE — declare the outage and start the "
                       "outage clock.",
                       "configuration and cabling match the §" + f"{wi}.3 captured state; the outage is "
                       "declared and timed."))
        else:
            if mbb_hosts and hard_hosts:   # MIXED: only the dual-homed legs roll back non-disruptively
                rb.append(("Dual-homed members roll back onto their still-live legacy leg (non-disruptive); "
                           "the single-homed members need the §" + f"{wi}.3 pre-change config re-applied (a "
                           "brief outage), since their path was torn down at cutover.",
                           "dual-homed members re-home with no loss; single-homed members restored from baseline."))
            rb.append(("Re-apply the pre-change configuration captured in §" + f"{wi}.3 to <devices>; "
                       "move uplinks/endpoints back to the legacy ports.",
                       "configuration and cabling match the §" + f"{wi}.3 captured state."))
        rb.append(("Re-run the §" + f"{wi}.6 checks against the legacy path to confirm service is "
                   "restored, then close the window as rolled-back and schedule a retro.",
                   "legacy-path checks match the §" + f"{wi}.3 baseline; window closed as rolled-back."))
        steps(rb)

        # 2.x.8 sign-off
        doc.add_heading(f"{wi}.8 Sign-off", level=2)
        table(["Role", "Name", "Time", "Result (proceed / rolled-back)"],
              [("Implementing engineer", "", "", ""),
               ("Validation / war-room lead", "", "", ""),
               ("Change owner", "", "", "")], widths=[2.2, 1.8, 1.2, 1.6])

    # ===== final acceptance =====
    doc.add_heading(f"{len(waves) + 3}. Post-Migration Acceptance", level=1)
    doc.add_paragraph(
        "The migration is complete when every wave above has been cut over and its validation passed, "
        "the consolidated punch-list has no open Critical/High item, and a final fleet re-assessment "
        "(re-run this toolkit and compare via the campaign-trend report) confirms the target-state "
        "health bands. Record the final sign-off and archive the per-wave evidence.")
    # closing signature gate under the existing section (heading=None: no new H1)
    # Deliverable Excellence (DE-01): shared glossary before the sign-off gate.
    add_glossary(doc)

    add_acceptance(
        doc, heading=None,
        scope_note="Sign-off below closes the change as a whole; each wave's in-window sign-off is "
                   "recorded in its own section.")

    # numbered sections only — the unnumbered furniture H1s (Contents, Document Control) are not
    # content sections and inflating the figure made run logs drift across versions (review fix)
    n_sections = len([p for p in doc.paragraphs
                      if p.style.name == "Heading 1" and p.text[:1].isdigit()])
    doc.save(output_path)
    logger.info(f"[Phase 34] MOP (DOCX) written: {output_path} ({len(waves)} wave(s), {n_sections} sections)")
