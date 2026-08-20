"""Shared document-control / acceptance "furniture" for the DOCX deliverable family.

NEW-V3.23.150: a Cisco Advanced Services deliverable (HLD/LLD, MOP, NRFU/ATP, runbook) always carries
the same furniture around the technical content — a document-control block after the cover (revision
history, intended audience, related documents, assumptions & caveats) and a signature acceptance block
at the end. That review-and-approval gate is how an AS engagement actually runs: no further work until
the deliverable is accepted. The engine writers (runbook / design / MOP) and AssessHub's web-layer
writers (cutover plan; the NRFU and PIR carry their own control tables and reuse only the
related-documents list) call these helpers so the wording, structure and cross-references stay
consistent across the whole set.

python-docx is OPTIONAL repo-wide: nothing here imports it at module load; the helpers import it
inside the call, exactly like the writers themselves, and operate on an already-open Document. The
writers only call in after their own `from docx import …` has succeeded, so the fail-soft path is
unchanged.
"""
from dataclasses import dataclass
from datetime import datetime
import json
import os
import zipfile
from typing import Iterable
from xml.etree import ElementTree

from cisco_toolkit.textutils import xml_safe   # shared XML-illegal-char sanitizer (every cell text routes through it)

# Severity ordering shared by the writer family (V3.23.159: hoisted from per-writer copies so the
# documents sort findings identically; engagement/crd import it — older writers' private copies
# migrate as they are next touched).
SEV_RANK = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3, "Info": 4}

_DOCX_MEDIA = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_PPTX_MEDIA = "application/vnd.openxmlformats-officedocument.presentationml.presentation"
_XLSX_MEDIA = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

PRE_CUTOVER = "pre-cutover"
POST_EXECUTION = "post-execution"
ENGINE_PRODUCER = "engine-cli"
ASSESSHUB_SNAPSHOT_PRODUCER = "assesshub-snapshot"
ASSESSHUB_EXECUTION_PRODUCER = "assesshub-execution"
SNAPSHOT_DOWNLOAD = "snapshot-download"
SNAPSHOT_EXPLORER = "snapshot-explorer"
EXECUTION_REPORT = "execution-report"


@dataclass(frozen=True)
class ArtifactSpec:
    """One canonical artifact contract shared by the engine, AssessHub and portable build.

    ``cli_suffix`` means the engine CLI writes the artifact and is the exact suffix used by the
    completeness verifier. ``assesshub_surface`` is independent: Cutover/NRFU are snapshot
    downloads but have no CLI suffix, while PIR is a conditional execution report and is neither a
    pre-cutover family member nor a CLI artifact.
    """

    key: str
    related_name: str
    label: str
    role: str
    ext: str
    media: str
    producer: str
    stage: str = PRE_CUTOVER
    cli_suffix: str | None = None
    assesshub_surface: str | None = None
    assesshub_suffix: str | None = None
    needs: str | None = None
    dependency_module: str | None = None
    writer_module: str | None = None
    writer_name: str | None = None
    conditional: bool = False

    @property
    def cli_produced(self) -> bool:
        return self.cli_suffix is not None

    @property
    def assesshub_download(self) -> bool:
        return self.assesshub_surface == SNAPSHOT_DOWNLOAD

    @property
    def web_only_pre_cutover(self) -> bool:
        return self.stage == PRE_CUTOVER and not self.cli_produced and self.assesshub_download

    @property
    def download_suffix(self) -> str:
        """The validated AssessHub filename suffix for a generating download/report surface."""
        if self.assesshub_suffix is None:
            raise ValueError(f"artifact {self.key!r} has no AssessHub download suffix")
        return self.assesshub_suffix


# The ONE artifact registry. Order is the recommended reading order of the pre-cutover set; the
# conditional PIR follows it so whole-registry callers see the lifecycle in order.
ARTIFACT_SPECS = (
    ArtifactSpec(
        "engagement", "Engagement Workflow & Plan of Record (.docx)",
        "Engagement Workflow & Plan of Record",
        "Phase-gated engagement workflow: verdict, gate calendar, RAID log and the next-action queue",
        "docx", _DOCX_MEDIA, ENGINE_PRODUCER, cli_suffix="_engagement.docx",
        assesshub_surface=SNAPSHOT_DOWNLOAD, assesshub_suffix="_engagement.docx",
        needs="python-docx", dependency_module="docx",
        writer_module="cisco_toolkit.engagement", writer_name="write_engagement_docx"),
    ArtifactSpec(
        "crd", "Customer Requirements Document (.docx)", "Customer Requirements Document (CRD)",
        "Plan-phase requirements capture: REQ-IDs, owners, and traceability into design and acceptance",
        "docx", _DOCX_MEDIA, ENGINE_PRODUCER, cli_suffix="_crd.docx",
        assesshub_surface=SNAPSHOT_DOWNLOAD, assesshub_suffix="_crd.docx",
        needs="python-docx", dependency_module="docx",
        writer_module="cisco_toolkit.crd", writer_name="write_crd_docx"),
    ArtifactSpec(
        "workbook", "Assessment workbook (.xlsx)", "Assessment workbook",
        "Full per-sheet evidence; every number in the narrative documents reconciles to it",
        "xlsx", _XLSX_MEDIA, ENGINE_PRODUCER, cli_suffix=".xlsx"),
    ArtifactSpec(
        "explorer", "Blast-Radius Explorer (.html)", "Blast-Radius Explorer",
        "Interactive topology, health and what-if view of the same snapshot",
        "html", "text/html", ENGINE_PRODUCER, cli_suffix="_explorer.html",
        assesshub_surface=SNAPSHOT_EXPLORER),
    ArtifactSpec(
        "runbook", "Assessment & Migration Runbook (.docx)", "Assessment & Migration Runbook",
        "Narrative findings, risk register and remediation detail",
        "docx", _DOCX_MEDIA, ENGINE_PRODUCER, cli_suffix="_runbook.docx",
        assesshub_surface=SNAPSHOT_DOWNLOAD, assesshub_suffix="_runbook.docx",
        needs="python-docx", dependency_module="docx",
        writer_module="cisco_toolkit.runbook", writer_name="write_runbook_docx"),
    ArtifactSpec(
        "design", "As-Built Network Design Document (.docx)", "As-Built Network Design Document",
        "HLD + LLD design record recovered from the fleet evidence",
        "docx", _DOCX_MEDIA, ENGINE_PRODUCER, cli_suffix="_design.docx",
        assesshub_surface=SNAPSHOT_DOWNLOAD, assesshub_suffix="_design.docx",
        needs="python-docx", dependency_module="docx",
        writer_module="cisco_toolkit.design", writer_name="write_design_doc_docx"),
    ArtifactSpec(
        "archreview", "Architecture Review & Conformance Report (.docx)",
        "Architecture Review & Conformance Report",
        "Leading-practice conformance scorecard, availability analysis and the design-review verdicts",
        "docx", _DOCX_MEDIA, ENGINE_PRODUCER, cli_suffix="_archreview.docx",
        assesshub_surface=SNAPSHOT_DOWNLOAD, assesshub_suffix="_archreview.docx",
        needs="python-docx", dependency_module="docx",
        writer_module="cisco_toolkit.archreview", writer_name="write_archreview_docx"),
    ArtifactSpec(
        "mop", "Per-Wave Method of Procedure (.docx)", "Per-Wave Method of Procedure",
        "Maintenance-window change procedure, validation and rollback per wave",
        "docx", _DOCX_MEDIA, ENGINE_PRODUCER, cli_suffix="_mop.docx",
        assesshub_surface=SNAPSHOT_DOWNLOAD, assesshub_suffix="_mop.docx",
        needs="python-docx", dependency_module="docx",
        writer_module="cisco_toolkit.mop", writer_name="write_mop_docx"),
    ArtifactSpec(
        "cutover", "Cutover Plan / Run-of-Show (.docx)", "Cutover Plan (Run-of-Show)",
        "Pilot-first wave sequence, go/no-go gates and window estimates",
        "docx", _DOCX_MEDIA, ASSESSHUB_SNAPSHOT_PRODUCER,
        assesshub_surface=SNAPSHOT_DOWNLOAD, assesshub_suffix="_cutover.docx",
        needs="python-docx", dependency_module="docx",
        writer_module="webapp.backend.cutover_docx", writer_name="write_cutover_docx"),
    ArtifactSpec(
        "nrfu", "NRFU / Acceptance Test Plan (.docx)",
        "Network Ready-For-Use (NRFU) Test Plan",
        "Production-acceptance test cases with expected baselines",
        "docx", _DOCX_MEDIA, ASSESSHUB_SNAPSHOT_PRODUCER,
        assesshub_surface=SNAPSHOT_DOWNLOAD, assesshub_suffix="_nrfu.docx",
        needs="python-docx", dependency_module="docx",
        writer_module="webapp.backend.nrfu_docx", writer_name="write_nrfu_docx"),
    ArtifactSpec(
        "opshandbook", "Operations Handbook (.docx)", "Operations Handbook",
        "Day-2 monitoring baseline, drift control, software governance and escalation readiness",
        "docx", _DOCX_MEDIA, ENGINE_PRODUCER, cli_suffix="_ops_handbook.docx",
        assesshub_surface=SNAPSHOT_DOWNLOAD, assesshub_suffix="_opshandbook.docx",
        needs="python-docx", dependency_module="docx",
        writer_module="cisco_toolkit.ops", writer_name="write_ops_handbook_docx"),
    ArtifactSpec(
        "deck", "Executive Presentation Deck (.pptx)", "Executive Presentation Deck",
        "Stakeholder summary of posture, risk and the migration plan",
        "pptx", _PPTX_MEDIA, ENGINE_PRODUCER, cli_suffix="_executive_deck.pptx",
        assesshub_surface=SNAPSHOT_DOWNLOAD, assesshub_suffix="_deck.pptx",
        needs="python-pptx", dependency_module="pptx",
        writer_module="cisco_toolkit.deck", writer_name="write_executive_deck_pptx"),
    ArtifactSpec(
        "pir", "Post-Implementation Review (.docx)", "Post-Implementation Review (PIR)",
        "Conditional as-executed record generated only after an AssessHub execution run exists",
        "docx", _DOCX_MEDIA, ASSESSHUB_EXECUTION_PRODUCER, stage=POST_EXECUTION,
        assesshub_surface=EXECUTION_REPORT, assesshub_suffix="_pir.docx",
        needs="python-docx", dependency_module="docx",
        writer_module="webapp.backend.pir_docx", writer_name="write_pir_docx", conditional=True),
)


def validate_artifact_registry(specs: Iterable[ArtifactSpec] = ARTIFACT_SPECS) -> None:
    """Fail closed on an empty, duplicate, contradictory or producer-orphaned registry.

    The explicit ``specs`` input lets tests prove each guard fails on a counterexample; validating
    only the live constant would permit a vacuous empty registry to pass every ``all`` check.
    """
    rows = tuple(specs)
    if not rows:
        raise ValueError("artifact registry is empty")
    keys = [spec.key for spec in rows]
    if len(keys) != len(set(keys)):
        raise ValueError("artifact registry contains duplicate keys")
    for spec in rows:
        if not all((spec.key, spec.related_name, spec.label, spec.role, spec.ext, spec.media,
                    spec.producer, spec.stage)):
            raise ValueError(f"artifact {spec.key!r} has an empty required field")
        if spec.stage not in (PRE_CUTOVER, POST_EXECUTION):
            raise ValueError(f"artifact {spec.key!r} has an unknown lifecycle stage")
        if spec.producer not in (ENGINE_PRODUCER, ASSESSHUB_SNAPSHOT_PRODUCER,
                                 ASSESSHUB_EXECUTION_PRODUCER):
            raise ValueError(f"artifact {spec.key!r} has an unknown producer")
        if spec.assesshub_surface not in (None, SNAPSHOT_DOWNLOAD, SNAPSHOT_EXPLORER,
                                          EXECUTION_REPORT):
            raise ValueError(f"artifact {spec.key!r} has an unknown AssessHub surface")
        if spec.stage == PRE_CUTOVER and not (spec.cli_produced or spec.assesshub_surface):
            raise ValueError(f"pre-cutover artifact {spec.key!r} has no producing channel")
        if spec.cli_suffix is not None:
            if spec.producer != ENGINE_PRODUCER:
                raise ValueError(f"artifact {spec.key!r} has a CLI suffix but is not engine-produced")
            if not spec.cli_suffix.endswith(f".{spec.ext}"):
                raise ValueError(f"artifact {spec.key!r} CLI suffix disagrees with its extension")
        if spec.assesshub_surface in (SNAPSHOT_DOWNLOAD, EXECUTION_REPORT):
            if not all((spec.assesshub_suffix, spec.needs, spec.dependency_module,
                        spec.writer_module, spec.writer_name)):
                raise ValueError(f"artifact {spec.key!r} has an orphaned AssessHub producer")
            assesshub_suffix = spec.assesshub_suffix
            if assesshub_suffix is None or not assesshub_suffix.endswith(f".{spec.ext}"):
                raise ValueError(f"artifact {spec.key!r} AssessHub suffix disagrees with its extension")
        elif spec.assesshub_suffix or spec.writer_module or spec.writer_name:
            raise ValueError(f"artifact {spec.key!r} declares a writer outside a generating surface")
        if bool(spec.needs) != bool(spec.dependency_module):
            raise ValueError(f"artifact {spec.key!r} has an incomplete optional-dependency contract")
        if spec.producer == ASSESSHUB_SNAPSHOT_PRODUCER and not spec.web_only_pre_cutover:
            raise ValueError(f"AssessHub snapshot artifact {spec.key!r} is not a web-only pre-cutover output")
        if spec.producer == ASSESSHUB_EXECUTION_PRODUCER and spec.stage != POST_EXECUTION:
            raise ValueError(f"AssessHub execution artifact {spec.key!r} is not post-execution")
        if spec.stage == POST_EXECUTION and (spec.assesshub_surface != EXECUTION_REPORT
                                             or not spec.conditional or spec.cli_produced):
            raise ValueError(f"post-execution artifact {spec.key!r} must be conditional and non-CLI")
        if spec.conditional and spec.stage != POST_EXECUTION:
            raise ValueError(f"pre-cutover artifact {spec.key!r} cannot be conditional")


validate_artifact_registry()

ARTIFACT_BY_KEY = {spec.key: spec for spec in ARTIFACT_SPECS}
PRE_CUTOVER_SPECS = tuple(spec for spec in ARTIFACT_SPECS if spec.stage == PRE_CUTOVER)
ASSESSHUB_DOWNLOAD_SPECS = tuple(spec for spec in ARTIFACT_SPECS if spec.assesshub_download)
CONDITIONAL_ARTIFACT_SPECS = tuple(spec for spec in ARTIFACT_SPECS if spec.conditional)

# Backward-compatible views. Existing writers and redaction tooling may keep importing these names,
# but they no longer own membership or filenames: every value is derived from ARTIFACT_SPECS.
FAMILY = tuple((spec.key, spec.related_name, spec.role) for spec in PRE_CUTOVER_SPECS)
CLI_ARTIFACT_SUFFIX = {
    spec.key: spec.cli_suffix for spec in PRE_CUTOVER_SPECS if spec.cli_suffix is not None
}
WEB_ONLY_KINDS = frozenset(spec.key for spec in PRE_CUTOVER_SPECS if spec.web_only_pre_cutover)


def artifact_spec(key: str) -> ArtifactSpec:
    """Return the canonical artifact contract for ``key`` (raises ``KeyError`` if unknown)."""
    return ARTIFACT_BY_KEY[key]


def artifact_family_metadata() -> dict:
    """Portable/UI-safe family denominators, always derived from the canonical registry."""
    return {
        "pre_cutover": len(PRE_CUTOVER_SPECS),
        "engine_cli": len(CLI_ARTIFACT_SUFFIX),
        "assesshub_only_pre_cutover": len(WEB_ONLY_KINDS),
        "conditional_post_execution": len(CONDITIONAL_ARTIFACT_SPECS),
    }


def artifact_dependency_modules() -> tuple[str, ...]:
    """Optional import modules needed to render the full portable artifact lifecycle."""
    return tuple(dict.fromkeys(spec.dependency_module for spec in ARTIFACT_SPECS
                               if spec.dependency_module))


def artifact_writer_modules() -> tuple[str, ...]:
    """Dynamically dispatched producer modules the frozen portable build must retain."""
    return tuple(dict.fromkeys(spec.writer_module for spec in ARTIFACT_SPECS if spec.writer_module))


#: Appended to a WEB_ONLY_KINDS entry's role in the Related-Documents table. The table lists the
#: whole FAMILY, so a CLI/Atlas delivery — whose complete set is CLI_ARTIFACT_SUFFIX, which excludes
#: these two — used to advertise a Cutover Plan and an NRFU/ATP that no engine writer produces and
#: that are not in the folder. Naming the producer keeps the cross-reference true in BOTH channels
#: (the web set really does contain them) instead of dropping rows in one.
WEB_ONLY_ROLE_NOTE = (" — rendered in AssessHub from the stored snapshot; not produced by the "
                      "offline engine run, so it may not be part of this delivery")


def cli_artifacts(stem):
    """``[(key, name, filename)]`` a complete engine CLI run writes for output stem ``stem``.

    ``stem`` is the --output path without its extension (``.../Assessment_redacted``); the
    filenames returned are basenames, in FAMILY reading order."""
    base = str(stem).replace("\\", "/").rsplit("/", 1)[-1]
    return [(key, name, base + CLI_ARTIFACT_SUFFIX[key])
            for key, name, _role in FAMILY if key in CLI_ARTIFACT_SUFFIX]


def artifact_candidate_paths(stem):
    """Return the finite, producer-owned output set for one engine output ``stem``.

    This is deliberately not a prefix glob.  A prefix glob turns an older run, an operator's
    similarly-named note, or another concurrent run into evidence for the current assessment.
    Callers use this exact set only to identify files which already existed and were *not* emitted
    by the current run; the manifest itself is built exclusively from the live producer registry.
    """
    base = os.path.abspath(str(stem))
    out_dir = os.path.dirname(base) or "."
    paths = [base + suffix for suffix in CLI_ARTIFACT_SUFFIX.values()]
    paths.extend([
        base + ".snapshot.json",
        base + ".phase_timings.json",
        base + ".run_manifest.json",
        os.path.join(out_dir, "topology.mmd"),
        os.path.join(out_dir, "topology.dot"),
    ])
    # Preserve order while protecting against a future suffix collision.
    return list(dict.fromkeys(os.path.normpath(p) for p in paths))


def artifact_kind(path):
    """Return the structural kind used to validate a generated artifact."""
    name = os.path.basename(str(path)).lower()
    if name.endswith(".phase_timings.json"):
        return "phase-timings"
    if name.endswith(".snapshot.json"):
        return "snapshot"
    if name.endswith(".run_manifest.json"):
        return "manifest"
    ext = os.path.splitext(name)[1]
    return {
        ".xlsx": "xlsx",
        ".docx": "docx",
        ".pptx": "pptx",
        ".html": "html",
        ".json": "json",
        ".mmd": "mermaid",
        ".dot": "graphviz",
    }.get(ext, "file")


_OFFICE_REQUIRED = {
    "xlsx": ("[Content_Types].xml", "xl/workbook.xml"),
    "docx": ("[Content_Types].xml", "word/document.xml"),
    "pptx": ("[Content_Types].xml", "ppt/presentation.xml"),
}


def validate_artifact(path, kind=None):
    """Structurally validate one generated file before it may enter custody.

    Returns ``(ok, reason)`` and never raises for a malformed/partial file.  A successful writer
    return is insufficient evidence of delivery: an interrupted ZIP write can leave a non-empty
    ``.docx``/``.xlsx``/``.pptx`` which exists but cannot be opened, and a truncated JSON/HTML
    writer can do the same.  The checks intentionally use only the standard library so the custody
    gate works in the base installation and frozen Atlas build.
    """
    p = os.path.abspath(str(path))
    k = str(kind or artifact_kind(p)).strip().lower()
    try:
        st = os.stat(p)
    except OSError as exc:
        return False, f"missing or unreadable: {exc}"
    if not os.path.isfile(p):
        return False, "not a regular file"
    if st.st_size <= 0:
        return False, "empty file"

    if k in _OFFICE_REQUIRED:
        try:
            with zipfile.ZipFile(p, "r") as zf:
                names = set(zf.namelist())
                missing = [n for n in _OFFICE_REQUIRED[k] if n not in names]
                if missing:
                    return False, "invalid Office package; missing " + ", ".join(missing)
                bad = zf.testzip()
                if bad:
                    return False, f"corrupt Office package member: {bad}"
                for member in _OFFICE_REQUIRED[k]:
                    ElementTree.fromstring(zf.read(member))
        except (OSError, zipfile.BadZipFile, RuntimeError, ValueError,
                ElementTree.ParseError) as exc:
            return False, f"invalid Office package: {exc}"
        return True, "valid Office Open XML package"

    if k in ("json", "snapshot", "phase-timings", "manifest"):
        try:
            with open(p, encoding="utf-8-sig") as fh:
                obj = json.load(fh)
        except (OSError, UnicodeError, ValueError) as exc:
            return False, f"invalid JSON: {exc}"
        if k in ("snapshot", "phase-timings", "manifest") and not isinstance(obj, dict):
            return False, f"{k} root is {type(obj).__name__}, not an object"
        if k == "phase-timings" and not isinstance(obj.get("phases"), list):
            return False, "phase-timings object has no phases list"
        if k == "manifest" and not isinstance(obj.get("chain"), list):
            return False, "manifest object has no chain list"
        return True, "valid JSON"

    if k == "html":
        try:
            with open(p, encoding="utf-8") as fh:
                text = fh.read()
        except (OSError, UnicodeError) as exc:
            return False, f"invalid HTML text: {exc}"
        low = text.lower()
        if "<html" not in low or "</html>" not in low:
            return False, "HTML document is missing its root/open or closing tag"
        return True, "complete HTML document"

    if k in ("mermaid", "graphviz"):
        try:
            with open(p, encoding="utf-8") as fh:
                text = fh.read().strip()
        except (OSError, UnicodeError) as exc:
            return False, f"invalid diagram text: {exc}"
        if k == "mermaid" and not (
                text.startswith("graph ") or text.startswith("flowchart ")):
            return False, "Mermaid document has no graph/flowchart declaration"
        if k == "graphviz" and not (
                (text.startswith("graph ") or text.startswith("digraph "))
                and text.endswith("}")):
            return False, "Graphviz document has no complete graph declaration"
        return True, f"valid {k} document"

    return True, "non-empty regular file"

# Caveats every generated deliverable shares; writers append doc-specific ones via extra_assumptions.
_ASSUMPTIONS = (
    "This document is generated from the offline assessment snapshot named on the cover; it reflects "
    "the network as collected at that time, not live state.",
    "The evidence base is read-only device command output (configuration plus CDP/LLDP neighbour "
    "data); nothing was changed on the network during collection.",
    "Every <placeholder> marker must be completed and the document peer-reviewed before it is "
    "treated as an approved record.",
    "Recommendations must be validated against the engagement's statement of work and the customer's "
    "change-control process before execution.",
)

DEFAULT_ROLES = ("Customer network owner", "Customer operations lead",
                 "Delivery engineer (author)", "Project / change manager")


def related_rows(exclude=()):
    """(name, role) rows for the Related-Documents table, excluding the document being written.
    `exclude` is a family key or an iterable of keys; a bare string is treated as one key (set("mop")
    would otherwise dissolve into letters and silently keep the document in its own related list).

    A WEB_ONLY_KINDS row carries WEB_ONLY_ROLE_NOTE: it is a real member of the family, but no engine
    writer produces it, so an offline delivery must not advertise it as if it were enclosed."""
    ex = {exclude} if isinstance(exclude, str) else set(exclude)
    return [(name, role + (WEB_ONLY_ROLE_NOTE if key in WEB_ONLY_KINDS else ""))
            for key, name, role in FAMILY if key not in ex]


def add_table(doc, headers, rows, widths=None, *, fixed=True):
    """House-style banded table (bold header). Public: the engine writers (runbook/design/MOP)
    delegate their local table() closures here so the family has ONE table builder instead of
    per-writer copies (V3.23.152 review cleanup).

    `fixed=True` (the furniture tables): autofit off + widths on both column and cells so renderers
    honour them exactly — these widths were authored to fit. `fixed=False` (the writers' content
    tables): widths are hints on cells only, preserving the writers' historical auto-layout — their
    width sets predate fixed layout and can exceed the portrait text column, so forcing fixed layout
    clips the last column (caught in V3.23.152 visual QA)."""
    from docx.shared import Inches

    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, hd in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = xml_safe(str(hd))            # sanitize: one XML-illegal char in any cell aborts the whole .docx
        for para in c.paragraphs:
            for run in para.runs:
                run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = "" if v is None else xml_safe(str(v))
    if widths:
        if fixed:
            t.autofit = False
            for i, w in enumerate(widths):
                t.columns[i].width = Inches(w)
                for row in t.rows:
                    row.cells[i].width = Inches(w)
        else:
            for i, w in enumerate(widths):
                for row in t.rows:
                    row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def _kv(doc, label, value):
    """Bold navy lead-in + plain value, matching the writers' _label_run/kv convention."""
    from docx.shared import RGBColor

    p = doc.add_paragraph()
    r = p.add_run(xml_safe(label))
    r.bold = True
    from cisco_toolkit.brand_tokens import DOC_NAVY_RGB
    r.font.color.rgb = RGBColor(*DOC_NAVY_RGB)
    p.add_run(" " + (xml_safe(str(value)) if value not in (None, "") else "—"))
    return p


def add_related_documents(doc, *, exclude=(), audience=None, intro=False):
    """Render the family cross-reference block (level-2 heading + Related-Documents table, with an
    optional intro paragraph and an optional trailing Intended-audience line). One implementation for
    the whole set: add_document_control uses it for the engine writers' front matter, and the NRFU /
    PIR writers (which carry their own control tables) call it directly (V3.23.152 review cleanup)."""
    doc.add_heading("Related documents", level=2)
    if intro:
        # NOT "the set is internally consistent (every number reconciles)". That was asserted
        # unconditionally, in the same front matter where add_excellence_front's single-source-of-truth
        # line may report the opposite (figures published but never reconciled, or reconciled with
        # violations) — a verification claimed by furniture that ran no verification. The reconcile
        # VERDICT has one owner (ssot.summary, rendered in 'At a Glance'); this block states only the
        # provenance fact it can vouch for and points at that owner (SSOT Law 1).
        doc.add_paragraph(
            "This document is one of a set generated from the same assessment snapshot, and should "
            "travel together with the engagement's statement of work and change records. Whether the "
            "set's headline figures were reconciled against the raw evidence is reported per document "
            "on the single-source-of-truth line under 'At a Glance' — this block does not assert it.")
    add_table(doc, ["Document", "Role in the set"], related_rows(exclude), widths=[2.7, 4.0])
    if audience:
        _kv(doc, "Intended audience:", audience)


def add_document_control(doc, *, document, label, engine_version="", generated_at=None,
                         collected_at=None, audience="", exclude=(), extra_assumptions=()):
    """Render the Document Control front matter (control table, revision history, intended audience,
    related documents, assumptions & caveats). Call it between the cover page and the table of
    contents; the heading is deliberately unnumbered so the writers' numbered sections are untouched.

    Provenance (wave R2-3-01): "Snapshot captured" is the COLLECTION instant (when the evidence was
    sampled), distinct from "Document generated" (when this DOCX was rendered). On a --no-collect
    re-analysis these differ; collected_at falls back to generated_at for snapshots predating the field."""
    doc.add_heading("Document Control", level=1)
    add_table(doc, ["Field", "Value"], [
        ("Document", document),
        ("Subject", label or "—"),
        ("Snapshot captured", collected_at or generated_at or "—"),
        ("Document generated", generated_at or "—"),
        ("Generated by", ("cisco-migration-assessment-toolkit " + str(engine_version)).strip()),
        ("Status", "DRAFT — generated; not yet reviewed"),
    ], widths=[1.8, 4.9])

    doc.add_heading("Revision history", level=2)
    add_table(doc, ["Version", "Date", "Author", "Notes"], [
        ("0.1", datetime.now().strftime("%Y-%m-%d"), "Generated draft",
         "Initial draft generated from the assessment snapshot"),
        ("", "", "", "<record each review/amendment here before acceptance>"),
    ], widths=[0.8, 1.1, 1.9, 2.9])

    # Distribution list (wave DE-01, Law 2 "Document Control"): a generated deliverable is a
    # controlled record — who holds a copy and in what capacity is part of the control front matter,
    # not an afterthought. Static roles (the review-and-approval chain); names/orgs are placeholders
    # the delivery lead completes before the document is issued.
    doc.add_heading("Distribution list", level=2)
    add_table(doc, ["Name", "Role", "Organisation", "Purpose"],
              [("<name>", r, "<organisation>",
                "Review & acceptance" if "author" not in r.lower() else "Author / issue")
               for r in DEFAULT_ROLES],
              widths=[1.6, 2.2, 1.7, 1.2])

    _kv(doc, "Intended audience:", audience)
    doc.add_paragraph()

    add_related_documents(doc, exclude=exclude, intro=True)

    doc.add_heading("Assumptions & caveats", level=2)
    for a in tuple(_ASSUMPTIONS) + tuple(extra_assumptions):
        doc.add_paragraph(a, style="List Bullet")


def add_acceptance(doc, *, heading="Document Acceptance", scope_note="", roles=DEFAULT_ROLES):
    """Render the closing acceptance / sign-off block: the review-and-approval gate. Pass
    heading=None to append the block under an existing section (the MOP's final acceptance)."""
    if heading:
        doc.add_heading(heading, level=1)
    doc.add_paragraph(
        "By signing below, the signatories confirm that this document has been reviewed, that review "
        "comments have been addressed, and that it is accepted as the working record for this phase "
        "of the engagement." + ((" " + scope_note) if scope_note else ""))
    add_table(doc, ["Role", "Name", "Signature", "Date"],
           [(r, "", "", "") for r in roles], widths=[2.2, 1.9, 1.6, 1.0])


def enable_update_fields(doc):
    """Mark the document so Word / LibreOffice rebuild every field (the TOC especially) on open —
    instead of shipping the client an empty placeholder they must Right-click → Update Field / F9.

    Sets ``<w:updateFields w:val="true"/>`` in word/settings.xml, idempotently (V3.23.172). Without
    it a field-code TOC renders as its placeholder text until manually refreshed, so a client who
    opens-and-prints (or a headless docx→pdf render) gets a deliverable with no table of contents.

    Placed in the schema-correct slot (``CT_Settings`` is an ordered sequence: ``updateFields`` must
    precede ``compat``/``rsids``/``mathPr``/…). Appending at the end is what Word/LibreOffice tolerate
    but a strict XSD validator rejects — insert before the first element that must follow it so the
    deliverable is schema-valid, not merely Word-openable.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    after = {qn("w:" + t) for t in (
        "compat", "docVars", "rsids", "attachedSchema", "themeFontLang", "clrSchemeMapping",
        "doNotAutoCompressPictures", "shapeDefaults", "decimalSymbol", "listSeparator",
        "hdrShapeDefaults", "footnotePr", "endnotePr")}
    after.add(qn("m:mathPr"))
    settings = doc.settings.element
    el = settings.find(qn("w:updateFields"))
    if el is None:
        el = OxmlElement("w:updateFields")
        anchor = next((c for c in settings if c.tag in after), None)
        anchor.addprevious(el) if anchor is not None else settings.append(el)
    el.set(qn("w:val"), "true")


def add_toc(doc, *, heading="Contents", depth="1-2"):
    """Render the 'Contents' heading + the Word TOC field code (V3.23.171 — was a verbatim
    copy in every generator; one home means one place to change heading depth or the
    update-field placeholder). Fields are marked to rebuild on open (``enable_update_fields``,
    V3.23.172), so the TOC populates for the client automatically — no manual F9."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc.add_heading(heading, level=1)
    p = doc.add_paragraph()
    run = p.add_run()
    fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = rf'TOC \o "{depth}" \h \z \u'
    fs = OxmlElement("w:fldChar"); fs.set(qn("w:fldCharType"), "separate")
    ft = OxmlElement("w:t"); ft.text = "This table builds automatically when the document opens (or Right-click → Update Field)."
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end")
    for el in (fb, instr, fs, ft, fe):
        run._r.append(el)
    enable_update_fields(doc)
    doc.add_page_break()


def as_dict(x) -> dict:
    """Coerce a possibly-malformed snapshot value to a dict. `snap.get(k) or {}` does NOT
    guard a truthy non-dict (the V3.23.159 review's recurring bug class) — every deliverable
    reader should use these shared coercers instead of carrying a private copy."""
    return x if isinstance(x, dict) else {}


def as_list(x) -> list:
    """Coerce a possibly-malformed snapshot value to a list (see as_dict)."""
    return x if isinstance(x, list) else []


# ---- Deliverable Excellence furniture (wave DE-01) --------------------------------------------
# The narrative documents were strong on evidence but weak on ORIENTATION: no "at a glance"
# register, no glossary, and — the recurring drift risk — they hand-recounted headline numbers
# instead of reading the single source of truth. These shared helpers close Laws 1/6/8/10 + the
# glossary for the WHOLE docx family in one place, exactly as add_document_control/add_acceptance
# already do for Law 2. Every value is read via cisco_toolkit.ssot.canonical_facts (never a recount)
# and is coverage-honest (an unpublished fact reads "[NOT OBSERVED]", never a fabricated 0).

# Curated shared glossary — terms + evidence conventions common to every deliverable. A writer
# appends doc-specific terms via extra_terms; the coverage-honesty rows are the load-bearing ones.
DEFAULT_GLOSSARY = (
    ("SVI", "Switched Virtual Interface — a VLAN's Layer-3 gateway on a switch."),
    ("FHRP", "First-Hop Redundancy Protocol (HSRP/VRRP/GLBP) — gateway redundancy across two devices."),
    ("LDoS / EoS", "Last Day of Support / End of Sale — hardware lifecycle milestones. Past-LDoS is migration-critical."),
    ("Blast radius", "The set of devices/endpoints affected if a given node or link fails."),
    ("Keystone device", "A device whose failure strands a disproportionate share of the fleet."),
    ("Move group", "A set of devices migrated together within one maintenance window."),
    ("vPC / MLAG", "Multi-chassis link aggregation — two switches presenting as one to a downstream device."),
    ("VRF", "Virtual Routing and Forwarding — an isolated routing table on a device."),
    ("Inferred-high", "Derived from CDP/LLDP or configuration rather than directly observed, but high-confidence."),
    ("[NOT OBSERVED]", "The evidence needed to judge this was NOT collected. It is NOT a healthy or passing "
                       "result — coverage-honesty: not observed never silently becomes 'fine'."),
)


def _glance_rows(snap):
    """Default 'at a glance' Q&A rows built from the canonical facts (coverage-honest). Kept private
    so a writer's extra_rows compose with a single, drift-proof headline block."""
    from cisco_toolkit import ssot   # lazy: avoids any module-load cycle (matches the docx lazy-import style)

    f = ssot.canonical_facts(snap if isinstance(snap, dict) else {})

    def v(x):
        return "[NOT OBSERVED]" if x is None else x

    nd, nc = f.get("n_devices"), f.get("n_collected")
    if nd is not None and nc is not None and nc != nd:
        scope = f"{nc} of {nd} inventoried devices collected (read-only evidence; {nd - nc} not reached)"
    elif nd is not None:
        scope = f"{nd} inventoried devices, collected read-only"
    else:
        scope = "[NOT OBSERVED]"

    rows = [
        ("What was assessed?", scope),
        ("How healthy is the fleet?",
         f"average health {v(f.get('avg_health'))}/100 — {v(f.get('n_critical'))} Critical, {v(f.get('n_poor'))} Poor"),
        # The unknown count rides WITH the headline, not in a footnote. `v()` maps only None to
        # [NOT OBSERVED], so a real 0 passed straight through and the first page of every
        # deliverable led with "0 device(s) past last-day-of-support" for a fleet whose platforms
        # the offline EoX KB never matched. Zero findings and zero coverage are different facts.
        ("Biggest lifecycle exposure?",
         f"{v(f.get('n_past_ldos'))} device(s) past last-day-of-support (migration-critical)"
         f" · {v(f.get('n_near'))} within one year of LDoS"
         f" · {v(f.get('n_past_eos'))} past end-of-sale with LDoS still future "
         "(date band only; entitlement not inferred)"
         + (f" — plus {f['n_unknown']} device(s) NOT ASSESSED: no authoritative lifecycle band "
            f"was assigned because either no exact EoX row matched or the matched row's source/date "
            f"authority was withheld; their support state is undetermined, not clear"
            if f.get("n_unknown") else
            " — lifecycle coverage count [NOT OBSERVED]: the summary did not publish n_unknown, so "
            "zero adverse findings cannot be read as complete assessment"
            if f.get("n_unknown") is None else "")),
        ("Scale (VLANs / endpoints)?",
         f"{v(f.get('n_vlans'))} production VLANs · {v(f.get('n_endpoints'))} evidenced endpoints"),
    ]
    if f.get("n_design_decisions") is not None:
        rows.append(("Design decisions recorded?",
                     f"{f['n_design_decisions']} (each traced to a requirement or a gap)"))
    return rows


def add_excellence_front(doc, snap, *, extra_rows=(), heading="At a Glance"):
    """Answer-first 'at a glance' register + the single-source-of-truth trust signal, for EVERY
    deliverable (Laws 6/1/10). Reads headline facts via ssot.canonical_facts (never a recount),
    renders the ssot.summary self-verification line, and is coverage-honest ('[NOT OBSERVED]', not 0,
    for an unpublished fact). extra_rows lets a writer pre-answer its own doc-specific questions.
    Unnumbered heading so the writers' numbered sections are untouched. Call it right after the TOC."""
    from cisco_toolkit import ssot

    doc.add_heading(heading, level=1)
    s = ssot.summary(snap if isinstance(snap, dict) else {})
    if s.get("n_facts"):
        if s.get("verified"):
            # `n_facts` and `n_checked` are NOT the same unit and must never share an "X of Y" frame:
            # n_facts counts the canonical facts PUBLISHED (14), n_checked counts the reconcile CHECKS
            # that RAN (20 on the Meridian reference fleet — several facts carry more than one independent basis, and
            # reconcile also verifies non-canonical siblings such as lifecycle_risk.summary.by_band).
            # Rendered as "{n_checked} of {n_facts}" this shipped "✓ 20 of 14 headline figures
            # self-verified" into all seven DOCX deliverables: a client-facing trust badge claiming more
            # figures verified than exist. eval_harness._BADGE_COUNT_RE reads the number immediately
            # before "headline figures self-verified", so that slot stays n_facts and the L1 rendered-
            # citation guard keeps reconciling; the check count is disclosed in its own unit beside it.
            _kv(doc, "Single source of truth:",
                f"✓ {s['n_facts']} headline figures self-verified against the raw evidence "
                f"({s['n_checked']} independent reconciliation check(s), none disagreed) — every "
                "number checked reconciles to one source.")
        elif not s.get("n_checked"):
            # PUBLISHED but never RECONCILED. Distinct from both other states and it must say so: a
            # snapshot can publish every canonical block while carrying none of the raw arrays the
            # checks derive from, and the old text claimed "self-verified ... every number in this
            # document reconciles to one source" over exactly that. Claiming a verification that did
            # not run is worse than reporting none.
            _kv(doc, "Single source of truth:",
                f"⚠ {s['n_facts']} headline figures published, but NONE could be reconciled against "
                "raw evidence in this snapshot — they are unverified, not confirmed.")
        else:
            # Same unit discipline as the verified branch: both numbers here are CHECKS (a violation is
            # emitted per failed check, and one fact can carry several), so the label says checks —
            # "reconciled headline figures" read them as facts and could imply more figures than exist.
            _kv(doc, "Single source of truth:",
                f"⚠ {s['n_violations']} of {s['n_checked']} reconciliation check(s) over the "
                f"{s['n_facts']} published headline figures did not agree; treat flagged numbers with "
                "caution (see the assessment-integrity disclosure).")
    rows = list(_glance_rows(snap)) + list(extra_rows)
    add_table(doc, ["Question", "Answer"], rows, widths=[2.6, 4.1], fixed=False)


def add_glossary(doc, *, extra_terms=(), heading="Glossary & Conventions"):
    """Shared term/convention glossary for the deliverable family (Technique T2). extra_terms
    appends doc-specific terms. Unnumbered; place before the acceptance block."""
    doc.add_heading(heading, level=1)
    doc.add_paragraph(
        "Terms and evidence conventions used across the assessment deliverable set. "
        "'[NOT OBSERVED]' always means the evidence was not collected — never that the item is healthy.")
    add_table(doc, ["Term", "Meaning"], list(DEFAULT_GLOSSARY) + list(extra_terms),
              widths=[1.7, 5.0], fixed=False)


def add_inputs_required(doc, snap, *, extra_rows=(), heading="Inputs Required"):
    """Consolidated 'what the customer and delivery team must still provide or confirm' register
    (Law 9): the reader should never have to hunt scattered <placeholders> to know what is owed.
    Universal, coverage-honest base rows from the snapshot — the not-collected devices are the #1
    real gap (a blind spot, [NOT OBSERVED], never assumed healthy) — plus the standard
    placeholder/peer-review gates; extra_rows lets a writer add its own unconfirmed items.
    Unnumbered; place in the front matter so the outstanding asks are answer-first."""
    snap = snap if isinstance(snap, dict) else {}
    cc = snap.get("collection_completeness")
    cc = cc if isinstance(cc, dict) else {}
    summ = cc.get("summary")
    summ = summ if isinstance(summ, dict) else {}
    rows = []
    n_notcoll = summ.get("not_collected")
    if isinstance(n_notcoll, int) and n_notcoll > 0:
        rows.append(("Collect the not-reached devices",
                     f"{n_notcoll} inventoried device(s) were not collected — their state is a blind spot "
                     "([NOT OBSERVED]). Collect them read-only to close it before relying on this document.",
                     "Customer / delivery"))
    rows.append(("Complete every <placeholder>",
                 "Fill every <…> marker in this document before it is treated as an approved record.",
                 "Delivery engineer"))
    rows.append(("Peer review & accept",
                 "This is a generated DRAFT; it must be peer-reviewed and accepted (see Document "
                 "Acceptance) before use.", "Review owner"))
    rows = rows + list(extra_rows)
    doc.add_heading(heading, level=1)
    doc.add_paragraph(
        "What the customer and delivery team must provide or confirm to close the assessment's known "
        "gaps. Coverage-honest: a not-collected device is a blind spot, never assumed healthy.")
    add_table(doc, ["Item", "What is needed", "Owner"], rows, widths=[1.9, 3.8, 1.0], fixed=False)
