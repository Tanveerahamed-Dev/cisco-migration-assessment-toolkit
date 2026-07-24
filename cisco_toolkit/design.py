"""As-Built Network Design Document — HLD + LLD (.docx).

NEW-V3.23.148: the Design-phase twin of the workbook / explorer / runbook / deck. Cisco Advanced
Services engagements produce a High-Level Design (architecture intent) and a Low-Level Design
(device-by-device build detail) in the PPDIOO Design phase. This module reconstructs BOTH from the
SAME snapshot the other deliverables read — i.e. an *as-built* design recovered from collected
configuration + CDP/LLDP evidence, not a greenfield design — and folds in the assessment's findings
as target-state design recommendations.

One source of truth: every number agrees with the workbook/explorer because it is the same
snap_dict. Evidence discipline is preserved (this is config-/neighbour-derived, not live telemetry):
gateway PLACEMENT is Confirmed where an SVI is configured; the active forwarding owner is Unknown;
inter-switch links are Inferred-high; endpoint presence is Confirmed at snapshot time only.

python-docx is an OPTIONAL dependency (like the runbook): imported inside the function so the package
imports fine without it, and a missing library is a warning + skip, never a crash. Every snapshot read
is defensive (.get / tolerant of missing keys) so an older snapshot degrades gracefully. Deterministic
content; no network, no device data beyond the snapshot.
"""
import logging
import re
from collections import Counter, defaultdict
from datetime import datetime

from cisco_toolkit.analyze import vlan_inventory
from cisco_toolkit.docmeta import add_acceptance, add_document_control, add_excellence_front, add_glossary, add_inputs_required, add_table, add_toc
# The family's shared falsy-guard coercers (archreview/deck import the same pair): `x or {}` / `x or []`
# substitutes only for FALSY values, so a TRUTHY non-dict/non-list survives and crashes the next deref.
from cisco_toolkit.docmeta import as_dict as _as_dict
from cisco_toolkit.docmeta import as_list as _as_list
from cisco_toolkit.textutils import _as_num, xml_safe, xml_safe_deep

logger = logging.getLogger(__name__)

_LC_BAND_RANK = {"Past-LDoS": 0, "Past-EoS": 1, "Near-LDoS": 2, "Active": 3, "Unknown": 4}


def _bom_rows(disposition: str, seq) -> list:
    """§5.1 replacement-BoM rows from a (model, qty) pair list. `bom.get(k, [])` substitutes only when
    the key is ABSENT, so a null / scalar value -- or a row that is not a 2-element pair -- raised on the
    `for m, q in ...` unpack and aborted the WHOLE design document. Coerce the container AND each row's
    arity; a malformed row is dropped (degrade to empty) rather than crashing the deliverable."""
    out = []
    for item in _as_list(seq):
        if isinstance(item, (list, tuple)) and len(item) == 2:
            out.append((disposition, item[0], item[1]))
    return out


def _is_l3(host: str, l3_forwarding, routing_neighbors) -> bool:
    """A device participates in L3 if it owns an SVI (appears in l3_forwarding) or runs a routing
    protocol (has any adjacency record). Used for the core/distribution vs access tiering — honest,
    evidence-based (a configured SVI / a routing neighbour is Confirmed), not a guess from hostname."""
    if any((r.get("switch") == host) for r in (l3_forwarding or []) if isinstance(r, dict)):
        return True
    rn = (routing_neighbors if isinstance(routing_neighbors, dict) else {}).get(host)
    rn = rn if isinstance(rn, dict) else {}   # the per-host value can itself be a truthy non-dict (same class)
    return any(rn.get(p) for p in ("ospf", "eigrp", "bgp"))


def _vlan_inventory(snap: dict):
    """Canonical VLAN inventory — delegates to analyze.vlan_inventory so the design doc and the CRD
    (and any other reader) share ONE derivation and cannot drift. Returns (vid:int, name:str) list."""
    return vlan_inventory(snap)


def _segmentation_facts(snap: dict):
    """Derive a segmentation posture directly from the interfaces (known shape), so the section is
    accurate regardless of the segmentation compute's internal layout: the set of non-default VRFs in
    use, and the count of SVIs carrying an ingress/egress ACL. Returns (vrfs:set, n_acl_svis:int,
    n_svis:int)."""
    vrfs: set = set()
    n_acl_svis = 0
    n_svis = 0
    _ifaces = snap.get("interfaces")
    for host, ports in (_ifaces if isinstance(_ifaces, dict) else {}).items():
        for p, d in (ports if isinstance(ports, dict) else {}).items():
            if not isinstance(d, dict):
                continue   # a truthy non-dict port-detail slips `or {}` and crashes d.get() (same class)
            vrf = str(d.get("vrf") or "").strip()   # str()-coerce a wrong-typed device leaf before .strip() (audit-6 #3 class)
            if vrf and vrf.lower() not in ("default", "global"):
                vrfs.add(vrf)
            if re.match(r"^Vlan\d+$", p, re.IGNORECASE) and (d.get("svi_ip") or ""):
                n_svis += 1
                if str(d.get("acl_in") or "").strip() or str(d.get("acl_out") or "").strip():
                    n_acl_svis += 1
    return vrfs, n_acl_svis, n_svis


def build_design_traceability(snap_dict: dict) -> list:
    """W3-7: the design TRACEABILITY MATRIX -- one row per RECOMMENDED design decision, tracing it to its CCDE
    principle (id + title), the published citation, and the collected evidence (summary + the snapshot field paths
    + devices). A pure renderer over data each decision already carries (principle + citation + evidence); NO new
    compute. COVERAGE-HONEST: a decision with no citation is shown '(uncited)', never given a fabricated reference.
    Total -- safe on None / a non-dict design_blueprint."""
    bp = (snap_dict or {}).get("design_blueprint") if isinstance(snap_dict, dict) else None
    bp = bp if isinstance(bp, dict) else {}
    rows = []
    for d in _as_list(bp.get("decisions")):   # `or []` passes a TRUTHY non-list straight into `for d in ...`
        if not isinstance(d, dict) or d.get("status") != "recommended":
            continue
        pr = d.get("principle") if isinstance(d.get("principle"), dict) else {}
        ev = d.get("evidence") if isinstance(d.get("evidence"), dict) else {}
        pid = str(pr.get("id") or "").strip()
        ptitle = str(pr.get("title") or "").strip()
        rows.append({
            "decision": str(d.get("title") or "").strip(),
            "priority": str(d.get("priority") or "").strip(),
            "principle": (f"{pid} — {ptitle}".strip(" —")) or "(no principle)",
            "citation": str(pr.get("citation") or "").strip() or "(uncited)",
            "evidence": str(ev.get("summary") or "").strip(),
            # _as_list: a truthy non-list 'fields' would be walked by `for f in ...`, and a truthy
            # non-list 'devices' would crash the [:8] slice -- both survive `or []` (audit-7 class).
            "fields": ", ".join(f for f in _as_list(ev.get("fields")) if isinstance(f, str)),
            "devices": ", ".join(str(x) for x in _as_list(ev.get("devices"))[:8]),
        })
    return rows


def write_design_doc_docx(output_path: str, snap_dict: dict, label: str) -> None:
    """Emit the As-Built Network Design Document (HLD + LLD) to `output_path`.

    Fail-soft applies to the OPTIONAL DEPENDENCY only: a missing python-docx is a warning + skip. There
    is deliberately NO blanket try/except around the render, so an unexpected error DOES propagate to the
    caller (the CLI reports it; the webapp /deliverable/design route maps it to an HTTP 500). Swallowing
    a render failure would ship a missing or truncated document while reporting success, which this
    repo's coverage-honesty doctrine forbids. (An earlier revision of this docstring claimed the error
    was "logged, never raised" -- it never was: the only `try` here is the import guard below.)

    Robustness therefore comes from GUARDING EVERY SNAPSHOT READ instead: `x or {}` / `x or []` catches
    only FALSY values, so every read is routed through _D/_R/_as_dict/_as_list and a malformed section
    degrades to EMPTY rather than raising."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor
    except ImportError:
        logger.warning("  Design document (DOCX) skipped: python-docx not installed "
                       "(pip install python-docx to enable the HLD/LLD design deliverable).")
        return

    # Deep-sanitize ALL device-derived text up front so a Unicode noncharacter (U+FFFE/U+FFFF) or lone surrogate
    # in any field cannot abort the whole DOCX at XML serialization ("All strings must be XML compatible"). Every
    # string the generator writes -- direct paragraphs (_label_run), tables, the title label -- flows from these.
    snap = xml_safe_deep(snap_dict if isinstance(snap_dict, dict) else {})
    label = xml_safe(label) if isinstance(label, str) else (str(label) if label is not None else "")
    from cisco_toolkit.brand_tokens import DOC_NAVY_RGB
    NAVY = RGBColor(*DOC_NAVY_RGB)
    GREY = RGBColor(0x59, 0x59, 0x59)
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    def _label_run(p, label_text, value, color=NAVY):
        r = p.add_run(label_text); r.bold = True; r.font.color.rgb = color
        p.add_run(" " + (str(value) if value not in (None, "") else "—"))

    def table(headers, rows, widths=None):
        # Delegates to the family's single table builder (docmeta.add_table) — V3.23.152 dedup.
        return add_table(doc, headers, rows, widths, fixed=False)

    # ---- snapshot-derived facts (reconciled to the workbook) ----
    def _D(x): return x if isinstance(x, dict) else {}    # audit-5 totality: a truthy non-dict section -> {}/[]
    def _R(x): return [r for r in (x if isinstance(x, list) else []) if isinstance(r, dict)]
    devices = _D(snap.get("devices"))
    l3f = _R(snap.get("l3_forwarding"))
    rn = _D(snap.get("routing_neighbors"))    # _D: a truthy non-dict -> {} (feeds _is_l3 AND the rn.items() render in §2.3)
    stp_roots = _D(snap.get("stp_roots"))    # _D: a truthy non-dict -> {} instead of crashing stp_roots.items() (§2.2)
    redist = _D(snap.get("redistribution"))    # _D: a truthy non-dict -> {} instead of crashing redist.values() (§2.3)
    fhrp = _R(snap.get("fhrp"))    # _R: a truthy non-list (malformed snapshot) -> [] instead of crashing `for g in fhrp`
    # FHRP candidate/configured counts come from the FULL gateway register (l3_forwarding), NOT snap['fhrp'] --
    # which is compute_fhrp_consistency()'s PROBLEMS-ONLY list (multi-gateway VLANs with an inconsistency).
    # Deriving the counts from it undercounts candidates and can NEVER credit a cleanly-redundant VLAN, so a
    # fully-redundant fabric reported 0 candidates / 0 with FHRP. snap['fhrp'] stays the source of the
    # 'lacking/inconsistent' row only. (Matches the CRD, which counts FHRP from l3_forwarding[].fhrp.)
    _vlan_gw_count: dict = {}
    for _r in l3f:
        _v = str(_r.get("vlan") or "")
        if _v:
            _vlan_gw_count[_v] = _vlan_gw_count.get(_v, 0) + 1
    n_multi_gw = sum(1 for _c in _vlan_gw_count.values() if _c >= 2)
    n_fhrp_cfg = len({str(_r.get("vlan") or "") for _r in l3f
                      if str(_r.get("fhrp") or "").strip().lower() not in ("", "none", "-", "—")})
    capacity = _R(snap.get("capacity"))
    lifecycle = _D(snap.get("lifecycle_risk"))    # _D: a truthy non-dict -> {} instead of crashing lifecycle.get('per_device')
    vpc = _D(snap.get("vpc"))    # _D: a truthy non-dict -> {} instead of crashing vpc.items() (§1 scale table + §2.4)
    failure_impact = _R(snap.get("failure_impact"))
    punchlist = _R(snap.get("punchlist"))    # _R: dict rows only, AND a truthy non-list -> [] (sorted()/.get() in §4 fallback)
    subnet_intel = _D(snap.get("subnet_intelligence"))    # _D: a truthy non-dict -> {} instead of crashing .get('per_device') (§3.2)
    svc = snap.get("service_map") or {}
    # isinstance-guard, not `or {}`: a TRUTHY non-dict (a string/list in a malformed or slimmed snapshot) slips
    # through `or {}` and then crashes .get('scale') -- the fail-soft siblings (engagement/ops) degrade via
    # _as_dict; design.py was the asymmetric gap (multi-domain audit L4).
    _eb = snap.get("executive_brief"); eb = _eb if isinstance(_eb, dict) else {}
    _bp = snap.get("design_blueprint"); bp = _bp if isinstance(_bp, dict) else {}   # canonical design blueprint

    lc_by_host = {}
    for r in _R(lifecycle.get("per_device")):   # _R: a truthy non-list -> [] AND dict rows only (r.get below)
        h = r.get("host") or r.get("hostname")
        if h:
            lc_by_host[h] = r
    cap_by_host = {r.get("hostname"): r for r in capacity}

    l3_hosts = [h for h in devices if _is_l3(h, l3f, rn)]
    l2_hosts = [h for h in devices if h not in l3_hosts]
    vlans = _vlan_inventory(snap)
    # SSOT scale: prefer the canonical executive_brief.scale (the published single source the
    # explorer, deck and webapp read) over a local recount, with len() only as a fallback — the HLD
    # was the last surface still recomputing scale, a latent drift risk even though they coincide
    # today (A5 SSOT fix). Lines that render the actual VLAN LIST keep len(vlans).
    _sc = eb.get("scale"); _scale = _sc if isinstance(_sc, dict) else {}
    n_dev = _scale.get("n_devices") or len(devices)
    n_vlan = _scale.get("n_vlans") or len(vlans)
    vrfs, n_acl_svis, n_svis = _segmentation_facts(snap)

    # ---- title page ----
    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("As-Built Network Design Document"); tr.bold = True
    tr.font.size = Pt(26); tr.font.color.rgb = NAVY
    sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s2 = sub2.add_run("High-Level Design (HLD) + Low-Level Design (LLD)")
    s2.font.size = Pt(14); s2.font.color.rgb = GREY
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(label); sr.font.size = Pt(13); sr.font.color.rgb = GREY
    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')}  ·  "
                 f"{n_dev} devices in scope  ·  script {snap.get('script_version', '')}"
                 ).font.color.rgb = GREY
    status = doc.add_paragraph(); status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st = status.add_run("DRAFT — as-built design recovered from collected evidence; review before reuse.")
    st.italic = True; st.font.color.rgb = GREY
    doc.add_paragraph()
    note = doc.add_paragraph()
    _label_run(note, "Design basis:",
               "This document reconstructs the current (as-built) design from collected configuration "
               "and CDP/LLDP neighbour data, not from live forwarding telemetry. It is intended as the "
               "design baseline for the migration: the HLD captures architecture intent recovered from "
               "the fleet, the LLD captures per-device build detail, and §4 lists the target-state "
               "design changes the assessment recommends. Gateway PLACEMENT is Confirmed where an SVI "
               "is configured; the active forwarding owner is Unknown. Inter-switch links are "
               "Inferred-high. Endpoint presence is Confirmed at snapshot time only.", GREY)
    doc.add_page_break()

    # ---- document control (AS-style front matter; unnumbered so §1–§4 are untouched) ----
    add_document_control(
        doc, document="As-Built Network Design Document (HLD + LLD)", label=label,
        engine_version=str(snap.get("script_version", "")), generated_at=snap.get("generated_at"),
        collected_at=snap.get("collected_at"),
        audience="Customer architecture and engineering owners, and the build engineers who derive "
                 "device configurations from the LLD detail.",
        exclude=("design",))
    doc.add_page_break()

    # ---- table of contents (shared field-code helper, V3.23.171) ----
    add_toc(doc)


    # Deliverable Excellence (DE-01): answer-first at-a-glance register + single-source-of-truth signal.
    add_excellence_front(doc, snap_dict)
    # Deliverable Excellence (DE-01): consolidated Inputs-Required register (Law 9).
    add_inputs_required(doc, snap_dict)
    # ===== 1. Executive design summary =====
    doc.add_heading("1. Executive Design Summary", level=1)
    posture = eb.get("posture_statement") or (
        f"{n_dev} devices, {n_vlan} VLANs and {n_svis} gateway SVIs across the assessed fabric.")
    doc.add_paragraph(posture)
    # _as_dict, not `or {}`: `bp.get("summary", {})` only substitutes when ABSENT (a None value would
    # raise) and `or {}` still lets a TRUTHY non-dict through into _sm.get(...) -- the nested class.
    _sm = _as_dict(bp.get("summary"))
    if _sm.get("headline"):
        _label_run(doc.add_paragraph(), "Design headline:", _sm["headline"])
        mp = doc.add_paragraph()
        mr = mp.add_run(
            "Design method: this design reasons top-down from requirements (the WHY) and manages the "
            "trade-offs between availability, convergence, scalability, modularity, security, simplicity, "
            "optimal routing, load-balancing, manageability and cost. §4 carries the evidence-grounded "
            "target-state decisions, their CCDE basis, the trade-off scorecard, and the open requirement "
            "questions the design still needs answered.")
        mr.italic = True; mr.font.color.rgb = GREY
    table(["Design attribute", "As-built value"], [
        ("Devices in scope", n_dev),
        ("L3 nodes (own an SVI or a routing adjacency)", len(l3_hosts)),
        ("L2-only access nodes", len(l2_hosts)),
        ("VLANs in use", n_vlan),
        ("Gateway SVIs", n_svis),
        ("Routing VRFs (non-default)", len(vrfs) if vrfs else "0 (single global table)"),
        ("Gateway SVIs with an ACL applied", f"{n_acl_svis} of {n_svis}"),
        ("Multi-gateway VLANs (FHRP candidates)", n_multi_gw),
        ("vPC / MLAG peerings", len([h for h, v in vpc.items() if v])),
    ], widths=[4.6, 2.2])

    # ===== 2. High-Level Design (HLD) =====
    doc.add_heading("2. High-Level Design (HLD)", level=1)

    doc.add_heading("2.1 Topology tiers", level=2)
    doc.add_paragraph(
        "Devices are tiered by their L3 participation recovered from the configuration: a node that owns "
        "a gateway SVI or forms a routing adjacency is treated as a core/distribution (L3) node; the "
        "remainder are L2 access nodes. This is an evidence-based tiering, not an assumption from naming.")
    table(["Tier", "Count", "Devices"], [
        ("Core / Distribution (L3)", len(l3_hosts), ", ".join(sorted(l3_hosts)[:30]) or "—"),
        ("Access (L2-only)", len(l2_hosts), ", ".join(sorted(l2_hosts)[:30]) or "—"),
    ], widths=[2.2, 0.8, 4.0])
    # keystone devices (concentrated dependency) from failure_impact
    keystones = sorted((r for r in failure_impact if isinstance(r, dict) and _as_num(r.get("stranded")) > 0),
                       key=lambda r: -_as_num(r.get("stranded")))[:5]
    if keystones:
        _label_run(doc.add_paragraph(), "Concentrated dependency:",
                   "the fabric leans on " + ", ".join(
                       f"{r.get('host')} ({r.get('stranded')} endpoints)" for r in keystones) +
                   " — protect and sequence these first.")

    doc.add_heading("2.2 Layer-2 domain", level=2)
    root_count: Counter = Counter()
    for host, vmap in stp_roots.items():
        for vid, info in _D(vmap).items():   # _D: a truthy non-dict per-host value (stp_roots={host:5}) -> {} (same class)
            if isinstance(info, dict) and info.get("is_root"):
                root_count[host] += 1
    doc.add_paragraph(
        # SSOT: the headline VLAN COUNT reads the canonical executive_brief.scale.n_vlans (n_vlan), not a
        # recomputed len(vlans), so it can't drift from the rest of the doc (multi-domain audit L3). The VLAN
        # LIST render below keeps len(vlans) -- it counts the rows actually shown.
        f"The fabric carries {n_vlan} VLANs. Spanning-tree root placement is recovered from "
        f"'show spanning-tree': " + (
            "; ".join(f"{h} roots {n} VLAN(s)" for h, n in root_count.most_common(6))
            if root_count else "no explicit root bridge was observed in the collected output") + ".")
    if vlans:
        table(["VLAN", "Name"], [(v, nm or "—") for v, nm in vlans[:40]], widths=[1.0, 4.5])
        if len(vlans) > 40:
            doc.add_paragraph(f"… and {len(vlans) - 40} more VLAN(s); full list in the workbook.")

    doc.add_heading("2.3 Layer-3 design", level=2)
    proto_use: Counter = Counter()
    for host, d in rn.items():
        for proto in ("ospf", "eigrp", "bgp"):
            if isinstance(d, dict) and d.get(proto):   # a truthy non-dict per-host value slips `or {}` (same class)
                proto_use[proto.upper()] += 1
    doc.add_paragraph(
        "Routing protocols recovered from neighbour state: " + (
            ", ".join(f"{p} on {n} device(s)" for p, n in proto_use.most_common())
            if proto_use else "no dynamic-routing adjacencies were observed (static / connected only)") + ".")
    n_redist = sum(len(_R(v)) for v in redist.values())   # _R: a truthy non-list per-host value (redist={host:5}) -> [] not len(scalar)
    if n_redist:
        _label_run(doc.add_paragraph(), "Redistribution boundaries:",
                   f"{n_redist} route-redistribution edge(s) across the fleet — each is a protocol "
                   "boundary the migration must preserve or consciously retire.")
    if l3f:
        rows = [(r.get("vlan"), r.get("svi_ip") or "—", r.get("switch"),
                 r.get("fhrp") or "—", r.get("risk") or "—") for r in l3f[:40]]
        table(["VLAN", "Gateway IP", "Owner", "FHRP", "Risk"], rows, widths=[0.7, 1.4, 1.4, 1.6, 1.4])
        if len(l3f) > 40:
            doc.add_paragraph(f"… and {len(l3f) - 40} more gateway(s); full register in the workbook.")

    doc.add_heading("2.4 Resilience & redundancy", level=2)
    n_single_gw = sum(1 for r in l3f if "single-gateway" in (r.get("risk") or ""))
    # 'lacking/inconsistent' is the only row sourced from snap['fhrp'] (the problems-only consistency list);
    # n_multi_gw (candidates) and n_fhrp_cfg (running FHRP) are derived from l3_forwarding above so a cleanly
    # FHRP-redundant fabric is credited instead of misrepresented as 0-candidate / 0-FHRP.
    n_fhrp_issue = sum(1 for g in fhrp if g.get("issues"))
    table(["Resilience attribute", "As-built value"], [
        ("Multi-gateway VLANs (FHRP candidates)", n_multi_gw),
        ("…with first-hop redundancy (FHRP) configured", n_fhrp_cfg),
        ("…lacking FHRP (missing / inconsistent)", n_fhrp_issue),
        ("Single-gateway VLANs (no FHRP peer)", n_single_gw),
        ("vPC / MLAG peerings", len([h for h, v in vpc.items() if v])),
        ("Keystone devices (strand endpoints if lost)", len([r for r in failure_impact
                                                             if isinstance(r, dict) and _as_num(r.get("stranded")) > 0])),
    ], widths=[4.6, 2.2])

    doc.add_heading("2.5 Multicast & timing design", level=2)
    # _as_dict/_as_list at EVERY level: a truthy non-dict service_map.multicast crashes mc.get(...), and a
    # truthy non-list classified_groups/igmp_queriers or non-dict ptp crashes len()/ptp.values() below.
    mc = _as_dict(svc.get("multicast")) if isinstance(svc, dict) else {}
    groups = _as_list(mc.get("classified_groups"))
    queriers = _as_list(mc.get("igmp_queriers"))
    ptp = _as_dict(mc.get("ptp"))
    ptp_active = sum(1 for v in ptp.values() if isinstance(v, dict) and v.get("operational"))
    # Render the section only when there is actual activity to report; a multicast dict that exists
    # but carries all-zero counts (a non-media fabric, or commands not collected) gets the fallback
    # instead of an all-zeros paragraph that reads as filler in a client deliverable.
    if mc.get("active_switch_count") or mc.get("active_interfaces") or groups or queriers or ptp:
        doc.add_paragraph(
            f"Active multicast was observed on {mc.get('active_switch_count', 0)} switch(es) across "
            f"{mc.get('active_interfaces', 0)} interface(s); {len(groups)} group(s) were classified and "
            f"{len(queriers)} IGMP-snooping querier(s) were found. PTP (IEEE 1588): {ptp_active} of "
            f"{len(ptp)} device(s) report an operational clock — the rest are not boundary-clocked. "
            "On a broadcast fabric these are the production media, Dante/AES67 and ST-2110/PTP flows; "
            "the migration design must preserve querier coverage and the PTP clock hierarchy.")
        if groups:
            table(["Group", "Name", "Category"],
                  [(g.get("group"), g.get("name") or "—", g.get("category") or "—")
                   for g in (_as_dict(x) for x in groups[:15])],   # per-element: a scalar row -> {} not g.get() crash
                  widths=[1.6, 2.2, 1.8])
    else:
        doc.add_paragraph("No multicast/PTP activity was classified in the collected output "
                          "(either not a media fabric, or the relevant show commands were not captured).")

    doc.add_heading("2.6 Segmentation & security posture", level=2)
    doc.add_paragraph(
        ("The fabric uses a single global routing table (no non-default VRFs were observed); "
         if not vrfs else
         f"The fabric uses {len(vrfs)} non-default VRF(s): {', '.join(sorted(vrfs)[:10])}. ") +
        f"{n_acl_svis} of {n_svis} gateway SVIs carry an ingress/egress ACL. "
        "Where neither a dedicated VRF nor a gateway ACL is present, inter-VLAN traffic is openly "
        "routed — the migration is an opportunity to introduce intended segmentation.")

    # NEW-V3.23.165: QoS is a named HLD design domain; render the audited CONFIGURED posture
    # when the snapshot carries the qos_audit axis (older snapshots simply skip the section).
    qa = _D(snap_dict.get("qos_audit"))    # _D: a truthy non-dict -> {} instead of crashing qa.get('summary') (§2.7)
    qsum = _as_dict(qa.get("summary"))    # _as_dict: a truthy non-dict summary slips `or {}` -> qsum.get() crash
    if qsum.get("n_devices"):
        doc.add_heading("2.7 Quality of service (configured posture)", level=2)
        if not qsum.get("n_assessable"):
            doc.add_paragraph(
                "QoS posture is not assessable — no full running-config captures were available. "
                "The target design must state the QoS intent (trust boundary, marking, queuing) "
                "explicitly rather than inherit an unknown current state.")
        else:
            qmodes = _as_dict(qsum.get("modes"))    # _as_dict: a truthy non-dict crashes qmodes.items()
            mtxt = ", ".join(f"{v} device(s) {k}" for k, v in qmodes.items())
            doc.add_paragraph(
                f"Configured posture across the {qsum.get('n_assessable', 0)} assessable device(s): "
                f"{mtxt}. {qsum.get('n_voice_ports', 0)} voice-VLAN port(s) observed. The campus "
                "leading practice is a trust boundary at the access edge (trust detected phones/APs, "
                "mark at ingress) with consistent per-hop queuing fleet-wide; the table below lists "
                "where the observed configuration departs from that intent.")
            # _as_list: a truthy non-list 'findings' crashes the [:12] slice; _as_dict per element: a
            # scalar finding row crashes f.get(...).
            qrows = [(f.get("host"), f.get("label"), f.get("severity"))
                     for f in (_as_dict(x) for x in _as_list(qa.get("findings"))[:12])]
            if qrows:
                table(["Device", "Departure from leading practice", "Severity"],
                      qrows, widths=[1.6, 4.0, 1.0])
            else:
                doc.add_paragraph("No departures: the configured posture is consistent on every "
                                  "assessable device.")

    # ===== 3. Low-Level Design (LLD) =====
    doc.add_heading("3. Low-Level Design (LLD)", level=1)

    doc.add_heading("3.1 Device inventory & roles", level=2)
    inv_rows = []
    for h in sorted(devices):
        d = _D(devices[h])    # _D: a truthy non-dict device value -> {} instead of crashing d.get(...)
        lc = lc_by_host.get(h, {})
        role = "Core/Dist (L3)" if h in l3_hosts else "Access (L2)"
        cap = cap_by_host.get(h, {})
        _ap = d.get("active_ports", cap.get("active_ports", "—"))
        _ap = "—" if _ap is None else _ap          # None = active count not observed -> dash, never a false 0
        ports = f"{_ap}/{d.get('total_ports', cap.get('total_ports', '—'))}"
        inv_rows.append((h, d.get("model") or "—", d.get("serial_number") or d.get("chassis_serial") or "—",
                         d.get("sw_version") or "—", role, lc.get("band") or "—", ports))
    table(["Hostname", "Model", "Serial", "SW version", "Role", "EoL band", "Active/Total ports"],
          inv_rows, widths=[1.3, 1.3, 1.2, 0.9, 1.1, 0.9, 0.9])

    doc.add_heading("3.2 Addressing & VLAN plan", level=2)
    # vlan -> subnet from subnet_intelligence served_subnets; vlan -> gateway/fhrp from l3_forwarding
    subnet_by_vlan: dict = {}
    for pd in _R(subnet_intel.get("per_device")):    # _R: dict rows only, AND a truthy non-list -> [] (pd.get() below)
        for s in _R(pd.get("served_subnets")):    # _R: same class one level down (s.get() below)
            v = str(s.get("vlan") or "")
            if v and s.get("subnet"):
                subnet_by_vlan.setdefault(v, s.get("subnet"))
    addr_rows = []
    for r in l3f:
        v = str(r.get("vlan") or "")
        nm = next((nm for vv, nm in vlans if str(vv) == v), "")
        addr_rows.append((v or "—", nm or "—", subnet_by_vlan.get(v, "—"),
                          r.get("svi_ip") or "—", r.get("switch") or "—", r.get("fhrp") or "—"))
    if addr_rows:
        table(["VLAN", "Name", "Subnet", "Gateway IP", "Gateway device", "FHRP"],
              addr_rows[:50], widths=[0.7, 1.4, 1.3, 1.2, 1.3, 1.2])
        if len(addr_rows) > 50:
            doc.add_paragraph(f"… and {len(addr_rows) - 50} more; full plan in the workbook.")
    else:
        doc.add_paragraph("No gateway SVIs were observed — addressing is delegated off-fabric (pure L2).")

    doc.add_heading("3.3 Per-device build detail", level=2)
    doc.add_paragraph(
        "Per device: the gateway SVIs it owns (with FHRP role), the trunk uplinks to neighbours, and the "
        "port-channels it terminates — the build facts a re-implementation must reproduce.")
    shown = 0
    for h in sorted(l3_hosts) + sorted(l2_hosts):
        ports = _D(snap.get("interfaces")).get(h) or {}
        svis, uplinks, pos = [], [], []
        for p, d in _D(ports).items():
            if not isinstance(d, dict):
                continue   # a truthy non-dict port-detail slips `or {}` and crashes d.get() (same class)
            if re.match(r"^Vlan\d+$", p, re.IGNORECASE) and (d.get("svi_ip") or ""):
                svis.append(f"{p} {d.get('svi_ip')}" + (f" [{d.get('hsrp_behavior')}]"
                                                        if d.get("hsrp_behavior") else ""))
            if str(d.get("switchport_mode") or "").lower() == "trunk" and d.get("cdp_neighbor"):
                uplinks.append(f"{p}→{d.get('cdp_neighbor')}")
            if p.startswith("Po"):
                pos.append(p)
        if not (svis or uplinks or pos):
            continue
        doc.add_heading(h, level=3)
        if svis:
            _label_run(doc.add_paragraph(), "Gateway SVIs:", "; ".join(sorted(svis)[:12]))
        if uplinks:
            _label_run(doc.add_paragraph(), "Trunk uplinks:", ", ".join(sorted(uplinks)[:14]))
        if pos:
            _label_run(doc.add_paragraph(), "Port-channels:", ", ".join(sorted(set(pos))))
        shown += 1
        if shown >= 40:
            doc.add_paragraph(f"… per-device detail truncated at {shown} devices; full per-port data "
                              "is in the workbook.")
            break

    doc.add_heading("3.4 Equipment list (Bill of Materials)", level=2)
    by_model = defaultdict(lambda: {"count": 0, "band": "Unknown"})
    for h, d in devices.items():
        m = _D(d).get("model") or "Unknown"    # _D: a truthy non-dict device value -> {} (same class as devices[h])
        by_model[m]["count"] += 1
        b = lc_by_host.get(h, {}).get("band")
        if b and _LC_BAND_RANK.get(b, 9) < _LC_BAND_RANK.get(by_model[m]["band"], 9):
            by_model[m]["band"] = b
    bom_rows = sorted(((m, v["count"], v["band"]) for m, v in by_model.items()),
                      key=lambda r: (-r[1], r[0]))
    table(["Model", "Qty", "Worst EoL band"], bom_rows, widths=[3.0, 1.0, 2.0])

    doc.add_heading("3.5 Software plan & recommendations", level=2)
    doc.add_paragraph(
        "Per AS low-level-design convention, software versions are recorded as the MINIMUM for the "
        "solution, and images per platform should vary as little as possible — consistent software "
        "eases troubleshooting and root-cause analysis. The table reads the observed (as-built) image "
        "per model; where a model runs mixed versions, the most widely deployed image is the natural "
        "standardization candidate.")
    sw_by_model: dict = defaultdict(Counter)
    for h, d in devices.items():
        m = _D(d).get("model") or "Unknown"    # _D: a truthy non-dict device value -> {} (same class as devices[h])
        sw_by_model[m][str(_D(d).get("sw_version") or "").strip() or "—"] += 1   # str(): a dict sw_version would be an unhashable key
    sw_rows, mixed = [], []
    for m, cnt in sorted(sw_by_model.items(), key=lambda kv: (-sum(kv[1].values()), kv[0])):
        images = ", ".join(f"{v} ×{n}" for v, n in cnt.most_common())
        known = [v for v, _ in cnt.most_common() if v != "—"]
        if len(known) > 1:
            status = f"MIXED — standardize on {known[0]}"
            mixed.append((m, known[0]))
        elif known:
            status = "Consistent"
        else:
            status = "No version data"
        sw_rows.append((m, images, status, by_model[m]["band"]))
    table(["Model", "Images observed", "Status", "Worst EoL band"],
          sw_rows[:30], widths=[1.7, 2.2, 1.8, 1.0])
    recs = []
    for m, cand in mixed[:8]:
        recs.append(f"{m}: consolidate to one image. The most widely deployed ({cand}) is the "
                    "standardization candidate; validate it against Cisco's published recommended "
                    "release for the platform before adoption.")
    eol_models = sorted(m for m, v in by_model.items() if v["band"] in ("Past-EoS", "Past-LDoS"))
    if eol_models:
        recs.append("Hardware past end-of-support gets replacement, not an image upgrade: "
                    + ", ".join(eol_models[:8])
                    + ". Plan these as new-platform builds in the target design (§4).")
    if recs:
        for r_ in recs:
            doc.add_paragraph(r_, style="List Bullet")
    else:
        doc.add_paragraph(
            "Software is consistent per platform and no past-end-of-support hardware was observed — "
            "carry the current images forward as the minimum-version baseline for the target design.")

    # ===== 4. Target-state design recommendations =====
    doc.add_heading("4. Target-State Design Recommendations", level=1)
    # isinstance-filter (like crd.py:355 / deck.py:376): a single non-dict element in a corrupt/hand-edited
    # design_blueprint.decisions otherwise raised AttributeError on d.get(...) and aborted the WHOLE As-Built
    # Design Document, despite the 'never raised' contract -- design.py was the asymmetric gap vs the deck/CRD.
    decisions = [d for d in _as_list(bp.get("decisions")) if isinstance(d, dict)]
    if decisions:
        doc.add_paragraph(
            "The senior-design view: each target-state decision below is gated on collected evidence and "
            "traced to a network-design principle (CCDE doctrine) plus the trade-off axes it serves. This is "
            "design INTENT — the prescriptive changes the rebuilt fabric should adopt — distinct from the "
            "as-built punch-list (full per-device evidence remains in the runbook and the workbook).")

        doc.add_heading("4.1 Design trade-off scorecard", level=2)
        doc.add_paragraph(
            "Where the current fabric stands on each axis a senior designer balances (0 = weak, 4 = strong). "
            "The target design raises the weak axes without overspending the others.")
        sc = _as_list(bp.get("tradeoff_scorecard"))    # _as_list: a truthy non-list would be walked by `for s in sc`
        table(["Trade-off axis", "Score", "Posture", "Current-state evidence"],
              [(s.get("label", ""), f"{s.get('score')}/4", s.get("posture", ""), s.get("evidence", ""))
               for s in (_as_dict(x) for x in sc)],    # per-element: a scalar axis row -> {} not s.get() crash
              widths=[1.7, 0.7, 1.1, 3.3])

        rec = [d for d in decisions if d.get("status") == "recommended"]
        doc.add_heading("4.2 Recommended target-state design decisions", level=2)
        if rec:
            # _as_dict on evidence/principle: a truthy non-dict per-decision value (a corrupt or
            # hand-edited blueprint) slips `or {}` and crashes .get() -- here and in the detail loop below.
            table(["Priority", "Domain", "Design decision", "Why (driver)", "Evidence", "CCDE basis"],
                  [(d.get("priority"), d.get("domain"), d.get("title"), d.get("driver"),
                    _as_dict(d.get("evidence")).get("summary", ""),
                    _as_dict(d.get("principle")).get("citation", ""))
                   for d in rec], widths=[0.75, 0.85, 1.55, 1.6, 1.7, 1.3])
            for d in rec[:10]:
                _label_run(doc.add_paragraph(), f"{d.get('title')}",
                           f"[{d.get('priority')} · {d.get('domain')}]")
                _label_run(doc.add_paragraph(), "Recommended pattern:", d.get("recommended_action"))
                _label_run(doc.add_paragraph(), "Alternatives weighed:", d.get("alternatives"))
                _label_run(doc.add_paragraph(), "Trade-offs:", d.get("tradeoffs"))
                _label_run(doc.add_paragraph(), "Evidence:", _as_dict(d.get("evidence")).get("summary"))
                _label_run(doc.add_paragraph(), "CCDE principle:",
                           _as_dict(d.get("principle")).get("citation"), GREY)
        else:
            doc.add_paragraph("No evidence-grounded target-state changes — the as-built design carries no "
                              "flagged gaps to redesign.")

        needs = [d for d in decisions if d.get("status") == "needs-requirement"]
        if needs:
            doc.add_heading("4.3 Open design questions (requirements to confirm)", level=2)
            doc.add_paragraph(
                "Design top-down from the WHY: these decisions depend on requirements the assessment cannot "
                "observe. Confirm them and the target design right-sizes accordingly — the engine never "
                "assumes an answer.")
            # _as_list + str(): a truthy non-list crashes ", ".join(), and a list carrying a non-str
            # element crashes it too ("sequence item 0: expected str, int found"). str() is identity on
            # every well-formed string, so the rendered text is unchanged.
            table(["Open design question", "Requirement needed", "Trade-off axes"],
                  [(d.get("title"),
                    ", ".join(str(x) for x in _as_list(d.get("requirements_needed"))) or "—",
                    ", ".join(str(x) for x in _as_list(d.get("axes"))) or "—")
                   for d in needs], widths=[3.0, 2.0, 1.8])

        cov = _as_dict(bp.get("coverage"))    # _as_dict: a truthy non-dict slips `or {}` -> cov.get() crash
        if cov.get("caveat"):
            _label_run(doc.add_paragraph(), "Coverage:", cov.get("caveat"), GREY)

        doctrine = _as_dict(bp.get("doctrine"))    # _as_dict: a truthy non-dict crashes doctrine.values()
        if doctrine:
            doc.add_heading("4.4 Design doctrine applied (CCDE-grounded reference)", level=2)
            # Normalise ONCE at every level: a truthy non-list per-domain value crashed len(v)/`for it in v`,
            # and a scalar principle element crashed it.get(...) -- the same class two levels down.
            _doctrine = {dom: [_as_dict(it) for it in _as_list(v)] for dom, v in doctrine.items()}
            n_tot = sum(len(v) for v in _doctrine.values())
            n_act = sum(1 for v in _doctrine.values() for it in v if it.get("engine_actionable"))
            doc.add_paragraph(
                f"The full body of design doctrine this engine reasons from — {n_tot} principles across "
                f"{len(_doctrine)} domains, each an original re-expression of leading-practice design with a "
                f"source citation. {n_act} are evidence-actionable (the engine auto-detects their trigger and, "
                f"when the evidence is present, surfaces them as decisions in §4.2/4.3); the remainder are "
                f"reference doctrine for design areas the L1–L4 assessment does not collect (firewall, BGP, "
                f"MPLS-TE, IPv6, …) — they inform the design narrative and are never asserted from absent evidence.")
            for dom in sorted(_doctrine):
                items = _doctrine[dom]
                doc.add_heading(f"{dom} ({len(items)})", level=3)
                table(["Design principle", "Recommended pattern", "Source"],
                      [(it.get("title"), it.get("recommended_action"), it.get("citation")) for it in items],
                      widths=[1.9, 3.0, 1.9])

        trace = build_design_traceability(snap_dict)        # W3-7: design traceability matrix (audit trail)
        if trace:
            doc.add_heading("4.5 Design traceability matrix", level=2)
            doc.add_paragraph(
                "Audit trail — every recommended decision above, traced to the CCDE principle and published source "
                "that justify it, plus the collected evidence (and the snapshot field paths) it stands on. A "
                "decision with no citation is shown '(uncited)', never given a manufactured reference.")
            table(["Design decision", "CCDE principle", "Citation", "Evidence", "Source fields"],
                  [(r["decision"], r["principle"], r["citation"], r["evidence"], r["fields"] or "—") for r in trace],
                  widths=[1.5, 1.6, 1.2, 1.9, 1.6])
    else:
        doc.add_paragraph(
            "The assessment's consolidated punch-list, read as design intent: the changes the target design "
            "should adopt so the rebuilt fabric does not inherit the current gaps. Severity-ranked; full "
            "evidence and per-device scope are in the runbook and the Migration Punch-List workbook sheet.")
        sev_rank = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3, "Info": 4}
        top = sorted(punchlist, key=lambda i: sev_rank.get(i.get("severity"), 5))[:12]
        if top:
            pl_rows = []
            for i in top:
                # _as_list ONCE per row (mirrors deck.py): a truthy non-list 'devices' slips `or []` and
                # then crashes the [:4] slice / len() -- and the old code re-read i['devices'] a third time.
                _devs = _as_list(i.get("devices"))
                pl_rows.append((i.get("severity"), i.get("category") or "—", i.get("title") or "—",
                                ", ".join(str(x) for x in _devs[:4])
                                + ("" if len(_devs) <= 4 else f" +{len(_devs) - 4}")))
            table(["Severity", "Category", "Design recommendation", "Devices"],
                  pl_rows, widths=[0.9, 1.3, 3.0, 1.6])
        else:
            doc.add_paragraph("No punch-list items — the as-built design carries no flagged gaps to redesign.")

    # ---- §5: generated candidate target-state architecture (synthesised from evidence + requirements) ----
    # Every §5 read is coerced at its own level: `or {}` / `or []` guards only the ABSENT/None case, so a
    # truthy non-dict target_state (or non-list dimensions) survived and crashed the next dereference.
    ts = _as_dict(bp.get("target_state"))
    ts_dims = _as_list(ts.get("dimensions"))
    if ts_dims:
        doc.add_heading("5. Proposed Target-State Architecture (candidate)", level=1)
        _ts_sum = _as_dict(ts.get("summary"))    # local, so the bare ts["summary"]["headline"] subscript goes away
        if _ts_sum.get("headline"):
            _label_run(doc.add_paragraph(), "Synthesis:", _ts_sum["headline"])
        doc.add_paragraph(
            "A candidate target architecture synthesised from the collected evidence, the requirements "
            "register and the design doctrine — a proposal to validate, not a final design. Each dimension "
            "traces current state to a recommended target with its rationale; scale-driven choices are "
            "requirement-gated (resolve them and the target firms up), and uncollected devices remain an "
            "explicit unknown.")
        table(["Design area", "Current (observed)", "Recommended target", "Rationale", "Confidence"],
              [(d.get("area"), d.get("current"), d.get("target"), d.get("rationale"), d.get("confidence"))
               for d in (_as_dict(x) for x in ts_dims)],   # per-element: a scalar dimension -> {} not d.get() crash
              widths=[1.2, 1.55, 1.95, 1.7, 0.9])

        bom = _as_dict(ts.get("replacement_bom"))
        if bom.get("n_replace") or bom.get("n_refresh"):
            doc.add_heading("5.1 Replacement Bill of Materials (target procurement)", level=2)
            doc.add_paragraph(
                f"{bom.get('n_replace', 0)} asset(s) at end-of-support to replace and "
                f"{bom.get('n_refresh', 0)} approaching it to refresh, grouped by current model — the "
                "procurement the target build requires. " + (bom.get("note") or ""))
            rows = (_bom_rows("Replace (past-LDoS)", bom.get("replace_now"))
                    + _bom_rows("Refresh (near-LDoS / past-EoS)", bom.get("refresh_soon")))
            if rows:
                table(["Disposition", "Current model", "Qty"], rows, widths=[1.9, 3.0, 0.9])

        segp = _as_dict(ts.get("segmentation_plan"))
        if segp.get("observed"):
            doc.add_heading("5.2 Target segmentation", level=2)
            _label_run(doc.add_paragraph(), "Observed:", segp.get("observed"))
            _label_run(doc.add_paragraph(), "Target:", segp.get("target"))
            if segp.get("requirement_needed"):
                _label_run(doc.add_paragraph(), "Requirement needed:", segp.get("requirement_needed"), GREY)

        ap = _as_dict(ts.get("addressing_plan"))
        if ap.get("status") == "candidate" and ap.get("subnets"):
            doc.add_heading("5.3 Net-new IP addressing plan (candidate)", level=2)
            doc.add_paragraph(
                f"Candidate per-VLAN allocation from {ap.get('supernet')} ({ap.get('n_allocated')} subnet(s)). "
                + (ap.get("note") or ""))
            if ap.get("mode") == "zone-aware" and ap.get("zones"):
                _label_run(doc.add_paragraph(), "Per-zone summarization:",
                           "each zone is one contiguous prefix toward the core.")
                # BARE SUBSCRIPT, not an `or []` guard: the `and ap.get("zones")` gate above only proves the
                # value is TRUTHY -- `zones: 5` reached `for z in 5`. Same for subnets and its [:60] slice.
                table(["Security zone", "Summary prefix", "VLANs"],
                      [(z.get("zone"), z.get("summary"), z.get("n_vlans"))
                       for z in (_as_dict(x) for x in _as_list(ap.get("zones")))],
                      widths=[2.2, 2.0, 0.9])
            _subnets = _as_list(ap.get("subnets"))    # read ONCE: the slice and both len() reads share it
            table(["VLAN", "Hosts (obs.)", "Target subnet", "Note"],
                  [(s.get("vlan"), s.get("hosts"), s.get("subnet"), s.get("note", ""))
                   for s in (_as_dict(x) for x in _subnets[:60])],
                  widths=[0.8, 1.3, 1.7, 2.7])
            if len(_subnets) > 60:
                doc.add_paragraph(f"… and {len(_subnets) - 60} more; full plan in the workbook.")
        elif ap.get("status") == "candidate":
            # candidate plan that allocated NO subnets (supernet too small / fully overflowed) — render the
            # section + its note so the "enlarge the address_space" disclosure is never silently dropped.
            doc.add_heading("5.3 Net-new IP addressing plan", level=2)
            doc.add_paragraph(ap.get("note")
                              or "The supplied address_space could not accommodate the candidate plan; enlarge it.")
        elif ap.get("requirement_needed"):
            doc.add_heading("5.3 Net-new IP addressing plan", level=2)
            _label_run(doc.add_paragraph(), "Requirement needed:", ap.get("requirement_needed"), GREY)
            if ap.get("n_census_vlans") is not None:
                nq = ap.get("n_unsizable") or 0
                _label_run(
                    doc.add_paragraph(), "VLAN census:",
                    f"{ap.get('n_census_vlans')} VLAN(s) total — "
                    f"{ap.get('observed_vlans')} sizeable (observed access port or L3 SVI)"
                    + (f", {nq} querier-only/VLAN-1 with no auto-sized subnet" if nq else "")
                    + ". The full census is sized once an address_space is supplied.")
            doc.add_paragraph(ap.get("note") or "")

        wp = _as_dict(ts.get("wave_plan"))
        if wp.get("waves"):
            doc.add_heading("5.4 Candidate migration wave plan", level=2)
            doc.add_paragraph(
                f"{wp.get('n_move_groups')} L2-coupling move-group(s) (largest {wp.get('largest_group')} "
                f"switches) → {wp.get('n_waves')} candidate wave(s) of ≤ {wp.get('wave_cap')}. "
                + (wp.get("note") or ""))
            # Another bare subscript behind a truthy-only gate: `waves: 5` reached wp["waves"][:40].
            table(["Wave", "Kind", "Switches"],
                  [(w.get("wave"), w.get("kind"), w.get("n_switches"))
                   for w in (_as_dict(x) for x in _as_list(wp.get("waves"))[:40])],
                  widths=[0.8, 2.4, 1.0])

        if ts.get("scope_note"):
            _label_run(doc.add_paragraph(), "Scope:", ts["scope_note"], GREY)
        _ts_cov = _as_dict(ts.get("coverage"))    # local, so the bare ts["coverage"]["caveat"] subscript goes away
        if _ts_cov.get("caveat"):
            _label_run(doc.add_paragraph(), "Coverage:", _ts_cov["caveat"], GREY)

    # ===== 6. Interoperability & platform footprint =====
    # What the target design must keep interoperating with — the observed NOS-family spread and the
    # multi-vendor endpoint estate (OUI-derived). HONEST: endpoint_identity is a per-MAC vendor/class
    # observation, NOT a service-dependency / HA-cluster graph (that over-claim was dropped earlier).
    nos_mix = {}
    for dd in devices.values():
        fam = str(_D(dd).get("platform") or "").strip().lower() or "unknown"    # _D: same devices-value class as §3.1/§3.4/§3.5
        nos_mix[fam] = nos_mix.get(fam, 0) + 1
    ei = _R(snap.get("endpoint_identity"))    # _R: dict rows only, AND a truthy non-list -> [] (not iterated & crashed)
    if len(nos_mix) > 1 or ei:
        doc.add_heading("6. Interoperability & Platform Footprint", level=1)
        doc.add_paragraph(
            "What the target design must keep interoperating with. A migration changes the network, not the "
            "estate attached to it — the rebuilt fabric inherits this NOS-family spread and endpoint population, "
            "so feature, automation and edge-compatibility parity must hold across all of it.")
        if len(nos_mix) > 1:
            doc.add_heading("6.1 NOS-family footprint", level=2)
            doc.add_paragraph(
                f"The fleet spans {len(nos_mix)} switch NOS families — design intent, automation and the MOP "
                "must cover every family, and any cross-family migration (e.g. IOS → NX-OS) is an explicit "
                "interoperability step, not a like-for-like swap.")
            table(["NOS family", "Switches"],
                  sorted(((k.upper(), v) for k, v in nos_mix.items()), key=lambda r: -r[1]),
                  widths=[2.4, 1.0])
        if ei:
            vend, cls = {}, {}
            for e in ei:
                vk = (e.get("vendor") or "Unknown").strip() or "Unknown"
                ck = (e.get("endpoint_class") or "Unknown").strip() or "Unknown"
                vend[vk] = vend.get(vk, 0) + 1
                cls[ck] = cls.get(ck, 0) + 1
            n_known_cls = len([c for c in cls if c != "Unknown"])
            doc.add_heading("6.2 Endpoint interoperability surface", level=2)
            doc.add_paragraph(
                f"{len(ei)} evidenced endpoints span {len(vend)} distinct hardware vendors (OUI-derived) across "
                f"{n_known_cls} identified device class(es). The access-edge design (port profiles, PoE, "
                "multicast/AV, security posture) must accommodate this heterogeneity, not only the switching "
                "tier. These are per-MAC vendor/class observations, not a service-dependency graph.")
            table(["Top endpoint vendor (OUI)", "Endpoints"],
                  sorted(vend.items(), key=lambda r: -r[1])[:8], widths=[3.2, 1.0])
            top_c = [(k, v) for k, v in sorted(cls.items(), key=lambda r: -r[1]) if k != "Unknown"][:8]
            if top_c:
                table(["Device class", "Endpoints"], top_c, widths=[3.2, 1.0])

    # ---- closing acceptance gate (AS-style back matter) ----
    # Deliverable Excellence (DE-01): shared glossary before the sign-off gate.
    add_glossary(doc)

    add_acceptance(
        doc, scope_note="Acceptance confirms the as-built record as the design baseline for the "
                        "migration; the §4 target-state items proceed to detailed design under "
                        "their own approvals.")

    n_sections = len([p for p in doc.paragraphs if p.style.name == "Heading 1"])
    doc.save(output_path)
    logger.info(f"[Phase 33] Design document (DOCX) written: {output_path} ({n_sections} sections)")
