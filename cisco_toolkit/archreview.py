"""Architecture Review & Conformance Report — the senior-engineer design-review deliverable (.docx).

NEW-V3.23.160: the document chain (CRD → HLD/LLD → MOP → cutover → NRFU → PIR) records WHAT the
engagement produces and the engagement plan records HOW it runs; nothing rendered the judgment a
senior Cisco Advanced Services engineer applies to the ARCHITECTURE ITSELF. Cisco's Plan & Design
Review services "review infrastructure device configurations based on Cisco leading-practice
recommendations" and perform an "Availability Analysis ... to identify potential infrastructure
design and configuration issues that could affect network resiliency and availability". This module
is that review, automated: `compute_architecture_review` walks the snapshot through ~24 deterministic
leading-practice checks across 8 design domains (hierarchy, resiliency, L2 design, L3/gateway design,
capacity & oversubscription, operational readiness, security & segmentation, lifecycle & software),
renders a per-check verdict (conforms / advisory / deviation / critical / not-assessable), rolls the
verdicts up into a domain scorecard + an overall conformance grade, and `write_archreview_docx` emits
the review as an AS-style report with the shared document furniture.

Evidence discipline (same doctrine as the rest of the toolkit):
  * Every check synthesizes ALREADY-COLLECTED snapshot sections; no new collection, no telemetry.
  * A check whose evidence was not captured returns NOT-ASSESSABLE — absence of evidence is never
    scored as either conforming or deviating, and the report says so explicitly (§2 scope).
  * Config-derived state is treated as placement, not live behaviour (e.g. the FHRP/STP alignment
    check is capped at *advisory* because the active forwarding owner is Unknown from configs).
  * Verdict thresholds cite their source rule in each check's `reference` (the 20:1 / 4:1
    oversubscription rules of thumb, distribution-layer root placement, VLAN-1 hygiene, CIS
    management-plane hardening, EoX support posture).

One source of truth: the CLI attaches the computed review to `snap_dict["architecture_review"]`; the
writer prefers that attached section and recomputes it only when reading an older snapshot that
predates the key (deterministic — same snapshot, same review). python-docx is OPTIONAL exactly like
the other writers: imported inside the function, missing library is a warning + skip, never a crash.
Every snapshot read is defensive (`_as_dict` / `_as_list` / tolerant of malformed rows) so arbitrary
stored JSON degrades instead of crashing (the V3.23.159 review doctrine).
"""
import logging
import re
from collections import defaultdict
from datetime import datetime

from cisco_toolkit.docmeta import add_acceptance, add_document_control, add_excellence_front, add_glossary, add_inputs_required, add_table, add_toc
from cisco_toolkit.docmeta import as_dict as _docmeta_as_dict
from cisco_toolkit.docmeta import as_list as _docmeta_as_list
from cisco_toolkit.textutils import (   # entry deep-sanitize of device text (audit-5) + fail-soft numeric
    NATIVE1_CFG_BASIS, _as_num, bpduguard_state, is_trunk_mode, xml_safe, xml_safe_deep)   # coercion + shared token owners
from cisco_toolkit import ssot as _ssot_mod   # Law 1 accessors (canonical facts + segmentation posture)

logger = logging.getLogger(__name__)

# Verdict severity order (worst first) and the conformance weight each verdict contributes.
V_RANK = {"critical": 0, "deviation": 1, "advisory": 2, "conforms": 3, "not-assessable": 4}
_V_WEIGHT = {"conforms": 1.0, "advisory": 0.7, "deviation": 0.35, "critical": 0.0}

# Grade bands over the weighted conformance score (percent of assessable checks).
_GRADES = ((90, "A", "Conforms to leading practice"),
           (80, "B", "Minor deviations"),
           (70, "C", "Material deviations"),
           (55, "D", "Significant deviations"),
           (0,  "F", "Critical rework required"))

# Review items the offline snapshot CANNOT evidence — listed verbatim in §2 so the report is honest
# about its reach (a senior review states its scope; it does not silently skip).
NOT_ASSESSED = (
    ("QoS policy effectiveness", "Queuing/drop behaviour needs live counters, not configuration."),
    ("Wireless RF design", "No WLC/AP data is part of the wired collection."),
    ("WAN / SD-WAN policy", "Edge routing policy and transport SLAs are outside the campus scan."),
    ("Control-plane scale (TCAM / FIB / MAC table pressure)", "Needs platform show-tech counters."),
    ("Performance baselining (throughput, latency, loss)", "Needs telemetry over time, not a snapshot."),
    ("Physical plant (cabling, optics budget, cooling)", "Site-survey evidence is not collected."),
)

# Interface-name prefix → nominal Gbps, used only when the parsed speed column is empty/auto.
# Longest-prefix-first where prefixes share a stem ("twe"=TwentyFiveGigE before "tw"=TwoGigabitEthernet;
# "fif"=FiftyGigE before "fi"=FiveGigE). NX-OS "Ethernet" carries no speed in its name → None (the
# port is excluded from ratio math rather than guessed).
_NAME_GBPS = (("hu", 100.0), ("fif", 50.0), ("fo", 40.0), ("twe", 25.0), ("te", 10.0),
              ("fi", 5.0), ("tw", 2.5), ("gi", 1.0), ("fa", 0.1))

_VIRTUAL_PORT = re.compile(r"^(vlan|po\d|port-channel|lo\d|loopback|tu\d|tunnel|mgmt|null)", re.I)


# V3.23.171: the truthy-non-dict coercers are shared family furniture now (docmeta.as_dict /
# as_list) -- local aliases keep the call sites unchanged.
_as_dict = _docmeta_as_dict
_as_list = _docmeta_as_list


def _as_int(v, default: int = 0) -> int:
    # Delegate to the shared fail-soft coercer: rejects the JSON Infinity / NaN too (inf -> OverflowError,
    # which the old `except (TypeError, ValueError)` let through and 500'd the unwrapped /archreview endpoint).
    # None sentinel so an unparseable value returns `default` verbatim, never int(default).
    n = _as_num(v, None)
    return int(n) if n is not None else default


def _canon(name: str) -> str:
    """CDP/LLDP ids often carry the domain ('CORE.broadcast.x') while the inventory key does not —
    compare on the first label, lower-cased (same canonicalization the explorer applies)."""
    return str(name or "").split(".")[0].strip().lower()


def _port_gbps(port: str, speed: str):
    """Nominal port bandwidth in Gbps: parse the collected speed column first ('1000', 'a-1000',
    '10G'), fall back to the interface-name prefix, else None (excluded from ratio math)."""
    s = str(speed or "").strip().lower()
    m = re.search(r"(\d+(?:\.\d+)?)\s*g\b", s)
    if m:
        return float(m.group(1))
    m = re.search(r"(\d{2,6})$", s)
    if m:
        return float(m.group(1)) / 1000.0
    p = str(port or "").strip().lower()
    for pre, g in _NAME_GBPS:
        if p.startswith(pre):
            return g
    return None


def _ev(hosts, cap: int = 12, total=None) -> str:
    """Render an evidence host/item list, capped (cry-wolf doctrine: aggregate, don't enumerate).

    `total` is the TRUE population size for the case where `hosts` is ITSELF already a sample: a
    check's stored `evidence` is capped at `_EVIDENCE_CAP` by `add()`, so a "+N more" computed off
    the stored list understates the scope by everything the first cut dropped -- on the Meridian reference fleet
    L2-3 carries 230 evidence hosts, and the §3 scorecard rendered "5 shown +15 more" for it.
    Shown + hidden must reconcile to the real total, so the caller passes the check's
    `evidence_total` (absent when nothing was dropped, hence the None default)."""
    items = sorted({str(h) for h in hosts if h is not None})
    if not items:
        return "—"
    shown = items[:cap]
    n_total = max(_as_int(total), len(items)) if total else len(items)
    hidden = n_total - len(shown)
    return ", ".join(shown) + (f" +{hidden} more" if hidden > 0 else "")


#: How many evidence hosts a check RETAINS (the snapshot section is consumed by the explorer + the
#: webapp, so an unbounded fleet-wide list is not free). The cut is disclosed via `evidence_total`.
_EVIDENCE_CAP = 20


def _clip(text, width: int) -> str:
    """Cut a free-text field to a display width WITHOUT manufacturing a token: trim back to the last
    word boundary and mark the cut with '…'.

    A raw ``str(...)[:60]`` on a device-derived string ends mid-word, and the reader has no way to
    tell a clipped title from a complete one (the same failure the Findings-Delta devices column had:
    a hostname's front half read as a hostname). Measured on the Meridian reference fleet: 1 of 3 rendered
    High/Critical drift titles is 69 chars and was cut mid-word."""
    s = str(text if text is not None else "")
    if len(s) <= width:
        return s
    cut = s[:width]
    sp = cut.rfind(" ")
    if sp >= width // 2:                    # only honour a boundary that keeps most of the text
        cut = cut[:sp]
    return cut.rstrip(" ,;:-") + "…"


# =============================================================================
# The review engine
# =============================================================================
def _fleet_model(snap: dict):
    """One pass over the snapshot into the shared facts every check reads. Defensive throughout."""
    devices = _as_dict(snap.get("devices"))
    interfaces = _as_dict(snap.get("interfaces"))
    l3f = [r for r in _as_list(snap.get("l3_forwarding")) if isinstance(r, dict)]
    rn = _as_dict(snap.get("routing_neighbors"))

    # Evidence-based tiering (same rule as the design document): owns a gateway SVI or forms a
    # routing adjacency → L3 (core/distribution); otherwise L2 access.
    l3_hosts = set()
    for r in l3f:
        sw = r.get("switch")
        # switch may be a poisoned dict/list on an uploaded snapshot; only a str can be a device key, so a
        # non-str never matches -- guard the `in devices` membership (which would else hash an unhashable dict).
        if isinstance(sw, str) and sw in devices:
            l3_hosts.add(sw)
    for h, d in rn.items():
        if h in devices and any(_as_list(_as_dict(d).get(p)) for p in ("ospf", "eigrp", "bgp")):
            l3_hosts.add(h)
    l2_hosts = {h for h in devices if h not in l3_hosts}

    canon_fleet = {_canon(h): h for h in devices}

    # Per-host adjacency (in-fleet CDP neighbours) + uplink bandwidth (physical trunk ports whose
    # neighbour — directly or via their port-channel — is in-fleet) + provisioned edge bandwidth.
    adj = defaultdict(set)            # host -> set(real neighbour hostnames)
    uplink_gbps = defaultdict(float)  # host -> sum Gbps of physical in-fleet trunk uplinks
    edge_gbps = defaultdict(float)    # host -> sum Gbps of physical access-mode ports
    uplink_unknown = set()            # hosts with an uplink whose speed could not be resolved
    for host, ports in interfaces.items():
        ports = _as_dict(ports)
        # first pass: which port-channels on this host face an in-fleet neighbour
        po_in_fleet = set()
        for p, d in ports.items():
            d = _as_dict(d)
            nb = canon_fleet.get(_canon(d.get("cdp_neighbor")))
            if nb and nb != host:
                adj[host].add(nb)
                if re.match(r"^po\d", str(p).strip(), re.I):
                    po_in_fleet.add(re.sub(r"[^0-9]", "", str(p)))
        for p, d in ports.items():
            d = _as_dict(d)
            name = str(p).strip()
            if _VIRTUAL_PORT.match(name):
                continue                                   # physical ports only (no double-count via Po)
            mode = str(d.get("switchport_mode") or "").lower()
            g = _port_gbps(name, d.get("speed"))
            if mode == "access":
                if g:
                    edge_gbps[host] += g
            elif mode == "trunk":
                nb = canon_fleet.get(_canon(d.get("cdp_neighbor")))
                pc = re.sub(r"[^0-9]", "", str(d.get("port_channel") or ""))
                if (nb and nb != host) or (pc and pc in po_in_fleet):
                    if g:
                        uplink_gbps[host] += g
                    else:
                        uplink_unknown.add(host)

    return {"devices": devices, "interfaces": interfaces, "l3f": l3f, "rn": rn,
            "l3_hosts": l3_hosts, "l2_hosts": l2_hosts, "canon_fleet": canon_fleet,
            "adj": adj, "uplink_gbps": uplink_gbps, "edge_gbps": edge_gbps,
            "uplink_unknown": uplink_unknown}


def compute_architecture_review(snap: dict) -> dict:
    """The automated senior-engineer design review: ~24 leading-practice checks across 8 domains,
    each with a verdict + evidence + recommendation + the rule it cites, rolled up into a domain
    scorecard, an overall conformance grade and an ordered remediation queue. Pure snapshot
    synthesis; deterministic; tolerant of any missing/malformed section."""
    snap = _as_dict(snap)
    fm = _fleet_model(snap)
    devices, interfaces = fm["devices"], fm["interfaces"]
    l3_hosts, l2_hosts, adj = fm["l3_hosts"], fm["l2_hosts"], fm["adj"]
    l3f, rn = fm["l3f"], fm["rn"]
    checks = []

    def add(cid, domain, title, verdict, observed, implication, recommendation, reference,
            evidence=()):
        # The evidence cut is DISCLOSED, never silent: every renderer of `evidence` shows a
        # "+N more" tail, and computing that N from the already-capped list is a lie about scope
        # (Meridian reference fleet: 8 of 25 checks exceed the cap -- L2-3 at 230 hosts rendered as "+15 more",
        # LC-1 152, HIER-1 116, SEC-1 74, HIER-2 67, L2-1 42, RES-1 21, L2-5 21).
        # `evidence_total` is emitted ONLY when something was actually dropped -- an unconditional
        # key would restate `len(evidence)` on every check for no reader.
        _ev_all = sorted({str(h) for h in evidence if h is not None})
        rec = {"id": cid, "domain": domain, "title": title, "verdict": verdict,
               "observed": observed, "implication": implication,
               "recommendation": recommendation, "reference": reference,
               "evidence": _ev_all[:_EVIDENCE_CAP]}
        if len(_ev_all) > _EVIDENCE_CAP:
            rec["evidence_total"] = len(_ev_all)
        checks.append(rec)

    # ---------------- D1 · Hierarchy & modularity ----------------
    D1 = "Hierarchy & modularity"
    if not l3f and not any(any(_as_list(_as_dict(d).get(p)) for p in ("ospf", "eigrp", "bgp"))
                           for d in rn.values()):
        add("HIER-1", D1, "Distribution/core tier present", "not-assessable",
            "No gateway SVI or routing adjacency was observed in the collected output.",
            "The L3 tier may sit outside the collected scope (firewall/router on a stick) or the "
            "relevant commands were not captured — the tiering cannot be judged from this snapshot.",
            "Extend collection to the L3 boundary devices, then re-run the review.",
            "Cisco campus design — hierarchical access/distribution/core model")
    elif l3_hosts:
        add("HIER-1", D1, "Distribution/core tier present", "conforms",
            f"{len(l3_hosts)} L3 node(s) (gateway SVI or routing adjacency) front "
            f"{len(l2_hosts)} L2 access node(s).",
            "A defined L3 tier bounds failure domains and gives the migration a stable anchor.",
            "Carry the tiering into the target design; keep gateway functions at the distribution tier.",
            "Cisco campus design — hierarchical access/distribution/core model",
            evidence=l3_hosts)
    else:
        add("HIER-1", D1, "Distribution/core tier present", "deviation",
            "Gateway/routing evidence exists but no in-scope device owns it.",
            "A flat L2 fabric concentrates the whole campus in one failure domain.",
            "Introduce (or bring into scope) a distribution tier owning the gateways.",
            "Cisco campus design — hierarchical access/distribution/core model")

    chains = sorted({tuple(sorted((h, nb))) for h in l2_hosts for nb in adj.get(h, ())
                     if nb in l2_hosts})
    if not interfaces:
        add("HIER-2", D1, "No daisy-chained access switches", "not-assessable",
            "No interface/neighbour data in the snapshot.", "—",
            "Re-collect with CDP/LLDP output included.",
            "Cisco campus design — each access switch homes to the distribution layer")
    elif not any(adj.get(h) for h in l2_hosts):
        # conforms-by-silence guard: interfaces are present but NO access-tier CDP/LLDP neighbour was
        # observed, so daisy-chaining cannot be assessed — never grade 'conforms' off absent evidence.
        add("HIER-2", D1, "No daisy-chained access switches", "not-assessable",
            "No in-fleet CDP/LLDP neighbour was observed on the access tier — daisy-chaining cannot be "
            "assessed from this evidence.", "—",
            "Re-collect with CDP/LLDP output on the access switches.",
            "Cisco campus design — each access switch homes to the distribution layer")
    elif chains:
        add("HIER-2", D1, "No daisy-chained access switches", "deviation",
            f"{len(chains)} access-to-access inter-switch link(s): "
            + _ev(f"{a}–{b}" for a, b in chains),
            "A daisy-chained closet extends the STP/broadcast domain and rides through a peer "
            "access switch — losing the mid-chain switch strands everything behind it.",
            "Re-home each chained switch directly to the distribution tier during the migration.",
            "Cisco campus design — each access switch homes to the distribution layer",
            evidence=[a for pair in chains for a in pair])
    else:
        add("HIER-2", D1, "No daisy-chained access switches", "conforms",
            "Every observed access uplink lands on the L3 tier.",
            "Failure domains stay bounded per closet.",
            "Preserve the star topology in the target design.",
            "Cisco campus design — each access switch homes to the distribution layer")

    # ---------------- D2 · Resiliency & availability ----------------
    D2 = "Resiliency & availability"
    homed = {h: adj.get(h, set()) for h in l2_hosts if adj.get(h)}
    single = sorted(h for h, nbs in homed.items() if len(nbs) == 1)
    if not homed:
        add("RES-1", D2, "Dual-homed access uplinks", "not-assessable",
            "No in-fleet uplink neighbours were observed for the access tier.",
            "Uplink redundancy cannot be judged from this snapshot.",
            "Re-collect with CDP/LLDP output included.",
            "Cisco campus HA design — redundant uplinks from every access switch")
    elif single:
        add("RES-1", D2, "Dual-homed access uplinks", "deviation",
            f"{len(single)} of {len(homed)} access switch(es) reach exactly ONE in-fleet "
            f"neighbour: {_ev(single)}.",
            "A single uplink (or single upstream switch) is a per-closet single point of failure — "
            "one optic, one linecard or one upstream chassis takes the closet down.",
            "Dual-home each listed switch to two distribution nodes (or at minimum two linecards) "
            "as part of the migration build.",
            "Cisco campus HA design — redundant uplinks from every access switch",
            evidence=single)
    else:
        add("RES-1", D2, "Dual-homed access uplinks", "conforms",
            f"All {len(homed)} access switch(es) with observed uplinks reach ≥2 in-fleet neighbours.",
            "No single upstream device isolates a closet.",
            "Preserve dual-homing in the target design.",
            "Cisco campus HA design — redundant uplinks from every access switch")

    single_gw = sorted({str(r.get("vlan")) for r in l3f if "single-gateway" in str(r.get("risk") or "")})
    fhrp = [g for g in _as_list(snap.get("fhrp")) if isinstance(g, dict)]
    fhrp_issues = [g for g in fhrp if _as_list(g.get("issues"))]
    if not l3f:
        add("RES-2", D2, "First-hop redundancy on every gateway", "not-assessable",
            "No gateway SVIs in the collected scope.", "—",
            "Bring the gateway tier into collection scope.",
            "Cisco campus HA design — FHRP (HSRP/VRRP/GLBP) on every shared gateway")
    elif single_gw:
        add("RES-2", D2, "First-hop redundancy on every gateway", "critical",
            f"VLAN(s) {_ev(single_gw, 10)} are served by a SINGLE gateway with no FHRP peer"
            + (f"; {len(fhrp_issues)} FHRP group(s) also carry consistency issues" if fhrp_issues else "") + ".",
            "Losing that one gateway device black-holes every endpoint on the VLAN — the highest-"
            "impact availability exposure a campus carries, and a migration-readiness FAIL.",
            "Add an FHRP peer (or move the VLAN behind a redundant gateway pair) BEFORE any "
            "migration wave touches these VLANs.",
            "Cisco campus HA design — FHRP (HSRP/VRRP/GLBP) on every shared gateway",
            evidence=[r.get("switch") for r in l3f if "single-gateway" in str(r.get("risk") or "")])
    elif fhrp_issues:
        add("RES-2", D2, "First-hop redundancy on every gateway", "deviation",
            f"{len(fhrp_issues)} of {len(fhrp)} FHRP group(s) carry a consistency issue.",
            "An inconsistent FHRP group fails over unpredictably under load or maintenance.",
            "Reconcile priorities/timers/preemption on the flagged groups before cutover.",
            "Cisco campus HA design — consistent FHRP group configuration",
            evidence=[m.get("host") for g in fhrp_issues
                      for m in _as_list(g.get("members")) if isinstance(m, dict)])
    else:
        add("RES-2", D2, "First-hop redundancy on every gateway", "conforms",
            f"{len(fhrp)} FHRP group(s) observed; no single-gateway VLANs, no consistency issues.",
            "Gateway failover is in place across the shared VLANs.",
            "Preserve the FHRP design (and its addressing) through the migration.",
            "Cisco campus HA design — FHRP (HSRP/VRRP/GLBP) on every shared gateway")

    psu_known = {h: _as_int(_as_dict(d).get("num_power_supplies")) for h, d in devices.items()}
    psu_single = sorted(h for h, n in psu_known.items() if n == 1)
    # _as_int coerces a MISSING inventory to 0, so `n > 0` is the coverage test, not a PSU count.
    n_psu_known = sum(1 for n in psu_known.values() if n > 0)
    n_psu_unknown = len(psu_known) - n_psu_known
    if n_psu_known == 0:
        add("RES-3", D2, "Redundant power supplies", "not-assessable",
            "Power-supply inventory was not captured for any device.", "—",
            "Include 'show environment power' / inventory output in collection.",
            "Cisco campus HA design — N+1 power on aggregation nodes")
    elif psu_single:
        worst = sorted(set(psu_single) & l3_hosts)
        add("RES-3", D2, "Redundant power supplies", "deviation" if worst else "advisory",
            f"{len(psu_single)} device(s) run on a single PSU: {_ev(psu_single)}"
            + (f" — including L3 node(s) {_ev(worst, 6)}" if worst else "")
            + (f" (power inventory captured for {n_psu_known} of {len(psu_known)} device(s); the "
               f"remaining {n_psu_unknown} are unassessed)" if n_psu_unknown else "") + ".",
            "A single supply turns a PSU/feed fault into a chassis outage; on a distribution node "
            "that is a multi-closet event.",
            "Fit second supplies (separate feeds) on the listed devices — distribution tier first.",
            "Cisco campus HA design — N+1 power on aggregation nodes", evidence=psu_single)
    elif n_psu_unknown:
        # conforms-by-silence guard (matches HIER-2 / L2-3): the not-assessable gate above is
        # FLEET-WIDE (`no device reported a count`), so a single device reporting 2 supplies used to
        # carry the whole fleet to CONFORMS and assert "no single feed/PSU fault takes a switch
        # down" over devices whose power inventory was never captured. Partial coverage is a blind
        # spot, never health.
        add("RES-3", D2, "Redundant power supplies", "not-assessable",
            f"Power-supply inventory was captured for only {n_psu_known} of {len(psu_known)} "
            f"device(s) (all of them carry ≥2 supplies); the remaining {n_psu_unknown} report none, "
            "so fleet power redundancy cannot be graded from this evidence.", "—",
            "Include 'show environment power' / inventory output for every device, then re-review.",
            "Cisco campus HA design — N+1 power on aggregation nodes")
    else:
        add("RES-3", D2, "Redundant power supplies", "conforms",
            f"All {len(psu_known)} device(s) report captured power inventory and carry ≥2 supplies.",
            "No single feed/PSU fault takes a switch down.",
            "Keep dual feeds through the migration (verify per site during NRFU).",
            "Cisco campus HA design — N+1 power on aggregation nodes")

    fi = [r for r in _as_list(snap.get("failure_impact")) if isinstance(r, dict)]
    # conforms-by-silence guard (matches HIER-2 / RES-3 / L2-2 / L2-3): the old gate was
    # `snap.get("failure_impact") is None`, which catches ONLY a literally absent key. A simulation
    # section that is present-but-EMPTY (or a truthy non-list from an uploaded snapshot, which
    # `_as_list` normalises to []), or rows that carry no `stranded` figure at all, produced an empty
    # keystone set and fell straight through to CONFORMS — asserting "the failure simulation strands
    # no endpoints behind any single device / redundancy absorbs any one device loss" over evidence
    # that was never collected. `stranded: 0` is a real observed zero; a MISSING one is not.
    fi_measured = [r for r in fi if r.get("stranded") is not None]
    keystones = sorted((r for r in fi_measured if _as_int(r.get("stranded")) > 0),
                       key=lambda r: -_as_int(r.get("stranded")))
    if not fi_measured:
        add("RES-4", D2, "No keystone single point of failure", "not-assessable",
            ("The failure-impact simulation is absent from this snapshot."
             if not fi else
             f"The failure-impact section carries {len(fi)} row(s) but none reports a stranded-"
             "endpoint figure, so single-device blast radius cannot be graded from this evidence."),
            "—",
            "Re-run the assessment with the current engine.",
            "Availability analysis — single-device blast radius")
    elif keystones:
        # This names 5 of len(keystones) — 193 on the Meridian reference fleet, 19 on the sample — and without the
        # tail sentence it reads as the complete keystone set, i.e. "the fleet has five keystones".
        # A migration that hardens those five leaves the rest unhardened and sequences waves off the
        # wrong blast-radius population. The band is read off the DESCENDING tail so it can never
        # overstate a hidden row ("at least the 5th value" would).
        _tail = keystones[5:]
        _more = ""
        if _tail:
            _hi = _as_int(_tail[0].get("stranded"))
            _lo = _as_int(_tail[-1].get("stranded"))
            _more = (f"; and {len(_tail)} further device(s) strand "
                     + (f"{_lo}-{_hi}" if _lo != _hi else f"{_lo}") + " endpoint(s) each")
        add("RES-4", D2, "No keystone single point of failure", "advisory",
            "Losing " + "; ".join(f"{r.get('host')} strands {_as_int(r.get('stranded'))} endpoint(s)"
                                  for r in keystones[:5]) + _more + ".",
            "Endpoint dependency is concentrated — these devices ARE the availability budget.",
            "Verify each keystone's redundancy (uplinks, power, supervisor) and sequence them with "
            "the most conservative cutover plan (their waves carry the widest blast radius).",
            "Availability analysis — single-device blast radius",
            # Pass the FULL keystone list: pre-cutting at 8 here happens BEFORE add()'s own 20-item
            # evidence cap, so `evidence_total` never fired for this check and its "+N more" marker
            # reconciled against 8 instead of the real population. Let the disclosed cap do the
            # bounding.
            evidence=[r.get("host") for r in keystones])
    elif len(fi_measured) < len(fi):
        # partial coverage is a blind spot, never health — the same rule RES-3 applies to PSU inventory
        add("RES-4", D2, "No keystone single point of failure", "not-assessable",
            f"Only {len(fi_measured)} of {len(fi)} simulated device(s) report a stranded-endpoint "
            "figure (none of those strands any endpoint); the rest carry no figure, so fleet blast "
            "radius cannot be graded from this evidence.", "—",
            "Re-run the assessment with the current engine so every device is simulated.",
            "Availability analysis — single-device blast radius")
    else:
        add("RES-4", D2, "No keystone single point of failure", "conforms",
            f"The failure simulation strands no endpoints behind any single device "
            f"(all {len(fi_measured)} simulated device(s) report a stranded-endpoint figure).",
            "Redundancy absorbs any one device loss.",
            "Re-verify after each migration wave (the topology changes underneath the simulation).",
            "Availability analysis — single-device blast radius")

    # ---------------- D3 · Layer-2 design ----------------
    D3 = "Layer-2 design"
    stp_roots = _as_dict(snap.get("stp_roots"))
    roots_by_host = {}
    for h, vmap in stp_roots.items():
        n = sum(1 for v, info in _as_dict(vmap).items() if _as_dict(info).get("is_root"))
        if n:
            roots_by_host[h] = n
    access_roots = sorted(h for h in roots_by_host if h in l2_hosts)
    if not roots_by_host:
        add("L2-1", D3, "Deterministic STP root at the distribution tier", "not-assessable",
            "No explicit root bridge was observed in the collected spanning-tree output.", "—",
            "Include 'show spanning-tree' in collection, then re-review.",
            "Cisco campus design — pin the STP root (priority) at the distribution layer")
    elif access_roots:
        add("L2-1", D3, "Deterministic STP root at the distribution tier", "deviation",
            "Access switch(es) hold the STP root: "
            # This sentence is the check's whole statement of SCOPE — how many closets the design
            # must re-pin. Naming 6 of len(access_roots) (42 on the Meridian reference fleet, 17 on the sample) sizes
            # an LLD task list and a per-wave STP risk off a display cap.
            + "; ".join(f"{h} roots {roots_by_host[h]} VLAN(s)" for h in access_roots[:6])
            + (f"; and {len(access_roots) - 6} further access switch(es) rooting "
               f"{sum(roots_by_host[h] for h in access_roots[6:])} VLAN(s) between them"
               if len(access_roots) > 6 else "") + ".",
            "A root on an access closet means the L2 tree converges around the smallest box in the "
            "fabric — default (un-pinned) priorities usually caused it, and any new switch can "
            "silently steal the root.",
            "Pin root primary/secondary priorities on the distribution pair for every VLAN.",
            "Cisco campus design — pin the STP root (priority) at the distribution layer",
            evidence=access_roots)
    else:
        add("L2-1", D3, "Deterministic STP root at the distribution tier", "conforms",
            "Every observed root bridge sits on an L3 (distribution/core) node: "
            + _ev(roots_by_host) + ".",
            "The L2 tree converges around the aggregation tier as designed.",
            "Carry the pinned priorities into the target design.",
            "Cisco campus design — pin the STP root (priority) at the distribution layer",
            evidence=roots_by_host)

    n_access_ports = guarded = bpdu_data = 0
    for host, ports in interfaces.items():
        for p, d in _as_dict(ports).items():
            d = _as_dict(d)
            if str(d.get("switchport_mode") or "").lower() != "access":
                continue
            n_access_ports += 1
            # textutils.bpduguard_state is the ONE owner of this token decision (see there). The
            # hand-rolled `v not in ("no","disabled","off","false","-","--")` this replaced did not
            # contain the producer's own "Disable", so an edge port with BPDU Guard explicitly OFF
            # was counted as GUARDED here while design_advisor counted it unguarded off the same
            # field -- a false-health inversion in the direction that matters.
            st = bpduguard_state(d.get("stp_bpduguard"))
            if st is not None:
                bpdu_data += 1              # assessable: the token proves one state or the other
                if st:
                    guarded += 1
    # Two DIFFERENT denominators are in play: `bpdu_data` = access ports whose BPDU-Guard state was
    # actually captured (the assessable set), `n_access_ports` = every access port. The old
    # `guarded * 2 < n_access_ports` advisory test graded CONFORMS — asserting "the edge is
    # protected" — as soon as guarded ports passed HALF of ALL access ports, so ports with NO
    # evidence (and even ports positively observed UNGUARDED) were carried into a pass. The cited
    # rule is PortFast + BPDU Guard on EVERY edge port, so conformance needs full capture AND no
    # unguarded port; anything less abstains or deviates (matching HIER-2 / L2-3).
    n_unguarded = bpdu_data - guarded
    n_uncaptured = n_access_ports - bpdu_data
    if n_access_ports == 0 or bpdu_data == 0:
        add("L2-2", D3, "Edge protection (BPDU Guard) on access ports", "not-assessable",
            "BPDU Guard state was not captured for the access ports.", "—",
            "Include 'show spanning-tree interface detail' (or run-config) in collection.",
            "Cisco campus design — PortFast + BPDU Guard on every edge port")
    elif n_unguarded:
        add("L2-2", D3, "Edge protection (BPDU Guard) on access ports", "advisory",
            f"{n_unguarded} of the {bpdu_data} access port(s) whose BPDU-Guard state was captured "
            f"do NOT carry it ({guarded} do)"
            + (f"; a further {n_uncaptured} of {n_access_ports} access port(s) carry no BPDU-Guard "
               "evidence at all" if n_uncaptured else f" — all {n_access_ports} access port(s) captured") + ".",
            "An unguarded edge port lets a plugged-in switch (or a looped wall jack) join and "
            "reshape the spanning tree — the classic cause of campus-wide L2 meltdowns.",
            "Enable PortFast + BPDU Guard as the access-port default in the target configuration "
            "standard (the remediation plan carries the snippets).",
            "Cisco campus design — PortFast + BPDU Guard on every edge port")
    elif n_uncaptured:
        # conforms-by-silence guard: every CAPTURED port is guarded, but the rest were never
        # evidenced — "the edge is protected" would be asserted off silence.
        add("L2-2", D3, "Edge protection (BPDU Guard) on access ports", "not-assessable",
            f"All {bpdu_data} access port(s) with a captured BPDU-Guard state carry it, but "
            f"{n_uncaptured} of {n_access_ports} access port(s) carry no BPDU-Guard evidence — edge "
            "protection cannot be graded from this evidence (the rule is EVERY edge port).", "—",
            "Include 'show spanning-tree interface detail' (or run-config) for every access port, "
            "then re-review.",
            "Cisco campus design — PortFast + BPDU Guard on every edge port")
    else:
        add("L2-2", D3, "Edge protection (BPDU Guard) on access ports", "conforms",
            f"All {n_access_ports} access port(s) carry BPDU Guard (state captured on every one).",
            "The edge is protected against accidental switch insertion.",
            "Keep the standard through the migration touch.",
            "Cisco campus design — PortFast + BPDU Guard on every edge port")

    v1_access, v1_native = [], []
    has_access_vlan = has_native = False                  # evidence-capture: was the field present at all?
    for host, ports in interfaces.items():
        for p, d in _as_dict(ports).items():
            d = _as_dict(d)
            mode = str(d.get("switchport_mode") or "").lower()
            _av = str(d.get("vlan") or "").strip()
            _nv = str(d.get("trunk_native_vlan") or "").strip()
            if mode == "access":
                if _av:
                    has_access_vlan = True
                if _av == "1":
                    v1_access.append(host)
            if is_trunk_mode(mode):   # shared predicate: MUST match design_advisor._signals (PR-#396 review)
                if _nv:
                    has_native = True
                if _nv == "1":
                    v1_native.append(host)
    if not interfaces:
        add("L2-3", D3, "VLAN 1 not used for user traffic", "not-assessable",
            "No interface data in the snapshot.", "—", "Re-collect.",
            "Cisco hardening guidance — keep user traffic and trunk native VLAN off VLAN 1")
    elif not (has_access_vlan or has_native):
        # conforms-by-silence guard: interfaces present but neither the access-VLAN nor the trunk
        # native-VLAN field was captured anywhere — VLAN-1 usage cannot be assessed off absent evidence.
        add("L2-3", D3, "VLAN 1 not used for user traffic", "not-assessable",
            "Interfaces are present but no access-VLAN or trunk native-VLAN field was captured — "
            "VLAN-1 usage cannot be assessed from this evidence.", "—",
            "Re-collect with access-VLAN / trunk native-VLAN detail.",
            "Cisco hardening guidance — keep user traffic and trunk native VLAN off VLAN 1")
    elif v1_access:
        add("L2-3", D3, "VLAN 1 not used for user traffic", "deviation",
            f"Access ports on VLAN 1 on {len(set(v1_access))} switch(es): {_ev(v1_access)}"
            + (f"; native VLAN 1 trunks ({NATIVE1_CFG_BASIS} -- as in the design gap) on "
               f"{len(set(v1_native))} switch(es)" if v1_native else "") + ".",
            "VLAN 1 carries control-plane protocols on every trunk by default; user endpoints on it "
            "ride the widest possible broadcast domain and complicate segmentation.",
            "Move user ports to purpose VLANs and set trunk native VLAN to a dedicated unused VLAN "
            "as part of the per-wave port mapping.",
            "Cisco hardening guidance — keep user traffic and trunk native VLAN off VLAN 1",
            evidence=set(v1_access) | set(v1_native))
    elif v1_native:
        add("L2-3", D3, "VLAN 1 not used for user traffic", "advisory",
            f"Trunks with native VLAN 1 ({NATIVE1_CFG_BASIS} -- as in the design gap) on "
            f"{len(set(v1_native))} switch(es): {_ev(set(v1_native))}.",
            "Native VLAN 1 on inter-switch trunks exposes untagged control-plane crosstalk and is "
            "flagged by the hardening baseline.",
            "Set the native VLAN to a dedicated unused VLAN on the listed trunks.",
            "Cisco hardening guidance — keep user traffic and trunk native VLAN off VLAN 1",
            evidence=set(v1_native))
    else:
        add("L2-3", D3, "VLAN 1 not used for user traffic", "conforms",
            "No access port and no trunk native VLAN uses VLAN 1.",
            "VLAN 1 is confined to its default control-plane role.",
            "Keep the standard in the target build.",
            "Cisco hardening guidance — keep user traffic and trunk native VLAN off VLAN 1")

    macs_per_vlan = defaultdict(set)
    for host, ports in interfaces.items():
        for p, d in _as_dict(ports).items():
            d = _as_dict(d)
            v = str(d.get("vlan") or "").strip()
            mac = str(d.get("end_host_mac") or "").strip()
            if v.isdigit() and mac:
                macs_per_vlan[int(v)].add(mac)
    big = sorted(((v, len(s)) for v, s in macs_per_vlan.items() if len(s) > 512), key=lambda t: -t[1])
    mid = sorted(((v, len(s)) for v, s in macs_per_vlan.items() if 256 < len(s) <= 512), key=lambda t: -t[1])
    if not macs_per_vlan:
        add("L2-4", D3, "Bounded broadcast domains", "not-assessable",
            "No endpoint MACs were captured, so per-VLAN endpoint counts are unknown.", "—",
            "Include MAC/ARP table output in collection.",
            "Campus design guidance — constrain broadcast domains (≈ a /23 of endpoints per VLAN)")
    elif big:
        add("L2-4", D3, "Bounded broadcast domains", "deviation",
            "Oversized VLAN(s): " + "; ".join(f"VLAN {v} ≈{n} observed endpoints" for v, n in big[:5]) + ".",
            "A broadcast domain past ~512 hosts amplifies every ARP/broadcast event fleet-wide and "
            "widens the failure (and migration blast) radius.",
            "Split the listed VLANs along closet or function lines in the target design; the move-"
            "group plan already isolates them as their own waves.",
            "Campus design guidance — constrain broadcast domains (≈ a /23 of endpoints per VLAN)")
    elif mid:
        add("L2-4", D3, "Bounded broadcast domains", "advisory",
            "; ".join(f"VLAN {v} ≈{n} observed endpoints" for v, n in mid[:5])
            + " — above the comfortable /24 envelope.",
            "Broadcast load grows quadratically with flat-domain size; these are trending large.",
            "Plan a split when the migration re-addresses these segments.",
            "Campus design guidance — constrain broadcast domains (≈ a /23 of endpoints per VLAN)")
    else:
        add("L2-4", D3, "Bounded broadcast domains", "conforms",
            f"Largest observed VLAN carries {max((len(s) for s in macs_per_vlan.values()), default=0)} "
            "endpoint(s) — within the conventional envelope.",
            "Broadcast domains are bounded.",
            "Keep VLAN sizing within a /24-/23 of endpoints in the target design.",
            "Campus design guidance — constrain broadcast domains (≈ a /23 of endpoints per VLAN)")

    trunks = allow_all = pruned_data = 0
    allow_all_hosts = set()
    for host, ports in interfaces.items():
        for p, d in _as_dict(ports).items():
            d = _as_dict(d)
            if str(d.get("switchport_mode") or "").lower() != "trunk" or _VIRTUAL_PORT.match(str(p)):
                continue
            trunks += 1
            av = str(d.get("trunk_allowed_vlans") or "").strip().lower()
            if av:
                pruned_data += 1
                if av in ("all", "1-4094") or av.startswith("1-4094"):   # 'none' = allows NO VLANs (the INVERSE of allow-all): never an un-pruned trunk
                    allow_all += 1
                    allow_all_hosts.add(host)
    if trunks == 0 or pruned_data == 0:
        add("L2-5", D3, "Trunks pruned to required VLANs", "not-assessable",
            "Allowed-VLAN lists were not captured for the trunks.", "—",
            "Include 'show interfaces trunk' in collection.",
            "Cisco campus design — prune trunk allowed-VLAN lists to the required set")
    elif allow_all:
        add("L2-5", D3, "Trunks pruned to required VLANs", "advisory",
            f"{allow_all} of {pruned_data} trunk(s) with captured lists allow ALL VLANs "
            f"({_ev(allow_all_hosts)}).",
            "Un-pruned trunks extend every VLAN's broadcast/STP domain across the whole fabric and "
            "turn any VLAN fault into a campus-wide event.",
            "Prune each trunk to the VLANs its closet actually serves (the workbook's per-port "
            "evidence gives the required set).",
            "Cisco campus design — prune trunk allowed-VLAN lists to the required set",
            evidence=allow_all_hosts)
    else:
        add("L2-5", D3, "Trunks pruned to required VLANs", "conforms",
            f"All {pruned_data} trunk(s) with captured lists are pruned.",
            "VLAN scope stays bounded per closet.",
            "Maintain pruning in the per-wave port maps.",
            "Cisco campus design — prune trunk allowed-VLAN lists to the required set")

    # ---------------- D4 · Layer-3 & gateway design ----------------
    D4 = "Layer-3 & gateway design"
    access_vlans = set()
    for host, ports in interfaces.items():
        for p, d in _as_dict(ports).items():
            d = _as_dict(d)
            v = str(d.get("vlan") or "").strip()
            if str(d.get("switchport_mode") or "").lower() == "access" and v.isdigit():
                access_vlans.add(v)
    gw_vlans = {str(r.get("vlan")) for r in l3f}
    orphan = sorted(access_vlans - gw_vlans, key=lambda s: int(s))
    if not l3f:
        add("L3-1", D4, "Every access VLAN has a gateway in scope", "not-assessable",
            "No gateway SVIs in the collected scope — gateways may live off-fabric.", "—",
            "Bring the gateway tier into collection scope to close this blind spot.",
            "Campus L3 design — every user VLAN terminates on a known, redundant gateway")
    elif orphan:
        add("L3-1", D4, "Every access VLAN has a gateway in scope", "deviation",
            f"{len(orphan)} access VLAN(s) have no SVI anywhere in the scan: VLAN(s) {_ev(orphan, 10)}.",
            "An unaccounted gateway is a migration black hole: endpoints move and their default "
            "gateway is not where the plan thinks it is.",
            "Locate each VLAN's real gateway (off-fabric router/firewall?) and add it to the design "
            "record before the wave that carries it.",
            "Campus L3 design — every user VLAN terminates on a known, redundant gateway")
    else:
        add("L3-1", D4, "Every access VLAN has a gateway in scope", "conforms",
            f"All {len(access_vlans)} access VLAN(s) terminate on an in-scope gateway SVI.",
            "The L2/L3 boundary is fully accounted for.",
            "Keep the gateway register in the LLD as the addressing source of truth.",
            "Campus L3 design — every user VLAN terminates on a known, redundant gateway")

    svi_owner = {}
    for r in l3f:
        v = str(r.get("vlan") or "")
        sw = r.get("switch")
        if v and isinstance(sw, str) and sw in devices:
            svi_owner.setdefault(v, sw)
    no_relay = []
    relay_data = False
    for v, owner in svi_owner.items():
        rec = _as_dict(_as_dict(interfaces.get(owner)).get(f"Vlan{v}"))
        if rec:
            relay_data = True
            if v in access_vlans and macs_per_vlan.get(int(v) if v.isdigit() else -1) \
                    and not str(rec.get("dhcp_helpers") or "").strip():
                no_relay.append(v)
    if not svi_owner or not relay_data:
        add("L3-2", D4, "DHCP relay on endpoint-serving gateways", "not-assessable",
            "No in-scope SVI interface records to read helpers from.", "—",
            "Bring gateway SVIs into collection scope.",
            "Campus L3 design — ip helper-address on every client-serving SVI")
    elif no_relay:
        add("L3-2", D4, "DHCP relay on endpoint-serving gateways", "advisory",
            f"{len(no_relay)} endpoint-carrying VLAN(s) have a gateway with no DHCP relay "
            f"configured: VLAN(s) {_ev(sorted(no_relay, key=lambda s: int(s) if s.isdigit() else 0), 10)}.",
            "If those clients use central DHCP, the missing helper is a silent post-cutover black "
            "hole (the classic day-1 failure); if addressing is static/local this is by design.",
            "Confirm the addressing model per VLAN; add helpers where central DHCP is intended "
            "before the carrying wave.",
            "Campus L3 design — ip helper-address on every client-serving SVI")
    else:
        add("L3-2", D4, "DHCP relay on endpoint-serving gateways", "conforms",
            "Every endpoint-carrying VLAN's gateway carries a DHCP relay (or serves no clients).",
            "Client addressing survives the cutover.",
            "Verify relay targets remain routable from the new gateway placement (NRFU case).",
            "Campus L3 design — ip helper-address on every client-serving SVI")

    igp_hosts = {p: sorted(h for h, d in rn.items() if _as_list(_as_dict(d).get(p)))
                 for p in ("ospf", "eigrp", "bgp")}
    redist = _as_dict(snap.get("redistribution"))
    n_redist = sum(len(_as_list(v)) for v in redist.values())
    dual_igp = bool(igp_hosts["ospf"]) and bool(igp_hosts["eigrp"])
    if not any(igp_hosts.values()):
        add("L3-3", D4, "Single IGP, controlled redistribution", "not-assessable",
            "No dynamic-routing adjacencies were observed (static/connected only, or the commands "
            "were not collected).", "—",
            "If dynamic routing exists off-fabric, extend collection to it.",
            "Campus L3 design — one IGP where possible; document every redistribution boundary")
    elif dual_igp or n_redist:
        add("L3-3", D4, "Single IGP, controlled redistribution", "advisory",
            (f"Two IGPs run side-by-side (OSPF on {len(igp_hosts['ospf'])}, EIGRP on "
             f"{len(igp_hosts['eigrp'])} device(s)). " if dual_igp else "")
            + (f"{n_redist} redistribution edge(s) observed." if n_redist else ""),
            "Every extra protocol and redistribution point is a loop/suboptimal-routing risk and a "
            "troubleshooting tax; migrations are when undocumented redistribution bites.",
            "Plan convergence to a single IGP in the target design; until then, document each "
            "redistribution edge (route-map, direction, owner) in the LLD.",
            "Campus L3 design — one IGP where possible; document every redistribution boundary",
            evidence=set(igp_hosts["ospf"]) | set(igp_hosts["eigrp"]) | set(redist))
    else:
        add("L3-3", D4, "Single IGP, controlled redistribution", "conforms",
            "A single routing protocol domain with no redistribution edges.",
            "The L3 control plane is as simple as it can be.",
            "Keep it that way through the migration.",
            "Campus L3 design — one IGP where possible; document every redistribution boundary")

    misaligned = []
    align_data = False
    root_host_by_vlan = {}
    for h, vmap in stp_roots.items():
        for v, info in _as_dict(vmap).items():
            if _as_dict(info).get("is_root"):
                root_host_by_vlan.setdefault(str(v), h)
    for r in l3f:
        v = str(r.get("vlan") or "")
        if "active" in str(r.get("fhrp") or "").lower() and v in root_host_by_vlan:
            align_data = True
            if r.get("switch") != root_host_by_vlan[v]:
                misaligned.append(v)
    if not align_data:
        add("L3-4", D4, "FHRP active aligned with the STP root", "not-assessable",
            "No VLAN has both a known FHRP active gateway and an observed root bridge.", "—",
            "Collect 'show standby'/'show spanning-tree' together to enable this check.",
            "Cisco campus design — co-locate the FHRP active gateway with the STP root per VLAN")
    elif misaligned:
        add("L3-4", D4, "FHRP active aligned with the STP root", "advisory",
            f"VLAN(s) {_ev(misaligned, 10)}: the configured FHRP active sits on a different switch "
            "than the STP root.",
            "Upstream traffic crosses the inter-distribution link twice — it works, but burns the "
            "interconnect and doubles the failure surface. (Config placement, not live state — "
            "hence advisory.)",
            "Align FHRP priorities with the root-bridge priorities per VLAN in the target design.",
            "Cisco campus design — co-locate the FHRP active gateway with the STP root per VLAN")
    else:
        add("L3-4", D4, "FHRP active aligned with the STP root", "conforms",
            "Every VLAN with both facts known has its FHRP active on the root bridge.",
            "L2 and L3 forwarding converge on the same node per VLAN.",
            "Preserve the pairing in the target design.",
            "Cisco campus design — co-locate the FHRP active gateway with the STP root per VLAN")

    # ---------------- D5 · Capacity & oversubscription ----------------
    D5 = "Capacity & oversubscription"
    edge_gbps, uplink_gbps = fm["edge_gbps"], fm["uplink_gbps"]
    ratios = {h: edge_gbps[h] / uplink_gbps[h]
              for h in l2_hosts if edge_gbps.get(h) and uplink_gbps.get(h)}
    over = sorted(((h, r) for h, r in ratios.items() if r > 20.0), key=lambda t: -t[1])
    if not ratios:
        add("CAP-1", D5, "Access uplink oversubscription within 20:1", "not-assessable",
            "No access switch had both edge and uplink bandwidth resolvable from the snapshot.", "—",
            "Ensure interface speed columns (or recognizable interface names) are collected.",
            "Cisco campus rule of thumb — ≤20:1 access→distribution, ≤4:1 distribution→core")
    elif over:
        add("CAP-1", D5, "Access uplink oversubscription within 20:1", "deviation",
            "Provisioned edge:uplink ratios exceed 20:1 on "
            + "; ".join(f"{h} ({r:.0f}:1)" for h, r in over[:6]) + ".",
            "Past the 20:1 rule of thumb, transient congestion on the uplink under simultaneous "
            "load is probable — and a migration consolidates traffic exactly there.",
            "Add or upgrade uplinks (or spread closets) on the listed switches in the target build; "
            "re-check distribution→core at ≤4:1 once the L3 tier is fully in scope.",
            "Cisco campus rule of thumb — ≤20:1 access→distribution, ≤4:1 distribution→core",
            evidence=[h for h, _ in over])
    else:
        worst = max(ratios.values())
        add("CAP-1", D5, "Access uplink oversubscription within 20:1", "conforms",
            f"{len(ratios)} access switch(es) measured; worst provisioned ratio {worst:.0f}:1.",
            "Uplink capacity is within the conventional envelope.",
            "Re-validate ratios whenever the target design changes uplink counts or speeds.",
            "Cisco campus rule of thumb — ≤20:1 access→distribution, ≤4:1 distribution→core")

    cap_rows = [r for r in _as_list(snap.get("capacity")) if isinstance(r, dict)]
    # conforms-by-silence guard: `port_util` is absent-prone (the runbook already excludes rows whose
    # active-port count was never observed, "port_util can read 0.0 from absent evidence"). The old
    # test filtered non-numeric rows into oblivion and then graded CONFORMS off whatever survived —
    # so a capacity section whose utilisation column was never captured asserted "All N device(s)
    # hold port headroom below the 90% line" over N unmeasured closets, and RAISED the overall grade.
    cap_measured = [r for r in cap_rows if isinstance(r.get("port_util"), (int, float))]
    tight = sorted((r for r in cap_measured if r["port_util"] > 90), key=lambda r: -r["port_util"])
    if not cap_measured:
        add("CAP-2", D5, "Port-capacity headroom for growth", "not-assessable",
            ("No capacity rows in the snapshot." if not cap_rows else
             f"{len(cap_rows)} capacity row(s) carry no numeric port-utilisation figure, so port "
             "headroom cannot be graded from this evidence."),
            "—", "Re-run with the current engine.",
            "Capacity planning — keep ~10% port headroom per closet")
    elif tight:
        add("CAP-2", D5, "Port-capacity headroom for growth", "advisory",
            # The BoM sizes replacement hardware off this list, so the count past the 90% line must
            # not be the display cap: 6 named of len(tight) (16 on the sample fleet) means ten full
            # closets get no spare-port allowance.
            "; ".join(f"{r.get('hostname')} at {r.get('port_util')}% port utilisation"
                      for r in tight[:6])
            + (f"; and {len(tight) - 6} further closet(s) above the 90% line"
               if len(tight) > 6 else "") + ".",
            "A full closet forces ad-hoc daisy-chaining (see HIER-2) the day a port is needed.",
            "Size replacement hardware with ≥10% spare ports per listed closet.",
            "Capacity planning — keep ~10% port headroom per closet",
            evidence=[r.get("hostname") for r in tight])
    elif len(cap_measured) < len(cap_rows):
        # partial coverage is a blind spot, never health (same rule as RES-3 / L2-2)
        add("CAP-2", D5, "Port-capacity headroom for growth", "not-assessable",
            f"All {len(cap_measured)} device(s) with a captured port-utilisation figure sit below the "
            f"90% line, but {len(cap_rows) - len(cap_measured)} of {len(cap_rows)} capacity row(s) "
            "carry no figure — fleet port headroom cannot be graded from this evidence.", "—",
            "Capture port utilisation for every device, then re-review.",
            "Capacity planning — keep ~10% port headroom per closet")
    else:
        add("CAP-2", D5, "Port-capacity headroom for growth", "conforms",
            f"All {len(cap_measured)} device(s) hold port headroom below the 90% line "
            "(utilisation captured on every one).",
            "Growth fits without topology workarounds.",
            "Carry the same headroom rule into the BoM for the target build.",
            "Capacity planning — keep ~10% port headroom per closet")

    # ---------------- D6 · Operational readiness & hygiene ----------------
    D6 = "Operational readiness & hygiene"
    drift_raw = snap.get("operational_drift")
    drift = [d for d in _as_list(drift_raw) if isinstance(d, dict)]
    d_hi = [d for d in drift if d.get("severity") in ("Critical", "High")]
    if drift_raw is None:
        add("OPS-1", D6, "No operational drift / false health", "not-assessable",
            "The operational-drift detector did not run on this snapshot.", "—",
            "Re-run the assessment with the current engine.",
            "Assessment doctrine — configured is not healthy; up is not healthy")
    elif d_hi:
        add("OPS-1", D6, "No operational drift / false health", "deviation",
            f"{len(d_hi)} High/Critical drift finding(s) (of {len(drift)} total): "
            + "; ".join(_clip(d.get("title") or d.get("detail") or "?", 60) for d in d_hi[:3]) + ".",
            "Drift findings are the traps a green control plane hides — they surface on the first "
            "change, which is exactly what a migration is.",
            "Close the High/Critical drift items before their carrying waves (punch-list owners).",
            "Assessment doctrine — configured is not healthy; up is not healthy")
    elif drift:
        add("OPS-1", D6, "No operational drift / false health", "advisory",
            f"{len(drift)} lower-severity drift finding(s).",
            "Low-grade drift compounds; the migration touch is the cheapest moment to clear it.",
            "Fold the items into the per-wave work orders.",
            "Assessment doctrine — configured is not healthy; up is not healthy")
    else:
        add("OPS-1", D6, "No operational drift / false health", "conforms",
            "The drift detector found no divergence between configured intent and observed state.",
            "What the configs promise is what the fleet does.",
            "Re-run the detector as each wave's post-checks.",
            "Assessment doctrine — configured is not healthy; up is not healthy")

    hyg = _as_dict(snap.get("config_hygiene"))

    def _undef_n(v):
        """Undefined-reference count for one device, or None when the record evidences NEITHER a
        count nor a list. `_as_int(<missing>)` coerced to 0, so a hygiene record that never reported
        the figure was carried to CONFORMS ("No undefined references across the fleet") off silence —
        the same conforms-by-silence hole guarded at RES-3 / RES-4 / L2-2 / L2-3 / CAP-2. The list is
        read when the summary lacks the count, so a partial record still yields a real number."""
        v = _as_dict(v)
        s = _as_dict(v.get("summary")).get("undefined")
        if s is not None:
            return _as_int(s)
        return len(v["undefined"]) if isinstance(v.get("undefined"), list) else None

    hyg_measured = {h: n for h, n in ((h, _undef_n(v)) for h, v in hyg.items()) if n is not None}
    undef_hosts = {h: n for h, n in hyg_measured.items() if n > 0}
    if not hyg or not hyg_measured:
        add("OPS-2", D6, "Clean configuration hygiene", "not-assessable",
            ("Config-hygiene analysis is absent from this snapshot." if not hyg else
             f"{len(hyg)} config-hygiene record(s) report no undefined-reference count, so hygiene "
             "cannot be graded from this evidence."), "—",
            "Re-run the assessment with the current engine.",
            "Operational hygiene — no undefined references in running configuration")
    elif undef_hosts:
        add("OPS-2", D6, "Clean configuration hygiene", "advisory",
            f"{sum(undef_hosts.values())} undefined reference(s) across {len(undef_hosts)} device(s): "
            + _ev(undef_hosts) + ".",
            "A referenced-but-undefined ACL/route-map fails OPEN or CLOSED depending on platform — "
            "either way it is latent surprise during cutover.",
            "Resolve or remove each undefined reference before its device migrates.",
            "Operational hygiene — no undefined references in running configuration",
            evidence=undef_hosts)
    elif len(hyg_measured) < len(hyg):
        # partial coverage is a blind spot, never health
        add("OPS-2", D6, "Clean configuration hygiene", "not-assessable",
            f"All {len(hyg_measured)} device(s) with a captured undefined-reference count are clean, "
            f"but {len(hyg) - len(hyg_measured)} of {len(hyg)} hygiene record(s) report no count — "
            "fleet hygiene cannot be graded from this evidence.", "—",
            "Re-run the assessment with the current engine so every device reports a count.",
            "Operational hygiene — no undefined references in running configuration")
    else:
        add("OPS-2", D6, "Clean configuration hygiene", "conforms",
            f"No undefined references across the fleet ({len(hyg_measured)} device(s) reported a "
            "count).",
            "Config intent is internally consistent.",
            "Keep the hygiene gate in the target config standard.",
            "Operational hygiene — no undefined references in running configuration")

    coll = _as_dict(snap.get("collection_completeness"))
    csum = _as_dict(coll.get("summary"))
    n_blind = _as_int(csum.get("partial")) + _as_int(csum.get("not_collected"))
    if not csum:
        add("OPS-3", D6, "Collection confidence", "not-assessable",
            "Completeness evidence is absent — review confidence is UNKNOWN.",
            "Absence of evidence is not evidence of completeness.",
            "Re-assess with a current engine collection before relying on the verdicts.",
            "Assessment doctrine — state collection blind spots explicitly")
    elif n_blind:
        add("OPS-3", D6, "Collection confidence", "advisory",
            f"{n_blind} device(s) have partial or missing collection.",
            "Every verdict in this review is provisional on those devices.",
            "Re-collect the flagged devices, then regenerate this review.",
            "Assessment doctrine — state collection blind spots explicitly")
    else:
        add("OPS-3", D6, "Collection confidence", "conforms",
            "Collection is complete across the fleet.",
            "The review stands on full evidence.",
            "—",
            "Assessment doctrine — state collection blind spots explicitly")

    # OPS-4 · cross-domain dependencies & interoperability (N37): declare the multi-NOS interop surface
    _nos = {}
    for d in _as_dict(snap.get("devices")).values():
        fam = str(_as_dict(d).get("platform") or "").strip().lower() or "unknown"
        _nos[fam] = _nos.get(fam, 0) + 1
    _fams = {k: v for k, v in _nos.items() if k != "unknown"}
    if len(_fams) >= 2:
        _mix = ", ".join(f"{v}× {k.upper()}" for k, v in sorted(_fams.items(), key=lambda x: -x[1]))
        add("OPS-4", D6, "Cross-domain dependencies & interoperability declared", "advisory",
            f"The fleet spans {len(_fams)} switch NOS families ({_mix}).",
            "A multi-NOS estate must keep feature / CLI / automation parity across families, and any "
            "cross-family migration (e.g. IOS->NX-OS) is an explicit interoperability step, not a "
            "like-for-like swap — undeclared, it surfaces as a cutover surprise.",
            "Declare the platform-dependency + interoperability surface in the HLD (the design "
            "interoperability footprint) and validate cross-family behaviour during NRFU.",
            "Cisco Advanced Services design-review practice — declare platform dependencies & interop")
    elif _fams:
        add("OPS-4", D6, "Cross-domain dependencies & interoperability declared", "conforms",
            f"The fleet is a single NOS family ({next(iter(_fams)).upper()}), simplifying interoperability.",
            "A homogeneous NOS estate minimises cross-platform interoperability risk.",
            "Keep the single-family footprint where practical; declare any new platform's dependencies.",
            "Cisco Advanced Services design-review practice — declare platform dependencies & interop")

    # ---------------- D7 · Security & segmentation ----------------
    D7 = "Security & segmentation"
    sec = _as_dict(snap.get("security"))
    grades = {h: str(_as_dict(_as_dict(v).get("summary")).get("grade") or "") for h, v in sec.items()}
    weak = sorted(h for h, g in grades.items() if g == "weak")
    partial = sorted(h for h, g in grades.items() if g == "partial")
    if not grades:
        add("SEC-1", D7, "Management-plane hardening (CIS baseline)", "not-assessable",
            "No security/compliance posture in the snapshot (run-configs not captured).", "—",
            "Capture running configurations to enable the CIS-aligned posture checks.",
            "CIS Cisco IOS Benchmark — AAA, NTP, logging, SSHv2/VTY, SNMPv3 hardening")
    elif weak:
        add("SEC-1", D7, "Management-plane hardening (CIS baseline)", "deviation",
            f"{len(weak)} device(s) grade WEAK (a high-severity hardening check fails): {_ev(weak)}"
            + (f"; {len(partial)} more grade partial" if partial else "") + ".",
            "A weak management plane (open HTTP/Smart Install, missing AAA, default communities) is "
            "the attack surface that turns a migration window into an incident.",
            "Apply the generated remediation snippets to the WEAK devices first — before any "
            "maintenance window publicises change activity.",
            "CIS Cisco IOS Benchmark — AAA, NTP, logging, SSHv2/VTY, SNMPv3 hardening",
            evidence=weak)
    elif partial:
        add("SEC-1", D7, "Management-plane hardening (CIS baseline)", "advisory",
            f"{len(partial)} device(s) grade partial; none weak.",
            "Medium/low hardening gaps remain.",
            "Fold the remaining CIS items into the target configuration standard.",
            "CIS Cisco IOS Benchmark — AAA, NTP, logging, SSHv2/VTY, SNMPv3 hardening",
            evidence=partial)
    else:
        add("SEC-1", D7, "Management-plane hardening (CIS baseline)", "conforms",
            f"All {len(grades)} device(s) with captured posture grade hardened.",
            "The management plane meets the baseline.",
            "Re-grade after the migration (new platforms, new defaults).",
            "CIS Cisco IOS Benchmark — AAA, NTP, logging, SSHv2/VTY, SNMPv3 hardening")

    # SSOT (Law 1): the gateway-tier segmentation posture has ONE owner -- analyze.compute_segmentation
    # -> snap['segmentation'] -- read here via ssot.segmentation_facts (interface-derived fallback for
    # pre-segmentation snapshots). The local recount this replaced asked a DIFFERENT question under the
    # owner's label: it counted every non-default VRF configured ANYWHERE (mgmt0's management VRF, a vPC
    # keepalive VRF) as a "VRF in use", so on the Meridian reference fleet it rendered "4 VRF(s) in use" beside the
    # runbook's "1 VRF(s) across 231 gateway SVI(s)" off the same snapshot -- and, worse, those 4
    # gateway-less VRFs made `not vrfs` False, downgrading a fabric the owner calls FLAT from the
    # deviation branch to "advisory / partial enforcement". Only VRFs that actually CARRY a gateway SVI
    # can segment user traffic, so the gate reads `gateway_vrfs`; the management-only VRFs are still
    # reported, under their own name, so nothing is dropped.
    _seg = _ssot_mod.segmentation_facts(snap)
    gw_vrfs = list(_seg.get("gateway_vrfs") or [])
    other_vrfs = list(_seg.get("other_vrfs") or [])
    n_svis = _as_int(_seg.get("n_gateways"))
    n_acl_svis = _as_int(_seg.get("n_with_acl"))
    _other = (f" ({len(other_vrfs)} further non-default VRF(s) are configured — "
              f"{', '.join(other_vrfs[:6])} — but carry no gateway SVI, so they segment no user traffic)"
              if other_vrfs else "")
    if n_svis == 0:
        add("SEC-2", D7, "Segmentation enforced at the L3 boundary", "not-assessable",
            "No in-scope gateway SVIs to read policy from.", "—",
            "Bring the gateway tier into collection scope.",
            "Segmentation guidance — enforce inter-segment policy at the L3 boundary (VRF / gateway ACL)")
    elif not gw_vrfs and n_acl_svis * 4 < n_svis:
        add("SEC-2", D7, "Segmentation enforced at the L3 boundary", "deviation",
            f"No gateway SVI sits in a non-default VRF and only {n_acl_svis} of {n_svis} carry an ACL"
            + _other + ".",
            "Inter-VLAN traffic is openly routed: any compromised segment reaches every other "
            "segment at wire speed.",
            "Use the migration to introduce intended segmentation — VRF separation for true "
            "isolation, gateway ACLs as the minimum policy line.",
            "Segmentation guidance — enforce inter-segment policy at the L3 boundary (VRF / gateway ACL)")
    elif n_acl_svis < n_svis or not gw_vrfs:
        add("SEC-2", D7, "Segmentation enforced at the L3 boundary", "advisory",
            (f"{len(gw_vrfs)} VRF(s) carry a gateway SVI; " if gw_vrfs
             else "No gateway SVI sits in a non-default VRF; ")
            + f"{n_acl_svis} of {n_svis} gateway SVIs carry an ACL" + _other + ".",
            "Partial enforcement leaves unpoliced inter-segment paths.",
            "Close the unpoliced SVIs (or fold them into a VRF) in the target design.",
            "Segmentation guidance — enforce inter-segment policy at the L3 boundary (VRF / gateway ACL)")
    else:
        add("SEC-2", D7, "Segmentation enforced at the L3 boundary", "conforms",
            f"{len(gw_vrfs)} VRF(s) carry a gateway SVI; every gateway SVI carries policy." + _other,
            "Inter-segment traffic crosses an enforced boundary.",
            "Preserve policy parity through the migration (NRFU verifies).",
            "Segmentation guidance — enforce inter-segment policy at the L3 boundary (VRF / gateway ACL)")

    # ---------------- D8 · Lifecycle & software ----------------
    D8 = "Lifecycle & software"
    lc_sum = _as_dict(_as_dict(snap.get("lifecycle_risk")).get("summary"))
    n_ldos, n_eos, n_near = (_as_int(lc_sum.get("n_past_ldos")), _as_int(lc_sum.get("n_past_eos")),
                             _as_int(lc_sum.get("n_near")))
    # Published by analyze.compute_lifecycle_risk (analyze.py:6199) and, until now, read by nothing
    # in this chain — which is exactly how an all-Unknown fleet reached the `conforms` branch below.
    n_unknown = _as_int(lc_sum.get("n_unknown"))
    lc_rows = _as_list(_as_dict(snap.get("lifecycle_risk")).get("per_device"))
    lc_ldos_dev = [str(_as_dict(r).get("host")) for r in lc_rows
                   if _as_dict(r).get("band") == "Past-LDoS"]
    lc_eos_dev = [str(_as_dict(r).get("host")) for r in lc_rows
                  if _as_dict(r).get("band") == "Past-EoS"]
    lc_near_dev = [str(_as_dict(r).get("host")) for r in lc_rows
                   if _as_dict(r).get("band") == "Near-LDoS"]
    # COVERAGE is ORTHOGONAL to what was FOUND, so it must not be chained behind the findings.
    # Subordinated as an `elif` (the shape crd.py carried until this round), the coverage disclosure
    # fired only when every adverse count was zero — so the BROWNFIELD fleet, some gear banded and
    # most not, lost it entirely. Measured: 1 Past-LDoS + 20 Unknown emitted verdict `critical`
    # reading "1 device(s) are past LAST-DAY-OF-SUPPORT: sw0." and never mentioned the 20 devices
    # whose support state was never determined — while LC-1 feeds domains[].score_pct, the
    # conformance grade, the Architecture Review DOCX and its workbook sheet.
    #
    # Build the coverage sentence ONCE and append it through a wrapper, so EVERY LC-1 exit carries
    # it by construction rather than by each branch remembering to. A future branch cannot forget.
    _lc_cov_obs = (f" Separately, {n_unknown} device(s) could NOT be lifecycle-banded — either no exact "
                   "EoX bulletin row matched or the matched row's retained source authority/complete dates "
                   "did not verify, so their support state is UNKNOWN and this check is unanswered for them."
                   if n_unknown else "")
    _lc_cov_imp = (" The undetermined device(s) carry no evidence of vendor support in either "
                   "direction; absent lifecycle data is not a pass, and a plan that assumes it is "
                   "may stage a platform with no TAC path into a change window."
                   if n_unknown else "")
    _lc_cov_rec = (" Resolve every undetermined platform against Cisco's EoX portal and re-run, or "
                   "record them as an accepted risk in the engagement RAID log."
                   if n_unknown else "")

    def add_lc(verdict, observed, implication, recommendation, evidence=()):
        add("LC-1", D8, "Hardware lifecycle exposure is controlled", verdict,
            observed + _lc_cov_obs, implication + _lc_cov_imp, recommendation + _lc_cov_rec,
            "Cisco EoX policy — exact EoS/LDoS date bands; support entitlement assessed separately",
            evidence=evidence)

    if not lc_sum:
        add_lc("not-assessable", "Lifecycle analysis is absent from this snapshot.", "—",
               "Re-run the assessment with the current engine (offline EoL KB).")
    elif n_ldos:
        add_lc("critical",
               f"{n_ldos} device(s) are past LAST-DAY-OF-SUPPORT: {_ev(lc_ldos_dev)}."
               + (f" Separately, {n_near} device(s) are within one year of LDoS: "
                  f"{_ev(lc_near_dev)}." if n_near else "")
               + (f" Separately, {n_eos} device(s) are past end-of-sale with LDoS still future: "
                  f"{_ev(lc_eos_dev)}." if n_eos else ""),
               "For the Past-LDoS devices there is NO standard TAC escalation path or software fix — "
               "a fault during a migration window on those devices has no vendor backstop. The "
               "Past-EoS-only date band does not establish contract entitlement.",
               "Replace (not upgrade) the LDoS devices as the first migration waves; stage spares for "
               "any that must briefly remain. Schedule Near-LDoS and EoS-only devices for refresh "
               "before their recorded LDoS dates.",
               evidence=lc_ldos_dev)
    elif n_near:
        add_lc("deviation",
               f"{n_near} device(s) are within one year of LDoS: {_ev(lc_near_dev)}."
               + (f" Separately, {n_eos} device(s) are past end-of-sale with LDoS still future: "
                  f"{_ev(lc_eos_dev)}." if n_eos else ""),
               "The Near-LDoS replacement window is short; Past-EoS also requires refresh planning. "
               "Neither date band establishes contract entitlement.",
               "Schedule Near-LDoS devices before their recorded deadline and place EoS-only devices "
               "in the refresh plan.", evidence=lc_near_dev)
    elif n_eos:
        add_lc("advisory",
               f"{n_eos} device(s) are past end-of-sale with LDoS still future: {_ev(lc_eos_dev)}.",
               "The recorded refresh window is finite; support entitlement is not inferred.",
               "Track exact LDoS dates and schedule the EoS platforms in the replacement plan.",
               evidence=lc_eos_dev)
    else:
        # Nothing adverse was found. Whether that is a PASS depends entirely on coverage: with any
        # undetermined device the honest verdict is `not-assessable`, because the sentence "every
        # lifecycle-banded device is in the pre-EoS date band" is VACUOUSLY true of a fleet
        # nothing could band — a true sentence under a false verdict, the hardest form of
        # "absence rendered as health" to spot. Measured on Catalyst 6500s years past support,
        # banded Unknown because no retained bulletin covers them.
        add_lc("not-assessable" if n_unknown else "conforms",
               ("Every device that COULD be lifecycle-banded is in the pre-EoS date band "
                "(public schema value: Active)."
                if n_unknown else
                "Every lifecycle-banded device is in the pre-EoS date band (public schema value: Active)."),
               ("For the banded devices, retained bulletin dates show EoS has not passed; support "
                "entitlement was not assessed."
                if n_unknown else
                "Retained bulletin dates show no assessed device past EoS or within one year of LDoS; "
                "support entitlement was not assessed."),
               "Re-check EoX bulletins at each engagement gate (statuses move).")

    sw_by_model = defaultdict(set)
    for h, d in devices.items():
        v = str(_as_dict(d).get("sw_version") or "").strip()
        if v:
            sw_by_model[str(_as_dict(d).get("model") or "Unknown")].add(v)
    mixed = sorted(m for m, vs in sw_by_model.items() if len(vs) > 1)
    if not sw_by_model:
        add("LC-2", D8, "One software image per platform", "not-assessable",
            "No software-version data captured.", "—", "Include 'show version' in collection.",
            "AS LLD convention — minimise image variance per platform")
    elif mixed:
        add("LC-2", D8, "One software image per platform", "advisory",
            f"{len(mixed)} platform(s) run mixed images: {_ev(mixed, 8)}.",
            "Version sprawl multiplies the defect surface and invalidates like-for-like rollback "
            "assumptions between devices of the same model.",
            "Standardise each platform on one image (the LLD §3.5 names the candidates) before or "
            "during the migration waves.",
            "AS LLD convention — minimise image variance per platform")
    else:
        add("LC-2", D8, "One software image per platform", "conforms",
            "Every platform runs a single image.",
            "Troubleshooting and rollback are uniform per model.",
            "Pin the images as the minimum-version baseline in the LLD.",
            "AS LLD convention — minimise image variance per platform")

    # ---------------- rollups ----------------
    domain_order = (D1, D2, D3, D4, D5, D6, D7, D8)
    domains = []
    for dom in domain_order:
        dc = [c for c in checks if c["domain"] == dom]
        assessable = [c for c in dc if c["verdict"] != "not-assessable"]
        verdict = (min(assessable, key=lambda c: V_RANK[c["verdict"]])["verdict"]
                   if assessable else "not-assessable")
        score = (round(100 * sum(_V_WEIGHT[c["verdict"]] for c in assessable) / len(assessable))
                 if assessable else None)
        domains.append({"key": dom, "verdict": verdict, "score_pct": score,
                        "checks": [c["id"] for c in dc]})

    assessable = [c for c in checks if c["verdict"] != "not-assessable"]
    n_by = {v: sum(1 for c in checks if c["verdict"] == v) for v in V_RANK}
    score_pct = (round(100 * sum(_V_WEIGHT[c["verdict"]] for c in assessable) / len(assessable))
                 if assessable else None)
    if score_pct is None:
        grade, grade_label = "N/A", "Insufficient evidence to grade"
    else:
        grade, grade_label = next((g, lbl) for thr, g, lbl in _GRADES if score_pct >= thr)

    actions = sorted((c for c in checks if c["verdict"] in ("critical", "deviation", "advisory")),
                     key=lambda c: (V_RANK[c["verdict"]], -len(c["evidence"]), c["id"]))
    top_actions = []
    for i, c in enumerate(actions[:10]):
        a = {"rank": i + 1, "id": c["id"], "domain": c["domain"], "verdict": c["verdict"],
             "action": c["recommendation"], "evidence": c["evidence"]}
        if c.get("evidence_total"):        # carry the true scope with the sampled list (see `add`)
            a["evidence_total"] = c["evidence_total"]
        top_actions.append(a)
    # NB the queue is the top 10 of `len(actions)`; §1.1 discloses that ratio, recomputed there from
    # `checks` rather than published as a new summary key (the key would restate a derivable count).

    if score_pct is None:
        statement = ("The snapshot carries too little evidence to grade the architecture — every "
                     "check returned not-assessable. Re-collect and re-run the review.")
    else:
        worst_doms = [d["key"] for d in domains if d["verdict"] in ("critical", "deviation")]
        statement = (f"Architecture conformance grade {grade} ({score_pct}%) — {grade_label}. "
                     f"{n_by['conforms']} of {len(assessable)} assessable checks conform to leading "
                     f"practice; {n_by['critical']} critical and {n_by['deviation']} material "
                     f"deviation(s) require remediation"
                     + (f", concentrated in: {', '.join(worst_doms[:4])}"
                        # "concentrated in: A, B, C, D" over 7 affected domains (Meridian reference fleet) inverts
                        # the finding -- the deviations are FLEET-WIDE, not concentrated.
                        + (f" and {len(worst_doms) - 4} further domain(s)"
                           if len(worst_doms) > 4 else "")
                        if worst_doms else "")
                     + f". {n_by['not-assessable']} check(s) were not assessable from this snapshot "
                       "and are listed in the review scope.")

    return {"domains": domains, "checks": checks, "top_actions": top_actions,
            "summary": {"n_checks": len(checks), "n_assessable": len(assessable),
                        "n_conforms": n_by["conforms"], "n_advisory": n_by["advisory"],
                        "n_deviation": n_by["deviation"], "n_critical": n_by["critical"],
                        "n_not_assessable": n_by["not-assessable"],
                        "score_pct": score_pct, "grade": grade, "grade_label": grade_label,
                        "statement": statement}}


# =============================================================================
# The DOCX writer
# =============================================================================
def write_archreview_docx(output_path: str, snap_dict: dict, label: str) -> None:
    """Emit the Architecture Review & Conformance Report to `output_path`. Prefers the review
    already attached at snap['architecture_review'] (one source of truth with the CLI run);
    recomputes only for older snapshots that predate the key. Fail-soft: a missing python-docx is
    a warning + skip; any unexpected render error is logged, never raised by the CLI path."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor
    except ImportError:
        logger.warning("  Architecture review (DOCX) skipped: python-docx not installed "
                       "(pip install python-docx to enable the architecture-review deliverable).")
        return

    snap = xml_safe_deep(_as_dict(snap_dict))   # one bad device byte (noncharacter/surrogate) must not abort the docx
    label = xml_safe(label) if isinstance(label, str) else (str(label) if label is not None else "")
    ar = _as_dict(snap.get("architecture_review"))
    if not _as_list(ar.get("checks")):
        ar = compute_architecture_review(snap)
    summary = _as_dict(ar.get("summary"))
    checks = {c.get("id"): c for c in _as_list(ar.get("checks")) if isinstance(c, dict) and isinstance(c.get("id"), str)}

    from cisco_toolkit.brand_tokens import DOC_NAVY_RGB
    NAVY = RGBColor(*DOC_NAVY_RGB)
    GREY = RGBColor(0x59, 0x59, 0x59)
    VERDICT_LABEL = {"conforms": "CONFORMS", "advisory": "ADVISORY", "deviation": "DEVIATION",
                     "critical": "CRITICAL DEVIATION", "not-assessable": "NOT ASSESSABLE"}

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    def _label_run(p, label_text, value, color=NAVY):
        r = p.add_run(label_text); r.bold = True; r.font.color.rgb = color
        p.add_run(" " + (str(value) if value not in (None, "") else "—"))

    def table(headers, rows, widths=None):
        return add_table(doc, headers, rows, widths, fixed=False)

    # ---- title page ----
    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Architecture Review & Conformance Report"); tr.bold = True
    tr.font.size = Pt(26); tr.font.color.rgb = NAVY
    sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s2 = sub2.add_run("Leading-practice design review · availability analysis · conformance scorecard")
    s2.font.size = Pt(13); s2.font.color.rgb = GREY
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(label); sr.font.size = Pt(13); sr.font.color.rgb = GREY
    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    from cisco_toolkit import ssot as _ssot   # DE-01/Law 1: cover shows the CANONICAL inventory count,
    _nd = _ssot.canonical_facts(snap).get("n_devices")   # not a recount of the (collected-only) devices map
    _nd = _nd if _nd is not None else len(_as_dict(snap.get("devices")))   # coverage-honest fallback
    meta.add_run(f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')}  ·  "
                 f"{_nd} devices in scope  ·  "
                 f"script {snap.get('script_version', '')}").font.color.rgb = GREY
    gp = doc.add_paragraph(); gp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gr = gp.add_run(f"Conformance grade: {summary.get('grade', 'N/A')}"
                    + (f"  ({summary.get('score_pct')}%)" if summary.get("score_pct") is not None else "")
                    + f" — {summary.get('grade_label', '')}")
    gr.bold = True; gr.font.size = Pt(15); gr.font.color.rgb = NAVY
    status = doc.add_paragraph(); status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st = status.add_run("DRAFT — generated design review; a senior engineer must validate the "
                        "verdicts before customer presentation.")
    st.italic = True; st.font.color.rgb = GREY
    doc.add_page_break()

    # ---- document control ----
    add_document_control(
        doc, document="Architecture Review & Conformance Report", label=label,
        engine_version=str(snap.get("script_version", "")), generated_at=snap.get("generated_at"),
        collected_at=snap.get("collected_at"),
        audience="Customer architecture owners and the engagement's design authority; consumed by "
                 "the HLD/LLD as target-state input and by the engagement plan as gate evidence.",
        exclude=("archreview",),
        extra_assumptions=(
            "Verdicts are derived from configuration and neighbour evidence only — live forwarding "
            "state, telemetry and performance counters are out of scope and listed in §2.",
            "Leading-practice references cite Cisco campus design guidance and the CIS benchmark as "
            "rules of thumb; the customer's own standards take precedence where they conflict.",
        ))
    doc.add_page_break()

    # ---- TOC (shared field-code helper, V3.23.171) ----
    add_toc(doc)


    # Deliverable Excellence (DE-01): answer-first at-a-glance register + single-source-of-truth signal.
    add_excellence_front(doc, snap_dict)
    # Deliverable Excellence (DE-01): consolidated Inputs-Required register (Law 9).
    add_inputs_required(doc, snap_dict)
    # ===== 1. Executive verdict =====
    doc.add_heading("1. Executive Verdict", level=1)
    doc.add_paragraph(str(summary.get("statement") or ""))
    table(["Design domain", "Verdict", "Domain score"],
          [(d.get("key"), VERDICT_LABEL.get(d.get("verdict"), str(d.get("verdict"))),
            f"{d.get('score_pct')}%" if d.get("score_pct") is not None else "not assessable")
           for d in _as_list(ar.get("domains")) if isinstance(d, dict)],
          widths=[3.2, 2.0, 1.6])
    top = _as_list(ar.get("top_actions"))
    if top:
        # The queue is the top 10 of every actionable check; unqualified, a 10-row table titled
        # "Priority remediation queue" reads as the WHOLE remediation scope (Meridian reference fleet: 19 actionable
        # checks, sample fleet: 14). Recomputed from `checks`, so an older snapshot discloses too.
        n_actionable = sum(1 for c in _as_list(ar.get("checks"))
                           if isinstance(c, dict)
                           and c.get("verdict") in ("critical", "deviation", "advisory"))
        doc.add_heading("1.1 Priority remediation queue", level=2)
        doc.add_paragraph(
            "Ordered by verdict severity then blast radius — the order a senior engineer would work "
            "them. Items overlapping the migration punch-list keep their punch-list owners."
            + (f" Showing the top {len(top)} of {n_actionable} check(s) requiring remediation; the "
               f"remaining {n_actionable - len(top)} carry the same verdicts and are listed in the "
               "§3 scorecard and §4 findings."
               if n_actionable > len(top) else ""))
        table(["#", "Check", "Severity", "Action", "Evidence"],
              [(a.get("rank"), a.get("id"), VERDICT_LABEL.get(a.get("verdict"), ""),
                _clip(a.get("action"), 220),
                _ev(_as_list(a.get("evidence")), 4, a.get("evidence_total")))
               for a in top if isinstance(a, dict)],
              widths=[0.4, 0.7, 1.2, 3.3, 1.4])

    # ===== 2. Review scope & method =====
    doc.add_heading("2. Review Scope & Method", level=1)
    doc.add_paragraph(
        f"The review walks the assessment snapshot through {summary.get('n_checks', 0)} deterministic "
        "leading-practice checks across eight design domains, mirroring the design-review activity of "
        "a Cisco Advanced Services engagement: each check renders a verdict (conforms / advisory / "
        "deviation / critical deviation), its evidence, why it matters, the remediation, and the rule "
        "it cites. The conformance score weights each assessable check (conforms 1.0, advisory 0.7, "
        "deviation 0.35, critical 0.0); checks whose evidence was not captured score NOTHING — they "
        "are declared below rather than silently skipped.")
    na = [c for c in _as_list(ar.get("checks"))
          if isinstance(c, dict) and c.get("verdict") == "not-assessable"]
    if na:
        doc.add_heading("2.1 Checks not assessable from this snapshot", level=2)
        table(["Check", "Why not assessable", "How to close the gap"],
              [(f"{c.get('id')} — {c.get('title')}", _clip(c.get("observed"), 160),
                _clip(c.get("recommendation"), 160)) for c in na],
              widths=[2.2, 2.5, 2.2])
    doc.add_heading("2.2 Outside the reach of an offline snapshot", level=2)
    doc.add_paragraph(
        "A complete architecture review also covers the following; they need data this offline "
        "assessment deliberately does not collect, and belong to a follow-on engagement:")
    table(["Review area", "Why it is out of scope here"], list(NOT_ASSESSED), widths=[2.6, 4.2])

    # ===== 3. Conformance scorecard =====
    doc.add_heading("3. Conformance Scorecard", level=1)
    doc.add_paragraph("Every check, its verdict and its evidence — the one-page audit trail. "
                      "§4 carries the engineering detail for everything that is not CONFORMS.")
    table(["Check", "Domain", "Verdict", "Evidence"],
          [(f"{c.get('id')} — {c.get('title')}", c.get("domain"),
            VERDICT_LABEL.get(c.get("verdict"), str(c.get("verdict"))),
            _ev(_as_list(c.get("evidence")), 5, c.get("evidence_total")))
           for c in _as_list(ar.get("checks")) if isinstance(c, dict)],
          widths=[2.6, 1.7, 1.3, 1.5])

    # ===== 4. Findings by design domain =====
    doc.add_heading("4. Findings by Design Domain", level=1)
    for i, d in enumerate(_as_list(ar.get("domains")), start=1):
        if not isinstance(d, dict):
            continue
        doc.add_heading(f"4.{i} {d.get('key')}", level=2)
        dom_checks = [checks.get(cid) for cid in _as_list(d.get("checks"))]
        dom_checks = [c for c in dom_checks if isinstance(c, dict)]
        ok = [c for c in dom_checks if c.get("verdict") == "conforms"]
        if ok:
            doc.add_paragraph("Conforms: " + "; ".join(
                f"{c.get('id')} ({c.get('title')})" for c in ok) + ".")
        for c in dom_checks:
            if c.get("verdict") in ("conforms", "not-assessable"):
                continue
            p = doc.add_paragraph()
            r = p.add_run(f"{c.get('id')} — {c.get('title')}  [{VERDICT_LABEL.get(c.get('verdict'))}]")
            r.bold = True; r.font.color.rgb = NAVY
            _label_run(doc.add_paragraph(), "Observed:", c.get("observed"), GREY)
            _label_run(doc.add_paragraph(), "Why it matters:", c.get("implication"), GREY)
            _label_run(doc.add_paragraph(), "Recommendation:", c.get("recommendation"), GREY)
            _label_run(doc.add_paragraph(), "Reference:", c.get("reference"), GREY)
            doc.add_paragraph()

    # ===== 5. Design decision traceability — the audit trail BEHIND the conformance grade =====
    from cisco_toolkit.design import build_design_traceability   # local import: avoid any module-load cycle
    trace = build_design_traceability(snap)                      # reads the xml-safe `snap` copy, not the raw arg
    handoff_sec = 5
    if trace:
        doc.add_heading("5. Design Decision Traceability", level=1)
        doc.add_paragraph(
            "The audit trail behind the conformance grade — every recommended design decision traced to the CCDE "
            "principle and published source that justify it, plus the collected evidence it stands on. A decision "
            "with no citation is shown '(uncited)', never given a manufactured reference.")
        table(["Design decision", "CCDE principle", "Citation", "Evidence", "Source fields"],
              [(r["decision"], r["principle"], r["citation"], r["evidence"], r["fields"] or "—") for r in trace],
              widths=[1.6, 1.7, 1.2, 1.9, 1.5])
        handoff_sec = 6                                          # renumber Hand-Off only when §5 actually rendered

    # ===== Hand-off into the document set (§5 if no traceability, §6 if it rendered — never a numbering gap) =====
    doc.add_heading(f"{handoff_sec}. Hand-Off Into the Document Set", level=1)
    doc.add_paragraph(
        "Deviations feed the Target-State Design Recommendations (design document §4) and the "
        "migration punch-list; the priority queue in §1.1 is the working order. The engagement plan "
        "of record carries each item's owner and gate; the NRFU plan re-verifies the affected checks "
        "post-cutover. Regenerate this review after each wave — the verdicts move as the fleet does.")

    # Deliverable Excellence (DE-01): shared glossary before the sign-off gate.
    add_glossary(doc)

    add_acceptance(
        doc, scope_note="Acceptance confirms the review's verdicts as the agreed design-gap baseline; "
                        "each remediation proceeds under the engagement's change control.")

    doc.save(output_path)
    logger.info(f"[Phase 37] Architecture review (DOCX) written: {output_path} "
                f"(grade {summary.get('grade', 'N/A')}, "
                f"{summary.get('n_checks', 0)} checks)")
