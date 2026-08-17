"""Assessment & migration RUNBOOK (DOCX) — the narrative sign-off twin of the workbook.

NEW-V3.23.93: the cisco-network-assessment doctrine mandates TWO complementary deliverables for
an engagement — an operational workbook (XLSX, which this toolkit already emits) AND a formal
assessment & migration runbook (DOCX). This module adds the second. It is PURE PRESENTATION of the
already-computed snapshot (one source of truth: it reads the same snap_dict the workbook + explorer
consume, so every number agrees) and follows the skill's report-standards 12-section structure with
evidence-disciplined finding blocks (Observed / Interpretation / Impact / Confidence / Unknowns /
Next validation).

Evidence discipline (false-health doctrine) is baked in, because this dataset is config- and
CDP/LLDP-derived, NOT live forwarding telemetry:
  * gateway PLACEMENT (an SVI is configured) is Confirmed; the active forwarding owner over time is
    Unknown -- "gateway active is not service proof".
  * inter-switch links are Inferred-high (a neighbour was advertised), never Confirmed; off-scan
    links are Unknown -- "no inferred link is presented as confirmed".
  * endpoint presence is Confirmed at snapshot time only; MAC aging means absence is Unknown.

python-docx is an OPTIONAL dependency: imported inside the function so the package imports fine
without it, and a missing library is a warning + skip (the run's workbook/explorer/JSON already
saved), exactly like the HTML explorer's template-missing path.
"""
import logging
import re
import textwrap
from collections import Counter
from datetime import datetime

from cisco_toolkit.docmeta import add_acceptance, add_document_control, add_excellence_front, add_glossary, add_inputs_required, add_table, add_toc
from cisco_toolkit.docmeta import as_dict as _as_dict, as_list as _as_list   # shared snapshot-section coercers (ops.py uses the same): a truthy non-dict/non-list section must degrade, not crash
from cisco_toolkit.textutils import _as_num   # fail-soft numeric coercion of device-derived leaf counts

logger = logging.getLogger(__name__)

# Confidence vocabulary (governance/terminology-and-confidence-glossary.md).
_CONF_CONFIRMED = "Confirmed"
_CONF_HIGH = "Inferred-high"
_CONF_MED = "Inferred-medium"
_CONF_UNKNOWN = "Unknown"

_SEV_ORDER = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3, "Info": 4}


def _dossier_coverage(d) -> tuple:
    """(n_na, n_axes, thin) for one device-dossier row — how much of its risk band rests on evidence.

    MIRROR of the canonical rule in cisco_toolkit/excel.py::dossier_coverage. It is duplicated rather
    than imported because excel.py pulls in openpyxl + analyze + parse, and this module must stay a thin
    presentation layer; the duplication is caught by EXECUTION, not by hope --
    tests/test_runbook.py::test_dossier_coverage_mirrors_the_canonical_rule drives this function over
    excel.DOSSIER_COVERAGE_CASES and fails the moment the two disagree.

    Why it exists: compute_device_dossiers ABSTAINS on an axis it has no evidence for (state 'na') and an
    abstention is weighted ZERO in the exposure score, so an asset whose EoL / software / control-plane /
    log / CIS / hygiene / drift / QoS axes were ALL un-assessed scores risk_index 0 -> band 'Low' ->
    "No stacked risk — routine migration handling." That is a collection gap rendered as a clean bill of
    health. FAIL-CLOSED on a missing census: no census means no denominator, so coverage is itself NOT
    ASSESSED (thin=True), never 'fine'."""
    dd = d if isinstance(d, dict) else {}
    ax = dd.get("exposures")
    axes = [e for e in ax if isinstance(e, dict)] if isinstance(ax, list) else []
    n_axes = len(axes)
    n_na = sum(1 for e in axes if e.get("state") == "na")
    if not n_axes:                                    # no axis census published -> fall back to the count
        try:
            n_na = max(0, int(dd.get("n_na") or 0))
        except (TypeError, ValueError):
            n_na = 0
        return n_na, 0, True                          # unknown denominator -> NOT ASSESSED, never "fine"
    return n_na, n_axes, n_na * 2 >= n_axes


def _coverage_cell(d) -> str:
    """The per-asset coverage cell: how many risk axes were NOT ASSESSED, and whether the band therefore
    rests on absent evidence. Never renders an empty/blank cell — absence has to be legible."""
    n_na, n_axes, thin = _dossier_coverage(d)
    cell = (f"{n_na} of {n_axes} NOT ASSESSED" if n_axes
            else f"axis census ABSENT — coverage NOT ASSESSED ({n_na} recorded not-assessed)")
    return cell + (" — band computed on absent evidence" if thin else "")


def _av_authority(mi) -> str:
    """The authority qualifier that must travel with any broadcast/AV (on-air) multicast count.

    review r9 F2. compute_multicast_intelligence classifies a group as on-air from the offline registry's
    CURATED media semantics and publishes the basis alongside it (`summary.n_av_groups_authoritative`,
    per-group `on_air_authoritative`). On the shipped pack that count is ZERO -- every multicast row is
    curated -- so a headline reading "44 broadcast/AV group(s)" with no qualifier presents a curated
    judgement in the same voice as an IANA-assigned fact, and that same flag is what escalates a
    mac-alias finding from Medium to High. FAIL-CLOSED: a snapshot that carries no authority census at
    all reads NOT ASSESSED, never 'authoritative'."""
    s = _as_dict(_as_dict(mi).get("summary"))
    if "n_av_groups_authoritative" not in s:
        return (" — on-air classification authority NOT ASSESSED in this snapshot; treat the "
                "broadcast/AV label as curated, not authoritative")
    n_auth = _as_num(s.get("n_av_groups_authoritative"))
    n_av = _as_num(s.get("n_av_groups"))
    if n_auth <= 0:
        return (" — NONE of them classified on-air by an authoritative source: the broadcast/AV label is "
                "a CURATED offline-registry classification, not a measurement")
    if n_auth > n_av:
        # `n_av_groups_authoritative` and `n_av_groups` are published independently, so "N of M" is
        # only meaningful when N <= M. Unchecked, an incoherent snapshot renders "7 of 3", and a
        # reader has no way to tell that from a real ratio. Disclose the incoherence instead.
        return (f" — census INCOHERENT: {n_auth:g} classification(s) reported as authoritative out of "
                f"{n_av:g} on-air group(s); the authority split cannot be stated for this snapshot")
    return (f" — {n_auth:g} of {n_av:g} on-air classification(s) rest on an authoritative source; the "
            "remainder are CURATED, not measurements")


def _integrity_failures(snap: dict) -> list:
    s = snap if isinstance(snap, dict) else {}
    ai = _as_dict(s.get("assessment_integrity"))
    failed = ai.get("failed_phases")
    failures = [str(x) for x in failed] if isinstance(failed, list) else ([str(failed)] if failed else [])
    for key, value in ai.items():
        if key not in ("failed_phases", "n_violations") and \
                str(value).strip().lower() in ("failed", "compute_failed", "unavailable", "error"):
            failures.append(f"{key}: {value}")
    for key, value in s.items():
        if isinstance(value, dict) and value.get("_unavailable"):
            failures.append(f"{key}: unavailable")
    return list(dict.fromkeys(failures))


def _phase_failed(failures, *tokens) -> bool:
    def _key(value) -> str:
        return "".join(ch for ch in str(value).lower() if ch.isalnum())
    return any(any(_key(token) in _key(item) for token in tokens) for item in failures)


def _disclose(doc, total, shown, noun, where="", note=""):
    """Emit the family's display-cap disclosure when a table/block list renders only the first
    `shown` of `total` records — the house rule every writer follows (`excel._xls_cell_value`
    states it; mop.py §x.2 and html.py's Findings-Delta column were fixed to it).

    An UNDISCLOSED cap is the silent-truncation bug class: the reader takes the rendered rows for
    the whole population, and this document states the FULL totals in §1's metric register and in
    each section's own lead-in sentence — so a capped list makes the runbook contradict itself
    (§6.2 rendered 8 blocks under a §1 headline of 407 Critical/High). Shown + hidden always
    reconcile to the real total. Emits nothing when nothing was dropped; returns the hidden count."""
    try:
        hidden = max(0, int(total) - int(shown))
    except (TypeError, ValueError):        # fail-soft like every other read here
        return 0
    if hidden:
        doc.add_paragraph(f"…and {hidden} further {noun} not shown ({shown} of {total} rendered)"
                          + (f" — see the workbook's '{where}' sheet." if where else ".")
                          + (f" {note}" if note else ""))
    return hidden


def _join_cap(items, n, sep="; "):
    """Inline join bounded to `n` items WITH the family's trailing '(+N more)' marker — the
    subnet-reachability pattern in excel.py and html.py's device-list column. A bare `[:n]` join
    reads as the complete set."""
    vals = [str(x) for x in items]
    return sep.join(vals[:n]) + (f" (+{len(vals) - n} more)" if len(vals) > n else "")


def _endpoint_census(snap: dict):
    """Derive endpoint counts the SAME way the workbook does (access ports carrying a host MAC) so the
    runbook reconciles to it. Returns (total, per_vlan Counter, per_switch Counter). Trunk MAC-table
    entries are excluded by construction (only access ports are counted). Pure read of the snapshot."""
    import re
    per_vlan: Counter = Counter()
    per_switch: Counter = Counter()
    total = 0
    for host, ports in _as_dict(snap.get("interfaces")).items():   # a truthy non-dict 'interfaces' (or per-host value) must not crash .items()/.values()
        for d in (v for v in _as_dict(ports).values() if isinstance(v, dict)):   # skip a truthy non-dict per-interface value
            if (d.get("switchport_mode") or "") != "Access":
                continue
            _ehm = d.get("end_host_mac")            # tolerate a list-valued / non-string end_host_mac (L2)
            _ehm = " ".join(str(x) for x in _ehm) if isinstance(_ehm, list) else str(_ehm or "")
            macs = [m for m in re.split(r"[,\s]+", _ehm) if m]
            if not macs:
                continue
            n = len(macs)
            total += n
            per_switch[host] += n
            vlan = str(d.get("vlan") or "")         # tolerate a non-string vlan (L2)
            if vlan.isdigit():
                per_vlan[int(vlan)] += n
    return total, per_vlan, per_switch


def _gateways(snap: dict):
    """Per-VLAN gateway register from l3_forwarding (SVIs). Returns list of dicts with the
    false-health separation the doctrine requires: placement (Confirmed) vs active-owner /
    forwarding (Unknown). FHRP role from config is Inferred-high."""
    rows = []
    for r in _as_list(snap.get("l3_forwarding")):   # a truthy non-list 'l3_forwarding' must not crash iteration
        if not isinstance(r, dict):
            continue
        rows.append({
            "vlan": r.get("vlan", ""), "svi_ip": r.get("svi_ip", ""), "owner": r.get("switch", ""),
            "fhrp": r.get("fhrp", ""), "role": r.get("role", ""), "risk": r.get("risk", ""),
        })
    return rows


def write_runbook_docx(output_path: str, snap_dict: dict, label: str, flow_paths: dict = None) -> None:
    """Emit the assessment & migration runbook (.docx) from the live snapshot. Safe/fail-soft:
    a missing python-docx is a warning + skip; never crashes a run whose other outputs are saved."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.section import WD_ORIENT
        from docx.shared import Pt, RGBColor, Inches
    except ImportError:
        logger.warning("  Runbook (DOCX) skipped: python-docx not installed "
                       "(pip install python-docx to enable the narrative runbook deliverable).")
        return

    # Input normalization (mirrors the deck's guards): a None / non-dict snapshot must degrade, not raise (L2);
    # and an XML-illegal control byte (BEL, U+FFFE/U+FFFF) in ANY device-derived string -- collected with
    # errors='ignore' -- otherwise aborts the WHOLE runbook at docx save (the deck ships from the same data). Strip
    # every illegal char from the snapshot tree + the label ONCE, up front, reusing the workbook's sanitizer (#5).
    from cisco_toolkit.excel import _xls_sanitize

    def _clean(o):
        if isinstance(o, dict):
            return {k: _clean(v) for k, v in o.items()}
        if isinstance(o, list):
            return [_clean(v) for v in o]
        return _xls_sanitize(o)

    snap_dict = _clean(snap_dict if isinstance(snap_dict, dict) else {})
    if not isinstance(snap_dict.get("executive_brief"), dict):    # a truthy non-dict slips '(... or {}).get' (L2)
        snap_dict["executive_brief"] = {}
    integrity_failures = _integrity_failures(snap_dict)
    label = _xls_sanitize(label) if isinstance(label, str) else (str(label) if label is not None else "")

    from cisco_toolkit.brand_tokens import DOC_NAVY_RGB
    NAVY = RGBColor(*DOC_NAVY_RGB)
    GREY = RGBColor(0x59, 0x59, 0x59)
    doc = Document()

    # base style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    def _label_run(p, label_text, value, color=NAVY):
        r = p.add_run(label_text)
        r.bold = True
        r.font.color.rgb = color
        # str()-guard the value (mirrors docmeta._kv): a truthy NON-string value (an int count, a list)
        # would raise `TypeError: can only concatenate str` here -- the lone family writer missing the guard.
        p.add_run(" " + (str(value) if value not in (None, "") else "—"))

    def finding_block(title, *, severity, scope, observed, interpretation, impact,
                      confidence, unknowns, next_validation, remediation=""):
        """One evidence-disciplined finding (report-standards.md finding format). Each line is
        explicitly typed (Observed / Interpretation / Impact / Confidence / Unknowns / Next), and
        Remediation is omitted when not evidence-justified."""
        doc.add_heading(title, level=3)
        meta = doc.add_paragraph()
        _label_run(meta, "Severity:", severity); meta.add_run("    ")
        _label_run(meta, "Scope:", scope)
        for lbl, val in (("Observed Evidence:", observed), ("Interpretation:", interpretation),
                         ("Impact / Blast Radius:", impact), ("Confidence:", confidence),
                         ("Unknowns:", unknowns), ("Next Validation:", next_validation)):
            p = doc.add_paragraph(); _label_run(p, lbl, val)
        if remediation:
            p = doc.add_paragraph(); _label_run(p, "Remediation:", remediation)

    def table(headers, rows, widths=None):
        # Delegates to the family's single table builder (docmeta.add_table) — V3.23.152 dedup.
        return add_table(doc, headers, rows, widths, fixed=False)

    # ---- snapshot-derived headline numbers (reconciled; all sections cite these) ----
    # audit-5 totality: a hostile/malformed upload (the webapp deliverable route) can make a section a truthy
    # non-list / non-dict, or carry a non-dict ELEMENT -> degrade, never raise (like the mop/crd/ops siblings).
    def _R(x): return [r for r in _as_list(x) if isinstance(r, dict)]   # record list: coerce container + keep only dict rows
    devices = {_h: _d for _h, _d in _as_dict(snap_dict.get("devices")).items() if isinstance(_d, dict)}   # drop a truthy non-dict per-device value
    hs = _R(snap_dict.get("health_scores"))
    bands = Counter(r.get("band", "") for r in hs)
    move_groups = _R(snap_dict.get("move_groups"))
    mr = _R(snap_dict.get("migration_readiness"))
    ws = _R(snap_dict.get("wave_sequencing"))
    cross_layer = _R(snap_dict.get("cross_layer"))
    punchlist = _R(snap_dict.get("punchlist"))
    failure_impact = _R(snap_dict.get("failure_impact"))
    link_centrality = _R(snap_dict.get("link_centrality"))
    gw = _gateways(snap_dict)
    ep_total, ep_per_vlan, ep_per_switch = _endpoint_census(snap_dict)
    _scale_dev = _as_dict(_as_dict(snap_dict.get("executive_brief")).get("scale")).get("n_devices")
    n_dev = _scale_dev if isinstance(_scale_dev, int) else len(devices)   # canonical-first (SSOT), like the endpoint/VLAN reads below
    n_links = len([r for r in link_centrality])  # link records == inter-switch links
    bridges = [r for r in link_centrality if r.get("is_bridge")]
    crit_cl = [f for f in cross_layer if f.get("severity") == "Critical"]
    high_cl = [f for f in cross_layer if f.get("severity") == "High"]

    # ---- title page ----
    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Network Assessment & Migration Runbook"); tr.bold = True
    tr.font.size = Pt(26); tr.font.color.rgb = NAVY
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(label); sr.font.size = Pt(13); sr.font.color.rgb = GREY
    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')}  ·  "
                 f"{n_dev} devices in scope  ·  script {snap_dict.get('script_version', '')}").font.color.rgb = GREY
    status = doc.add_paragraph(); status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st = status.add_run("DRAFT for war-room review — evidence-led; numbers reconcile to the workbook.")
    st.italic = True; st.font.color.rgb = GREY
    doc.add_paragraph()
    note = doc.add_paragraph()
    _label_run(note, "Evidence posture:",
               "This runbook is derived from collected configuration and CDP/LLDP neighbour data, not "
               "from live forwarding telemetry. Gateway PLACEMENT is Confirmed where an SVI is "
               "configured; the active forwarding owner over time is Unknown. Inter-switch links are "
               "Inferred-high (a neighbour was advertised), never Confirmed; off-scan links are Unknown. "
               "Endpoint presence is Confirmed at snapshot time only. Configured is not healthy; up is "
               "not healthy; gateway-active is not service proof.", GREY)
    doc.add_page_break()

    # ---- document control (AS-style front matter; unnumbered so §1–§13 are untouched) ----
    add_document_control(
        doc, document="Network Assessment & Migration Runbook", label=label,
        engine_version=str(snap_dict.get("script_version", "")),
        generated_at=snap_dict.get("generated_at"),
        collected_at=snap_dict.get("collected_at"),
        audience="The customer's network engineering and operations teams and the migration war "
                 "room; review owner: the customer's network architecture owner.",
        exclude=("runbook",))
    doc.add_page_break()

    # ---- table of contents (shared field-code helper, V3.23.171) ----
    add_toc(doc)


    # Deliverable Excellence (DE-01): answer-first at-a-glance register + single-source-of-truth signal.
    add_excellence_front(doc, snap_dict)
    # Deliverable Excellence (DE-01): consolidated Inputs-Required register (Law 9).
    add_inputs_required(doc, snap_dict)
    if integrity_failures:
        p = doc.add_paragraph()
        r0 = p.add_run("ASSESSMENT INTEGRITY - UNVERIFIED: ")
        r0.bold = True
        r0.font.color.rgb = RGBColor(0xB0, 0x2A, 0x1E)
        p.add_run(
            f"{len(integrity_failures)} phase/sentinel failure(s) were recorded ("
            + "; ".join(integrity_failures)
            + "). Empty fallback sections are not clean results. Any affected zero count or empty "
              "register below is an abstention until the failed phase is repaired and the runbook regenerated.")
    # ===== 1. Assessment Header & Executive Summary =====
    doc.add_heading("1. Assessment Header & Executive Summary", level=1)
    _scale = _as_dict(_as_dict(snap_dict.get("executive_brief")).get("scale"))
    _ep_canon = _scale.get("n_endpoints"); _vl_canon = _scale.get("n_vlans")
    # Coverage-honesty on the headline register: each of these three rows is a COUNT over a snapshot
    # section, and an absent section counts 0 — which renders as a measured all-clear ("0 bridge
    # links", "0 / 0 cross-layer findings") on the runbook's own front page. Absence of the analysis
    # is not absence of the exposure, so an unpublished axis says so (same doctrine as §6.1's
    # single-gateway abstention and §6.10's syslog "no evidence either way").
    _NOBS = "[NOT OBSERVED]"
    def _link_v(value):
        return value if link_centrality else _NOBS

    table(["Metric", "Value"], [
        ("Devices in scope", n_dev),
        ("Inter-switch links (CDP/LLDP-derived; Inferred-high)", _link_v(n_links)),
        ("Bridge links (single points of fabric partition)", _link_v(len(bridges))),
        ("Evidenced endpoints (canonical assessment scale)", _ep_canon if isinstance(_ep_canon, int) else ep_total),
        ("Endpoints — access-port host MACs at snapshot (superset of canonical)", ep_total),
        ("Production VLANs (canonical assessment scale)", _vl_canon if isinstance(_vl_canon, int) else len(ep_per_vlan)),
        ("VLANs carrying endpoints (subset)", len(ep_per_vlan)),
        ("Gateways (SVIs) observed", len(gw)),
        ("Migration move groups",
         _NOBS if _phase_failed(integrity_failures, "move groups") else len(move_groups)),
        ("Health bands (Critical/Poor/Fair/Good/Excellent)",
         (_NOBS if _phase_failed(integrity_failures, "health score") else
          f"{bands.get('Critical',0)} / {bands.get('Poor',0)} / {bands.get('Fair',0)} / "
          f"{bands.get('Good',0)} / {bands.get('Excellent',0)}")),
        ("Cross-layer findings (Critical / High)",
         f"{len(crit_cl)} / {len(high_cl)}" if cross_layer else _NOBS),
        ("Punch-list items",
         _NOBS if _phase_failed(integrity_failures, "punch-list", "punchlist") else len(punchlist)),
    ], widths=[4.8, 2.0])
    # …and the same discipline on the gating SENTENCE, which is the one line a war-room reader takes
    # away. With the analyses absent it used to read "0 Critical cross-layer single-point(s) of
    # failure and 0 bridge link(s) must be resolved ... 0 of 0 move group(s) are NOT READY" — a
    # complete all-clear manufactured entirely from missing sections.
    _cl_txt = (f"{len(crit_cl)} Critical cross-layer single-point(s) of failure" if cross_layer else
               "Critical cross-layer single-point(s) of failure [NOT OBSERVED — no cross-layer "
               "analysis in this snapshot; absence of the analysis is not absence of the exposure]")
    _br_txt = (f"{len(bridges)} bridge link(s)" if link_centrality else
               "bridge link(s) [NOT OBSERVED — no link-centrality analysis in this snapshot]")
    _rd_txt = (f"{sum(1 for r in mr if r.get('readiness') == 'NOT READY')} of {len(mr)} move group(s) "
               "are NOT READY on the pre-migration checklist." if mr else
               "Migration readiness was NOT computed for this snapshot, so no move group can be "
               "called ready — treat every group as no-go until the pre-migration checklist is run.")
    p = doc.add_paragraph()
    _label_run(p, "Top gating decisions:",
               f"{_cl_txt} and {_br_txt} must be resolved or risk-accepted before cutover. {_rd_txt}")

    # cross-axis migration brief (NEW-V3.23.120)
    eb = snap_dict.get("executive_brief") or {}
    if eb.get("axes"):
        _label_run(doc.add_paragraph(), "Migration brief:", eb.get("posture_statement", ""))
        tg = _as_list(eb.get("top_gating"))
        if tg:
            # '(+N more)': a bare [:6] join presented a 9-item gating list as the whole of "address first".
            _label_run(doc.add_paragraph(), "Address first:", _join_cap(tg, 6))
        table(["Axis", "Severity", "Headline"],
              [[a.get("axis"), a.get("severity"), a.get("headline")] for a in _R(eb.get("axes"))],
              widths=[1.8, 0.9, 4.0])

    # ===== 2. Scope & Data Completeness =====
    doc.add_heading("2. Scope & Data Completeness", level=1)
    # coverage-honesty: n_dev is the INVENTORY (303), not the collected count (253). Do not call the
    # inventory total "the collected dataset" — state how many of the inventoried switches were collected.
    _n_coll = _as_dict(_as_dict(snap_dict.get("collection_completeness")).get("summary")).get("complete")
    _scope = (f"{_n_coll} of {n_dev} inventoried Cisco switches were collected"
              if isinstance(_n_coll, int) and _n_coll != n_dev
              else f"the {n_dev} inventoried Cisco switches")
    doc.add_paragraph(
        f"In scope: {_scope} and their configuration, interface, CDP/LLDP, STP, FHRP/SVI, "
        "and routing-adjacency state.")
    doc.add_paragraph("Dataset limits (these bound every claim below):", style="List Bullet")
    _mc_collected = _as_dict(_as_dict(snap_dict.get("service_map")).get("multicast")).get("group_level_collected")
    _mc_limit = ("Multicast / PTP group membership IS collected; QoS baselines are not — and PTP lock state "
                 "is read from `show ptp clock`, not observed over time."
                 if _mc_collected else
                 "Multicast / PTP / QoS baselines are not collected — broadcast-media timing health is Unknown.")
    for t in (
        "No live forwarding/ARP-aging telemetry — endpoint presence is a snapshot, not a census over time.",
        "Topology is CDP/LLDP-derived: links to off-scan or non-advertising devices are Unknown.",
        _mc_limit,
        "Active FHRP ownership and failover behaviour are inferred from config, not observed over a failover.",
    ):
        doc.add_paragraph(t, style="List Bullet 2")

    # Collection completeness / blind spots (NEW-V3.23.109): make the un-/partially-collected inventory
    # devices explicit -- every finding is only as trustworthy as the data behind it.
    cc = _as_dict(snap_dict.get("collection_completeness"))
    cs = _as_dict(cc.get("summary"))
    blind = _R(cc.get("devices"))
    if cs.get("inventory"):
        doc.add_heading("2.1 Collection completeness (assessment blind spots)", level=2)
        doc.add_paragraph(
            f"Of {cs.get('inventory', 0)} inventory device(s): {cs.get('complete', 0)} fully collected, "
            f"{cs.get('partial', 0)} partial, {cs.get('not_collected', 0)} NOT collected. The not-/partially-"
            f"collected devices below are assessment blind spots ({_CONF_UNKNOWN} state) — re-collect them "
            "before relying on the findings for those devices; a missing device is not a healthy device.")
        if blind:
            table(["Status", "Device", "Data quality", "Missing essential commands"],
                  [[d.get("status", ""), d.get("host", ""), f"{d.get('data_quality', 0)}%",
                    # `.get("missing", [])` defaulted only on ABSENT: a present null/scalar reached
                    # join() (TypeError), and a list of non-str reached it too. Coerce container AND
                    # elements — the §6.9 golden-drift twin below reads the same field the same way.
                    ", ".join(str(x) for x in _as_list(d.get("missing")))] for d in blind[:40]],
                  widths=[1.3, 2.8, 1.0, 3.3])
            if len(blind) > 40:
                doc.add_paragraph(f"…and {len(blind) - 40} more — see the 'Collection Completeness' sheet.")

    # ===== 3. Scenario Classification =====
    doc.add_heading("3. Scenario Classification", level=1)
    doc.add_paragraph(
        "Primary mode: Migration planning (brownfield). Secondary: steady-state health assessment. "
        "The plan rests on an endpoint-move architecture: access-layer switches and their endpoints "
        "move in dependency-scoped groups while L3 gateways remain on the legacy core. That premise is "
        "validated per VLAN in §6; VLANs whose gateway cannot be evidenced are gated, not assumed.")
    scen = _as_dict(snap_dict.get("migration_scenarios"))
    if scen.get("fleet_recommendation"):
        p = doc.add_paragraph()
        _label_run(p, "Fleet recommendation:", scen["fleet_recommendation"])
    pg = _R(scen.get("per_group"))
    if pg:
        doc.add_paragraph("Recommended cutover scenario per move-group (the group's shape decides; the "
                          "war-room confirms). Playbooks per scenario are in §11:")
        table(["Move group", "Switches", "Dual-homed %", "Hard cutovers", "Scenario", "Why"],
              [[g.get("group") or "(group)", g.get("switches"), g.get("dual_homed_pct"),
                g.get("hard_cutover"), g.get("recommended_scenario"), g.get("rationale")]
               for g in pg[:12]], widths=[1.6, 0.9, 1.1, 1.1, 1.4, 3.4])
        _disclose(doc, len(pg), 12, "move-group scenario row(s)", "Migration Scenarios",
                  "Every group needs a chosen scenario before its wave is scheduled — an unlisted "
                  "group is unplanned, not scenario-free.")

    # ===== 4. Platform Intelligence =====
    doc.add_heading("4. Platform Intelligence", level=1)
    model_mix = Counter(d.get("model", "") or "unknown" for d in devices.values())
    doc.add_paragraph("Hardware/software mix (device-reported; cross-check labels against behaviour "
                      "before relying on them):")
    table(["Model / PID", "Count"],
          sorted(([m or "unknown", c] for m, c in model_mix.items()), key=lambda r: -r[1])[:15],
          widths=[4.5, 1.5])
    _disclose(doc, len(model_mix), 15, "model/PID row(s)", "Switch Inventory",
              "The hidden models are the long tail the hardware refresh and spares plan still has to cover.")
    # EoL / environmental flags surfaced from capacity + device fields
    env_flags = [(d.get("hostname", h), d.get("ps_status", ""), d.get("fan_status", ""),
                  d.get("temperature_status", ""))
                 for h, d in devices.items()
                 if any(str(d.get(k) or "").lower() not in ("", "ok", "normal", "good")
                        for k in ("ps_status", "fan_status", "temperature_status"))]
    if env_flags:
        doc.add_paragraph("Environmental exceptions (PSU/fan/temperature not nominal):")
        table(["Device", "PSU", "Fan", "Temp"], env_flags[:20], widths=[3.5, 1.3, 1.3, 1.3])

    # 4.1 Hardware lifecycle (EoL / End-of-Support) — NEW-V3.23.117
    lr = _as_dict(snap_dict.get("lifecycle_risk"))
    ls = _as_dict(lr.get("summary"))
    if ls.get("n_devices"):
        doc.add_heading("4.1 Hardware lifecycle (EoL / End-of-Support)", level=2)
        doc.add_paragraph(
            f"As of the assessment date {lr.get('asof', '')} (when the evidence was collected), of {ls.get('n_devices', 0)} device(s): "
            f"{ls.get('n_past_ldos', 0)} are PAST Cisco's last day of support (no software fixes / no TAC), "
            f"{ls.get('n_near', 0)} reach end-of-support within a year, "
            f"{ls.get('n_past_eos', 0)} are past end-of-sale with LDoS still future (date band only; "
            f"support entitlement not inferred), {ls.get('n_active', 0)} are in the "
            f"pre-EoS date band (schema: Active; not a support-entitlement claim), and "
            f"{ls.get('n_unknown', 0)} device(s) are NOT ASSESSED. End-of-support hardware is a hard migration driver.")
        if ls.get("n_unknown"):
            doc.add_paragraph(
                "NOT ASSESSED means either no exact EoX row matched the collected PID or the matched "
                "row's retained source/date authority was withheld or incomplete. Lifecycle position "
                "and support entitlement remain undetermined until both are verified.")
        _label_run(doc.add_paragraph(), "Reference note:", lr.get("note", ""), GREY)
        rows = [[p.get("platform"), p.get("count"), p.get("band"), p.get("ldos") or "—"]
                for p in _R(ls.get("by_platform"))]
        if rows:
            table(["Platform", "Devices", "Lifecycle band", "LDoS"], rows, widths=[2.6, 1.0, 1.7, 1.4])

    # ===== 5. Topology & Current State =====
    doc.add_heading("5. Topology & Current State", level=1)
    p = doc.add_paragraph()
    _label_run(p, "Confidence:",
               (f"{_CONF_HIGH}. {n_links} inter-switch links reconstructed from CDP/LLDP. "
                f"{len(bridges)} are bridges (their loss partitions the fabric). Off-scan links are "
                f"{_CONF_UNKNOWN}."
                if link_centrality else
                # an Inferred-high confidence label asserted over ZERO reconstructed links claimed a
                # structural conclusion the snapshot never carried
                f"{_CONF_UNKNOWN}. No link-centrality analysis is present in this snapshot, so the "
                "inter-switch link and bridge (fabric-partition) counts are [NOT OBSERVED] — not "
                "zero. Re-run the assessment with the current engine before treating this fabric as "
                "having no chokepoints."))
    if bridges:
        doc.add_paragraph("Chokepoint links (ranked by structural blast radius; live fault state is "
                          "Unknown until validated):")
        table(["Rank", "Link", "Betweenness", "Switch-pairs cut if it fails"],
              [[b.get("rank"), f"{b.get('a_host')} {b.get('a_port')} ↔ {b.get('b_host')} {b.get('b_port')}",
                b.get("betweenness"), b.get("pairs_cut")] for b in
               sorted(bridges, key=lambda r: -_as_num(r.get("pairs_cut")))[:12]],
              widths=[0.6, 4.4, 1.2, 1.6])
        _disclose(doc, len(bridges), 12, "bridge link(s)", "Link Centrality",
                  f"§1's metric register and the confidence line above both count all {len(bridges)} — "
                  "each is a single point of fabric partition, so the unlisted ones gate cutover too.")

    # ===== 6. L1-L4 Findings (the gateway premise + evidence finding blocks) =====
    doc.add_heading("6. L1–L4 Findings", level=1)
    doc.add_heading("6.1 Gateway placement (the migration premise)", level=2)
    doc.add_paragraph(
        "Per-VLAN gateway register. Placement (an SVI is configured on the named switch) is "
        f"{_CONF_CONFIRMED}; the active forwarding owner and failover behaviour are {_CONF_UNKNOWN} "
        "without a live ARP/HSRP-state capture. VLANs flagged below carry an L3 risk.")
    gw_rows = [[g["vlan"], g["svi_ip"] or "—", g["owner"], g["fhrp"] or "none",
                (g["risk"] or "ok")] for g in
               sorted(gw, key=lambda g: (0 if g["risk"] and g["risk"] != "ok" else 1, str(g["vlan"])))[:25]]
    if gw_rows:
        table(["VLAN", "SVI IP (Confirmed)", "Owner (Confirmed)", "FHRP (Inferred-high)", "L3 risk"],
              gw_rows, widths=[0.8, 1.8, 2.6, 1.5, 1.5])
        _disclose(doc, len(gw), 25, "gateway SVI record(s)", "SVI Gateway",
                  "The register is sorted risk-first, so the flagged VLANs are the ones rendered; "
                  "the FHRP counts below are computed over ALL records, not the 25 listed.")
    # No real FHRP. The parser writes the literal "none" (truthy) when no HSRP/VRRP/GLBP exists, so a
    # bare `not fhrp.strip()` counted ZERO single-gateway gateways for an all-"none" fleet — the inverse
    # of the truth. Mirror the engine's canonical `(fhrp or "none") == "none"` no-FHRP gate.
    n_no_fhrp = sum(1 for g in gw if (g.get("fhrp", "none") or "none") == "none")
    doc.add_paragraph(
        f"{n_no_fhrp} of {len(gw)} gateways have no FHRP peer in scope (counted per gateway SVI "
        "record — a VLAN served by two switches contributes two rows). "
        "Any VLAN whose gateway is off-scan is gated as Unknown in §12.")
    # …and that total is NOT the single-gateway count. The producer (excel.write_l3_forwarding_sheet)
    # emits TWO distinct risks: 'single-gateway' when the VLAN has <=1 gateway AT ALL, and 'no-FHRP'
    # when it has two or more gateways but no FHRP protocol. Labelling the merged no-FHRP total
    # "single-gateway exposure" reported the highest-severity availability exposure many times larger
    # than archreview RES-2 and the design document's §2.4 — both of which read the risk field — while
    # this runbook claims its numbers reconcile to the workbook. Read the same owner they read, and
    # ABSTAIN when the register carries no risk classification (0 must never render as clean).
    risked = [g for g in gw if str(g.get("risk") or "").strip()]
    if risked:
        n_single_gw = sum(1 for g in risked if "single-gateway" in str(g.get("risk")))
        doc.add_paragraph(
            f"Of those, {n_single_gw} carry the SINGLE-GATEWAY risk — the VLAN has no second gateway "
            "in scope at all, the highest-severity availability exposure a VLAN can hold (the same "
            "count the Architecture Review's RES-2 and the design document's §2.4 report from this "
            "snapshot). The remainder have a redundant gateway but no FHRP protocol configured — a "
            "different, lesser condition; do not merge the two.")
    elif gw:
        doc.add_paragraph(
            "This snapshot's gateway register carries no L3-risk classification, so the "
            "single-gateway subset of the gateways above cannot be counted from this evidence — "
            "no-FHRP and single-gateway are distinct conditions and only the risk field separates "
            "them. Re-run the assessment with the current engine before treating the count above as "
            "a single-gateway exposure figure.")

    doc.add_heading("6.2 Material cross-layer findings", level=2)
    # The cap stays (8 evidence-disciplined blocks is a readable section) but it must DISCLOSE, because
    # §1's metric register states the FULL Critical/High cross-layer population two pages earlier. On the
    # 303-device production snapshot §1 reads "3 / 404" while this section rendered 3 Critical + 5 High and
    # dropped 399 High findings with no marker at all — a headline the reader cannot reconcile with the
    # list beneath it, and the reader's only cue that anything is missing. Rank, render, then reconcile.
    _CL_CAP = 8
    _cl_ranked = sorted(cross_layer, key=lambda f: _SEV_ORDER.get(f.get("severity"), 9))
    for f in _cl_ranked[:_CL_CAP]:
        hosts = ", ".join(str(_h) for _h in _as_list(f.get("hosts"))) or "—"
        finding_block(
            f.get("title", f.get("id", "Cross-layer finding")),
            severity=f.get("severity", "Medium"),
            scope=f"{hosts}  (layers {f.get('layers', '?')})",
            observed=f.get("detail", "—"),
            interpretation=f"A single condition is implicated across {f.get('layers', 'multiple')} — "
                           "structural exposure, independent of whether a fault is active now.",
            impact="Loss of the named element removes the only path/gateway for the dependent "
                   "population (see Risk Register §10 for the scoped count).",
            confidence=f"{_CONF_HIGH} for the structural exposure (config + topology agree); "
                       f"{_CONF_UNKNOWN} for live fault state.",
            unknowns="Whether redundancy exists off-scan; whether the fault is active at cutover time.",
            next_validation="Confirm the second path/gateway on the live devices before the wave that "
                            "moves this element.",
            remediation=f.get("recommendation", ""))
    if len(_cl_ranked) > _CL_CAP:
        _shown_ch = sum(1 for f in _cl_ranked[:_CL_CAP] if f.get("severity") in ("Critical", "High"))
        doc.add_paragraph(
            f"…and {len(_cl_ranked) - _CL_CAP} further cross-layer finding(s) are NOT rendered as blocks "
            f"above: the {_CL_CAP} highest-severity of {len(_cl_ranked)} are shown, covering {_shown_ch} of "
            f"the {len(crit_cl) + len(high_cl)} Critical/High finding(s) that §1's metric register counts "
            f"({len(crit_cl)} Critical / {len(high_cl)} High). The unshown findings are not lower-risk by "
            "omission — read the workbook's 'Cross-Layer Analysis' sheet for the complete list before "
            "treating this section as the full cross-layer exposure.")

    drift = _R(snap_dict.get("operational_drift"))
    if drift:
        doc.add_heading("6.3 False-health / operational drift", level=2)
        doc.add_paragraph(
            "Conditions a green control plane hides — surfaced because configured/up is not healthy. "
            "Each is a snapshot-evidenced exposure; whether it is biting right now is Unknown until "
            "validated live.")
        # Why/next-step cell: budget the two parts SEPARATELY -- the old single (detail+remediation)[:160]
        # slice let a long detail evict the entire remediation and its own tail (PR-#396 review: the
        # 270-char native-VLAN-1 detail cut mid-word and the 'next step' vanished). docx wraps, so the
        # budgets only bound pathological rows; shorten() cuts at a word boundary and shows an ellipsis.
        def _why_next(d: dict) -> str:
            why = textwrap.shorten(str(d.get("detail", "") or ""), width=400, placeholder=" …") \
                if d.get("detail") else ""
            nxt = str(d.get("remediation", "") or "").strip()
            if nxt:
                nxt = textwrap.shorten(nxt, width=160, placeholder=" …")
            return (why + (" Fix: " + nxt if nxt else "")).strip()

        table(["Severity", "Finding", "Devices", "Why it matters / next step"],
              [[d.get("severity", ""), d.get("title", ""),
                str(len(_as_list(d.get("devices")))),
                _why_next(d)]
               for d in sorted(drift, key=lambda d: _SEV_ORDER.get(d.get("severity"), 9))],
              widths=[0.9, 2.8, 0.9, 4.0])

    si = _as_dict(snap_dict.get("subnet_intelligence"))
    pdv = _R(si.get("per_device"))
    if pdv:
        doc.add_heading("6.4 Subnet & routing reachability", level=2)
        n_l3 = sum(1 for r in pdv if r.get("is_l3"))
        bgp_note = ("BGP received-prefixes were collected." if si.get("bgp_received_collected")
                    else f"BGP received-prefixes were NOT collected ({_CONF_UNKNOWN}); re-run with the "
                         "new 'show ip bgp' command to populate them.")
        doc.add_paragraph(
            f"{n_l3} L3 device(s) terminate/route subnets; {len(pdv) - n_l3} L2 access switch(es) reach "
            "their subnets transitively via a gateway SVI. Destination = subnets a device terminates "
            f"({_CONF_CONFIRMED} from connected routes/SVIs); reachable = its route table "
            f"({_CONF_CONFIRMED}-at-snapshot from 'show ip route'). {bgp_note}")
        top = sorted([r for r in pdv if r.get("is_l3")],
                     key=lambda r: -(_as_num(r.get("destination_count")) + _as_num(r.get("reachable_count"))))[:12]
        doc.add_paragraph("L3 devices — subnets terminated vs reachable:")
        table(["Switch", "Dest subnets", "Reachable", "Sources", "Default next-hop"],
              [[r.get("host"), r.get("destination_count"), r.get("reachable_count"),
                ", ".join(f"{k}:{v}" for k, v in _as_dict(r.get("reachable_sources")).items()),
                r.get("default_next_hop", "")] for r in top], widths=[3.0, 1.1, 1.1, 2.4, 1.4])
        _disclose(doc, n_l3, 12, "L3 device(s)", "Subnet Reachability",
                  f"The sentence above counts all {n_l3}; the table ranks by subnets terminated + reachable.")
        mg = _R(si.get("move_groups"))
        if mg:
            doc.add_paragraph("Per move-group source↔destination (local subnets move with the group; "
                              "remote subnets must stay reachable across the cutover):")
            table(["Move group", "Switches", "Local subnets", "Remote subnets to preserve"],
                  [[m.get("group") or "(group)", m.get("switches"), m.get("local_count"),
                    m.get("remote_count")] for m in mg[:12]], widths=[2.0, 1.0, 1.4, 2.2])
            _disclose(doc, len(mg), 12, "move-group subnet row(s)", "Subnet Reachability",
                      "A group missing here still has remote subnets that must stay reachable across "
                      "its cutover.")

    fpaths = _R(_as_dict(flow_paths).get("flows"))
    if fpaths:
        fsum = _as_dict(_as_dict(flow_paths).get("summary"))
        doc.add_heading("6.4.1 Representative end-to-end flow paths", level=3)
        doc.add_paragraph(
            f"{fsum.get('n_flows', 0)} representative flow(s) — the lowest-IP endpoint per VLAN, traced "
            "L1–L3 (the static twin of the explorer's interactive Flow Simulator): "
            f"{fsum.get('n_at_risk', 0)} at HIGH/CRITICAL risk, {fsum.get('n_partitioned', 0)} partitioned. "
            "Each path is a lower bound (scan-bound topology); the explorer walks any chosen pair, all OSI "
            "layers, with the backup path. See the workbook 'Flow Paths' sheet for the full hop list.")
        table(["Flow (VLAN → VLAN)", "Type", "Risk", "SPOFs on path"],
              [[f.get("label", ""), _as_dict(f.get("summary")).get("flow_type", ""),
                _as_dict(f.get("summary")).get("risk", ""),
                "; ".join(str(_s) for _s in _as_list(_as_dict(f.get("summary")).get("spofs"))) or "none"]
               for f in fpaths], widths=[2.8, 1.4, 0.9, 2.4])

    # Runtime protocol coverage is a first-class receipt, not something inferred from the
    # sparse health/advisory arrays.  Render the section even when all three artifacts are
    # empty so a legacy or failed run cannot disappear into apparent protocol cleanliness.
    phealth = _R(snap_dict.get("protocol_health"))
    pintel = _R(snap_dict.get("protocol_intelligence"))
    passess = _as_dict(snap_dict.get("protocol_assessability"))
    pa_rows = _R(passess.get("rows"))
    pa_valid = passess.get("schema") == "protocol_assessability/1" and bool(pa_rows)
    doc.add_heading("6.5 Protocol behaviour & remediation", level=2)

    # VTP safety is a closed local-status audit receipt, distinct from sparse
    # protocol health/intelligence.  The runbook accepts only the serialized
    # embedded projection; it never promotes a persisted current-run claim.
    vtp_input = snap_dict.get("vtp_safety_baseline")
    try:
        from cisco_toolkit.vtp_safety import validate_vtp_safety_baseline

        vtp_view = validate_vtp_safety_baseline(vtp_input, require_current_run=False)
    except Exception:
        vtp_view = {}
    vtp_valid = bool(
        isinstance(vtp_view, dict)
        and vtp_view.get("valid") is True
        and isinstance(vtp_view.get("baseline"), dict)
        and vtp_view["baseline"].get("schema") == "vtp_safety_baseline/1"
        and vtp_view["baseline"].get("projection_custody") == "embedded_unverified"
        and isinstance(vtp_view.get("rows"), list)
        and isinstance(vtp_view["baseline"].get("coverage"), list)
        and isinstance(vtp_view["baseline"].get("summary"), dict)
    )

    def _vtp_text(value) -> str:
        if value is None:
            return ""
        if isinstance(value, bool):
            return "yes" if value else "no"
        return str(value) if isinstance(value, (str, int, float)) else ""

    def _vtp_findings(value) -> str:
        if not isinstance(value, list):
            return ""
        parts = []
        for item in value:
            if not isinstance(item, dict):
                continue
            part = " — ".join(filter(None, (
                _vtp_text(item.get("code")), _vtp_text(item.get("issue")),
            )))
            if part:
                parts.append(part)
        return "; ".join(parts)

    if vtp_valid:
        vtp_receipt = vtp_view["baseline"]
        vtp_rows = [row for row in vtp_view["rows"] if isinstance(row, dict)]
        vtp_coverage = [cell for cell in vtp_receipt["coverage"] if isinstance(cell, dict)]
        if len(vtp_rows) != len(vtp_view["rows"]) or len(vtp_coverage) != len(vtp_receipt["coverage"]):
            vtp_valid = False
    if vtp_valid:
        vtp_summary = vtp_receipt["summary"]
        vtp_by_status = vtp_summary["by_status"]
        vtp_by_coverage = vtp_summary["by_coverage_status"]
        vtp_coverage_by_host = {cell["switch"]: cell for cell in vtp_coverage}
        vtp_verdict = _vtp_text(vtp_receipt.get("verdict")) or "NOT ASSESSED"
        doc.add_paragraph(
            "VTP safety gate — observed local show vtp status. "
            f"Verdict: {vtp_verdict.upper()}. Across {vtp_summary['n_rows']} subject row(s), "
            f"{vtp_by_status['review']} require review, {vtp_by_status['not_verified']} are not "
            f"verified, and {vtp_by_status['assessed']} are boundedly assessed; "
            f"{vtp_summary['n_high_revision_servers']} VTP Server row(s) carry configuration revision "
            "100 or higher. That threshold is a conservative heuristic requiring explicit REVIEW, "
            "not proof that an overwrite will occur. Matching a high revision is NOT ACCEPTANCE. "
            "Receipt custody: embedded_unverified."
        )
        coverage_copy = (
            "Host coverage (producer-owned fleet census; distinct from subject rows): "
            f"{vtp_summary['n_hosts']} host(s), {vtp_summary['n_subject_hosts']} subject host(s); "
            f"coverage cells — {vtp_by_coverage['review']} review, "
            f"{vtp_by_coverage['not_verified']} not verified, {vtp_by_coverage['assessed']} assessed, "
            f"and {vtp_by_coverage['not_applicable']} not applicable. Coverage detail is used only "
            "for the corresponding host; fleet totals come from summary.by_coverage_status."
        )
        if vtp_verdict.upper() == "NOT_APPLICABLE":
            coverage_copy += (
                " NOT_APPLICABLE means no positive local VTP-status subject was identified in the "
                "bounded input; it is not proof that VTP is absent from the platform or network."
            )
        coverage_run = doc.add_paragraph().add_run(coverage_copy)
        if vtp_by_coverage["review"] or vtp_by_coverage["not_verified"] \
                or vtp_verdict.upper() == "NOT_APPLICABLE":
            coverage_run.bold = True
            coverage_run.font.color.rgb = RGBColor(0x7F, 0x60, 0x00)
        doc.add_paragraph(
            "Boundaries: this receipt observes one local status record only. It establishes no "
            "VLAN-database equality or contents, advertisement or per-VLAN propagation, intended "
            "version/compatibility, pruning, password/authentication, freshness, simultaneous state, "
            "failover, revision-reset safety, or cutover-authorization proof. Raw captures, hashes, "
            "source locators, and filesystem paths are not rendered. The protocol-health and "
            "protocol-intelligence detail below remains a separate observed/advisory projection."
        )
        vtp_blockers = [
            row for row in vtp_rows if row.get("status") in {"review", "not_verified"}
        ]
        vtp_ordinary = [row for row in vtp_rows if row.get("status") == "assessed"]
        vtp_shown = vtp_blockers + vtp_ordinary[:50]
        if vtp_shown:
            doc.add_paragraph(
                "VTP safety subjects (every REVIEW and NOT VERIFIED row, including each subject "
                "coverage gap, is rendered outside the 50-row assessed display cap):"
            )

            def _vtp_observed(row):
                domain = _vtp_text(row.get("domain")) if row.get("domain_present") is True else "not reported"
                if row.get("domain_present") is True and not domain:
                    domain = "empty"
                revision = _vtp_text(row.get("revision")) if row.get("revision_present") is True else "not reported"
                version = _vtp_text(row.get("version")) if row.get("version_present") is True else "not reported"
                return (
                    f"mode {_vtp_text(row.get('mode')) or 'unknown'} · domain {domain} · "
                    f"revision {revision} · running version {version}"
                )

            table(
                ["Status", "Device / platform", "Observed local status", "Capture / parser",
                 "Command / custody", "Acceptance / finding"],
                [[
                    _vtp_text(row.get("status")).replace("_", " ").upper(),
                    f"{_vtp_text(row.get('switch'))} · {_vtp_text(row.get('platform')) or '—'}",
                    _vtp_observed(row),
                    f"{_vtp_text(vtp_coverage_by_host.get(row.get('switch'), {}).get('capture_status')) or 'not reported'} / "
                    f"{_vtp_text(vtp_coverage_by_host.get(row.get('switch'), {}).get('parser_status')) or 'not reported'}",
                    f"{_vtp_text(row.get('command')) or 'show vtp status'} · "
                    f"{_vtp_text(row.get('projection_custody')) or 'embedded_unverified'}",
                    "\n".join(filter(None, (
                        _vtp_text(row.get("acceptance")), _vtp_findings(row.get("findings")),
                    ))) or "—",
                ] for row in vtp_shown],
                widths=[1.1, 1.8, 2.4, 1.6, 2.1, 4.0],
            )
        if len(vtp_ordinary) > 50:
            doc.add_paragraph(
                f"All {len(vtp_blockers)} blocker row(s) are shown; 50 of "
                f"{len(vtp_ordinary)} assessed row(s) are shown. The VTP Safety workbook sheet uses "
                "the same bounded projection."
            )
    if not vtp_valid:
        doc.add_paragraph(
            "VTP safety gate — Verdict: NOT ASSESSED. vtp_safety_baseline/1 is absent, malformed, "
            "or not an embedded-unverified serialized receipt. No rejected receipt leaves, counts, "
            "or VTP safety conclusion are shown. Re-collect and validate show vtp status before "
            "acceptance; zero visible rows is not proof that VTP is absent."
        )

    # IPv6 routing adjacency is a separately owned observed-runtime receipt.
    # Only a validated serialized audit copy is projected here; persisted JSON
    # cannot claim the process-local custody required by the decision owner.
    ipv6_input = snap_dict.get("ipv6_routing_adjacency_baseline")
    try:
        from cisco_toolkit.ipv6_routing import validate_ipv6_routing_adjacency_baseline

        ipv6_view = validate_ipv6_routing_adjacency_baseline(
            ipv6_input, require_current_run=False,
        )
    except Exception:
        ipv6_view = {}
    ipv6_valid = bool(
        isinstance(ipv6_view, dict)
        and ipv6_view.get("valid") is True
        and isinstance(ipv6_view.get("baseline"), dict)
        and ipv6_view["baseline"].get("schema") == "ipv6_routing_adjacency_baseline/1"
        and ipv6_view["baseline"].get("projection_custody") == "embedded_unverified"
        and isinstance(ipv6_view.get("rows"), list)
        and isinstance(ipv6_view["baseline"].get("coverage"), list)
        and isinstance(ipv6_view["baseline"].get("summary"), dict)
    )

    def _ipv6_text(value) -> str:
        if value is None:
            return ""
        if isinstance(value, bool):
            return "yes" if value else "no"
        return str(value) if isinstance(value, (str, int, float)) else ""

    def _ipv6_findings(value) -> str:
        if not isinstance(value, list):
            return ""
        parts = []
        for item in value:
            if not isinstance(item, dict):
                continue
            part = " — ".join(filter(None, (
                _ipv6_text(item.get("code")), _ipv6_text(item.get("issue")),
            )))
            if part:
                parts.append(part)
        return "; ".join(parts)

    if ipv6_valid:
        ipv6_receipt = ipv6_view["baseline"]
        ipv6_rows = [row for row in ipv6_view["rows"] if isinstance(row, dict)]
        ipv6_coverage = [
            cell for cell in ipv6_receipt["coverage"] if isinstance(cell, dict)
        ]
        if (
            len(ipv6_rows) != len(ipv6_view["rows"])
            or len(ipv6_coverage) != len(ipv6_receipt["coverage"])
        ):
            ipv6_valid = False
    if ipv6_valid:
        ipv6_summary = ipv6_receipt["summary"]
        ipv6_by_status = ipv6_summary["by_status"]
        ipv6_by_coverage = ipv6_summary["by_coverage_status"]
        ipv6_coverage_by_subject = {
            (cell.get("switch"), cell.get("protocol")): cell
            for cell in ipv6_coverage
        }
        ipv6_verdict = _ipv6_text(ipv6_receipt.get("verdict")) or "NOT ASSESSED"
        doc.add_paragraph(
            "IPv6 routing adjacency gate — observed default/global OSPFv3 and IPv6-unicast BGP. "
            f"Verdict: {ipv6_verdict.upper()}. Across {ipv6_summary['n_rows']} observed adjacency "
            f"row(s), {ipv6_by_status['degraded']} are degraded, {ipv6_by_status['review']} "
            f"require review, {ipv6_by_status['not_verified']} are not verified, and "
            f"{ipv6_by_status['assessed']} are boundedly assessed; "
            f"{ipv6_summary['n_ospfv3_rows']} OSPFv3 row(s) and "
            f"{ipv6_summary['n_bgpv6_rows']} BGPv6 row(s) were published. Every DEGRADED, "
            "REVIEW, and NOT VERIFIED row is a pre-cutover blocker; matching a degraded state "
            "is NOT ACCEPTANCE. Receipt custody: embedded_unverified."
        )
        coverage_copy = (
            "Fleet coverage (producer-owned census; distinct from adjacency rows): "
            f"{ipv6_summary['n_hosts']} host(s), {ipv6_summary['n_subject_hosts']} subject host(s); "
            "host-family/input cells — "
            f"{ipv6_by_coverage['degraded']} degraded, {ipv6_by_coverage['review']} review, "
            f"{ipv6_by_coverage['not_verified']} not verified, "
            f"{ipv6_by_coverage['assessed']} assessed, and "
            f"{ipv6_by_coverage['not_applicable']} not applicable. Coverage detail qualifies only "
            "the corresponding host and input; fleet totals come from summary.by_coverage_status."
        )
        if ipv6_verdict.upper() == "NOT_APPLICABLE":
            coverage_copy += (
                " NOT_APPLICABLE means no positive bounded subject was identified; it is not proof "
                "that IPv6 routing or an expected adjacency is absent."
            )
        coverage_run = doc.add_paragraph().add_run(coverage_copy)
        if (
            ipv6_by_coverage["degraded"]
            or ipv6_by_coverage["review"]
            or ipv6_by_coverage["not_verified"]
            or ipv6_verdict.upper() == "NOT_APPLICABLE"
        ):
            coverage_run.bold = True
            coverage_run.font.color.rgb = RGBColor(0x7F, 0x60, 0x00)
        doc.add_paragraph(
            "Boundaries: observed default/global runtime peers only; no configured or expected-peer "
            "denominator. Empty or NOT_APPLICABLE evidence is not proof of absence. OSPFv3 process, "
            "area, network type, timers, authentication, LSDB correctness, persistence, and convergence "
            "history are incomplete or outside this receipt. BGPv6 policy, activation intent, route "
            "correctness, other VRFs/address families, and configured-but-unobserved peers remain outside "
            "it; RIB/FIB/path selection, freshness, simultaneity, interoperability, and cutover "
            "authorization are unproved. IPv6 route-summary counts are point-in-time census context; "
            "BGP prefix counts are informational and are not pinned acceptance targets. Raw captures, hashes, "
            "source locators, and filesystem paths are not rendered. The before/after "
            "protocol_adjacency_delta/1 comparator remains bounded to IPv4 OSPF, BGP, and EIGRP; "
            "this current-state IPv6 receipt still reaches the shared current-baseline gate."
        )
        ipv6_blockers = [
            row for row in ipv6_rows
            if row.get("status") in {"degraded", "review", "not_verified"}
        ]
        ipv6_ordinary = [row for row in ipv6_rows if row.get("status") == "assessed"]
        ipv6_shown = ipv6_blockers + ipv6_ordinary[:50]
        if ipv6_shown:
            doc.add_paragraph(
                "IPv6 routing adjacency subjects (every DEGRADED, REVIEW, and NOT VERIFIED row is "
                "rendered outside the 50-row assessed display cap):"
            )

            def _ipv6_observed(row):
                protocol = _ipv6_text(row.get("protocol")) or "unknown protocol"
                peer = _ipv6_text(row.get("peer")) or "peer not reported"
                raw_state = _ipv6_text(row.get("state_raw")) or "state not reported"
                normalized = _ipv6_text(row.get("state")) or "not reported"
                interface = _ipv6_text(row.get("interface")) or "—"
                process = _ipv6_text(row.get("process")) or "—"
                remote_as = _ipv6_text(row.get("remote_as")) or "—"
                role = _ipv6_text(row.get("role")) or "—"
                prefixes = (
                    _ipv6_text(row.get("prefix_count"))
                    if row.get("prefix_count_present") is True else "not reported"
                )
                return (
                    f"{protocol} peer {peer} · exact state {raw_state} · normalized {normalized} · "
                    f"process {process} · interface {interface} · remote AS {remote_as} · "
                    f"role {role} · prefixes {prefixes}"
                )

            table(
                ["Status", "Device / platform", "Observed peer / state", "Capture / parser",
                 "Command / custody", "Acceptance / finding"],
                [[
                    _ipv6_text(row.get("status")).replace("_", " ").upper(),
                    f"{_ipv6_text(row.get('switch'))} · {_ipv6_text(row.get('platform')) or '—'}",
                    _ipv6_observed(row),
                    f"{_ipv6_text(ipv6_coverage_by_subject.get((row.get('switch'), row.get('protocol')), {}).get('capture_status')) or 'not reported'} / "
                    f"{_ipv6_text(ipv6_coverage_by_subject.get((row.get('switch'), row.get('protocol')), {}).get('parser_status')) or 'not reported'}",
                    f"{_ipv6_text(row.get('command'))} · "
                    f"{_ipv6_text(row.get('projection_custody')) or 'embedded_unverified'}",
                    "\n".join(filter(None, (
                        _ipv6_text(row.get("acceptance")),
                        _ipv6_findings(row.get("findings")),
                    ))) or "—",
                ] for row in ipv6_shown],
                widths=[1.1, 1.8, 3.5, 1.5, 2.2, 4.0],
            )
        if len(ipv6_ordinary) > 50:
            doc.add_paragraph(
                f"All {len(ipv6_blockers)} blocker row(s) are shown; 50 of "
                f"{len(ipv6_ordinary)} assessed row(s) are shown. The IPv6 Routing workbook sheet "
                "and Explorer use the same bounded projection."
            )
    if not ipv6_valid:
        doc.add_paragraph(
            "IPv6 routing adjacency gate — Verdict: NOT ASSESSED. "
            "ipv6_routing_adjacency_baseline/1 is absent, malformed, or not an "
            "embedded-unverified serialized receipt. No rejected receipt leaves, counts, or IPv6 "
            "routing conclusion are shown. Re-collect and validate IPv6 route-summary, OSPFv3 "
            "neighbor, and BGP IPv6-unicast summary evidence before acceptance; zero visible rows "
            "is not proof that IPv6 routing or expected peers are absent."
        )

    # The configured-peer receipt is a distinct denominator from sparse observed protocol
    # health.  Render it even when absent so an older/failed snapshot cannot read as zero
    # configured peers.  Only typed receipt leaves are projected; raw configuration/debug
    # objects never enter this document.
    bgp_baseline = _as_dict(snap_dict.get("bgp_configured_peer_baseline"))

    bgp_coverage_status_order = (
        "degraded", "review", "not_verified", "assessed", "not_applicable",
    )

    def _bgp_coverage_census(receipt):
        """Validate the producer-owned fleet census without using rows as its display owner."""

        summary = _as_dict(receipt.get("summary"))
        by_coverage = summary.get("by_coverage_status")
        coverage = receipt.get("coverage")
        if (
            not isinstance(by_coverage, dict)
            or set(by_coverage) != set(bgp_coverage_status_order)
            or not isinstance(coverage, list)
        ):
            return None
        counts = {}
        for status in bgp_coverage_status_order:
            value = by_coverage.get(status)
            if not isinstance(value, int) or isinstance(value, bool) or value < 0:
                return None
            counts[status] = value
        n_hosts = summary.get("n_hosts")
        n_subject_hosts = summary.get("n_subject_hosts")
        if any(
            not isinstance(value, int) or isinstance(value, bool) or value < 0
            for value in (n_hosts, n_subject_hosts)
        ):
            return None
        if n_subject_hosts > n_hosts or n_hosts != len(coverage) or sum(counts.values()) != n_hosts:
            return None
        row_counts = {status: 0 for status in bgp_coverage_status_order}
        seen_switches = set()
        subject_hosts = 0
        for row in coverage:
            if not isinstance(row, dict):
                return None
            switch, subject, status = row.get("switch"), row.get("subject"), row.get("status")
            if (
                not isinstance(switch, str) or not switch
                or switch in seen_switches
                or not isinstance(subject, bool)
                or status not in row_counts
                or (subject and status == "not_applicable")
            ):
                return None
            seen_switches.add(switch)
            subject_hosts += int(subject)
            row_counts[status] += 1
        if subject_hosts != n_subject_hosts or row_counts != counts:
            return None
        return {
            "n_hosts": n_hosts,
            "n_subject_hosts": n_subject_hosts,
            "by_coverage_status": counts,
        }

    bgp_coverage_census = _bgp_coverage_census(bgp_baseline)
    bgp_valid = (
        bgp_baseline.get("schema") == "bgp_configured_peer_baseline/1"
        and isinstance(bgp_baseline.get("rows"), list)
        and bgp_coverage_census is not None
    )
    bgp_rows = _R(bgp_baseline.get("rows")) if bgp_valid else []

    def _bgp_text(value) -> str:
        if value is None:
            return ""
        if isinstance(value, bool):
            text = "yes" if value else "no"
        elif isinstance(value, (str, int, float)):
            text = str(value)
        else:
            return ""
        return re.sub(
            r"(?i)(\b(?:password|secret|community)\b(?:\s+\d+)?\s+)([^\s,;]+)",
            r"\1[REDACTED]", text,
        )

    def _bgp_findings(value) -> str:
        if isinstance(value, list):
            return "; ".join(filter(None, (_bgp_findings(item) for item in value)))
        if isinstance(value, dict):
            return " — ".join(filter(None, (
                _bgp_text(value.get(key))
                for key in ("code", "title", "message", "finding", "detail", "issue")
            )))
        return _bgp_text(value)

    if bgp_valid:
        bgp_active = [
            row for row in bgp_rows
            if str(row.get("activation") or "").strip().lower() == "active"
        ]
        bgp_established = sum(
            str(row.get("status") or "").strip().lower() == "assessed" for row in bgp_rows
        )
        bgp_configured = sum(bool(str(row.get("configured_remote_as") or "").strip()) for row in bgp_rows)
        bgp_by_status = Counter(str(row.get("status") or "").strip().lower() for row in bgp_rows)
        bgp_degraded = bgp_by_status["degraded"]
        bgp_review = bgp_by_status["review"]
        bgp_unverified = bgp_by_status["not_verified"]
        bgp_verdict = _bgp_text(bgp_baseline.get("verdict")) or "NOT ASSESSED"
        bgp_custody = _bgp_text(bgp_baseline.get("projection_custody")) or "not verified"
        bgp_summary = _as_dict(bgp_baseline.get("summary"))

        def _bgp_summary_count(key, fallback):
            value = bgp_summary.get(key)
            return value if isinstance(value, int) and not isinstance(value, bool) and value >= 0 else fallback

        bgp_configured = _bgp_summary_count("n_configured_peers", bgp_configured)
        bgp_active_count = _bgp_summary_count("n_active_peers", len(bgp_active))
        bgp_established = _bgp_summary_count("n_established", bgp_established)
        bgp_degraded = _bgp_summary_count("n_degraded", bgp_degraded)
        bgp_review = _bgp_summary_count("n_review", bgp_review)
        bgp_unverified = _bgp_summary_count("n_not_verified", bgp_unverified)
        doc.add_paragraph(
            "Configured BGP peer gate — default/global IPv4 unicast. "
            f"Verdict: {bgp_verdict.upper()}. Across {bgp_configured} configured literal peers, "
            f"{bgp_established} of {bgp_active_count} configured-active "
            f"literal peers are Established. Peer rows: {bgp_degraded} degraded, "
            f"{bgp_review} require review, and "
            f"{bgp_unverified} are not verified. Receipt custody: {bgp_custody}. A configured-active "
            "peer absent from a usable summary is a blocker. This does not validate VRFs, IPv6, "
            "VPNv4/EVPN, peer groups/templates, dynamic peers, policy, routes, best path, RPKI, "
            "convergence, freshness, interoperability, or cutover authorization. CLEAR applies only "
            "to this bounded configured-peer denominator."
        )
        coverage_counts = bgp_coverage_census["by_coverage_status"]
        coverage_copy = (
            "Host coverage (distinct from peer rows): "
            f"{bgp_coverage_census['n_hosts']} host(s) · "
            f"{bgp_coverage_census['n_subject_hosts']} subject host(s) · coverage cells — "
            f"{coverage_counts['degraded']} degraded, {coverage_counts['review']} review, "
            f"{coverage_counts['not_verified']} not verified, "
            f"{coverage_counts['assessed']} assessed, and "
            f"{coverage_counts['not_applicable']} not applicable."
        )
        if bgp_verdict.upper() == "NOT_APPLICABLE":
            coverage_copy += (
                " NOT_APPLICABLE means no in-scope literal peer subject was identified; it is not "
                "proof that BGP is absent or that configuration coverage is complete."
            )
        coverage_qualified = bool(
            coverage_counts["degraded"]
            or coverage_counts["review"]
            or coverage_counts["not_verified"]
            or bgp_verdict.upper() == "NOT_APPLICABLE"
        )
        coverage_paragraph = doc.add_paragraph()
        coverage_run = coverage_paragraph.add_run(("⚠ " if coverage_qualified else "") + coverage_copy)
        coverage_run.bold = coverage_qualified
        if coverage_qualified:
            coverage_run.font.color.rgb = RGBColor(0x7F, 0x60, 0x00)
        bgp_blockers = [
            row for row in bgp_rows
            if str(row.get("status") or "").strip().lower() in {"degraded", "review", "not_verified"}
        ]
        if bgp_blockers:
            doc.add_paragraph(
                "Configured BGP peer blockers (all rendered; these rows are outside the ordinary "
                "protocol display cap):"
            )
            table(
                ["Status", "Device", "Peer", "Configured / runtime AS", "Runtime state",
                 "Custody / source", "Acceptance / finding"],
                [[
                    _bgp_text(row.get("status")).replace("_", " ").upper(),
                    _bgp_text(row.get("switch")),
                    _bgp_text(row.get("peer")),
                    f"{_bgp_text(row.get('configured_remote_as')) or '—'} / "
                    f"{_bgp_text(row.get('runtime_remote_as')) or '—'}",
                    _bgp_text(row.get("runtime_state")) or _bgp_text(row.get("runtime_state_raw")) or "—",
                    f"{_bgp_text(row.get('projection_custody')) or bgp_custody} · "
                    f"{_bgp_text(row.get('source_key')) or '—'}",
                    "\n".join(filter(None, (
                        _bgp_text(row.get("acceptance")), _bgp_findings(row.get("findings"))
                    ))) or "—",
                ] for row in bgp_blockers],
                widths=[1.2, 1.4, 1.4, 1.5, 1.4, 2.4, 3.3],
            )
    else:
        doc.add_paragraph(
            "Configured BGP peer gate — default/global IPv4 unicast. Verdict: NOT ASSESSED. "
            "Established/configured-active/degraded/review/not-verified counts are unavailable because "
            "bgp_configured_peer_baseline/1 is absent or malformed. No configured-peer completeness or "
            "health conclusion is asserted. The bounded gate excludes VRFs, IPv6, VPNv4/EVPN, peer "
            "groups/templates, dynamic peers, policy, routes, best path, RPKI, convergence, freshness, "
            "interoperability, and cutover authorization."
        )

    # Configured FHRP groups are a second denominator, distinct from the sparse runtime-health
    # rows and from the configured BGP peer gate above. Coverage is host × subtype (HSRP, VRRP,
    # GLBP), while group rows count concrete local group subjects; never alias those two units.
    fhrp_baseline = _as_dict(snap_dict.get("fhrp_configured_group_baseline"))
    fhrp_group_states = (
        "degraded", "review", "not_verified", "assessed", "administratively_disabled",
    )
    fhrp_coverage_states = (
        "degraded", "review", "not_verified", "assessed", "not_applicable",
    )
    fhrp_protocols = ("HSRP", "VRRP", "GLBP")

    def _fhrp_counts(raw, vocabulary):
        if not isinstance(raw, dict) or list(raw) != list(vocabulary):
            return None
        out = {}
        for status in vocabulary:
            count = raw.get(status)
            if not isinstance(count, int) or isinstance(count, bool) or count < 0:
                return None
            out[status] = count
        return out

    def _fhrp_view(receipt):
        try:
            from cisco_toolkit.fhrp_intent import validate_fhrp_configured_group_baseline

            validated = validate_fhrp_configured_group_baseline(receipt)
        except Exception:
            return None
        if not isinstance(validated, dict) or validated.get("valid") is not True:
            return None
        receipt = _as_dict(validated.get("baseline"))
        summary = _as_dict(receipt.get("summary"))
        rows, coverage = receipt.get("rows"), receipt.get("coverage")
        if (
            receipt.get("schema") != "fhrp_configured_group_baseline/1"
            or _as_dict(receipt.get("scope")) != {
                "routing_instance": "default", "afi": "ipv4", "group_kind": "direct_literal_local",
            }
            or not isinstance(rows, list) or not isinstance(coverage, list)
        ):
            return None
        by_status = _fhrp_counts(summary.get("by_status"), fhrp_group_states)
        by_coverage = _fhrp_counts(summary.get("by_coverage_status"), fhrp_coverage_states)
        if by_status is None or by_coverage is None:
            return None
        safe_rows = [row for row in rows if isinstance(row, dict)]
        safe_coverage = [row for row in coverage if isinstance(row, dict)]
        if len(safe_rows) != len(rows) or len(safe_coverage) != len(coverage):
            return None
        row_counts = {status: 0 for status in fhrp_group_states}
        identities = set()
        for row in safe_rows:
            identity = (row.get("switch"), row.get("protocol"), row.get("interface"), row.get("group"))
            status = row.get("status")
            if (
                status not in row_counts or identity in identities
                or not all(isinstance(part, str) and part for part in identity)
                or row.get("protocol") not in fhrp_protocols
            ):
                return None
            identities.add(identity)
            row_counts[status] += 1
        if row_counts != by_status:
            return None
        row_summary_expected = {
            "n_configured_groups": sum(row.get("configured") is True for row in safe_rows),
            "n_active_groups": sum(row.get("activation") == "active" for row in safe_rows),
            "n_runtime_groups": sum(row.get("runtime_observed") is True for row in safe_rows),
            "n_assessed": row_counts["assessed"],
            "n_degraded": row_counts["degraded"],
            "n_review": row_counts["review"],
            "n_not_verified": row_counts["not_verified"],
            "n_disabled": row_counts["administratively_disabled"],
        }
        if any(
            not isinstance(summary.get(key), int)
            or isinstance(summary.get(key), bool)
            or summary.get(key) != value
            for key, value in row_summary_expected.items()
        ):
            return None
        coverage_counts = {status: 0 for status in fhrp_coverage_states}
        keys, hosts, subject_hosts, subject_cells = set(), set(), set(), 0
        for row in safe_coverage:
            switch, protocol = row.get("switch"), row.get("protocol")
            subject, status = row.get("subject"), row.get("status")
            key = (switch, protocol)
            if (
                not isinstance(switch, str) or not switch or protocol not in fhrp_protocols
                or key in keys or not isinstance(subject, bool) or status not in coverage_counts
                or (subject and status == "not_applicable")
            ):
                return None
            keys.add(key)
            hosts.add(switch)
            coverage_counts[status] += 1
            if subject:
                subject_cells += 1
                subject_hosts.add(switch)
        if coverage_counts != by_coverage:
            return None
        expected = {
            "n_hosts": len(hosts), "n_coverage_cells": len(safe_coverage),
            "n_subject_hosts": len(subject_hosts), "n_subject_cells": subject_cells,
        }
        if any(
            not isinstance(summary.get(key), int) or isinstance(summary.get(key), bool)
            or summary.get(key) != value for key, value in expected.items()
        ) or (safe_coverage and len(safe_coverage) != len(hosts) * len(fhrp_protocols)):
            return None
        return {
            "rows": safe_rows, "summary": summary, "by_status": by_status,
            "by_coverage_status": by_coverage,
        }

    def _fhrp_text(value):
        if value is None:
            return ""
        if isinstance(value, bool):
            text = "yes" if value else "no"
        elif isinstance(value, (str, int, float)):
            text = str(value)
        else:
            return ""
        return re.sub(
            r"(?i)(\b(?:password|secret|community|authentication(?:\s+(?:text|md5))?|key-string)\b"
            r"(?:\s+\d+)?\s+)([^\s,;]+)", r"\1[REDACTED]", text,
        )

    def _fhrp_findings(value):
        if isinstance(value, list):
            return "; ".join(filter(None, (_fhrp_findings(item) for item in value)))
        if isinstance(value, dict):
            return " — ".join(filter(None, (
                _fhrp_text(value.get(key))
                for key in ("code", "title", "message", "finding", "detail", "issue")
            )))
        return _fhrp_text(value)

    fhrp_view = _fhrp_view(fhrp_baseline)
    if fhrp_view:
        fhrp_rows, fhrp_summary = fhrp_view["rows"], fhrp_view["summary"]
        group_counts, coverage_counts = fhrp_view["by_status"], fhrp_view["by_coverage_status"]
        fhrp_verdict = _fhrp_text(fhrp_baseline.get("verdict")) or "NOT ASSESSED"
        fhrp_custody = _fhrp_text(fhrp_baseline.get("projection_custody")) or "not verified"
        configured = fhrp_summary.get("n_configured_groups", sum(row.get("configured") is True for row in fhrp_rows))
        active = fhrp_summary.get("n_active_groups", sum(row.get("activation") == "active" for row in fhrp_rows))
        runtime = fhrp_summary.get("n_runtime_groups", sum(row.get("runtime_observed") is True for row in fhrp_rows))
        doc.add_paragraph(
            "Configured FHRP group gate — default/global IPv4 direct literal local groups. "
            f"Verdict: {fhrp_verdict.upper()}. Group rows: {configured} configured, {active} configured-active, "
            f"{runtime} runtime-observed, {group_counts['assessed']} assessed, {group_counts['degraded']} degraded, "
            f"{group_counts['review']} review, {group_counts['not_verified']} not verified, and "
            f"{group_counts['administratively_disabled']} administratively disabled. Receipt custody: "
            f"{fhrp_custody}. A configured-active local group absent from usable subtype runtime evidence is "
            "a blocker. CLEAR applies only to this bounded denominator. Excludes VRFs, IPv6, templates, "
            "inheritance, dynamic constructs, secondary VIPs, expected member count, timers, authentication, "
            "preemption, tracking behavior, simultaneous election, failover, convergence, freshness, "
            "interoperability, and cutover authorization; NX-OS configured-group parsing is limited to "
            "nested HSRP."
        )
        coverage_copy = (
            "Host × subtype coverage (distinct from group rows): "
            f"{fhrp_summary['n_hosts']} host(s) · {fhrp_summary['n_coverage_cells']} HSRP/VRRP/GLBP cell(s) · "
            f"{fhrp_summary['n_subject_hosts']} subject host(s) · {fhrp_summary['n_subject_cells']} subject cell(s) · "
            f"coverage cells — {coverage_counts['degraded']} degraded, {coverage_counts['review']} review, "
            f"{coverage_counts['not_verified']} not verified, {coverage_counts['assessed']} assessed, and "
            f"{coverage_counts['not_applicable']} not applicable."
        )
        if fhrp_verdict.upper() == "NOT_APPLICABLE":
            coverage_copy += (
                " NOT_APPLICABLE means no in-scope literal local group subject was identified; it is not "
                "proof that FHRP is absent or that configuration coverage is complete."
            )
        qualified = bool(
            coverage_counts["degraded"] or coverage_counts["review"]
            or coverage_counts["not_verified"] or fhrp_verdict.upper() == "NOT_APPLICABLE"
        )
        run = doc.add_paragraph().add_run(("⚠ " if qualified else "") + coverage_copy)
        run.bold = qualified
        if qualified:
            run.font.color.rgb = RGBColor(0x7F, 0x60, 0x00)
        blockers = [
            row for row in fhrp_rows
            if str(row.get("status") or "").strip().lower() in {"degraded", "review", "not_verified"}
        ]
        if blockers:
            doc.add_paragraph(
                "Configured FHRP group blockers (all rendered; these rows are outside the ordinary "
                "protocol display cap):"
            )
            table(
                ["Status", "Device", "Subtype / local identity", "Configured / runtime VIP",
                 "Runtime state", "Custody / source", "Acceptance / finding"],
                [[
                    _fhrp_text(row.get("status")).replace("_", " ").upper(),
                    _fhrp_text(row.get("switch")),
                    f"{_fhrp_text(row.get('protocol'))} {_fhrp_text(row.get('interface'))} "
                    f"group {_fhrp_text(row.get('group'))}",
                    f"{_fhrp_text(row.get('configured_vip')) or '—'} / "
                    f"{_fhrp_text(row.get('runtime_vip')) or '—'}",
                    _fhrp_text(row.get("runtime_state"))
                    or _fhrp_text(row.get("runtime_state_raw")) or "—",
                    f"{_fhrp_text(row.get('projection_custody')) or fhrp_custody} · "
                    f"{_fhrp_text(row.get('source_key')) or '—'}",
                    "\n".join(filter(None, (
                        _fhrp_text(row.get("acceptance")), _fhrp_findings(row.get("findings"))
                    ))) or "—",
                ] for row in blockers],
                widths=[1.0, 1.2, 2.0, 1.6, 1.4, 2.3, 3.2],
            )
    else:
        doc.add_paragraph(
            "Configured FHRP group gate — default/global IPv4 direct literal local groups. Verdict: NOT "
            "ASSESSED. Group-row and host × subtype coverage counts are unavailable because "
            "fhrp_configured_group_baseline/1 is absent or malformed. No configured-group completeness or "
            "health conclusion is asserted. Absence of rows is not proof that FHRP is absent."
        )

    # The redundancy-domain owner is distinct from configured local-group truth:
    # it joins every current normalized SVI member inside the exact observed IPv4
    # VLAN/VRF/subnet. Persisted DOCX input must remain embedded-unverified.
    fhrp_domain_input = _as_dict(snap_dict.get("fhrp_redundancy_domain_baseline"))
    try:
        from cisco_toolkit.fhrp_redundancy import (
            validate_fhrp_redundancy_domain_baseline,
        )

        fhrp_domain_view = validate_fhrp_redundancy_domain_baseline(fhrp_domain_input)
    except Exception:
        fhrp_domain_view = {}
    fhrp_domain_valid = (
        isinstance(fhrp_domain_view, dict)
        and fhrp_domain_view.get("valid") is True
        and isinstance(fhrp_domain_view.get("baseline"), dict)
        and fhrp_domain_view["baseline"].get("schema") == "fhrp_redundancy_domain_baseline/1"
        and fhrp_domain_view["baseline"].get("projection_custody") == "embedded_unverified"
        and isinstance(fhrp_domain_view.get("rows"), list)
        and isinstance(fhrp_domain_view["baseline"].get("domains"), list)
    )
    if fhrp_domain_valid:
        fhrp_domain_receipt = fhrp_domain_view["baseline"]
        fhrp_domain_rows = [
            row for row in fhrp_domain_view["rows"] if isinstance(row, dict)
        ]
        if len(fhrp_domain_rows) != len(fhrp_domain_view["rows"]):
            fhrp_domain_valid = False
    if fhrp_domain_valid:
        fhrp_domain_source_valid = bool(
            isinstance(fhrp_domain_receipt.get("source_receipt"), dict)
            and fhrp_domain_receipt["source_receipt"].get("valid") is True
        )
        fhrp_domain_verdict = _fhrp_text(fhrp_domain_receipt.get("verdict")) or "NOT ASSESSED"
        fhrp_domain_statuses = {
            status: sum(str(row.get("status") or "") == status for row in fhrp_domain_rows)
            for status in ("degraded", "review", "not_verified", "assessed")
        }
        doc.add_paragraph(
            "FHRP redundancy-domain composition gate — exact observed IPv4 VLAN + normalized VRF + "
            f"subnet. Verdict: {'NOT VERIFIED' if not fhrp_domain_source_valid else fhrp_domain_verdict.upper()}. "
            f"{len(fhrp_domain_rows)} current normalized SVI member row(s) across "
            f"{len(fhrp_domain_receipt['domains'])} observed domain(s): "
            f"{fhrp_domain_statuses['degraded']} degraded, {fhrp_domain_statuses['review']} review, "
            f"{fhrp_domain_statuses['not_verified']} not verified, and "
            f"{fhrp_domain_statuses['assessed']} assessed. Receipt custody: embedded_unverified. "
            "Candidate identity is subtype + group + virtual IP. Zero positive participation is REVIEW "
            "because intended membership is unresolved, not a proven unprotected gateway or failure. "
            "Matching unresolved composition is NOT ACCEPTANCE."
        )
        if not fhrp_domain_source_valid:
            run = doc.add_paragraph().add_run(
                "The closed owner receipt is structurally valid, but its current-run configured-group/SVI "
                "source receipt is not verified. Zero published rows is not evidence of absence or complete "
                "intent; re-collect both inputs before acceptance."
            )
            run.bold = True
            run.font.color.rgb = RGBColor(0x7F, 0x60, 0x00)
        if fhrp_domain_verdict.upper() == "NOT_APPLICABLE":
            run = doc.add_paragraph().add_run(
                "NOT_APPLICABLE means no subject was identified; it is not proof that FHRP is absent "
                "or that intended membership is complete."
            )
            run.bold = True
            run.font.color.rgb = RGBColor(0x7F, 0x60, 0x00)
        doc.add_paragraph(
            "Boundaries: current normalized in-scope SVI members and observed IPv4 domains only. "
            "No expected/off-scan member count or identity, simultaneous roles, timers, authentication, "
            "tracking, failover behavior, convergence, freshness, interoperability, or cutover "
            "authorization is established."
        )
        domain_blockers = [
            row for row in fhrp_domain_rows
            if str(row.get("status") or "") in {"degraded", "review", "not_verified"}
        ]
        domain_ordinary = [
            row for row in fhrp_domain_rows
            if str(row.get("status") or "") not in {"degraded", "review", "not_verified"}
        ]
        domain_shown = domain_blockers + domain_ordinary[:50]
        if domain_shown:
            doc.add_paragraph(
                "FHRP redundancy-domain members (every degraded, review, and not-verified row is "
                "rendered outside the 50-row ordinary display cap):"
            )
            table(
                ["Status", "Device / SVI", "Observed domain", "Participation / candidate",
                 "Role", "Custody / source / command", "Acceptance / finding"],
                [[
                    _fhrp_text(row.get("status")).replace("_", " ").upper(),
                    f"{_fhrp_text(row.get('switch'))} · {_fhrp_text(row.get('interface'))} · "
                    f"{_fhrp_text(row.get('svi_ip'))}",
                    f"VLAN {_fhrp_text(row.get('vlan'))} · VRF "
                    f"{_fhrp_text(row.get('vrf')) or 'global'} · {_fhrp_text(row.get('subnet'))}",
                    f"{_fhrp_text(row.get('participation')).replace('_', ' ').upper()} · "
                    f"{_fhrp_text(row.get('candidate_key')) or 'no positive candidate'}",
                    f"{_fhrp_text(row.get('protocol')) or '—'} group "
                    f"{_fhrp_text(row.get('group')) or '—'} VIP "
                    f"{_fhrp_text(row.get('virtual_ip')) or '—'} · "
                    f"{_fhrp_text(row.get('role')) or '—'}",
                    f"{_fhrp_text(row.get('projection_custody')) or 'embedded_unverified'} · "
                    f"{_fhrp_text(row.get('source_key')) or '—'} · "
                    f"{_fhrp_text(row.get('command')) or '—'}",
                    "\n".join(filter(None, (
                        _fhrp_text(row.get("acceptance")),
                        _fhrp_findings(row.get("findings")),
                    ))) or "—",
                ] for row in domain_shown],
                widths=[1.0, 1.9, 2.1, 2.0, 1.8, 2.8, 3.4],
            )
        if len(domain_ordinary) > 50:
            doc.add_paragraph(
                f"All {len(domain_blockers)} blocker row(s) are shown; 50 of "
                f"{len(domain_ordinary)} ordinary row(s) are shown. Inspect the uncapped FHRP Domains "
                "workbook sheet for the remainder."
            )
    else:
        doc.add_paragraph(
            "FHRP redundancy-domain composition gate — Verdict: NOT ASSESSED. "
            "fhrp_redundancy_domain_baseline/1 is absent, malformed, or not an embedded-unverified "
            "serialized receipt. No rejected receipt leaves, member counts, or domain-composition "
            "conclusion are shown. Re-collect and validate the source evidence before acceptance."
        )

    if pa_valid:
        pa_states = Counter(str(row.get("state") or "unknown") for row in pa_rows)
        n_cells = len(pa_rows)
        n_assessed = pa_states["assessed"]
        n_partial = pa_states["partial"]
        n_not_assessed = max(0, n_cells - n_assessed - n_partial)
        n_health_rows = sum(bool(row.get("health_row_emitted")) for row in pa_rows)
        doc.add_paragraph(
            f"Runtime assessability: {n_health_rows} of {n_cells} device × protocol-family cells "
            f"emitted a bounded health row; {n_assessed} assessed, {n_partial} partial, and "
            f"{n_not_assessed} not assessable. Missing rows are never interpreted as healthy or as "
            "proof that the protocol is absent. This receipt reports current-run collection/parser "
            "reachability, not an expected-neighbor or configured-protocol denominator.")

        family_order = [str(family.get("protocol") or "")
                        for family in _R(passess.get("families"))]
        if not family_order:
            family_order = sorted({str(row.get("protocol") or "") for row in pa_rows})
        family_rows = []
        for protocol in family_order:
            cells = [row for row in pa_rows if row.get("protocol") == protocol]
            states = Counter(str(row.get("state") or "unknown") for row in cells)
            partial = states["partial"]
            assessed = states["assessed"]
            not_assessed = max(0, len(cells) - assessed - partial)
            family_rows.append([
                protocol,
                assessed,
                partial,
                not_assessed,
                sum(bool(row.get("health_row_emitted")) for row in cells),
            ])
        table(["Protocol family", "Assessed", "Partial", "Not assessable", "Health rows"],
              family_rows, widths=[1.5, 0.9, 0.9, 1.2, 1.0])

        gaps = [row for row in pa_rows if row.get("state") != "assessed"]
        if gaps:
            doc.add_paragraph("Current-run protocol evidence gaps requiring collection or parser review:")
            table(["State", "Switch", "Protocol", "Capture / input states"],
                  [[str(row.get("state") or "").replace("_", " ").upper(),
                    row.get("switch", ""), row.get("protocol", ""),
                    "; ".join(f"{name}={value}" for name, value in
                              _as_dict(row.get("input_states")).items())]
                   for row in gaps[:18]], widths=[1.4, 1.8, 1.2, 3.8])
            _disclose(doc, len(gaps), 18, "protocol evidence gap(s)",
                      "Collection Completeness",
                      "Use its Protocol assessability block. The family summary above counts the full "
                      "receipt; each omitted gap remains "
                      "not assessable until its stated capture is supplied.")
    else:
        doc.add_paragraph(
            f"Runtime protocol assessability receipt unavailable ({len(phealth)} observed health "
            "row(s) carried separately). Observed health rows or a lack of "
            "advisories cannot establish device × family coverage; no fleet-wide protocol health "
            "conclusion is asserted.")

    # Keep the sparse observed-state evidence visible independently of advisory doctrine.  A
    # family such as EIGRP can emit a useful presence-only health row while correctly emitting no
    # protocol_intelligence advisory; hiding it made the runbook disagree with the workbook and
    # snapshot.  Seed the bounded table with one severity-ranked row per observed family, then use
    # the remaining budget for the next most severe rows, so a noisy family cannot crowd out all
    # evidence from another family.
    if phealth:
        _ph_cap = 18
        _ph_ranked = sorted(
            phealth,
            key=lambda row: (
                _SEV_ORDER.get(row.get("severity"), 9),
                str(row.get("protocol") or ""),
                str(row.get("switch") or ""),
            ),
        )
        _ph_family_first = []
        _ph_remaining = []
        _ph_seen_families = set()
        for row in _ph_ranked:
            family = str(row.get("protocol") or "Unknown")
            if family not in _ph_seen_families:
                _ph_seen_families.add(family)
                _ph_family_first.append(row)
            else:
                _ph_remaining.append(row)
        _ph_shown = (_ph_family_first + _ph_remaining)[:_ph_cap]
        _ph_family_word = "family" if len(_ph_seen_families) == 1 else "families"
        doc.add_paragraph(
            f"Observed protocol health: {len(phealth)} current-run row(s) across "
            f"{len(_ph_seen_families)} observed {_ph_family_word}. These are bounded observations; "
            "Info means no supported issue was observed in that row, not fleet-wide protocol health.")

        def _ph_text(value, width=180) -> str:
            text = str(value or "").strip()
            return textwrap.shorten(text, width=width, placeholder=" …") if text else "—"

        table(["Severity", "Switch", "Protocol", "Observed summary", "Detail"],
              [[row.get("severity", ""), row.get("switch", ""), row.get("protocol", ""),
                _ph_text(row.get("summary")), _ph_text(row.get("detail"))]
               for row in _ph_shown],
              widths=[0.9, 1.7, 1.1, 3.2, 3.1])
        _disclose(doc, len(phealth), len(_ph_shown), "observed protocol-health row(s)",
                  "Protocol Health",
                  "The bounded runbook table preserves at least one observed row per family before "
                  "using its remaining row budget.")

    if pintel:
        n_high = sum(1 for r in pintel if r.get("severity") == "High")
        by_proto = sorted({r.get("protocol", "") for r in pintel})
        doc.add_paragraph(
            f"{len(pintel)} abnormal control-plane state(s) across {len(by_proto)} protocol(s) "
            f"[{', '.join(by_proto)}], {n_high} High. Each observed state is a fact (read from the "
            f"device); the likely cause is Inferred per Cisco doctrine ({_CONF_UNKNOWN} until validated "
            "live). Fix these before the affected switches are cut over.")
        # Budget the two parts SEPARATELY — the identical fix §6.3's _why_next already carries. The old
        # single (likely_cause + remediation)[:200] slice cut from the END, so it ate the REMEDIATION,
        # the only actionable half of the cell, and cut mid-word with no marker: on the 303-device
        # production snapshot 10 of 84 rows exceeded 200 chars and lost 44–79 characters, e.g. the cell
        # ended "…check the" having dropped " peer port-channel and any `lacp min-links`." shorten()
        # cuts at a word boundary and shows an ellipsis, and per-part budgets keep both halves present.
        def _cause_fix(r: dict) -> str:
            why = str(r.get("likely_cause", "") or "")
            fix = str(r.get("remediation", "") or "")
            if why:
                why = textwrap.shorten(why, width=200, placeholder=" …")
            if fix:
                fix = textwrap.shorten(fix, width=200, placeholder=" …")
            return (why + ("  →  " + fix if fix else "")).strip()

        _pi_ranked = sorted(pintel, key=lambda r: _SEV_ORDER.get(r.get("severity"), 9))
        table(["Severity", "Switch", "Protocol", "State", "Likely cause (Inferred) → Remediation"],
              [[r.get("severity", ""), r.get("switch", ""), r.get("protocol", ""), r.get("state", ""),
                _cause_fix(r)] for r in _pi_ranked[:18]],
              widths=[0.9, 2.0, 1.1, 0.7, 4.3])
        _disclose(doc, len(pintel), 18, "abnormal control-plane state(s)", "Protocol Intelligence",
                  f"The paragraph above counts all {len(pintel)}; each one is to be fixed before its "
                  "switch is cut over, including the unlisted ones.")

    sm = _as_dict(snap_dict.get("service_map"))
    svcs = _R(sm.get("services"))
    mc = _as_dict(sm.get("multicast"))
    if svcs or mc.get("active_interfaces"):
        doc.add_heading("6.6 Services & multicast (observed vs intent)", level=2)
        doc.add_paragraph(
            f"L4 services referenced in ACLs are migration *design intent* ({_CONF_UNKNOWN} as live "
            "traffic — there is no flow telemetry); preserve these reachabilities across the cutover. "
            f"Multicast forwarding presence (PIM/mroute on an interface) is {_CONF_CONFIRMED}; when the "
            "IGMP/mroute group collection is present the broadcast/PTP groups are classified below.")
        if svcs:
            table(["Port", "Proto", "Service", "Category", "ACL refs", "Switches"],
                  [[s.get("port"), s.get("proto"), s.get("service"), s.get("category"),
                    s.get("refs"), s.get("host_count")] for s in svcs[:15]],
                  widths=[0.8, 0.8, 1.7, 1.4, 1.0, 1.0])
        # media-fabric intelligence (NEW-V3.23.115) — read up here because the broadcast/AV headline
        # below needs its authority census (see _av_authority).
        mi = _as_dict(snap_dict.get("multicast_intelligence"))
        groups = _R(mc.get("classified_groups"))
        bcast = [g for g in groups if g.get("broadcast")]
        gnote = (f" {len(groups)} multicast group(s) classified ({len(bcast)} broadcast/AV"
                 f"{_av_authority(mi)})." if groups
                 else f" Per-group (S,G)/IGMP classification {_CONF_UNKNOWN} until re-collection.")
        doc.add_paragraph(
            f"Multicast: {mc.get('active_interfaces', 0)} interface(s) across "
            f"{mc.get('active_switch_count', 0)} switch(es) run PIM/mroute ({_CONF_CONFIRMED} "
            f"forwarding presence).{gnote}")
        # PTP is a per-switch map {host: clock-summary}; report operational boundary clocks vs dormant.
        ptp = _as_dict(mc.get("ptp"))
        if ptp:
            oper = [h for h, v in ptp.items() if _as_dict(v).get("operational")]
            gms = sorted({_as_dict(v).get("grandmaster") for v in ptp.values() if _as_dict(v).get("grandmaster")})
            if oper:
                doc.add_paragraph(
                    f"PTP (IEEE 1588): {len(oper)} of {len(ptp)} switch(es) act as boundary/transparent "
                    f"clocks" + (f" (grandmaster(s) {', '.join(gms)})" if gms else "") +
                    f"; the rest have PTP present but dormant ({_CONF_CONFIRMED}). PTP lock gates any "
                    "ST 2110 / AES67 / Dante move-group cutover.")
            else:
                doc.add_paragraph(
                    f"PTP (IEEE 1588): present on {len(ptp)} switch(es) but NONE are active boundary/"
                    "transparent clocks (Device Type Unknown / 0 ports / no parent) — PTP is flowing as "
                    f"plain multicast, not boundary-clocked ({_CONF_CONFIRMED}). For ST 2110/AES67/Dante "
                    "timing accuracy, confirm whether boundary-clock mode is required on the media-path "
                    "switches before the cutover.")
        queriers = _R(mc.get("igmp_queriers"))
        if queriers:
            _q_vlans = len({q.get("vlan") for q in queriers if isinstance(q, dict) and q.get("vlan") is not None})
            doc.add_paragraph(f"IGMP snooping querier present on {_q_vlans} VLAN(s) "
                              f"({len(queriers)} querier records; {_CONF_CONFIRMED}); VLANs without a querier "
                              f"risk multicast flooding/pruning.")
        # media-fabric intelligence (NEW-V3.23.115): MAC-aliasing / querier coverage / PTP tree
        aliases = _R(mi.get("mac_aliases"))
        if aliases:
            # review r9 F2: the on-air / Broadcast-AV flag is what raises a mac-alias finding from Medium
            # to High, and it is a CURATED classification from the offline registry, not a measurement.
            # Render the flag with its authority, never the flag alone.
            def _alias_line(a: dict) -> str:
                line = (f"{a.get('mac')} ← "
                        f"{', '.join(str(_g) for _g in _as_list(a.get('groups')))}")
                if a.get("has_av"):
                    line += (" [on-air/Broadcast-AV — "
                             + ("authoritative registry semantics" if a.get("has_av_authoritative")
                                else "CURATED classification, NOT an authoritative source")
                             + "; this flag is what raises the finding to High]")
                return line

            doc.add_paragraph(
                f"Multicast MAC-address aliasing ({_CONF_CONFIRMED}, RFC 4541): {len(aliases)} case(s) where "
                "groups collapse to one L2 MAC (IPv4 multicast is 32:1 into Ethernet MACs) so a MAC-level switch "
                "forwards them together — " + "; ".join(_alias_line(a) for a in aliases[:4])
                + ". The overlap itself is arithmetic on the observed group addresses ("
                + _CONF_CONFIRMED + "); the on-air classification is not. "
                + "Re-address one of each overlapping pair (or use IGMPv3 SSM end-to-end) before the cutover.")
        gaps = _as_list(_as_dict(mi.get("querier")).get("gap_vlans"))
        if gaps:
            doc.add_paragraph(
                f"Multicast VLANs without an IGMP querier ({_CONF_CONFIRMED}, RFC 4541): VLAN(s) "
                f"{', '.join(gaps[:20])} carry multicast but have no querier — flooding/blackhole risk; "
                "configure exactly one querier per multicast VLAN.")
        ptp_tree = _as_dict(mi.get("ptp"))
        if ptp_tree.get("n_clocks"):
            doc.add_paragraph(
                f"PTP timing tree (ST 2059): {ptp_tree.get('n_operational', 0)} of {ptp_tree.get('n_clocks', 0)} "
                f"clock(s) are active boundary clocks, {ptp_tree.get('n_dormant', 0)} dormant"
                + (f"; grandmaster(s) {', '.join(str(_m) for _m in _as_list(ptp_tree.get('grandmasters')))}."
                   if ptp_tree.get("grandmasters") else "; no grandmaster observed."))

    # ===== 6.7 Application domains (workload synthesis & migration playbook) — NEW-V3.23.112 =====
    appi = _as_dict(snap_dict.get("application_intelligence"))
    appd = _R(appi.get("domains"))
    if appd:
        doc.add_heading("6.7 Application domains & migration playbook", level=2)
        asum = _as_dict(appi.get("summary"))
        doc.add_paragraph(
            f"The fleet resolves into {asum.get('n_domains', len(appd))} application domain(s) "
            f"({asum.get('n_on_air_critical', 0)} on-air-critical, {asum.get('n_high_risk', 0)} carrying a "
            "High/Critical migration risk). Each switch is attributed by a deliberate hostname role token, by "
            f"observed PTP/multicast ({_CONF_CONFIRMED}), or by its dominant endpoint class ({_CONF_MED}); "
            "switches with no application signal fall to a General bucket rather than being over-claimed as a "
            "media fabric. Per-domain migration risk is grounded in broadcast-over-IP practice (SMPTE ST 2059 "
            "boundary clocks, RFC 4541 IGMP-querier continuity, ST 2022-7 dual-path, Avid NEXIS dual-leg).")
        table(["Domain", "Tier", "Switches", "Endpoints", "PTP", "Top migration risk"],
              [[d.get("domain"), d.get("tier"), d.get("switch_count"), d.get("endpoint_count"),
                ("boundary-clocked" if d.get("ptp_boundary_clocked")
                 else "present, NOT BC" if d.get("ptp_present") else "—"),
                # _R (not _as_list): a truthy non-dict ROW inside a well-formed risks LIST survived the
                # container coercion and crashed .get() — the list guard never looked at the elements.
                ((_R(d.get("risks")) or [{}])[0].get("title", "—") if d.get("risks") else "—")]
               for d in appd],
              widths=[2.2, 1.2, 0.8, 0.9, 1.1, 2.3])
        # the on-air-critical domain risks as evidence-disciplined findings (false-health doctrine)
        for d in appd:
            if d.get("tier") != "On-air critical":
                continue
            for r in _R(d.get("risks")):   # same row-level guard as the §6.7 table above
                if r.get("severity") not in ("Critical", "High"):
                    continue
                finding_block(
                    f"{d.get('domain')} — {r.get('title')}",
                    severity=r.get("severity"),
                    scope=f"{d.get('switch_count')} switch(es); tier {d.get('tier')}",
                    observed=d.get("evidence", ""),
                    interpretation=r.get("detail", ""),
                    impact="On-air media path — a timing or multicast break at cutover is audience-visible.",
                    confidence=f"{_CONF_CONFIRMED} (switch state) / {_CONF_MED} (media dependency)",
                    unknowns="Whether this domain's live traffic actually traverses the flagged switches at "
                             "cutover time (no flow telemetry).",
                    next_validation="; ".join(str(_v) for _v in _as_list(d.get("validation"))),
                    remediation=r.get("remediation", ""))
        cross = _R(appi.get("cross_domain_risks"))
        if cross:
            doc.add_paragraph(
                f"Cross-domain risks (IGMP querier continuity / RFC 4541 + dependency coupling): "
                f"{len(cross)} finding(s) — " + "; ".join(str(c.get("title", "")) for c in cross[:6]) + ".")
        # 6.7.1 inter-domain dependency graph (NEW-V3.23.113)
        edges = _R(appi.get("edges"))
        keystones = _R(appi.get("keystones"))
        if edges:
            doc.add_heading("6.7.1 Inter-domain dependencies (what couples to what)", level=3)
            kline = ""
            if keystones:
                k = keystones[0]
                kline = (f" The keystone domain is {k.get('domain')} (weighted coupling degree "
                         f"{k.get('degree')}, {len(_as_list(k.get('neighbors')))} neighbour domains) — the widest "
                         "cross-domain blast radius; sequence it deliberately, not as an early pilot.")
            doc.add_paragraph(
                f"The domains are coupled by {len(edges)} dependency edge(s): physical inter-switch links "
                f"({_CONF_CONFIRMED}), shared subnets ({_CONF_HIGH}), and dual-homed endpoints that span domains "
                f"({_CONF_CONFIRMED}). A coupling means two domains cannot be cut over in isolation without "
                f"coordination.{kline}")
            table(["Domain A", "Domain B", "Weight", "Coupling", "Migration note"],
                  [[e.get("source"), e.get("target"), e.get("weight"), ", ".join(str(_k) for _k in _as_list(e.get("kinds"))),
                    e.get("migration_note") or ("media-adjacent" if e.get("media") else "")]
                   for e in edges[:15]],
                  widths=[1.9, 1.9, 0.7, 1.4, 2.6])
            _disclose(doc, len(edges), 15, "dependency edge(s)", "Application Intelligence",
                      f"The paragraph above counts all {len(edges)}; an unlisted coupling still means "
                      "those two domains cannot be cut over in isolation.")
        # 6.7.2 recommended cutover order (NEW-V3.23.114)
        cutover = _R(appi.get("cutover_order"))
        if cutover:
            doc.add_heading("6.7.2 Recommended cutover order", level=3)
            doc.add_paragraph(
                "A migration-criticality score (tier + health + open risk + dependency coupling + timing/wave "
                f"flags) orders the domains lowest-risk first. Start with the {asum.get('pilot_domain', '')} "
                f"pilot to fail-fast and learn; migrate {asum.get('last_domain', '')} last with make-before-break "
                "and full validation. A recommendation — the war-room sets the final schedule.")
            table(["#", "Band", "Domain", "Tier", "Score", "Rationale"],
                  [[c.get("order"), c.get("band"), c.get("domain"), c.get("tier"), c.get("score"),
                    c.get("rationale")] for c in cutover],
                  widths=[0.4, 0.8, 2.0, 1.2, 0.6, 3.0])

    # 6.8 Segmentation & isolation (NEW-V3.23.118)
    seg = _as_dict(snap_dict.get("segmentation"))
    ss = _as_dict(seg.get("summary"))
    if ss.get("n_gateways"):
        doc.add_heading("6.8 Segmentation & isolation", level=2)
        ga = _as_dict(seg.get("gateway_acl"))
        doc.add_paragraph(
            ("The L3 fabric is FLAT: " if ss.get("flat") else "")
            + f"{ss.get('n_vrfs', 0)} VRF(s) across {ss.get('n_gateways', 0)} gateway SVI(s); only "
            f"{ga.get('n_with_acl', 0)} ({ga.get('coverage_pct', 0)}%) apply an ACL. "
            f"{ss.get('n_oncrit_exposed', 0)} on-air-critical domain(s) are not isolated — they share the global "
            "routing table with back-office and management and have no gateway ACL. Segmentation for the media "
            "fabric must be designed into the target.")
        rows = [[d.get("domain"), d.get("tier"), d.get("gateways"),
                 "yes" if d.get("isolated") else "NO", d.get("exposure")]
                for d in _R(seg.get("domains"))]
        if rows:
            table(["Domain", "Tier", "Gateways", "Isolated", "Exposure"], rows, widths=[2.2, 1.2, 0.8, 0.8, 3.0])

    # 6.9 golden-config drift (NEW-V3.23.146): per-device running-config compliance vs the baseline.
    gd = _as_dict(snap_dict.get("golden_drift"))
    gsum = _as_dict(gd.get("summary"))
    if gsum.get("n_baseline"):
        doc.add_heading("6.9 Configuration-standard drift", level=2)
        gmode = gsum.get("mode") or gd.get("mode") or "majority"
        gsrc = ("a supplied golden-config baseline" if gmode == "golden-file"
                else "the fleet's de-facto majority baseline (auto-derived)")
        doc.add_paragraph(
            f"Each device's running-config measured against {gsrc} of {gsum.get('n_baseline', 0)} required "
            f"directive(s): {gsum.get('n_drifting', 0)} of {gsum.get('n_devices', 0)} device(s) drift "
            f"(avg compliance {gsum.get('avg_compliance_pct', 0)}%). Bring the drifting devices up to the "
            "standard before — or as part of — their cutover wave.")
        drows = [[d.get("host"), f"{d.get('compliance_pct', 0)}%", d.get("n_missing", 0),
                  "; ".join(str(x) for x in _as_list(d.get("missing"))[:6])]
                 for d in _R(gd.get("per_device")) if d.get("n_missing")]
        if drows:
            table(["Device", "Compliance", "Missing", "Missing required directives"],
                  drows, widths=[1.6, 1.0, 0.8, 4.6])
        else:
            doc.add_paragraph("All devices match the baseline — no configuration drift.")

    # 6.10 syslog intelligence (NEW-V3.23.164): NOS-style operational log analysis.
    si = _as_dict(snap_dict.get("syslog_intelligence"))
    ssum = _as_dict(si.get("summary"))
    if ssum.get("n_devices"):
        doc.add_heading("6.10 Syslog intelligence (operational log analysis)", level=2)
        if not ssum.get("n_collected"):
            doc.add_paragraph(
                "No 'show logging' output was collected for this fleet, so the operational log "
                "axis carries no evidence either way. Re-run the collection (the command is part "
                "of the standard set) to add log-derived detections to this assessment.")
        else:
            doc.add_paragraph(
                f"{ssum.get('total_events', 0)} log event(s) parsed from the buffered 'show logging' "
                f"output of {ssum.get('n_collected', 0)} of {ssum.get('n_devices', 0)} device(s): "
                f"{ssum.get('crit_0_2', 0)} critical (severity 0–2), {ssum.get('err_3', 0)} error "
                f"(severity 3), {ssum.get('n_detections', 0)} operational detection(s). The log buffer "
                "is a bounded recent window, so counts are floors, not totals.")
            _sy_dets = _R(si.get("detections"))
            srows = [[d.get("host"), d.get("label"), d.get("severity"), d.get("count", 0),
                      d.get("recommendation")]
                     for d in _sy_dets[:20]]
            if srows:
                table(["Device", "Finding", "Severity", "Count", "Recommendation"],
                      srows, widths=[1.4, 1.5, 0.8, 0.6, 3.7])
                _disclose(doc, len(_sy_dets), 20, "operational detection(s)", "Syslog Intelligence",
                          "The paragraph above states the full detection count from this same axis.")
            else:
                doc.add_paragraph("No operational detections in the collected logs.")
            if ssum.get("n_not_collected"):
                doc.add_paragraph(
                    f"{ssum.get('n_not_collected', 0)} device(s) had no 'show logging' output "
                    f"({', '.join(_as_list(ssum.get('hosts_not_collected')))}); absence of logs is not "
                    "scored as absence of problems.")

    # 6.11 QoS audit (NEW-V3.23.165): configured QoS posture from the captured running-configs.
    qa = _as_dict(snap_dict.get("qos_audit"))
    qsum = _as_dict(qa.get("summary"))
    if qsum.get("n_devices"):
        doc.add_heading("6.11 QoS posture (configured)", level=2)
        if not qsum.get("n_assessable"):
            doc.add_paragraph(
                "No full 'show running-config' captures are available, so the QoS posture is not "
                "assessable for this fleet — declared, not scored.")
        else:
            qmodes = _as_dict(qsum.get("modes"))
            mtxt = ", ".join(f"{v}× {k}" for k, v in qmodes.items()) or "—"
            doc.add_paragraph(
                f"Configured QoS posture of the {qsum.get('n_assessable', 0)} device(s) with a full "
                f"running-config capture (of {qsum.get('n_devices', 0)} total): {mtxt}; "
                f"{qsum.get('n_voice_ports', 0)} voice-VLAN port(s); "
                f"{qsum.get('n_findings', 0)} finding(s). This is the configured posture — live "
                "queue counters are out of scope for an offline assessment.")
            _qa_find = _R(qa.get("findings"))
            qrows = [[f.get("host"), f.get("label"), f.get("severity"), f.get("recommendation")]
                     for f in _qa_find[:15]]
            if qrows:
                table(["Device", "Finding", "Severity", "Recommendation"],
                      qrows, widths=[1.3, 2.0, 0.8, 3.9])
                _disclose(doc, len(_qa_find), 15, "QoS finding(s)", "QoS Audit",
                          "The paragraph above states the full finding count from this same axis.")
            else:
                doc.add_paragraph("No QoS findings on the assessable devices.")
            if qsum.get("n_not_assessable"):
                doc.add_paragraph(
                    f"{qsum.get('n_not_assessable', 0)} device(s) not assessable "
                    f"({', '.join(_as_list(qsum.get('hosts_not_assessable')))}) — no full "
                    "running-config capture.")

    # 6.12 software risk screening (NEW-V3.23.166): advisory-surface + train lifecycle.
    sr = _as_dict(snap_dict.get("software_risk"))
    rsum = _as_dict(sr.get("summary"))
    if rsum.get("n_devices"):
        doc.add_heading("6.12 Software risk screening", level=2)
        doc.add_paragraph(
            "Attack-surface screening from configuration evidence joined to landmark public "
            "advisories (the exploited-in-the-wild surfaces an engineer checks first), plus a "
            "cautious software-train lifecycle classification. This is screening, NOT a "
            "vulnerability scan — validate every running release with the Cisco PSIRT Software "
            "Checker before drawing release-level conclusions.")
        _sr_find = _R(sr.get("findings"))
        rrows = []
        for f in _sr_find[:15]:
            adv = "; ".join(f"{a.get('cve')}" for a in _as_list(f.get("advisories"))) or "—"
            rrows.append([f.get("host"), f.get("surface"), f.get("severity"), adv,
                          f.get("recommendation")])
        if rrows:
            table(["Device", "Exposed surface", "Severity", "Advisory", "Recommendation"],
                  rrows, widths=[1.1, 1.7, 0.7, 1.1, 3.4])
            _disclose(doc, len(_sr_find), 15, "exposed advisory surface(s)", "Software Risk",
                      "An unlisted surface is an unscreened one, not a closed one.")
        else:
            doc.add_paragraph("No exposed advisory surfaces on the config-assessable devices."
                              if rsum.get("n_config_assessable") else
                              "No full running-config captures — surface screening not assessable.")
        tb = _as_dict(rsum.get("train_bands"))
        worst = [b for b in ("Replace/Upgrade", "Verify EoL") if tb.get(b)]
        if worst:
            _worst_dev = [d for d in _R(sr.get("per_device"))
                          if d.get("train_band") in ("Replace/Upgrade", "Verify EoL")]
            wrows = [[d.get("host"), d.get("platform"), d.get("sw_version"), d.get("train"),
                      d.get("train_band")]
                     for d in _worst_dev[:12]]
            doc.add_paragraph(
                "Software trains needing lifecycle attention (band = curated train-level "
                "knowledge; exact dates live in Cisco's published EoL notices):")
            table(["Device", "Platform", "Version", "Train", "Band"],
                  wrows, widths=[1.4, 1.0, 1.6, 2.0, 2.0])
            _disclose(doc, len(_worst_dev), 12, "device(s) on a Replace/Upgrade or Verify-EoL train",
                      "Software Risk",
                      "The upgrade programme is sized by the full count, not by the rows listed here.")
        if rsum.get("n_config_not_assessable"):
            doc.add_paragraph(
                f"{rsum.get('n_config_not_assessable', 0)} device(s) had no full running-config "
                f"({', '.join(_as_list(rsum.get('hosts_config_not_assessable')))}) — surface "
                "screening declared not assessable for them.")

    # 6.13 platform health (NEW-V3.23.167): control-plane CPU/memory capacity screening.
    ph = _as_dict(snap_dict.get("platform_health"))
    psum = _as_dict(ph.get("summary"))
    if psum.get("n_devices"):
        doc.add_heading("6.13 Platform health (control-plane capacity)", level=2)
        if not psum.get("n_collected"):
            doc.add_paragraph(
                "No CPU/memory capacity output was collected for this fleet — the control-plane "
                "capacity axis carries no evidence either way. Re-run the collection (the "
                "commands are part of the standard set).")
        else:
            pbands = _as_dict(psum.get("bands"))
            btxt = ", ".join(f"{v}× {k}" for k, v in pbands.items()) or "—"
            doc.add_paragraph(
                f"Control-plane capacity at collection time on {psum.get('n_collected', 0)} of "
                f"{psum.get('n_devices', 0)} device(s): {btxt}; {psum.get('n_findings', 0)} "
                "finding(s). This is a SINGLE point-in-time sample — a snapshot, not a trend; "
                "correlate with §6.10 (syslog CPUHOG/MALLOCFAIL events) and re-sample close to "
                "the migration window.")
            prow = [[f.get("host"), f.get("label"), f.get("severity"), f.get("detail"),
                     f.get("recommendation")]
                    for f in _R(ph.get("findings"))[:12]]
            if prow:
                table(["Device", "Finding", "Severity", "Detail", "Recommendation"],
                      prow, widths=[1.2, 1.5, 0.7, 1.6, 3.0])
            else:
                doc.add_paragraph("No capacity findings — every sampled control plane has headroom.")
            if psum.get("n_not_collected"):
                doc.add_paragraph(
                    f"{psum.get('n_not_collected', 0)} device(s) without capacity output "
                    f"({', '.join(_as_list(psum.get('hosts_not_collected')))}) — declared, not scored.")

    # ===== 7. Endpoint & VLAN Census =====
    doc.add_heading("7. Endpoint & VLAN Census", level=1)
    p = doc.add_paragraph()
    _label_run(p, "Counting rule:",
               "endpoints = access ports carrying at least one learned host MAC. Trunk/uplink MAC-table "
               f"entries are excluded. Presence is {_CONF_CONFIRMED} at snapshot time; absence is "
               f"{_CONF_UNKNOWN} (MAC aging).")
    ep_canon = _as_dict(_as_dict(snap_dict.get("executive_brief")).get("scale")).get("n_endpoints")
    if ep_canon is not None:
        doc.add_paragraph(f"Evidenced endpoints: {ep_canon} (canonical, from the assessment scale). "
                          f"Access-port host MACs: {ep_total} across {len(ep_per_vlan)} VLAN(s) and "
                          f"{len(ep_per_switch)} switch(es) — one endpoint may present more than one MAC.")
    else:
        doc.add_paragraph(f"Access-port host MACs: {ep_total} across {len(ep_per_vlan)} VLAN(s) and "
                          f"{len(ep_per_switch)} switch(es).")
    doc.add_paragraph("Top VLANs by access-port MAC count:")
    table(["VLAN", "MACs"], [[v, c] for v, c in ep_per_vlan.most_common(12)], widths=[1.5, 1.5])
    doc.add_paragraph("Top switches by access-port MAC count:")
    table(["Switch", "MACs"], [[s, c] for s, c in ep_per_switch.most_common(12)], widths=[4.5, 1.5])

    # vendor + endpoint-type grouping with confidence (a skill first-class output). NEW-V3.23.95.
    ident = _R(snap_dict.get("endpoint_identity"))
    if ident:
        doc.add_heading("7.1 Endpoint identity (vendor & type)", level=2)
        p = doc.add_paragraph()
        _label_run(p, "Method:",
                   "vendor is the MAC-OUI registered owner (a fact); the class is inferred from vendor + "
                   f"description + CDP platform and is labelled {_CONF_HIGH}/{_CONF_MED}/{_CONF_UNKNOWN} "
                   "per endpoint (no role is claimed as proven from passive data).")
        cls_c = Counter(r.get("endpoint_class", "Unknown") for r in ident)
        ven_c = Counter(r.get("vendor") for r in ident if r.get("vendor"))
        resolved = sum(1 for r in ident if r.get("vendor"))
        doc.add_paragraph(f"{len(ident)} endpoints; vendor resolved for {resolved} "
                          f"({round(100 * resolved / max(len(ident), 1))}%).")
        doc.add_paragraph("By inferred class:")
        table(["Class", "Endpoints"], [[k, v] for k, v in cls_c.most_common()], widths=[3.0, 1.5])
        doc.add_paragraph("Top vendors:")
        table(["Vendor", "Endpoints"], [[k, v] for k, v in ven_c.most_common(12)], widths=[4.5, 1.5])
        # Unlike §7's two "Top …" tables — whose populations ARE stated in the sentence above them
        # ("across N VLAN(s) and M switch(es)"), so a reader can reconcile a 12-row ranking — the
        # vendor spread is stated NOWHERE else in the document: 12 of 105 vendors on the production
        # snapshot, with no figure the reader could reconcile against.
        _disclose(doc, len(ven_c), 12, "vendor(s)", "Endpoint Intelligence",
                  "Endpoint-class planning is sized by the full vendor spread, not the top twelve.")

    # ===== 8. Shared Infrastructure & Consolidation =====
    doc.add_heading("8. Shared Infrastructure & Consolidation", level=1)
    cap = _R(snap_dict.get("capacity"))
    # HON-runbook-drop-1: exclude devices whose active-port count was never observed (port_util can read 0.0
    # from absent evidence) so an unmeasured switch can't top the "most spare / consolidate first" list.
    _cap_measured = [c for c in cap if isinstance(c.get("port_util"), (int, float))
                     and str(c.get("active_ports", "")).strip().isdigit()]
    low_util = sorted(_cap_measured, key=lambda c: c.get("port_util") or 0)[:10]
    doc.add_paragraph("Spare port capacity (consolidation candidates — lowest port utilisation):")
    table(["Switch", "Model", "Active/Total", "Port util %"],
          [[c.get("hostname"), c.get("model"), f"{c.get('active_ports')}/{c.get('total_ports')}",
            c.get("port_util")] for c in low_util], widths=[3.2, 2.2, 1.2, 1.2])
    _disclose(doc, len(_cap_measured), 10, "switch(es) with a measured port utilisation", "Capacity",
              "The consolidation business case is sized from the full measured set, not from these "
              "ten — more switches may sit below the threshold you choose.")
    dep = _as_dict(snap_dict.get("endpoint_dependencies"))
    clusters = _R(dep.get("clusters"))
    dual = _R(dep.get("dual_homed"))
    affinity = _R(dep.get("affinity"))
    if clusters or dual or affinity:
        doc.add_heading("8.1 Endpoint clusters & dependencies", level=2)
        doc.add_paragraph(
            "Cohesive units inferred from the endpoint identity model — a distributed system is a "
            f"(vendor, class) seen across many switches. Confidence travels with the identities "
            f"({_CONF_HIGH}/{_CONF_MED}); live cluster role (active/standby) stays {_CONF_UNKNOWN} "
            "until validated.")
        if clusters:
            doc.add_paragraph("Largest cohesive units (clusters):")
            table(["Class", "Vendor", "Endpoints", "Switches", "VLANs", "Spans move-groups?"],
                  [[c.get("endpoint_class"), c.get("vendor"), c.get("count"), c.get("switches"),
                    c.get("vlans"), "YES" if c.get("spans_groups") else "no"] for c in clusters[:12]],
                  widths=[1.7, 2.4, 1.0, 0.9, 0.8, 1.4])
            _disclose(doc, len(clusters), 12, "cohesive unit(s)", "Endpoint Dependencies",
                      "An unlisted cluster can still span move-groups — check the sheet before "
                      "sequencing a wave around this table.")
        if dual:
            split = sum(1 for d in dual if d.get("split_across_groups"))
            doc.add_paragraph(
                f"Dual-homed / NIC-team endpoints: {len(dual)} (same MAC on >=2 switches) — sequence each "
                f"make-before-break so the peer leg stays up. {split} are split across move-groups "
                "(coordinate those waves explicitly).")
        if affinity:
            doc.add_paragraph("Per-VLAN application tiers (dominant endpoint class in each VLAN):")
            table(["VLAN", "Dominant class", "Endpoints", "Class mix"],
                  [[a.get("vlan"), a.get("dominant"), a.get("total"),
                    ", ".join(f"{k} ({v})" for k, v in _as_dict(a.get("classes")).items())]
                   for a in affinity[:12]], widths=[0.9, 2.0, 1.0, 3.5])
            _disclose(doc, len(affinity), 12, "VLAN(s)", "Endpoint Dependencies",
                      "The lead-in says 'each VLAN' — it is every VLAN in the sheet, not only the "
                      "rows rendered here.")

    # ===== 9. Migration Dependency & Move Groups =====
    doc.add_heading("9. Migration Dependency & Move Groups", level=1)
    doc.add_paragraph(
        "Unit of movement: a shared-VLAN dependency group (switches coupled by a common access VLAN). "
        "Dual-homed switches migrate make-before-break; single-homed switches are hard cutovers "
        "(maintenance window; their endpoints take an outage). Numbers reconcile to the workbook's "
        "Move Groups and Migration Readiness sheets.")
    ws_by_group = {w.get("group"): w for w in ws}
    mg_rows = []
    unsequenced = []
    for r in mr:
        g = r.get("group", "")
        w = ws_by_group.get(g)
        # `ws_by_group.get(g, {})` turned a group with NO sequencing record into "0 make-before-break /
        # 0 hard cutover / 0 at-risk endpoints" — under an intro that says single-homed switches take
        # an outage, that reads as "this group needs no maintenance window". The homing split is the
        # ONE fact that decides whether a window and an outage are needed, so an absent record abstains.
        if w is None:
            unsequenced.append(g)
            mg_rows.append([g, len(_as_list(r.get("switches"))), r.get("endpoints", 0),
                            r.get("readiness", ""), "[NOT OBSERVED]", "[NOT OBSERVED]",
                            "[NOT OBSERVED]"])
            continue
        mg_rows.append([g, len(_as_list(r.get("switches"))), r.get("endpoints", 0),
                        r.get("readiness", ""), len(_as_list(w.get("make_before_break"))),
                        len(_as_list(w.get("hard_cutover"))), w.get("hard_cutover_endpoints", 0)])
    if mg_rows:
        table(["Group", "Switches", "Endpoints", "Readiness", "Make-before-break",
               "Hard cutover", "At-risk endpoints"], mg_rows, widths=[1.4, 1.0, 1.0, 1.3, 1.4, 1.1, 1.2])
    if unsequenced:
        doc.add_paragraph(
            f"{len(unsequenced)} group(s) carry no wave-sequencing record ({', '.join(str(g) for g in unsequenced[:8])}"
            + (" …" if len(unsequenced) > 8 else "") + "), so their dual-homed / single-homed split is "
            f"{_CONF_UNKNOWN}: the make-before-break, hard-cutover and at-risk-endpoint columns above "
            "are [NOT OBSERVED], NOT zero. Do not plan those groups as outage-free — re-run the "
            "assessment (or confirm homing per switch) before scheduling their windows.")

    # ===== 10. Risk Register (behaviour-based blast radius) =====
    doc.add_heading("10. Risk Register", level=1)
    doc.add_paragraph(
        "Per-switch migration blast radius from the removal simulation, ranked by stranded endpoints. "
        "Structural exposure (who depends on the element) is separated from live fault state, which "
        f"is {_CONF_UNKNOWN} until validated.")
    fi_rows = [[r.get("host"), r.get("severity"), r.get("vlans_impacted"), r.get("stranded"),
                r.get("hard"), r.get("backup"), r.get("fhrp")]
               for r in sorted(failure_impact,
                               key=lambda r: (_SEV_ORDER.get(r.get("severity"), 9),
                                              -_as_num(r.get("stranded"))))[:15]]
    if fi_rows:
        table(["Switch (remove/migrate)", "Severity", "VLANs", "Stranded eps", "Hard",
               "Backup-covered", "FHRP-covered"], fi_rows, widths=[3.0, 1.0, 0.8, 1.1, 0.8, 1.2, 1.2])
        _disclose(doc, len(failure_impact), 15, "switch blast-radius row(s)", "Failure Impact",
                  "This is the Risk Register: the simulation covers every in-scope switch, and the "
                  "rows below the cut still strand endpoints when they move.")

    # NEW-V3.23.174: §10.1 -- the per-asset compound-risk register (compute_device_dossiers).
    # The structural table above answers "who depends on this box"; the register answers the
    # senior-engineer question "which boxes scare you most when ALL the risks are stacked".
    dossiers = _as_dict(snap_dict.get("device_dossiers"))
    dd_rows = _R(dossiers.get("per_device"))
    if dd_rows:
        doc.add_heading("10.1 Per-asset compound risk (the Device Risk Register)", level=2)
        dsum = _as_dict(dossiers.get("summary"))
        dbands = _as_dict(dsum.get("bands"))
        # review r9 F1: the register rendered `risk_band` with NO coverage qualification, so an asset
        # whose axes were never assessed ("Low", index 0, verdict "No stacked risk") read exactly like a
        # genuinely clean one. Disclose the census fleet-wide here and per asset in the table below --
        # the same disclosure the workbook's Device Risk Register sheet already carries.
        _cov = [_dossier_coverage(d) for d in dd_rows]
        _n_na_tot = sum(na for na, _ax, _t in _cov)
        _n_ax_tot = sum(ax for _na, ax, _t in _cov)
        _n_half = sum(1 for _na, ax, thin in _cov if thin and ax)
        _n_nocensus = sum(1 for _na, ax, _t in _cov if not ax)
        _gap = ""
        if _n_ax_tot:
            _gap += (f" COVERAGE: {_n_na_tot} of {_n_ax_tot} risk axes fleet-wide were NOT ASSESSED "
                     "(per-asset column 'Coverage' below).")
        if _n_nocensus:
            _gap += (f" {_n_nocensus} asset(s) published NO risk-axis census at all — for those rows how much "
                     "of the band rests on evidence is itself NOT ASSESSED, so they are treated as "
                     "un-evidenced, never as clean.")
        if _n_half or dbands.get("Unassessed", 0):
            _gap += (f" {dbands.get('Unassessed', 0)} asset(s) banded Unassessed and {_n_half} asset(s) have "
                     "HALF OR MORE of their risk axes NOT ASSESSED — an abstaining axis scores ZERO "
                     "exposure, so those rows can band 'Low' on absent evidence. Read the Coverage column "
                     "before treating a Low row as clean; not assessed is a collection gap, never a clean "
                     "result.")
        doc.add_paragraph(
            f"{dbands.get('Severe', 0)} Severe and {dbands.get('Elevated', 0)} Elevated of "
            f"{dsum.get('n_devices', 0)} asset(s); {dsum.get('n_compound', 0)} compound pattern(s). "
            "Risk index = topology impact (1–10) × stacked exposure (0–10) per asset -- "
            "independent risks coinciding on one box outrank any single finding. An axis without "
            "evidence is 'not assessed', never scored." + _gap)
        # The verdict is the engineer's summary judgement on the asset; a bare [:220] cut it mid-word
        # with no marker (3 of 303 rows on the production snapshot). shorten() cuts at a word boundary
        # and shows an ellipsis — the same treatment §6.3 and §6.5 give their prose cells.
        def _verdict(d: dict) -> str:
            v = str(d.get("verdict") or "")
            return textwrap.shorten(v, width=220, placeholder=" …") if v else ""

        table(["Asset", "Band", "Index", "Impact × exposure", "Compound", "Coverage", "Engineer's verdict"],
              [[d.get("host"), d.get("risk_band"), d.get("risk_index"),
                f"{d.get('impact_score')} × {d.get('exposure_score')}",
                ", ".join(c.get("code", "") for c in _as_list(d.get("compound"))) or "—",
                _coverage_cell(d),
                _verdict(d)]
               for d in dd_rows[:15]],
              widths=[1.3, 0.8, 0.6, 1.1, 0.9, 1.8, 2.8])
        # Same self-contradiction as §6.2: the paragraph above states the FULL banded population
        # (production snapshot: 250 Severe of 303 assets) while the register renders 15 rows — and
        # because per_device is ranked risk-first, all 15 are Severe, so the omission is invisible.
        _disclose(doc, len(dd_rows), 15, "asset(s)", "Device Risk Register",
                  f"The paragraph above counts the full population ({dbands.get('Severe', 0)} Severe / "
                  f"{dbands.get('Elevated', 0)} Elevated of {dsum.get('n_devices', len(dd_rows))}); the "
                  "register is ranked risk-first, so the unlisted assets are lower-index — not unassessed.")
        comp = [(d.get("host"), c) for d in dd_rows
                for c in _as_list(d.get("compound")) if isinstance(c, dict)]
        if comp:
            doc.add_paragraph("Compound patterns (why these assets outrank their single findings):")
            for host, c in comp[:10]:
                doc.add_paragraph(
                    f"{c.get('code', '')} {c.get('title', '')} — {host}: {c.get('basis', '')}",
                    style="List Bullet")
            _disclose(doc, len(comp), 10, "compound pattern(s)", "Device Risk Register",
                      "The paragraph above states the full compound-pattern count for this fleet.")

    # ===== 11. Validation & Rollback Logic =====
    doc.add_heading("11. Validation & Rollback Logic", level=1)
    doc.add_paragraph("Per-phase, governed by the false-health rule (a green control plane is not "
                      "service proof — validate forwarding for a real endpoint).")
    table(["Stage", "Pre-check", "Validation (forwarding-proof)", "Rollback trigger", "Rollback action"],
          [["Per group, before move",
            "Confirm second uplink/path live; confirm gateway reachable from a test endpoint",
            "Ping/trace from a moved endpoint to its gateway AND to one cross-VLAN service",
            "Any moved endpoint loses gateway reachability, or a bridge link drops",
            "Re-home the group's uplinks to the legacy path; re-validate"]],
          widths=[1.6, 2.4, 2.4, 2.0, 2.0])

    # per-class service-validation derived from the endpoint classes actually present (NEW-V3.23.96).
    psv = _as_dict(_as_dict(snap_dict.get("endpoint_dependencies")).get("per_switch_validation"))
    if psv:
        present = sorted({line for lines in psv.values() for line in _as_list(lines)})
        doc.add_heading("11.1 Service validation by endpoint class (what to prove per switch)", level=2)
        doc.add_paragraph(
            f"The workbook / snapshot carry a per-switch checklist ({len(psv)} switches). The distinct "
            "service-validation items across the fleet — attach the ones matching a switch's hosted "
            "classes to its move:")
        for line in present:
            doc.add_paragraph(line, style="List Bullet")

    # scenario playbooks for the recommended scenarios (NEW-V3.23.98).
    pg2 = _R(_as_dict(snap_dict.get("migration_scenarios")).get("per_group"))
    if pg2:
        seen_sc = {}
        for g in pg2:
            sc = g.get("recommended_scenario")
            pb0 = _as_dict(g.get("playbook"))   # a truthy non-dict playbook must degrade, not crash pb.get() below
            if sc and sc not in seen_sc and pb0:
                seen_sc[sc] = pb0
        if seen_sc:
            doc.add_heading("11.2 Cutover playbook by scenario", level=2)
            table(["Scenario", "Pre-check", "Validate", "Rollback"],
                  [[sc, pb.get("pre", ""), pb.get("validate", ""), pb.get("rollback", "")]
                   for sc, pb in sorted(seen_sc.items())], widths=[1.3, 2.7, 2.7, 2.7])

    # 11.3 generated per-wave post-cutover verification plan (NEW-V3.23.143): the concrete checks +
    # commands + observed baseline / acceptance condition to run after each wave.
    vp = _as_dict(snap_dict.get("validation_plan"))
    vp_by_wave = _as_dict(vp.get("by_wave"))
    if vp_by_wave:
        vs = _as_dict(vp.get("summary"))

        def _validation_blocker(it):
            state = str(it.get("evidence_state") or "").strip().lower()
            expected = str(it.get("expect") or "").strip().upper()
            return state in {"degraded", "review", "not_verified"} or expected.startswith(
                ("PRE-CUTOVER DEGRADED — BLOCKER:", "PRE-CUTOVER REVIEW — BLOCKER:",
                 "FHRP CONFIGURED GROUP NOT VERIFIED — BLOCKER:",
                 "ROUTING BASELINE NOT VERIFIED — BLOCKER:",
                 "ETHERCHANNEL BASELINE NOT VERIFIED — BLOCKER:")
            )

        def _validation_acceptance(it):
            acceptance = str(it.get("expect") or "")
            state = str(it.get("evidence_state") or "").strip()
            custody = str(it.get("projection_custody") or "").strip()
            source = str(it.get("source_key") or "").strip()
            if not (state or custody or source):
                return acceptance
            if custody == "embedded_unverified":
                custody = "embedded_unverified (embedded projection; integrity unverified)"
            evidence = (f"Evidence state: {state or '—'} | Projection custody: {custody or '—'} | "
                        f"Source: {source or '—'}")
            return acceptance + ("\n" if acceptance else "") + evidence

        # 'validation group(s)', not 'wave(s)': validation_plan.by_wave buckets are per move-group ("Group N",
        # the 3 groups with checks), a distinct unit from the sequenced wave_plan waves (the 9). Reserve 'wave'
        # for those; here the count is the validation-group count so the noun must match it.
        doc.add_heading("11.3 Post-cutover verification plan (per validation group)", level=2)
        doc.add_paragraph(
            f"{vs.get('n_items', 0)} check(s) across {vs.get('n_waves', 0)} validation group(s) "
            # 'each wave's cutover' below is the RESERVED sense: cutover happens in the sequenced wave_plan
            # waves. Only the COUNT noun above is corrected (it counts validation groups, not waves).
            f"({vs.get('n_high', 0)} High/Critical). After each wave's cutover run these and compare the "
            "result with the observed baseline / acceptance column. A PRE-CUTOVER DEGRADED row is a blocker "
            "to resolve or explicitly risk-accept before the window; matching that degraded state after "
            "cutover is not success. A PRE-CUTOVER REVIEW row is also a blocker and requires live simultaneous "
            "verification. Any BASELINE NOT VERIFIED row is a blocker, not an acceptance target: re-collect "
            "the protocol evidence or explicitly disposition the gap before the window. Evidence-bearing rows "
            "name their evidence state, embedded-projection custody boundary, and source locator. Full detail "
            "is in the 'Cutover Validation' workbook sheet.")
        for wave, vits in vp_by_wave.items():
            vits = [_as_dict(it) for it in _R(vits)]
            # Evidence-state blockers are cutover gates, not optional detail. Keep every one outside the bounded
            # 14-row allowance used for ordinary checks, then disclose only the ordinary rows that remain.
            evidence_blockers = [it for it in vits if _validation_blocker(it)]
            ordinary = [it for it in vits if not _validation_blocker(it)]
            shown = evidence_blockers + ordinary[:14]
            doc.add_heading(f"Validation group: {wave}", level=3)
            table(["Device", "Category", "Check", "Command", "Observed baseline / acceptance"],
                  [[it.get("device"), it.get("category"), it.get("check"), it.get("command"),
                    _validation_acceptance(it)] for it in shown],
                  widths=[1.6, 1.1, 2.6, 2.2, 2.5])
            omitted = len(ordinary) - min(len(ordinary), 14)
            if omitted:
                doc.add_paragraph(f"…and {omitted} more ordinary check(s) for {wave} — see the 'Cutover "
                                  "Validation' workbook sheet. Every evidence-state blocker is retained above.")

    # ===== 12. War-Room Decision Logic & Open Unknowns =====
    doc.add_heading("12. War-Room Decision Logic & Open Unknowns", level=1)
    doc.add_paragraph("GO / HOLD / ROLLBACK matrix:")
    table(["State after a move", "Decision"],
          [["Endpoint forwarding to gateway + cross-VLAN service proven", "GO (proceed to next group)"],
           ["Control plane up but forwarding not yet proven", "HOLD (validate before proceeding)"],
           ["Moved endpoint stranded, or a bridge link / sole gateway dropped", "ROLLBACK"]],
          widths=[5.0, 2.0])
    doc.add_heading("Open unknowns (these gate the phases above — no confidence theater)", level=2)
    unknowns = [
        ("Active forwarding owner per FHRP VLAN", _CONF_UNKNOWN,
         "Gates any VLAN relying on a specific gateway staying active during cutover.",
         "Capture 'show standby brief' / ARP for the gateway VIP at cutover time."),
        ("Off-scan links and gateways", _CONF_UNKNOWN,
         "A VLAN gatewayed off-scan looks orphaned here; a redundant path off-scan looks like a bridge.",
         "Extend collection to the legacy core / any non-responding device."),
        ("Live cluster role (active/standby) & application dependencies", _CONF_MED,
         "Endpoint vendor/class and cohesive units are now inferred (see §7.1/§8.1), but which node is "
         "active vs standby, and which app depends on which, is not provable from passive data.",
         "Confirm cluster roles + app dependencies with the application owners using the §8.1 clusters "
         "and the per-switch service checklist (§11.1) as the starting map."),
    ]
    for name, conf, gates, nextstep in unknowns:
        p = doc.add_paragraph(style="List Bullet")
        _label_run(p, name + ":", f"[{conf}] {gates}")
        sp = doc.add_paragraph(style="List Bullet 2"); _label_run(sp, "Next validation:", nextstep)

    # ===== 13. Remediation Appendix (generated config, review-only) — NEW-V3.23.116 =====
    rp = _as_dict(snap_dict.get("remediation_plan"))
    rem_items = _R(rp.get("items"))
    if rem_items:
        doc.add_heading("13. Remediation Appendix — generated config (review only)", level=1)
        bp = doc.add_paragraph()
        _label_run(bp, "⚠ Review banner:", rp.get("banner") or
                   "Generated for review — validate against the running-config and change-control before applying.",
                   GREY)
        rs = _as_dict(rp.get("summary"))
        doc.add_paragraph(
            f"{rs.get('n_items', 0)} generated fix snippet(s) across {rs.get('n_devices', 0)} device(s) "
            f"({rs.get('n_high', 0)} High/Critical), per device in severity order. Placeholders (<...>) mark "
            "values that require local intent — do not paste them literally.")
        by_dev = _as_dict(rp.get("by_device"))
        ranked = sorted(by_dev.items(),
                        key=lambda kv: (-sum(1 for it in _R(kv[1]) if it.get("severity") in ("Critical", "High")),
                                        str(kv[0])))
        for host, its in ranked[:40]:
            its = _R(its)   # a per-device item list that is a truthy non-list must not crash slicing/len
            doc.add_heading(str(host), level=2)
            for it in its[:8]:
                hp = doc.add_paragraph()
                hr = hp.add_run(f"[{it.get('severity')}] {it.get('category')} — {it.get('title')}"); hr.bold = True
                cp = doc.add_paragraph(); cp.paragraph_format.left_indent = Inches(0.3)
                cr = cp.add_run("\n".join(_as_list(it.get("commands"))))
                cr.font.name = "Consolas"; cr.font.size = Pt(9)
                _label_run(doc.add_paragraph(style="List Bullet 2"), "Verify:", it.get("verify", ""))
                _label_run(doc.add_paragraph(style="List Bullet 2"), "Caution:", it.get("caution", ""))
            if len(its) > 8:
                doc.add_paragraph(f"…and {len(its) - 8} more item(s) for {host} — see the Remediation Plan sheet.")
        if len(ranked) > 40:
            doc.add_paragraph(f"…and {len(ranked) - 40} more device(s) — see the Remediation Plan workbook sheet.")

    # ---- closing acceptance gate (AS-style back matter) ----
    # Deliverable Excellence (DE-01): shared glossary before the sign-off gate.
    add_glossary(doc)

    add_acceptance(
        doc, scope_note="Acceptance of this runbook confirms its findings and migration approach as "
                        "the agreed basis for the MOP and cutover planning that follow.")

    # landscape for the wide tables; US Letter
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Inches(11), Inches(8.5)
    # running footer
    footer = sec.footer.paragraphs[0]
    footer.text = f"{label} — Assessment & Migration Runbook (DRAFT, evidence-led)"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(8); run.font.color.rgb = GREY

    doc.save(output_path)
    logger.info(f"[Phase 31] Runbook (DOCX) written: {output_path}")
