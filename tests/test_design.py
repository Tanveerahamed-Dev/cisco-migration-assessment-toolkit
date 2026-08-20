"""NEW-V3.23.148: the As-Built Network Design Document (HLD + LLD, DOCX) deliverable. python-docx is
optional, so the whole module is skipped when it is absent (the generator itself fails soft the same
way). These tests pin the HLD/LLD section structure, the reconciliation of the as-built facts to the
snapshot, the BoM aggregation, and the fail-soft path."""
import pytest

docx = pytest.importorskip("docx")  # skip the file if the optional dep is absent
from docx import Document  # noqa: E402

from cisco_toolkit.design import write_design_doc_docx  # noqa: E402


def _snap():
    """A representative snapshot exercising every section the design generator builds."""
    return {
        "script_version": "V3.23.0",
        "devices": {
            "core1": {"hostname": "core1", "model": "N9K-C93180YC", "serial_number": "SAL001",
                      "sw_version": "10.2", "platform": "nxos", "total_ports": 48, "active_ports": 30},
            "dist1": {"hostname": "dist1", "model": "C9300-48T", "serial_number": "SAL002",
                      "sw_version": "17.9", "platform": "ios", "total_ports": 48, "active_ports": 20},
            "acc1": {"hostname": "acc1", "model": "WS-C2960X-48FPD-L", "serial_number": "SAL003",
                     "sw_version": "15.2", "platform": "ios", "total_ports": 48, "active_ports": 12},
        },
        "interfaces": {
            "core1": {"Vlan10": {"svi_ip": "10.0.10.1", "hsrp_behavior": "HSRP grp10 active", "vrf": "PROD",
                                 "acl_in": "PROD_IN"},
                      "Vlan20": {"svi_ip": "10.0.20.1"},
                      "Po1": {"switchport_mode": "Trunk", "cdp_neighbor": "dist1"},
                      "Te1/1": {"switchport_mode": "Trunk", "cdp_neighbor": "dist1"}},
            "dist1": {"Vlan10": {"svi_ip": "10.0.10.2", "hsrp_behavior": "HSRP grp10 standby"},
                      "Gi1/0/1": {"switchport_mode": "Trunk", "cdp_neighbor": "core1"},
                      "Gi1/0/24": {"switchport_mode": "Access", "vlan": "10", "vlan_name": "PROD",
                                   "end_host_mac": "aaaa.0000.0001"}},
            "acc1": {"Gi1/0/1": {"switchport_mode": "Trunk", "cdp_neighbor": "dist1"},
                     "Gi1/0/5": {"switchport_mode": "Access", "vlan": "20", "vlan_name": "VOICE",
                                 "end_host_mac": "bbbb.0000.0001"}},
        },
        "l3_forwarding": [
            {"switch": "core1", "vlan": "10", "svi_ip": "10.0.10.1", "fhrp": "HSRP active",
             "role": "primary", "risk": ""},
            {"switch": "dist1", "vlan": "10", "svi_ip": "10.0.10.2", "fhrp": "HSRP standby",
             "role": "secondary", "risk": ""},
            {"switch": "core1", "vlan": "20", "svi_ip": "10.0.20.1", "fhrp": "",
             "role": "", "risk": "single-gateway"},
        ],
        "routing_neighbors": {"core1": {"ospf": [{"neighbor": "10.0.0.2", "state": "FULL"}],
                                        "eigrp": [], "bgp": [{"neighbor": "10.0.0.3", "as": "65001",
                                                              "state": "Established"}]}},
        "redistribution": {"core1": [{"into_proto": "bgp", "from_proto": "ospf", "route_map": "OSPF_TO_BGP"}]},
        "stp_roots": {"core1": {"10": {"is_root": True, "root_priority": "24576"},
                                "20": {"is_root": True, "root_priority": "24576"}}},
        "fhrp": [{"vid": "10", "members": [{"host": "core1"}, {"host": "dist1"}], "issues": []},
                 {"vid": "20", "members": [{"host": "core1"}], "issues": ["no FHRP peer"]}],
        "vpc": {"core1": {"domain_id": "1", "role": "primary", "peer_status": "peer-ok", "vpcs": []}},
        "lifecycle_risk": {
            "per_device": [
                {"host": "core1", "model": "N9K-C93180YC", "band": "Active", "eos": "", "ldos": ""},
                {"host": "dist1", "model": "C9300-48T", "band": "Active", "eos": "", "ldos": ""},
                {"host": "acc1", "model": "WS-C2960X-48FPD-L", "band": "Past-EoS",
                 "eos": "2025-04-30", "ldos": "2030-04-30"},
            ],
            "summary": {"n_devices": 3, "by_band": {"Active": 2, "Past-EoS": 1}, "n_past_eos": 1,
                        "n_past_ldos": 0, "n_near": 0},
        },
        "capacity": [{"hostname": "core1", "model": "N9K-C93180YC", "total_ports": 48,
                      "active_ports": 30, "free_ports": 18, "port_util": 62.5}],
        "failure_impact": [{"host": "core1", "severity": "High", "vlans_impacted": 2, "stranded": 12,
                            "hard": 4, "detail": "core1 fronts both gateways"}],
        "punchlist": [{"severity": "Critical", "category": "L3 design", "devices": ["core1"],
                       "title": "VLAN 20 has a single gateway (no FHRP)", "detail": "x"},
                      {"severity": "High", "category": "Lifecycle", "devices": ["acc1"],
                       "title": "acc1 is past end-of-sale", "detail": "y"}],
        "subnet_intelligence": {
            "per_device": [
                {"host": "dist1", "is_l3": False, "served_subnets": [
                    {"subnet": "10.0.10.0/24", "gateway": "core1", "vlan": "10"}]},
                {"host": "acc1", "is_l3": False, "served_subnets": [
                    {"subnet": "10.0.20.0/24", "gateway": "core1", "vlan": "20"}]},
            ],
            "move_groups": [],
        },
        "service_map": {
            "multicast": {"active_interfaces": 4, "active_switch_count": 2,
                          "classified_groups": [{"group": "224.0.1.129", "name": "PTP-primary",
                                                 "category": "Broadcast-AV", "broadcast": True}],
                          "igmp_queriers": [{"switch": "core1", "vlan": "10", "querier": "10.0.10.1"}],
                          "ptp": {"core1": {"operational": True}, "dist1": {"operational": False}}},
        },
        "executive_brief": {"posture_statement": "A two-tier campus with a single-gateway exposure on VLAN 20.",
                            "axes": [], "top_gating": []},
    }


def _all_text(doc):
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            parts.extend(c.text for c in row.cells)
    return "\n".join(parts)


def test_design_has_hld_and_lld_sections(tmp_path):
    out = str(tmp_path / "d.docx")
    write_design_doc_docx(out, _snap(), "Unit Test Fleet")
    d = Document(out)
    h1 = [p.text for p in d.paragraphs if p.style.name == "Heading 1"]
    for token in ("1. Executive Design Summary", "2. High-Level Design (HLD)",
                  "3. Low-Level Design (LLD)", "4. Target-State Design Recommendations"):
        assert any(t == token for t in h1), f"missing section: {token}; have {h1}"
    h2 = [p.text for p in d.paragraphs if p.style.name == "Heading 2"]
    for token in ("2.3 Layer-3 design", "2.5 Multicast & timing design",
                  "3.1 Device inventory & roles", "3.4 Equipment list (Bill of Materials)"):
        assert any(t == token for t in h2), f"missing subsection: {token}"


def test_design_scale_reads_canonical_executive_brief(tmp_path):
    """A5 (SSOT): the HLD scale figures (devices / VLANs in scope) must read the canonical
    executive_brief.scale — the published single source the explorer/deck/webapp read — not a local
    recount. Discriminating fixture: scale says 999 devices / 777 VLANs while the raw arrays hold a
    handful each, so a recompute regression would render the small count, not the canonical 999/777."""
    snap = _snap()
    snap["executive_brief"]["scale"] = {"n_devices": 999, "n_vlans": 777, "n_endpoints": 5000}
    out = str(tmp_path / "d.docx")
    write_design_doc_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "999" in text       # "Devices in scope" reads canonical scale
    assert "777" in text       # "VLANs in use" reads canonical scale


def test_design_fhrp_candidates_not_mislabeled_as_configured_groups(tmp_path):
    """FHRP false-redundancy class: snap['fhrp'] is the set of MULTI-GATEWAY VLANs (FHRP candidates),
    most/all WITHOUT a configured first-hop-redundancy protocol (every member fhrp=False — the Meridian shape).
    The §2.4 resilience table must NOT label them 'FHRP gateway groups' (which reads as 'N FHRP groups
    protect the gateways'); it labels them candidates and surfaces the honest count actually running
    FHRP (0 here), so the design doc can never imply redundancy that is not configured."""
    snap = _snap()
    snap["fhrp"] = [
        {"vid": 3, "issues": ["3 gateways but no FHRP — no first-hop redundancy"],
         "members": [{"host": "a", "fhrp": False}, {"host": "b", "fhrp": False}, {"host": "c", "fhrp": False}]},
        {"vid": 4, "issues": ["2 gateways but no FHRP — no first-hop redundancy"],
         "members": [{"host": "a", "fhrp": False}, {"host": "b", "fhrp": False}]},
    ]
    out = str(tmp_path / "d.docx")
    write_design_doc_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "FHRP gateway groups" not in text, "must not imply configured FHRP groups exist (false redundancy)"
    assert "Multi-gateway VLANs (FHRP candidates)" in text
    assert "first-hop redundancy (FHRP) configured" in text


def test_design_interoperability_footprint_section(tmp_path):
    """N14: the HLD must surface what the target design has to keep INTEROPERATING with — the
    observed NOS-family spread and the multi-vendor endpoint estate (OUI-derived) — grounded in
    devices[].platform + endpoint_identity, framed honestly as per-MAC observations (NOT a
    service-dependency / HA-cluster claim)."""
    snap = _snap()
    hosts = list(snap["devices"])
    for i, h in enumerate(hosts):
        snap["devices"][h]["platform"] = "nxos" if i == 0 else "ios"   # two NOS families
    snap["endpoint_identity"] = (
        [{"vendor": "Hewlett Packard", "endpoint_class": "Server", "mac": f"aaaa.0000.{i:04x}"} for i in range(5)]
        + [{"vendor": "APC by Schneider Electric", "endpoint_class": "UPS/PDU", "mac": "bbbb.0000.0001"}])
    out = str(tmp_path / "d.docx")
    write_design_doc_docx(out, snap, "Unit Test Fleet")
    d = Document(out)
    h1 = [p.text for p in d.paragraphs if p.style.name == "Heading 1"]
    assert any("Interoperability" in t for t in h1), h1
    text = _all_text(d)
    assert "IOS" in text and "NXOS" in text                          # both NOS families surfaced
    assert "Hewlett Packard" in text and "APC by Schneider Electric" in text   # vendor mix
    assert "Server" in text and "UPS/PDU" in text                    # class mix
    assert "not a service-dependency" in text.lower() or "not a service-dependency graph" in text  # honest framing


def test_design_reconciles_to_snapshot(tmp_path):
    out = str(tmp_path / "d.docx")
    write_design_doc_docx(out, _snap(), "Unit Test Fleet")
    text = _all_text(Document(out))
    # device inventory carries every device + model + lifecycle band
    for token in ("core1", "dist1", "acc1", "N9K-C93180YC", "C9300-48T", "WS-C2960X-48FPD-L", "Past-EoS"):
        assert token in text, token
    # HLD L3: routing protocols recovered from neighbours
    assert "OSPF" in text and "BGP" in text
    # addressing plan joins the served subnet to the VLAN/gateway
    assert "10.0.10.0/24" in text and "10.0.10.1" in text
    # multicast/timing design surfaces the classified group + PTP
    assert "PTP-primary" in text or "224.0.1.129" in text
    # target-state recommendations are the punch-list read as design intent
    assert "single gateway" in text.lower()


def test_design_bom_aggregates_by_model(tmp_path):
    out = str(tmp_path / "d.docx")
    write_design_doc_docx(out, _snap(), "Unit Test Fleet")
    d = Document(out)
    # the BoM table has one row per distinct model (3 here) + header
    bom = None
    for t in d.tables:
        if t.rows[0].cells[0].text == "Model":
            bom = t
            break
    assert bom is not None, "BoM table not found"
    models = {r.cells[0].text for r in bom.rows[1:]}
    assert models == {"N9K-C93180YC", "C9300-48T", "WS-C2960X-48FPD-L"}


def test_collected_vs_inventoried_is_reconciled_in_scope_inventory_and_bom(tmp_path):
    """§1 'Devices in scope' reads the canonical INVENTORIED count while §3.1, §3.4 (BoM) and §3.5
    all iterate snap['devices'] — the hosts that COLLECTED. On a 303/253 fleet the BoM totalled 253
    chassis and went to procurement 50 short, with nothing in the document naming the gap. Both
    counts must appear, and every table over the narrower population must declare which one it
    covers (absence is never zero)."""
    snap = _snap()
    n_collected = len(snap["devices"])
    snap["executive_brief"]["scale"] = {"n_devices": n_collected + 50, "n_vlans": 3}
    out = str(tmp_path / "d_scope.docx")
    write_design_doc_docx(out, snap, "Unit Test Fleet")
    d = Document(out)
    text = _all_text(d)
    # §1 still reads the canonical inventoried count (SSOT) AND now names the enumerated population
    assert str(n_collected + 50) in text
    assert "50 inventoried device(s) did NOT collect" in text
    # the BoM totals the collected population and says so where procurement will read it
    bom = next(t for t in d.tables if t.rows[0].cells[0].text == "Model"
               and t.rows[0].cells[1].text == "Qty")
    assert sum(int(r.cells[1].text) for r in bom.rows[1:]) == n_collected
    assert "DO NOT ORDER FROM THIS TABLE UNRECONCILED" in text
    assert text.count("Scope reconciliation:") >= 3          # 3.1 inventory, 3.4 BoM, 3.5 software
    assert f"out of {n_collected + 50} inventoried device(s) in scope" in text
    # a fleet where every inventoried device collected carries no gap note (no cry-wolf)
    snap["executive_brief"]["scale"] = {"n_devices": n_collected, "n_vlans": 3}
    out2 = str(tmp_path / "d_full.docx")
    write_design_doc_docx(out2, snap, "Unit Test Fleet")
    assert "Scope reconciliation:" not in _all_text(Document(out2))


def test_design_doc_renders_w37_traceability_matrix(tmp_path):
    """[W3-7] the generated As-Built .docx carries the §4.5 traceability matrix when recommended decisions exist --
    each traced to its CCDE principle id + published citation + the snapshot source fields."""
    snap = _snap()
    snap["design_blueprint"] = {"tradeoff_scorecard": [], "doctrine": {}, "decisions": [
        {"status": "recommended", "title": "Enforce SNMPv3", "priority": "Critical", "domain": "Security",
         "driver": "management-plane integrity",
         "principle": {"id": "mgmt-secure-protocols", "title": "Secure management protocols", "citation": "CCDE Session 19"},
         "evidence": {"summary": "3 device(s) fail management hardening", "devices": ["core1"],
                      "fields": ["security[host].findings[].status"]},
         "recommended_action": "Standardize SNMPv3 + SSH", "alternatives": "local AAA", "tradeoffs": "op lift",
         "axes": ["security"]},
    ]}
    out = str(tmp_path / "trace.docx")
    write_design_doc_docx(out, snap, "Meridian")
    doc = Document(out)
    assert any("Design traceability matrix" in p.text for p in doc.paragraphs)   # the §4.5 section renders
    cells = [c.text for t in doc.tables for row in t.rows for c in row.cells]
    assert any("mgmt-secure-protocols" in c for c in cells)      # decision traced to its CCDE principle id
    assert any("CCDE Session 19" in c for c in cells)            # ...and the published citation
    assert any("security[host].findings[].status" in c for c in cells)   # ...and the snapshot source field


def test_design_multicast_section_suppressed_when_no_activity(tmp_path):
    """A multicast dict that exists but carries all-zero counts (non-media fabric, or the commands were
    not collected) gets the fallback message — not an all-zeros 'PTP: 0 of 0' filler paragraph."""
    snap = _snap()
    snap["service_map"] = {"multicast": {"active_interfaces": 0, "active_switch_count": 0,
                                         "classified_groups": [], "igmp_queriers": [], "ptp": {}}}
    out = str(tmp_path / "d.docx")
    write_design_doc_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "No multicast/PTP activity was classified" in text
    assert "0 of 0 device(s)" not in text


def test_design_tolerates_malformed_stp_roots(tmp_path):
    """A stp_roots VLAN value that is not a dict (malformed snapshot) must not abort the design doc."""
    snap = _snap()
    snap["stp_roots"] = {"core1": {"10": None, "20": "garbage", "30": {"is_root": True}}}
    out = str(tmp_path / "d.docx")
    write_design_doc_docx(out, snap, "Unit Test Fleet")   # must not raise
    import os
    assert os.path.exists(out)


def test_design_carries_document_furniture(tmp_path):
    """V3.23.150: AS-style front/back matter — Document Control between cover and TOC, and the
    closing acceptance signature gate."""
    out = str(tmp_path / "d.docx")
    write_design_doc_docx(out, _snap(), "Unit Test Fleet")
    d = Document(out)
    h1 = [p.text for p in d.paragraphs if p.style.name == "Heading 1"]
    assert "Document Control" in h1 and "Document Acceptance" in h1


def test_design_renders_canonical_blueprint_section(tmp_path):
    """When the snapshot carries the canonical design_blueprint, §4 renders the CCDE-grounded decisions,
    the trade-off scorecard and the open requirement questions (not the punch-list fallback) — proving the
    design DOCX reads the single source of truth rather than re-deriving design intent."""
    snap = _snap()
    snap["design_blueprint"] = {
        "summary": {"headline": "1 critical target-state design decision(s); leading: introduce FHRP.",
                    "n_decisions": 2, "n_recommended": 1, "n_needs_requirement": 1, "n_critical": 1,
                    "by_domain": {}, "requirements_provided": False},
        "tradeoff_scorecard": [{"axis": "availability", "label": "High availability & resiliency",
                                "score": 0, "posture": "Weak", "evidence": "no FHRP; SPOFs present"}],
        "decisions": [
            {"id": "fhrp-first-hop-gateway-redundancy", "title": "Introduce first-hop redundancy",
             "domain": "methodology", "priority": "Critical", "status": "recommended",
             "confidence": "Observed", "driver": "Gateway resilience",
             "evidence": {"summary": "52 VLAN(s) without FHRP", "count": 52, "devices": []},
             "principle": {"id": "fhrp-first-hop-gateway-redundancy", "title": "First-hop redundancy",
                           "citation": "CCDE In Depth — High Availability"},
             "recommended_action": "Deploy HSRP/VRRP across dual gateways", "alternatives": "GLBP; anycast gateway",
             "tradeoffs": "HA vs simplicity", "axes": ["availability", "convergence"], "requirements_needed": []},
            {"id": "availability-right-sized-per-tier", "title": "Right-size availability per tier",
             "domain": "methodology", "priority": "High", "status": "needs-requirement",
             "confidence": "Requirement-needed", "driver": "WHY-first",
             "evidence": {"summary": "redundancy posture varies by tier", "count": 0, "devices": []},
             "principle": {"id": "availability-right-sized-per-tier", "title": "Right-size availability",
                           "citation": "CCDE In Depth — Fundamentals"},
             "recommended_action": "Assign an availability target per class", "alternatives": "uniform SLA fleet-wide",
             "tradeoffs": "HA vs cost", "axes": ["availability", "cost"], "requirements_needed": ["availability_tier"]},
        ],
        "requirements_model": {"fields": [], "open_questions": [], "provided": False, "note": ""},
        "coverage": {"inventory": 3, "collected": 3, "not_collected": 0,
                     "caveat": "grounded only in collected evidence"},
        "methodology": "WHY-first.", "axes": [],
    }
    out = str(tmp_path / "d.docx")
    write_design_doc_docx(out, snap, "Unit Test Fleet")
    d = Document(out)
    h2 = [p.text for p in d.paragraphs if p.style.name == "Heading 2"]
    for token in ("4.1 Design trade-off scorecard", "4.2 Recommended target-state design decisions",
                  "4.3 Open design questions (requirements to confirm)"):
        assert token in h2, f"missing §4 subsection: {token}; have {h2}"
    text = _all_text(d)
    assert "Introduce first-hop redundancy" in text                 # the recommended decision
    assert "CCDE In Depth — High Availability" in text              # traced to its CCDE basis
    assert "Right-size availability per tier" in text               # the open requirement question
    assert "availability_tier" in text
    text = _all_text(d)
    assert "Revision history" in text and "Assumptions & caveats" in text
    assert "Customer network owner" in text                      # acceptance signature roles
    assert "Per-Wave Method of Procedure (.docx)" in text        # related-documents cross-reference


def test_design_renders_full_doctrine_catalogue(tmp_path):
    """§4.4 surfaces the FULL CCDE doctrine reference (incl. non-actionable domains the L1-L4 assessment
    cannot trigger, e.g. firewall), so the HLD reasons with the whole knowledge base, not only the
    auto-emitted decisions -- proving the 51 'dark' principles reach the page."""
    from cisco_toolkit.design_advisor import compute_design_blueprint
    snap = _snap()
    snap["design_blueprint"] = compute_design_blueprint(snap)        # real blueprint -> carries the doctrine catalogue
    out = str(tmp_path / "d.docx")
    write_design_doc_docx(out, snap, "Unit Test Fleet")
    d = Document(out)
    h2 = [p.text for p in d.paragraphs if p.style.name == "Heading 2"]
    assert "4.4 Design doctrine applied (CCDE-grounded reference)" in h2, f"missing §4.4; have {h2}"
    text = _all_text(d)
    assert ("screened-subnet" in text) or ("DMZ" in text), "firewall doctrine must surface in the HLD §4.4"


def test_design_renders_target_state_architecture(tmp_path):
    """C: §5 renders the generated CANDIDATE target-state architecture (tier model, L2/L3 boundary,
    resilience, lifecycle disposition, migration), each row current->target with its rationale."""
    from cisco_toolkit.design_advisor import compute_design_blueprint
    snap = _snap()
    snap["design_blueprint"] = compute_design_blueprint(snap)
    out = str(tmp_path / "d.docx")
    write_design_doc_docx(out, snap, "Unit Test Fleet")
    d = Document(out)
    h1 = [p.text for p in d.paragraphs if p.style.name == "Heading 1"]
    assert any("proposed target-state architecture" in t.lower() for t in h1), f"missing §5; have {h1}"
    text = _all_text(d)
    assert "Topology / tier model" in text                       # a target-state dimension
    assert "candidate" in text.lower()                            # honest framing


def test_design_renders_replacement_bom(tmp_path):
    """C next-layer: §5 surfaces the target-state replacement BoM (EoL gear to procure), grounded in
    lifecycle evidence — the target-procurement detail the as-built §3.4 BoM does not give."""
    from cisco_toolkit.design_advisor import compute_design_blueprint
    snap = _snap()
    snap["lifecycle_risk"] = {"per_device": [
        {"host": "a", "band": "Past-LDoS", "model": "WS-C4948E"},
        {"host": "b", "band": "Past-LDoS", "model": "WS-C4948E"}]}
    snap["design_blueprint"] = compute_design_blueprint(snap)
    out = str(tmp_path / "d.docx")
    write_design_doc_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "Replacement" in text and "WS-C4948E" in text         # target procurement surfaced in §5


def test_design_renders_addressing_plan_when_supernet_supplied(tmp_path):
    """F1: with an address_space requirement, §5.3 renders a candidate per-VLAN IP allocation from it."""
    from cisco_toolkit.design_advisor import compute_design_blueprint
    snap = _snap()
    snap["design_blueprint"] = compute_design_blueprint(snap, {"address_space": "10.50.0.0/16"})
    out = str(tmp_path / "d.docx")
    write_design_doc_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "Net-new IP addressing plan" in text and "10.50." in text


def test_design_addressing_plan_needs_requirement_discloses_census(tmp_path):
    """SSOT + coverage-honesty: with NO address_space, HLD §5.3 still discloses the VLAN census
    context (total / sizeable-observed / querier-only un-sizable) the snapshot carries — not just
    'requirement needed'. Mirrors the explorer's addressing-plan block so the deliverable matches the
    canonical addressing_plan instead of silently dropping the census the engine computed."""
    from cisco_toolkit.design_advisor import compute_design_blueprint
    snap = _snap()
    # a querier-only VLAN (live per IGMP, no access port / SVI) -> census(3) > sizeable(2), 1 un-sizable
    snap["service_map"]["multicast"]["igmp_queriers"].append(
        {"switch": "core1", "vlan": "777", "querier": "10.0.77.1"})
    bp = compute_design_blueprint(snap)  # no address_space -> needs-requirement path
    snap["design_blueprint"] = bp
    ap = bp["target_state"]["addressing_plan"]
    assert ap["status"] == "needs-requirement"
    assert (ap["n_census_vlans"], ap["observed_vlans"], ap["n_unsizable"]) == (3, 2, 1)
    out = str(tmp_path / "d.docx")
    write_design_doc_docx(out, snap, "Unit Test Fleet")
    paras = [p.text for p in Document(out).paragraphs]
    assert any("Net-new IP addressing plan" in p for p in paras)
    census_line = next((p for p in paras if "census" in p.lower()), None)
    assert census_line is not None, "§5.3 must disclose the VLAN census on the needs-requirement path"
    # the disclosed sentence carries all three canonical counts
    assert "3" in census_line and "2" in census_line and "1" in census_line, census_line


def test_design_renders_wave_plan(tmp_path):
    """F2: §5.4 renders the candidate migration wave plan derived from move-groups."""
    from cisco_toolkit.design_advisor import compute_design_blueprint
    snap = _snap()
    snap["move_groups"] = [{"switches": [f"SW{i:03d}" for i in range(95)], "spanning_vlans": [[124, "X", 95]]}]
    snap["design_blueprint"] = compute_design_blueprint(snap)
    out = str(tmp_path / "d.docx")
    write_design_doc_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "Candidate migration wave plan" in text and "coupled-subwave" in text


def test_design_renders_zone_aware_ip_plan(tmp_path):
    """AUDIT: §5.3 renders the per-zone summary table when the IP plan is zone-aware (vlan_zones supplied)."""
    from cisco_toolkit.design_advisor import compute_design_blueprint
    snap = _snap()
    snap["l3_forwarding"] = [{"switch": "d", "vlan": "10", "svi_ip": "10.0.10.1"},
                             {"switch": "d", "vlan": "20", "svi_ip": "10.0.20.1"}]
    snap["design_blueprint"] = compute_design_blueprint(
        snap, {"address_space": "10.20.0.0/16", "vlan_zones": {10: "PCI", 20: "corp"}})
    out = str(tmp_path / "d.docx")
    write_design_doc_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "Per-zone summarization" in text and "PCI" in text
    """V3.23.153: LLD §3.5 — version sprawl per model is flagged with the most widely deployed
    image as the standardization candidate; Past-EoS gets refresh planning, not a lost-support claim."""
    snap = _snap()
    # second + third 2960X: two devices on 15.2(7)E, the original acc1 on 15.2 → MIXED, candidate 15.2(7)E
    for n in ("acc2", "acc3"):
        snap["devices"][n] = {"hostname": n, "model": "WS-C2960X-48FPD-L", "serial_number": f"S{n}",
                              "sw_version": "15.2(7)E", "platform": "ios"}
    out = str(tmp_path / "d.docx")
    write_design_doc_docx(out, snap, "Unit Test Fleet")
    d = Document(out)
    h2 = [p.text for p in d.paragraphs if p.style.name == "Heading 2"]
    assert any(t == "3.5 Software plan & recommendations" for t in h2), h2
    text = _all_text(d)
    assert "MIXED — standardize on 15.2(7)E" in text          # majority image is the candidate
    assert "validate it against Cisco's published recommended release" in text
    assert "past-end-of-sale date band needs a planned refresh decision" in text
    assert "does not prove lost support entitlement" in text and "WS-C2960X-48FPD-L" in text


def test_design_software_plan_separates_past_ldos_replacement_from_past_eos_refresh(tmp_path):
    snap = _snap()
    snap["devices"] = {
        "old": {"model": "WS-C4948E", "sw_version": "1"},
        "sale": {"model": "WS-C2960X-48FPD-L", "sw_version": "2"},
    }
    snap["lifecycle_risk"]["per_device"] = [
        {"host": "old", "model": "WS-C4948E", "band": "Past-LDoS"},
        {"host": "sale", "model": "WS-C2960X-48FPD-L", "band": "Past-EoS"},
    ]
    out = str(tmp_path / "mixed-lifecycle.docx")
    write_design_doc_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "past last-day-of-support gets replacement" in text
    assert "WS-C4948E" in text
    assert "past-end-of-sale date band needs a planned refresh decision" in text
    assert "date band alone does not prove lost support entitlement" in text


def test_design_software_plan_consistent_fleet_message(tmp_path):
    """A fleet with one image per model and no past-EoS hardware gets the carry-forward message."""
    snap = _snap()
    snap["lifecycle_risk"]["per_device"][2]["band"] = "Active"   # acc1 no longer Past-EoS
    out = str(tmp_path / "d.docx")
    write_design_doc_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "carry the current images forward as the minimum-version baseline" in text
    assert "MIXED" not in text


def test_design_failsoft_without_python_docx(monkeypatch, tmp_path):
    import builtins, os
    real_import = builtins.__import__

    def _blocked(name, *a, **k):
        if name == "docx" or name.startswith("docx."):
            raise ImportError("simulated missing python-docx")
        return real_import(name, *a, **k)

    monkeypatch.setattr(builtins, "__import__", _blocked)
    out = str(tmp_path / "d.docx")
    write_design_doc_docx(out, _snap(), "Unit Test Fleet")   # must not raise
    assert not os.path.exists(out)


def test_design_vlan_inventory_is_canonical_single_source():
    """Single source of truth: the design doc's VLAN inventory must BE analyze.vlan_inventory — the
    same derivation the CRD reads — so 'VLANs in use' agrees across deliverables, and an L3-only VLAN
    (a gateway SVI with no access port) is counted, not silently dropped by an access-port-only tally."""
    from cisco_toolkit.analyze import vlan_inventory
    from cisco_toolkit.design import _vlan_inventory
    snap = {
        "interfaces": {
            "acc1": {"Gi1/0/1": {"switchport_mode": "Access", "vlan": "10",
                                 "end_host_mac": "aaaa.0000.0001"}},
        },
        # VLAN 40 has an L3 gateway but sits on NO access port -> only the canonical inventory sees it
        "l3_forwarding": [{"switch": "core1", "vlan": "40", "svi_ip": "10.0.40.1"}],
    }
    canonical = vlan_inventory(snap)
    assert {v for v, _ in canonical} == {10, 40}        # access 10 + L3-only 40
    assert _vlan_inventory(snap) == canonical           # design delegates to the ONE derivation (no drift)


def test_design_bom_labels_past_eos_as_refresh_not_replace(tmp_path):
    """REVIEW #3: a Past-EoS device with LDoS still future and entitlement unassessed must appear
    under REFRESH (label names past-EoS), never
    in a 'Replace (past-LDoS)' row -- the §5.1 BoM must not call it past-LDoS/replace-now."""
    from cisco_toolkit.design_advisor import compute_design_blueprint
    snap = _snap()                                       # acc1 is Past-EoS WS-C2960X-48FPD-L; no Past-LDoS
    snap["design_blueprint"] = compute_design_blueprint(snap)   # §5.1 reads the canonical blueprint
    out = str(tmp_path / "d.docx")
    write_design_doc_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "Refresh (near-LDoS / past-EoS)" in text     # refresh label now names past-EoS
    assert "Replace (past-LDoS)" not in text            # no replace row (the fleet has 0 Past-LDoS)
    assert "0 Near-LDoS asset(s)" in text
    assert "1 Past-EoS asset(s) to place in refresh planning while LDoS is still future" in text
    assert "approaching it to refresh" not in text


def test_design_bom_renders_inventory_residue_with_no_lifecycle_rows(tmp_path):
    """The HLD must retain both named devices and anonymous inventory residue when the axis emits no rows."""
    from cisco_toolkit.design_advisor import compute_design_blueprint
    snap = _snap()
    snap["devices"] = {
        "a": {"hostname": "a", "model": "MODEL-A", "serial_number": "SER-A"},
        "b": {"hostname": "b", "model": "MODEL-B", "serial_number": "SER-B"},
    }
    snap["collection_completeness"] = {
        "summary": {"inventory": 4, "complete": 4, "partial": 0, "not_collected": 0}}
    snap["lifecycle_risk"] = {"per_device": []}
    snap["design_blueprint"] = compute_design_blueprint(snap)
    bom = snap["design_blueprint"]["target_state"]["replacement_bom"]
    assert bom["n_undetermined"] == 4 and bom["n_not_assessed"] == 4
    assert ["(lifecycle row missing)", 2] in bom["undetermined"]

    out = str(tmp_path / "no-lifecycle-rows.docx")
    write_design_doc_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "5.1 Replacement Bill of Materials" in text
    assert "A further 4 asset(s) could NOT be banded" in text
    assert "MODEL-A" in text and "MODEL-B" in text
    assert "(lifecycle row missing)" in text
    assert "UNDETERMINED — resolve before procurement" in text


def test_design_53_renders_when_candidate_but_no_subnets(tmp_path):
    """REVIEW #9: a candidate addressing plan that allocated NO subnets (supernet too small / overflow) must
    still render §5.3 with its 'enlarge the address_space' note, not be silently dropped."""
    from cisco_toolkit.design_advisor import compute_design_blueprint
    snap = _snap()
    snap["design_blueprint"] = compute_design_blueprint(snap, {"address_space": "10.0.0.0/30"})
    ap = snap["design_blueprint"]["target_state"]["addressing_plan"]
    assert ap["status"] == "candidate" and not ap.get("subnets") and not ap.get("requirement_needed")
    out = str(tmp_path / "d.docx")
    write_design_doc_docx(out, snap, "Unit Test Fleet")
    paras = [p.text for p in Document(out).paragraphs]
    assert any("Net-new IP addressing plan" in p for p in paras)
    assert any("smaller than a /24" in p or "enlarge" in p.lower() for p in paras)


def test_design_fhrp_counts_credit_a_clean_redundant_fabric(tmp_path):
    """DECK_-01: the FHRP candidate/configured counts must come from the FULL gateway register
    (l3_forwarding), NOT the PROBLEMS-ONLY snap['fhrp'] list -- else a cleanly-redundant multi-gateway VLAN
    (absent from snap['fhrp']) is never counted and a fully-redundant fabric is misrepresented as 0 candidates /
    0 with FHRP. Here snap['fhrp']=[] (no inconsistencies) yet VLAN 10 has 2 gateways running HSRP."""
    snap = _snap()
    snap["fhrp"] = []
    snap["l3_forwarding"] = [
        {"switch": "core1", "vlan": "10", "svi_ip": "10.0.10.1", "fhrp": "HSRP active"},
        {"switch": "core2", "vlan": "10", "svi_ip": "10.0.10.2", "fhrp": "HSRP standby"},
        {"switch": "core1", "vlan": "20", "svi_ip": "10.0.20.1", "fhrp": "none"},
    ]
    out = str(tmp_path / "d.docx")
    write_design_doc_docx(out, snap, "Unit Test Fleet")
    rows = {}
    for t in Document(out).tables:
        for row in t.rows:
            cells = [c.text for c in row.cells]
            if len(cells) == 2:
                rows[cells[0]] = cells[1]
    cand = next(v for k, v in rows.items() if "Multi-gateway VLANs (FHRP candidates)" in k)
    cfg = next(v for k, v in rows.items() if "first-hop redundancy (FHRP) configured" in k)
    assert cand == "1", rows     # VLAN 10 has 2 gateways -> 1 candidate (was 0 from len(snap['fhrp']))
    assert cfg == "1", rows      # VLAN 10 runs HSRP -> 1 configured (was 0)


def test_design_doc_robust_to_non_dict_decision(tmp_path):
    """DECK_-03: a single non-dict element in design_blueprint.decisions (corrupt/hand-edited snapshot) must
    not abort the whole As-Built Design Document with AttributeError (deck.py/crd.py already isinstance-guard)."""
    import os
    snap = _snap()
    snap["design_blueprint"] = {"decisions": ["junk", {"status": "recommended", "title": "ok",
                                "priority": "High", "domain": "x"}], "tradeoff_scorecard": []}
    out = str(tmp_path / "d.docx")
    write_design_doc_docx(out, snap, "Unit Test Fleet")   # must not raise
    assert os.path.getsize(out) > 0


def test_design_doc_tolerates_non_dict_executive_brief(tmp_path):
    """[multi-domain audit L4] design.py used 'snap.get("executive_brief") or {}' -- a TRUTHY non-dict (a string in
    a malformed/slimmed snapshot) slipped through and crashed .get('scale'), where the fail-soft siblings degrade.
    The builder must not raise."""
    out = str(tmp_path / "d.docx")
    write_design_doc_docx(out, {"executive_brief": "oops", "devices": {"sw1": {}}}, "Test")   # must not raise
    import os
    assert os.path.exists(out)


def test_design_doc_survives_xml_illegal_chars(tmp_path):
    """A device-derived Unicode noncharacter (U+FFFE/U+FFFF) or lone surrogate (U+D800-DFFF) passes through but
    aborts the WHOLE design DOCX at XML serialization ('All strings must be XML compatible') -- the same class the
    excel + runbook generators were hardened against. Every written string must be sanitized. Exercises the
    direct-paragraph path (keystone host, like the repro), the docmeta table path (model cell), a CDP neighbour
    name, and the label arg."""
    import os
    snap = _snap()
    bad = chr(0xFFFF) + chr(0xFFFE) + chr(0xD800)   # two noncharacters + a lone surrogate (ASCII-safe source)
    snap.setdefault("failure_impact", []).insert(
        0, {"host": "core1" + bad, "stranded": 99, "severity": "High", "vlans_impacted": 3, "detail": "VLAN 10" + bad})
    snap["interfaces"]["core1"]["Po1"]["cdp_neighbor"] = "dist1" + bad     # direct topology text
    snap["devices"]["core1"]["model"] = "N9K" + bad                       # inventory TABLE cell (docmeta)
    out = str(tmp_path / "d.docx")
    write_design_doc_docx(out, snap, "Meridian" + bad)                          # must NOT raise
    assert os.path.exists(out)
    Document(out)                                                         # and be a valid, openable .docx


# --- Truthy-non-dict section hardening (the `or {}` / `or []` falsy-guard class) ----------------------
# A section whose value is a TRUTHY non-dict / non-list scalar (a corrupt or slimmed --no-collect
# snapshot) slips the `or {}` / `or []` guards and is then .items()-ed / .get()-ed / iterated, aborting
# the whole As-Built Design DOCX -> HTTP 500 on the webapp /design route + a CLI exception. Same class
# already hardened in analyze.vlan_inventory, the executive_brief read, stp_roots and the decisions loop.
# Each malformed section must degrade to EMPTY (coverage-honest), never crash.

def _poison_base():
    """Minimal snapshot that still drives every crash site under test: one device, so the §3.1 inventory,
    §3.4 BoM, §3.5 software and §6 NOS-mix loops all run and every guarded site is actually reached."""
    return {"devices": {"sw1": {"hostname": "sw1", "model": "C9300-48T", "platform": "ios",
                                "sw_version": "17.9"}}}


@pytest.mark.parametrize("section, value", [
    ("routing_neighbors", 5),                    # _is_l3: (routing_neighbors or {}).get(host)
    ("routing_neighbors", {"sw1": 5}),           # §2.3: rn.items() then (d or {}).get(proto) on a scalar per-host value
    ("interfaces", 5),                           # _segmentation_facts + LLD §3.3: (interfaces or {}).items()/.get(h)
    ("interfaces", {"sw1": 5}),                  # scalar ports row -> (ports or {}).items()
    ("interfaces", {"sw1": {"Vlan10": 5}}),      # scalar port-detail -> d.get('vrf') / d.get('svi_ip')
    ("fhrp", 5),                                 # §2.4 resilience: for g in fhrp: g.get('issues')
    ("lifecycle_risk", 5),                       # lifecycle.get('per_device')
    ("endpoint_identity", 5),                    # §6: [e for e in endpoint_identity ...] (scalar not iterable)
])
def test_design_tolerates_truthy_non_dict_section(tmp_path, section, value):
    """A truthy non-dict/non-list section must degrade to empty, never abort the design DOCX."""
    import os
    snap = _poison_base()
    snap[section] = value
    out = str(tmp_path / "d.docx")
    write_design_doc_docx(out, snap, "Poison")          # must not raise
    assert os.path.exists(out)


def test_design_tolerates_scalar_device_value(tmp_path):
    """The devices-value class: a device whose value is a truthy non-dict is dereferenced with .get() in
    FOUR loops -- §3.1 inventory (devices[h]), §3.4 BoM + §3.5 software ((d or {}).get), and §6 NOS-mix
    (devices.values()). Guarding only one of the four leaves the other three crashing."""
    import os
    out = str(tmp_path / "d.docx")
    write_design_doc_docx(out, {"devices": {"sw1": 5}}, "Poison")   # must not raise
    assert os.path.exists(out)


def test_design_tolerates_every_in_scope_section_non_dict_at_once(tmp_path):
    """Kitchen-sink: every section in the hardened crash class a truthy scalar at once -- the builder must
    still emit an openable .docx (l3_forwarding/capacity are already _R-guarded; the rest are fixed here)."""
    import os
    snap = {"devices": {"sw1": 5}, "interfaces": 5, "routing_neighbors": 5, "fhrp": 5,
            "lifecycle_risk": 5, "endpoint_identity": 5, "l3_forwarding": 5, "capacity": 5}
    out = str(tmp_path / "d.docx")
    write_design_doc_docx(out, snap, "Poison")          # must not raise
    assert os.path.exists(out)
    Document(out)                                        # ...and it opens as a valid document


# --- Comprehensive crash-class sweep: EVERY top-level section x EVERY poison shape --------------------
# The falsy-guard (`or {}` / `or []`) crash class is closed section-by-section only if EVERY section the
# generator dereferences is proven against EVERY malformed shape a corrupt/slimmed snapshot can carry.
# The four poison shapes are chosen to reach different guards: a scalar / a string (iterable! a bare
# `for x in section` would walk its chars) / a list-of-scalars (slips a `_R`-less row loop) / a
# dict-with-a-scalar-VALUE ({"k":1} -- survives the OUTER `_D` and reaches the INNER deref, e.g.
# stp_roots->vmap.items(), redistribution->len(v); a top-level-only guard is NOT enough).

_ALL_DESIGN_SECTIONS = [
    "devices", "interfaces", "l3_forwarding", "routing_neighbors", "stp_roots", "redistribution",
    "fhrp", "capacity", "lifecycle_risk", "vpc", "failure_impact", "punchlist", "subnet_intelligence",
    "service_map", "executive_brief", "design_blueprint", "qos_audit", "endpoint_identity",
    "multicast_intelligence",   # §2.5 on-air authority census (mi.get("summary") / mi.get("groups"))
]
_DESIGN_POISONS = [5, "x", [1, 2], {"k": 1}]


@pytest.mark.parametrize("poison", _DESIGN_POISONS)
@pytest.mark.parametrize("section", _ALL_DESIGN_SECTIONS)
def test_design_poison_sweep_every_section_x_every_poison(tmp_path, section, poison):
    """Every section the As-Built Design DOCX reads, set to each malformed shape, must degrade to empty
    and still emit an OPENABLE .docx -- never abort the whole deliverable (HTTP 500 on /design + a CLI
    exception). 18 sections x 4 shapes = 72 cases. Non-vacuity was verified out-of-band by reverting the
    guards: stp_roots/redistribution/vpc/punchlist/subnet_intelligence/qos_audit crash without them,
    incl. redistribution+{"k":1} -> TypeError(len(scalar)) and stp_roots+{"k":1} -> vmap.items()."""
    import os
    snap = {"devices": {"sw1": {}}}
    snap[section] = poison
    out = str(tmp_path / "d.docx")
    write_design_doc_docx(out, snap, "l")   # must not raise
    assert os.path.exists(out)
    Document(out)                            # ...and open as a valid document


def test_design_wellformed_still_renders_newly_guarded_sections(tmp_path):
    """Golden-unchanged: routing the newly-hardened sections through _D/_R is IDENTITY on well-formed
    input, so a well-formed snapshot must STILL render the STP root placement (stp_roots) and the
    redistribution boundaries (redistribution) -- the two sections no prior test pinned. An over-broad
    guard that emptied them would fail here rather than silently ship a blank section. (punchlist and
    subnet_intelligence golden are already pinned by test_design_reconciles_to_snapshot.)"""
    out = str(tmp_path / "d.docx")
    write_design_doc_docx(out, _snap(), "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "core1 roots 2 VLAN(s)" in text                                  # stp_roots -> §2.2 root placement
    assert "Redistribution boundaries" in text and "route-redistribution edge" in text  # redistribution -> §2.3
    assert "1 route-redistribution edge" in text                            # the one edge in _snap() is counted


# --- NESTED falsy-guard poison: the crash class the top-level sweep CANNOT reach ---------------------
# The sweep above sets snap[section] = poison. For a CONTAINER section that is enough, but for a section
# the generator only reads THROUGH (design_blueprint, service_map.multicast, qos_audit, lifecycle_risk,
# punchlist rows) it is VACUOUS: with design_blueprint = {"k": 1} every inner bp.get(...) returns None,
# and `or {}` / `or []` DOES catch None -- so every §4/§5 gate goes falsy and the guarded code never runs.
#
# The real stored availability DoS lives one to three levels DOWN: a WELL-FORMED blueprint whose
# `decisions` / `doctrine.<domain>` / `target_state.addressing_plan.subnets` (…) value is a TRUTHY
# non-dict / non-list. It survives `or {}` / `or []`, survives the webapp's only validation
# (isinstance(snap, dict) and "devices" in snap) and survives xml_safe_deep (int/float/bool/None pass
# through unchanged), is STORED by the accepted POST, and then aborts every later
# GET /api/snapshots/{id}/deliverable/design with an HTTP 500.
#
# So the base below is a FULLY-POPULATED snapshot that satisfies every render gate, and exactly ONE
# nested value is poisoned per case (poisoning several at once lets an upstream gate skip the row and
# silently re-vacuates the test).

def _rich_design_snap():
    """A snapshot that reaches EVERY gated render site the nested sweep poisons: the full
    design_blueprint (summary / scorecard / decisions / coverage / doctrine / target_state and all five
    target-state sub-plans), plus multicast, qos_audit, lifecycle_risk. Pinned as non-vacuous by
    test_rich_design_snap_reaches_every_nested_guard_site."""
    return {
        "devices": {"sw1": {"hostname": "sw1", "model": "C9300-48T", "platform": "ios",
                            "sw_version": "17.9"}},
        "lifecycle_risk": {"per_device": [{"host": "sw1", "model": "C9300-48T", "band": "Past-EoS"}]},
        "service_map": {"multicast": {
            "active_switch_count": 1, "active_interfaces": 2,
            "classified_groups": [{"group": "224.0.1.129", "name": "PTP-primary",
                                   "category": "Broadcast-AV"}],
            "igmp_queriers": [{"switch": "sw1", "vlan": "10"}],
            "ptp": {"sw1": {"operational": True}}}},
        # §2.5's on-air authority census, in the shape compute_multicast_intelligence publishes it.
        "multicast_intelligence": {
            "groups": [{"group": "224.0.1.129", "name": "PTP-primary", "category": "Broadcast-AV",
                        "on_air": True, "on_air_authoritative": False,
                        "semantics_authoritative": False, "overlay_status": "curated-only"}],
            "summary": {"n_groups": 1, "n_av_groups": 1, "n_av_groups_authoritative": 0}},
        "qos_audit": {
            "summary": {"n_devices": 1, "n_assessable": 1, "modes": {"trust-dscp": 1},
                        "n_voice_ports": 3},
            "findings": [{"host": "sw1", "label": "no ingress marking", "severity": "Medium"}]},
        "design_blueprint": {
            "summary": {"headline": "Two-tier campus, single-gateway exposure"},
            "tradeoff_scorecard": [{"label": "Availability", "score": 2, "posture": "Fair",
                                    "evidence": "one gateway on VLAN 20"}],
            "decisions": [
                {"status": "recommended", "priority": "P1", "domain": "L3", "title": "Add FHRP",
                 "driver": "single point of failure", "recommended_action": "HSRP pair",
                 "alternatives": "anycast gateway", "tradeoffs": "cost vs availability",
                 "evidence": {"summary": "VLAN 20 has one SVI", "fields": ["l3_forwarding"],
                              "devices": ["sw1"]},
                 "principle": {"id": "CCDE-AV-1", "title": "First-hop redundancy",
                               "citation": "Cisco Campus Design Guide"}},
                {"status": "needs-requirement", "title": "What is the RTO?",
                 "requirements_needed": ["recovery time objective"], "axes": ["availability"]},
            ],
            "coverage": {"caveat": "1 of 1 device collected"},
            "doctrine": {"Campus": [{"title": "Redundant first hop", "recommended_action": "HSRP/VRRP",
                                     "citation": "CVD", "engine_actionable": True}]},
            "target_state": {
                "summary": {"headline": "Collapsed core with redundant gateways"},
                "dimensions": [{"area": "L3", "current": "single gateway", "target": "HSRP pair",
                                "rationale": "availability", "confidence": "Medium"}],
                "replacement_bom": {"n_replace": 1, "n_refresh": 1, "note": "indicative",
                                    "replace_now": [["WS-C2960X-48FPD-L", 2]],
                                    "refresh_soon": [["C9300-48T", 1]]},
                "segmentation_plan": {"observed": "single global table", "target": "two VRFs",
                                      "requirement_needed": "zone list"},
                "addressing_plan": {"status": "candidate", "supernet": "10.20.0.0/16",
                                    "n_allocated": 1, "note": "candidate only", "mode": "zone-aware",
                                    "zones": [{"zone": "PROD", "summary": "10.20.0.0/20",
                                               "n_vlans": 1}],
                                    "subnets": [{"vlan": "10", "hosts": 5,
                                                 "subnet": "10.20.10.0/24", "note": "sized"}]},
                "wave_plan": {"n_move_groups": 1, "largest_group": 1, "n_waves": 1, "wave_cap": 5,
                              "note": "candidate", "waves": [{"wave": 1, "kind": "L2-coupled",
                                                              "n_switches": 1}]},
                "scope_note": "1 device in scope",
                "coverage": {"caveat": "target state is a proposal"},
            },
        },
    }


def _set_path(obj, path, value):
    """Set one dotted path inside a nested dict/list structure; an integer segment indexes a list."""
    keys = path.split(".")
    cur = obj
    for k in keys[:-1]:
        cur = cur[int(k)] if isinstance(cur, list) else cur[k]
    last = keys[-1]
    if isinstance(cur, list):
        cur[int(last)] = value
    else:
        cur[last] = value
    return obj


# Every nested read the design writer dereferences after a `or {}` / `or []` (or a bare subscript /
# absent-only `.get(k, [])` default, which the falsy guard does not even attempt). Grouped by cluster.
_NESTED_DESIGN_POISON_PATHS = [
    # --- design_blueprint.* (§1 headline, §4 decisions/doctrine, §5 target state) ---
    "design_blueprint.summary",                                   # §1 _sm.get("headline")
    "design_blueprint.decisions",                                 # §4 list-comp + build_design_traceability
    "design_blueprint.decisions.0.evidence",                      # §4.2 (evidence or {}).get("summary")
    "design_blueprint.decisions.0.evidence.fields",               # traceability: join over ev["fields"]
    "design_blueprint.decisions.0.evidence.devices",              # traceability: ev["devices"][:8] slice
    "design_blueprint.decisions.0.principle",                     # §4.2 (principle or {}).get("citation")
    "design_blueprint.decisions.1.requirements_needed",           # §4.3 ", ".join(...)
    "design_blueprint.decisions.1.axes",                          # §4.3 ", ".join(...)
    "design_blueprint.tradeoff_scorecard",                        # §4.1 `for s in sc`
    "design_blueprint.tradeoff_scorecard.0",                      # §4.1 per-element s.get(...)
    "design_blueprint.coverage",                                  # §4 cov.get("caveat")
    "design_blueprint.doctrine",                                  # §4.4 doctrine.values()
    "design_blueprint.doctrine.Campus",                           # §4.4 len(v) / `for it in v`
    "design_blueprint.doctrine.Campus.0",                         # §4.4 per-element it.get(...)
    "design_blueprint.target_state",                              # §5 ts.get("dimensions")
    "design_blueprint.target_state.summary",                      # §5 (summary or {}).get("headline")
    "design_blueprint.target_state.dimensions",                   # §5 `for d in ts_dims`
    "design_blueprint.target_state.dimensions.0",                 # §5 per-element d.get(...)
    "design_blueprint.target_state.replacement_bom",              # §5.1 bom.get("n_replace")
    "design_blueprint.target_state.replacement_bom.replace_now",  # §5.1 `for m, q in bom.get(k, [])`
    "design_blueprint.target_state.replacement_bom.refresh_soon",  # §5.1 same, absent-only default
    "design_blueprint.target_state.segmentation_plan",            # §5.2 segp.get("observed")
    "design_blueprint.target_state.addressing_plan",              # §5.3 ap.get("status")
    "design_blueprint.target_state.addressing_plan.zones",        # §5.3 bare `for z in ap["zones"]`
    "design_blueprint.target_state.addressing_plan.zones.0",      # §5.3 per-element z.get(...)
    "design_blueprint.target_state.addressing_plan.subnets",      # §5.3 bare ap["subnets"][:60] + len()
    "design_blueprint.target_state.addressing_plan.subnets.0",    # §5.3 per-element s.get(...)
    "design_blueprint.target_state.wave_plan",                    # §5.4 wp.get("waves")
    "design_blueprint.target_state.wave_plan.waves",              # §5.4 bare wp["waves"][:40]
    "design_blueprint.target_state.wave_plan.waves.0",            # §5.4 per-element w.get(...)
    "design_blueprint.target_state.coverage",                     # §5 (coverage or {}).get("caveat")
    # --- service_map.multicast.* (§2.5) ---
    "service_map.multicast",                                      # mc.get("classified_groups")
    "service_map.multicast.classified_groups",                    # len(groups) / groups[:15]
    "service_map.multicast.classified_groups.0",                  # per-element g.get(...)
    "service_map.multicast.igmp_queriers",                        # len(queriers)
    "service_map.multicast.ptp",                                  # ptp.values() / len(ptp)
    # --- multicast_intelligence.* (§2.5 classification-basis column + census note) ---
    "multicast_intelligence",                                     # mi.get("summary") / mi.get("groups")
    "multicast_intelligence.summary",                             # `"n_av..." in s` / s.get(...)
    "multicast_intelligence.groups",                              # `for g in groups`
    "multicast_intelligence.groups.0",                            # per-element g.get(...)
    # --- qos_audit.* (§2.7) ---
    "qos_audit.summary",                                          # qsum.get("n_devices")
    "qos_audit.summary.modes",                                    # qmodes.items()
    "qos_audit.findings",                                         # (findings or [])[:12]
    "qos_audit.findings.0",                                       # per-element f.get(...)
    # --- lifecycle_risk.per_device (feeds §3.1 EoL band + §3.4 BoM) ---
    "lifecycle_risk.per_device",                                  # `for r in (per_device or [])`
    "lifecycle_risk.per_device.0",                                # per-element r.get("host")
]


@pytest.mark.parametrize("poison", _DESIGN_POISONS)
@pytest.mark.parametrize("path", _NESTED_DESIGN_POISON_PATHS)
def test_design_tolerates_nested_truthy_non_container(tmp_path, path, poison):
    """One nested value set to a truthy non-dict / non-list must degrade to EMPTY and still emit an
    openable .docx — never abort the whole As-Built Design Document (a stored HTTP 500 on the webapp
    /deliverable/design route, and a CLI exception)."""
    import os
    snap = _set_path(_rich_design_snap(), path, poison)
    out = str(tmp_path / "d.docx")
    # This call used to be wrapped in a try/except that xfail-ed exactly one case
    # (service_map.multicast.igmp_queriers = 5), because the identical falsy-guard bug lived OUT OF UNIT
    # in cisco_toolkit/analyze.py::vlan_inventory -- the ONE canonical VLAN derivation that
    # design.py:_vlan_inventory delegates to and crd.py shares. analyze.py is guarded now, so that
    # except-branch is dead and the tolerance is removed: every path/poison is asserted directly.
    write_design_doc_docx(out, snap, "Poison")          # must not raise
    assert os.path.exists(out)
    Document(out)                                        # ...and open as a valid document


@pytest.mark.parametrize("poison", _DESIGN_POISONS)
def test_design_tolerates_nested_poison_punchlist_devices(tmp_path, poison):
    """§4 punch-list fallback (rendered only when design_blueprint carries NO decisions): a row whose
    'devices' is a truthy non-list crashes the [:4] slice / len(). deck.py already fixed this exact
    shape; the design doc was the remaining surface."""
    import os
    snap = {"devices": {"sw1": {"model": "C9300-48T"}},
            "punchlist": [{"severity": "Critical", "category": "L3 design",
                           "title": "VLAN 20 has a single gateway", "devices": poison}]}
    out = str(tmp_path / "d.docx")
    write_design_doc_docx(out, snap, "Poison")          # must not raise
    assert os.path.exists(out)
    Document(out)


def test_rich_design_snap_reaches_every_nested_guard_site(tmp_path):
    """ANTI-VACUITY + golden-unchanged, in one pin. The nested sweep is only meaningful if the
    UN-poisoned base actually executes every guarded site, so this asserts each gated section renders
    AND that its well-formed values survive the coercers unchanged (as_list/as_dict are identity on a
    well-formed list/dict). An over-broad guard that emptied a section, or a gate that stopped firing,
    fails here instead of silently re-vacuating the sweep."""
    out = str(tmp_path / "d.docx")
    write_design_doc_docx(out, _rich_design_snap(), "Rich Fleet")
    d = Document(out)
    text = _all_text(d)
    heads = [p.text for p in d.paragraphs if p.style.name in ("Heading 1", "Heading 2", "Heading 3")]
    for token in ("2.5 Multicast & timing design", "2.7 Quality of service (configured posture)",
                  "4.1 Design trade-off scorecard", "4.2 Recommended target-state design decisions",
                  "4.3 Open design questions (requirements to confirm)",
                  "4.4 Design doctrine applied (CCDE-grounded reference)", "Campus (1)",
                  "4.5 Design traceability matrix",
                  "5. Proposed Target-State Architecture (candidate)",
                  "5.1 Replacement Bill of Materials (target procurement)",
                  "5.2 Target segmentation", "5.3 Net-new IP addressing plan (candidate)",
                  "5.4 Candidate migration wave plan"):
        assert any(h == token for h in heads), f"gate never fired: {token}; have {heads}"
    for token in (
        "Two-tier campus, single-gateway exposure",   # bp.summary.headline           (§1)
        "PTP-primary",                                # multicast.classified_groups   (§2.5)
        "1 IGMP-snooping querier",                    # multicast.igmp_queriers len() (§2.5)
        "CURATED — unverified",                       # multicast_intelligence.groups  (§2.5 basis cell)
        "NONE of them classified on-air",             # multicast_intelligence.summary (§2.5 census note)
        "1 of 1 device(s) report an operational",     # multicast.ptp values()/len()  (§2.5)
        "1 device(s) trust-dscp",                     # qos_audit.summary.modes       (§2.7)
        "no ingress marking",                         # qos_audit.findings            (§2.7)
        "Past-EoS",                                   # lifecycle_risk.per_device     (§3.1/§3.4)
        "Availability",                               # tradeoff_scorecard            (§4.1)
        "Add FHRP",                                   # decisions (recommended)       (§4.2)
        "VLAN 20 has one SVI",                        # decisions[].evidence.summary  (§4.2)
        "Cisco Campus Design Guide",                  # decisions[].principle.citation(§4.2)
        "recovery time objective",                    # decisions[].requirements_needed (§4.3)
        "availability",                               # decisions[].axes              (§4.3)
        "1 of 1 device collected",                    # bp.coverage.caveat            (§4)
        "Redundant first hop",                        # doctrine.<domain>[]           (§4.4)
        "l3_forwarding",                              # traceability evidence.fields  (§4.5)
        "Collapsed core with redundant gateways",     # target_state.summary.headline (§5)
        "HSRP pair",                                  # target_state.dimensions       (§5)
        "WS-C2960X-48FPD-L",                          # replacement_bom.replace_now   (§5.1)
        "Refresh (near-LDoS / past-EoS)",             # replacement_bom.refresh_soon  (§5.1)
        "single global table",                        # segmentation_plan.observed    (§5.2)
        "10.20.0.0/20",                               # addressing_plan.zones         (§5.3)
        "10.20.10.0/24",                              # addressing_plan.subnets       (§5.3)
        "L2-coupled",                                 # wave_plan.waves               (§5.4)
        "target state is a proposal",                 # target_state.coverage.caveat  (§5)
    ):
        assert token in text, f"well-formed value lost after guarding: {token!r}"


def test_design_deliverable_route_does_not_500_on_nested_poison(tmp_path):
    """E2E through the REAL stored-DoS route. The upload is accepted (201) and stored verbatim — the
    only validation is isinstance(snap, dict) and 'devices' in snap — so a nested truthy non-container
    that raises inside write_design_doc_docx makes EVERY later GET of the design deliverable a 500
    (webapp/backend/deliverables.py re-raises after unlinking the temp file; app.py maps it to 500).
    The route must answer 200 with a real .docx instead."""
    import json
    pytest.importorskip("fastapi")
    pytest.importorskip("httpx")
    from fastapi.testclient import TestClient

    from webapp.backend.app import create_app

    snap = _set_path(_rich_design_snap(),
                     "design_blueprint.target_state.addressing_plan.subnets", 5)
    app = create_app(db_path=str(tmp_path / "e2e.db"))
    # base_url=localhost so the DNS-rebinding Host allowlist passes; raise_server_exceptions=False so a
    # stored DoS shows up as a 500 STATUS (how a real client experiences it) rather than a raised error.
    with TestClient(app, base_url="http://localhost", raise_server_exceptions=False) as c:
        cid = c.post("/api/campaigns", json={"name": "c"}).json()["id"]
        r = c.post(f"/api/campaigns/{cid}/snapshots",
                   files={"file": ("s.json", json.dumps(snap).encode(), "application/json")},
                   data={"label": "poison"})
        assert r.status_code == 201, r.text          # accepted + STORED first -> a stored DoS
        sid = r.json()["id"]
        g = c.get(f"/api/snapshots/{sid}/deliverable/design")
    assert g.status_code == 200, f"stored DoS: {g.status_code} {g.text[:400]}"
    assert g.content[:2] == b"PK"                    # a real .docx (zip), not an error page


# --- leaf-type class: a non-str leaf reaching a string operation ------------------------------------
# DISTINCT from the truthy-non-container class: `l3_forwarding[]` is a well-formed list of well-formed
# dicts and `risk` is PRESENT -- it is the LEAF type. `(x or "")` does NOT coerce a non-str, so
# `"single-gateway" in 5` raises `TypeError: argument of type 'int' is not iterable`. The snapshot is
# attacker-controllable (upload -> deliverables.generate("design", snap) -> HTTP 500).
@pytest.mark.parametrize("bad", [5, 1.5, True, [1, 2], {"a": 1}],
                         ids=["int", "float", "bool", "list", "dict"])
def test_design_survives_nonstr_l3_risk_leaf(tmp_path, bad):
    """A non-str `risk` leaf must degrade, not crash the single-gateway tally."""
    snap = {"devices": {"sw1": {"hostname": "sw1"}},
            "l3_forwarding": [{"switch": "sw1", "vlan": "10", "risk": bad}]}
    out = tmp_path / "leafrisk.docx"
    write_design_doc_docx(str(out), snap, "Poison")      # must not raise
    assert out.is_file(), f"no design doc for risk={bad!r}"


def test_design_nonstr_risk_coercion_keeps_single_gateway_tally(tmp_path):
    """Non-vacuity companion: str() must be identity for well-formed risk strings -- a real
    single-gateway exposure must still be counted, so an over-broad guard fails here."""
    snap = {"devices": {"sw1": {"hostname": "sw1"}, "sw2": {"hostname": "sw2"}},
            "l3_forwarding": [{"switch": "sw1", "vlan": "10", "risk": "single-gateway exposure"},
                              {"switch": "sw2", "vlan": "20", "risk": "single-gateway exposure"}]}
    out = tmp_path / "wfrisk.docx"
    write_design_doc_docx(str(out), snap, "WF")
    assert out.is_file()
    txt = "\n".join(p.text for p in Document(str(out)).paragraphs)
    txt += "\n".join(c.text for t in Document(str(out)).tables for r in t.rows for c in r.cells)
    assert "2" in txt, "the single-gateway tally (2) must survive the str() coercion"


_BOM_HEADING = "5.1 Replacement Bill of Materials"


def _hld_text_for_bands(tmp_path, per_device):
    from cisco_toolkit.design_advisor import compute_design_blueprint
    snap = _snap()
    # This helper models a fully enumerated lifecycle census. Keep the device owner census aligned with
    # the synthetic host rows; missing-row behavior has its own explicit regression below.
    snap["devices"] = {
        row["host"]: {"hostname": row["host"], "model": row.get("model") or "Unknown"}
        for row in per_device
    }
    snap["collection_completeness"] = {"summary": {
        "inventory": len(per_device), "complete": len(per_device), "partial": 0, "not_collected": 0,
    }}
    snap["lifecycle_risk"] = {"per_device": per_device}
    snap["design_blueprint"] = compute_design_blueprint(snap)
    out = str(tmp_path / "d.docx")
    write_design_doc_docx(out, snap, "Unit Test Fleet")
    return _all_text(Document(out))


def test_hld_bom_section_appears_for_an_entirely_UNDETERMINED_fleet(tmp_path):
    """Absence-as-health at document level: §5.1 used to VANISH when nothing could be banded.

    `_replacement_bom` bucketed only `past-ldos` -> replace_now and the refresh bands -> refresh_soon.
    `Unknown` -- no exact EoX row OR a row whose source/date authority was withheld -- fell through
    both, so an undetermined fleet produced n_replace=0 and n_refresh=0, the §5.1 gate was false, and the whole
    procurement section was OMITTED. The reader could not tell that from a fleet needing no
    procurement at all: an omitted section leaves nothing on the page to disagree with.

    Asserted at DOCUMENT level, not on the bucket dict, because the producer's new key is inert
    unless the consumer's heading gate and row builder both carry it through.
    """
    text = _hld_text_for_bands(tmp_path, [
        {"host": "a", "band": "Unknown", "model": "WS-C6509-E"},
        {"host": "b", "band": "Unknown", "model": "WS-C6509-E"},
        {"host": "c", "band": "Unknown", "model": "WS-C3560-48PS"}])
    assert _BOM_HEADING in text, "the procurement section is omitted for an all-undetermined fleet"
    assert "WS-C6509-E" in text and "WS-C3560-48PS" in text, "undetermined models not listed"
    assert "could NOT be banded" in text, "the reason for the gap is not disclosed"
    assert "UNDETERMINED — resolve before procurement" in text, "rows not labelled as undetermined"
    assert "no exact EoX row matched" in text
    assert "source/date authority was withheld" in text
    assert "offline KB matched" not in text


def test_hld_bom_section_stays_absent_for_a_fully_assessed_clean_fleet(tmp_path):
    """Non-vacuity: the heading gate must still be a gate, not always-true.

    Keyed on the §5.1 HEADING rather than the word "UNDETERMINED", which occurs elsewhere in the
    document -- asserting on the bare word would pass here for the wrong reason.
    """
    text = _hld_text_for_bands(tmp_path, [{"host": "a", "band": "Active", "model": "C9300-48P"}])
    assert _BOM_HEADING not in text, "an empty procurement section was rendered for a clean fleet"


def test_hld_bom_keeps_real_replacements_leading_the_undetermined_rows(tmp_path):
    """A real past-LDoS asset must still be costed and named; the coverage bucket must not absorb it."""
    text = _hld_text_for_bands(tmp_path, [
        {"host": "a", "band": "Past-LDoS", "model": "WS-C4948E"},
        {"host": "b", "band": "Unknown", "model": "WS-C6509-E"}])
    assert _BOM_HEADING in text and "WS-C4948E" in text and "WS-C6509-E" in text
    assert "1 Past-LDoS asset(s) to replace now" in text, "the real replacement was miscounted"
    assert "A further 1 asset(s) could NOT be banded" in text


def test_design_software_plan_discloses_unknown_alongside_real_recommendation(tmp_path):
    snap = _snap()
    snap["devices"] = {
        "old": {"model": "WS-C4948E", "sw_version": "1"},
        "blind": {"model": "WS-C6509-E", "sw_version": "2"},
    }
    snap["lifecycle_risk"] = {"per_device": [
        {"host": "old", "band": "Past-LDoS", "model": "WS-C4948E"},
        {"host": "blind", "band": "Unknown", "model": "WS-C6509-E"},
    ]}
    out = str(tmp_path / "mixed-known-unknown.docx")
    write_design_doc_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "past last-day-of-support gets replacement" in text
    assert "NOT ASSESSED — 1 device(s) across 1 platform(s)" in text
    assert "Either no exact EoX row matched" in text
    assert "retained source/date authority was withheld or incomplete" in text


def test_design_model_band_uses_canonical_near_ldos_precedence_over_past_eos(tmp_path):
    snap = _snap()
    snap["devices"] = {
        "near": {"model": "SHARED", "sw_version": "1"},
        "eos": {"model": "SHARED", "sw_version": "1"},
    }
    snap["lifecycle_risk"] = {"per_device": [
        {"host": "near", "band": "Near-LDoS", "model": "SHARED"},
        {"host": "eos", "band": "Past-EoS", "model": "SHARED"},
    ]}
    out = str(tmp_path / "band-order.docx")
    write_design_doc_docx(out, snap, "Unit Test Fleet")
    doc = Document(out)
    rows = [[c.text for c in row.cells] for table in doc.tables for row in table.rows]
    model_rows = [r for r in rows if r and r[0] == "SHARED"]
    assert model_rows and all("Near-LDoS" in r[-1] for r in model_rows)


# =====================================================================================================
# §2.5 multicast classification AUTHORITY (review r9 EXIT E).
#
# analyze.compute_multicast_intelligence classifies a group as on-air / broadcast-AV from the offline
# port registry's CURATED media semantics. On the SHIPPED pack the authoritative on-air count is ZERO --
# every multicast row is curated -- and that same curated flag ESCALATES a mac-alias finding from Medium
# to High (analyze.py:2774). The HLD's §2.5 table rendered (group, name, category) with no authority
# column and no note, so a curated hint reached the reader in the same voice as an IANA-assigned fact.
#
# These pin the DISCLOSURE (not a re-scoring): the group ADDRESS + its observed activity are evidence,
# the NAME/CATEGORY are registry semantics, and the per-group basis + the fleet census say which.
# =====================================================================================================
_BASIS_HEADER = "Classification basis"
_BASIS_NOTE_LEAD = "Classification basis — the Group address and its observed multicast activity"


def _real_multicast_snap(groups=("224.0.1.129", "239.1.1.1", "239.255.255.250"), authoritative=()):
    """A snapshot whose multicast sections come from the REAL producers, not a hand-shaped dict.

    `compute_service_map(igmp_groups=...)` classifies the addresses through the SHIPPED offline port
    registry (so `classified_groups` carries the registry's own authority labels verbatim), and
    `compute_multicast_intelligence` derives the per-group `on_air_authoritative` + the
    `summary.n_av_groups_authoritative` census from it. `authoritative` names the addresses whose
    registry semantics are flipped to authoritative BEFORE the intelligence pass runs, so the census is
    still computed by the producer -- that is the only way to obtain the authoritative fleet the shipped
    pack cannot produce (every multicast row in it is curated)."""
    from cisco_toolkit.analyze import compute_multicast_intelligence, compute_service_map
    sm = compute_service_map({}, {}, igmp_groups=list(groups))
    mc = sm["multicast"]
    for g in mc["classified_groups"]:
        if g.get("group") in authoritative:
            g["semantics_authoritative"] = True
            g["semantics_source"] = "IANA multicast-addresses registry"
    mc["active_switch_count"], mc["active_interfaces"] = 2, 4
    mc["igmp_queriers"] = [{"switch": "core1", "vlan": "10", "querier": "10.0.10.1"}]
    mc["ptp"] = {"core1": {"operational": True}}
    return {"devices": {"core1": {"hostname": "core1", "model": "C9300-48T", "platform": "ios"}},
            "service_map": sm,
            "multicast_intelligence": compute_multicast_intelligence(sm)}


def _design_text(snap, tmp_path, name="d.docx"):
    out = str(tmp_path / name)
    write_design_doc_docx(out, snap, "Multicast Fleet")
    return _all_text(Document(out))


def test_design_multicast_table_discloses_the_curated_classification_basis(tmp_path):
    """Shipped-pack reality: every classified group is CURATED, so the §2.5 table must say so per row
    and the note must state the fleet census. Pre-fix the table was (Group, Name, Category) only, and
    the words 'CURATED' / 'Classification basis' appeared nowhere in the document."""
    snap = _real_multicast_snap()
    # The producer's own verdict, asserted first so this test cannot pass against a changed fixture.
    assert snap["multicast_intelligence"]["summary"]["n_av_groups_authoritative"] == 0
    assert snap["multicast_intelligence"]["summary"]["n_av_groups"] == 1
    text = _design_text(snap, tmp_path)
    assert _BASIS_HEADER in text, "§2.5 still renders (group, name, category) with no authority column"
    assert "CURATED — unverified" in text, "a curated classification is not labelled as curated"
    assert _BASIS_NOTE_LEAD in text, "the address-vs-semantics distinction is not stated"
    assert ("NONE of them classified on-air by an authoritative source: the broadcast/AV label is a "
            "CURATED offline-registry classification, not a measurement") in text, \
        "the zero-authoritative census is not disclosed in the wording runbook.py already uses"
    assert "224.0.1.129" in text and "PTP-primary" in text     # the row itself still renders


def test_design_multicast_unmatched_group_is_labelled_unclassified_not_curated(tmp_path):
    """239.1.1.1 has NO registry entry (overlay_status 'unclassified'), so its generic Name/Category are
    not a curation anybody performed. Labelling it 'CURATED' would claim a judgement that never
    happened; it must read UNCLASSIFIED."""
    snap = _real_multicast_snap()
    by_group = {g["group"]: g for g in snap["multicast_intelligence"]["groups"]}
    assert by_group["239.1.1.1"]["overlay_status"] == "unclassified"      # producer's verdict
    assert by_group["239.255.255.250"]["overlay_status"] == "curated-only"
    text = _design_text(snap, tmp_path)
    assert "UNCLASSIFIED — no registry match" in text
    assert "CURATED — unverified" in text            # ...and the curated-only row is still CURATED


def test_design_multicast_basis_is_not_assessed_when_the_census_is_absent(tmp_path):
    """FAIL-CLOSED. An older snapshot carries classified_groups but no `multicast_intelligence` at all.
    Absent must read NOT ASSESSED -- never 'authoritative', never silently omitted, never coerced to the
    (assertive) CURATED verdict by a key-presence guard plus a falsy default."""
    snap = _real_multicast_snap()
    snap.pop("multicast_intelligence")
    text = _design_text(snap, tmp_path)
    assert "NOT ASSESSED" in text
    assert ("on-air classification authority NOT ASSESSED in this snapshot; treat the broadcast/AV "
            "label as curated, not authoritative") in text
    assert "AUTHORITATIVE (registry source)" not in text, "an absent census was rendered as authority"


def test_design_multicast_basis_is_not_assessed_when_a_group_has_no_authority_keys(tmp_path):
    """Per-GROUP fail-closed: the census exists, but this group's record predates the authority keys.
    It must degrade to NOT ASSESSED on its own row while its siblings keep their real verdict."""
    snap = _real_multicast_snap()
    for g in snap["multicast_intelligence"]["groups"]:
        if g["group"] == "224.0.1.129":
            g.pop("on_air_authoritative")
            g.pop("semantics_authoritative")
    text = _design_text(snap, tmp_path)
    assert "NOT ASSESSED" in text                     # the stripped row
    assert "CURATED — unverified" in text        # ...siblings unaffected


@pytest.mark.parametrize("bad", ["unknown", "false", {}, [], 1, 0, None])
def test_design_multicast_basis_fails_closed_on_a_malformed_authority_flag(tmp_path, bad):
    """The OTHER fail-open direction. A bare truthiness read of a malformed flag ("unknown", a dict, a
    list) reports AUTHORITATIVE — the worst error available here. The producer publishes real booleans,
    so a non-bool is UNUSABLE, and an unusable verdict is NOT ASSESSED, not a verdict."""
    snap = _real_multicast_snap()
    for g in snap["multicast_intelligence"]["groups"]:
        if g["group"] == "224.0.1.129":
            g["on_air_authoritative"] = bad
            g["semantics_authoritative"] = bad
    text = _design_text(snap, tmp_path)
    assert "AUTHORITATIVE (registry source)" not in text, f"{bad!r} was read as authority"
    assert "NOT ASSESSED" in text, f"{bad!r} did not degrade to NOT ASSESSED"


def test_design_multicast_authoritative_fleet_does_NOT_acquire_the_caveat(tmp_path):
    """NON-VACUITY 1. A fleet whose on-air classification IS authoritative must be reported as such --
    the disclosure is not an unconditional string bolted onto every document."""
    snap = _real_multicast_snap(authoritative=("224.0.1.129",))
    s = snap["multicast_intelligence"]["summary"]
    assert (s["n_av_groups"], s["n_av_groups_authoritative"]) == (1, 1)   # producer's verdict
    text = _design_text(snap, tmp_path)
    assert "AUTHORITATIVE (registry source)" in text
    assert "1 of 1 on-air classification(s) rest on an authoritative source" in text
    assert "NONE of them classified on-air by an authoritative source" not in text
    assert "authority NOT ASSESSED in this snapshot" not in text


def test_design_fleet_with_no_media_estate_gains_no_multicast_caveat(tmp_path):
    """NON-VACUITY 2. A fabric with no multicast at all still gets the existing fallback sentence and
    NO classification-basis text -- the caveat may not leak into a document that classifies nothing."""
    snap = {"devices": {"sw1": {"hostname": "sw1", "model": "C9300-48T"}},
            "service_map": {"multicast": {"active_interfaces": 0, "active_switch_count": 0,
                                          "classified_groups": [], "igmp_queriers": [], "ptp": {}}}}
    text = _design_text(snap, tmp_path)
    assert "No multicast/PTP activity was classified" in text
    for token in (_BASIS_HEADER, _BASIS_NOTE_LEAD, "CURATED — unverified",
                  "AUTHORITATIVE (registry source)", "UNCLASSIFIED — no registry match"):
        assert token not in text, f"a non-media fabric acquired multicast authority text: {token!r}"


def test_design_multicast_incoherent_census_refuses_to_state_a_ratio(tmp_path):
    """authoritative > total is incoherent (the two counts are published independently). Rendering
    '7 of 3' is indistinguishable from a real ratio, so the incoherence is disclosed instead -- the same
    check runbook.py carries because an earlier fix printed a NEGATIVE count into a client workbook."""
    snap = _real_multicast_snap()
    snap["multicast_intelligence"]["summary"].update(n_av_groups=3, n_av_groups_authoritative=7)
    text = _design_text(snap, tmp_path)
    assert "census INCOHERENT: 7 classification(s) reported as authoritative out of 3 on-air group(s)" in text
    assert "7 of 3" not in text


def test_av_authority_mirrors_the_runbook_wording():
    """design._av_authority is a deliberate MIRROR of runbook._av_authority (documented in both
    docstrings). Drive both over one case table so a divergence fails here instead of shipping two
    different authority claims about the same fact in one deliverable set."""
    from cisco_toolkit.design import _av_authority as design_av
    from cisco_toolkit.runbook import _av_authority as runbook_av
    cases = [
        None, {}, 5, "x", {"summary": 5}, {"summary": {}},
        {"summary": {"n_av_groups": 3, "n_av_groups_authoritative": 0}},
        {"summary": {"n_av_groups": 3, "n_av_groups_authoritative": 2}},
        {"summary": {"n_av_groups": 3, "n_av_groups_authoritative": 3}},
        {"summary": {"n_av_groups": 3, "n_av_groups_authoritative": 7}},
        {"summary": {"n_av_groups": 3, "n_av_groups_authoritative": None}},
        {"summary": {"n_av_groups": 3, "n_av_groups_authoritative": "lots"}},
        {"summary": {"n_av_groups": 3, "n_av_groups_authoritative": -4}},
        {"summary": {"n_av_groups": None, "n_av_groups_authoritative": 2}},
    ]
    for case in cases:
        assert design_av(case) == runbook_av(case), f"wording diverged on {case!r}"
    # ...and the table is not vacuous: the branches really do produce different strings.
    assert len({runbook_av(c) for c in cases}) >= 4


@pytest.mark.parametrize("bad", _DESIGN_POISONS + [None])
def test_design_multicast_authority_tolerates_a_malformed_census(tmp_path, bad):
    """A malformed `multicast_intelligence` (or a malformed `groups` list inside it) must degrade to
    NOT ASSESSED and still emit an openable document -- never claim authority, never crash."""
    import os
    snap = _real_multicast_snap()
    snap["multicast_intelligence"] = bad
    out = str(tmp_path / "d.docx")
    write_design_doc_docx(out, snap, "Poison")
    assert os.path.exists(out)
    text = _all_text(Document(out))
    assert "AUTHORITATIVE (registry source)" not in text
    assert "NOT ASSESSED" in text
