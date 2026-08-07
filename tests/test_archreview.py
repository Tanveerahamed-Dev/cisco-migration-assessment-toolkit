"""NEW-V3.23.160: the Architecture Review & Conformance Report — the automated senior-engineer
design review. Tests pin the check verdict logic (the engineering judgment), the not-assessable
honesty (absent evidence is never scored), the rollup math, and the DOCX structure/furniture.
The compute is pure-stdlib so the verdict tests run without python-docx; only the writer tests
skip when the optional dep is absent."""
import pytest

from cisco_toolkit.archreview import compute_architecture_review, V_RANK


def _snap():
    """A two-tier fleet with deliberate, known deviations: VLAN 20 single-gateway (critical),
    acc1 single-homed + single-PSU (deviations), a daisy-chained acc2, VLAN 1 native trunk,
    a 48:1 oversubscribed acc1, mixed 2960X images, and one past-EoS device."""
    return {
        "script_version": "V3.23.0",
        "devices": {
            "core1": {"hostname": "core1", "model": "C9500-24Y4C", "sw_version": "17.9",
                      "num_power_supplies": 2},
            "dist1": {"hostname": "dist1", "model": "C9300-48T", "sw_version": "17.9",
                      "num_power_supplies": 2},
            "acc1": {"hostname": "acc1", "model": "WS-C2960X-48FPD-L", "sw_version": "15.2",
                     "num_power_supplies": 1},
            "acc2": {"hostname": "acc2", "model": "WS-C2960X-48FPD-L", "sw_version": "15.2(7)E",
                     "num_power_supplies": 2},
        },
        "interfaces": {
            "core1": {"Vlan10": {"svi_ip": "10.0.10.1", "hsrp_behavior": "HSRP grp10 active",
                                 "acl_in": "PROD_IN", "dhcp_helpers": "10.9.9.9"},
                      "Vlan20": {"svi_ip": "10.0.20.1"},
                      "Te1/1": {"switchport_mode": "Trunk", "cdp_neighbor": "dist1",
                                "trunk_allowed_vlans": "10,20", "trunk_native_vlan": "999"}},
            "dist1": {"Vlan10": {"svi_ip": "10.0.10.2", "hsrp_behavior": "HSRP grp10 standby"},
                      "Te1/0/1": {"switchport_mode": "Trunk", "cdp_neighbor": "core1",
                                  "trunk_allowed_vlans": "10,20", "trunk_native_vlan": "999"},
                      "Gi1/0/48": {"switchport_mode": "Trunk", "cdp_neighbor": "acc1.corp.example",
                                   "trunk_allowed_vlans": "1-4094", "trunk_native_vlan": "1"}},
            # acc1: 48 Gi access ports vs ONE Gi uplink -> 48:1 oversubscription; single-homed.
            "acc1": dict(
                {"Gi1/0/1": {"switchport_mode": "Trunk", "cdp_neighbor": "dist1",
                             "trunk_allowed_vlans": "1-4094", "trunk_native_vlan": "1"}},
                **{f"Gi1/0/{i}": {"switchport_mode": "Access", "vlan": "10",
                                  "stp_bpduguard": "Yes",
                                  "end_host_mac": f"aaaa.0000.{i:04x}"} for i in range(2, 50)},
            ),
            # acc2 daisy-chains off acc1 (access-to-access link) — the HIER-2 deviation.
            "acc2": {"Gi1/0/1": {"switchport_mode": "Trunk", "cdp_neighbor": "acc1",
                                 "trunk_allowed_vlans": "20", "trunk_native_vlan": "999"},
                     "Gi1/0/5": {"switchport_mode": "Access", "vlan": "20",
                                 "stp_bpduguard": "Yes", "end_host_mac": "bbbb.0000.0001"}},
        },
        "l3_forwarding": [
            {"switch": "core1", "vlan": "10", "svi_ip": "10.0.10.1", "fhrp": "HSRP active", "risk": ""},
            {"switch": "dist1", "vlan": "10", "svi_ip": "10.0.10.2", "fhrp": "HSRP standby", "risk": ""},
            {"switch": "core1", "vlan": "20", "svi_ip": "10.0.20.1", "fhrp": "", "risk": "single-gateway"},
        ],
        "routing_neighbors": {"core1": {"ospf": [{"neighbor": "10.0.0.2", "state": "FULL"}],
                                        "eigrp": [], "bgp": []}},
        "redistribution": {},
        "stp_roots": {"core1": {"10": {"is_root": True}, "20": {"is_root": True}}},
        "fhrp": [{"vid": "10", "members": [{"host": "core1"}, {"host": "dist1"}], "issues": []}],
        "capacity": [{"hostname": "acc1", "total_ports": 48, "active_ports": 47, "port_util": 97.9}],
        "failure_impact": [{"host": "core1", "stranded": 12, "vlans_impacted": 2}],
        "lifecycle_risk": {
            "per_device": [{"host": "acc1", "model": "WS-C2960X-48FPD-L", "band": "Past-EoS"}],
            "summary": {"n_devices": 4, "n_past_eos": 1, "n_past_ldos": 0, "n_near": 0},
        },
        "security": {"core1": {"findings": [], "summary": {"fail": 0, "pass": 9, "na": 0,
                                                           "grade": "hardened"}},
                     "acc1": {"findings": [], "summary": {"fail": 3, "pass": 5, "na": 1,
                                                          "grade": "weak"}}},
        "config_hygiene": {"acc1": {"undefined": [], "unused": [],
                                    "summary": {"undefined": 2, "unused": 1, "structures": 9}}},
        "operational_drift": [],
        "collection_completeness": {"summary": {"complete": 4, "partial": 0, "not_collected": 0},
                                    "devices": []},
        "punchlist": [],
    }


def _check(ar, cid):
    return next(c for c in ar["checks"] if c["id"] == cid)


def test_review_shape_and_rollup():
    ar = compute_architecture_review(_snap())
    assert len(ar["domains"]) == 8
    assert ar["summary"]["n_checks"] == len(ar["checks"]) >= 20
    assert all(c["verdict"] in V_RANK for c in ar["checks"])
    # every check carries the full senior-engineer block
    for c in ar["checks"]:
        for k in ("observed", "implication", "recommendation", "reference"):
            assert c[k], (c["id"], k)
    s = ar["summary"]
    assert s["n_assessable"] == s["n_checks"] - s["n_not_assessable"]
    assert 0 <= s["score_pct"] <= 100
    assert s["grade"] in ("A", "B", "C", "D", "F")
    assert "grade" in s["statement"].lower() or s["grade"] in s["statement"]


def test_archreview_interop_check_flags_multi_nos():
    """N37: a multi-NOS estate gets an OPS-4 interoperability check (advisory) telling the design to
    DECLARE its cross-platform dependency surface; a single-NOS fleet conforms."""
    ar = compute_architecture_review({"devices": {"a": {"platform": "ios"}, "b": {"platform": "ios"},
                                                  "c": {"platform": "nxos"}}})
    c = _check(ar, "OPS-4")
    assert c["verdict"] == "advisory"
    assert "interoperability" in c["title"].lower()
    assert "NOS famil" in c["observed"]
    # a single-NOS fleet conforms
    ar2 = compute_architecture_review({"devices": {"a": {"platform": "ios"}, "b": {"platform": "ios"}}})
    assert _check(ar2, "OPS-4")["verdict"] == "conforms"


def test_single_gateway_vlan_is_critical_and_tops_the_queue():
    ar = compute_architecture_review(_snap())
    res2 = _check(ar, "RES-2")
    assert res2["verdict"] == "critical"
    assert "20" in res2["observed"]
    # priority queue is severity-ordered: the critical item ranks first
    assert ar["top_actions"][0]["id"] in ("RES-2",)
    assert ar["top_actions"][0]["rank"] == 1


def test_daisy_chain_and_single_homing_detected():
    ar = compute_architecture_review(_snap())
    hier2 = _check(ar, "HIER-2")
    assert hier2["verdict"] == "deviation"
    assert "acc1" in hier2["observed"] and "acc2" in hier2["observed"]
    res1 = _check(ar, "RES-1")
    assert res1["verdict"] == "deviation"
    assert "acc1" in res1["evidence"]          # one in-fleet neighbour (dist1) = single-homed


def test_cdp_fqdn_neighbors_canonicalized():
    """dist1's downlink names 'acc1.corp.example' — the adjacency must still bind to acc1."""
    ar = compute_architecture_review(_snap())
    assert _check(ar, "RES-1")["verdict"] == "deviation"   # acc1 resolved, counted, single-homed


def test_oversubscription_math():
    snap = _snap()
    ar = compute_architecture_review(snap)
    cap1 = _check(ar, "CAP-1")
    assert cap1["verdict"] == "deviation"                  # 47 Gi edge / 1 Gi uplink ≈ 47:1
    assert "acc1" in cap1["observed"]
    # upgrade the uplink to TenGig -> 4.7:1, inside the 20:1 rule of thumb
    up = snap["interfaces"]["acc1"].pop("Gi1/0/1")
    snap["interfaces"]["acc1"]["Te1/0/1"] = up
    ar2 = compute_architecture_review(snap)
    assert _check(ar2, "CAP-1")["verdict"] == "conforms"


def test_stp_root_placement_judged_by_tier():
    snap = _snap()
    ar = compute_architecture_review(snap)
    assert _check(ar, "L2-1")["verdict"] == "conforms"     # roots on core1 (an L3 node)
    snap["stp_roots"] = {"acc2": {"20": {"is_root": True}}}
    ar2 = compute_architecture_review(snap)
    l21 = _check(ar2, "L2-1")
    assert l21["verdict"] == "deviation" and "acc2" in l21["observed"]


def test_vlan1_hygiene_levels():
    snap = _snap()
    ar = compute_architecture_review(snap)
    assert _check(ar, "L2-3")["verdict"] == "advisory"     # native-1 trunks only
    snap["interfaces"]["acc2"]["Gi1/0/6"] = {"switchport_mode": "Access", "vlan": "1"}
    ar2 = compute_architecture_review(snap)
    assert _check(ar2, "L2-3")["verdict"] == "deviation"   # user traffic ON vlan 1


def test_weak_cis_grade_is_a_security_deviation():
    ar = compute_architecture_review(_snap())
    sec1 = _check(ar, "SEC-1")
    assert sec1["verdict"] == "deviation" and "acc1" in sec1["evidence"]


def test_lifecycle_bands_map_to_verdicts():
    snap = _snap()
    assert _check(compute_architecture_review(snap), "LC-1")["verdict"] == "advisory"  # past-EoS
    snap["lifecycle_risk"] = {
        "per_device": [{"host": "near1", "band": "Near-LDoS"}],
        "summary": {"n_past_eos": 0, "n_past_ldos": 0, "n_near": 1},
    }
    assert _check(compute_architecture_review(snap), "LC-1")["verdict"] == "deviation"
    snap["lifecycle_risk"]["summary"]["n_past_ldos"] = 1
    snap["lifecycle_risk"]["per_device"] = [{"host": "ldos1", "band": "Past-LDoS"}]
    assert _check(compute_architecture_review(snap), "LC-1")["verdict"] == "critical"
    snap["lifecycle_risk"]["summary"] = {"n_past_eos": 0, "n_past_ldos": 0, "n_near": 0}
    assert _check(compute_architecture_review(snap), "LC-1")["verdict"] == "conforms"


def test_lc1_detail_omits_misleading_zero_past_eos():
    """Coverage-honesty (EoS/LDoS class, same as the campaign-trend fix): the lifecycle bands are
    EXCLUSIVE, so n_past_eos is the Past-EoS-ONLY count. Every Past-LDoS device is ALSO past end-of-sale,
    so when n_past_eos==0 the LC-1 detail must NOT print '(and 0 past end-of-sale)' (reads as 'nothing is
    past end-of-sale'); when there ARE EoS-only devices it says 'N more past end-of-sale'."""
    snap = _snap()
    snap["lifecycle_risk"]["summary"] = {"n_past_eos": 0, "n_past_ldos": 2, "n_near": 0}
    snap["lifecycle_risk"]["per_device"] = [{"host": "d1", "band": "Past-LDoS"},
                                            {"host": "d2", "band": "Past-LDoS"}]
    obs = _check(compute_architecture_review(snap), "LC-1")["observed"]
    assert "past LAST-DAY-OF-SUPPORT" in obs
    assert "0 past end-of-sale" not in obs and "and 0" not in obs
    snap["lifecycle_risk"]["summary"] = {"n_past_eos": 3, "n_past_ldos": 2, "n_near": 0}
    obs2 = _check(compute_architecture_review(snap), "LC-1")["observed"]
    assert "Separately, 3 device(s) are past end-of-sale with LDoS still future" in obs2


def test_lc1_mixed_bands_reserve_no_vendor_backstop_for_past_ldos_hosts():
    snap = _snap()
    snap["lifecycle_risk"] = {
        "summary": {"n_past_ldos": 1, "n_past_eos": 1, "n_near": 1, "n_unknown": 1},
        "per_device": [
            {"host": "ldos1", "band": "Past-LDoS"},
            {"host": "eos1", "band": "Past-EoS"},
            {"host": "near1", "band": "Near-LDoS"},
            {"host": "unknown1", "band": "Unknown"},
        ],
    }
    check = _check(compute_architecture_review(snap), "LC-1")
    assert "past LAST-DAY-OF-SUPPORT: ldos1" in check["observed"]
    assert "past end-of-sale with LDoS still future: eos1" in check["observed"]
    assert "within one year of LDoS: near1" in check["observed"]
    assert "1 device(s) could NOT be lifecycle-banded" in check["observed"]
    assert check["evidence"] == ["ldos1"]
    assert "For the Past-LDoS devices" in check["implication"]
    assert "Past-EoS-only date band does not establish contract entitlement" in check["implication"]


def test_hier2_not_assessable_without_neighbour_evidence():
    """Conforms-by-silence guard: HIER-2's evidence is the in-fleet CDP/LLDP adjacency map. With
    interfaces present but NO access-tier neighbour observed, it must grade 'not-assessable', never
    'conforms' (which would assert a star topology off absent evidence)."""
    snap = _snap()
    for ports in snap["interfaces"].values():
        for d in ports.values():
            if isinstance(d, dict):
                d.pop("cdp_neighbor", None); d.pop("lldp_neighbor", None)
    assert _check(compute_architecture_review(snap), "HIER-2")["verdict"] == "not-assessable"


def test_l2_3_not_assessable_without_vlan_evidence():
    """Conforms-by-silence guard: L2-3's evidence is the access-VLAN / trunk native-VLAN fields. With
    interfaces present but neither field captured anywhere, it must grade 'not-assessable', never
    'conforms' (which would assert VLAN-1 hygiene off absent evidence)."""
    snap = _snap()
    for ports in snap["interfaces"].values():
        for d in ports.values():
            if isinstance(d, dict):
                d.pop("vlan", None); d.pop("trunk_native_vlan", None)
    assert _check(compute_architecture_review(snap), "L2-3")["verdict"] == "not-assessable"


def _access_ports(n_total, n_captured, n_guarded):
    """`n_total` access ports on one switch: the first `n_captured` carry a BPDU-Guard field
    (`n_guarded` of those enabled, the rest explicitly disabled); the remainder carry NO field."""
    ports = {}
    for i in range(1, n_total + 1):
        d = {"switchport_mode": "Access", "vlan": "10"}
        if i <= n_captured:
            d["stp_bpduguard"] = "Yes" if i <= n_guarded else "disabled"
        ports[f"Gi1/0/{i}"] = d
    return {"devices": {"acc1": {}}, "interfaces": {"acc1": ports}}


def test_l2_2_conforms_only_when_every_access_port_is_captured_and_guarded():
    """L2-2's not-assessable gate counts ports whose BPDU-Guard state was CAPTURED, but the pass
    threshold compared `guarded` against ALL access ports — so CONFORMS ('the edge is protected
    against accidental switch insertion') was reached with half the fleet's edge unguarded or
    unevidenced. The cited rule is BPDU Guard on EVERY edge port."""
    # half the captured ports are positively UNGUARDED -> a deviation, never a pass
    c = _check(compute_architecture_review(_access_ports(100, 100, 50)), "L2-2")
    assert c["verdict"] == "advisory"
    assert "50" in c["observed"] and "do NOT carry it" in c["observed"]
    # every captured port is guarded but 40 ports carry NO evidence -> conforms-by-silence guard
    c2 = _check(compute_architecture_review(_access_ports(100, 60, 60)), "L2-2")
    assert c2["verdict"] == "not-assessable"
    assert "no BPDU-Guard evidence" in c2["observed"]
    assert "protected" not in c2["implication"]
    # full capture, every port guarded -> the only shape that conforms
    c3 = _check(compute_architecture_review(_access_ports(100, 100, 100)), "L2-2")
    assert c3["verdict"] == "conforms" and "All 100" in c3["observed"]
    # one unguarded port out of a fully-captured fleet still fails the EVERY-edge-port rule
    assert _check(compute_architecture_review(_access_ports(100, 100, 99)), "L2-2")["verdict"] == "advisory"


def test_l2_2_reads_the_producers_own_bpduguard_tokens():
    """Cross-module SSOT audit 2026-07-28: L2-2 hand-rolled its own token set and it did not contain
    the PRODUCER's own value. parse.py writes exactly "Enable"/"Disable" (and build.py promotes a
    global `portfast bpduguard default` to "Enable"), but the old test was
    `v not in ("no","disabled","off","false","-","--")` — "Disable" is not in that list, so an edge
    port with BPDU Guard EXPLICITLY OFF counted as GUARDED, while design_advisor counted the same
    port unguarded off the same field. Both consumers now share textutils.bpduguard_state."""
    from cisco_toolkit.textutils import bpduguard_state
    snap = _access_ports(10, 10, 10)
    for i in (1, 2, 3):                      # the producer's literal token, not the test's "disabled"
        snap["interfaces"]["acc1"][f"Gi1/0/{i}"]["stp_bpduguard"] = "Disable"
    c = _check(compute_architecture_review(snap), "L2-2")
    assert c["verdict"] == "advisory", c     # pre-fix: "conforms" — 3 disabled ports read as protected
    assert "3 of the 10" in c["observed"] and "do NOT carry it" in c["observed"]

    # and an UNRECOGNISED token is not evidence of protection either: it leaves the assessable set
    # rather than being counted as guarded (absence of evidence is never health).
    snap2 = _access_ports(10, 10, 10)
    snap2["interfaces"]["acc1"]["Gi1/0/1"]["stp_bpduguard"] = "??"
    assert bpduguard_state("??") is None
    c2 = _check(compute_architecture_review(snap2), "L2-2")
    assert c2["verdict"] == "not-assessable" and "no BPDU-Guard evidence" in c2["observed"]


def test_bpduguard_state_is_the_one_token_owner_for_both_consumers():
    """The two consumers must agree, port for port, on the producer's whole vocabulary."""
    from cisco_toolkit import design_advisor
    from cisco_toolkit.textutils import bpduguard_state
    assert (bpduguard_state("Enable"), bpduguard_state("Disable"), bpduguard_state("")) == (True, False, None)
    ports = {"Gi1/0/1": {"switchport_mode": "Access", "end_host_mac": "00:11:22:33:44:55",
                         "stp_bpduguard": "Disable"},
             "Gi1/0/2": {"switchport_mode": "Access", "end_host_mac": "00:11:22:33:44:66",
                         "stp_bpduguard": "Enable"}}
    snap = {"devices": {"acc1": {}}, "interfaces": {"acc1": ports}}
    sig = design_advisor._signals(snap)
    assert sig["bpdu_unguarded"] == 1 and sig["bpdu_not_assessed"] == 0
    c = _check(compute_architecture_review(snap), "L2-2")
    # archreview's assessable/guarded split must reconcile with design_advisor's verdict on the SAME
    # two ports: 2 captured, exactly 1 guarded (pre-fix archreview said 2 guarded, 0 unguarded).
    assert "1 of the 2 access port(s)" in c["observed"] and "(1 do)" in c["observed"]


def test_res_3_not_assessable_when_only_some_devices_report_psu_inventory():
    """RES-3's not-assessable gate was FLEET-WIDE (`no device reported a count`), so ONE device
    reporting 2 supplies graded the whole fleet CONFORMS — 'no single feed/PSU fault takes a switch
    down' asserted over 49 devices whose power inventory was never captured."""
    devices = {"dist1": {"num_power_supplies": 2}}
    devices.update({f"acc{i}": {"model": "WS-C2960X"} for i in range(1, 50)})
    c = _check(compute_architecture_review({"devices": devices}), "RES-3")
    assert c["verdict"] == "not-assessable"
    assert "1 of 50" in c["observed"]
    assert "No single feed" not in c["implication"]
    # full coverage still conforms
    every = {h: {"num_power_supplies": 2} for h in devices}
    assert _check(compute_architecture_review({"devices": every}), "RES-3")["verdict"] == "conforms"
    # a single-PSU device is still the deviation, and the partial coverage is disclosed with it
    devices["acc1"] = {"num_power_supplies": 1}
    c2 = _check(compute_architecture_review({"devices": devices}), "RES-3")
    assert c2["verdict"] in ("deviation", "advisory")
    assert "2 of 50" in c2["observed"] and "unassessed" in c2["observed"]


def test_mixed_images_flagged_per_platform():
    ar = compute_architecture_review(_snap())
    lc2 = _check(ar, "LC-2")
    assert lc2["verdict"] == "advisory" and "WS-C2960X-48FPD-L" in lc2["observed"]


def test_empty_snapshot_is_honestly_ungraded():
    """Absence of evidence must never score — an empty snapshot grades N/A, with every check
    not-assessable and a statement that says so (never a fabricated 'A')."""
    ar = compute_architecture_review({})
    s = ar["summary"]
    assert s["n_assessable"] == 0 and s["score_pct"] is None and s["grade"] == "N/A"
    assert all(c["verdict"] == "not-assessable" for c in ar["checks"])
    assert ar["top_actions"] == []
    assert "not-assessable" in s["statement"] or "too little evidence" in s["statement"]


def test_malformed_sections_degrade_not_crash():
    snap = _snap()
    snap["fhrp"] = "garbage"
    snap["stp_roots"] = {"core1": {"10": None, "20": "junk"}}
    snap["capacity"] = [{"port_util": "high"}, "junk"]
    snap["security"] = {"core1": "junk"}
    snap["lifecycle_risk"] = 7
    ar = compute_architecture_review(snap)              # must not raise
    assert ar["summary"]["n_checks"] >= 20


# ---------------------------------------------------------------------------
# Workbook scorecard sheet (V3.23.161) — openpyxl is a hard engine dep, no skip
# ---------------------------------------------------------------------------
def _sheet_text(ws):
    return "\n".join(str(c.value) for row in ws.iter_rows() for c in row if c.value is not None)


def test_workbook_scorecard_sheet_renders_the_review():
    from openpyxl import Workbook

    from cisco_toolkit.excel import ARCHREVIEW_SHEET_NAME, write_architecture_review_sheet

    wb = Workbook()
    write_architecture_review_sheet(wb, compute_architecture_review(_snap()))
    text = _sheet_text(wb[ARCHREVIEW_SHEET_NAME])
    assert "Conformance grade" in text
    for cid in ("HIER-1", "RES-2", "CAP-1", "SEC-1", "LC-1"):
        assert cid in text, cid
    assert "CRITICAL DEVIATION" in text                       # RES-2 single-gateway verdict
    assert "By domain" in text and "All checks" in text
    assert "Priority remediation queue" in text


def test_workbook_scorecard_sheet_handles_empty_review_and_rewrites():
    from openpyxl import Workbook

    from cisco_toolkit.excel import ARCHREVIEW_SHEET_NAME, write_architecture_review_sheet

    wb = Workbook()
    write_architecture_review_sheet(wb, compute_architecture_review({}))
    text = _sheet_text(wb[ARCHREVIEW_SHEET_NAME])
    assert "Conformance grade N/A" in text                    # honest, never a fabricated grade
    assert "Priority remediation queue" not in text           # nothing to remediate -> no section
    # idempotent: a re-run replaces the sheet instead of erroring or duplicating
    write_architecture_review_sheet(wb, compute_architecture_review(_snap()))
    assert wb.sheetnames.count(ARCHREVIEW_SHEET_NAME) == 1


# ---------------------------------------------------------------------------
# DOCX writer (optional python-docx, exactly like the other deliverables)
# ---------------------------------------------------------------------------
docx = pytest.importorskip("docx")
from docx import Document  # noqa: E402

from cisco_toolkit.archreview import write_archreview_docx  # noqa: E402


def _all_text(doc):
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            parts.extend(c.text for c in row.cells)
    return "\n".join(parts)


def test_archreview_docx_structure(tmp_path):
    out = str(tmp_path / "ar.docx")
    write_archreview_docx(out, _snap(), "Unit Test Fleet")
    d = Document(out)
    h1 = [p.text for p in d.paragraphs if p.style.name == "Heading 1"]
    for token in ("1. Executive Verdict", "2. Review Scope & Method", "3. Conformance Scorecard",
                  "4. Findings by Design Domain", "5. Hand-Off Into the Document Set",
                  "Document Control", "Document Acceptance"):
        assert any(t == token for t in h1), f"missing section: {token}; have {h1}"
    text = _all_text(d)
    assert "Conformance grade:" in text
    assert "CRITICAL DEVIATION" in text                  # the single-gateway verdict renders
    assert "20:1" in text                                # the oversubscription rule is cited
    assert "Priority remediation queue" in text
    # honesty furniture: the out-of-scope table is always present
    assert "QoS policy effectiveness" in text and "Wireless RF design" in text


def test_archreview_docx_carries_family_furniture(tmp_path):
    out = str(tmp_path / "ar.docx")
    write_archreview_docx(out, _snap(), "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "Revision history" in text and "Assumptions & caveats" in text
    assert "Customer network owner" in text                          # acceptance roles
    assert "As-Built Network Design Document (.docx)" in text        # related documents…
    assert "Architecture Review & Conformance Report (.docx)" not in text  # …excluding self


def test_archreview_docx_renders_w37_traceability_section(tmp_path):
    """[W3-7 follow-on] When the snapshot carries recommended design decisions, the Architecture Review gains a
    'Design Decision Traceability' section — the audit trail behind the conformance grade: each decision traced to
    its CCDE principle + published citation. COVERAGE-HONEST: a decision with no citation renders '(uncited)', never
    a fabricated reference; and Hand-Off renumbers to §6 (no numbering gap) only because §5 actually rendered."""
    snap = _snap()
    snap["design_blueprint"] = {"tradeoff_scorecard": [], "doctrine": {}, "decisions": [
        {"status": "recommended", "title": "Enforce SNMPv3 fleet-wide", "priority": "Critical", "domain": "Security",
         "principle": {"id": "mgmt-secure-protocols", "title": "Secure management protocols",
                       "citation": "CCDE Session 19"},
         "evidence": {"summary": "3 device(s) fail management hardening", "devices": ["core1"],
                      "fields": ["security[host].findings[].status"]}},
        {"status": "recommended", "title": "Collapse the access daisy-chain", "priority": "High", "domain": "Topology",
         "principle": {"id": "topo-no-daisy-chain", "title": "No access daisy-chains", "citation": ""},   # uncited
         "evidence": {"summary": "acc2 is single-homed via acc1", "devices": ["acc2"], "fields": []}},
    ]}
    out = str(tmp_path / "ar_trace.docx")
    write_archreview_docx(out, snap, "Unit Test Fleet")
    d = Document(out)
    h1 = [p.text for p in d.paragraphs if p.style.name == "Heading 1"]
    assert "5. Design Decision Traceability" in h1                    # the new section rendered
    assert "6. Hand-Off Into the Document Set" in h1                  # Hand-Off renumbered (no gap)
    text = _all_text(d)
    assert "Enforce SNMPv3 fleet-wide" in text and "mgmt-secure-protocols" in text   # decision -> principle traced
    assert "CCDE Session 19" in text                                  # the published citation surfaced
    assert "(uncited)" in text                                        # honest: the citation-less decision is flagged


def test_archreview_docx_prefers_attached_section(tmp_path):
    """One source of truth: when the CLI attached architecture_review, the writer renders THAT
    (here: a sentinel statement), rather than recomputing."""
    snap = _snap()
    snap["architecture_review"] = {
        "domains": [{"key": "Hierarchy & modularity", "verdict": "conforms", "score_pct": 100,
                     "checks": ["HIER-1"]}],
        "checks": [{"id": "HIER-1", "domain": "Hierarchy & modularity", "title": "t",
                    "verdict": "conforms", "observed": "o", "implication": "i",
                    "recommendation": "r", "reference": "ref", "evidence": []}],
        "top_actions": [],
        "summary": {"n_checks": 1, "n_assessable": 1, "n_conforms": 1, "n_advisory": 0,
                    "n_deviation": 0, "n_critical": 0, "n_not_assessable": 0, "score_pct": 100,
                    "grade": "A", "grade_label": "SENTINEL-ATTACHED-REVIEW",
                    "statement": "SENTINEL-ATTACHED-REVIEW"},
    }
    out = str(tmp_path / "ar.docx")
    write_archreview_docx(out, snap, "Unit Test Fleet")
    assert "SENTINEL-ATTACHED-REVIEW" in _all_text(Document(out))


def test_archreview_docx_renders_empty_snapshot(tmp_path):
    """An evidence-free snapshot still renders a valid (honestly N/A-graded) document."""
    out = str(tmp_path / "ar.docx")
    write_archreview_docx(out, {}, "Empty Fleet")
    text = _all_text(Document(out))
    assert "Conformance grade: N/A" in text
    assert "Checks not assessable from this snapshot" in text


def test_archreview_failsoft_without_python_docx(monkeypatch, tmp_path):
    import builtins, os
    real_import = builtins.__import__

    def _blocked(name, *a, **k):
        if name == "docx" or name.startswith("docx."):
            raise ImportError("simulated missing python-docx")
        return real_import(name, *a, **k)

    monkeypatch.setattr(builtins, "__import__", _blocked)
    out = str(tmp_path / "ar.docx")
    write_archreview_docx(out, _snap(), "Unit Test Fleet")   # must not raise
    assert not os.path.exists(out)


def test_trunk_allowed_none_is_not_classified_allow_all():
    """L2-5 inverted cry-wolf: a trunk allowed-VLAN list of 'none' allows NO VLANs -- the INVERSE of
    'allow ALL VLANs'. It must never be counted as an un-pruned allow-all trunk. With only a 'none' trunk
    present, the captured-list trunk is maximally pruned, so the check CONFORMS and the observed text never
    says 'allow ALL VLANs'."""
    snap = {"devices": {"SW1": {"platform": "ios"}},
            "interfaces": {"SW1": {"Gi1/0/1": {"switchport_mode": "trunk", "trunk_allowed_vlans": "none"}}}}
    c = _check(compute_architecture_review(snap), "L2-5")
    assert "allow ALL VLANs" not in c["observed"], c["observed"]
    assert c["verdict"] == "conforms", c["verdict"]


def _lc_snap(**summary):
    """A snapshot whose ONLY lifecycle signal is the summary band counts."""
    return {"devices": {f"sw{i}": {"platform": "ios"} for i in range(4)},
            "lifecycle_risk": {"summary": summary, "per_device": []}}


def _lc1(ar):
    return next(c for c in ar["checks"] if c["id"] == "LC-1")


def test_lc1_does_not_certify_a_fleet_whose_lifecycle_is_UNKNOWN():
    """An all-Unknown fleet must not read as 'conforms'.

    `analyze.py:6199` publishes `n_unknown` in the lifecycle summary, and the LC-1 chain
    (`if n_past_ldos / elif n_past_eos / elif n_near / else conforms`) never consulted it. So a fleet
    of Catalyst 6500s years past support — every one banded Unknown because no authoritative lifecycle
    band could be assigned — fell to the `else` and was certified:

        VERDICT: conforms — "Every device with lifecycle data is in an Active support band."
                            "Vendor support backs the whole migration."

    That reaches the Architecture Review DOCX, its workbook sheet, and the conformance grade. It is
    guardrail 3's exact wording ("Not observed" never silently becomes "healthy") in a signed
    deliverable, and it is the most consequential form of it: the sentence is VACUOUSLY true —
    zero devices have lifecycle data, so all zero of them are Active — while the verdict it carries
    is false.
    """
    ar = compute_architecture_review(_lc_snap(n_past_ldos=0, n_past_eos=0, n_near=0, n_unknown=4))
    f = _lc1(ar)
    assert f["verdict"] != "conforms", (
        f"LC-1 certified a fleet with 4 UNKNOWN-lifecycle devices: {f['verdict']} / {f['observed']}"
    )
    assert "unknown" in (f["observed"] + f["implication"]).lower(), (
        f"the unknown count is not disclosed to the reader: {f['observed']!r}"
    )


def test_lc1_still_conforms_when_every_device_IS_assessed_and_active():
    """Non-vacuity: a fully banded pre-EoS fleet still conforms, without claiming entitlement."""
    f = _lc1(compute_architecture_review(_lc_snap(n_past_ldos=0, n_past_eos=0, n_near=0, n_unknown=0)))
    assert f["verdict"] == "conforms", f"a fully-assessed Active fleet must still conform: {f}"
    rendered = f["observed"] + " " + f["implication"]
    assert "pre-EoS date band" in rendered
    assert "public schema value: Active" in rendered
    assert "support entitlement was not assessed" in rendered
    assert "Active support band" not in rendered
    assert "Vendor support backs" not in rendered


def test_lc1_unknown_wording_covers_no_match_and_provenance_withheld_states():
    f = _lc1(compute_architecture_review(
        _lc_snap(n_past_ldos=0, n_past_eos=0, n_near=0, n_unknown=2)))
    assert "either no exact EoX bulletin row matched" in f["observed"]
    assert "retained source authority/complete dates did not verify" in f["observed"]
    assert "no EoX bulletin in the offline KB covers" not in f["observed"]


def test_lc1_reports_the_real_finding_when_both_unknown_and_past_ldos_exist():
    """A real past-LDoS finding must not be displaced by the unknown-coverage branch."""
    f = _lc1(compute_architecture_review(_lc_snap(n_past_ldos=2, n_past_eos=0, n_near=0, n_unknown=3)))
    assert f["verdict"] not in ("conforms", "not-assessable"), f
    assert "2" in f["observed"], f"the past-LDoS count must lead: {f['observed']!r}"


def test_lc1_discloses_coverage_even_when_a_real_finding_fires():
    """The coverage disclosure was SUBORDINATED to the findings (`elif n_unknown`), so it fired only
    when every adverse count was zero -- covering the all-unbanded fleet and silently dropping the
    BROWNFIELD one, which is the fleet this instrument exists for.

    Measured before the fix: 1 Past-LDoS + 20 Unknown produced verdict `critical` reading
    "1 device(s) are past LAST-DAY-OF-SUPPORT: sw0." with the 20 undetermined devices never
    mentioned -- and LC-1 feeds domains[].score_pct, the conformance grade, the Architecture Review
    DOCX and its workbook sheet. What was FOUND and what could NOT BE ASSESSED are orthogonal facts.
    """
    for adverse in (dict(n_past_ldos=1), dict(n_past_eos=1), dict(n_near=1)):
        base = dict(n_past_ldos=0, n_past_eos=0, n_near=0, n_unknown=20)
        base.update(adverse)
        f = _lc1(compute_architecture_review(_lc_snap(**base)))
        assert f["verdict"] not in ("conforms", "not-assessable"), (adverse, f)
        assert "20 device(s) could NOT be lifecycle-banded" in f["observed"], (adverse, f["observed"])
        assert "UNKNOWN" in f["observed"], (adverse, f["observed"])

    # NON-VACUITY: with the same real finding and FULL coverage, no coverage clause may appear --
    # otherwise the disclosure is always-on and tells the reader nothing.
    clean_cov = _lc1(compute_architecture_review(
        _lc_snap(n_past_ldos=1, n_past_eos=0, n_near=0, n_unknown=0)))
    assert "could NOT be lifecycle-banded" not in clean_cov["observed"], clean_cov["observed"]
    assert clean_cov["verdict"] == "critical", clean_cov
