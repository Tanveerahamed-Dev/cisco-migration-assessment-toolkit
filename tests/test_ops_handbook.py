"""Tests for the Operations Handbook deliverable (NEW-V3.23.168, cisco_toolkit/ops.py) —
the PPDIOO Operate-phase document: evidence-derived baselines, evidence-gated sections,
family furniture, fail-soft behaviour."""
import json
import os

import pytest

docx = pytest.importorskip("docx")  # the deliverable needs the optional python-docx
from docx import Document  # noqa: E402

from cisco_toolkit.ops import write_ops_handbook_docx  # noqa: E402

_GOLDEN = os.path.join(os.path.dirname(__file__), "golden", "snapshot.json")


def _all_text(doc):
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            parts.extend(c.text for c in row.cells)
    return "\n".join(parts)


@pytest.fixture(scope="module")
def golden_snap():
    with open(_GOLDEN, encoding="utf-8") as f:
        return json.load(f)


def test_rich_snapshot_renders_evidence_derived_baselines(tmp_path, golden_snap):
    out = tmp_path / "ops.docx"
    write_ops_handbook_docx(str(out), golden_snap, "golden-fleet")
    doc = Document(str(out))
    text = _all_text(doc)
    heads = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    # the section skeleton
    for h in ("1. Purpose & Audience", "2. Network Quick Reference",
              "3. Monitoring & Alerting Baseline", "4. Operational Standards & Drift Control",
              "5. Software & Lifecycle Governance", "6. Backup & Recovery",
              "7. Known Issues & Operating Caveats", "8. Routine Operations Calendar",
              "9. Escalation & TAC Readiness"):
        assert any(h in x for x in heads), h
    # baselines really come from the golden fleet's own evidence
    assert "core1" in text                               # quick reference inventory
    assert "MAC address flapping" in text                # syslog axis detection (3.1)
    assert "Cisco PSIRT Software Checker" in text        # software governance (5)
    assert "single point in time" in text                # capacity honesty (3.2)
    assert "Right-click" in text                         # TOC field placeholder
    # furniture: document control + related docs excluding self + acceptance
    assert "Document control" in text or "Document Control" in text
    assert "Operations Handbook" in text
    assert "Architecture Review & Conformance Report" in text   # related-docs table
    assert "Acceptance" in text


def test_ops_routing_adjacency_and_fhrp_day2_section(tmp_path):
    """N36: the handbook carries a Day-2 routing-adjacency + first-hop-redundancy monitoring
    subsection — observed routing protocols get adjacency monitoring, and the FHRP coverage (or its
    absence) is a NAMED monitored item, not a silent gap."""
    snap = {
        "script_version": "V3.23.0", "devices": {"r1": {}},
        "protocol_health": [
            {"switch": "r1", "protocol": "OSPF", "severity": "Info", "summary": "1 nbr Full"},
            {"switch": "r1", "protocol": "EtherChannel", "severity": "High", "summary": "member down"}],
        "l3_forwarding": [{"switch": "r1", "vlan": "10", "svi_ip": "10.0.10.1", "fhrp": "none"},
                          {"switch": "r1", "vlan": "20", "svi_ip": "10.0.20.1", "fhrp": "none"}],
    }
    out = str(tmp_path / "ops.docx")
    write_ops_handbook_docx(out, snap, "Unit Test Fleet")
    d = Document(out)
    heads = [p.text for p in d.paragraphs if p.style.name.startswith("Heading")]
    text = _all_text(d)
    assert any("routing-adjacency" in h.lower() for h in heads), heads
    assert "OSPF" in text                                   # observed routing protocol monitored
    assert "0 of 2" in text                                 # FHRP coverage named: 0 of 2 gateways


def test_related_docs_exclude_self(tmp_path, golden_snap):
    out = tmp_path / "ops.docx"
    write_ops_handbook_docx(str(out), golden_snap, "x")
    doc = Document(str(out))
    rel = None
    for t in doc.tables:
        hdr = " ".join(c.text for c in t.rows[0].cells)
        if "Document" in hdr and "Role in the set" in hdr:
            rel = t
            break
    assert rel is not None
    names = "\n".join(c.text for row in rel.rows[1:] for c in row.cells)
    assert "Operations Handbook" not in names            # never lists itself
    assert "Executive Presentation Deck" in names        # the rest of the 12-doc family


def test_empty_snapshot_declares_absent_evidence(tmp_path):
    out = tmp_path / "ops_empty.docx"
    write_ops_handbook_docx(str(out), {}, "empty")
    doc = Document(str(out))
    text = _all_text(doc)
    # honest declarations, not invented baselines
    assert text.count("Not in this snapshot:") >= 3
    assert "re-run the collection" in text or "re-run the assessment" in text
    # the always-true content still renders
    assert "Routine Operations Calendar" in text
    assert "Daily" in text and "Quarterly" in text
    assert "P1" in text                                  # escalation skeleton


def test_malformed_sections_degrade_not_crash(tmp_path):
    snap = {
        "devices": {"sw1": None},                         # null device record
        "failure_impact": ["not-a-dict", {"host": "sw1", "stranded": "oops"}],
        "syslog_intelligence": "truthy-non-dict",
        "platform_health": {"summary": None, "per_device": "nope"},
        "golden_drift": {"summary": {"n_baseline": "x"}},
        "software_risk": [],
        "security": {"sw1": {"summary": {"fail": "many"}}},
    }
    out = tmp_path / "ops_bad.docx"
    write_ops_handbook_docx(str(out), snap, "bad")        # must not raise
    assert out.exists() and out.stat().st_size > 1000


@pytest.mark.parametrize("poison", [5, "boom", True, 3.14])
def test_ops_survives_truthy_scalar_platform_health_bands(tmp_path, poison):
    """§3.2's `pb = ph_sum.get("bands") or {}` was the file's last truthy-non-dict residue: `or` only
    catches FALSY, so a scalar `bands` reached `.items()` -> AttributeError. Its own twins already use
    the coercer (qos_audit.summary.modes, software_risk.summary.train_bands render the identical
    'Nx band' join). test_malformed_sections_degrade_not_crash cannot reach this line — it sets
    `platform_health.summary = None`, which is FALSY, so the `if ph_sum.get("n_collected")` gate at
    ops.py:390 short-circuits first; the summary here is deliberately TRUTHY and gate-passing."""
    snap = {
        "script_version": "V3.23.0", "devices": {"sw1": {}},
        "platform_health": {"summary": {"n_collected": 3, "bands": poison},
                            "per_device": [{"host": "sw1", "collected": True, "cpu_5min": 4,
                                            "mem_free_pct": 60, "band": "Healthy"}]},
    }
    out = str(tmp_path / f"ops_bands_{type(poison).__name__}.docx")
    write_ops_handbook_docx(out, snap, "Unit Test Fleet")     # must not raise
    text = _all_text(Document(out))
    # the section still renders from the readable evidence; the unreadable band mix degrades to empty
    assert "3.2 Control-plane capacity baseline" in text
    assert "single point in time" in text


def test_missing_docx_is_warning_not_crash(monkeypatch, tmp_path):
    import builtins
    real_import = builtins.__import__

    def boom(name, *a, **k):
        if name == "docx" or name.startswith("docx."):
            raise ImportError("no docx")
        return real_import(name, *a, **k)
    monkeypatch.setattr(builtins, "__import__", boom)
    out = tmp_path / "never.docx"
    write_ops_handbook_docx(str(out), {}, "x")            # warns + returns
    assert not out.exists()


def test_ops_handbook_tolerates_non_dict_protocol_health_row(tmp_path):
    """[audit-2 #14] a non-dict row in snap['protocol_health'] crashed write_ops_handbook_docx (.get on a str/None)
    despite the 'never raises' contract."""
    from cisco_toolkit.ops import write_ops_handbook_docx
    out = str(tmp_path / "ops.docx")
    write_ops_handbook_docx(out, {"devices": {"sw1": None}, "protocol_health": ["STP up", None]}, "bad")  # no raise
    import os
    assert os.path.exists(out)


# ---- P2 (deliverable-excellence): Backup & Recovery + Known-Issues sections ----

def _heads(doc):
    return [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]


def test_backup_recovery_section_carries_restore_test_discipline(tmp_path, golden_snap):
    """P2: the handbook carries a Backup & Recovery section — strategy, cadence, restore procedure and,
    load-bearing, the restore-TEST discipline (a backup never restore-tested is not a backup)."""
    out = tmp_path / "ops.docx"
    write_ops_handbook_docx(str(out), golden_snap, "golden-fleet")
    doc = Document(str(out))
    text = _all_text(doc)
    assert any("Backup & Recovery" in h for h in _heads(doc)), _heads(doc)
    # the four pillars are all present
    assert "restore" in text.lower()                     # restore procedure
    assert "retention" in text.lower()                   # storage / retention
    # the discipline that separates a backup from a false sense of security
    assert "restore-test" in text.lower() or "restore test" in text.lower()
    assert "is not a backup" in text                     # the exact doctrine line


def test_backup_evidence_is_coverage_honest(tmp_path, golden_snap):
    """P2: the backup section is evidence-gated — it reconciles device counts to the canonical scale and,
    when no backup-evidence axis was collected, DECLARES that as a blind spot rather than asserting the
    fleet is backed up (the false-health class)."""
    out = tmp_path / "ops.docx"
    write_ops_handbook_docx(str(out), golden_snap, "golden-fleet")
    doc = Document(str(out))
    text = _all_text(doc)
    # coverage-honest: the assessment does not collect backup state, so it must NOT claim devices are backed up
    assert "does not directly evidence" in text or "not directly evidence" in text \
        or "blind spot" in text.lower()
    # count reconciled to the canonical scale (SSOT), not a silent fabricated number
    assert "device(s) in scope" in text or "devices in scope" in text or "of the fleet" in text


def test_backup_ndfc_discipline_gated_on_evpn_target(tmp_path):
    """P2: when the target is the NX-OS VXLAN-EVPN / NDFC fabric, the backup section notes NDFC
    config-backup + the 'no out-of-band CLI once managed' discipline; on a non-EVPN engagement it stays
    silent (gated on the same design_blueprint.evpn_migration.applicable flag the MOP uses)."""
    on = {"script_version": "V", "devices": {"leaf1": {}},
          "design_blueprint": {"evpn_migration": {"applicable": True,
                                                  "model_basis": "requirement-confirmed"}}}
    off = {"script_version": "V", "devices": {"sw1": {}}}
    a = tmp_path / "on.docx"
    b = tmp_path / "off.docx"
    write_ops_handbook_docx(str(a), on, "evpn")
    write_ops_handbook_docx(str(b), off, "legacy")
    da, db = Document(str(a)), Document(str(b))
    ta, tb = _all_text(da), _all_text(db)
    # the gated §6.4 fabric-backup subsection is present on-target
    assert any("Fabric backup discipline" in h for h in _heads(da)), _heads(da)
    assert "NDFC" in ta                                   # NDFC config-backup discipline surfaces on-target
    assert "out-of-band CLI" in ta or "out-of-band cli" in ta.lower()
    # ...and is SILENT on a legacy (non-EVPN) engagement — no §6.4, no "no out-of-band CLI" doctrine
    assert not any("Fabric backup discipline" in h for h in _heads(db)), _heads(db)
    assert "out-of-band CLI" not in tb and "out-of-band cli" not in tb.lower()


def test_known_issues_synthesizes_from_real_axes_with_citation(tmp_path, golden_snap):
    """P2: the Known-Issues section synthesizes the assessment's OWN findings across >=3 real axes and
    cites each issue's source axis + affected devices — it is not a generic caveat list."""
    out = tmp_path / "ops.docx"
    write_ops_handbook_docx(str(out), golden_snap, "golden-fleet")
    doc = Document(str(out))
    text = _all_text(doc)
    assert any("Known Issues" in h for h in _heads(doc)), _heads(doc)
    # the golden fleet fires >=3 distinct axes — each must appear as a cited known-issue
    assert "Syslog Intelligence" in text                 # recurring-signature axis cited
    assert "Software Risk" in text                        # advisory-surface axis cited
    assert "Platform Health" in text                      # hot control-plane axis cited
    # affected device is named, not just an abstract count
    assert "core1" in text
    # a real recurring signature is surfaced as a known-issue (mac-flap / err-disable class)
    assert "flap" in text.lower() or "err-disable" in text.lower()


def test_known_issues_declares_not_assessable_axis_no_silent_clean(tmp_path):
    """P2 (the load-bearing honesty test): an axis that was NOT collected is DECLARED not-assessable in
    the Known-Issues section — it never silently reads as 'no known issues'. The golden lacks
    lifecycle_risk, so lifecycle must be named as a blind spot, not omitted."""
    # a snapshot with syslog + software evidence but NO lifecycle_risk axis at all
    snap = {
        "script_version": "V", "devices": {"c1": {}},
        "syslog_intelligence": {"summary": {"n_devices": 1, "n_collected": 1, "n_detections": 1},
                                "detections": [{"host": "c1", "label": "MAC address flapping",
                                                "severity": "High", "count": 2}]},
    }
    out = tmp_path / "ops.docx"
    write_ops_handbook_docx(str(out), snap, "partial")
    doc = Document(str(out))
    text = _all_text(doc)
    # STRUCTURAL (non-vacuous): the SPECIFIC uncollected axes must land in `absent`, not be dropped.
    # ("not-assessable" as a bare substring is vacuous — it is always present from the §7.1 heading.)
    from cisco_toolkit.ops import _facts, _known_issues
    _issues, absent = _known_issues(_facts(snap))
    absent_axes = " | ".join(a for a, _how in absent).lower()
    assert "lifecycle risk" in absent_axes, absent          # the uncollected axis IS declared
    assert "security posture" in absent_axes, absent        # the axis whose silent-drop was the 2026-07-05 bug
    # a not-collected axis must never appear as a clean/issue row
    issue_axes = " | ".join(a for a, *_ in _issues).lower()
    assert "lifecycle risk" not in issue_axes and "security posture" not in issue_axes
    assert "no known issues" not in text.lower()            # never a false all-clear in the rendered doc


def test_known_issues_all_axes_absent_is_honest_not_empty(tmp_path):
    """P2: with NO finding-bearing axes collected at all, the Known-Issues section still renders and is
    coverage-honest — every axis declared not-assessable BY NAME, never an empty section that reads clean."""
    from cisco_toolkit.ops import _facts, _known_issues
    issues, absent = _known_issues(_facts({"devices": {"x": {}}}))
    absent_axes = " | ".join(a for a, _how in absent).lower()
    for axis in ("syslog", "software risk", "platform health", "lifecycle risk", "qos audit", "security posture"):
        assert axis in absent_axes, f"{axis!r} not declared not-assessable: {absent}"
    assert not issues, f"nothing was collected, so there can be no positive known-issue rows: {issues}"
    out = tmp_path / "ops.docx"
    write_ops_handbook_docx(str(out), {"devices": {"x": {}}}, "bare")
    doc = Document(str(out))
    text = _all_text(doc)
    assert any("Known Issues" in h for h in _heads(doc)), _heads(doc)
    assert "no known issues" not in text.lower()          # never a false all-clear


def test_known_issues_security_affected_names_only_failing_hosts():
    """Adversarial-review HIGH (2026-07-05): the Security-Posture 'Affected' set must be ONLY the devices
    that actually failed a CIS check — never every device that merely carries a security block (that would
    tell a change board that clean boxes have open hardening failures)."""
    from cisco_toolkit.ops import _facts, _known_issues
    snap = {"devices": {"c1": {}, "c2": {}, "c3": {}}, "security": {
        "c1": {"summary": {"fail": 0}}, "c2": {"summary": {"fail": 5}}, "c3": {"summary": {"fail": 0}}}}
    issues, _absent = _known_issues(_facts(snap))
    sec = [row for row in issues if row[0] == "Security Posture"]
    assert sec, "security collected-with-failures must produce a Security Posture issue row"
    affected = sec[0][2]
    assert "c2" in affected and "c1" not in affected and "c3" not in affected, affected
    # and a collected-but-all-clean fleet reads screened-clean, never silently absent
    clean = {"devices": {"c1": {}}, "security": {"c1": {"summary": {"fail": 0}}}}
    cissues, cabsent = _known_issues(_facts(clean))
    assert any(r[0] == "Security Posture" and "no failing check" in r[1].lower() for r in cissues)
    assert "security posture" not in " | ".join(a for a, _ in cabsent).lower()
