"""[audit-5 / the flagged follow-up to the design.py fix] every DOCX generator must survive a Unicode
noncharacter (U+FFFE/U+FFFF) or lone surrogate (U+D800-DFFF) in any device-derived string. The design + deck
fixes (and the runbook fix) entry-sanitize their snapshot, and docmeta.add_table/_kv route table text through
xml_safe -- but mop/crd/engagement/archreview/ops also write device text via DIRECT add_run paths that bypass the
table hardening, so one bad byte (collected with errors='ignore') aborted the whole document at XML serialization.
"""
import importlib
import os

import pytest

pytest.importorskip("docx", reason="python-docx not installed (optional deliverable dependency)")
from docx import Document  # noqa: E402


def _poison_snap():
    bad = chr(0xFFFF) + chr(0xFFFE) + chr(0xD800)   # two noncharacters + a lone surrogate (ASCII-safe source)
    return {
        "devices": {"core1" + bad: {"model": "N9K" + bad, "hostname": "core1" + bad, "serial_number": "SAL1" + bad}},
        "executive_brief": {"scale": {"n_devices": 1, "n_collected": 1, "n_endpoints": 5, "n_vlans": 2},
                            "posture": {"avg_health": 50, "n_critical": 1}, "axes": [], "top_gating": []},
        "failure_impact": [{"host": "core1" + bad, "stranded": 9, "severity": "High", "vlans_impacted": 2,
                            "detail": "VLAN 10" + bad}],
        "migration_readiness": [{"group": "G1" + bad, "switches": ["core1" + bad], "readiness": "READY",
                                 "endpoints": 5, "checks": [], "n_fail": 0, "n_warn": 0}],
        "punchlist": [{"severity": "High", "category": "X" + bad, "title": "issue " + bad, "devices": ["core1" + bad]}],
        "health_scores": [{"switch": "core1" + bad, "band": "Critical", "score": 10}],
        "migration_scenarios": {"per_group": [{"group": "G1" + bad, "recommended_scenario": "phased" + bad}]},
        "architecture_review": {"checks": [{"domain": "X" + bad, "title": "t" + bad, "verdict": "deviation",
                                            "evidence": "e" + bad}], "summary": {}},
        "interfaces": {"core1" + bad: {"Gi1/0/1": {"switchport_mode": "Access", "vlan": "10",
                                                   "end_host_mac": "aaaa.0000.0001", "cdp_neighbor": "nbr" + bad}}},
    }


@pytest.mark.parametrize("mod,fn", [
    ("mop", "write_mop_docx"), ("crd", "write_crd_docx"), ("engagement", "write_engagement_docx"),
    ("archreview", "write_archreview_docx"), ("ops", "write_ops_handbook_docx"),
    ("design", "write_design_doc_docx"), ("runbook", "write_runbook_docx"),
])
def test_docx_generator_survives_xml_illegal_chars(mod, fn, tmp_path):
    f = getattr(importlib.import_module("cisco_toolkit." + mod), fn)
    out = str(tmp_path / f"{mod}.docx")
    f(out, _poison_snap(), "[HISTORY-REDACTED]" + chr(0xFFFF))   # must not raise
    assert os.path.exists(out)
    Document(out)                                # and be a valid, openable .docx
