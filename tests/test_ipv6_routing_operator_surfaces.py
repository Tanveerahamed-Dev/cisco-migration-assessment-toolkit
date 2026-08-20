"""Closed operator projections for the serialized IPv6 routing receipt."""

from __future__ import annotations

import copy
import hashlib
import json
from pathlib import Path
import re
import shutil
import subprocess

import pytest
from openpyxl import Workbook

import cisco_toolkit.ipv6_routing as ipv6_routing_module
from cisco_toolkit.capture_integrity import compute_capture_integrity_from_paths
from cisco_toolkit.excel import (
    IPV6_ROUTING_SHEET_NAME,
    write_ipv6_routing_adjacency_sheet,
)
from cisco_toolkit.ipv6_routing import (
    compute_ipv6_routing_adjacency_baseline,
    embedded_ipv6_routing_adjacency_baseline,
    validate_ipv6_routing_adjacency_baseline,
)
from cisco_toolkit.runbook import write_runbook_docx


ROOT = Path(__file__).resolve().parents[1]
EXPLORER = ROOT / "cisco_toolkit" / "blast_radius_explorer.html"
NODE = shutil.which("node")


def _route_summary(*, ospf: int = 0, bgp: int = 0) -> str:
    total = ospf + bgp + 4
    return (
        f"IPv6 Routing Table - default - {total} entries\n"
        "Route Source    Networks    Subnets     Overhead    Memory (bytes)\n"
        "connected       2           0           192         288\n"
        "local           2           0           192         288\n"
        f"ospf 1          {ospf}           0           96          144\n"
        f"bgp 65001       {bgp}           0           96          144\n"
        f"Total           {total}           0           576         864\n"
    )


def _ospf_table(rows: list[tuple[str, str, str, str]], *, malformed: bool = False) -> str:
    lines = [
        "OSPFv3 1 address-family ipv6 (router-id 10.255.255.1)",
        "",
        "Neighbor ID     Pri   State           Dead Time   Interface ID    Interface",
    ]
    for index, (peer, state, role, interface) in enumerate(rows, 1):
        lines.append(
            f"{peer:<20} 1   {state}/{role:<10} 00:00:37    {index:<15} {interface}"
        )
    if malformed:
        lines.append("10.254.254.254 1 BROKEN-CANDIDATE")
    return "\n".join(lines) + "\n"


def _bgp_table(rows: list[tuple[str, str, str]]) -> str:
    lines = [
        "BGP router identifier 10.255.255.1, local AS number 65001",
        "Neighbor                  V         AS  MsgRcvd  MsgSent  TblVer  InQ OutQ Up/Down  State/PfxRcd",
    ]
    for peer, remote_as, state in rows:
        lines.append(
            f"{peer:<25} 4      {remote_as:<8} 3421     3418      15    0    0 1d02h          {state}"
        )
    return "\n".join(lines) + "\n"


def _canonical_sha(value: object) -> str:
    return hashlib.sha256(
        json.dumps(
            value, sort_keys=True, separators=(",", ":"), ensure_ascii=True,
            allow_nan=False,
        ).encode("utf-8")
    ).hexdigest()


def _canonical_digest(value: dict) -> str:
    payload = copy.deepcopy(value)
    payload["summary"].pop("baseline_sha256", None)
    return _canonical_sha(payload)


def _rehash_coverage_cell(baseline: dict, cell: dict) -> None:
    source_payload = {
        key: value for key, value in cell.items()
        if key not in {"source_sha256", "projection_sha256"}
    }
    row_fields = (
        "switch", "platform", "protocol", "routing_instance", "process",
        "peer", "peer_key", "interface", "remote_as", "role", "state_raw",
        "state", "prefix_count", "prefix_count_present", "status", "findings",
    )
    projected_rows = [] if cell["input"] == "route_summary" else [
        {key: row[key] for key in row_fields}
        for row in baseline["rows"]
        if row["switch"] == cell["switch"] and row["protocol"] == cell["protocol"]
    ]
    projection_payload = {
        "switch": cell["switch"], "input": cell["input"],
        "protocol": cell["protocol"], "subject": cell["subject"],
        "status": cell["status"], "active_route_count": cell["active_route_count"],
        "rows": projected_rows,
    }
    cell["source_sha256"] = _canonical_sha(source_payload)
    cell["projection_sha256"] = _canonical_sha(projection_payload)


@pytest.fixture(scope="module")
def receipts(tmp_path_factory):
    root = tmp_path_factory.mktemp("ipv6-routing-operator")
    healthy_rows = [
        (f"10.{index}.200.1", "FULL", "DR", f"Vlan{index}")
        for index in range(1, 56)
    ]
    bodies = {
        "bulk-edge": {
            "show ipv6 route summary": _route_summary(ospf=57, bgp=2),
            "show ospfv3 neighbor": _ospf_table(healthy_rows + [
                ("10.250.0.1", "EXSTART", "-", "GigabitEthernet0/1"),
                ("10.250.0.2", "INIT", "DROTHER", "GigabitEthernet0/2"),
            ]),
            "show bgp ipv6 unicast summary": _bgp_table([
                ("2001:db8:0:1::1", "65001", "12"),
                ("2001:db8:0:9::9", "65009", "Active"),
                ("2001:db8:0:a::a", "65010", "Idle (Admin-Shut)"),
            ]),
        },
        "review-edge": {
            "show ipv6 route summary": _route_summary(ospf=1),
            "show ospfv3 neighbor": _ospf_table(
                [("10.251.0.1", "FULL", "DR", "Vlan251")], malformed=True,
            ),
        },
        "notverified-edge": {
            "show ipv6 route summary": _route_summary(ospf=1),
        },
        "coverage-only-edge": {
            "show ipv6 route summary": _route_summary(),
        },
    }
    paths: dict[str, dict[str, str]] = {}
    for host, commands in bodies.items():
        paths[host] = {}
        for index, (command, body) in enumerate(commands.items()):
            path = root / f"{host}-{index}.txt"
            path.write_text(body, encoding="utf-8")
            paths[host][command] = str(path)
    current = compute_ipv6_routing_adjacency_baseline(
        paths,
        compute_capture_integrity_from_paths(paths),
        {host: {"platform": "ios"} for host in paths},
    )
    current_view = validate_ipv6_routing_adjacency_baseline(
        current, require_current_run=True,
    )
    assert current_view["valid"] is True, current_view
    embedded = embedded_ipv6_routing_adjacency_baseline(current)
    embedded_view = validate_ipv6_routing_adjacency_baseline(embedded)
    assert embedded_view["valid"] is True, embedded_view
    return {"root": root, "current": current, "embedded": embedded}


def _sheet_text(sheet) -> str:
    return "\n".join(
        str(cell.value)
        for row in sheet.iter_rows()
        for cell in row
        if cell.value is not None
    )


def _docx_text(path: Path) -> str:
    from docx import Document

    document = Document(path)
    return "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    )


def _explorer_result(baseline: object, host: str | None = None) -> dict:
    assert NODE is not None
    script = r"""
const fs=require('fs');
const html=fs.readFileSync(process.argv[1],'utf8');
const start=html.indexOf('const _FHRP_GROUP_SCOPE=');
const end=html.indexOf('/* the dynamic routing protocols',start);
const detailStart=html.indexOf('function protocolDetailHosts(');
const detailEnd=html.indexOf('function drawProtocols()',detailStart);
if(start<0||end<0||detailStart<0||detailEnd<0)throw new Error('protocol receipt block not found');
const input=JSON.parse(fs.readFileSync(0,'utf8'));
let SNAP={ipv6_routing_adjacency_baseline:input.baseline};
function esc(v){return String(v??'').replace(/[&<>"']/g,c=>({'&':'&amp;','<':'&lt;','>':'&gt;','"':'&quot;',"'":'&#39;'}[c]));}
function shortName(v){return String(v??'');}
eval(html.slice(start,end));
eval(html.slice(detailStart,detailEnd));
process.stdout.write(JSON.stringify({
  valid:ipv6RoutingAdjacencyReceiptValid(input.baseline),
  rendered:ipv6RoutingAdjacencySection(input.host),
  hosts:protocolDetailHosts([],[],[],[]),
  debug:input.baseline&&typeof input.baseline==='object'?{
    badRows:Array.isArray(input.baseline.rows)?input.baseline.rows.map((row,index)=>[index,_ipv6RoutingRowShape(row,input.baseline.projection_custody)]).filter(pair=>!pair[1]).slice(0,5):[],
    badCoverage:Array.isArray(input.baseline.coverage)?input.baseline.coverage.map((cell,index)=>[index,_ipv6RoutingCoverageShape(cell)]).filter(pair=>!pair[1]).slice(0,5):[],
    badFindings:Array.isArray(input.baseline.findings)?input.baseline.findings.map((finding,index)=>[index,_ipv6RoutingGlobalFinding(finding)]).filter(pair=>!pair[1]).slice(0,5):[],
    digest:input.baseline.summary&&_ipv6RoutingBaselineDigest(input.baseline)===input.baseline.summary.baseline_sha256,
  }:null,
}));
"""
    proc = subprocess.run(
        [NODE, "-e", script, str(EXPLORER)],
        input=json.dumps({"baseline": baseline, "host": host}),
        capture_output=True,
        check=True,
        text=True,
        encoding="utf-8",
        timeout=60,
    )
    return json.loads(proc.stdout)


def _explorer_validation_result(encoded_baseline: str) -> dict:
    assert NODE is not None
    script = r"""
const fs=require('fs');
const html=fs.readFileSync(process.argv[1],'utf8');
const start=html.indexOf('const _FHRP_GROUP_SCOPE=');
const end=html.indexOf('/* the dynamic routing protocols',start);
if(start<0||end<0)throw new Error('protocol receipt block not found');
const baseline=JSON.parse(fs.readFileSync(0,'utf8'));
let SNAP={ipv6_routing_adjacency_baseline:baseline};
function esc(v){return String(v??'');}
function shortName(v){return String(v??'');}
eval(html.slice(start,end));
const started=performance.now();
const valid=ipv6RoutingAdjacencyReceiptValid(baseline);
const elapsedMs=performance.now()-started,memory=process.memoryUsage();
process.stdout.write(JSON.stringify({
  valid,elapsedMs,rssBytes:memory.rss,heapUsedBytes:memory.heapUsed,
}));
"""
    proc = subprocess.run(
        [NODE, "--max-old-space-size=512", "-e", script, str(EXPLORER)],
        input=encoded_baseline,
        capture_output=True,
        check=True,
        text=True,
        encoding="utf-8",
        timeout=90,
    )
    return json.loads(proc.stdout)


def _explorer_section_cap_result(encoded_baseline: str) -> dict:
    assert NODE is not None
    script = r"""
const fs=require('fs');
const html=fs.readFileSync(process.argv[1],'utf8');
const start=html.indexOf('const _FHRP_GROUP_SCOPE=');
const end=html.indexOf('/* the dynamic routing protocols',start);
if(start<0||end<0)throw new Error('protocol receipt block not found');
const baseline=JSON.parse(fs.readFileSync(0,'utf8'));
let SNAP={ipv6_routing_adjacency_baseline:baseline};
function esc(v){return String(v??'').replace(/[&<>"']/g,c=>({'&':'&amp;','<':'&lt;','>':'&gt;','"':'&quot;',"'":'&#39;'}[c]));}
function shortName(v){return String(v??'');}
eval(html.slice(start,end));
const started=performance.now(),section=ipv6RoutingAdjacencySection();
const elapsedMs=performance.now()-started;
const exported=ipv6RoutingAdjacencyBlockerExportRows();
const rowKeys=['acceptance','command','findings','interface','peer','platform','process','projection_custody','protocol','remote_as','state','state_raw','status','switch'].sort().join('\0');
const findingKeys=['code','issue','kind'].sort().join('\0');
const memory=process.memoryUsage();
process.stdout.write(JSON.stringify({
  elapsedMs,sectionLength:section.length,
  renderedCards:(section.match(/data-jump=/g)||[]).length,
  exactCounts:section.includes('blockers rendered 200 / total 20000 / omitted 19800'),
  exportControl:section.includes('Export all IPv6 blocker rows (JSON)'),
  exportCount:exported.length,
  exportKeysSafe:exported.every(row=>Object.keys(row).sort().join('\0')===rowKeys&&
    Array.isArray(row.findings)&&row.findings.every(finding=>Object.keys(finding).sort().join('\0')===findingKeys)),
  rssBytes:memory.rss,heapUsedBytes:memory.heapUsed,
}));
"""
    proc = subprocess.run(
        [NODE, "--max-old-space-size=512", "-e", script, str(EXPLORER)],
        input=encoded_baseline,
        capture_output=True,
        check=True,
        text=True,
        encoding="utf-8",
        timeout=90,
    )
    return json.loads(proc.stdout)


def _schema_max_ipv6_receipt(receipts: dict, *, degraded: bool = False) -> dict:
    base = receipts["embedded"]
    row_template = next(
        row for row in base["rows"]
        if row["protocol"] == "OSPFv3"
        and row["status"] == ("degraded" if degraded else "assessed")
        and row["state_raw"] == ("EXSTART/-" if degraded else "FULL/DR")
    )
    route_template = next(
        cell for cell in base["coverage"]
        if cell["switch"] == "review-edge" and cell["input"] == "route_summary"
    )
    ospf_template = next(
        cell for cell in base["coverage"]
        if cell["switch"] == "bulk-edge"
        and cell["input"] == "ospfv3_neighbors"
    )
    bgp_template = next(
        cell for cell in base["coverage"]
        if cell["switch"] == "review-edge"
        and cell["input"] == "bgp_ipv6_neighbors"
    )
    assert route_template["parser_status"] == "complete"
    assert ospf_template["capture_status"] == "ok"
    assert bgp_template["capture_status"] != "ok"

    n_hosts, n_rows = 4096, 20_000
    rows: list[dict] = []
    coverage: list[dict] = []
    remaining = n_rows
    for index in range(n_hosts):
        host = f"h{index:04d}"
        host_row_count = min(5, remaining - (n_hosts - index - 1) * 4)
        remaining -= host_row_count
        host_rows = []
        for peer_index in range(1, host_row_count + 1):
            peer = f"10.{index // 256}.{index % 256}.{peer_index}"
            interface = f"Vlan{peer_index}"
            row = copy.deepcopy(row_template)
            row.update({
                "switch": host, "peer": peer, "interface": interface,
                "peer_key": (
                    f"ospfv3|default|1|{peer}|{interface.casefold()}"
                ),
            })
            row["acceptance"] = ipv6_routing_module._acceptance(row)
            host_rows.append(row)
        rows.extend(host_rows)

        route_cell = copy.deepcopy(route_template)
        route_cell.update({"switch": host, "active_route_count": host_row_count})
        ospf_cell = copy.deepcopy(ospf_template)
        ospf_cell.update({
            "switch": host, "subject": True,
            "status": "degraded" if degraded else "assessed",
            "candidate_count": host_row_count,
            "parsed_count": host_row_count, "rejected_count": 0,
            "active_route_count": host_row_count,
            "finding_codes": ["ospfv3_state_degraded"] if degraded else [],
        })
        bgp_cell = copy.deepcopy(bgp_template)
        bgp_cell["switch"] = host
        for cell in (route_cell, ospf_cell, bgp_cell):
            _rehash_coverage_cell({"rows": host_rows}, cell)
            coverage.append(cell)
    assert remaining == 0

    baseline = copy.deepcopy(base)
    findings = [
        {
            "switch": row["switch"], "protocol": row["protocol"],
            "peer_key": row["peer_key"], "kind": finding["kind"],
            "code": finding["code"], "issue": finding["issue"],
        }
        for row in rows for finding in row["findings"]
    ]
    baseline.update({
        "rows": rows, "coverage": coverage, "findings": findings,
        "verdict": "BLOCKED" if degraded else "CLEAR", "assessed": True,
    })
    baseline["summary"].update({
        "n_hosts": n_hosts, "n_subject_hosts": n_hosts, "n_rows": n_rows,
        "n_ospfv3_rows": n_rows, "n_bgpv6_rows": 0,
        "n_assessed": 0 if degraded else n_rows,
        "n_degraded": n_rows if degraded else 0, "n_review": 0,
        "n_not_verified": 0,
        "by_status": {
            "degraded": n_rows if degraded else 0, "review": 0,
            "not_verified": 0, "assessed": 0 if degraded else n_rows,
        },
        "by_coverage_status": {
            "degraded": n_hosts if degraded else 0, "review": 0,
            "not_verified": 0,
            "assessed": n_hosts if degraded else n_hosts * 2,
            "not_applicable": n_hosts,
        },
    })
    payload = dict(baseline)
    payload["summary"] = dict(baseline["summary"])
    payload["summary"].pop("baseline_sha256", None)
    baseline["summary"]["baseline_sha256"] = _canonical_sha(payload)
    return baseline


@pytest.mark.skipif(NODE is None, reason="node is required for Explorer receipt tests")
def test_valid_embedded_broken_and_healthy_rows_survive_caps_without_private_leaves(receipts):
    baseline = receipts["embedded"]
    by_status = baseline["summary"]["by_status"]
    assert by_status["assessed"] >= 56
    assert by_status["degraded"] >= 3
    assert by_status["review"] >= 1
    assert by_status["not_verified"] >= 1

    workbook = Workbook()
    write_ipv6_routing_adjacency_sheet(workbook, baseline)
    sheet = workbook[IPV6_ROUTING_SHEET_NAME]
    workbook_text = _sheet_text(sheet)
    assert sheet.max_row == 5 + sum(
        by_status[state] for state in ("degraded", "review", "not_verified")
    ) + 50
    assert sheet.freeze_panes == "A6"

    output = receipts["root"] / "ipv6-routing-runbook.docx"
    write_runbook_docx(
        str(output),
        {
            "ipv6_routing_adjacency_baseline": baseline,
            "protocol_health": [{
                "switch": "bulk-edge", "protocol": "OSPF", "severity": "Info",
                "summary": "separate health", "detail": "HEALTH-DETAIL-PRESERVED",
            }],
            "protocol_intelligence": [{
                "switch": "bulk-edge", "protocol": "OSPF", "state": "EXSTART",
                "severity": "High", "meaning": "INTELLIGENCE-DETAIL-PRESERVED",
                "likely_cause": "INTELLIGENCE-DETAIL-PRESERVED",
                "remediation": "review", "confidence": "Inferred",
            }],
        },
        "fixture",
    )
    runbook_text = _docx_text(output)
    explorer = _explorer_result(baseline)
    assert explorer["valid"] is True, explorer["debug"]

    blockers = [
        row for row in baseline["rows"]
        if row["status"] in {"degraded", "review", "not_verified"}
    ]
    assessed = [row for row in baseline["rows"] if row["status"] == "assessed"]
    assert len(assessed) > 50
    for text in (workbook_text, runbook_text, explorer["rendered"]):
        for row in blockers:
            assert row["switch"] in text
            assert row["peer"] in text
            assert row["state_raw"] in text
        assert assessed[49]["peer"] in text
        assert assessed[50]["peer"] not in text
        assert "matching a degraded state is NOT ACCEPTANCE" in text
        assert "summary.by_coverage_status" in text
        assert baseline["summary"]["baseline_sha256"] not in text
        assert baseline["coverage"][0]["source_sha256"] not in text
        assert baseline["coverage"][0]["projection_sha256"] not in text
        assert baseline["rows"][0]["source_key"] not in text
        assert str(receipts["root"]) not in text

    blocker_count = len(blockers)
    assert f"All {blocker_count} blocker row(s) are shown; 50 of" in runbook_text
    assert (
        f"blockers rendered {blocker_count} / total {blocker_count} / omitted 0"
        in explorer["rendered"]
    )
    assert (
        f"assessed rendered 50 / total {len(assessed)} / omitted "
        f"{len(assessed) - 50}" in explorer["rendered"]
    )
    assert "Export all IPv6 blocker rows (JSON)" in explorer["rendered"]
    assert "HEALTH-DETAIL-PRESERVED" in runbook_text
    assert "INTELLIGENCE-DETAIL-PRESERVED" in runbook_text


@pytest.mark.skipif(NODE is None, reason="node is required for Explorer receipt tests")
def test_malformed_tampered_and_serialized_current_claims_fail_closed(receipts):
    current_json = json.loads(json.dumps(receipts["current"]))
    assert current_json["projection_custody"] == "current_run_source_bound"
    tampered = copy.deepcopy(receipts["embedded"])
    tampered["rows"][0]["state_raw"] = "HOSTILE-STATE"
    hostile = {
        "schema": "ipv6_routing_adjacency_baseline/1",
        "projection_custody": "embedded_unverified",
        "rows": [{"switch": "HOSTILE-SWITCH", "acceptance": "HOSTILE PASS"}],
        "raw": "password HOSTILE-SECRET",
    }

    for index, rejected in enumerate((current_json, tampered, hostile)):
        workbook = Workbook()
        write_ipv6_routing_adjacency_sheet(workbook, rejected)
        workbook_text = _sheet_text(workbook[IPV6_ROUTING_SHEET_NAME])
        output = receipts["root"] / f"ipv6-rejected-{index}.docx"
        write_runbook_docx(
            str(output), {"ipv6_routing_adjacency_baseline": rejected}, "fixture",
        )
        runbook_text = _docx_text(output)
        explorer = _explorer_result(rejected)
        assert explorer["valid"] is False
        for text in (workbook_text, runbook_text, explorer["rendered"]):
            assert "NOT ASSESSED" in text.upper()
            assert "No rejected receipt leaves" in text
            assert "HOSTILE" not in text
            assert "bulk-edge" not in text


@pytest.mark.skipif(NODE is None, reason="node is required for Explorer receipt tests")
def test_explorer_python_js_semantics_match_and_semantic_tamper_fails(receipts):
    valid = receipts["embedded"]

    def coverage_only_route_receipt(
        parser_status: str,
        candidate_count: int,
        parsed_count: int,
        rejected_count: int,
        finding_codes: list[str],
    ) -> dict:
        baseline = copy.deepcopy(valid)
        cell = next(
            item for item in baseline["coverage"]
            if item["switch"] == "coverage-only-edge"
            and item["input"] == "route_summary"
        )
        assert cell["subject"] is False
        assert cell["status"] == "not_applicable"
        cell.update({
            "parser_status": parser_status,
            "candidate_count": candidate_count,
            "parsed_count": parsed_count,
            "rejected_count": rejected_count,
            "finding_codes": finding_codes,
        })
        _rehash_coverage_cell(baseline, cell)
        baseline["summary"]["baseline_sha256"] = _canonical_digest(baseline)
        return baseline

    def non_subject_family_receipt(
        parser_status: str,
        candidate_count: int,
        parsed_count: int,
        rejected_count: int,
    ) -> dict:
        baseline = copy.deepcopy(valid)
        cell = next(
            item for item in baseline["coverage"]
            if item["switch"] == "coverage-only-edge"
            and item["input"] == "ospfv3_neighbors"
        )
        assert cell["subject"] is False
        assert cell["status"] == "not_applicable"
        cell.update({
            "capture_status": "ok", "parser_status": parser_status,
            "candidate_count": candidate_count, "parsed_count": parsed_count,
            "rejected_count": rejected_count, "active_route_count": 0,
            "finding_codes": [],
        })
        _rehash_coverage_cell(baseline, cell)
        baseline["summary"]["baseline_sha256"] = _canonical_digest(baseline)
        return baseline

    def family_process_receipt(protocol: str, mutation: str) -> dict:
        baseline = copy.deepcopy(valid)
        candidates = [
            row for row in baseline["rows"]
            if row["switch"] == "bulk-edge" and row["protocol"] == protocol
            and row["peer_key"] and row["status"] == "assessed"
            and not row["findings"]
        ]
        assert candidates
        target = min(candidates, key=lambda row: row["peer_key"])
        if mutation == "split":
            target = max(candidates, key=lambda row: row["peer_key"])
        target["process"] = "" if mutation == "remove" else (
            "2" if protocol == "OSPFv3" else "65009"
        )
        if protocol == "OSPFv3":
            parts = target["peer_key"].split("|")
            parts[2] = target["process"] or "-"
            target["peer_key"] = "|".join(parts)
        target["acceptance"] = ipv6_routing_module._acceptance(target)
        protocol_order = {"OSPFv3": 0, "BGPv6": 1}
        baseline["rows"].sort(key=lambda row: (
            row["switch"].casefold(), row["switch"],
            protocol_order[row["protocol"]], row["peer_key"].casefold(),
            row["peer_key"],
        ))
        input_name = (
            "ospfv3_neighbors" if protocol == "OSPFv3"
            else "bgp_ipv6_neighbors"
        )
        cell = next(
            item for item in baseline["coverage"]
            if item["switch"] == "bulk-edge" and item["input"] == input_name
        )
        _rehash_coverage_cell(baseline, cell)
        baseline["summary"]["baseline_sha256"] = _canonical_digest(baseline)
        return baseline

    digest_tamper = copy.deepcopy(valid)
    digest_tamper["rows"][0]["peer"] = "2001:db8::dead"
    semantic_tamper = copy.deepcopy(valid)
    semantic_tamper["rows"][0]["acceptance"] = "SEMANTIC HOSTILE ACCEPTANCE"
    semantic_tamper["summary"]["baseline_sha256"] = _canonical_digest(semantic_tamper)
    semantic_view = validate_ipv6_routing_adjacency_baseline(semantic_tamper)
    assert semantic_view["valid"] is False
    assert semantic_view["reason"] == "baseline_row_semantics_mismatch"
    assert _explorer_result(semantic_tamper)["valid"] is False
    duplicate_identity = copy.deepcopy(valid)
    duplicate_identity["rows"][1] = copy.deepcopy(duplicate_identity["rows"][0])
    duplicate_identity["summary"]["baseline_sha256"] = _canonical_digest(
        duplicate_identity
    )
    row_order = copy.deepcopy(valid)
    row_order["rows"][0], row_order["rows"][1] = row_order["rows"][1], row_order["rows"][0]
    row_order["summary"]["baseline_sha256"] = _canonical_digest(row_order)
    verdict_tamper = copy.deepcopy(valid)
    verdict_tamper["verdict"] = "CLEAR"
    verdict_tamper["summary"]["baseline_sha256"] = _canonical_digest(verdict_tamper)
    assessed_tamper = copy.deepcopy(valid)
    assessed_tamper["assessed"] = not assessed_tamper["assessed"]
    assessed_tamper["summary"]["baseline_sha256"] = _canonical_digest(
        assessed_tamper
    )
    source_hash_tamper = copy.deepcopy(valid)
    source_hash_tamper["coverage"][0]["source_sha256"] = "0" * 64
    source_hash_tamper["summary"]["baseline_sha256"] = _canonical_digest(
        source_hash_tamper
    )
    projection_hash_tamper = copy.deepcopy(valid)
    projection_hash_tamper["coverage"][0]["projection_sha256"] = "f" * 64
    projection_hash_tamper["summary"]["baseline_sha256"] = _canonical_digest(
        projection_hash_tamper
    )
    rejected_route_tamper = copy.deepcopy(valid)
    route_cell = next(
        cell for cell in rejected_route_tamper["coverage"]
        if cell["switch"] == "bulk-edge" and cell["input"] == "route_summary"
    )
    route_cell.update({
        "status": "not_verified", "parser_status": "rejected",
        "candidate_count": 1, "parsed_count": 0, "rejected_count": 1,
        "finding_codes": ["route_census_not_verified", "route_census_rejected"],
    })
    _rehash_coverage_cell(rejected_route_tamper, route_cell)
    rejected_route_tamper["summary"]["baseline_sha256"] = _canonical_digest(
        rejected_route_tamper
    )
    ospf_process_tamper = copy.deepcopy(valid)
    ospf_row = next(
        row for row in ospf_process_tamper["rows"]
        if row["switch"] == "review-edge" and row["protocol"] == "OSPFv3"
        and row["peer_key"]
    )
    old_peer_key = ospf_row["peer_key"]
    ospf_row["process"] = "bad process"
    ospf_row["peer_key"] = (
        f"ospfv3|default|bad process|{ospf_row['peer']}|"
        f"{ospf_row['interface'].casefold()}"
    )
    for finding in ospf_process_tamper["findings"]:
        if (
            finding["switch"] == ospf_row["switch"]
            and finding["protocol"] == "OSPFv3"
            and finding["peer_key"] == old_peer_key
        ):
            finding["peer_key"] = ospf_row["peer_key"]
    ospf_cell = next(
        cell for cell in ospf_process_tamper["coverage"]
        if cell["switch"] == ospf_row["switch"]
        and cell["input"] == "ospfv3_neighbors"
    )
    _rehash_coverage_cell(ospf_process_tamper, ospf_cell)
    ospf_process_tamper["summary"]["baseline_sha256"] = _canonical_digest(
        ospf_process_tamper
    )
    ospf_process_view = validate_ipv6_routing_adjacency_baseline(
        ospf_process_tamper
    )
    assert ospf_process_view["valid"] is False
    assert ospf_process_view["reason"] == "baseline_ospfv3_facts_invalid"

    bgp_state_raw_tamper = copy.deepcopy(valid)
    bgp_row = next(
        row for row in bgp_state_raw_tamper["rows"]
        if row["protocol"] == "BGPv6" and row["state_raw"] == "Active"
    )
    bgp_row["state_raw"] = "Ac tive"
    bgp_row["acceptance"] = bgp_row["acceptance"].replace(
        "state Active", "state Ac tive",
    )
    bgp_cell = next(
        cell for cell in bgp_state_raw_tamper["coverage"]
        if cell["switch"] == bgp_row["switch"]
        and cell["input"] == "bgp_ipv6_neighbors"
    )
    _rehash_coverage_cell(bgp_state_raw_tamper, bgp_cell)
    bgp_state_raw_tamper["summary"]["baseline_sha256"] = _canonical_digest(
        bgp_state_raw_tamper
    )
    bgp_state_raw_view = validate_ipv6_routing_adjacency_baseline(
        bgp_state_raw_tamper
    )
    assert bgp_state_raw_view["valid"] is False
    assert bgp_state_raw_view["reason"] == "baseline_bgpv6_state_invalid"

    route_complete_zero_tamper = copy.deepcopy(valid)
    complete_route_cell = next(
        cell for cell in route_complete_zero_tamper["coverage"]
        if cell["switch"] == "bulk-edge" and cell["input"] == "route_summary"
    )
    assert complete_route_cell["capture_status"] == "ok"
    assert complete_route_cell["parser_status"] == "complete"
    complete_route_cell.update({
        "candidate_count": 0, "parsed_count": 0, "rejected_count": 0,
    })
    _rehash_coverage_cell(route_complete_zero_tamper, complete_route_cell)
    route_complete_zero_tamper["summary"]["baseline_sha256"] = _canonical_digest(
        route_complete_zero_tamper
    )
    route_complete_zero_view = validate_ipv6_routing_adjacency_baseline(
        route_complete_zero_tamper
    )
    assert route_complete_zero_view["valid"] is False
    assert route_complete_zero_view["reason"] == (
        "baseline_coverage_parser_count_invalid"
    )

    non_ok_route_active_count_tamper = copy.deepcopy(valid)
    non_ok_route_cell = next(
        cell for cell in non_ok_route_active_count_tamper["coverage"]
        if cell["switch"] == "notverified-edge"
        and cell["input"] == "route_summary"
    )
    assert non_ok_route_cell["capture_status"] == "ok"
    assert non_ok_route_cell["parser_status"] == "complete"
    assert non_ok_route_cell["active_route_count"] == 1
    non_ok_route_cell.update({
        "status": "not_verified",
        "capture_status": "not_observed",
        "parser_status": "not_verified",
        "candidate_count": 0,
        "parsed_count": 0,
        "rejected_count": 0,
        "finding_codes": ["route_census_not_verified"],
    })
    # Deliberately retain active_route_count == 1. The family cell also has
    # one active route, so only the producer-reachability rule can reject it.
    _rehash_coverage_cell(
        non_ok_route_active_count_tamper, non_ok_route_cell,
    )
    coverage_counts = non_ok_route_active_count_tamper["summary"][
        "by_coverage_status"
    ]
    coverage_counts["assessed"] -= 1
    coverage_counts["not_verified"] += 1
    non_ok_route_active_count_tamper["summary"]["baseline_sha256"] = (
        _canonical_digest(non_ok_route_active_count_tamper)
    )
    non_ok_route_view = validate_ipv6_routing_adjacency_baseline(
        non_ok_route_active_count_tamper
    )
    assert non_ok_route_view["valid"] is False
    assert non_ok_route_view["reason"] == (
        "baseline_coverage_source_parser_invalid"
    )

    valid_route_review = coverage_only_route_receipt(
        "review", 2, 1, 1, ["route_context_review"],
    )
    assert validate_ipv6_routing_adjacency_baseline(valid_route_review)[
        "valid"
    ] is True

    unreachable_route_rejected_tampers = {
        "route_na_ok_rejected_with_code_and_recomputed_hashes": (
            coverage_only_route_receipt(
                "rejected", 1, 0, 1, ["route_census_rejected"],
            )
        ),
        "route_na_ok_rejected_without_code_and_recomputed_hashes": (
            coverage_only_route_receipt("rejected", 1, 0, 1, [])
        ),
    }
    for tampered in unreachable_route_rejected_tampers.values():
        view = validate_ipv6_routing_adjacency_baseline(tampered)
        assert view["valid"] is False
        assert view["reason"] == "baseline_coverage_source_parser_invalid"

    route_review_without_parsed_tamper = coverage_only_route_receipt(
        "review", 1, 0, 1, ["route_context_review"],
    )
    route_review_without_parsed_view = validate_ipv6_routing_adjacency_baseline(
        route_review_without_parsed_tamper
    )
    assert route_review_without_parsed_view["valid"] is False
    assert route_review_without_parsed_view["reason"] == (
        "baseline_coverage_parser_count_invalid"
    )

    unreachable_family_tampers = {
        "family_na_ok_rejected_with_recomputed_hashes": (
            non_subject_family_receipt("rejected", 1, 0, 1)
        ),
        "family_na_ok_not_verified_with_recomputed_hashes": (
            non_subject_family_receipt("not_verified", 0, 0, 0)
        ),
        "family_na_ok_explicit_none_with_recomputed_hashes": (
            non_subject_family_receipt("explicit_no_subject", 0, 0, 0)
        ),
        "family_na_ok_complete_zero_with_recomputed_hashes": (
            non_subject_family_receipt("complete", 0, 0, 0)
        ),
    }
    for tampered in unreachable_family_tampers.values():
        view = validate_ipv6_routing_adjacency_baseline(tampered)
        assert view["valid"] is False
        assert view["reason"] == "baseline_coverage_source_parser_invalid"

    family_process_tampers = {
        "ospf_process_removed_with_recomputed_hashes": (
            family_process_receipt("OSPFv3", "remove")
        ),
        "bgp_process_removed_with_recomputed_hashes": (
            family_process_receipt("BGPv6", "remove")
        ),
        "ospf_process_split_with_recomputed_hashes": (
            family_process_receipt("OSPFv3", "split")
        ),
        "bgp_process_split_with_recomputed_hashes": (
            family_process_receipt("BGPv6", "split")
        ),
    }
    for name, tampered in family_process_tampers.items():
        view = validate_ipv6_routing_adjacency_baseline(tampered)
        assert view["valid"] is False
        expected_reason = (
            "baseline_family_process_mismatch" if "split" in name
            else "baseline_ospfv3_facts_invalid" if name.startswith("ospf")
            else "baseline_bgpv6_identity_invalid"
        )
        assert view["reason"] == expected_reason

    current_json = json.loads(json.dumps(receipts["current"]))
    cases = {
        "valid_embedded": valid,
        "valid_route_review": valid_route_review,
        "digest_tamper": digest_tamper,
        "semantic_tamper_with_recomputed_digest": semantic_tamper,
        "duplicate_identity_with_recomputed_digest": duplicate_identity,
        "row_order_with_recomputed_digest": row_order,
        "verdict_with_recomputed_digest": verdict_tamper,
        "assessed_with_recomputed_digest": assessed_tamper,
        "coverage_source_hash_with_recomputed_digest": source_hash_tamper,
        "coverage_projection_hash_with_recomputed_digest": projection_hash_tamper,
        "route_census_rejected_with_recomputed_hashes": rejected_route_tamper,
        "ospf_process_with_recomputed_hashes": ospf_process_tamper,
        "bgp_state_raw_with_recomputed_hashes": bgp_state_raw_tamper,
        "route_complete_zero_with_recomputed_hashes": route_complete_zero_tamper,
        "non_ok_route_positive_active_count_with_recomputed_hashes": (
            non_ok_route_active_count_tamper
        ),
        "route_review_without_parsed_with_recomputed_hashes": (
            route_review_without_parsed_tamper
        ),
        **unreachable_route_rejected_tampers,
        **unreachable_family_tampers,
        **family_process_tampers,
        "serialized_current_claim": current_json,
        "missing": None,
    }
    for name, baseline in cases.items():
        python = validate_ipv6_routing_adjacency_baseline(baseline)
        expected = bool(
            python["valid"]
            and python["baseline"].get("projection_custody") == "embedded_unverified"
        )
        assert _explorer_result(baseline)["valid"] is expected, name
        if "tamper" in name or "recomputed" in name:
            assert expected is False, name

    html = EXPLORER.read_text(encoding="utf-8")
    block = html[
        html.index("const _IPV6_ROUTING_SCHEMA="):
        html.index("/* the dynamic routing protocols", html.index("const _IPV6_ROUTING_SCHEMA="))
    ]
    assert "b.rows.length>20000" in block
    assert "b.coverage.length>12288" in block
    assert "b.findings.length>20000" in block
    assert "row.findings.length>16" in block
    assert "cell.finding_codes.length>16" in block
    assert "8192" in block
    assert 'b.projection_custody!=="embedded_unverified"' in block
    renderer = block[block.index("function ipv6RoutingAdjacencyCoverageRow"):]
    assert "source_\"+\"sha256" not in renderer
    assert "projection_\"+\"sha256" not in renderer
    assert "source_key" not in renderer


@pytest.mark.skipif(NODE is None, reason="node is required for Explorer cap tests")
def test_explorer_accepts_schema_max_receipt_with_bounded_linear_validation(
    receipts,
):
    baseline = _schema_max_ipv6_receipt(receipts)
    assert len(baseline["rows"]) == 20_000
    assert len(baseline["coverage"]) == 12_288
    encoded = json.dumps(
        baseline, sort_keys=True, separators=(",", ":"), ensure_ascii=True,
        allow_nan=False,
    )
    assert 20 * 1024 * 1024 < len(encoded) <= 32 * 1024 * 1024

    result = _explorer_validation_result(encoded)
    assert result["valid"] is True
    assert result["elapsedMs"] < 30_000
    assert result["rssBytes"] < 768 * 1024 * 1024

    # The IPv6 digest gets the larger schema-specific chunked path; the
    # shared FHRP fragment guard remains unchanged.
    html = EXPLORER.read_text(encoding="utf-8")
    fhrp_canonical = html[
        html.index("function _fhrpCanonicalJson"):
        html.index("function _fhrpSha256Ascii")
    ]
    assert "parts.length>=2000000" in fhrp_canonical


@pytest.mark.skipif(NODE is None, reason="node is required for Explorer cap tests")
def test_explorer_bounds_schema_max_blocker_dom_and_exports_every_safe_row(
    receipts,
):
    baseline = _schema_max_ipv6_receipt(receipts, degraded=True)
    encoded = json.dumps(
        baseline, sort_keys=True, separators=(",", ":"), ensure_ascii=True,
        allow_nan=False,
    )
    assert 25 * 1024 * 1024 < len(encoded) <= 32 * 1024 * 1024

    result = _explorer_section_cap_result(encoded)
    assert result["sectionLength"] < 1_000_000
    assert result["renderedCards"] == 200
    assert result["exactCounts"] is True
    assert result["exportControl"] is True
    assert result["exportCount"] == 20_000
    assert result["exportKeysSafe"] is True
    assert result["elapsedMs"] < 15_000
    assert result["rssBytes"] < 768 * 1024 * 1024


@pytest.mark.skipif(NODE is None, reason="node is required for Explorer receipt tests")
def test_explorer_fails_closed_on_python_casefold_host_collision(receipts):
    shared_commands = {
        "show ipv6 route summary": str(receipts["root"] / "bulk-edge-0.txt"),
        "show ospfv3 neighbor": str(receipts["root"] / "bulk-edge-1.txt"),
        "show bgp ipv6 unicast summary": str(receipts["root"] / "bulk-edge-2.txt"),
    }
    paths = {
        "alpha-edge": dict(shared_commands),
        "beta-edge": dict(shared_commands),
    }
    current = compute_ipv6_routing_adjacency_baseline(
        paths,
        compute_capture_integrity_from_paths(paths),
        {host: {"platform": "ios"} for host in paths},
    )
    baseline = embedded_ipv6_routing_adjacency_baseline(current)
    assert validate_ipv6_routing_adjacency_baseline(baseline)["valid"] is True

    renames = {"alpha-edge": "strasse", "beta-edge": "straße"}
    for row in baseline["rows"]:
        old_host = row["switch"]
        row["switch"] = renames[old_host]
        row["acceptance"] = row["acceptance"].replace(old_host, row["switch"])
    for cell in baseline["coverage"]:
        cell["switch"] = renames[cell["switch"]]
    for finding in baseline["findings"]:
        finding["switch"] = renames[finding["switch"]]

    protocol_order = {"OSPFv3": 0, "BGPv6": 1}
    input_order = {
        "route_summary": 0, "ospfv3_neighbors": 1, "bgp_ipv6_neighbors": 2,
    }
    baseline["rows"].sort(key=lambda row: (
        row["switch"].casefold(), row["switch"], protocol_order[row["protocol"]],
        row["peer_key"].casefold(), row["peer_key"],
    ))
    baseline["coverage"].sort(key=lambda cell: (
        cell["switch"].casefold(), cell["switch"], input_order[cell["input"]],
    ))
    baseline["findings"].sort(key=lambda finding: (
        finding["switch"].casefold(), finding["switch"],
        protocol_order[finding["protocol"]], finding["peer_key"].casefold(),
        finding["code"], finding["issue"],
    ))
    for cell in baseline["coverage"]:
        _rehash_coverage_cell(baseline, cell)
    baseline["summary"]["baseline_sha256"] = _canonical_digest(baseline)

    python = validate_ipv6_routing_adjacency_baseline(baseline)
    assert python["valid"] is False
    assert python["reason"] == "baseline_row_identity_collision"
    assert _explorer_result(baseline)["valid"] is False


@pytest.mark.skipif(NODE is None, reason="node is required for Explorer receipt tests")
def test_coverage_only_host_is_in_union_and_has_input_drilldown(receipts):
    baseline = receipts["embedded"]
    assert not any(row["switch"] == "coverage-only-edge" for row in baseline["rows"])
    coverage = [
        cell for cell in baseline["coverage"] if cell["switch"] == "coverage-only-edge"
    ]
    assert len(coverage) == 3
    assert all(cell["subject"] is False for cell in coverage)
    result = _explorer_result(baseline, "coverage-only-edge")
    assert "coverage-only-edge" in result["hosts"]
    assert "Host-family/input coverage" in result["rendered"]
    assert "route_summary" in result["rendered"]
    assert "No IPv6 adjacency rows were published for this host" in result["rendered"]
    for cell in coverage:
        assert cell["source_sha256"] not in result["rendered"]
        assert cell["projection_sha256"] not in result["rendered"]


def _validation_plan(items: list[dict]) -> dict:
    by_wave: dict[str, list[dict]] = {}
    by_category: dict[str, int] = {}
    for item in items:
        by_wave.setdefault(item["wave"], []).append(dict(item))
        by_category[item["category"]] = by_category.get(item["category"], 0) + 1
    return {
        "items": [dict(item) for item in items],
        "by_wave": by_wave,
        "summary": {
            "n_items": len(items), "n_waves": len(by_wave),
            "n_high": sum(item["severity"] in {"Critical", "High"} for item in items),
            "by_category": by_category,
        },
        "banner": "Run after each wave.",
    }


@pytest.mark.skipif(NODE is None, reason="node is required for current-baseline parity")
def test_ipv6_not_verified_marker_python_js_current_gate_parity(tmp_path):
    from cisco_toolkit.analyze import compute_current_baseline_gate

    item = {
        "device": "ipv6-edge", "platform": "ios", "wave": "(unscheduled)",
        "category": "Routing", "severity": "High",
        "check": "IPv6 routing evidence is not verified",
        "command": "show ospfv3 neighbor",
        "expect": (
            "IPV6 ROUTING BASELINE NOT VERIFIED — BLOCKER: re-collect the OSPFv3 "
            "and BGPv6 evidence before acceptance."
        ),
        "why": "The current IPv6 routing adjacency denominator is incomplete.",
        "evidence_state": "not_verified",
        "projection_custody": "embedded_unverified",
        "source_key": "ipv6_routing_adjacency_baseline.rows[ipv6-edge]",
    }
    plan = _validation_plan([item])
    expected = compute_current_baseline_gate(plan)
    html = EXPLORER.read_text(encoding="utf-8")
    match = re.search(r"REASONING-CORE-PORT START.*?REASONING-CORE-PORT END", html, re.S)
    assert match
    block = match.group(0)
    core = block[block.index("*/") + 2:block.rindex("/*")]
    script = (
        core
        + "\nconst fs=require('fs');"
        "const plan=JSON.parse(fs.readFileSync(0,'utf8'));"
        "process.stdout.write(JSON.stringify(computeCurrentBaselineGate(plan)));"
    )
    driver = tmp_path / "ipv6-current-baseline-driver.js"
    driver.write_text(script, encoding="utf-8")
    proc = subprocess.run(
        [NODE, str(driver)], input=json.dumps(plan), capture_output=True, check=True,
        text=True, encoding="utf-8", timeout=60,
    )
    actual = json.loads(proc.stdout)
    assert actual == expected
    assert actual["verdict"] == "INDETERMINATE"
    assert actual["summary"]["by_state"]["not_verified"] == 1
    assert actual["blockers"][0]["expect"].startswith(
        "IPV6 ROUTING BASELINE NOT VERIFIED — BLOCKER:"
    )


def test_explorer_mounts_ipv6_receipt_without_replacing_generic_health_or_waves():
    html = EXPLORER.read_text(encoding="utf-8")
    draw = html[html.index("function drawProtocols()"):html.index("function drawProtocolsDetail")]
    detail = html[html.index("function drawProtocolsDetail"):html.index("CAUSALITY DRAWER")]
    model = html[html.index("function buildModel(snap)"):html.index("function linkCarries(")]
    assert draw.count("${ipv6RoutingAdjacencySection()}") == 2
    assert "${ipv6RoutingAdjacencySection(host)}" in detail
    assert "${protoIntelSection()}" in draw and "${protoHealthSection()}" in draw
    assert "ipv6RoutingRows" in model and "ipv6RoutingCoverageSubjects" in model
    assert "computeCurrentBaselineGate" in html
    assert "IPV6 ROUTING BASELINE NOT VERIFIED — BLOCKER:" in html
