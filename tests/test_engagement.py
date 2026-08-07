"""NEW-V3.23.157: the Engagement Workflow & Plan of Record (DOCX) — the engagement-management
layer over the document set. python-docx is optional, so the module is skipped when it is absent
(the generator fails soft the same way). These tests pin the skeleton sections, the evidence-led
verdict (HOLD on blockers / PROCEED when clean), the pilot-first wave selection, the wave-schedule
reconciliation, the RAID seeding from real findings, the no-waves fallback, and the fail-soft path."""
import pytest

docx = pytest.importorskip("docx")  # skip the file if the optional dep is absent
from docx import Document  # noqa: E402

from cisco_toolkit.engagement import write_engagement_docx  # noqa: E402


def _snap():
    """A compact snapshot exercising every evidence-driven workflow section."""
    return {
        "script_version": "V3.23.0",
        "devices": {"core1": {}, "acc1": {}, "acc2": {}, "acc3": {}},
        "punchlist": [
            {"severity": "Critical", "category": "L3 design",
             "title": "VLAN 30 has a single gateway", "devices": ["core1"]},
            {"severity": "High", "category": "Physical",
             "title": "acc3 uplink is a single fiber", "devices": ["acc3"]},
        ],
        "migration_readiness": [
            {"group": "Group 1", "switches": ["acc1", "acc2"], "endpoints": 40,
             "readiness": "READY", "n_fail": 0},
            {"group": "Group 2", "switches": ["acc3"], "endpoints": 12,
             "readiness": "READY", "n_fail": 0},
            {"group": "Group 3", "switches": ["core1"], "endpoints": 90,
             "readiness": "NOT READY", "n_fail": 2},
        ],
        "wave_sequencing": [
            {"group": "Group 1", "make_before_break": ["acc1", "acc2"], "hard_cutover": [],
             "hard_cutover_endpoints": 0, "sequence": "make-before-break"},
            {"group": "Group 2", "make_before_break": [], "hard_cutover": ["acc3"],
             "hard_cutover_endpoints": 12, "sequence": "hard cutover"},
        ],
        "migration_scenarios": {
            # [audit-4 #9] the engine emits 'recommended_scenario' (analyze.compute_migration_scenarios), NOT
            # 'scenario' -- this fixture used the fabricated key that masked the §4.2 wave-table bug.
            "per_group": [{"group": "Group 1", "recommended_scenario": "parallel-run", "why": "dual-homed"},
                          {"group": "Group 2", "recommended_scenario": "phased", "why": "single-homed"}],
            "fleet_recommendation": "Parallel-run the dual-homed estate; window the rest.",
            "scenario_counts": {"parallel-run": 1, "phased": 1},
        },
        "validation_plan": {"by_wave": {"Group 1": [{}, {}, {}], "Group 2": [{}]},
                            "items": [], "summary": {}, "banner": ""},
        "collection_completeness": {
            "summary": {"inventory": 4, "complete": 3, "partial": 1, "not_collected": 0},
            "devices": [{"host": "acc3", "status": "partial", "data_quality": 60,
                         "missing": ["show cdp neighbors detail"]}],
        },
        "executive_brief": {"posture": "Fair", "posture_statement": "Fleet is migratable with care.",
                            "axes": [], "top_gating": ["1 Critical L3-design finding gates cutover"]},
        "lifecycle_risk": {"summary": {"n_devices": 4, "n_past_ldos": 1, "n_past_eos": 0}},
        "remediation_plan": {"summary": {"n_items": 5, "n_devices": 2}},
    }


def _all_text(doc):
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            parts.extend(c.text for c in row.cells)
    return "\n".join(parts)


def test_engagement_has_skeleton_sections_and_furniture(tmp_path):
    out = str(tmp_path / "e.docx")
    write_engagement_docx(out, _snap(), "Unit Test Fleet")
    d = Document(out)
    h1 = [p.text for p in d.paragraphs if p.style.name == "Heading 1"]
    for token in ("1. Engagement Verdict", "2. Engagement Phase Tracker", "3. Next Actions",
                  "4. Wave Gate Calendar", "5. RAID Log (seeded from the assessment)",
                  "6. Operating Rhythm"):
        assert any(t == token for t in h1), f"missing section: {token}; have {h1}"
    assert "Document Control" in h1 and "Document Acceptance" in h1
    text = _all_text(d)
    assert "Customer network owner" in text                     # acceptance roles
    assert "Assessment workbook (.xlsx)" in text                # related-documents cross-reference
    # self-exclusion: the doc never lists itself in its own related-documents table
    assert "Engagement Workflow & Plan of Record (.docx)" not in text


def test_engagement_inventoried_count_reads_canonical_scale(tmp_path):
    """SSOT: the 'N collected of M inventoried' line must take BOTH figures from the canonical
    executive_brief.scale (n_collected AND n_devices), not mix a canonical n_collected with a
    len(devices) recount — which renders a self-contradictory 'collected of <smaller> inventoried'
    when the inventory exceeds the (smaller) device map. Discriminating fixture: scale says 303
    devices while the devices map holds 4."""
    snap = _snap()
    snap.setdefault("executive_brief", {})["scale"] = {"n_devices": 303, "n_collected": 253}
    out = str(tmp_path / "e.docx")
    write_engagement_docx(out, snap, "Test Fleet")
    assert "253 collected of 303 inventoried" in _all_text(Document(out))


def test_engagement_verdict_holds_on_blockers_and_cites_evidence(tmp_path):
    out = str(tmp_path / "e.docx")
    write_engagement_docx(out, _snap(), "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "HOLD" in text
    assert "1 Critical punch-list item(s)" in text              # condition cites the punch-list
    assert "Group 3" in text                                    # the NOT READY group is named
    assert "Fleet is migratable with care." in text             # exec-brief posture statement carried
    assert "1 Critical L3-design finding gates cutover" in text  # top_gating headline carried


def test_engagement_verdict_proceeds_on_clean_fleet(tmp_path):
    snap = _snap()
    snap["punchlist"] = []
    snap["migration_readiness"] = [r for r in snap["migration_readiness"]
                                   if r["readiness"] == "READY"]
    snap["collection_completeness"] = {"summary": {"inventory": 4, "complete": 4,
                                                   "partial": 0, "not_collected": 0}, "devices": []}
    snap["lifecycle_risk"] = {"summary": {"n_devices": 4, "n_past_eos": 0}}
    out = str(tmp_path / "e.docx")
    write_engagement_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "PROCEED" in text and "HOLD" not in text
    # 'PROCEED' is a substring of the conditioned verdict — pin the UNCONDITIONAL one explicitly
    assert "PROCEED WITH CONDITIONS" not in text
    assert "No conditions attached" in text                     # no placeholder dressed as a condition
    assert "No open issues seeded" in text                      # the empty-issues path


def test_engagement_pilot_is_smallest_ready_group(tmp_path):
    out = str(tmp_path / "e.docx")
    write_engagement_docx(out, _snap(), "Unit Test Fleet")
    text = _all_text(Document(out))
    # Group 2 (READY, 12 endpoints) beats Group 1 (READY, 40); Group 3 (NOT READY) never qualifies
    assert "Group 2 (PILOT)" in text
    assert "Group 3 (PILOT)" not in text and "Group 1 (PILOT)" not in text
    assert "smallest blast radius" in text


def test_engagement_wave_schedule_reconciles_to_evidence(tmp_path):
    out = str(tmp_path / "e.docx")
    write_engagement_docx(out, _snap(), "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "parallel-run" in text and "phased" in text          # scenarios joined per wave
    assert "1 switch(es) / 12 endpoint(s)" in text              # hard-cutover exposure (Group 2)
    assert "Parallel-run the dual-homed estate; window the rest." in text   # fleet recommendation
    # T-minus gate cadence present
    for gate in ("T-28", "T-14", "T-1", "Go / No-Go", "Hypercare exit"):
        assert gate in text, gate


def test_engagement_raid_is_seeded_from_findings(tmp_path):
    out = str(tmp_path / "e.docx")
    write_engagement_docx(out, _snap(), "Unit Test Fleet")
    d = Document(out)
    text = _all_text(d)
    assert "RSK-001" in text and "VLAN 30 has a single gateway" in text     # risk from punch-list
    assert "past last-day-of-support (LDoS)" in text                        # risk from lifecycle (LDoS band)
    assert "ISS-001" in text and "show cdp neighbors detail" in text        # issue from blind spot
    assert "ASM-001" in text and "DEP-001" in text                          # assumptions + dependencies
    assert "DEC-001" in text and "PROPOSED" in text                         # decision log seeded
    lifecycle_rows = [
        [cell.text for cell in row.cells]
        for table in d.tables for row in table.rows
        if any("past last-day-of-support (LDoS)" in cell.text for cell in row.cells)
    ]
    assert lifecycle_rows and "Critical" in lifecycle_rows[0][0]


def test_engagement_gate_and_raid_key_on_past_ldos_not_past_eos(tmp_path):
    """A2 (coverage-honesty): the date band beyond LDoS is Past-LDoS, not Past-EoS (end-of-sale,
    with LDoS still future and entitlement unassessed). A fleet with 152 past-LDoS and 0 Past-EoS must surface the 152 in
    BOTH the go/no-go conditions and the RAID risk — reading n_past_eos (=0) silently drops the 152
    (the exact Meridian defect). Discriminating fixture: n_past_eos=0 so only an n_past_ldos reader passes."""
    snap = _snap()
    snap["lifecycle_risk"] = {"summary": {"n_devices": 303, "n_past_ldos": 152, "n_past_eos": 0,
                                          "n_near": 61}}
    out = str(tmp_path / "e.docx")
    write_engagement_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "152 device(s) are past last-day-of-support (LDoS)" in text       # go/no-go condition
    assert "152 device(s) past last-day-of-support (LDoS)" in text           # RAID risk row
    assert "152 device(s) are past end-of-support" not in text               # old conflated wording gone


def test_engagement_unknown_lifecycle_names_match_and_authority_causes(tmp_path):
    snap = _snap()
    snap["lifecycle_risk"] = {"summary": {"n_devices": 2, "n_unknown": 2}}
    out = str(tmp_path / "unknown-lifecycle.docx")
    write_engagement_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "either no exact EoX row matched" in text
    assert "matched row's retained source/date authority was incomplete" in text
    assert "support entitlement are UNDETERMINED" in text


def test_engagement_past_eos_is_refresh_window_not_support_entitlement(tmp_path):
    snap = _snap()
    snap["lifecycle_risk"] = {"summary": {"n_devices": 2, "n_past_eos": 2}}
    out = str(tmp_path / "past-eos.docx")
    write_engagement_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "finite recorded refresh window before LDoS" in text
    assert "date band does not establish support entitlement" in text
    assert "past end-of-sale with LDoS still future" in text
    assert "verify serial-numbered support entitlement separately" in text
    assert "support window closing" not in text


def test_engagement_near_ldos_is_seeded_into_conditions_and_raid(tmp_path):
    snap = _snap()
    snap["lifecycle_risk"] = {"summary": {"n_devices": 2, "n_near": 2}}
    out = str(tmp_path / "near-ldos.docx")
    write_engagement_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "2 device(s) reach last-day-of-support within the planning horizon" in text
    assert "2 device(s) are within one year of recorded LDoS" in text
    assert "date band does not establish it" in text


def test_engagement_no_waves_fallback(tmp_path):
    snap = _snap()
    snap["migration_readiness"] = []
    snap["wave_sequencing"] = []
    snap["migration_scenarios"] = {}
    snap["validation_plan"] = {}
    out = str(tmp_path / "e.docx")
    write_engagement_docx(out, snap, "Unit Test Fleet")
    d = Document(out)
    text = _all_text(d)
    assert "No migration waves are derivable" in text
    assert "(PILOT)" not in text
    h1 = [p.text for p in d.paragraphs if p.style.name == "Heading 1"]
    assert "Document Acceptance" in h1                          # furniture survives the fallback


def test_engagement_gate_record_renders_as_signed_trail(tmp_path):
    """V3.23.158: gate dispositions fed back from AssessHub land in §4.3, keyed by GATE_SEQUENCE."""
    record = {"Group 2": {
        "go_no_go": {"decision": "go", "signed_by": "A. Engineer",
                     "decided_at": "2026-06-10T21:00:00+00:00", "note": "all checks green"},
        "commit": {"decision": "slipped", "signed_by": "PM", "decided_at": "2026-06-01T09:00:00+00:00",
                   "note": ""},
    }}
    out = str(tmp_path / "e.docx")
    write_engagement_docx(out, _snap(), "Unit Test Fleet", gate_record=record)
    d = Document(out)
    h2 = [p.text for p in d.paragraphs if p.style.name == "Heading 2"]
    assert any("4.3 Gate record (as signed)" in t for t in h2)
    text = _all_text(d)
    assert "GO" in text and "SLIPPED" in text and "A. Engineer" in text
    assert "all checks green" in text
    assert "2026-06-10 21:00 UTC" in text                   # ISO stamp labeled, offset never dropped


def test_engagement_without_gate_record_has_no_signed_section(tmp_path):
    """No recorded state -> no §4.3; the document must not invent project state. An all-empty
    record (waves mapped to nothing) is the same as no record."""
    out = str(tmp_path / "e.docx")
    write_engagement_docx(out, _snap(), "Unit Test Fleet", gate_record={"Group 1": {}})
    d = Document(out)
    h2 = [p.text for p in d.paragraphs if p.style.name == "Heading 2"]
    assert not any("Gate record" in t for t in h2)


def test_gate_sequence_is_the_cadence_contract():
    """The webapp gate board and §4.1 both render from GATE_SEQUENCE — pin its keys and order."""
    from cisco_toolkit.engagement import GATE_SEQUENCE
    assert [k for k, *_ in GATE_SEQUENCE] == [
        "commit", "checkpoint", "readiness", "go_no_go", "window", "hypercare_exit"]
    for _key, label, when, purpose, criteria in GATE_SEQUENCE:
        assert label and when and purpose and criteria


def test_engagement_survives_malformed_snapshot(tmp_path):
    """V3.23.159: every read is defensive — junk shapes in user-uploaded snapshots degrade to
    placeholders instead of crashing the writer (the webapp would otherwise 500)."""
    snap = _snap()
    snap["punchlist"] = ["orphan note", None,
                         {"severity": "Critical", "title": "real finding", "category": "L3"}]
    snap["migration_scenarios"] = ["not", "a", "dict"]      # truthy non-dict block
    snap["validation_plan"] = "garbage"                     # truthy non-dict block
    snap["executive_brief"] = 7                             # truthy non-dict block
    snap["lifecycle_risk"] = {"summary": {"n_past_eos": "two"}}   # non-numeric count
    snap["migration_readiness"][0]["endpoints"] = "unknown"       # non-numeric endpoints
    snap["collection_completeness"]["devices"][0]["missing"] = "show cdp neighbors detail"
    out = str(tmp_path / "e.docx")
    write_engagement_docx(out, snap, "Unit Test Fleet")     # must not raise
    text = _all_text(Document(out))
    assert "real finding" in text                           # the one valid punch item survives
    assert "show cdp neighbors detail" in text              # string `missing` is ONE item,
    assert "s, h, o, w" not in text                         # never char-joined


@pytest.mark.parametrize("poison", [5, "boom", True, 3.14])
def test_engagement_survives_truthy_scalar_brief_scale(tmp_path, poison):
    """The INNER executive_brief.scale was the residue of the truthy-non-dict class: `_as_dict()`
    guarded only the OUTER section, so `(... .get("scale")) or {}` passed a truthy scalar straight
    into `.get("n_devices")` -> AttributeError inside _workflow_facts, BEFORE any rendering. The
    section-level test above cannot reach it (a scalar `executive_brief` is absorbed by _as_dict
    and never yields a `scale`). Reachable as a stored DoS through the webapp's uploaded snapshot."""
    snap = _snap()
    snap.setdefault("executive_brief", {})["scale"] = poison   # truthy non-dict INNER block
    out = str(tmp_path / "e.docx")
    write_engagement_docx(out, snap, "Unit Test Fleet")        # must not raise
    text = _all_text(Document(out))
    # scale unreadable -> the counts degrade to the len(devices) fallback, they are not invented
    assert "4 collected of 4 inventoried" in text


def test_engagement_absent_completeness_is_unknown_not_complete(tmp_path):
    """V3.23.159: a snapshot with NO collection_completeness block must read UNKNOWN — absence of
    evidence is not evidence of completeness, and the verdict picks up a condition."""
    snap = _snap()
    snap["punchlist"] = []
    snap["migration_readiness"] = [r for r in snap["migration_readiness"]
                                   if r["readiness"] == "READY"]
    snap["lifecycle_risk"] = {"summary": {"n_past_eos": 0}}
    del snap["collection_completeness"]
    out = str(tmp_path / "e.docx")
    write_engagement_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "completeness UNKNOWN" in text                   # §2 Assess status
    assert "collection is complete" not in text             # the old false claim
    assert "PROCEED WITH CONDITIONS" in text                # the gap is a verdict condition
    assert "Collection-completeness evidence is absent" in text


def test_engagement_pilot_never_won_by_unknown_endpoints(tmp_path):
    """V3.23.159: a READY group with endpoints null/non-numeric must not masquerade as the
    smallest blast radius — and the pilot row never renders 'None endpoint(s)'."""
    snap = _snap()
    snap["migration_readiness"] = [
        {"group": "Group 1", "switches": ["a"], "endpoints": None, "readiness": "READY", "n_fail": 0},
        {"group": "Group 2", "switches": ["b"], "endpoints": 12, "readiness": "READY", "n_fail": 0},
    ]
    out = str(tmp_path / "e.docx")
    write_engagement_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "Group 2 (PILOT)" in text and "Group 1 (PILOT)" not in text
    assert "None endpoint(s)" not in text


def test_engagement_gate_record_calendar_order_and_stray_waves(tmp_path):
    """V3.23.159: §4.3 follows the §4.2 calendar order (natural-numeric, never lexicographic),
    and rows signed against waves outside this snapshot's calendar are fenced off, not mixed in."""
    record = {
        "Group 2": {"commit": {"decision": "go", "signed_by": "A", "decided_at": "", "note": ""}},
        "Group 10": {"commit": {"decision": "go", "signed_by": "B", "decided_at": "", "note": ""}},
    }
    out = str(tmp_path / "e.docx")
    write_engagement_docx(out, _snap(), "Unit Test Fleet", gate_record=record)
    text = _all_text(Document(out))
    # Group 2 is in the fixture's calendar; Group 10 is not -> stray section with the caveat.
    # Scope the ordering assertion to the trail itself (the doc mentions Group 2 earlier too).
    trail = text[text.index("Gate record (as signed)"):]
    assert "NOT in this snapshot's calendar" in trail
    assert trail.index("Group 2") < trail.index("Group 10")


def test_engagement_phase_tracker_never_claims_generation(tmp_path):
    """V3.23.159: the tracker must not assert drafts exist that this run cannot evidence."""
    out = str(tmp_path / "e.docx")
    write_engagement_docx(out, _snap(), "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "Draft generated" not in text and "Drafts generated" not in text
    assert "Pending — requirements workshop not yet held" in text


def test_engagement_failsoft_without_python_docx(monkeypatch, tmp_path):
    import builtins, os
    real_import = builtins.__import__

    def _blocked(name, *a, **k):
        if name == "docx" or name.startswith("docx."):
            raise ImportError("simulated missing python-docx")
        return real_import(name, *a, **k)

    monkeypatch.setattr(builtins, "__import__", _blocked)
    out = str(tmp_path / "e.docx")
    write_engagement_docx(out, _snap(), "Unit Test Fleet")   # must not raise
    assert not os.path.exists(out)


def test_engagement_title_does_not_falsely_claim_full_collection(tmp_path):
    """[audit-2 #15] 'n_collected or n_devices' turned n_collected==0 into n_devices -> a NOT-collected fleet read
    '303 collected of 303 inventoried'. Use a proper None check."""
    from cisco_toolkit.engagement import write_engagement_docx
    from docx import Document
    out = str(tmp_path / "eng.docx")
    write_engagement_docx(out, {"executive_brief": {"scale": {"n_devices": 303, "n_collected": 0}},
                                "devices": {f"s{i}": {} for i in range(303)}}, "Meridian")
    txt = " ".join(p.text for p in Document(out).paragraphs)
    assert "0 collected of 303" in txt


def test_engagement_wave_endpoint_column_is_per_port_qualified(tmp_path):
    """QA row-13: §4.2's bare 'Endpoints' header invited summing the per-wave per-port MAC counts
    against the fleet-unique endpoint figure (a MAC behind several waves counts in each, so the column
    overlaps). The column must carry the per-port relabel its workbook/MOP siblings use, plus the
    overlap note."""
    from cisco_toolkit.engagement import write_engagement_docx
    from docx import Document
    out = str(tmp_path / "e.docx")
    write_engagement_docx(out, _snap(), "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "Endpoint MACs (per-port)" in text
    assert "counts in each" in text and "fleet-unique" in text
