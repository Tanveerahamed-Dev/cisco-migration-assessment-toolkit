"""Round-6 silent-truncation sweep of the runbook + operations-handbook display caps.

Silent truncation was found four separate times in four different deliverable writers during the
whole-repo review. The house rule is that a display cap must DISCLOSE what it dropped — a trailing
``(+N)`` marker or an "…and N further …" sentence (`excel._xls_cell_value` states the rule; mop.py
§x.2 and html.py's Findings-Delta column were fixed to it). `cisco_toolkit/runbook.py` and
`cisco_toolkit/ops.py` carried 27 undisclosed reachable caps.

The worst was runbook §6.2: §1's metric register states the FULL Critical/High cross-layer total
while §6.2 renders at most 8 finding blocks. On the 303-device production snapshot that is
"3 / 404" in the headline over 8 rendered blocks — a document contradicting itself, with 505
findings dropped and no marker anywhere for the reader to notice.

Fixtures come from the REAL producers (`tests/golden/snapshot.json` and
`webapp/sample_data/sample_fleet.snapshot.json`); rows are REPLICATED from those snapshots to push a
population past its cap rather than hand-authored into the shape the writer expects. Neither
snapshot file is modified. (The 303-device `Migration_Assessment_AUTOFILLED_*.snapshot.json` that
the caps were MEASURED on is gitignored, so it is deliberately not a test input.)
"""
import copy
import json
import os
import re

import pytest

docx = pytest.importorskip("docx")            # the deliverables need the optional python-docx
from docx import Document                     # noqa: E402

from cisco_toolkit.runbook import write_runbook_docx    # noqa: E402
from cisco_toolkit.ops import write_ops_handbook_docx   # noqa: E402

_GOLDEN = os.path.join(os.path.dirname(__file__), "golden", "snapshot.json")
_SAMPLE = os.path.join(os.path.dirname(__file__), os.pardir, "webapp", "sample_data",
                       "sample_fleet.snapshot.json")

# The disclosure sentence every capped table now emits, e.g.
#   "…and 288 further asset(s) not shown (15 of 303 rendered) — see the workbook's '…' sheet."
_DISCLOSURE = re.compile(r"…and (\d+) further (.+?) not shown \((\d+) of (\d+) rendered\)")


def _all_text(doc):
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            parts.extend(c.text for c in row.cells)
    return "\n".join(parts)


@pytest.fixture(scope="module")
def golden():
    with open(_GOLDEN, encoding="utf-8") as f:
        return json.load(f)


@pytest.fixture(scope="module")
def sample():
    with open(_SAMPLE, encoding="utf-8") as f:
        return json.load(f)


def _grow(rows, n, vary=("host",)):
    """Replicate REAL producer rows until there are exactly `n`, giving each copy a distinct identity
    in `vary` so the writer treats them as separate records (a fixture manufactured in the shape the
    writer expects would agree with the writer's own bugs; these rows come from the snapshot)."""
    rows = [r for r in rows if isinstance(r, dict)]
    assert rows, "the fixture snapshot must supply at least one real row to replicate"
    out = []
    for i in range(n):
        r = copy.deepcopy(rows[i % len(rows)])
        for k in vary:
            if k in r or i >= len(rows):
                r[k] = f"{r.get(k, 'row')}-r6-{i:03d}"
        out.append(r)
    return out


def _runbook_text(tmp_path, snap, name="rb.docx"):
    out = str(tmp_path / name)
    write_runbook_docx(out, snap, "R6 Cap Fleet")
    return _all_text(Document(out))


def _ops_text(tmp_path, snap, name="ops.docx"):
    out = str(tmp_path / name)
    write_ops_handbook_docx(out, snap, "R6 Cap Fleet")
    return _all_text(Document(out))


# --------------------------------------------------------------------------------------------
# §6.2 — the flagship: a headline the reader cannot reconcile with the list beneath it
# --------------------------------------------------------------------------------------------
def _cross_layer_20(golden):
    """20 Critical/High cross-layer findings built from the golden snapshot's own rows: 3 Critical +
    17 High. §1's register then reads '3 / 17' while §6.2 renders 8 blocks (3 Critical + 5 High)."""
    snap = copy.deepcopy(golden)
    rows = [r for r in snap.get("cross_layer") or [] if isinstance(r, dict)]
    crit = [r for r in rows if r.get("severity") == "Critical"]
    high = [r for r in rows if r.get("severity") == "High"]
    assert crit and high, "golden must carry a real Critical and a real High cross-layer finding"
    snap["cross_layer"] = _grow(crit, 3, vary=("id",)) + _grow(high, 17, vary=("id",))
    return snap


def test_cross_layer_section_discloses_the_findings_its_cap_dropped(tmp_path, golden):
    """§6.2 renders at most 8 finding blocks while §1's metric register counts the FULL Critical/High
    population. Without the disclosure the runbook states 20 Critical/High on its front page and
    shows 8 blocks with no marker — the reader concludes the 8 blocks ARE the cross-layer exposure."""
    text = _runbook_text(tmp_path, _cross_layer_20(golden), "rb_cl.docx")

    # the §1 headline is unchanged and still states the full population
    assert "3 / 17" in text, "§1's metric register must still carry the full Critical/High total"

    # …and §6.2 now reconciles to it, naming the true remainder
    assert "…and 12 further cross-layer finding(s) are NOT rendered as blocks above" in text
    assert "the 8 highest-severity of 20 are shown" in text
    assert "covering 8 of the 20 Critical/High finding(s) that §1's metric register counts" in text
    assert "(3 Critical / 17 High)" in text
    assert "'Cross-Layer Analysis' sheet" in text          # where the complete list lives
    # shown + hidden reconcile to the real total
    assert 8 + 12 == len(_cross_layer_20(golden)["cross_layer"])


def test_cross_layer_disclosure_is_silent_when_nothing_was_dropped(tmp_path, golden):
    """The golden fleet carries fewer findings than the cap — a disclosure there would be noise, and
    worse, would imply an omission that never happened."""
    text = _runbook_text(tmp_path, golden, "rb_cl_small.docx")
    assert "further cross-layer finding(s) are NOT rendered" not in text


# --------------------------------------------------------------------------------------------
# every capped runbook table
# --------------------------------------------------------------------------------------------
def _b_scenarios(g, _s):
    snap = copy.deepcopy(g)
    snap["migration_scenarios"]["per_group"] = _grow(snap["migration_scenarios"]["per_group"], 20,
                                                     vary=("group",))
    return snap, 20


def _b_models(g, _s):
    snap = copy.deepcopy(g)
    base = dict(next(iter(snap["devices"].values())))
    for i in range(20):
        d = dict(base); d["hostname"] = f"r6sw{i:02d}"; d["model"] = f"C9300-R6-{i:02d}"
        snap["devices"][f"r6sw{i:02d}"] = d
    return snap, 20 + len({(d or {}).get("model") or "unknown"
                           for d in g["devices"].values() if isinstance(d, dict)})


def _b_bridges(g, _s):
    snap = copy.deepcopy(g)
    br = [r for r in snap["link_centrality"] if r.get("is_bridge")]
    snap["link_centrality"] = _grow(br, 20, vary=("a_host", "b_host"))
    return snap, 20


def _b_gateways(g, _s):
    snap = copy.deepcopy(g)
    snap["l3_forwarding"] = _grow(snap["l3_forwarding"], 30, vary=("vlan",))
    return snap, 30


def _b_l3_devices(g, _s):
    snap = copy.deepcopy(g)
    l3 = [r for r in snap["subnet_intelligence"]["per_device"] if r.get("is_l3")]
    snap["subnet_intelligence"]["per_device"] = _grow(l3, 20)
    return snap, 20


def _b_si_move_groups(g, _s):
    snap = copy.deepcopy(g)
    snap["subnet_intelligence"]["move_groups"] = _grow(snap["subnet_intelligence"]["move_groups"],
                                                       20, vary=("group",))
    return snap, 20


def _b_protocol_intel(g, _s):
    snap = copy.deepcopy(g)
    snap["protocol_intelligence"] = _grow(snap["protocol_intelligence"], 25, vary=("switch",))
    return snap, 25


def _b_domain_edges(g, _s):
    snap = copy.deepcopy(g)
    snap["application_intelligence"]["edges"] = _grow(snap["application_intelligence"]["edges"],
                                                      20, vary=("source", "target"))
    return snap, 20


def _b_syslog(g, _s):
    snap = copy.deepcopy(g)
    snap["syslog_intelligence"]["detections"] = _grow(snap["syslog_intelligence"]["detections"], 30)
    return snap, 30


def _b_qos(g, _s):
    snap = copy.deepcopy(g)
    snap["qos_audit"]["findings"] = _grow(snap["qos_audit"]["findings"], 20)
    return snap, 20


def _b_software(g, _s):
    snap = copy.deepcopy(g)
    snap["software_risk"]["findings"] = _grow(snap["software_risk"]["findings"], 20)
    return snap, 20


def _b_trains(g, _s):
    snap = copy.deepcopy(g)
    worst = [d for d in snap["software_risk"]["per_device"]
             if d.get("train_band") in ("Replace/Upgrade", "Verify EoL")]
    snap["software_risk"]["per_device"] = _grow(worst, 20)
    return snap, 20


def _b_capacity(g, _s):
    snap = copy.deepcopy(g)
    snap["capacity"] = _grow(snap["capacity"], 20, vary=("hostname",))
    return snap, 20


def _b_clusters(g, _s):
    snap = copy.deepcopy(g)
    # neither tracked snapshot carries a cohesive-unit cluster, so the rows are built to the shape
    # the §8.1 table reads; the CAP under test is count-driven, not shape-driven.
    snap.setdefault("endpoint_dependencies", {})["clusters"] = [
        {"endpoint_class": "VM / Hypervisor", "vendor": f"Vendor-{i:02d}", "count": 10 + i,
         "switches": 3, "vlans": 2, "spans_groups": bool(i % 2)} for i in range(20)]
    return snap, 20


def _b_affinity(g, s):
    snap = copy.deepcopy(g)
    aff = [r for r in (s.get("endpoint_dependencies") or {}).get("affinity") or []
           if isinstance(r, dict)]
    snap.setdefault("endpoint_dependencies", {})["affinity"] = _grow(aff, 20, vary=("vlan",))
    return snap, 20


def _b_failure_impact(g, _s):
    snap = copy.deepcopy(g)
    snap["failure_impact"] = _grow(snap["failure_impact"], 20)
    return snap, 20


def _b_dossiers(g, s):
    snap = copy.deepcopy(g)
    per = [r for r in (s.get("device_dossiers") or {}).get("per_device") or []
           if isinstance(r, dict)]
    rows = _grow(per, 20)
    snap["device_dossiers"] = {"per_device": rows,
                               "summary": dict((s.get("device_dossiers") or {}).get("summary") or {},
                                               n_devices=20)}
    return snap, 20


def _b_compound(g, s):
    """20 compound patterns spread over 4 assets — §10.1's paragraph states the full count."""
    snap = copy.deepcopy(g)
    per = [r for r in (s.get("device_dossiers") or {}).get("per_device") or []
           if isinstance(r, dict) and r.get("compound")]
    assert per, "the sample fleet must carry at least one asset with a compound pattern"
    rows = _grow(per, 4)
    for i, r in enumerate(rows):
        base = [c for c in r.get("compound") or [] if isinstance(c, dict)]
        r["compound"] = [dict(base[0], code=f"CMP-{i}{j}") for j in range(5)]
    snap["device_dossiers"] = {"per_device": rows,
                               "summary": {"n_devices": 4, "n_compound": 20, "bands": {}}}
    return snap, 20


_RUNBOOK_CAPS = [
    ("§3   scenarios",        _b_scenarios,      12, "move-group scenario row(s)"),
    ("§4   model mix",        _b_models,         15, "model/PID row(s)"),
    ("§5   bridge links",     _b_bridges,        12, "bridge link(s)"),
    ("§6.1 gateways",         _b_gateways,       25, "gateway SVI record(s)"),
    ("§6.4 L3 devices",       _b_l3_devices,     12, "L3 device(s)"),
    ("§6.4 group subnets",    _b_si_move_groups, 12, "move-group subnet row(s)"),
    ("§6.5 protocol intel",   _b_protocol_intel, 18, "abnormal control-plane state(s)"),
    ("§6.7.1 domain edges",   _b_domain_edges,   15, "dependency edge(s)"),
    ("§6.10 syslog",          _b_syslog,         20, "operational detection(s)"),
    ("§6.11 qos",             _b_qos,            15, "QoS finding(s)"),
    ("§6.12 advisories",      _b_software,       15, "exposed advisory surface(s)"),
    ("§6.12 trains",          _b_trains,         12, "device(s) on a Replace/Upgrade or Verify-EoL train"),
    ("§8   capacity",         _b_capacity,       10, "switch(es) with a measured port utilisation"),
    ("§8.1 clusters",         _b_clusters,       12, "cohesive unit(s)"),
    ("§8.1 vlan affinity",    _b_affinity,       12, "VLAN(s)"),
    ("§10  blast radius",     _b_failure_impact, 15, "switch blast-radius row(s)"),
    ("§10.1 assets",          _b_dossiers,       15, "asset(s)"),
    ("§10.1 compound",        _b_compound,       10, "compound pattern(s)"),
]


@pytest.mark.parametrize("section,builder,cap,noun",
                         _RUNBOOK_CAPS, ids=[c[0] for c in _RUNBOOK_CAPS])
def test_every_runbook_display_cap_discloses_its_remainder(tmp_path, golden, sample,
                                                           section, builder, cap, noun):
    """Each of these tables silently rendered its first N rows of a larger population. The reader of
    a capped table has no cue that anything is missing, and every one of these sections states its
    FULL count in the paragraph above the table."""
    snap, total = builder(golden, sample)
    text = _runbook_text(tmp_path, snap, f"rb_{abs(hash(section)):x}.docx")
    expected = f"…and {total - cap} further {noun} not shown ({cap} of {total} rendered)"
    assert expected in text, f"{section}: missing disclosure — expected {expected!r}"


def test_endpoint_vendor_table_discloses_the_vendors_it_dropped(tmp_path, golden):
    """§7.1's 'Top vendors' renders 12 rows. §7's two 'Top …' tables are reconcilable — the sentence
    above them states the VLAN and switch populations — but the distinct-vendor count appears NOWHERE
    else in the runbook, so 12 rows read as the fleet's whole vendor spread (105 on the production
    snapshot)."""
    snap = copy.deepcopy(golden)
    ident = [r for r in snap.get("endpoint_identity") or [] if isinstance(r, dict)]
    rows = _grow(ident, 20, vary=("mac",))
    for i, r in enumerate(rows):
        r["vendor"] = f"Vendor {i:02d}"
    snap["endpoint_identity"] = rows
    text = _runbook_text(tmp_path, snap, "rb_vendors.docx")
    assert "…and 8 further vendor(s) not shown (12 of 20 rendered)" in text
    assert "'Endpoint Intelligence' sheet" in text


def test_no_runbook_cap_fires_on_a_fleet_that_does_not_reach_it(tmp_path, golden):
    """The golden fleet is under every cap; a disclosure there would assert an omission that never
    happened (the mirror-image false claim)."""
    text = _runbook_text(tmp_path, golden, "rb_none.docx")
    assert not _DISCLOSURE.search(text), _DISCLOSURE.findall(text)


def test_runbook_disclosures_reconcile_shown_plus_hidden_to_the_real_total(tmp_path, golden, sample):
    """Every disclosure must be arithmetically honest: hidden + shown == total. A disclosure naming
    the wrong remainder is a new false claim, not a fix."""
    # _b_dossiers and _b_compound both own snap["device_dossiers"], so they cannot share one
    # snapshot — the cumulative build takes every other cap and compound renders on its own.
    stacked = [b for b in _RUNBOOK_CAPS if b[1] is not _b_compound]
    snap = copy.deepcopy(golden)
    for _, builder, _, _ in stacked:
        snap, _total = builder(snap, sample)
    texts = [_runbook_text(tmp_path, snap, "rb_all.docx"),
             _runbook_text(tmp_path, _b_compound(golden, sample)[0], "rb_all_cmp.docx")]
    found = [m for t in texts for m in _DISCLOSURE.findall(t)]
    fired = {noun for _h, noun, _s, _t in found}
    assert fired >= {noun for _sec, _b, _cap, noun in _RUNBOOK_CAPS}, \
        f"caps that did not fire: {{n for _s, _b, _c, n in _RUNBOOK_CAPS}} - {fired}"
    for hidden, noun, shown, total in found:
        assert int(hidden) + int(shown) == int(total), (noun, hidden, shown, total)


def test_top_gating_list_marks_the_items_it_dropped(tmp_path, golden):
    """§1's 'Address first:' line joined the first 6 gating items and presented them as the whole
    list — the one line a war-room reader takes away."""
    snap = copy.deepcopy(golden)
    eb = dict(snap.get("executive_brief") or {})
    eb["axes"] = eb.get("axes") or [{"axis": "Resilience", "severity": "High", "headline": "h"}]
    eb["top_gating"] = [f"gating item {i}" for i in range(9)]
    snap["executive_brief"] = eb
    text = _runbook_text(tmp_path, snap, "rb_gating.docx")
    assert "(+3 more)" in text
    assert "gating item 5" in text and "gating item 6" not in text   # 6 shown, 3 disclosed


# --------------------------------------------------------------------------------------------
# string caps: a cut that eats the actionable half of the cell, and a cut with no marker
# --------------------------------------------------------------------------------------------
def test_protocol_intelligence_cell_keeps_the_remediation_the_cap_used_to_evict(tmp_path, golden):
    """§6.5 rendered (likely_cause + '→' + remediation)[:200]: the slice cuts from the END, so a long
    cause ate the REMEDIATION — the only actionable half. This is the identical bug §6.3's _why_next
    was already fixed for. On the production snapshot 10 of 84 rows lost 44–79 characters."""
    snap = copy.deepcopy(golden)
    row = copy.deepcopy([r for r in snap["protocol_intelligence"] if isinstance(r, dict)][0])
    row["likely_cause"] = ("a long inferred cause that fills the cell " * 5).strip()   # ~205 chars
    row["remediation"] = "configure channel-group mode active on BOTH ends and check lacp min-links"
    snap["protocol_intelligence"] = [row]
    text = _runbook_text(tmp_path, snap, "rb_pi_cell.docx")
    assert "lacp min-links" in text, "the remediation must survive the cell budget"
    assert "…" in text                       # and the cause's own cut is marked, not silent


def test_device_verdict_cut_is_marked_instead_of_ending_mid_word(tmp_path, golden, sample):
    """§10.1's 'Engineer's verdict' cell used a bare [:220]: a mid-word cut with no marker, so the
    verdict simply stopped and read as a finished sentence."""
    snap, _ = _b_dossiers(golden, sample)
    long_verdict = ("This asset stacks an end-of-support platform with a sole gateway and a "
                    "single uplink; treat every change to it as high risk and sequence it late "
                    "in the migration after its redundancy gap has been closed and independently "
                    "validated against the pre-cutover baseline.")
    assert len(long_verdict) > 220
    snap["device_dossiers"]["per_device"][0]["verdict"] = long_verdict
    text = _runbook_text(tmp_path, snap, "rb_verdict.docx")
    cell = next(line for line in text.splitlines() if line.startswith("This asset stacks"))
    assert cell.endswith("…"), f"verdict cut left unmarked: {cell[-60:]!r}"
    assert long_verdict[:60] in cell          # the head of the verdict is intact


# --------------------------------------------------------------------------------------------
# the operations handbook
# --------------------------------------------------------------------------------------------
def _o_keystones(g, _s):
    snap = copy.deepcopy(g)
    fi = [r for r in snap["failure_impact"] if isinstance(r, dict) and (r.get("stranded") or 0) > 0]
    assert fi, "golden must carry a real endpoint-stranding failure-impact row"
    snap["failure_impact"] = _grow(fi, 20)
    return snap, 20


def _o_detections(g, _s):
    snap = copy.deepcopy(g)
    snap["syslog_intelligence"]["detections"] = _grow(snap["syslog_intelligence"]["detections"], 30)
    return snap, 30


def _o_capacity_baseline(g, _s):
    snap = copy.deepcopy(g)
    per = [d for d in snap["platform_health"]["per_device"] if d.get("collected")]
    assert per, "golden must carry a real collected platform-health sample"
    snap["platform_health"]["per_device"] = _grow(per, 40)
    return snap, 40


def _o_trains(g, s):
    return _b_trains(g, s)


_OPS_CAPS = [
    ("§2.1 keystones", _o_keystones,          5,  "device(s) whose loss strands endpoints"),
    ("§3.1 detections", _o_detections,        12, "detection(s)"),
    ("§3.2 baselines", _o_capacity_baseline,  30, "sampled device(s)"),
    ("§5   trains",    _o_trains,             12, "device(s) on a Replace/Upgrade or Verify-EoL train"),
]


@pytest.mark.parametrize("section,builder,cap,noun", _OPS_CAPS, ids=[c[0] for c in _OPS_CAPS])
def test_every_ops_handbook_display_cap_discloses_its_remainder(tmp_path, golden, sample,
                                                                section, builder, cap, noun):
    """The handbook is the NOC's working document: §3.2's table IS the per-device alerting baseline
    and §2.1's table IS the high-risk-change list, so a silently short table under-scopes day-2 work."""
    snap, total = builder(golden, sample)
    text = _ops_text(tmp_path, snap, f"ops_{abs(hash(section)):x}.docx")
    expected = f"…and {total - cap} further {noun} not shown ({cap} of {total} rendered)"
    assert expected in text, f"{section}: missing disclosure — expected {expected!r}"


def test_ops_known_issue_signature_list_marks_the_classes_it_dropped(tmp_path, golden):
    """§7's Known-Issues register names the top 6 signature CLASSES; §3.1's alert list is built from
    exactly that register, so an undisclosed cut under-scopes the alerting the NOC configures."""
    snap = copy.deepcopy(golden)
    dets = [d for d in snap["syslog_intelligence"]["detections"] if isinstance(d, dict)]
    grown = _grow(dets, 12)
    for i, d in enumerate(grown):            # 12 DISTINCT signature classes
        d["label"] = f"Signature class {i:02d}"
        d["count"] = 100 - i
    snap["syslog_intelligence"]["detections"] = grown
    text = _ops_text(tmp_path, snap, "ops_sigs.docx")
    assert "(+6 further signature class(es) not listed" in text
    assert "Syslog Intelligence sheet" in text


def test_ops_facts_expose_the_pre_cap_keystone_total(golden):
    """_facts() caps `keystones` at 5; §2.1 cannot disclose what it never sees, so the pre-cap total
    travels with it."""
    from cisco_toolkit.ops import _facts
    snap = copy.deepcopy(golden)
    fi = [r for r in snap["failure_impact"] if isinstance(r, dict) and (r.get("stranded") or 0) > 0]
    snap["failure_impact"] = _grow(fi, 20)
    ev = _facts(snap)
    assert len(ev["keystones"]) == 5
    assert ev["n_keystones"] == 20


def test_no_ops_cap_fires_on_a_fleet_that_does_not_reach_it(tmp_path, golden):
    text = _ops_text(tmp_path, golden, "ops_none.docx")
    assert not _DISCLOSURE.search(text), _DISCLOSURE.findall(text)
