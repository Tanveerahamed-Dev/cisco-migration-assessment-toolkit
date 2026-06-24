"""Cross-surface single-source-of-truth (SSOT) reconciliation harness.

The engine publishes each headline fact ONCE (executive_brief.scale / .posture,
lifecycle_risk.summary, design_blueprint.summary); every deliverable and dashboard must READ
those canonical fields, never recompute or conflate them. Piecemeal per-surface locks already
exist (test_design_scale_reads_canonical_executive_brief, test_deck_title_slide_always_shows_
canonical_scale, the runbook endpoint-label lock, archreview's LC-1 EoS/LDoS exclusivity). What
those miss -- and what let the n_past_eos/n_past_ldos conflation slip surface-by-surface across
crd -> runbook -> engagement -> deck over several audit waves -- is a UNIFIED guard:

  1. a producer-side internal-consistency invariant (cisco_toolkit.ssot.reconcile) that proves
     every published canonical value still matches its raw-evidence derivation, and
  2. a cross-surface behavioural lock that renders the lifecycle/scale deliverables from a
     fixture carrying the exact AJ trap values (n_past_ldos=152, n_past_eos=0) and asserts each
     surface headlines the past-SUPPORT population (152), never the conflated sibling (0).

Together they convert "re-find the drift by hand every wave" into "CI catches it the moment it
is reintroduced." python-docx is optional, so the render tests skip cleanly when it is absent.
"""
import copy
import importlib

import pytest

from cisco_toolkit import ssot


# --------------------------------------------------------------------------------------------
# Producer-side internal-consistency invariant (no rendering -- fast, portable, mutation-proof)
# --------------------------------------------------------------------------------------------
def _consistent_snap():
    """A small, fully self-consistent snapshot: every published canonical value is derivable
    from the raw arrays it summarises. 6 devices (3 Critical / 2 Poor / 1 Good health band;
    4 Past-LDoS / 1 Near-LDoS / 1 Active lifecycle band), 40 endpoints, 5 collected."""
    health = (  # scores chosen so round(mean) == posture.avg_health (40): (15*3 + 50*2 + 95) / 6 = 40
        [{"switch": f"crit{i}", "band": "Critical", "score": 15} for i in range(3)]
        + [{"switch": f"poor{i}", "band": "Poor", "score": 50} for i in range(2)]
        + [{"switch": "good0", "band": "Good", "score": 95}]
    )
    per_device = (
        [{"host": f"crit{i}", "band": "Past-LDoS"} for i in range(4)]
        + [{"host": "near0", "band": "Near-LDoS"}]
        + [{"host": "good0", "band": "Active"}]
    )
    return {
        "health_scores": health,
        "endpoint_identity": [{"mac": f"00:00:00:00:00:{i:02x}"} for i in range(40)],
        "lifecycle_risk": {
            "summary": {
                "n_devices": 6, "n_past_ldos": 4, "n_past_eos": 0, "n_near": 1, "n_active": 1,
                "by_band": {"Past-LDoS": 4, "Near-LDoS": 1, "Active": 1},
            },
            "per_device": per_device,
        },
        "collection_completeness": {"summary": {"inventory": 6, "complete": 5, "partial": 0, "not_collected": 1}},
        "executive_brief": {
            "scale": {"n_devices": 6, "n_collected": 5, "n_endpoints": 40, "n_domains": 2},
            "posture": {"avg_health": 40, "n_critical": 3, "n_poor": 2, "worst_band": "Critical"},
        },
        "design_blueprint": {"summary": {"n_decisions": 2}, "decisions": [{"id": "a"}, {"id": "b"}]},
    }


def test_reconcile_is_clean_on_a_consistent_snapshot():
    assert ssot.reconcile(_consistent_snap()) == []


def test_reconcile_catches_device_count_recount():
    snap = _consistent_snap()
    snap["executive_brief"]["scale"]["n_devices"] = 999  # a len(devices) recount gone wrong
    viol = ssot.reconcile(snap)
    assert any("n_devices" in v and "len(health_scores)" in v for v in viol), viol


def test_reconcile_catches_the_ldos_eos_conflation():
    """THE recurring trap: a surface (or the producer) reports the past-support population as
    n_past_eos (0) instead of n_past_ldos (4). The producer-side guard must bite immediately."""
    snap = _consistent_snap()
    snap["lifecycle_risk"]["summary"]["n_past_ldos"] = 0  # silently drop the 4 past-LDoS devices
    viol = ssot.reconcile(snap)
    assert any("n_past_ldos" in v and "Past-LDoS" in v for v in viol), viol


def test_reconcile_catches_posture_band_drift():
    snap = _consistent_snap()
    snap["executive_brief"]["posture"]["n_critical"] = 99
    viol = ssot.reconcile(snap)
    assert any("n_critical" in v and "Critical" in v for v in viol), viol


def test_reconcile_catches_avg_health_and_worst_band_lies():
    """avg_health and worst_band are published aggregates COUNTED as self-verified by summary(), so
    a lie in either must be caught -- the gap an adversarial review found (counted but unchecked):
    worst_band='Excellent' with Critical devices present, or an out-of-range avg_health, slipped
    through verified=True. Both reconcile EXACTLY to compute_executive_brief's formula."""
    bad_avg = _consistent_snap()
    bad_avg["executive_brief"]["posture"]["avg_health"] = 999  # impossible (scores are 0-100)
    assert any("avg_health" in v for v in ssot.reconcile(bad_avg)), ssot.reconcile(bad_avg)
    assert ssot.summary(bad_avg)["verified"] is False
    bad_worst = _consistent_snap()
    bad_worst["executive_brief"]["posture"]["worst_band"] = "Excellent"  # while 3 Critical devices exist
    assert any("worst_band" in v for v in ssot.reconcile(bad_worst)), ssot.reconcile(bad_worst)
    assert ssot.summary(bad_worst)["verified"] is False


def test_reconcile_catches_endpoint_and_collected_drift():
    snap = _consistent_snap()
    snap["executive_brief"]["scale"]["n_endpoints"] = 7
    snap["executive_brief"]["scale"]["n_collected"] = 1
    viol = ssot.reconcile(snap)
    assert any("n_endpoints" in v for v in viol), viol
    assert any("n_collected" in v for v in viol), viol


def test_reconcile_catches_by_band_and_decision_drift():
    snap = _consistent_snap()
    snap["lifecycle_risk"]["summary"]["by_band"]["Past-LDoS"] = 1  # inconsistent with 4 in per_device
    snap["design_blueprint"]["summary"]["n_decisions"] = 5         # but only 2 decisions exist
    viol = ssot.reconcile(snap)
    assert any("by_band[Past-LDoS]" in v for v in viol), viol
    assert any("n_decisions" in v for v in viol), viol


def test_reconcile_is_coverage_honest_on_unpublished_blocks():
    """A minimal / partially-assembled snapshot (no scale/posture/lifecycle published -- like the
    golden) must produce ZERO violations, not a crash and not invented drift."""
    assert ssot.reconcile({}) == []
    assert ssot.reconcile({"executive_brief": {"scale": None}, "health_scores": []}) == []
    # Published scale but no raw basis present -> still skipped (both sides required).
    assert ssot.reconcile({"executive_brief": {"scale": {"n_devices": 10}}}) == []


def test_canonical_facts_reads_published_values():
    facts = ssot.canonical_facts(_consistent_snap())
    assert facts["n_devices"] == 6 and facts["n_collected"] == 5 and facts["n_endpoints"] == 40
    assert facts["n_critical"] == 3 and facts["worst_band"] == "Critical"
    assert facts["n_past_ldos"] == 4 and facts["n_past_eos"] == 0
    assert facts["n_design_decisions"] == 2


def test_canonical_facts_is_none_when_unpublished():
    """Coverage-honest: an absent canonical block reads as None, never a silent 0."""
    facts = ssot.canonical_facts({})
    assert facts["n_devices"] is None and facts["n_past_ldos"] is None and facts["worst_band"] is None


def test_dashboard_trend_point_reads_canonical_not_recount():
    """The DASHBOARD half of the SSOT contract: the explorer header + campaign trend metrics
    (html._trend_point) must read the canonical scale/posture/lifecycle facts, not recompute from
    the raw arrays. Discriminating fixture -- scale says 303 devices / posture says 117 Critical /
    lifecycle says 152 past-LDoS, while the raw arrays would tally 2 devices / 1 Critical and the
    conflated n_past_eos sibling is the empty 0. A regression to len(devices) / a band tally /
    n_past_eos renders 2 / 1 / 0 and trips this lock."""
    from cisco_toolkit.html import _trend_point
    snap = {
        "devices": {"a": {}, "b": {}},                            # only 2 raw devices
        "health_scores": [{"switch": "a", "band": "Critical"}],   # only 1 Critical in the raw band tally
        "executive_brief": {"scale": {"n_devices": 303}, "posture": {"n_critical": 117, "avg_health": 41}},
        "lifecycle_risk": {"summary": {"n_past_ldos": 152, "n_past_eos": 0}},
        "punchlist": [], "migration_readiness": [], "generated_at": "2026-06-13T00:00:00",
    }
    tp = _trend_point(snap)
    assert tp["n_switches"] == 303, "explorer header must read canonical scale.n_devices, not len(devices)"
    assert tp["n_critical"] == 117, "explorer header must read canonical posture.n_critical, not a band tally"
    assert tp["past_ldos"] == 152, "explorer header must read n_past_ldos (past-support), not n_past_eos (0)"


def test_audit_returns_none_on_a_clean_run():
    """The runtime self-check raises NO integrity alarm (and stamps NO field) when the published
    facts reconcile -- preserving the 'assessment_integrity not in snap' invariant a healthy run
    must hold (locked independently by tests/test_pipeline_failopen.py)."""
    assert ssot.audit(_consistent_snap()) is None
    assert ssot.audit({}) is None


def test_audit_discloses_drift_for_the_assessment_integrity_channel():
    """On real field data where the engine ever published an inconsistent fact, audit() returns a
    bounded, machine-readable disclosure to stamp into assessment_integrity (the same channel the
    deck/explorer already surface)."""
    snap = _consistent_snap()
    snap["lifecycle_risk"]["summary"]["n_past_ldos"] = 0  # the conflation, on field data
    disclosure = ssot.audit(snap)
    assert disclosure is not None
    assert disclosure["ssot_reconciliation"] == "failed"
    assert disclosure["n_violations"] >= 1
    assert isinstance(disclosure["violations"], list) and len(disclosure["violations"]) <= 20
    assert any("n_past_ldos" in v for v in disclosure["violations"])
    # The exact stamp the assembly line performs (COLLECT_PARSE main) must yield a well-formed
    # assessment_integrity the deck/explorer disclosure channel can read.
    snap.setdefault("assessment_integrity", {}).update(disclosure)
    assert snap["assessment_integrity"]["ssot_reconciliation"] == "failed"
    assert snap["assessment_integrity"]["n_violations"] >= 1


def test_summary_is_the_positive_dashboard_badge():
    """summary() is the always-present self-verification badge the dashboards render. Clean ->
    verified True with the count of published facts; drift -> verified False with the violation
    count (the explorer briefCard reads exactly these three fields)."""
    clean = ssot.summary(_consistent_snap())
    assert clean["verified"] is True
    assert clean["n_violations"] == 0
    assert clean["n_facts"] >= 10  # the consistent fixture publishes most canonical facts
    drifted = _consistent_snap()
    drifted["executive_brief"]["scale"]["n_endpoints"] = 1   # recount drift
    bad = ssot.summary(drifted)
    assert bad["verified"] is False and bad["n_violations"] >= 1


def test_summary_counts_only_published_facts():
    """Coverage-honest: n_facts counts canonical facts actually published, not a constant. An empty
    snapshot publishes none."""
    assert ssot.summary({})["n_facts"] == 0
    assert ssot.summary({})["verified"] is True   # nothing published -> nothing to contradict


# --------------------------------------------------------------------------------------------
# Cross-surface behavioural lock: render the deliverables from the AJ trap fixture and assert
# each surface reads the canonical past-support population, not the conflated sibling.
# --------------------------------------------------------------------------------------------
docx = pytest.importorskip("docx")  # the render tests below need python-docx
from docx import Document  # noqa: E402


def _trap_render_snap():
    """Union of the maintained per-deliverable fixtures (so each renderer has the keys it reads),
    with the canonical SSOT blocks overridden to the exact AJ trap values: the past-support
    population is n_past_ldos=152 while the conflated sibling n_past_eos=0. A surface that reads
    the wrong field renders 0 / omits the headline, so 152 vanishing from its output is the
    discriminator."""
    base = {}
    for mod_name in ("test_crd", "test_engagement", "test_design", "test_deck",
                     "test_runbook", "test_ops_handbook", "test_archreview", "test_mop"):
        try:
            snap_fn = getattr(importlib.import_module(mod_name), "_snap", None)
            if snap_fn:
                base.update(copy.deepcopy(snap_fn()))
        except Exception:
            pass  # a fixture that won't import is simply not contributed
    base["executive_brief"] = {
        "scale": {"n_devices": 303, "n_collected": 253, "n_endpoints": 5127, "n_vlans": 202, "n_domains": 11},
        "posture": {"avg_health": 41, "n_critical": 117, "n_poor": 85, "worst_band": "Critical"},
        "posture_statement": "Fleet is migratable with care; hardware EoL a primary driver.",
        # Deliberately NO canonical lifecycle number (152) in any passthrough string (posture_
        # statement / top_gating): a surface must surface 152 from lifecycle_risk itself, so the
        # assertion below cannot be satisfied by the fixture echoing a literal back (the adversarial
        # review found the old top_gating="152 ..." made the MOP row tautological).
        "axes": [], "top_gating": ["hardware past last-day-of-support gates the cutover window"],
    }
    base["lifecycle_risk"] = {
        "summary": {"n_devices": 303, "n_past_ldos": 152, "n_past_eos": 0, "n_near": 61, "n_active": 90,
                    "by_band": {"Past-LDoS": 152, "Near-LDoS": 61, "Active": 90}, "by_platform": {}, "asof": "2026-06-13"},
        "per_device": [], "risks": [], "asof": "2026-06-13",
    }
    return base


def _all_text(doc):
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(c.text for c in row.cells)
    return "\n".join(parts)


# Surfaces that headline the past-support population by READING lifecycle_risk.summary.n_past_ldos
# directly (verified: corrupting n_past_ldos makes 152 vanish from each). The MOP is deliberately
# EXCLUDED -- it reads no lifecycle field (only executive_brief.top_gating), so it does not headline
# this canonical fact and asserting on it would be a tautology, not an SSOT check.
_LIFECYCLE_SURFACES = [
    ("crd", "cisco_toolkit.crd", "write_crd_docx"),
    ("engagement", "cisco_toolkit.engagement", "write_engagement_docx"),
    ("runbook", "cisco_toolkit.runbook", "write_runbook_docx"),
    ("archreview", "cisco_toolkit.archreview", "write_archreview_docx"),
]


@pytest.mark.parametrize("name,module,fn_name", _LIFECYCLE_SURFACES)
def test_lifecycle_surface_headlines_ldos_not_eos(tmp_path, name, module, fn_name):
    """Render each lifecycle-bearing deliverable from the trap fixture and prove it surfaces the
    canonical past-support population (n_past_ldos=152) -- never the conflated n_past_eos (0).
    This is the unified cross-surface guard the per-surface locks lacked."""
    render = getattr(importlib.import_module(module), fn_name)
    out = str(tmp_path / f"{name}.docx")
    render(out, _trap_render_snap(), "SSOT Trap Fleet")
    text = _all_text(Document(out))
    assert "152" in text, f"{name} dropped the past-LDoS population (n_past_ldos=152) from its output"
    # No surface may render the false-health phrase that conflates the empty EoS sibling with the
    # past-support headline.
    lowered = text.lower()
    assert "0 device(s) are past last-day-of-support" not in lowered, f"{name} shows a false 0 past-LDoS"
    assert "0 devices past last-day-of-support" not in lowered, f"{name} shows a false 0 past-LDoS"


def test_scale_surfaces_read_canonical_counts(tmp_path):
    """The HLD, runbook and CRD must report the canonical fleet scale (devices / endpoints /
    VLANs) from executive_brief.scale, not a local recount of the raw arrays."""
    snap = _trap_render_snap()
    from cisco_toolkit.design import write_design_doc_docx
    from cisco_toolkit.runbook import write_runbook_docx
    out_d = str(tmp_path / "design.docx")
    out_r = str(tmp_path / "runbook.docx")
    write_design_doc_docx(out_d, copy.deepcopy(snap), "SSOT Trap Fleet")
    write_runbook_docx(out_r, copy.deepcopy(snap), "SSOT Trap Fleet")
    design_text = _all_text(Document(out_d))
    runbook_text = _all_text(Document(out_r))
    assert "303" in design_text, "HLD must surface the canonical device count (303), not a recount"
    assert "202" in design_text, "HLD must surface the canonical VLAN count (202)"
    assert "5127" in runbook_text, "runbook must surface the canonical endpoint count (5127)"
    assert "303" in runbook_text, "runbook must surface the canonical device count (303)"


def test_n_endpoints_reconciles_with_host_less_endpoint_rows():
    """SSOT producer/verifier agreement: executive_brief.scale.n_endpoints is published as
    len(endpoint_identity) -- the canonical evidenced-endpoint total ssot.reconcile checks -- NOT the
    per-domain endpoint_count sum, which silently drops any endpoint_identity row whose host is not an
    all_interfaces key. Those two formulas diverged on an off-pipeline snapshot (a host-less / foreign-host
    row), so reconcile/audit FALSE-FIRED a spurious 'ssot_reconciliation: failed' integrity alarm. Here two of
    four endpoint rows are off-host: the brief must publish 4 (not the domain-sum 2) and reconcile clean."""
    from cisco_toolkit.analyze import compute_application_intelligence, compute_executive_brief
    from cisco_toolkit.model import InterfaceData
    d = InterfaceData(port="Gi1/0/1", switchport_mode="Access", end_host_mac="00:11:22:33:44:55", vlan="10")
    aif = {"sw1": {"Gi1/0/1": d}}
    ei = [
        {"host": "sw1", "port": "Gi1/0/1", "vlan": "10", "mac": "00:11:22:33:44:55", "endpoint_class": "Server", "vendor": "Dell"},
        {"host": "sw1", "port": "Gi1/0/2", "vlan": "10", "mac": "00:11:22:33:44:66", "endpoint_class": "Server", "vendor": "Dell"},
        {"host": "GHOST", "port": "Gi9/9", "vlan": "10", "mac": "aa:bb:cc:dd:ee:ff", "endpoint_class": "Server", "vendor": "HP"},
        {"host": "", "port": "Gi9/8", "vlan": "10", "mac": "aa:bb:cc:dd:ee:00", "endpoint_class": "Server", "vendor": "HP"},
    ]
    app = compute_application_intelligence(aif, endpoint_identity=ei)
    brief = compute_executive_brief(health_scores=[{"switch": "sw1", "band": "Good", "score": 90}],
                                    application_intelligence=app, endpoint_identity=ei)
    assert brief["scale"]["n_endpoints"] == 4              # len(endpoint_identity), NOT the domain-sum (2)
    snap = {"executive_brief": brief, "endpoint_identity": ei,
            "health_scores": [{"switch": "sw1", "band": "Good", "score": 90}]}
    assert ssot.reconcile(snap) == []                      # producer == verifier -> no spurious violation
    assert ssot.audit(snap) is None
