"""The Deliverable Excellence ratchet (wave DE-01).

Every narrative deliverable in the family must carry the same 'excellence furniture': a Document
Control block (Law 2), an answer-first 'At a Glance' register with the single-source-of-truth trust
signal (Laws 6/1/10), a Glossary (Technique T2), and a Document Acceptance gate (Law 2). This test
renders each writer from a representative snapshot and asserts those sections are present — turning
"complete + answer-first + self-verified" into a BUILD CONDITION (the same shape as
test_ssot_reconciliation.py's cross-surface render lock and test_docx_family_xml_safe.py's fan-out).

A missing section fails the build; the required set only grows (the ratchet). Coverage-honest: the
'At a Glance' block reads canonical facts via cisco_toolkit.ssot — an unpublished fact renders
'[NOT OBSERVED]', never a fabricated 0.
"""
import importlib

import pytest

pytest.importorskip("docx", reason="python-docx not installed (optional deliverable dependency)")
from docx import Document  # noqa: E402

# Same writer roster as tests/test_docx_family_xml_safe.py — the whole narrative family.
WRITERS = [
    ("mop", "write_mop_docx"), ("crd", "write_crd_docx"), ("engagement", "write_engagement_docx"),
    ("archreview", "write_archreview_docx"), ("ops", "write_ops_handbook_docx"),
    ("design", "write_design_doc_docx"), ("runbook", "write_runbook_docx"),
]

# The furniture every deliverable must carry. Stable anchors: three unnumbered headings + the
# acceptance-gate invariant sentence (the MOP appends the gate headingless under its final section,
# so we anchor on add_acceptance's body text, present in every writer, not the heading).
REQUIRED_SECTIONS = ("Document Control", "At a Glance", "Glossary", "accepted as the working record")


def _snap():
    """A representative, render-complete snapshot (enough for every writer to produce a document).
    Canonical scale/posture/lifecycle are published so the 'At a Glance' register has real facts."""
    return {
        "devices": {"core1": {"model": "N9K-C93180", "hostname": "core1", "serial_number": "SAL1"},
                    "core2": {"model": "N9K-C93180", "hostname": "core2", "serial_number": "SAL2"}},
        "executive_brief": {
            "scale": {"n_devices": 2, "n_collected": 2, "n_endpoints": 40, "n_vlans": 6, "n_domains": 1},
            "posture": {"avg_health": 62, "n_critical": 1, "n_poor": 1, "worst_band": "Critical"},
            "axes": [{"axis": "Lifecycle", "severity": "High", "headline": "1 device past LDoS"}],
            "top_gating": ["Resolve the core SPOF before cutover"],
            "posture_statement": "Fleet is migratable with two gating items.",
        },
        "lifecycle_risk": {"summary": {"n_past_ldos": 1, "n_past_eos": 0, "n_near": 1, "n_active": 0}},
        "design_blueprint": {"summary": {"n_decisions": 3}, "decisions": [
            {"id": "D1", "title": "Collapse core", "status": "proposed"}]},
        "failure_impact": [{"host": "core1", "stranded": 9, "severity": "High", "vlans_impacted": 2,
                            "detail": "VLAN 10"}],
        "migration_readiness": [{"group": "G1", "switches": ["core1", "core2"], "readiness": "READY",
                                 "endpoints": 40, "checks": [], "n_fail": 0, "n_warn": 0}],
        "punchlist": [{"severity": "High", "category": "Redundancy", "title": "No FHRP on core",
                       "devices": ["core1"]}],
        "health_scores": [{"switch": "core1", "band": "Critical", "score": 20},
                          {"switch": "core2", "band": "Poor", "score": 45}],
        "migration_scenarios": {"per_group": [{"group": "G1", "recommended_scenario": "phased"}],
                                "fleet_recommendation": "phased"},
        "architecture_review": {"checks": [{"domain": "Availability", "title": "FHRP", "verdict": "deviation",
                                            "evidence": "no fhrp"}], "summary": {"grade": "C"}},
        "collection_completeness": {"summary": {"inventory": 2, "complete": 2, "partial": 0, "not_collected": 0}},
        "interfaces": {"core1": {"Gi1/0/1": {"switchport_mode": "Access", "vlan": "10",
                                             "end_host_mac": "aaaa.0000.0001", "cdp_neighbor": "core2"}}},
        "script_version": "9.9.9-test", "generated_at": "2026-07-04T00:00:00",
    }


def _all_text(doc) -> str:
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


@pytest.mark.parametrize("mod,fn", WRITERS)
def test_every_deliverable_carries_the_excellence_sections(mod, fn, tmp_path):
    writer = getattr(importlib.import_module("cisco_toolkit." + mod), fn)
    out = str(tmp_path / f"{mod}.docx")
    writer(out, _snap(), "[HISTORY-REDACTED]")
    txt = _all_text(Document(out))
    missing = [s for s in REQUIRED_SECTIONS if s not in txt]
    assert not missing, f"{mod} deliverable is missing excellence section(s): {missing}"


@pytest.mark.parametrize("mod,fn", WRITERS)
def test_every_deliverable_shows_the_single_source_of_truth_signal(mod, fn, tmp_path):
    """The 'At a Glance' block must render the ssot self-verification line (Law 10) — proof the doc
    reads canonical facts rather than recounting."""
    writer = getattr(importlib.import_module("cisco_toolkit." + mod), fn)
    out = str(tmp_path / f"{mod}.docx")
    writer(out, _snap(), "[HISTORY-REDACTED]")
    txt = _all_text(Document(out))
    assert "Single source of truth" in txt, f"{mod} deliverable does not show the SSOT trust signal"
