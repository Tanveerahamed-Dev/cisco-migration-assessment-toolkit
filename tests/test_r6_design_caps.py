"""R6 whole-repo sweep — display caps in the design / CRD / engagement writers must DISCLOSE what
they dropped.

The house rule is stated in ``excel.py :: _xls_cell_value`` and was applied in round 4 to ``mop.py``
§x.2 and ``html.py``'s Findings-Delta column: a ``[:N]`` display cap emits a trailing ``(+N more)``
marker or an "…and N further …" sentence, so what is SHOWN plus what is HIDDEN reconciles to the
total the reader is given elsewhere in the same document. These three writers are read by people who
ACT on the list: ``design.py`` writes the HLD/LLD a customer approves and an engineer builds from,
``crd.py`` the requirements the workshop takes a position on, ``engagement.py`` the verdict and its
RAID log.

Grounding: every fixture starts from ``tests/golden/snapshot.json`` — a REAL engine producer artifact
— deep-copied, never mutated on disk. Tests whose cap the golden fleet already exceeds (the
punch-list caps) run it UNMODIFIED; the rest state their augmentation explicitly. Each test fails on
the pre-fix writer.
"""
import copy
import json
import os
import re

import pytest

pytest.importorskip("docx")                      # optional dependency, like the generators themselves
from docx import Document                        # noqa: E402

from cisco_toolkit.crd import write_crd_docx                    # noqa: E402
from cisco_toolkit.design import write_design_doc_docx          # noqa: E402
from cisco_toolkit.engagement import write_engagement_docx      # noqa: E402

_GOLDEN = os.path.join(os.path.dirname(os.path.abspath(__file__)), "golden", "snapshot.json")


def _golden():
    """A deep copy of the real golden producer snapshot (the file itself is never written)."""
    with open(_GOLDEN, encoding="utf-8") as fh:
        return copy.deepcopy(json.load(fh))


def _blocks(path):
    """Every rendered text block: paragraphs plus one 'a | b | c' line per table row."""
    doc = Document(path)
    out = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            out.append(" | ".join(c.text for c in row.cells))
    return out


def _render(writer, snap, tmp_path, name):
    p = str(tmp_path / f"{name}.docx")
    writer(p, snap, "R6 cap sweep")
    return _blocks(p)


def _find(blocks, needle):
    return [b for b in blocks if needle in b]


# --------------------------------------------------------------------------------------------------
# the shared marker
# --------------------------------------------------------------------------------------------------
def test_capped_join_marker_reconciles_shown_and_hidden():
    # imported inside the test so REVERTING design.py leaves the other tests collectable (they then
    # fail on the missing disclosure, which is the point of this file)
    from cisco_toolkit.design import _capped_join
    assert _capped_join(["a", "b"], 5) == "a, b"                     # under the cap: no noise
    assert _capped_join(list("abcde"), 3) == "a, b, c (+2 more)"     # 3 shown + 2 hidden == 5
    assert _capped_join(list("abcde"), 3, "; ").startswith("a; b; c")


# --------------------------------------------------------------------------------------------------
# design.py §2.1 — the Count column states the tier size; the Devices cell must not contradict it
# --------------------------------------------------------------------------------------------------
def test_tier_table_device_cell_reconciles_with_its_own_count_column(tmp_path):
    snap = _golden()
    for i in range(45):                                   # +45 pure-L2 access switches (no SVI, no adjacency)
        snap["devices"][f"acc{i:02d}"] = {"hostname": f"acc{i:02d}", "model": "C9200-24T",
                                          "sw_version": "17.9", "platform": "ios"}
    row = _find(_render(write_design_doc_docx, snap, tmp_path, "tier"), "Access (L2-only)")
    assert row, "the §2.1 tier table did not render"
    _tier, count, devices = [c.strip() for c in row[0].split(" | ")]
    m = re.search(r"\(\+(\d+) more\)", devices)
    assert m, f"the Devices cell was cut with no disclosure while the Count column says {count}: {devices[:120]}"
    shown = len([n for n in devices.split(" (+")[0].split(", ") if n])
    assert shown + int(m.group(1)) == int(count)          # 30 shown + 16 hidden == 46


def test_keystone_sentence_names_the_population_section_2_4_counts(tmp_path):
    snap = _golden()
    snap["failure_impact"] = [{"host": f"sw{i:02d}", "severity": "High", "stranded": 100 - i}
                              for i in range(12)]
    blocks = _render(write_design_doc_docx, snap, tmp_path, "keystone")
    sentence = _find(blocks, "Concentrated dependency:")
    assert sentence, "§2.1 concentrated-dependency sentence missing"
    assert "the 5 largest of 12 device(s)" in sentence[0], sentence[0]
    # …and the §2.4 resilience row it points at states the same 12.
    assert _find(blocks, "Keystone devices (strand endpoints if lost) | 12")


# --------------------------------------------------------------------------------------------------
# design.py §2.5 / §2.7 — tables whose FULL count the surrounding prose already states
# --------------------------------------------------------------------------------------------------
def test_multicast_group_table_discloses_the_groups_it_dropped(tmp_path):
    snap = _golden()
    snap.setdefault("service_map", {}).setdefault("multicast", {})["classified_groups"] = [
        {"group": f"239.1.1.{i}", "name": f"feed{i}", "category": "AV"} for i in range(20)]
    blocks = _render(write_design_doc_docx, snap, tmp_path, "mcast")
    assert _find(blocks, "20 group(s) were classified"), "the headline count is the thing being contradicted"
    disc = _find(blocks, "further classified group(s)")
    assert disc, "15 of 20 multicast groups rendered with no disclosure"
    assert "…and 5 further" in disc[0] and "(20 in total" in disc[0], disc[0]


def test_qos_departure_table_discloses_the_findings_it_dropped(tmp_path):
    snap = _golden()
    snap["qos_audit"] = {
        "summary": {"n_devices": 20, "n_assessable": 20, "modes": {"none": 20}, "n_voice_ports": 4},
        "findings": [{"host": f"sw{i:02d}", "label": "no trust boundary at the access edge",
                      "severity": "Medium"} for i in range(15)],
    }
    disc = _find(_render(write_design_doc_docx, snap, tmp_path, "qos"), "further QoS departure(s)")
    assert disc, ("§2.7 promises 'the table below lists where the observed configuration departs' — "
                  "12 of 15 rendered silently makes absence read as conformance")
    assert "…and 3 further" in disc[0] and "(15 in total)" in disc[0], disc[0]


# --------------------------------------------------------------------------------------------------
# design.py §3.3 — the per-device build sheet an engineer reproduces the fabric from
# --------------------------------------------------------------------------------------------------
def test_build_detail_svi_and_uplink_lists_disclose_their_cut(tmp_path):
    snap = _golden()
    ports = {f"Vlan{100 + i}": {"svi_ip": f"10.1.{i}.1"} for i in range(15)}      # 15 SVIs vs the 12 cap
    ports.update({f"Te1/{i}": {"switchport_mode": "trunk", "cdp_neighbor": f"nbr{i:02d}"}
                  for i in range(20)})                                            # 20 uplinks vs the 14 cap
    snap["devices"]["aaa-bigcore"] = {"hostname": "aaa-bigcore", "model": "N9K-C9508",
                                      "sw_version": "10.2", "platform": "nxos"}
    snap["interfaces"]["aaa-bigcore"] = ports
    snap["l3_forwarding"].append({"switch": "aaa-bigcore", "vlan": "100", "svi_ip": "10.1.0.1",
                                  "fhrp": "none", "risk": ""})
    blocks = _render(write_design_doc_docx, snap, tmp_path, "builddetail")
    svi = [b for b in blocks if b.startswith("Gateway SVIs:") and "Vlan100" in b]
    upl = [b for b in blocks if b.startswith("Trunk uplinks:") and "nbr00" in b]
    assert svi and upl, "the §3.3 build-detail block for the augmented device did not render"
    assert "(+3 more)" in svi[0], f"3 gateway SVIs dropped from a build sheet with no mark: {svi[0][:160]}"
    assert "(+6 more)" in upl[0], f"6 trunk uplinks dropped from a build sheet with no mark: {upl[0][:160]}"


# --------------------------------------------------------------------------------------------------
# design.py §3.5 — the software / replacement recommendations
# --------------------------------------------------------------------------------------------------
def test_software_recommendations_disclose_the_models_they_dropped(tmp_path):
    snap = _golden()
    snap["devices"] = {}
    per_device = []
    for i in range(11):                       # 11 models running mixed images; 10 of them past-EoS
        for n, ver in enumerate(("1.0", "2.0")):
            host = f"m{i:02d}s{n}"
            snap["devices"][host] = {"hostname": host, "model": f"MDL-{i:02d}",
                                     "sw_version": ver, "platform": "ios"}
            per_device.append({"host": host, "model": f"MDL-{i:02d}",
                               "band": "Past-EoS" if i < 10 else "Active"})
    snap["lifecycle_risk"] = {"per_device": per_device, "summary": {"n_past_eos": 20, "n_past_ldos": 0}}
    blocks = _render(write_design_doc_docx, snap, tmp_path, "swplan")
    mixed = _find(blocks, "further model(s) running MIXED images")
    assert mixed, "3 of 11 mixed-image models got no standardization recommendation, undisclosed"
    assert "…and 3 further" in mixed[0] and "(11 in total)" in mixed[0], mixed[0]
    eol = _find(blocks, "Hardware past end-of-support gets replacement")
    assert eol and "(+2 more)" in eol[0] and "10 model(s) in total" in eol[0], eol[0] if eol else "missing"


# --------------------------------------------------------------------------------------------------
# design.py §4 — the target-state decisions, both branches
# --------------------------------------------------------------------------------------------------
def test_recommended_decision_detail_discloses_the_decisions_it_did_not_detail(tmp_path):
    snap = _golden()
    snap["design_blueprint"] = {"decisions": [
        {"status": "recommended", "priority": "P2", "domain": "L3", "title": f"Decision {i}",
         "driver": "evidence", "recommended_action": f"pattern {i}", "alternatives": "alt",
         "tradeoffs": "trade", "evidence": {"summary": "observed"}, "principle": {"citation": "CCDE"}}
        for i in range(14)]}
    blocks = _render(write_design_doc_docx, snap, tmp_path, "decisions")
    assert _find(blocks, "Decision 13"), "the §4.2 table should list every recommended decision"
    disc = _find(blocks, "further recommended decision(s)")
    assert disc, ("the detail blocks carry the recommended pattern / alternatives / trade-offs that "
                  "appear nowhere else; 10 of 14 rendered silently")
    assert "…and 4 further" in disc[0] and "(14 in total)" in disc[0], disc[0]


def test_punchlist_fallback_discloses_on_the_real_golden_snapshot(tmp_path):
    """UNMODIFIED real producer artifact: 37 punch-list items, no design blueprint — so §4 IS the
    punch-list table, and the 12-row cap is hit by the golden fleet itself."""
    snap = _golden()
    assert not snap.get("design_blueprint") and len(snap["punchlist"]) == 37   # pin the preconditions
    disc = _find(_render(write_design_doc_docx, snap, tmp_path, "fallback"), "further punch-list item(s)")
    assert disc, "§4 rendered 12 of 37 punch-list items as the whole target-state section, undisclosed"
    assert "…and 25 further" in disc[0] and "(37 in total)" in disc[0], disc[0]


# --------------------------------------------------------------------------------------------------
# design.py §6.2 — the endpoint interoperability surface the access-edge design must cover
# --------------------------------------------------------------------------------------------------
def test_endpoint_class_table_discloses_the_classes_it_dropped(tmp_path):
    snap = _golden()
    snap["endpoint_identity"] = [
        {"host": "access1", "port": f"Gi0/{i}", "mac": f"aabb.ccdd.{i:04x}",
         "vendor": "Cisco", "endpoint_class": f"Class {i % 11:02d}"} for i in range(44)]
    blocks = _render(write_design_doc_docx, snap, tmp_path, "classes")
    assert _find(blocks, "Top device class | Endpoints"), \
        "the class table must say 'Top …' like its sibling vendor table, not read as the full set"
    disc = _find(blocks, "further identified device class(es)")
    assert disc, "the prose states 11 identified classes and the table showed 8 with no mark"
    assert "…and 3 further" in disc[0] and "(11 in total" in disc[0], disc[0]


# --------------------------------------------------------------------------------------------------
# crd.py §2 — the known issues the requirements must take a position on
# --------------------------------------------------------------------------------------------------
def test_crd_known_issues_disclose_the_punchlist_on_the_real_golden_snapshot(tmp_path):
    snap = _golden()
    disc = _find(_render(write_crd_docx, snap, tmp_path, "crd"), "further punch-list item(s)")
    assert disc, ("the CRD's function is 'take a position on' these issues; 8 of 37 rendered with no "
                  "total anywhere in §2 reads as the complete known-issue set")
    assert "…and 29 further" in disc[0] and "(37 in total)" in disc[0], disc[0]


# --------------------------------------------------------------------------------------------------
# engagement.py — the verdict section, the RAID log, the issue log
# --------------------------------------------------------------------------------------------------
def test_engagement_gating_headlines_disclose_the_axes_they_dropped(tmp_path):
    snap = _golden()
    snap["executive_brief"] = {"posture_statement": "test posture",
                               "top_gating": [f"Axis {i} is Critical" for i in range(9)]}
    disc = _find(_render(write_engagement_docx, snap, tmp_path, "gating"), "further gating headline(s)")
    assert disc, "6 of 9 Critical/High gating axes rendered in the VERDICT section, undisclosed"
    assert "…and 3 further" in disc[0] and "(9 Critical/High axes in total)" in disc[0], disc[0]


def test_engagement_risk_log_reconciles_with_the_section_1_conditions(tmp_path):
    """UNMODIFIED real producer artifact: golden carries 3 Critical + 16 High punch-list findings.
    §1 states both counts as verdict conditions, so a §5.1 RAID log that seeds 8 rows and says
    nothing is the same document contradicting itself."""
    blocks = _render(write_engagement_docx, _golden(), tmp_path, "raid")
    assert _find(blocks, "3 Critical punch-list item(s)") and _find(blocks, "16 High punch-list item(s)")
    disc = _find(blocks, "further Critical/High punch-list finding(s)")
    assert disc, "§5.1 seeded 8 risks against §1's 3 Critical + 16 High, with no marker"
    assert "…and 11 further" in disc[0] and "(19 in total" in disc[0], disc[0]     # 8 seeded + 11 == 19


def test_engagement_issue_log_discloses_blind_spots_beyond_the_cap(tmp_path):
    snap = _golden()
    snap["collection_completeness"] = {
        "summary": {"complete": 0, "partial": 12, "not_collected": 0},
        "devices": [{"host": f"sw{i:02d}", "status": "partial", "data_quality": 40,
                     "missing": ["show cdp neighbors detail"]} for i in range(12)]}
    blocks = _render(write_engagement_docx, snap, tmp_path, "issues")
    assert _find(blocks, "12 device(s) with blind spots"), "§2's phase tracker states the full count"
    disc = _find(blocks, "further device(s) with collection blind spots")
    assert disc, "§5.3 listed 8 of 12 blind-spot devices while §2 stated 12"
    assert "…and 4 further" in disc[0] and "(12 in total" in disc[0], disc[0]


# --------------------------------------------------------------------------------------------------
# the disclosures must be silent when nothing was dropped (a marker on an uncut list is its own defect)
# --------------------------------------------------------------------------------------------------
def test_no_disclosure_fires_when_the_caps_are_not_reached(tmp_path):
    """The golden fleet is under every cap except the punch-list ones — nothing else may announce a
    truncation that did not happen."""
    snap = _golden()
    blocks = _render(write_design_doc_docx, snap, tmp_path, "quiet")
    for phrase in ("further classified group(s)", "further QoS departure(s)",
                   "further identified device class(es)", "further model(s) running MIXED images",
                   "further recommended decision(s)", "(+"):
        assert not _find(blocks, phrase), f"{phrase!r} fired on a fleet that is under the cap"
    eng = _render(write_engagement_docx, snap, tmp_path, "quiet_eng")
    assert not _find(eng, "further device(s) with collection blind spots")   # golden has no blind spots
    assert not _find(eng, "further gating headline(s)")                      # golden has no executive_brief
