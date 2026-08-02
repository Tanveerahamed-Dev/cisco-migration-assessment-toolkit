"""The golden-snapshot eval tier (Phase 0 of the autonomous-brain plan).

Two guarantees, tiered:

* **Smoke** (always — the ``verify-green`` gate runs it on every change): a fixed *known-good*
  deliverable scores a true 100 (every grounded Law verified, nothing tripped, nothing unverified),
  AND three deliberately-broken mutants each trip exactly their target Law. The mutants are the
  load-bearing half — they prove the harness is **non-vacuous** (a guard that cannot fail is red).
* **Full** (also always, since 2026-07-28): render the whole DOCX family from the known-good
  snapshot and re-score with the rendered text, exercising the render checks (excellence furniture +
  the rendered self-verification badge byte-matching the machine fact count); plus, opportunistically,
  score the real on-disk assessment snapshot if one is present. This tier was gated behind
  ``EVAL_FULL=1``, which no workflow, hook or script ever set — so the only tests covering the
  rendered furniture had never run. Un-gated for 2.1s of runtime.

Coverage-honesty is asserted as a first-class property: a partial snapshot's unverifiable Laws come
back ``unverified`` (disclosed), never a fabricated pass; an empty deliverable scores ``None``, never
a false 100.
"""
import glob
import importlib

import pytest

from cisco_toolkit import eval_harness as E
from cisco_toolkit import ssot

# The "full tier" used to be gated behind EVAL_FULL=1 and described as the release gate -- but that
# variable is set NOWHERE in this repository: not in .github/workflows, not in .claude/hooks, not in
# pytest.ini, not in any release script. So the two tests below had never run in any lane, and they
# are the only ones that render the whole DOCX family from the known-good snapshot and assert the
# rendered-tier checks: the [NOT OBSERVED] furniture (guardrail 3), the SSOT badge, and the
# byte-match between the rendered citation count and ssot.summary(snap)["n_facts"]. Any writer could
# have stopped emitting those and every lane would have stayed green.
#
# Measured before removing the gate: the two tests add 2.1s to a 0.6s file. A release gate nobody
# runs is not a gate; 2 seconds is not a reason to keep one.
def _full(fn):
    """Identity passthrough, kept only so the full-tier tests stay greppable as a group."""
    return fn

# The same writer roster the excellence ratchet renders (tests/test_deliverable_excellence.py).
WRITERS = [
    ("mop", "write_mop_docx"), ("crd", "write_crd_docx"), ("engagement", "write_engagement_docx"),
    ("archreview", "write_archreview_docx"), ("ops", "write_ops_handbook_docx"),
    ("design", "write_design_doc_docx"), ("runbook", "write_runbook_docx"),
]


def _ports(vlans):
    """Access ports carrying the given VLANs — the raw basis analyze.vlan_inventory counts, so the
    published n_vlans reconciles to len(these distinct VLANs)."""
    return {f"Gi1/0/{i + 1}": {"switchport_mode": "Access", "vlan": str(v),
                               "end_host_mac": f"aaaa.0000.00{v:02d}", "cdp_neighbor": "sw-dist-1"}
            for i, v in enumerate(vlans)}


def known_good_snapshot():
    """A fixed, fully-published, **reconcile-clean** known-good deliverable substrate. Every one of
    the 14 canonical facts is published and each reconciles to an independent raw basis (health
    bands, collection completeness, endpoint/VLAN/domain evidence, lifecycle per-device, decisions),
    so ssot.summary(snap) == {verified True, n_facts 14, n_violations 0}. Rich enough that all seven
    DOCX writers render. This is the deliverable the eval scores 100."""
    return {
        "script_version": "9.9.9-eval", "generated_at": "2026-07-06T00:00:00",
        "attestation": {"schema": "attestation/1", "claims": [{"id": "c1", "claim": "read-only collection"}]},
        "devices": {h: {"model": "C9300", "hostname": h, "serial_number": h[-1]}
                    for h in ("sw-core-1", "sw-dist-1", "sw-acc-1")},
        "health_scores": [{"switch": "sw-core-1", "band": "Critical", "score": 20},
                          {"switch": "sw-dist-1", "band": "Poor", "score": 50},
                          {"switch": "sw-acc-1", "band": "Good", "score": 80}],
        "collection_completeness": {"summary": {"inventory": 3, "complete": 2, "partial": 0, "not_collected": 1},
                                    "devices": [{"host": "sw-acc-1", "status": "not collected"}]},
        "endpoint_identity": [{"mac": f"aaaa.0000.100{i}"} for i in range(1, 6)],
        "application_intelligence": {"domains": [{"id": "d1", "name": "prod"}]},
        "executive_brief": {
            "scale": {"n_devices": 3, "n_collected": 2, "n_endpoints": 5, "n_vlans": 6, "n_domains": 1},
            "posture": {"avg_health": 50, "n_critical": 1, "n_poor": 1, "worst_band": "Critical"},
            "axes": [{"axis": "Lifecycle", "severity": "High", "headline": "1 device past LDoS"}],
            "top_gating": ["Resolve the core SPOF before cutover"],
            "posture_statement": "Fleet is migratable with two gating items."},
        "lifecycle_risk": {"summary": {"n_devices": 3, "n_past_ldos": 1, "n_past_eos": 0, "n_near": 1, "n_active": 1},
                           "per_device": [{"host": "sw-core-1", "band": "Past-LDoS"},
                                          {"host": "sw-dist-1", "band": "Near-LDoS"},
                                          {"host": "sw-acc-1", "band": "Active"}]},
        "design_blueprint": {"summary": {"n_decisions": 2},
                             "decisions": [{"id": "D1", "title": "Collapse core", "status": "proposed"},
                                           {"id": "D2", "title": "Add FHRP", "status": "proposed"}]},
        "failure_impact": [{"host": "sw-core-1", "stranded": 9, "severity": "High",
                            "vlans_impacted": 2, "detail": "VLAN 10"}],
        "migration_readiness": [{"group": "G1", "switches": ["sw-core-1", "sw-dist-1"], "readiness": "READY",
                                 "endpoints": 5, "checks": [], "n_fail": 0, "n_warn": 0}],
        "punchlist": [{"severity": "High", "category": "Redundancy", "title": "No FHRP on core",
                       "devices": ["sw-core-1"]}],
        "migration_scenarios": {"per_group": [{"group": "G1", "recommended_scenario": "phased"}],
                                "fleet_recommendation": "phased"},
        "architecture_review": {"checks": [{"domain": "Availability", "title": "FHRP", "verdict": "deviation",
                                            "evidence": "no fhrp"}], "summary": {"grade": "C"}},
        "interfaces": {"sw-core-1": _ports([10, 20, 30]), "sw-dist-1": _ports([40, 50, 60])},
    }


# --- smoke tier: the known-good deliverable scores a real 100 ----------------------------------

def test_known_good_reconciles_clean():
    """Precondition on the fixture itself: it must be a genuine SSOT-clean deliverable."""
    snap = known_good_snapshot()
    assert ssot.reconcile(snap) == []
    # n_checked is the count of facts actually RECONCILED against raw evidence, and is asserted
    # non-zero on purpose: n_facts alone counts what was PUBLISHED, so a snapshot stripped of its raw
    # arrays would publish 14 and reconcile 0 while still reporting an empty violation list.
    assert ssot.summary(snap) == {"verified": True, "n_facts": 14, "n_checked": 16,
                                  "n_violations": 0}


def test_known_good_scores_100_smoke():
    r = E.score(known_good_snapshot())            # snapshot-only (no rendered text)
    assert r.score == 100, [c.detail for c in r.checks if c.status != "pass"]
    assert r.laws_tripped == []
    assert r.counterexamples == 0
    # Every Law the harness claims to ground must be VERIFIED (not merely unverified) on the smoke
    # tier — the snapshot signals alone cover all seven; render checks stay unverified but never
    # leave a Law wholly unverified because each Law also has a snapshot signal.
    assert set(r.laws_verified) == set(E.GROUNDED_LAWS)
    assert r.laws_unverified == []


MUTANTS = [
    # (label, expected tripped law, mutation) — each proves one grounded check is non-vacuous.
    ("L1_reconcile", "L1", lambda s: s["executive_brief"]["scale"].__setitem__("n_devices", 999)),
    ("L2_provenance", "L2", lambda s: (s.pop("script_version"), s.pop("attestation"))),
    ("L3_coverage", "L3", lambda s: s["collection_completeness"]["summary"].__setitem__("not_collected", 0)),
]


@pytest.mark.parametrize("label,law,mutate", MUTANTS, ids=[m[0] for m in MUTANTS])
def test_harness_detects_injected_defect(label, law, mutate):
    """Non-vacuity: a deliverable broken on exactly one Law must drop below 100 and trip THAT Law.
    A harness that cannot fail here is a rubber stamp."""
    snap = known_good_snapshot()
    mutate(snap)
    r = E.score(snap)
    assert law in r.laws_tripped, f"{label}: expected {law} tripped, got {r.laws_tripped}"
    assert r.counterexamples >= 1
    assert r.score is not None and r.score < 100


def test_grounded_law_registry_matches_what_the_scorer_verifies():
    """The disclosed coverage (GROUNDED_LAWS) must equal the Laws the scorer actually exercises on a
    complete deliverable — so the 'which Laws are machine-checkable' claim can't rot."""
    r = E.score(known_good_snapshot())
    exercised = {c.law for c in r.checks}
    assert exercised == set(E.GROUNDED_LAWS)


def test_coverage_honest_partial_snapshot_reports_unverified_not_pass():
    """A snapshot with every canonical-fact block stripped (no published facts — the real shape of
    tests/golden/snapshot.json) must leave L1/L6/L10 UNVERIFIED: disclosed as unassessable, never
    silently passed and never counted as a defect. The coverage/provenance Laws it CAN prove stay
    verified."""
    snap = known_good_snapshot()
    for block in ("executive_brief", "lifecycle_risk", "design_blueprint"):
        del snap[block]
    r = E.score(snap)
    assert ssot.summary(snap)["n_facts"] == 0  # precondition: genuinely no published facts
    for law in ("L1", "L6", "L10"):
        assert law in r.laws_unverified, (law, r.laws_unverified)
        assert law not in r.laws_tripped
        assert law not in r.laws_verified


def test_empty_deliverable_scores_none_not_a_false_100():
    """Absence is reported as absence: an empty snapshot has nothing applicable to score, so the
    score is None (not 100, not 0) and every grounded Law is unverified."""
    r = E.score({})
    assert r.score is None
    assert r.laws_tripped == []
    assert r.n_applicable == 0
    assert set(r.laws_unverified) == set(E.GROUNDED_LAWS)


def test_scorecard_row_matches_schema():
    """to_scorecard_row emits exactly the docs/quality/README.md schema keys (so the appender and
    the eval suite write the same shape) — reconciled to scorecard.SCHEMA_KEYS, the owner, never a
    hardcoded copy (SSOT Law 1). A deterministic scored row carries judge_tnr null (no judge to
    measure) and provisional false — it is bias-free, never an advisory judge verdict (P0-6)."""
    from cisco_toolkit.scorecard import SCHEMA_KEYS
    r = E.score(known_good_snapshot())
    row = r.to_scorecard_row(deliverable="golden", commit="abc1234", date="2026-07-06")
    assert set(row) == set(SCHEMA_KEYS)
    assert row["verdict"] == "APPROVE" and row["score"] == 100 and row["counterexamples"] == 0
    assert row["judge_tnr"] is None and row["provisional"] is False
    assert row["authored_by"] is None and row["reviewed_by"] is None  # no provenance pair on a deterministic row
    # a tripped deliverable rows as BLOCK with the defect in notes
    bad = known_good_snapshot(); bad["executive_brief"]["scale"]["n_devices"] = 999
    brow = E.score(bad).to_scorecard_row(deliverable="golden", commit="abc1234", date="2026-07-06")
    assert brow["verdict"] == "BLOCK" and "L1" in brow["laws_tripped"] and brow["notes"].startswith("eval:")
    # the coverage-honest edge: an EMPTY deliverable's APPROVE (nothing applicable, score None)
    # quantifies nothing -> it self-marks provisional rather than reading as gate-worthy
    erow = E.score({}).to_scorecard_row(deliverable="golden", commit="abc1234", date="2026-07-06")
    assert erow["score"] is None and erow["verdict"] == "APPROVE" and erow["provisional"] is True


# --- full tier (release): render the family + rendered checks; real snapshot if present ---------

@_full
def test_known_good_renders_and_scores_100_full(tmp_path):
    """Render the whole DOCX family from the known-good snapshot, then re-score with the rendered
    text. This exercises the render checks (Document Control / At a Glance / Inputs Required /
    [NOT OBSERVED] / SSOT badge) and the rendered citation byte-match (badge count == n_facts)."""
    pytest.importorskip("docx", reason="python-docx not installed (optional deliverable dependency)")
    from docx import Document
    snap = known_good_snapshot()
    parts = []
    for mod, fn in WRITERS:
        writer = getattr(importlib.import_module("cisco_toolkit." + mod), fn)
        out = str(tmp_path / f"{mod}.docx")
        writer(out, snap, "Meridian")
        doc = Document(out)
        seg = [p.text for p in doc.paragraphs]
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    seg.append(cell.text)
        parts.append("\n".join(seg))
    texts = "\n".join(parts)
    r = E.score(snap, texts=texts)
    assert r.score == 100, [c.detail for c in r.checks if c.status != "pass"]
    assert r.laws_tripped == [] and r.laws_unverified == []
    # the rendered-tier checks must actually have run (not stayed unverified)
    rendered = [c for c in r.checks if c.check_id.endswith(".rendered")]
    assert rendered and all(c.status == "pass" for c in rendered)


@_full
def test_real_on_disk_snapshot_scores_clean_if_present():
    """Opportunistic: a real assessment snapshot in the repo root is a genuine known-good
    deliverable — it must reconcile and trip no Law. Coverage-honest: if none is present, SKIP with
    a reason (absent fixture reported as absence, never a silent pass)."""
    import json
    candidates = sorted(glob.glob("*_AUTOFILLED_*.snapshot.json") or glob.glob("*.snapshot.json"))
    if not candidates:
        pytest.skip("no on-disk assessment snapshot present to score (release runs against a real one)")
    snap = json.load(open(candidates[-1], encoding="utf-8"))
    r = E.score(snap)
    assert r.laws_tripped == [], f"{candidates[-1]} tripped {r.laws_tripped}: " \
                                 + "; ".join(c.detail for c in r.checks if c.status == "fail")
    assert r.n_applicable > 0


def test_render_checks_fail_when_furniture_is_absent():
    """The render tier normally needs a real DOCX render; a synthetic furniture-less `texts=` trips the
    .rendered checks to FAIL — exercised here without EVAL_FULL or a DOCX build (pure text)."""
    snap = known_good_snapshot()
    assert E.score(snap).score == 100                     # grounded-only (no texts): a clean 100
    bare = E.score(snap, texts="a plain paragraph with none of the required excellence furniture")
    failed = [c for c in bare.checks if c.status == "fail"]
    assert failed, "furniture-less rendered text must fail the render checks"
    assert bare.score is not None and bare.score < 100    # the furniture-less doc is penalized


def test_law3_flags_impossible_coverage_accounting():
    """Law-3: 'collected' exceeding inventory is an impossible-coverage FAIL — a coverage-honesty violation
    detector, previously untested (only the buckets-don't-sum branch was exercised)."""
    snap = known_good_snapshot()
    snap["collection_completeness"] = {"summary": {"inventory": 3, "complete": 5}}
    l3 = [c for c in E.score(snap).checks if c.check_id == "law3.accounting"]
    assert l3 and l3[0].status == "fail" and "exceeds inventory" in l3[0].detail
