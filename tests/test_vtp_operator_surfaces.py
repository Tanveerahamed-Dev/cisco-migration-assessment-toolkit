"""Operator projections for the closed serialized VTP safety receipt."""

from __future__ import annotations

import copy
import json
from pathlib import Path
import shutil
import subprocess

import pytest
from openpyxl import Workbook

from cisco_toolkit.capture_integrity import compute_capture_integrity_from_paths
from cisco_toolkit.excel import VTP_SAFETY_SHEET_NAME, write_vtp_safety_sheet
from cisco_toolkit.runbook import write_runbook_docx
from cisco_toolkit.vtp_safety import (
    compute_vtp_safety_baseline,
    embedded_vtp_safety_baseline,
    validate_vtp_safety_baseline,
)


ROOT = Path(__file__).resolve().parents[1]
EXPLORER = ROOT / "cisco_toolkit" / "blast_radius_explorer.html"
NODE = shutil.which("node")
BLOCKERS = ("notverified-000", "notverified-001", "review-000", "review-001")


def _status(*, revision: int | None, domain: str | None = "CAMPUS") -> str:
    lines = ["VTP version running             : 2"]
    if domain is not None:
        lines.append(f"VTP Domain Name                 : {domain}")
    lines.append("VTP Operating Mode              : Server")
    if revision is not None:
        lines.append(f"Configuration Revision          : {revision}")
    return "\n".join(lines) + "\n"


@pytest.fixture(scope="module")
def receipts(tmp_path_factory):
    root = tmp_path_factory.mktemp("vtp-operator")
    bodies = {
        **{f"assessed-{index:03d}": _status(revision=index) for index in range(55)},
        "review-000": _status(revision=100),
        "review-001": _status(revision=150),
        "notverified-000": _status(revision=None),
        "notverified-001": _status(revision=1, domain=None),
        "disabled-coverage": "VTP is disabled.\n",
    }
    paths = {}
    for host, body in bodies.items():
        path = root / f"{host}-show-vtp-status.txt"
        path.write_text(body, encoding="utf-8")
        paths[host] = {"show vtp status": str(path)}
    current = compute_vtp_safety_baseline(
        paths,
        compute_capture_integrity_from_paths(paths),
        {host: {"platform": "ios-xe"} for host in bodies},
    )
    assert validate_vtp_safety_baseline(current, require_current_run=True)["valid"] is True
    embedded = embedded_vtp_safety_baseline(current)
    assert validate_vtp_safety_baseline(embedded)["valid"] is True
    return {"root": root, "current": current, "embedded": embedded}


def _sheet_text(sheet) -> str:
    return "\n".join(
        str(cell.value)
        for row in sheet.iter_rows()
        for cell in row
        if cell.value is not None
    )


def _docx_text(path: Path) -> str:
    from docx import Document

    document = Document(path)
    return "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    )


def _explorer_result(baseline: object, host: str | None = None) -> dict:
    assert NODE is not None
    script = r"""
const fs=require('fs');
const html=fs.readFileSync(process.argv[1],'utf8');
const start=html.indexOf('const _FHRP_GROUP_SCOPE=');
const end=html.indexOf('/* the dynamic routing protocols',start);
if(start<0||end<0)throw new Error('protocol receipt block not found');
const input=JSON.parse(fs.readFileSync(0,'utf8'));
let SNAP={vtp_safety_baseline:input.baseline};
function esc(v){return String(v??'').replace(/[&<>"']/g,c=>({'&':'&amp;','<':'&lt;','>':'&gt;','"':'&quot;',"'":'&#39;'}[c]));}
function shortName(v){return String(v??'');}
eval(html.slice(start,end));
process.stdout.write(JSON.stringify({
  valid:vtpSafetyReceiptValid(input.baseline),
  rendered:vtpSafetySection(input.host),
}));
"""
    proc = subprocess.run(
        [NODE, "-e", script, str(EXPLORER)],
        input=json.dumps({"baseline": baseline, "host": host}),
        capture_output=True,
        check=True,
        text=True,
        encoding="utf-8",
        timeout=60,
    )
    return json.loads(proc.stdout)


def test_workbook_runbook_and_explorer_keep_all_blockers_and_first_50_assessed(receipts):
    baseline = receipts["embedded"]
    summary = baseline["summary"]
    assert summary["n_hosts"] == 60 and summary["n_subject_hosts"] == 59
    assert summary["by_coverage_status"]["not_applicable"] == 1

    workbook = Workbook()
    write_vtp_safety_sheet(workbook, baseline)
    sheet = workbook[VTP_SAFETY_SHEET_NAME]
    workbook_text = _sheet_text(sheet)
    assert sheet.max_row == 5 + len(BLOCKERS) + 50
    assert sheet.freeze_panes == "A6"

    output = receipts["root"] / "vtp-runbook.docx"
    write_runbook_docx(
        str(output),
        {
            "vtp_safety_baseline": baseline,
            "protocol_health": [{
                "switch": "review-001", "protocol": "VTP", "severity": "Info",
                "summary": "mode server; domain CAMPUS; rev 150",
                "detail": "HEALTH-DETAIL-PRESERVED",
            }],
            "protocol_intelligence": [{
                "switch": "review-001", "protocol": "VTP", "state": "HIGH-REVISION",
                "severity": "High", "meaning": "INTELLIGENCE-DETAIL-PRESERVED",
                "likely_cause": "INTELLIGENCE-DETAIL-PRESERVED", "remediation": "review before staging",
                "confidence": "High",
            }],
        },
        "fixture",
    )
    runbook_text = _docx_text(output)
    explorer = _explorer_result(baseline)
    assert explorer["valid"] is True

    for text in (workbook_text, runbook_text, explorer["rendered"]):
        for host in BLOCKERS:
            assert host in text
        assert "assessed-049" in text
        assert "assessed-050" not in text
        assert "100 or higher" in text
        assert "not proof that an overwrite will occur" in text
        assert "NOT ACCEPTANCE" in text
        assert "60 host(s)" in text and "59 subject host(s)" in text
        assert baseline["summary"]["baseline_sha256"] not in text
        assert baseline["coverage"][0]["source_sha256"] not in text
        assert "show vtp status#mode/domain/revision/version" not in text
        assert str(receipts["root"]) not in text

    assert "All 4 blocker row(s) are shown; 50 of 55 assessed" in runbook_text
    assert "All 4 blocker row(s) are rendered.</b> 50 of 55 assessed" in explorer["rendered"]
    assert "HEALTH-DETAIL-PRESERVED" in runbook_text
    assert "INTELLIGENCE-DETAIL-PRESERVED" in runbook_text


def test_serialized_current_run_claim_and_malformed_receipts_fail_closed(receipts):
    current_json = json.loads(json.dumps(receipts["current"]))
    assert current_json["projection_custody"] == "current_run_source_bound"
    hostile = {
        "schema": "vtp_safety_baseline/1",
        "projection_custody": "embedded_unverified",
        "rows": [{"switch": "HOSTILE-SWITCH", "acceptance": "HOSTILE PASS"}],
        "raw": "password HOSTILE-SECRET",
    }

    for rejected in (current_json, hostile):
        workbook = Workbook()
        write_vtp_safety_sheet(workbook, rejected)
        workbook_text = _sheet_text(workbook[VTP_SAFETY_SHEET_NAME])
        output = receipts["root"] / ("current-claim.docx" if rejected is current_json else "hostile.docx")
        write_runbook_docx(str(output), {"vtp_safety_baseline": rejected}, "fixture")
        runbook_text = _docx_text(output)
        explorer = _explorer_result(rejected)
        assert explorer["valid"] is False
        for text in (workbook_text, runbook_text, explorer["rendered"]):
            assert "NOT ASSESSED" in text.upper()
            assert "No rejected receipt leaves" in text
            assert "HOSTILE" not in text
            assert "assessed-000" not in text


def test_explorer_validator_matches_python_structural_result_plus_embedded_custody(receipts):
    valid = receipts["embedded"]
    tampered = copy.deepcopy(valid)
    tampered["rows"][0]["revision"] += 1
    current_json = json.loads(json.dumps(receipts["current"]))
    cases = {
        "valid_embedded": valid,
        "tampered_digest": tampered,
        "serialized_current_claim": current_json,
        "missing": None,
    }
    for name, baseline in cases.items():
        python = validate_vtp_safety_baseline(baseline)
        expected = bool(
            python["valid"]
            and python["baseline"].get("projection_custody") == "embedded_unverified"
        )
        actual = _explorer_result(baseline)["valid"]
        assert actual is expected, name

    html = EXPLORER.read_text(encoding="utf-8")
    block = html[html.index("const _VTP_SAFETY_SCHEMA="):html.index(
        "/* the dynamic routing protocols", html.index("const _VTP_SAFETY_SCHEMA="))]
    assert "b.rows.length>4096" in block
    assert "b.coverage.length>4096" in block
    assert "b.findings.length>65536" in block
    assert 'b.projection_custody!=="embedded_unverified"' in block
    assert "source_sha256" not in block[block.index("function vtpSafetyCoverageRow"):]
    assert "projection_sha256" not in block[block.index("function vtpSafetyCoverageRow"):]


def test_explorer_mounts_vtp_without_replacing_health_or_intelligence():
    html = EXPLORER.read_text(encoding="utf-8")
    draw = html[html.index("function drawProtocols()"):html.index("function drawProtocolsDetail")]
    detail = html[html.index("function drawProtocolsDetail"):html.index("CAUSALITY DRAWER")]
    assert draw.count("${vtpSafetySection()}") == 2
    assert "${vtpSafetySection(host)}" in detail
    assert "${protoIntelSection()}" in draw and "${protoHealthSection()}" in draw
    assert "vtpRows" in html[html.index("function protocolDetailHosts"):html.index("function drawProtocols()")]
