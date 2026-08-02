"""NEW-V3.23.93: the Assessment & Migration Runbook (DOCX) deliverable. python-docx is optional, so
the whole module is skipped when it is not installed (the generator itself fails soft the same way)."""
import pytest

docx = pytest.importorskip("docx")  # skip the file if the optional dep is absent
from docx import Document  # noqa: E402

from cisco_toolkit.runbook import write_runbook_docx  # noqa: E402


def _snap():
    """A minimal but representative snapshot exercising every section the generator builds."""
    return {
        "script_version": "V3.23.0",
        "devices": {
            "sw1": {"hostname": "sw1", "model": "C9300", "sw_version": "17.9", "platform": "ios",
                    "ps_status": "ok", "fan_status": "ok", "temperature_status": "ok",
                    "total_ports": 48, "active_ports": 10},
            "sw2": {"hostname": "sw2", "model": "N9K-C93180", "ps_status": "fail", "fan_status": "ok",
                    "temperature_status": "ok", "total_ports": 48, "active_ports": 5},
        },
        "interfaces": {
            "sw1": {"Gi1/0/1": {"switchport_mode": "Access", "vlan": "10", "end_host_mac": "aaaa.0000.0001"},
                    "Gi1/0/2": {"switchport_mode": "Access", "vlan": "10", "end_host_mac": "aaaa.0000.0002"},
                    "Te1/1": {"switchport_mode": "Trunk", "end_host_mac": "cccc.0000.0001"}},  # trunk excluded
            "sw2": {"Gi1/0/1": {"switchport_mode": "Access", "vlan": "20", "end_host_mac": "bbbb.0000.0001"}},
        },
        "health_scores": [{"switch": "sw1", "score": 50, "band": "Poor"},
                          {"switch": "sw2", "score": 10, "band": "Critical"}],
        "move_groups": [{"switches": ["sw1"], "endpoints": 2}, {"switches": ["sw2"], "endpoints": 1}],
        "migration_readiness": [{"group": "Group 1", "switches": ["sw1"], "endpoints": 2,
                                 "readiness": "NOT READY", "n_fail": 1, "n_warn": 0, "checks": []}],
        "wave_sequencing": [{"group": "Group 1", "make_before_break": ["sw1"], "hard_cutover": [],
                             "hard_cutover_endpoints": 0, "sequence": "make-before-break"}],
        "cross_layer": [{"id": "CL-01", "severity": "Critical", "layers": "L1+L3",
                         "title": "VLAN 10: single-fiber uplink to a sole gateway",
                         "detail": "single fiber fronts the sole gateway", "recommendation": "add redundancy",
                         "hosts": ["sw1"]}],
        "punchlist": [{"severity": "Critical", "category": "Cross-layer", "devices": ["sw1"],
                       "title": "x", "detail": "y"}],
        "failure_impact": [{"host": "sw1", "severity": "High", "vlans_impacted": 3, "stranded": 5,
                            "hard": 2, "backup": 1, "fhrp": 0, "detail": "..."}],
        "link_centrality": [{"a_host": "sw1", "a_port": "Te1/1", "b_host": "sw2", "b_port": "Te1/1",
                             "betweenness": 1.0, "is_bridge": True, "pairs_cut": 4, "rank": 1}],
        "l3_forwarding": [{"switch": "sw1", "vlan": "10", "svi_ip": "10.0.10.1", "fhrp": "",
                           "role": "", "risk": "single-gateway"}],
        "capacity": [{"hostname": "sw1", "model": "C9300", "total_ports": 48, "active_ports": 10,
                      "free_ports": 38, "port_util": 20.8}],
        "operational_drift": [{"severity": "High", "category": "False-health", "devices": ["sw1"],
                               "title": "Temporary L2 bridge on sw1",
                               "detail": "a temp bridge enlarges the STP domain", "remediation": "remove it"}],
        "endpoint_identity": [
            {"host": "sw1", "port": "Gi1/0/1", "vlan": "10", "ip": "", "mac": "00:50:56:aa:00:01",
             "mac_count": 1, "vendor": "VMware, Inc.", "endpoint_class": "VM / Hypervisor",
             "confidence": "Inferred-high", "evidence": "vendor 'VMware, Inc.'"},
            {"host": "sw1", "port": "Gi1/0/2", "vlan": "20", "ip": "", "mac": "00:c0:b7:00:00:01",
             "mac_count": 1, "vendor": "APC", "endpoint_class": "UPS/PDU",
             "confidence": "Inferred-high", "evidence": "vendor 'APC'"},
        ],
        "endpoint_dependencies": {
            "clusters": [{"vendor": "VMware, Inc.", "endpoint_class": "VM / Hypervisor", "count": 12,
                          "switches": 4, "vlans": 2, "move_groups": 1, "spans_groups": False}],
            "dual_homed": [{"mac": "bb:00:00:00:00:01", "ip": "10.0.0.5", "vendor": "NetApp",
                            "endpoint_class": "Storage", "switches": ["sw1", "sw2"], "ports": [],
                            "move_groups": ["Group 1"], "split_across_groups": False}],
            "shared_ip": [],
            "affinity": [{"vlan": "10", "total": 12, "classes": {"VM / Hypervisor": 12}, "dominant": "VM / Hypervisor"}],
            "per_switch_validation": {"sw1": ["Storage: datastore/LUN reachable AND the cluster peer is up"]},
        },
        "subnet_intelligence": {
            "per_device": [
                {"host": "sw1", "is_l3": True, "destination_subnets": ["10.0.10.0/24"],
                 "destination_count": 1, "reachable_count": 2, "reachable_sources": {"ospf": 2},
                 "reachable_sample": [], "default_next_hop": "10.0.0.1", "served_subnets": [],
                 "bgp_received_count": 0},
                {"host": "sw2", "is_l3": False, "destination_subnets": [], "destination_count": 0,
                 "reachable_count": 0, "reachable_sources": {}, "reachable_sample": [],
                 "default_next_hop": "", "served_subnets": [{"subnet": "10.0.10.0/24", "gateway": "sw1", "vlan": "10"}],
                 "bgp_received_count": 0},
            ],
            "move_groups": [{"group": "Group 1", "switches": 2, "local_subnets": ["10.0.10.0/24"],
                             "local_count": 1, "remote_count": 3}],
            "bgp_received_collected": False,
        },
        "migration_scenarios": {
            "per_group": [{"group": "Group 1", "switches": 5, "endpoints": 50, "readiness": "READY",
                           "make_before_break": 4, "hard_cutover": 1, "hard_cutover_endpoints": 2,
                           "dual_homed_pct": 80, "recommended_scenario": "parallel-run",
                           "rationale": "mostly dual-homed — build beside and cut leg-by-leg",
                           "playbook": {"pre": "build beside", "validate": "cut one leg, prove forwarding",
                                        "rollback": "fail back to legacy leg"}}],
            "fleet_recommendation": "89% of switches are Poor/Critical — consider a GREENFIELD rebuild.",
            "scenario_counts": {"parallel-run": 1},
        },
        "collection_completeness": {
            "summary": {"inventory": 3, "complete": 1, "partial": 1, "not_collected": 1},
            "devices": [
                {"host": "sw3-unreachable", "status": "not collected", "data_quality": 0,
                 "missing": ["interface status", "switchport", "version/inventory", "CDP/LLDP neighbors"]},
                {"host": "sw2", "status": "partial", "data_quality": 75, "missing": ["CDP/LLDP neighbors"]},
            ],
        },
        "protocol_intelligence": [
            {"switch": "sw1", "protocol": "EtherChannel", "state": "I", "severity": "High",
             "meaning": "Member is stand-alone (individual) — NOT bundled.",
             "likely_cause": "Incompatible port settings vs the bundle, or no LACP partner.",
             "remediation": "Make the member config identical to the bundle and verify LACP on the peer.",
             "confidence": "observed state = fact; likely cause = Inferred (Cisco doctrine)"},
        ],
        "service_map": {
            "services": [
                {"port": 319, "proto": "udp", "service": "PTP-event", "category": "Broadcast-AV",
                 "broadcast": True, "refs": 2, "host_count": 1,
                 "evidence_class": "Inferred (ACL design intent -- not active traffic; no flow telemetry)"},
            ],
            "categories": [{"category": "Broadcast-AV", "refs": 2}],
            "acl_rule_count": 12,
            "multicast": {"active_interfaces": 4, "active_switch_count": 2,
                          "active_switches": ["sw1", "sw2"],
                          "classified_groups": [{"group": "224.0.1.129", "name": "PTP-primary",
                                                 "category": "Broadcast-AV", "broadcast": True, "source": "IGMP/mroute"}],
                          "group_level_collected": True,
                          "ptp": {"sw1": {"device_type": "Unknown", "num_ports": 0, "grandmaster": "", "operational": False}},
                          "igmp_queriers": [{"switch": "sw1", "vlan": "10", "querier": "10.0.10.1"}]},
        },
    }


def test_runbook_has_12_sections_and_reconciles(tmp_path):
    out = str(tmp_path / "rb.docx")
    write_runbook_docx(out, _snap(), "Unit Test Fleet")
    d = Document(out)
    h1 = [p.text for p in d.paragraphs if p.style.name == "Heading 1"]
    # the 12 standard sections (plus a Contents heading)
    for n in range(1, 13):
        assert any(t.startswith(f"{n}.") for t in h1), f"missing section {n}: {h1}"

    # numbers reconcile to the snapshot (the workbook-vs-runbook agreement contract).
    # Locate the §1 metric table by its header — the document-control front matter
    # (V3.23.150) inserts tables before it, so index 0 is no longer the exec summary.
    exec_t = next((t for t in d.tables if t.rows[0].cells[0].text == "Metric"), None)
    assert exec_t is not None, "exec-summary Metric table not found in the runbook"
    exec_rows = {r.cells[0].text: r.cells[1].text for r in exec_t.rows}
    assert exec_rows["Devices in scope"] == "2"
    assert exec_rows["Migration move groups"] == "2"
    assert exec_rows["Punch-list items"] == "1"
    assert exec_rows["Endpoints — access-port host MACs at snapshot (superset of canonical)"] == "3"   # trunk MAC excluded
    assert "Evidenced endpoints (canonical assessment scale)" in exec_rows   # SSOT-endpoint-deliv-1: canonical headline present


def _all_text(doc):
    """All visible text — paragraphs AND table cells (findings live in both)."""
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            parts.extend(c.text for c in row.cells)
    return "\n".join(parts)


def test_runbook_scope_distinguishes_collected_from_inventory(tmp_path):
    """Coverage-honesty: §2 Scope must NOT call the device INVENTORY total 'the collected dataset' — it
    states how many of the inventoried switches were actually collected (collection_completeness.complete)."""
    snap = _snap()
    snap["collection_completeness"] = {"summary": {"complete": 2, "inventory": 9}}
    out = str(tmp_path / "rb_scope.docx")
    write_runbook_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "present in the collected dataset" not in text   # the inventory-mislabeled-as-collected wording is gone
    assert "inventoried" in text                            # the honest inventoried/collected framing


def test_runbook_renders_device_risk_register_section(tmp_path):
    """NEW-V3.23.174: a snapshot carrying the Device Risk Register renders §10.1 with the
    ranked table + compound bullets; without the key the section is absent (data-gated)."""
    snap = _snap()
    out = str(tmp_path / "rb_no_register.docx")
    write_runbook_docx(out, snap, "Unit Test Fleet")
    assert "10.1 Per-asset compound risk" not in _all_text(Document(out))

    snap["device_dossiers"] = {
        "per_device": [
            {"host": "sw2", "risk_band": "Severe", "risk_index": 60, "impact_score": 10,
             "exposure_score": 6,
             "compound": [{"code": "CR-02", "title": "Failing hardware past support",
                           "severity": "High", "basis": "Critical health on EoL hardware."}],
             "verdict": "Stabilize or replace before migration — health Critical."}],
        "summary": {"n_devices": 1, "bands": {"Severe": 1, "Elevated": 0, "Guarded": 0, "Low": 0},
                    "n_compound": 1, "worst": ["sw2"]}}
    out2 = str(tmp_path / "rb_register.docx")
    write_runbook_docx(out2, snap, "Unit Test Fleet")
    text = _all_text(Document(out2))
    assert "10.1 Per-asset compound risk" in text
    assert "CR-02" in text and "Stabilize or replace" in text
    assert "not assessed" in text          # the honesty rule is stated in the section intro


def test_runbook_is_evidence_disciplined(tmp_path):
    out = str(tmp_path / "rb.docx")
    write_runbook_docx(out, _snap(), "Unit Test Fleet")
    text = _all_text(Document(out))
    # every material finding carries the evidence frame
    for token in ("Observed Evidence:", "Interpretation:", "Impact / Blast Radius:",
                  "Confidence:", "Unknowns:", "Next Validation:"):
        assert token in text, token
    # false-health doctrine + confidence vocabulary are present
    assert "gateway-active is not service proof" in text
    # §6.5 protocol behaviour: the abnormal state's cause + remediation are present, evidence-framed
    assert "Protocol behaviour & remediation" in text
    assert "verify LACP on the peer" in text
    # §2.1 collection completeness: blind-spot devices are made explicit (not-collected listed)
    assert "Collection completeness" in text
    assert "sw3-unreachable" in text and "a missing device is not a healthy device" in text
    # §6.6 services & multicast: design-intent caveat + multicast forwarding presence
    assert "Services & multicast" in text
    assert "PTP-event" in text and "run PIM/mroute" in text
    # PTP operational-state finding: dormant (not boundary-clocked) is surfaced, not crashed on the {host:dict} map
    assert "NONE are active boundary" in text
    assert "Inferred-high" in text and "Unknown" in text
    # the cross-layer finding surfaced as a titled block
    assert "single-fiber uplink to a sole gateway" in text
    # the false-health / operational-drift section (§6.3) surfaced the drift finding
    assert "False-health / operational drift" in text
    assert "Temporary L2 bridge on sw1" in text
    # the endpoint identity section (§7.1) surfaced vendor + class grouping
    assert "Endpoint identity (vendor & type)" in text
    assert "VM / Hypervisor" in text and "VMware, Inc." in text
    # the dependency intelligence (§8.1 clusters/dual-homed + §11.1 per-class validation)
    assert "Endpoint clusters & dependencies" in text
    assert "Service validation by endpoint class" in text
    assert "datastore/LUN reachable" in text
    # subnet & routing reachability (§6.4) + source<->destination
    assert "Subnet & routing reachability" in text
    assert "Remote subnets to preserve" in text
    # migration scenario framework (§3 recommendation + §11.2 playbook)
    assert "GREENFIELD rebuild" in text and "parallel-run" in text
    assert "Cutover playbook by scenario" in text


def test_runbook_validation_section_uses_validation_group_noun(tmp_path):
    """QA row 12 (noun overload): §11.3 counts validation_plan.by_wave buckets (one per move-group, keyed
    'Group N'), a DISTINCT unit from the sequenced wave_plan waves. The noun must read 'validation group',
    reserving 'wave' for the wave_plan, so 'N check(s) across M ...' cannot be misread as the wave count."""
    snap = _snap()
    snap["validation_plan"] = {
        "banner": "Run after each cutover to confirm it succeeded.",
        "by_wave": {"Group 1": [{"device": "sw1", "category": "Gateway", "severity": "High",
                                 "check": "Default gateway for VLAN 10 is up", "command": "show ip int brief",
                                 "expect": "Vlan10 up/up", "why": "endpoints lose their gateway"}]},
        "summary": {"n_items": 1, "n_waves": 1, "n_high": 1, "by_category": {"Gateway": 1}},
    }
    out = str(tmp_path / "rb_val.docx")
    write_runbook_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "Post-cutover verification plan (per validation group)" in text
    assert "check(s) across 1 validation group(s)" in text          # the count noun matches the unit
    assert "Validation group: Group 1" in text                      # the per-bucket heading
    assert "1 wave(s)" not in text and "Wave: Group 1" not in text   # the old overloaded phrasings are gone


def test_runbook_carries_document_furniture(tmp_path):
    """V3.23.150: AS-style front/back matter — Document Control between cover and TOC, and the
    closing acceptance signature gate."""
    out = str(tmp_path / "rb.docx")
    write_runbook_docx(out, _snap(), "Unit Test Fleet")
    d = Document(out)
    h1 = [p.text for p in d.paragraphs if p.style.name == "Heading 1"]
    assert "Document Control" in h1 and "Document Acceptance" in h1
    text = _all_text(d)
    assert "Revision history" in text and "Assumptions & caveats" in text
    assert "Customer network owner" in text                      # acceptance signature roles
    assert "Assessment workbook (.xlsx)" in text                 # related-documents cross-reference


def test_runbook_failsoft_without_python_docx(monkeypatch, tmp_path):
    """If python-docx is not importable, the generator warns and returns -- it never crashes a run
    whose workbook/explorer/JSON already saved."""
    import builtins
    real_import = builtins.__import__

    def _blocked(name, *a, **k):
        if name == "docx" or name.startswith("docx."):
            raise ImportError("simulated missing python-docx")
        return real_import(name, *a, **k)

    monkeypatch.setattr(builtins, "__import__", _blocked)
    out = str(tmp_path / "rb.docx")
    write_runbook_docx(out, _snap(), "Unit Test Fleet")   # must not raise
    import os
    assert not os.path.exists(out)                          # nothing written, no crash


def test_runbook_endpoint_census_uses_canonical_not_mac_sum(tmp_path):
    """Single source of truth: §7 surfaces the canonical evidenced-endpoint count
    (executive_brief.scale.n_endpoints) labeled as endpoints, and labels the access-port MAC sum as
    MACs -- it must NOT publish the MAC sum under the bare word 'Total endpoints' (which diverges from
    the family-wide endpoint count)."""
    snap = _snap()
    snap["executive_brief"] = {"scale": {"n_devices": 2, "n_endpoints": 99, "n_domains": 1}}
    out = str(tmp_path / "rb_ep.docx")
    write_runbook_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "Evidenced endpoints: 99" in text          # canonical, labeled as endpoints
    assert "Access-port host MACs:" in text           # the MAC sum, labeled as MACs
    assert "Total endpoints:" not in text             # the old bare mislabel is gone


def test_runbook_no_fhrp_count_excludes_none_sentinel(tmp_path):
    """False-health guard (the 'none'-is-truthy bug class): gateways carry fhrp='none' when no FHRP is
    configured, and 'none'.strip() is truthy — so the §6 'N of M gateways have no FHRP peer' line
    counted ZERO single-gateway gateways for an all-'none' fleet, the exact inverse of the truth (every
    gateway single-homed). The no-FHRP test must treat 'none' as no-FHRP."""
    snap = _snap()
    snap["l3_forwarding"] = [
        {"switch": "sw1", "vlan": "10", "svi_ip": "10.0.10.1", "fhrp": "none"},        # no FHRP
        {"switch": "sw1", "vlan": "20", "svi_ip": "10.0.20.1", "fhrp": "none"},        # no FHRP
        {"switch": "sw1", "vlan": "30", "svi_ip": "10.0.30.1", "fhrp": "HSRP active"}, # real FHRP
    ]
    out = str(tmp_path / "rb_fhrp.docx")
    write_runbook_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "2 of 3 gateways have no FHRP peer" in text   # the two 'none' gateways, not hidden
    assert "0 of 3 gateways have no FHRP peer" not in text


def test_single_gateway_exposure_uses_the_risk_field_not_the_no_fhrp_count(tmp_path):
    """§6.1 counted `fhrp == "none"` and labelled the total 'single-gateway exposure', but the
    producer (excel.write_l3_forwarding_sheet) emits 'single-gateway' only when the VLAN has <=1
    gateway and 'no-FHRP' otherwise — two conditions of very different severity. On a fleet with one
    genuine single-gateway VLAN and seven dual-gateway no-FHRP VLANs the war room read 15 where
    archreview RES-2 and design §2.4 read 1."""
    snap = _snap()
    l3f = [{"switch": "sw1", "vlan": "10", "svi_ip": "10.0.10.1", "fhrp": "", "risk": "single-gateway"}]
    for v in range(20, 27):                       # 7 VLANs x 2 gateways, no FHRP protocol
        for host in ("sw1", "sw2"):
            l3f.append({"switch": host, "vlan": str(v), "svi_ip": f"10.0.{v}.1",
                        "fhrp": "", "risk": "no-FHRP"})
    snap["l3_forwarding"] = l3f
    out = str(tmp_path / "rb_sgw.docx")
    write_runbook_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    canonical = sum(1 for r in l3f if "single-gateway" in str(r.get("risk") or ""))
    assert canonical == 1
    assert "Of those, 1 carry the SINGLE-GATEWAY risk" in text
    # the merged 15 must never be presented as the single-gateway exposure
    assert "15 of 15 gateways have no FHRP peer in scope" in text     # the no-FHRP fact, correctly labelled
    assert "single-gateway exposure" not in text.replace("SINGLE-GATEWAY", "")
    # and with no risk classification at all the count ABSTAINS rather than reading 0 = clean
    for r in snap["l3_forwarding"]:
        r.pop("risk", None)
    out2 = str(tmp_path / "rb_norisk.docx")
    write_runbook_docx(out2, snap, "Unit Test Fleet")
    text2 = _all_text(Document(out2))
    assert "no L3-risk classification" in text2
    assert "carry the SINGLE-GATEWAY risk" not in text2


def test_runbook_drift_table_keeps_disambiguation_and_remediation_on_long_detail(tmp_path):
    """PR-#396 review: the §6.3 why/next-step cell was a single (detail+remediation)[:160] slice, so a
    long detail (the 270-char native-VLAN-1 row) evicted BOTH its own disambiguating tail and the
    ENTIRE remediation. The cell now budgets the two parts separately -- the detail's parenthetical
    and the 'Fix:' next step must both render."""
    snap = _snap()
    long_detail = ("4 operationally-trunking port(s) (live trunk status) across 3 live-trunking "
                   "switch(es) carry the default VLAN 1 as the native (untagged) VLAN -- a hygiene and "
                   "VLAN-hopping exposure. (The design gap counts trunk-mode port(s) -- switchport-mode "
                   "basis, administrative or negotiated -- so the two figures can differ.)")
    snap["operational_drift"] = [
        {"severity": "Low", "category": "False-health", "devices": ["sw1", "sw2"],
         "title": "Native VLAN 1 on 4 operationally-trunking port(s)", "detail": long_detail,
         "remediation": "Set a dedicated, unused native VLAN on every 802.1Q trunk."}]
    out = str(tmp_path / "rb_drift_long.docx")
    write_runbook_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "so the two figures can differ" in text                  # disambiguating tail survives
    assert "Fix: Set a dedicated, unused native VLAN" in text       # remediation survives


# every top-level snapshot section the runbook writer dereferences (reads via `.get(...)` then
# .items()/.get()/iterates/len()/slices). The webapp `--no-collect` path makes this JSON
# attacker-controllable, so a TRUTHY non-dict/non-list section (a scalar, a string) must degrade
# to empty, never raise — the same crash class fixed in parse/design_advisor/ssot/design.
_RUNBOOK_SECTIONS = (
    "devices", "interfaces", "health_scores", "move_groups", "migration_readiness",
    "wave_sequencing", "cross_layer", "punchlist", "failure_impact", "link_centrality",
    "l3_forwarding", "executive_brief", "collection_completeness", "service_map",
    "migration_scenarios", "lifecycle_risk", "operational_drift", "subnet_intelligence",
    "protocol_intelligence", "multicast_intelligence", "application_intelligence",
    "segmentation", "golden_drift", "syslog_intelligence", "qos_audit", "software_risk",
    "platform_health", "endpoint_identity", "capacity", "endpoint_dependencies",
    "device_dossiers", "validation_plan", "remediation_plan",
)


@pytest.mark.parametrize("section", _RUNBOOK_SECTIONS)
@pytest.mark.parametrize("poison", [5, "boom", True, 3.14])
def test_runbook_survives_truthy_scalar_section(tmp_path, section, poison):
    """Crash-hardening (audit-5 totality): setting ANY snapshot section to a truthy scalar used to slip
    the `snap.get(x) or {}` / `or []` guards (which catch only falsy values) and then crash on
    `.items()` / `.get()` / iteration / len() / slicing. Every section must now degrade to empty and
    the runbook must still render end-to-end (§1 present), not raise."""
    snap = _snap()
    snap[section] = poison
    out = str(tmp_path / f"rb_poison_{section}.docx")
    write_runbook_docx(out, snap, "Unit Test Fleet")   # must not raise
    import os
    assert os.path.exists(out), f"runbook not written for poisoned {section!r}={poison!r}"
    # the whole document still renders — the poisoned section degrades to empty, the rest is intact
    assert "1. Assessment Header & Executive Summary" in _all_text(Document(out))


def test_runbook_survives_truthy_scalar_flow_paths(tmp_path):
    """The flow_paths argument is read `(flow_paths or {}).get(...)` the same way; a truthy non-dict
    must degrade, not crash the §6.4.1 flow-path table."""
    for bad in (5, "boom", [1, 2]):
        out = str(tmp_path / f"rb_fp_{type(bad).__name__}.docx")
        write_runbook_docx(out, _snap(), "Unit Test Fleet", flow_paths=bad)   # must not raise
        assert "1. Assessment Header & Executive Summary" in _all_text(Document(out))


# --- audit-5 totality, NESTED level -----------------------------------------------------------------
# The section-scalar test above proves a truthy scalar *section* degrades. But a well-formed section
# (a dict / list-of-dicts) whose INNER value is a truthy scalar still slipped the row-level
# `.get(x) or []` / `or {}` guards (they catch only falsy) and crashed len() / .items() / iteration /
# subscription — e.g. migration_readiness[0]["switches"] = 7 -> len(7). The top-level test cannot reach
# these (a scalar section empties the row list first). Same reachable stored-DoS class, via
# deliverables.generate("runbook", snap) on an attacker-uploaded --no-collect snapshot.
def _rich_snap():
    """_snap() + the row-bearing sections whose inner values are dereferenced, so a poisoned inner
    value actually reaches the vulnerable read (application domains, software-risk per-device)."""
    s = _snap()
    s["application_intelligence"] = {
        "summary": {"n_domains": 1, "n_on_air_critical": 1, "n_high_risk": 1},
        "domains": [{"domain": "D1", "tier": "On-air critical", "switch_count": 1, "endpoint_count": 2,
                     "ptp_present": True, "evidence": "e",
                     "risks": [{"severity": "Critical", "title": "t", "detail": "d", "remediation": "r"}],
                     "validation": ["ping"]}],
        "cross_domain_risks": [{"title": "x"}],
        "edges": [{"source": "D1", "target": "D2", "weight": 1, "kinds": ["link"], "media": True,
                   "migration_note": "n"}],
        "keystones": [{"domain": "D1", "degree": 2, "neighbors": ["D2", "D3"]}],
        "cutover_order": [{"domain": "D1"}],
    }
    s["software_risk"] = {"summary": {"crit_0_2": 0, "err_3": 0},
                          "per_device": [{"host": "sw1", "advisories": [{"cve": "CVE-1"}],
                                          "compound": [{"code": "X"}], "missing": ["patch"]}]}
    return s


# depth 5 is EXHAUSTIVE for _rich_snap() — it bottoms out there (a depth-6 sweep yields the identical
# 129 paths). The old depth<=3 cap stopped ONE level short of `collection_completeness.devices[i].missing`
# and TWO short of `application_intelligence.domains[i].risks[j]`, both of which were live crashes; the
# cap, not the code, was what made the sweep green. Cost of going exhaustive: 99 -> 129 poison cases.
_MAX_NEST_DEPTH = 5


def _nested_container_paths(obj, path=(), depth=0):
    """Every path (depth<=_MAX_NEST_DEPTH) to a nested dict/list inside the snapshot, excluding the root."""
    if path and isinstance(obj, (dict, list)):
        yield path
    if depth >= _MAX_NEST_DEPTH:
        return
    if isinstance(obj, dict):
        for k, v in obj.items():
            yield from _nested_container_paths(v, path + (k,), depth + 1)
    elif isinstance(obj, list):
        for i, v in enumerate(obj):
            if isinstance(v, (dict, list)):
                yield from _nested_container_paths(v, path + (i,), depth + 1)


def test_runbook_survives_truthy_scalar_nested_value(tmp_path):
    """Replacing ANY nested container (a section's inner dict/list, to the fixture's full depth) with a
    truthy scalar must degrade to empty, never raise — the runbook still renders end-to-end.
    Revert-proof by construction: e.g. migration_readiness[0]["switches"]=7 -> len(7) crashes the
    pre-extension module."""
    import copy
    base = _rich_snap()
    # `executive_brief` is ALSO consumed by the shared docmeta.add_excellence_front -> ssot.reconcile
    # layer; the nested-scalar hardening of it lives in ssot.py (PR #452, verified:
    # `cc = _as_dict(_dotted(snap,"collection_completeness.summary"))`), so this runbook sweep
    # deliberately does not own it — a principled module boundary.
    # `collection_completeness` was excluded on that same reasoning and it was WRONG: runbook.py has
    # its OWN read of collection_completeness.devices[].missing (§2.1's blind-spot table), which the
    # exclusion hid — a live `", ".join(d.get("missing", []))` crash. It is back in the sweep.
    _SHARED_SSOT_SECTIONS = {"executive_brief"}
    paths = [p for p in _nested_container_paths(base) if p[0] not in _SHARED_SSOT_SECTIONS]
    assert len(paths) > 20, "sanity: the rich snapshot must expose many nested containers to poison"
    # REACH GUARD. The cap silently going stale is exactly how this sweep stayed green over four live
    # crashes: at depth<=3 it never generated these paths, so "no crashes" meant "never looked". Pin
    # the deepest sites explicitly — lowering _MAX_NEST_DEPTH, or a fixture that drops these keys,
    # must fail loudly here rather than quietly shrink the sweep's reach.
    for must_reach in (("collection_completeness", "devices", 0, "missing"),
                       ("migration_scenarios", "per_group", 0, "playbook"),
                       ("application_intelligence", "domains", 0, "risks", 0)):
        assert must_reach in paths, (
            f"sweep no longer reaches {'.'.join(map(str, must_reach))} — the poison set shrank "
            f"(depth cap {_MAX_NEST_DEPTH}, {len(paths)} paths); it would go green without looking")
    crashes = []
    for i, p in enumerate(paths):
        snap = copy.deepcopy(base)
        node = snap
        for key in p[:-1]:
            node = node[key]
        node[p[-1]] = 7   # a truthy scalar where a nested dict/list is expected
        # ONE output path, rewritten per case. Nothing reads these files — the assertion is purely
        # "the render did not raise" — but a distinct path per case left 129 .docx files
        # accumulating in one directory, and on Windows each new file is a fresh AV scan in a
        # directory that keeps growing. Measured: 262 ms/render into a near-empty directory vs
        # 174 ms rewriting one path, and this test was the single slowest in the suite at ~155 s
        # (22% of the whole run) against ~34 s for the same 129 renders measured standalone — the
        # gap is file accumulation, not the rendering.
        #
        # The sweep's REACH is untouched: same 129 cases, same poison, same assertion. Shrinking the
        # case set is what the reach guard above exists to forbid, and it is not what this does.
        out = str(tmp_path / "rb_nested.docx")
        try:
            write_runbook_docx(out, snap, "Unit Test Fleet")
        except Exception as e:   # noqa: BLE001 - the whole point is that NOTHING may raise
            crashes.append((".".join(map(str, p)), type(e).__name__, str(e)[:60]))
    assert not crashes, f"{len(crashes)} nested truthy-scalar case(s) still crash the runbook: {crashes[:10]}"


def test_runbook_survives_non_string_elements_in_joined_lists(tmp_path):
    """The container sweep above poisons CONTAINERS, so it structurally cannot reach this: a
    well-formed LIST whose ELEMENTS are non-str still crashed `", ".join(...)` with
    `TypeError: sequence item 0: expected str instance, int found`. Both `missing` renders (§2.1
    collection blind spots, §6.9 golden-config drift) read an attacker-supplied list straight into
    join(), so both must stringify their elements. Same uploaded-snapshot DoS route."""
    import copy
    base = _snap()
    base["golden_drift"] = {
        "summary": {"n_baseline": 3, "n_drifting": 1, "n_devices": 2, "avg_compliance_pct": 66,
                    "mode": "majority"},
        "per_device": [{"host": "sw1", "compliance_pct": 66, "n_missing": 1, "missing": ["aaa new-model"]}],
    }
    for i, (label, mutate) in enumerate((
        ("collection_completeness.devices[0].missing",
         lambda s: s["collection_completeness"]["devices"][0].__setitem__("missing", [1, 2, None])),
        ("golden_drift.per_device[0].missing",
         lambda s: s["golden_drift"]["per_device"][0].__setitem__("missing", [7, {"a": 1}])),
    )):
        snap = copy.deepcopy(base)
        mutate(snap)
        out = str(tmp_path / f"rb_join_{i}.docx")
        write_runbook_docx(out, snap, "Unit Test Fleet")   # must not raise
        assert "1. Assessment Header & Executive Summary" in _all_text(Document(out)), label


# --- leaf-type class: a non-str ELEMENT inside a well-formed list -----------------------------------
# DISTINCT from the truthy-non-container class above and NOT fixable by _as_list: `_as_list` guards the
# container's TYPE and is structurally blind to what is inside it. `", ".join([5])` raises
# `TypeError: sequence item 0: expected str instance, int found`. The snapshot is attacker-controllable
# (upload -> deliverables.generate("runbook", snap) -> HTTP 500), so each is a stored DoS.
_LEAF_JOIN_CASES = [
    ("cross_layer[].hosts", {"cross_layer": [{"id": "C1", "severity": "Critical", "title": "t",
                                              "detail": "d", "hosts": [5]}]}),
    ("app_int.cross_domain_risks[].title",
     {"application_intelligence": {"domains": [{"domain": "D1", "tier": "On-air critical"}],
                                   "cross_domain_risks": [{"title": 5}]}}),
    ("app_int.domains[].validation",
     {"application_intelligence": {"domains": [{"domain": "D1", "tier": "On-air critical", "evidence": "e",
                                                "risks": [{"severity": "Critical", "title": "t",
                                                           "detail": "d", "remediation": "r"}],
                                                "validation": [5]}]}}),
    ("app_int.edges[].kinds",
     {"application_intelligence": {"domains": [{"domain": "D1"}],
                                   "edges": [{"source": "a", "target": "b", "weight": 1, "kinds": [5]}]}}),
]


@pytest.mark.parametrize("label,extra", _LEAF_JOIN_CASES, ids=[c[0] for c in _LEAF_JOIN_CASES])
def test_runbook_survives_nonstr_list_element(tmp_path, label, extra):
    """A non-str element of an otherwise well-formed list must degrade, never crash the join."""
    snap = {"devices": {"sw1": {"hostname": "sw1"}}}
    snap.update(extra)
    out = tmp_path / "leaf.docx"
    write_runbook_docx(str(out), snap, "Unit Test Fleet")   # must not raise
    assert out.is_file(), f"no runbook for {label}"


def test_runbook_nonstr_join_coercion_is_identity_on_strings(tmp_path):
    """Non-vacuity companion: str()-per-element must be a NO-OP for well-formed string members --
    every host must still appear verbatim in the rendered cross-layer finding."""
    snap = {"devices": {"sw1": {"hostname": "sw1"}},
            "cross_layer": [{"id": "C1", "severity": "Critical", "title": "t", "detail": "d",
                             "hosts": ["alpha1", "bravo2"]}]}
    out = str(tmp_path / "wf.docx")
    write_runbook_docx(out, snap, "Unit Test Fleet")
    txt = _all_text(Document(out))
    assert "alpha1" in txt and "bravo2" in txt, "well-formed hosts must render verbatim"


# ===========================================================================================
# review r9 F1/F2 — the runbook's two un-qualified exits:
#   §10.1 rendered `risk_band` with no coverage qualification, and §6.6 rendered the curated
#   broadcast/AV classification with no authority. Both are "absence rendered as health".
# ===========================================================================================

def _real_dossiers(assessed: bool):
    """A dossier section built by the REAL producer (analyze.compute_device_dossiers), not a
    hand-shaped dict. A hand-built row in the shape the writer expects would only prove the writer
    agrees with itself; this proves the writer agrees with the ENGINE.

    assessed=False feeds ONLY health + blast radius + an Unknown-lifecycle row -> 8 of the 11 risk
    axes abstain -> exposure 0 -> risk_band 'Low', risk_index 0, verdict "No stacked risk — routine
    migration handling." That is the exact false-health shape this section fixes. assessed=True
    feeds most axes -> 5 of 11 na -> NOT thin (the non-vacuity control)."""
    from cisco_toolkit.analyze import compute_device_dossiers
    if not assessed:
        return compute_device_dossiers(
            health_scores=[{"switch": "sw0", "score": 70, "band": "Good", "role": "access"}],
            failure_impact=[{"host": "sw0", "stranded": 40, "vlans_impacted": 2}],
            lifecycle_risk={"per_device": [{"host": "sw0", "band": "Unknown"}]})
    return compute_device_dossiers(
        health_scores=[{"switch": "sw0", "score": 92, "band": "Excellent", "role": "access"}],
        failure_impact=[{"host": "sw0", "stranded": 4, "vlans_impacted": 1}],
        lifecycle_risk={"per_device": [{"host": "sw0", "band": "Active", "model": "C9300"}]},
        software_risk={"per_device": [{"host": "sw0", "band": "OK", "n_findings": 0}]},
        security={"sw0": {"checks": [{"id": "x", "status": "pass"}]}},
        config_hygiene={"sw0": {"findings": []}})


def test_dossier_coverage_mirrors_the_canonical_rule():
    """runbook._dossier_coverage is a MIRROR of cisco_toolkit/excel.py::dossier_coverage. Drift is
    caught by EXECUTION over the canonical case table, not by a comment asking for discipline."""
    from cisco_toolkit.excel import DOSSIER_COVERAGE_CASES, dossier_coverage
    from cisco_toolkit import runbook as R

    # NON-VACUITY of the table: it must exercise BOTH verdicts, else "they agree" is trivial.
    thins = {k for k, v in DOSSIER_COVERAGE_CASES.items() if v[1][2]}
    assert thins and (set(DOSSIER_COVERAGE_CASES) - thins)
    for label, (row, expected) in DOSSIER_COVERAGE_CASES.items():
        assert list(R._dossier_coverage(row)) == list(expected), label
        assert list(R._dossier_coverage(row)) == list(dossier_coverage(row)), label


def test_device_risk_register_discloses_not_assessed_axes(tmp_path):
    """review r9 F1. The §10.1 register rendered host/band/index/verdict and nothing else, so an
    asset whose 8-of-11 risk axes were NEVER ASSESSED printed as 'Low', index 0, "No stacked risk —
    routine migration handling" — indistinguishable from a genuinely clean asset. Pre-fix the
    rendered document contained neither the per-asset census nor the fleet-wide coverage sentence."""
    dd = _real_dossiers(assessed=False)
    row = dd["per_device"][0]
    # the producer really does band an un-assessed asset 'Low' (this is the defect's precondition)
    assert row["risk_band"] == "Low" and row["risk_index"] == 0 and row["n_na"] == 8

    snap = _snap()
    snap["device_dossiers"] = dd
    out = str(tmp_path / "rb_cov.docx")
    write_runbook_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "8 of 11 NOT ASSESSED" in text, "per-asset coverage cell missing from the register table"
    assert "band computed on absent evidence" in text, "the 'Low' band is not qualified"
    assert "8 of 11 risk axes fleet-wide were NOT ASSESSED" in text
    assert "HALF OR MORE of their risk axes NOT ASSESSED" in text
    assert "Coverage" in text                                  # the new column header renders


def test_device_risk_register_clean_asset_does_not_acquire_the_disclosure(tmp_path):
    """NON-VACUITY: an asset whose axes WERE assessed (5 of 11 na -> not thin) must keep its clean
    reading. If the qualifier appeared on every row it would carry no information."""
    dd = _real_dossiers(assessed=True)
    row = dd["per_device"][0]
    assert row["n_na"] == 5 and len(row["exposures"]) == 11     # below the half-abstained threshold
    snap = _snap()
    snap["device_dossiers"] = dd
    out = str(tmp_path / "rb_clean.docx")
    write_runbook_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "5 of 11 NOT ASSESSED" in text                       # census still stated (always legible)
    assert "band computed on absent evidence" not in text       # ... but NOT flagged as un-evidenced
    assert "HALF OR MORE of their risk axes NOT ASSESSED" not in text


def test_register_with_no_axis_census_fails_closed(tmp_path):
    """An OLDER snapshot whose dossier rows carry no `exposures` census must read NOT ASSESSED —
    an unknown denominator is never 'fine'."""
    snap = _snap()
    snap["device_dossiers"] = {
        "per_device": [{"host": "sw2", "risk_band": "Low", "risk_index": 0, "impact_score": 1,
                        "exposure_score": 0, "compound": [], "verdict": "No stacked risk."}],
        "summary": {"n_devices": 1, "bands": {"Severe": 0, "Elevated": 0, "Guarded": 0, "Low": 1},
                    "n_compound": 0}}
    out = str(tmp_path / "rb_nocensus.docx")
    write_runbook_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "axis census ABSENT — coverage NOT ASSESSED" in text
    assert "published NO risk-axis census at all" in text


def _real_multicast_intelligence():
    """multicast_intelligence from the REAL producer, over two groups that alias onto one L2 MAC,
    one of them curated Broadcast-AV. compute_multicast_intelligence escalates that finding to High
    purely on the curated classification."""
    from cisco_toolkit.analyze import compute_multicast_intelligence
    return compute_multicast_intelligence({"multicast": {
        "classified_groups": [
            {"group": "239.1.1.1", "name": "ST2110-video", "category": "Broadcast-AV",
             "broadcast": True, "source": "IGMP/mroute"},
            {"group": "224.1.1.1", "name": "generic", "category": "", "source": "IGMP/mroute"},
        ],
        "active_switch_count": 1, "active_interfaces": 2, "ptp": {}, "igmp_queriers": []}}, {})


def test_broadcast_av_classification_renders_with_its_authority(tmp_path):
    """review r9 F2. §6.6 printed '(N broadcast/AV)' and 'MAC ← groups' with no hint that the on-air
    label is a CURATED offline-registry judgement carrying no authoritative source — the same label
    that raises the MAC-alias finding from Medium to High. Pre-fix neither string below appeared
    anywhere in the document."""
    mi = _real_multicast_intelligence()
    # the producer's own position: the alias is High, and it rests on a NON-authoritative label
    assert mi["summary"]["n_av_groups_authoritative"] == 0 and mi["summary"]["n_av_groups"] == 1
    assert mi["risks"][0]["severity"] == "High" and mi["mac_aliases"][0]["has_av"] is True
    assert mi["mac_aliases"][0]["has_av_authoritative"] is False

    snap = _snap()
    snap["multicast_intelligence"] = mi
    out = str(tmp_path / "rb_av.docx")
    write_runbook_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "CURATED offline-registry classification, not a measurement" in text
    assert "CURATED classification, NOT an authoritative source" in text
    assert "this flag is what raises the finding to High" in text


def test_broadcast_av_authority_absent_reads_not_assessed(tmp_path):
    """FAIL-CLOSED on an older snapshot: no authority census must read NOT ASSESSED, never
    'authoritative'. NON-VACUITY companion: an AUTHORITATIVE census must NOT print either the
    'NOT ASSESSED' or the 'NONE of them' wording."""
    snap = _snap()
    snap["multicast_intelligence"] = {"summary": {"n_groups": 1, "n_av_groups": 1}}   # legacy: no authority key
    out = str(tmp_path / "rb_av_old.docx")
    write_runbook_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "on-air classification authority NOT ASSESSED in this snapshot" in text

    snap["multicast_intelligence"] = {"summary": {"n_groups": 1, "n_av_groups": 1,
                                                  "n_av_groups_authoritative": 1}}
    out2 = str(tmp_path / "rb_av_auth.docx")
    write_runbook_docx(out2, snap, "Unit Test Fleet")
    text2 = _all_text(Document(out2))
    assert "1 of 1 on-air classification(s) rest on an authoritative source" in text2
    assert "authority NOT ASSESSED" not in text2 and "NONE of them" not in text2


def test_av_authority_qualifier_refuses_to_state_an_incoherent_ratio():
    """`{n_auth} of {n_av}` is meaningful only when n_auth <= n_av. The two counts are published
    independently, so an incoherent snapshot rendered "7 of 3" and a reader had no way to tell that
    from a real ratio. Fail closed on the incoherence instead of formatting it."""
    from cisco_toolkit.runbook import _av_authority

    bad = _av_authority({"summary": {"n_av_groups": 3, "n_av_groups_authoritative": 7}})
    assert "INCOHERENT" in bad, bad
    assert "7 of 3" not in bad

    # NON-VACUITY in all three healthy directions.
    good = _av_authority({"summary": {"n_av_groups": 5, "n_av_groups_authoritative": 2}})
    assert "2 of 5" in good and "INCOHERENT" not in good, good
    none_auth = _av_authority({"summary": {"n_av_groups": 5, "n_av_groups_authoritative": 0}})
    assert "NONE of them" in none_auth
    absent = _av_authority({"summary": {"n_av_groups": 5}})
    assert "NOT ASSESSED" in absent, "an absent authority census must fail CLOSED"
    # a malformed value must not read as authoritative either
    malformed = _av_authority({"summary": {"n_av_groups": 5, "n_av_groups_authoritative": None}})
    assert "NONE of them" in malformed or "NOT ASSESSED" in malformed, malformed
