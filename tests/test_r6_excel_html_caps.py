"""R6 truncation sweep — every reachable, undisclosed `[:N]` display cap in excel.py / html.py /
archreview.py.

Silent truncation was found four separate times in earlier review rounds (the workbook's
subnet-reachability `served_subnets[:8]`, the Findings-Delta `[:60]` CHARACTER cut that manufactured
"DS03-DC" out of the front half of a hostname, …). This file pins the remaining reachable ones. The
house rule is stated in `excel._xls_cell_value`: where the module bounds a list it DISCLOSES the cut,
with a trailing `(+N)` marker or an "…and N further …" sentence — and the disclosure must name the
TRUE remainder, so shown + hidden reconciles to the real population.

html.py contributes no test: its ONLY upper-bounded slice is `ts[:10]` (an ISO timestamp -> date),
and its one display cap — the Findings-Delta devices column — was already fixed to the house pattern
in an earlier round (`_devices_cell`, covered by the existing suite).

Every cap pinned here was MEASURED against real producer artifacts before it was fixed
(`Migration_Assessment_*.snapshot.json` = the Meridian reference fleet, `webapp/sample_data/sample_fleet.snapshot.json`,
`tests/golden/snapshot.json`); the numbers quoted in each docstring are those measurements. Caps that
no real fleet reaches (`archreview` big/mid-VLAN `[:5]`, oversubscription `[:6]`, `other_vrfs[:6]`)
are deliberately NOT pinned here — they are unreachable, not broken.

Workbook assertions read the cell back from a SAVED-then-RELOADED file, never the in-memory object.
"""
import re

import pytest

from cisco_toolkit.archreview import _clip, _ev, compute_architecture_review

openpyxl = pytest.importorskip("openpyxl")
from openpyxl import Workbook, load_workbook  # noqa: E402

from cisco_toolkit.excel import write_executive_summary_sheet  # noqa: E402


# --------------------------------------------------------------------------- #
# helpers
# --------------------------------------------------------------------------- #
def _saved_sheet(wb, tmp_path, name="Executive Summary"):
    """Save the workbook and read the sheet back from DISK — the only view that proves what the
    customer opens (some truncation is visible only after the save/reload round trip)."""
    out = tmp_path / "wb.xlsx"
    wb.save(str(out))
    return load_workbook(str(out))[name]


def _cells(ws):
    return [str(c.value) for row in ws.iter_rows() for c in row if c.value is not None]


def _value_right_of(ws, key):
    """The cell to the right of a label cell (the sheet's `_kv` layout)."""
    for row in ws.iter_rows():
        for c in row:
            if str(c.value).strip() == key:
                return str(ws.cell(c.row, c.column + 1).value)
    raise AssertionError(f"label {key!r} not found in sheet")


#: the migration-brief block (and with it the "Address first" row) renders only when the brief
#: carries axes — one axis is enough to reach the row under test.
_AXES = [{"axis": "Fleet posture", "severity": "Critical", "headline": "h"}]


def _hidden_count(text):
    """The N a disclosure marker claims — `(+N)`, `(+N more …)` or `… +N more`."""
    m = re.search(r"\(\+(\d+)", text) or re.search(r"\+(\d+) more", text)
    assert m, f"no disclosure marker in {text!r}"
    return int(m.group(1))


# --------------------------------------------------------------------------- #
# excel.py — Executive Summary (the one-page landing tab)
# --------------------------------------------------------------------------- #
def test_address_first_discloses_the_gating_items_it_drops(tmp_path):
    """`" · ".join(tg[:6])` cut the executive brief's GATING list with no marker. Measured on the Meridian reference fleet: 9 top-gating headlines, 6 rendered — a reader planning around the six shown is three
    blockers short, and nothing else on the sheet carries them (the axis table below is one row per
    AXIS, not per gating item)."""
    tg = [f"gating-item-{i}" for i in range(9)]
    wb = Workbook()
    write_executive_summary_sheet(wb, [], [], [], [], brief={"axes": _AXES, "top_gating": tg})
    ws = _saved_sheet(wb, tmp_path)
    cell = _value_right_of(ws, "Address first")
    assert "gating-item-0" in cell and "gating-item-5" in cell        # the six shown
    assert "gating-item-6" not in cell                                # still capped at six
    shown = cell.split(" (+")[0].count("·") + 1
    assert shown + _hidden_count(cell) == len(tg), cell               # shown + hidden == the truth


def test_address_first_stays_unmarked_when_nothing_is_dropped(tmp_path):
    """The marker must be conditional — a cry-wolf "(+0)" on every workbook trains the reader to
    ignore it."""
    wb = Workbook()
    write_executive_summary_sheet(wb, [], [], [], [],
                                  brief={"axes": _AXES, "top_gating": ["a", "b", "c"]})
    assert "(+" not in _value_right_of(_saved_sheet(wb, tmp_path), "Address first")


def test_top_categories_discloses_the_categories_it_drops(tmp_path):
    """`sorted(...)[:5]` under the label "Top categories" named the ranking but not the population,
    so five entries read as the punch-list's whole category set. Measured: 17 categories on the Meridian reference fleet (the 12 hidden ones carry 176 of the 1,805 items), 15 on the sample fleet, 14 on the
    golden fixture — this cap is hit by EVERY artifact in the repo."""
    punchlist = [{"severity": "High", "category": f"cat{i}"} for i in range(17)]
    wb = Workbook()
    write_executive_summary_sheet(wb, [], punchlist, [], [])
    cell = _value_right_of(_saved_sheet(wb, tmp_path), "Top categories")
    assert cell.count("(1)") == 5                                     # still five entries shown
    assert 5 + _hidden_count(cell) == 17, cell                        # …of seventeen, disclosed


def test_top_categories_unmarked_when_all_categories_fit(tmp_path):
    wb = Workbook()
    write_executive_summary_sheet(wb, [], [{"severity": "High", "category": "only"}], [], [])
    assert "(+" not in _value_right_of(_saved_sheet(wb, tmp_path), "Top categories")


def test_keystone_table_header_names_the_population_it_ranks(tmp_path):
    """`fi[:10]` renders ten rows under "Keystone devices — fix-first"; on the Meridian reference fleet 193 of the
    303 simulated devices strand at least one endpoint. Ten rows with no ratio size the TABLE, not
    the problem. Rows carrying no `stranded` figure are unmeasured and must not be counted as
    keystones (coverage honesty: absent evidence is never a zero)."""
    fi = ([{"host": f"h{i}", "severity": "Critical", "stranded": 100 - i, "vlans_impacted": 3}
           for i in range(25)]
          + [{"host": f"u{i}", "severity": "Info", "stranded": None} for i in range(5)]
          + [{"host": f"z{i}", "severity": "Info", "stranded": 0} for i in range(5)])
    wb = Workbook()
    write_executive_summary_sheet(wb, [], [], [], fi)
    text = "\n".join(_cells(_saved_sheet(wb, tmp_path)))
    assert "Keystone devices" in text                                  # label preserved
    assert "top 10 of 25 device(s) that strand" in text, text[:400]     # 25 = stranded>0 only
    assert "h10" not in text                                           # table itself still 10 rows


def test_keystone_header_unmarked_when_every_keystone_is_shown(tmp_path):
    fi = [{"host": f"h{i}", "severity": "High", "stranded": 5} for i in range(4)]
    wb = Workbook()
    write_executive_summary_sheet(wb, [], [], [], fi)
    text = "\n".join(_cells(_saved_sheet(wb, tmp_path)))
    assert "Keystone devices" in text and "top 10 of" not in text


# --------------------------------------------------------------------------- #
# archreview.py — the evidence cap, whose "+N more" was itself wrong
# --------------------------------------------------------------------------- #
def _l3_snap(n=25):
    """A fleet whose HIER-1 evidence (the L3 tier) is larger than the 20-host evidence cap."""
    hosts = [f"core{i:02d}" for i in range(n)]
    return {"devices": {h: {"hostname": h} for h in hosts},
            "l3_forwarding": [{"switch": h, "vlan": 10} for h in hosts]}


def test_evidence_total_is_published_when_the_evidence_list_is_cut():
    """`add()` stores `sorted(evidence)[:20]`. Every renderer then appends "+N more" computed off
    that STORED list, so the disclosure understated the scope by everything the first cut dropped.
    Measured on the Meridian reference fleet: 8 of 25 checks exceed the cap (L2-3 230 hosts — which rendered as
    "+15 more" — LC-1 152, HIER-1 116, SEC-1 74, HIER-2 67, L2-1 42, RES-1 21, L2-5 21)."""
    res = compute_architecture_review(_l3_snap(25))
    hier1 = next(c for c in res["checks"] if c["id"] == "HIER-1")
    assert len(hier1["evidence"]) == 20                    # the cap itself is unchanged
    assert hier1["evidence_total"] == 25                   # …and the true population is published


def test_evidence_total_absent_when_nothing_was_dropped():
    """Conditional, so a snapshot that fits the cap is byte-identical to before (this is also what
    keeps tests/golden/snapshot.json unchanged — its largest evidence list is 3 hosts)."""
    res = compute_architecture_review(_l3_snap(4))
    hier1 = next(c for c in res["checks"] if c["id"] == "HIER-1")
    assert "evidence_total" not in hier1


def test_ev_marker_reconciles_to_the_true_total_not_the_sample():
    """The renderer contract: shown + hidden == the real population."""
    sample, total = [f"h{i:02d}" for i in range(20)], 116
    rendered = _ev(sample, 5, total)
    assert rendered.count(",") == 4                                    # five hosts shown
    assert _hidden_count(rendered) == total - 5, rendered
    # …and with no `total` it still behaves exactly as before (no total known == no inflation)
    assert _hidden_count(_ev(sample, 5)) == 15
    # a `total` smaller than what we hold is never trusted downward
    assert _hidden_count(_ev(sample, 5, 3)) == 15


def test_scorecard_evidence_column_reconciles_in_the_saved_docx(tmp_path):
    """Read back from the SAVED .docx: §3's Evidence column must claim the true remainder."""
    pytest.importorskip("docx")
    from docx import Document

    from cisco_toolkit.archreview import write_archreview_docx
    out = str(tmp_path / "ar.docx")
    write_archreview_docx(out, _l3_snap(25), "R6 Fleet")
    d = Document(out)
    cells = [c.text for t in d.tables for row in t.rows for c in row.cells]
    ev_cells = [c for c in cells if "core00" in c and "more" in c]
    assert ev_cells, cells[:40]
    for cell in ev_cells:
        assert _hidden_count(cell) == 25 - 5, cell     # +20, NOT the pre-fix +15


# --------------------------------------------------------------------------- #
# archreview.py — narrative list caps that named no total
# --------------------------------------------------------------------------- #
# --------------------------------------------------------------------------- #
# archreview.py — three narrative caps MEASURED as hit but whose fix is WITHHELD
#
# RES-4 / L2-1 / CAP-2 state their scope in `observed`, a string that lands in the published
# `architecture_review` snapshot section — and `tests/test_sample_fleet.py` locks that section
# against the committed demo snapshot, so disclosing the remainder is a legitimate change to
# webapp/sample_data/sample_fleet.snapshot.json (regenerate: webapp/sample_data/build_sample.py).
# The three tests below were left CHARACTERISING the undisclosed behaviour, because the pass that
# found them could not regenerate the demo snapshot. The fixes have since LANDED (the orchestrator
# owns the pin), so each now asserts the disclosure and reconciles shown + hidden against the true
# population — the form the docstrings specified.
# --------------------------------------------------------------------------- #
def test_res4_keystone_sentence_discloses_the_keystones_it_did_not_name():
    """RES-4 named 5 keystones and stopped, so the availability check read as "the fleet has five
    keystones". Measured: 193 devices strand endpoints on the Meridian reference fleet, 19 on the sample — a
    migration that hardens the five leaves the rest unhardened and sequences waves off the wrong
    blast-radius population.

    The band is read off the DESCENDING tail (``keystones[5:]``), never "at least the 5th value",
    which would overstate every hidden row."""
    fi = [{"host": f"h{i:03d}", "stranded": 500 - i} for i in range(30)]
    res = compute_architecture_review({"failure_impact": fi})
    obs = next(c for c in res["checks"] if c["id"] == "RES-4")["observed"]
    assert obs.count("strands") == 5                                   # 5 still named by name …
    assert len([r for r in fi if r["stranded"] > 0]) == 30             # … of 30 that qualify
    assert "and 25 further device(s) strand 471-495 endpoint(s) each" in obs, obs
    # the band must bound the TAIL, not the named five: 500 (the top) must never appear in it
    assert "500 endpoint(s) each" not in obs


def test_res4_keystone_sentence_is_silent_when_every_keystone_is_named():
    """Refute the fix: at or below the cap there is no tail, so the sentence must be unchanged."""
    fi = [{"host": f"h{i}", "stranded": 9 - i} for i in range(4)]
    obs = next(c for c in compute_architecture_review({"failure_impact": fi})["checks"]
               if c["id"] == "RES-4")["observed"]
    assert "further device" not in obs, obs


def test_l2_1_stp_root_sentence_discloses_the_switches_it_did_not_name():
    """L2-1 names 6 root-holding access switches; measured 42 on the Meridian reference fleet, 17 on the sample
    fleet. The sentence is the check's whole statement of SCOPE — how many closets the target design
    must re-pin, so a display cap was sizing an LLD task list."""
    hosts = [f"acc{i:02d}" for i in range(14)]
    snap = {"devices": {h: {"hostname": h} for h in hosts},
            "stp_roots": {h: {"10": {"is_root": True}} for h in hosts}}
    obs = next(c for c in compute_architecture_review(snap)["checks"]
               if c["id"] == "L2-1")["observed"]
    assert obs.count(" roots ") == 6                                   # 6 named of 14 rooted
    # 8 hidden, each rooting 1 VLAN -> shown + hidden reconciles to the real population
    assert "and 8 further access switch(es) rooting 8 VLAN(s) between them" in obs, obs


def test_cap2_tight_closet_sentence_discloses_the_closets_it_did_not_name():
    """CAP-2 names 6 closets past the 90% line; measured 16 on the sample fleet. The BoM sizes
    replacement hardware off this list, so the count past the 90% line must not be the display
    cap — ten full closets would get no spare-port allowance."""
    cap = ([{"hostname": f"acc{i:02d}", "port_util": 99.0 - i * 0.5} for i in range(11)]
           + [{"hostname": "ok01", "port_util": 40.0}])            # all 11 are past the 90% line
    obs = next(c for c in compute_architecture_review({"capacity": cap})["checks"]
               if c["id"] == "CAP-2")["observed"]
    assert obs.count("port utilisation") == 6                          # 6 named of 11 tight
    assert "and 5 further closet(s) above the 90% line" in obs, obs


def test_summary_statement_discloses_further_affected_domains():
    """"concentrated in: A, B, C, D" over 7 affected domains (Meridian reference fleet) inverts the finding: the
    deviations are fleet-wide, not concentrated."""
    hosts = [f"acc{i:02d}" for i in range(6)]
    snap = {
        "devices": {h: {"hostname": h, "num_power_supplies": 1} for h in hosts},
        "l3_forwarding": [{"switch": hosts[0], "vlan": 10}],                    # D2 (single-PSU L3)
        "stp_roots": {h: {"10": {"is_root": True}} for h in hosts[1:]},         # D3
        "operational_drift": [{"severity": "Critical", "title": "drift"}],      # D6
        "security": {h: {"summary": {"grade": "weak"}} for h in hosts},         # D7
        "lifecycle_risk": {"summary": {"n_past_eos": 2},
                           "per_device": [{"host": hosts[0], "band": "Past-EoS"}]},   # D8
    }
    res = compute_architecture_review(snap)
    worst = [d["key"] for d in res["domains"] if d["verdict"] in ("critical", "deviation")]
    assert len(worst) > 4, worst                                       # the fixture must reach it
    st = res["summary"]["statement"]
    m = re.search(r"concentrated in: (.+?) and (\d+) further domain\(s\)", st)
    assert m, st
    assert m.group(1).count(",") + 1 + int(m.group(2)) == len(worst), st


# --------------------------------------------------------------------------- #
# archreview.py — the CHARACTER slices (the class that manufactures a name)
# --------------------------------------------------------------------------- #
def test_clip_never_manufactures_a_token():
    """A raw `str(...)[:60]` ends mid-word and reads as complete text — the same defect as the
    Findings-Delta devices column, where a device prefix was the front half of a hostname. `_clip` cuts on
    a word boundary and marks it."""
    s = (
        "Temporary L2 bridge awaiting approval on "
        "MERIDIAN-DISTRIBUTION-SWITCH-192-LONGROLE pending closet consolidation"
    )
    out = _clip(s, 60)
    assert out.endswith("…") and len(out) <= 61
    assert s.startswith(out[:-1].rstrip(" ,;:-"))                      # a true prefix, nothing new
    assert not out[:-1].rstrip().endswith(
        "MERIDIAN-DISTRIBUTION-SWITCH-192-LONG"
    )  # no half identifier
    assert _clip(s, 500) == s and "…" not in _clip(s, 500)             # short values untouched
    assert _clip(None, 10) == "" and _clip(12345, 3) == "123…"         # non-str inputs are safe


def test_ops1_drift_titles_are_clipped_visibly():
    """Measured on the Meridian reference fleet: 1 of the 3 rendered High/Critical drift titles is 69 chars and was
    cut mid-word at 60."""
    long_title = "PoE fault on MERIDIAN-SW-STUDIO-062 powered endpoint dropped twice"
    snap = {"operational_drift": [{"severity": "High", "title": long_title}]}
    obs = next(c for c in compute_architecture_review(snap)["checks"]
               if c["id"] == "OPS-1")["observed"]
    assert "…" in obs, obs                                             # the cut is visible
    assert long_title[:40] in obs                                      # and it is still the title
    assert "endpoint dropped twice" not in obs


def test_priority_queue_discloses_that_it_is_only_the_top_ten(tmp_path):
    """`actions[:10]` feeds a table headed "Priority remediation queue" — which reads as the whole
    remediation scope. Measured: 19 actionable checks on the Meridian reference fleet, 14 on the sample fleet, 13 on
    the golden fixture — this cap is hit by every artifact in the repo. Disclosed in §1.1's prose so
    the published snapshot section is untouched."""
    pytest.importorskip("docx")
    from docx import Document

    from cisco_toolkit.archreview import write_archreview_docx
    checks = [{"id": f"C-{i}", "domain": "Layer-2 design", "title": f"t{i}",
               "verdict": "deviation", "observed": "o", "implication": "i",
               "recommendation": f"fix {i}", "reference": "ref", "evidence": []}
              for i in range(13)]
    snap = {"architecture_review": {
        "domains": [{"key": "Layer-2 design", "verdict": "deviation", "score_pct": 35,
                     "checks": [c["id"] for c in checks]}],
        "checks": checks,
        "top_actions": [{"rank": i + 1, "id": c["id"], "domain": c["domain"],
                         "verdict": c["verdict"], "action": c["recommendation"],
                         "evidence": []} for i, c in enumerate(checks[:10])],
        "summary": {"n_checks": 13, "statement": "s", "grade": "D", "score_pct": 35}}}
    out = str(tmp_path / "queue.docx")
    write_archreview_docx(out, snap, "R6 Fleet")
    text = "\n".join(p.text for p in Document(out).paragraphs)
    assert "Showing the top 10 of 13 check(s) requiring remediation" in text, text[:600]
    assert "remaining 3" in text


def test_priority_queue_unmarked_when_the_whole_queue_fits(tmp_path):
    pytest.importorskip("docx")
    from docx import Document

    from cisco_toolkit.archreview import write_archreview_docx
    checks = [{"id": f"C-{i}", "domain": "Layer-2 design", "title": f"t{i}",
               "verdict": "deviation", "observed": "o", "implication": "i",
               "recommendation": f"fix {i}", "reference": "ref", "evidence": []}
              for i in range(3)]
    snap = {"architecture_review": {
        "domains": [{"key": "Layer-2 design", "verdict": "deviation", "score_pct": 35,
                     "checks": [c["id"] for c in checks]}],
        "checks": checks,
        "top_actions": [{"rank": i + 1, "id": c["id"], "domain": c["domain"],
                         "verdict": c["verdict"], "action": c["recommendation"],
                         "evidence": []} for i, c in enumerate(checks)],
        "summary": {"n_checks": 3, "statement": "s", "grade": "D", "score_pct": 35}}}
    out = str(tmp_path / "queue2.docx")
    write_archreview_docx(out, snap, "R6 Fleet")
    assert "Showing the top" not in "\n".join(p.text for p in Document(out).paragraphs)


# --------------------------------------------------------------------------- #
# U1-2 — the Device Risk Register banded an EoL-Unknown asset GREEN
# --------------------------------------------------------------------------- #
def _dossiers(*, blind: bool):
    """Three healthy access switches from the REAL producer (compute_device_dossiers), differing only in
    how much evidence was collected. `blind=True` collects health + an unmatched EoX platform and nothing
    else, which is the Meridian shape: every other axis ABSTAINS ('na'), abstention weighs ZERO exposure,
    so risk_index is 0 and the engine bands the asset 'Low'."""
    from cisco_toolkit.analyze import compute_device_dossiers
    hosts = ["sw0", "sw1", "sw2"]
    hs = [{"switch": h, "band": "Good", "score": 88, "role": "access"} for h in hosts]
    if blind:
        return compute_device_dossiers(
            health_scores=hs,
            lifecycle_risk={"per_device": [{"host": h, "band": "Unknown", "model": "WS-XYZ",
                                            "platform": "cat", "sw_version": "16.9.1"} for h in hosts]})
    return compute_device_dossiers(
        health_scores=hs,
        lifecycle_risk={"per_device": [{"host": h, "band": "Active", "model": "C9300-48P",
                                        "platform": "C9300", "sw_version": "17.9.4"} for h in hosts]},
        software_risk={"per_device": [{"host": h, "band": "Current", "sw_version": "17.9.4"} for h in hosts],
                       "findings": []},
        platform_health={"per_device": [{"host": h, "collected": True, "cpu_5min": 12,
                                         "mem_free_pct": 55, "band": "Healthy"} for h in hosts],
                         "findings": []},
        syslog_intelligence={"per_device": [{"host": h, "n_events": 4} for h in hosts], "detections": []},
        qos_audit={"per_device": [{"host": h, "policies": 1} for h in hosts], "findings": []},
        golden_drift={"per_device": [{"host": h, "compliance_pct": 100, "n_missing": 0} for h in hosts],
                      "summary": {"n_baseline": 20}},
        security={h: {"findings": [{"status": "pass", "id": "x"}]} for h in hosts},
        config_hygiene={h: {"undefined_refs": [], "scanned": True, "findings": []} for h in hosts})


def test_device_risk_sheet_discloses_axes_that_were_never_assessed(tmp_path):
    """[U1-2 false-health] A fleet whose only collected evidence is a health score bands 'Low' with the
    GREEN 'Low' fill (D9EAD3) and the verdict 'No stacked risk - routine migration handling.' -- because an
    abstaining axis scores ZERO exposure. The writer already computed `n_na` per row and DROPPED it.

    Measured pre-fix (openpyxl, blind fixture): row 5 band='Low' fill=00D9EAD3, no column past 16, and the
    banner carried no coverage sentence at all -- byte-for-byte the fully-assessed rendering."""
    from cisco_toolkit.excel import DEVICE_RISK_SHEET_NAME, write_device_risk_sheet
    wb = Workbook()
    write_device_risk_sheet(wb, _dossiers(blind=True))
    ws = _saved_sheet(wb, tmp_path, DEVICE_RISK_SHEET_NAME)

    banner = str(ws.cell(1, 1).value)
    assert "NOT ASSESSED" in banner, banner
    assert "HALF OR MORE" in banner, banner
    assert "never a clean result" in banner, banner

    assert ws.cell(4, 17).value == "Not-assessed axes", "the n_na column is still dropped"
    cell = str(ws.cell(5, 17).value)
    assert "8 of 11 NOT ASSESSED" in cell, cell
    assert "absent evidence" in cell, cell

    # the band cell still carries the ENGINE's band (never re-derived here) but MUST NOT wear the
    # reassuring green fill -- the fill is the fastest read on the sheet.
    assert ws.cell(5, 6).value == "Low"
    assert (ws.cell(5, 6).fill.fgColor.rgb or "").endswith("D9D9D9"), ws.cell(5, 6).fill.fgColor.rgb


def test_device_risk_sheet_keeps_green_for_an_actually_assessed_low_asset(tmp_path):
    """NON-VACUITY for the test above: an asset whose axes really were collected keeps the Low band's
    green fill and acquires none of the coverage-gap wording. Without this the guard could be always-on
    (every row grey) and would prove nothing."""
    from cisco_toolkit.excel import DEVICE_RISK_SHEET_NAME, write_device_risk_sheet
    wb = Workbook()
    write_device_risk_sheet(wb, _dossiers(blind=False))
    ws = _saved_sheet(wb, tmp_path, DEVICE_RISK_SHEET_NAME)

    assert ws.cell(5, 6).value == "Low"
    assert (ws.cell(5, 6).fill.fgColor.rgb or "").endswith("D9EAD3"), ws.cell(5, 6).fill.fgColor.rgb
    assert "absent evidence" not in str(ws.cell(5, 17).value)
    banner = str(ws.cell(1, 1).value)
    assert "HALF OR MORE" not in banner, banner
    # the measured axis coverage is still stated -- a count of 0 that means "not measured" must never
    # read as "nothing wrong", so the denominator is published either way.
    assert "risk axes fleet-wide were NOT ASSESSED" in banner, banner


def test_device_risk_sheet_fails_closed_when_a_row_carries_no_axis_census(tmp_path):
    """[r8 F1] `dossier_coverage` fell back to `n_na` when a row published no `exposures` list -- and
    then returned `bool(n_axes and ...)`, which is FALSE for n_axes == 0. So the branch written for a
    census-less dossier concluded "not thin", i.e. fully assessed, and the row kept the Low band's
    reassuring green fill: absence rendered as health inside the guard that closes absence-rendered-as-
    health. With no denominator, how much of the band rests on evidence is itself NOT ASSESSED.

    Measured pre-fix on the fixture below: row 5 fill=00D9EAD3 (green), column Q = "0 not assessed",
    and the banner carried no census sentence."""
    from cisco_toolkit.excel import DEVICE_RISK_SHEET_NAME, write_device_risk_sheet
    dd = _dossiers(blind=False)                       # a genuinely assessed fleet from the REAL producer
    for row in dd["per_device"]:
        row.pop("exposures", None)                    # ... whose axis census did not survive to the writer
    wb = Workbook()
    write_device_risk_sheet(wb, dd)
    ws = _saved_sheet(wb, tmp_path, DEVICE_RISK_SHEET_NAME)

    assert ws.cell(5, 6).value == "Low", "the ENGINE's band must still be shown verbatim"
    assert (ws.cell(5, 6).fill.fgColor.rgb or "").endswith("D9D9D9"), (
        "a row with NO axis census kept the Low band's green fill — the fastest read on the sheet "
        f"claims evidence that was never published: {ws.cell(5, 6).fill.fgColor.rgb}")
    cell = str(ws.cell(5, 17).value)
    assert "census ABSENT" in cell and "NOT ASSESSED" in cell, cell
    assert "absent evidence" in cell, cell
    banner = str(ws.cell(1, 1).value)
    assert "published NO risk-axis census at all" in banner, banner
    assert "never" in banner and "clean" in banner, banner


# --------------------------------------------------------------------------- #
# service-map name authority (cross-lane contract: portdb authority labels)
# --------------------------------------------------------------------------- #
# The service-map rows come from the REAL producer chain, end to end:
#     device running-config text -> parse.parse_acls -> analyze.compute_service_map (-> portdb)
#     -> excel.write_service_map_sheet
# The previous revision of these tests hand-built the `sm` dict in the shape the writer expects, which is
# the fixture-agrees-with-the-bug trap: it asserted `semantics_authoritative: True` for 22/tcp ssh, a
# combination portdb NEVER emits (ssh carries a curated overlay, so its SEMANTICS are not authoritative).
# The fixture was therefore proving a rendering of a registry state that does not exist.
_REAL_ACL_TEXT = """ip access-list extended AV-CONTROL
 permit udp 10.10.10.0 0.0.0.255 any eq 4440
 permit tcp 10.10.10.0 0.0.0.255 host 10.20.20.5 eq 21
 deny   ip any any log
!
"""


def _real_service_map():
    """compute_service_map() over ACLs the real parser extracted from real running-config text.

    4440/udp resolves to the curated-overlay-only "Dante-audio" (assignment_authoritative False) and
    21/tcp to the IANA-assigned "ftp" with no overlay (both authority flags True) -- both facts come
    from portdb, not from this test."""
    from cisco_toolkit.analyze import compute_service_map
    from cisco_toolkit.parse import parse_acls
    acls = parse_acls(_REAL_ACL_TEXT)
    assert acls.get("AV-CONTROL"), f"the real ACL parser produced nothing: {acls!r}"
    # acl_hits upgrades the TRAFFIC evidence to "Confirmed" -- the exact pairing the disclaimer exists
    # for: confirmed traffic on a port whose NAME is only a curated hypothesis.
    return compute_service_map({"sw1": acls}, {"sw1": {}}, acl_hits={"4440:udp": 9})


def test_service_map_labels_a_curated_only_service_name(tmp_path):
    """A curated-overlay-only classification (4440/udp 'Dante-audio', assignment_authoritative=False)
    reached this sheet indistinguishable from an IANA assignment, and `evidence_class` -- which speaks only
    to TRAFFIC evidence -- then read as if it confirmed the NAME. Rendered defensively: the column exists
    only when the producer publishes the labels."""
    from cisco_toolkit.excel import SERVICE_MAP_SHEET_NAME, write_service_map_sheet
    sm = _real_service_map()
    rows = {(s["port"], s["proto"]): s for s in sm["services"]}
    # the producer really does hand the writer these two authority states (fixture non-vacuity)
    assert rows[(4440, "udp")]["assignment_authoritative"] is False, rows[(4440, "udp")]
    assert rows[(21, "tcp")]["assignment_authoritative"] is True, rows[(21, "tcp")]
    assert rows[(21, "tcp")]["semantics_authoritative"] is True, rows[(21, "tcp")]

    wb = Workbook()
    write_service_map_sheet(wb, sm)
    ws = _saved_sheet(wb, tmp_path, SERVICE_MAP_SHEET_NAME)
    assert ws.cell(1, 9).value == "Name authority"
    by_port = {ws.cell(r, 1).value: r for r in (2, 3)}
    assert set(by_port) == {21, 4440}, [ws.cell(r, 1).value for r in (2, 3)]

    dante, ftp = by_port[4440], by_port[21]
    # the row whose TRAFFIC evidence was upgraded to Confirmed is exactly the row whose NAME is not
    # an assignment of record -- which is the whole point of the column.
    assert "Confirmed" in str(ws.cell(dante, 8).value), ws.cell(dante, 8).value
    assert "NOT authoritative" in str(ws.cell(dante, 9).value), ws.cell(dante, 9).value
    # NON-VACUITY: the fully-authoritative row must NOT acquire the disclaimer.
    assert "NOT authoritative" not in str(ws.cell(ftp, 9).value), ws.cell(ftp, 9).value
    assert str(ws.cell(ftp, 9).value).startswith("authoritative"), ws.cell(ftp, 9).value


def test_service_map_makes_no_authority_claim_when_the_producer_publishes_none(tmp_path):
    """NON-VACUITY / back-compat: an older snapshot carries no authority labels. Absence of the labels is
    not evidence of authority -- the column is simply not rendered, and nothing is claimed either way.

    This fixture is hand-built ON PURPOSE: it is a LEGACY on-disk shape that the current producer can no
    longer emit. That claim is asserted rather than assumed -- if the producer ever stopped publishing
    the labels, this test would be guarding today's output, not yesterday's."""
    from cisco_toolkit.excel import (SERVICE_MAP_SHEET_NAME, _SERVICE_AUTHORITY_KEYS,
                                     write_service_map_sheet)
    live = _real_service_map()["services"]
    assert live and all(_SERVICE_AUTHORITY_KEYS & set(s) for s in live), (
        "the CURRENT producer no longer publishes authority labels — the 'older snapshot' fixture below "
        "is now the live shape, and this test is guarding the wrong thing")

    sm = {"services": [{"port": 67, "proto": "udp", "service": "DHCP-server", "category": "Infra",
                        "broadcast": False, "refs": 4, "host_count": 2,
                        "evidence_class": "Inferred (ACL design intent)"}], "multicast": {}}
    wb = Workbook()
    write_service_map_sheet(wb, sm)
    ws = _saved_sheet(wb, tmp_path, SERVICE_MAP_SHEET_NAME)
    assert ws.cell(1, 9).value is None
    assert ws.cell(2, 9).value is None


# --------------------------------------------------------------------------- #
# R8 / on-air renderers: the multicast "on-air" classification is CURATED
# --------------------------------------------------------------------------- #
# analyze.compute_multicast_intelligence derives `on_air` entirely from the offline registry's
# curated broadcast/category fields, and then PROMOTES a MAC-alias clash to High on the strength of
# it. The producer publishes the basis (`on_air_authoritative` / `has_av_authoritative` /
# `n_av_groups_authoritative` / `severity_basis`); the workbook rendered a bare "yes", a bare
# category and a bare "N broadcast/AV", so a curated hypothesis reached the customer in the same
# voice as an observed measurement.


def _real_multicast(groups):
    """compute_service_map -> compute_multicast_intelligence over REAL registry lookups.

    Nothing here is hand-shaped: portdb decides that 224.0.1.129 is the curated-only "PTP-primary"
    (Broadcast-AV, semantics_authoritative False) and that 225.0.1.129 / 239.70.70.70 match no
    record at all. The Vlan30 interface carries real `multicast_info` so the multicast block is
    non-empty the way a collected fabric makes it."""
    from cisco_toolkit.analyze import compute_multicast_intelligence, compute_service_map
    from cisco_toolkit.model import InterfaceData
    svi = InterfaceData(port="Vlan30")
    svi.multicast_info = "PIM sparse-mode"
    ifaces = {"sw1": {"Vlan30": svi}}
    sm = compute_service_map({}, ifaces, igmp_groups=list(groups))
    return sm, compute_multicast_intelligence(sm, ifaces)


def _mcast_sheet(mi, tmp_path):
    from cisco_toolkit.excel import MULTICAST_INTEL_SHEET_NAME, write_multicast_intelligence_sheet
    wb = Workbook()
    write_multicast_intelligence_sheet(wb, mi)
    return _saved_sheet(wb, tmp_path, MULTICAST_INTEL_SHEET_NAME)


def _row_starting(ws, first_cell):
    for row in ws.iter_rows():
        if str(row[0].value or "") == first_cell:
            return row[0].row
    raise AssertionError(f"no row starting {first_cell!r}")


def test_multicast_sheet_discloses_that_the_on_air_labels_are_curated(tmp_path):
    """The AV headline, the MAC-alias High and the group census all rested on a curated registry
    classification and said nothing about it. Each must now carry its basis at the point of use."""
    sm, mi = _real_multicast(["224.0.1.129", "225.0.1.129", "239.70.70.70"])
    # fixture non-vacuity: the PRODUCER really does report a High promoted by a non-authoritative label
    assert mi["summary"]["n_av_groups"] == 1 and mi["summary"]["n_av_groups_authoritative"] == 0
    assert mi["mac_aliases"][0]["has_av"] is True
    assert mi["mac_aliases"][0]["has_av_authoritative"] is False
    ws = _mcast_sheet(mi, tmp_path)

    headline = str(ws.cell(1, 2).value)
    assert "ALL curated, NOT an authoritative source" in headline, headline
    assert "2 unclassified (no registry match)" in headline, headline

    alias = _row_starting(ws, "01:00:5e:00:01:81")
    assert str(ws.cell(alias, 3).value) == "yes (CURATED, unverified)", ws.cell(alias, 3).value
    assert str(ws.cell(alias, 4).value) == "High", ws.cell(alias, 4).value    # NOT re-scored
    assert "NOT an authoritative source" in str(ws.cell(alias, 5).value), ws.cell(alias, 5).value
    assert "not in question" in str(ws.cell(alias, 5).value), ws.cell(alias, 5).value

    av = _row_starting(ws, "224.0.1.129")
    assert "CURATED, NOT authoritative" in str(ws.cell(av, 3).value), ws.cell(av, 3).value
    assert str(ws.cell(av, 4).value) == "yes (CURATED, unverified)", ws.cell(av, 4).value
    unk = _row_starting(ws, "239.70.70.70")
    # a group with NO registry match is a THIRD state -- not a quieter kind of curated
    assert "UNCLASSIFIED" in str(ws.cell(unk, 3).value), ws.cell(unk, 3).value
    assert str(ws.cell(unk, 4).value or "") == "", ws.cell(unk, 4).value
    # the sheet's row-1 width is pinned by tests/golden/sheet_schema.json -- 6 columns, unchanged
    assert ws.max_column == 6


def test_multicast_sheet_clean_case_acquires_no_curated_caveat(tmp_path):
    """NON-VACUITY: a clash with NO on-air member stays Medium and must NOT pick up the curated-AV
    disclosure, or the disclosure is always-on and carries no information."""
    sm, mi = _real_multicast(["239.70.70.70", "238.70.70.70"])
    assert mi["summary"]["n_av_groups"] == 0 and mi["mac_aliases"][0]["has_av"] is False
    ws = _mcast_sheet(mi, tmp_path)
    headline = str(ws.cell(1, 2).value)
    assert "(0 broadcast/AV)" in headline, headline
    assert "curated" not in headline.lower(), headline
    alias = _row_starting(ws, "01:00:5e:46:46:46")
    assert str(ws.cell(alias, 3).value or "") == "", ws.cell(alias, 3).value
    assert str(ws.cell(alias, 4).value) == "Medium", ws.cell(alias, 4).value
    assert "NOT an authoritative source" not in str(ws.cell(alias, 5).value), ws.cell(alias, 5).value


def test_multicast_sheet_says_registry_verified_when_the_registry_really_vouches(tmp_path):
    """NON-VACUITY (second axis): the tag must be a FUNCTION of the producer's labels, not a constant
    "curated". No group in the shipped pack is semantics-authoritative, so the only way to exercise
    the branch is to flip the producer's own labels on the producer's own record."""
    sm, mi = _real_multicast(["224.0.1.129", "225.0.1.129"])
    for g in mi["groups"]:
        if g["group"] == "224.0.1.129":
            g["assignment_authoritative"] = True
            g["semantics_authoritative"] = True
            g["on_air_authoritative"] = True
    mi["mac_aliases"][0]["has_av_authoritative"] = True
    mi["summary"]["n_av_groups_authoritative"] = 1
    ws = _mcast_sheet(mi, tmp_path)
    headline = str(ws.cell(1, 2).value)
    assert "1 registry-authoritative" in headline, headline
    assert "ALL curated" not in headline, headline
    av = _row_starting(ws, "224.0.1.129")
    assert "registry-authoritative" in str(ws.cell(av, 3).value), ws.cell(av, 3).value
    assert str(ws.cell(av, 4).value) == "yes (registry-verified)", ws.cell(av, 4).value
    alias = _row_starting(ws, "01:00:5e:00:01:81")
    assert str(ws.cell(alias, 3).value) == "yes (registry-verified)", ws.cell(alias, 3).value


def test_multicast_sheet_fails_closed_when_the_producer_published_no_basis(tmp_path):
    """COVERAGE-HONEST: an older snapshot carries no basis fields. Absence must read "NOT published",
    never as verified -- and never as a bare High that looks measured.

    The "older" shape is derived by DELETING the fields from real producer output, so it cannot drift
    into a shape the producer never emitted."""
    sm, mi = _real_multicast(["224.0.1.129", "225.0.1.129"])
    for g in mi["groups"]:
        g.pop("on_air_authoritative", None)
    mi["mac_aliases"][0].pop("has_av_authoritative", None)
    mi["summary"].pop("n_av_groups_authoritative", None)
    mi.pop("risks", None)
    ws = _mcast_sheet(mi, tmp_path)
    assert "classification basis NOT published by this snapshot" in str(ws.cell(1, 2).value)
    alias = _row_starting(ws, "01:00:5e:00:01:81")
    assert str(ws.cell(alias, 3).value) == "yes — basis NOT published", ws.cell(alias, 3).value
    assert "severity basis NOT published" in str(ws.cell(alias, 5).value), ws.cell(alias, 5).value
    av = _row_starting(ws, "224.0.1.129")
    assert str(ws.cell(av, 4).value) == "yes — basis NOT published", ws.cell(av, 4).value


def test_service_map_sheet_labels_the_multicast_group_classification(tmp_path):
    """The Service Map sheet printed `224.0.1.129 = PTP-primary (Broadcast-AV)` with nothing to say
    the name and category are a curated hypothesis -- the same defect the service rows were fixed for."""
    from cisco_toolkit.excel import SERVICE_MAP_SHEET_NAME, write_service_map_sheet
    sm, _mi = _real_multicast(["224.0.1.129", "239.70.70.70"])
    wb = Workbook()
    write_service_map_sheet(wb, sm)
    ws = _saved_sheet(wb, tmp_path, SERVICE_MAP_SHEET_NAME)
    lines = [str(c.value) for row in ws.iter_rows() for c in row
             if c.value is not None and str(c.value).startswith(("224.", "239."))]
    assert len(lines) == 2, lines
    curated = next(x for x in lines if x.startswith("224.0.1.129"))
    unmatched = next(x for x in lines if x.startswith("239.70.70.70"))
    assert "PTP-primary (Broadcast-AV)" in curated and "CURATED, NOT authoritative" in curated, curated
    assert "UNCLASSIFIED" in unmatched, unmatched
    # NON-VACUITY / back-compat: a record carrying no authority labels gets NO claim either way.
    live = sm["multicast"]["classified_groups"]
    assert live and all("overlay_status" in g for g in live), (
        "the CURRENT producer no longer publishes authority labels -- the legacy fixture below is now "
        "the live shape, and this assertion is guarding the wrong thing")
    legacy = {"multicast": {"classified_groups": [{"group": "224.0.1.129", "name": "PTP-primary",
                                                   "category": "Broadcast-AV"}]}, "services": []}
    wb2 = Workbook()
    write_service_map_sheet(wb2, legacy)
    (tmp_path / "legacy").mkdir()
    ws2 = _saved_sheet(wb2, tmp_path / "legacy", SERVICE_MAP_SHEET_NAME)
    legacy_line = next(str(c.value) for row in ws2.iter_rows() for c in row
                       if c.value is not None and str(c.value).startswith("224."))
    assert legacy_line == "224.0.1.129 = PTP-primary (Broadcast-AV)", legacy_line


def test_an_incoherent_authority_census_is_disclosed_not_turned_into_a_negative_count():
    """`n_av_groups` and `n_av_groups_authoritative` are published independently, so their
    subtraction is only meaningful when authoritative <= total. Unchecked, the Multicast Intelligence
    sheet rendered `{n_av - n_av_auth} curated/unverified` and printed a NEGATIVE count -- a number
    that cannot exist -- in a client workbook. Deriving a third figure from an incoherent pair and
    presenting it as a measurement is the failure; saying the split cannot be stated is the fix.
    """
    from openpyxl import Workbook
    from cisco_toolkit.excel import write_multicast_intelligence_sheet

    def line(summary):
        wb = Workbook(); wb.remove(wb.active)
        write_multicast_intelligence_sheet(
            wb, {"summary": summary, "groups": [], "mac_aliases": [], "risks": []})
        ws = wb[wb.sheetnames[0]]
        return next((c.value for r in ws.iter_rows() for c in r
                     if isinstance(c.value, str) and "broadcast/AV" in c.value), "")

    bad = line({"n_groups": 9, "n_av_groups": 3, "n_av_groups_authoritative": 7})
    assert "INCOHERENT" in bad, bad
    assert "-4" not in bad and "-" not in bad.split("broadcast/AV")[1][:20], \
        f"a negative derived count reached the workbook: {bad!r}"

    # NON-VACUITY: a coherent census must still state the real split, or the disclosure is always-on
    # and tells the reader nothing.
    good = line({"n_groups": 9, "n_av_groups": 5, "n_av_groups_authoritative": 2})
    assert "2 registry-authoritative" in good and "3 curated/unverified" in good, good
    assert "INCOHERENT" not in good


# ------------------------------------- the Punch-List severity must carry the basis that raised it
# review r10 EXIT X. compute_migration_punchlist's media fold puts the producer's `severity_basis` /
# `evidence_confidence` on the ITEM, and NO punch-list renderer read either key -- a contract with no
# consumer. The workbook's Migration Punch-List sheet is the one where it bites hardest: a High raised
# purely by a CURATED, explicitly non-authoritative on-air classification sat in the same orange band,
# in the same "Severity" column, as a High raised by an observed measurement.


def _punchlist_with_media(media_risks, **kw):
    """The fleet punch-list from the REAL producer, every other fold empty except the one the caller
    passes, so a note below cannot be satisfied by an unrelated row."""
    from cisco_toolkit.analyze import compute_migration_punchlist
    return compute_migration_punchlist([], {}, {}, [], [], [], {}, [], [],
                                       media_risks=media_risks, **kw)


def _real_media_risks(groups):
    """Media risks from the REAL producer chain, filtered the way COLLECT_PARSE_V3_23_0.py filters
    them before the fold. Nothing hand-shaped: portdb decides 224.0.1.129 is the curated-only on-air
    group, and compute_multicast_intelligence is what raises the clash to High because of it."""
    from cisco_toolkit.analyze import compute_multicast_intelligence, compute_service_map
    sm = compute_service_map({}, {}, igmp_groups=groups)
    mi = compute_multicast_intelligence(sm, {})
    return [r for r in mi["risks"] if r.get("kind") in ("mac-alias", "querier-gap")]


def _punch_sheet(punchlist, tmp_path):
    from cisco_toolkit.excel import PUNCHLIST_SHEET_NAME, write_punchlist_sheet
    wb = Workbook()
    write_punchlist_sheet(wb, punchlist)
    return _saved_sheet(wb, tmp_path, PUNCHLIST_SHEET_NAME)


def _punch_row(ws, category):
    for r in range(2, ws.max_row + 1):
        if str(ws.cell(r, 3).value or "") == category:
            return r
    raise AssertionError(f"no punch-list row in category {category!r}")


_ORDINARY_ROW = [{"severity": "High", "category": "False-health", "devices": ["sw1"],
                  "title": "Interface counters frozen", "detail": "d", "remediation": "r"}]


def test_punchlist_severity_cell_carries_the_basis_that_raised_it(tmp_path):
    """The Severity cell of a media High must carry the producer's own basis as a note, read back
    from a SAVED file -- a comment that does not survive serialization is no consumer at all."""
    risks = _real_media_risks(["224.0.1.129", "225.0.1.129"])
    # fixture non-vacuity: the PRODUCER really did raise this to High on a non-authoritative label
    src = next(r for r in risks if r["kind"] == "mac-alias")
    assert src["severity"] == "High" and "NOT an authoritative source" in src["severity_basis"]

    ws = _punch_sheet(_punchlist_with_media(risks, drift=_ORDINARY_ROW), tmp_path)
    media = _punch_row(ws, "Multicast/Media")
    sev = ws.cell(media, 2)
    assert sev.value == "High"                       # NOT re-scored -- the fix is disclosure
    assert sev.comment is not None, "the media High still carries no basis at the point of claim"
    note = sev.comment.text
    # the producer's own basis + confidence travel verbatim, un-paraphrased
    assert src["severity_basis"] in note, note
    assert src["evidence_confidence"] in note, note
    assert "NOT an authoritative source" in note and "curated, unverified" in note, note

    # NON-VACUITY: an ordinary punch-list row with no media basis acquires NO note whatsoever.
    ordinary = _punch_row(ws, "False-health")
    assert ws.cell(ordinary, 2).value == "High"      # same severity, so the note is the only signal
    assert ws.cell(ordinary, 2).comment is None, ws.cell(ordinary, 2).comment

    # and this is a NOTE, not a schema change: the golden-pinned header row is untouched.
    assert [c.value for c in ws[1]] == ["#", "Severity", "Category", "Device(s)", "Wave",
                                        "Issue", "Detail", "Remediation"]
    assert ws.max_column == 8


def test_punchlist_severity_note_distinguishes_an_unpublished_basis(tmp_path):
    """COVERAGE-HONEST: a risk whose basis the snapshot never published must be DISTINGUISHABLE from
    one carrying a real basis -- not silently identical, and not dressed as authoritative.

    The "older snapshot" shape is derived by DELETING the fields from real producer output, so it
    cannot drift into a shape the producer never emitted."""
    risks = _real_media_risks(["224.0.1.129", "225.0.1.129"])
    for r in risks:
        r.pop("severity_basis", None)
        r.pop("evidence_confidence", None)
    ws = _punch_sheet(_punchlist_with_media(risks), tmp_path)
    sev = ws.cell(_punch_row(ws, "Multicast/Media"), 2)
    assert sev.value == "High"
    assert sev.comment is not None, "an unpublished basis must still be DISCLOSED, not go silent"
    note = sev.comment.text
    assert "NOT published by this snapshot" in note, note
    # ...and it must not read like the curated-but-published row above, nor like a verified one
    assert "NOT an authoritative source" not in note, note
    assert "curated, unverified" not in note, note


def test_punchlist_severity_note_keys_on_a_usable_value_not_a_present_key(tmp_path):
    """FAIL CLOSED. `"severity_basis" in item` is TRUE for null / blank / non-string values, and the
    key-presence shape would render an EMPTY note -- which a reader takes for "nothing to disclose".
    An unusable value must be treated exactly like a missing one: no note at all."""
    from cisco_toolkit.excel import punchlist_severity_note
    for n, bad in enumerate((None, "", "   ", {"a": 1}, 0, [], False)):
        item = {"severity": "High", "severity_basis": bad, "evidence_confidence": bad}
        assert "severity_basis" in item                       # the fail-OPEN test would pass here
        assert punchlist_severity_note(item) == "", (bad, punchlist_severity_note(item))
        # ...and the sheet writes NO comment object rather than an empty one
        sub = tmp_path / f"case{n}"
        sub.mkdir()
        ws = _punch_sheet([dict(item, category="Multicast/Media", title="t",
                                detail="d", remediation="r")], sub)
        assert ws.cell(2, 2).comment is None, (bad, ws.cell(2, 2).comment)

    # half a disclosure is not a disclosure: a usable basis with an unusable confidence says so.
    half = punchlist_severity_note({"severity_basis": "B-TOKEN", "evidence_confidence": "  "})
    assert "B-TOKEN" in half and "Evidence: NOT published by this snapshot" in half, half
    # NON-VACUITY: a usable confidence is rendered instead of that sentinel.
    whole = punchlist_severity_note({"severity_basis": "B-TOKEN", "evidence_confidence": "C-TOKEN"})
    assert "Evidence: C-TOKEN" in whole and "NOT published" not in whole, whole


def test_punchlist_severity_note_is_bounded_against_a_pathological_snapshot(tmp_path):
    """A `--no-collect` re-analysis feeds this writer an untrusted snapshot section, so the note is
    length-bounded -- and the bound must clip visibly, never manufacture a truncated token."""
    from cisco_toolkit.excel import punchlist_severity_note
    note = punchlist_severity_note({"severity_basis": "HEAD-TOKEN " + "wordy " * 4000,
                                    "evidence_confidence": "C-TOKEN"})
    assert "HEAD-TOKEN" in note and "…" in note, note[:80]
    assert len(note) < 2200, len(note)
    assert "Evidence: C-TOKEN" in note, "the bound must not swallow the second half of the note"


def test_the_severity_note_box_is_sized_to_its_text_so_nothing_is_silently_clipped():
    """The note is bounded at `_PUNCH_NOTE_MAX` PER HALF, so it can reach ~1,800 characters, while the
    comment box was fixed at 460x190 -- roughly 850 visible. More than half the basis was invisible in
    a client workbook.

    That is worse than never writing the note: the reader sees a comment marker, opens it, and
    believes they have the whole reason for a High severity when they have the first half of it.
    """
    from openpyxl import Workbook
    from cisco_toolkit.excel import write_punchlist_sheet, punchlist_severity_note

    def box_for(basis):
        item = {"severity": "High", "category": "Multicast/Media", "title": "t", "detail": "d",
                "remediation": "r", "devices": ["sw1"], "severity_basis": basis,
                "evidence_confidence": "observed overlap = fact"}
        wb = Workbook(); wb.remove(wb.active)
        write_punchlist_sheet(wb, [item])
        ws = wb[wb.sheetnames[0]]
        c = next((cell.comment for row in ws.iter_rows() for cell in row if cell.comment), None)
        return len(punchlist_severity_note(item) or ""), c

    # a pathologically long basis must still fit: capacity is width/7 chars per line x height/14 lines
    n_long, c_long = box_for("wordy " * 190)
    assert c_long is not None
    capacity = int((c_long.width / 7) * (c_long.height / 14))
    assert capacity >= n_long, (
        f"{n_long - capacity} characters of the basis fall outside the {c_long.width}x{c_long.height} "
        "comment box and are invisible to the reader")

    # NON-VACUITY, both directions: a typical basis must NOT inflate the box (the note stays compact),
    # and an item with no usable basis must produce no comment at all rather than an empty box.
    n_short, c_short = box_for("classified Broadcast-AV / on-air; CURATED registry semantics.")
    assert c_short.height == 190, f"a short basis grew the box to {c_short.height}"
    assert int((c_short.width / 7) * (c_short.height / 14)) >= n_short
    assert box_for("")[1] is None, "an item with no basis produced a comment box anyway"
