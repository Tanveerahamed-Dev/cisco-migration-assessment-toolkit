"""Round-4 whole-repo review of the three highest-consequence deliverable writers:
`cisco_toolkit/mop.py` (the change steps + ROLLBACK an engineer executes at 2am on production),
`cisco_toolkit/runbook.py` (the operations handbook) and `cisco_toolkit/archreview.py` (the
architecture grade that justifies spend).

Every fixture here is derived from a REAL producer artifact — `tests/golden/snapshot.json` or
`webapp/sample_data/sample_fleet.snapshot.json` — by DELETING or blanking the section under test,
never by hand-building a dict in the shape the code expects (a hand-made stub agrees with the
writer's own bugs). Neither pinned artifact is modified.

The defects pinned:
  * MOP  — the pure make-before-break procedure decommissioned the legacy path BEFORE the §x.6
    go/no-go ran, while §x.7's rollback asserted the legacy path "was never torn down"; that branch
    also named no restore-from-§x.3 path at all.
  * MOP  — an absent failure-impact axis rendered "Max blast radius … 0", and that fabricated 0 is
    the threshold in the quantified "roll back if more endpoints than the §x.1 figure are affected"
    trigger, silently rewriting it to "roll back if ANY endpoint is affected".
  * MOP  — §x.2 "Blockers to clear before this window" display-capped at 10 punch-list + 8
    remediation items with no disclosure, while the precondition gate cites the FULL counts and
    points the engineer at that section (80 attributed, 10 shown, on the real sample fleet).
  * RUNBOOK — §1's gating headline and §5's confidence line counted absent sections as 0, so a
    snapshot missing cross-layer / link-centrality / readiness read as a complete all-clear.
  * RUNBOOK — §9's move-group table rendered a group with no wave-sequencing record as
    "0 make-before-break / 0 hard cutover / 0 at-risk endpoints", i.e. as needing no outage.
  * ARCHREVIEW — CAP-2 and RES-4 graded CONFORMS (raising the overall grade) off evidence that was
    never captured: capacity rows with no port-utilisation figure, and a failure-impact section that
    is empty / carries no `stranded` figure.
"""
import copy
import json
import os

import pytest

docx = pytest.importorskip("docx")  # skip the file if the optional dep is absent
from docx import Document  # noqa: E402

from cisco_toolkit.archreview import compute_architecture_review  # noqa: E402
from cisco_toolkit.mop import write_mop_docx  # noqa: E402
from cisco_toolkit.runbook import write_runbook_docx  # noqa: E402

_HERE = os.path.dirname(os.path.abspath(__file__))
_REPO = os.path.dirname(_HERE)
GOLDEN = os.path.join(_HERE, "golden", "snapshot.json")
SAMPLE = os.path.join(_REPO, "webapp", "sample_data", "sample_fleet.snapshot.json")


def _load(path):
    """Read a pinned producer artifact. Returns a deep copy so a test can delete sections from it
    without touching the file or leaking into the next test."""
    with open(path, encoding="utf-8") as fh:
        return json.load(fh)


@pytest.fixture(scope="module")
def sample():
    if not os.path.exists(SAMPLE):
        pytest.skip("sample_fleet.snapshot.json not present")
    return _load(SAMPLE)


@pytest.fixture(scope="module")
def golden():
    if not os.path.exists(GOLDEN):
        pytest.skip("golden snapshot not present")
    return _load(GOLDEN)


def _text(path):
    """All visible text: one line per paragraph, then one line per TABLE ROW with its cells joined
    by ' | ' (row-shaped, because several assertions below are about what a single row says)."""
    d = Document(path)
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            parts.append(" | ".join(c.text.replace("\n", " ") for c in row.cells))
    return "\n".join(parts)


def _mop_text(tmp_path, snap, name="mop.docx"):
    out = str(tmp_path / name)
    write_mop_docx(out, snap, "R4 Review Fleet")
    return _text(out)


def _runbook_text(tmp_path, snap, name="rb.docx"):
    out = str(tmp_path / name)
    write_runbook_docx(out, snap, "R4 Review Fleet")
    return _text(out)


def _steps_of(text, heading_prefix):
    """The numbered step lines of one MOP subsection (steps() renders '<n>. <action>')."""
    lines = text.split("\n")
    try:
        start = next(i for i, ln in enumerate(lines) if ln.startswith(heading_prefix))
    except StopIteration:  # pragma: no cover - guards a fixture change, not a code path
        raise AssertionError(f"section {heading_prefix!r} not rendered:\n{text[:2000]}")
    out = []
    for ln in lines[start + 1:]:
        if ln[:1].isdigit() and ". " in ln[:5]:
            out.append(ln)
        elif out and ln and not ln.startswith(" "):
            break
    return out


# ---------------------------------------------------------------------------------------------
# MOP · rollback completeness + step ordering (the single most dangerous class in this repo)
# ---------------------------------------------------------------------------------------------
def _pure_mbb_snap(sample):
    """The sample fleet's Group 2 is 100% dual-homed (a REAL producer wave_sequencing record with an
    empty hard_cutover list), so dropping design_blueprint — which otherwise bundles both groups into
    one mixed wave — yields a genuine PURE make-before-break MOP section."""
    s = copy.deepcopy(sample)
    s.pop("design_blueprint", None)
    return s


def test_mop_pure_mbb_decommissions_legacy_only_after_the_go_no_go_gate(tmp_path, sample):
    """ORDERING. §x.5 used to emit 'decommission the legacy path' as step 6 of 8, two steps BEFORE
    'Run §x.6 post-cutover validation in full'. §x.6 is the go/no-go whose failure triggers the
    rollback — so the procedure destroyed the fallback before the decision that needs it. The
    decommission step must come strictly AFTER the §x.6 step."""
    text = _mop_text(tmp_path, _pure_mbb_snap(sample), "mop_order.docx")
    steps = _steps_of(text, "4.5 Cutover procedure")
    assert steps, "no §4.5 procedure steps rendered"
    decom = next((i for i, s in enumerate(steps) if "decommission the legacy path" in s), None)
    gate = next((i for i, s in enumerate(steps) if "post-cutover validation in full" in s), None)
    assert decom is not None and gate is not None, steps
    assert gate < decom, (
        "the legacy path is decommissioned BEFORE the §4.6 go/no-go gate — the rollback's fallback "
        f"is gone by the time the decision is taken:\n" + "\n".join(steps))
    # and the step must say so, not merely be positioned there
    assert "ONLY once" in steps[decom] and "4.6" in steps[decom]


def test_mop_pure_mbb_phase_tags_are_monotonic(tmp_path, sample):
    """The [PRE]/[DURING]/[POST] tags exist so the window structure reads at a glance. The misplaced
    decommission step made them run PRE, PRE, DURING, DURING, DURING, POST, DURING, POST — a [DURING]
    step after a [POST] one. Ordering the procedure correctly makes the tags monotonic again."""
    text = _mop_text(tmp_path, _pure_mbb_snap(sample), "mop_phase.docx")
    rank = {"[PRE]": 0, "[DURING]": 1, "[POST]": 2}
    seen = [rank[t] for s in _steps_of(text, "4.5 Cutover procedure")
            for t in rank if t in s]
    assert seen == sorted(seen), f"phase tags out of order: {seen}"


def test_mop_pure_mbb_rollback_names_a_restore_path(tmp_path, sample):
    """ROLLBACK COMPLETENESS. On the pure make-before-break branch the whole rollback was 'move the
    endpoints back onto the still-live legacy leg … the legacy path was never torn down, this is
    non-disruptive'. If the trigger fires after the final forward step that claim is false and the
    branch offered NO way back — it never referenced the §x.3 configuration backup that §x.3 itself
    declares mandatory 'for the §x.7 rollback'. Every rollback must name a restore path."""
    text = _mop_text(tmp_path, _pure_mbb_snap(sample), "mop_rb.docx")
    rb = "\n".join(_steps_of(text, "4.7 Rollback"))
    assert "re-apply the pre-change configuration captured in §4.3" in rb.lower() \
        or "re-apply the pre-change configuration captured in §4.3" in rb, \
        "the pure make-before-break rollback never restores the §4.3 baseline:\n" + rb
    assert "ALREADY been decommissioned" in rb, \
        "the rollback does not handle the legacy path already being gone:\n" + rb
    # the unconditional falsehood must be gone
    assert "the legacy path was never torn down" not in rb


# ---------------------------------------------------------------------------------------------
# MOP · a rollback threshold sized on an axis that was never collected
# ---------------------------------------------------------------------------------------------
def test_mop_absent_blast_radius_is_not_rendered_as_zero(tmp_path, sample):
    """COVERAGE HONESTY. `max(..., default=0)` over an absent failure_impact rendered
    'Max blast radius … 0' — a measured-looking figure for an axis nobody collected. Delete the one
    section from the real sample fleet and the document must abstain."""
    s = copy.deepcopy(sample)
    s.pop("failure_impact", None)
    text = _mop_text(tmp_path, s, "mop_noblast.docx")
    row = next(ln for ln in text.split("\n") if ln.startswith("Max blast radius"))
    assert "[NOT OBSERVED]" in row, row
    assert not row.endswith("| 0"), "an uncollected blast radius still renders as a measured 0: " + row


def test_mop_blast_radius_rollback_trigger_withdraws_its_absent_threshold(tmp_path, sample):
    """The quantified trigger reads 'roll back if … more endpoints than the §x.1 max-blast-radius
    figure are affected'. With that figure fabricated as 0 the trigger becomes 'roll back if ANY
    endpoint is affected' — unconditionally true in a window whose purpose is moving endpoints, so
    the in-window engineer is handed a rollback condition that is already satisfied. When the figure
    is absent the endpoint clause must be withdrawn, not evaluated against zero."""
    s = copy.deepcopy(sample)
    s.pop("failure_impact", None)
    text = _mop_text(tmp_path, s, "mop_trigger.docx")
    row = next(ln for ln in text.split("\n") if ln.startswith("Blast-radius / outage overrun"))
    assert "more endpoints than the" not in row, \
        "the trigger still sizes on a max-blast-radius figure that was never observed: " + row
    assert "[NOT OBSERVED]" in row and "NOT a threshold of zero" in row, row


def test_mop_observed_blast_radius_still_renders_the_quantified_trigger(tmp_path, sample):
    """Refute-the-fix: the guard must be EVIDENCE-based, not a blanket withdrawal. The unmodified
    sample fleet carries a real failure_impact simulation, so the figure and the quantified endpoint
    clause must both survive."""
    text = _mop_text(tmp_path, sample, "mop_blast_ok.docx")
    row = next(ln for ln in text.split("\n") if ln.startswith("Max blast radius"))
    assert "[NOT OBSERVED]" not in row and row.rstrip().split("|")[-1].strip().isdigit(), row
    trig = next(ln for ln in text.split("\n") if ln.startswith("Blast-radius / outage overrun"))
    assert "more endpoints than the" in trig, trig


# ---------------------------------------------------------------------------------------------
# MOP · the precondition gate must not point at a silently truncated list
# ---------------------------------------------------------------------------------------------
def test_mop_blocker_section_discloses_its_display_cap(tmp_path, sample):
    """The pre-implementation checklist states the FULL attributed blocker counts and says 'clear or
    risk-accept before proceeding (see §blockers)'. §x.2 display-caps the punch list at 10 rows and
    the remediation snippets at 8 — on the real sample fleet that is 10 of 80 shown, with no
    disclosure. An engineer who clears everything printed there opens the window on 70 uncleared
    Critical/High blockers believing the gate is met."""
    text = _mop_text(tmp_path, sample, "mop_blockers.docx")
    gate = next(ln for ln in text.split("\n") if "risk-accept before proceeding" in ln)
    n_pl = int(gate.split("readiness")[1].split(";")[1].split("Critical/High punch-list")[0].strip())
    assert n_pl > 10, f"fixture no longer discriminates (only {n_pl} punch-list blockers)"
    assert f"…and {n_pl - 10} further Critical/High punch-list item(s)" in text, \
        (f"the gate cites {n_pl} blockers but §x.2 shows 10 with no disclosure of the "
         f"{n_pl - 10} it dropped")
    assert f"({n_pl} in total)" in text


# ---------------------------------------------------------------------------------------------
# RUNBOOK · absent analyses must not read as an all-clear
# ---------------------------------------------------------------------------------------------
def test_runbook_gating_headline_abstains_on_absent_analyses(tmp_path, sample):
    """§1's 'Top gating decisions' is the one line a war-room reader takes away. Every clause is a
    count over a snapshot section, and an absent section counts 0 — so deleting the three analyses
    used to render '0 Critical cross-layer single-point(s) of failure and 0 bridge link(s) must be
    resolved … 0 of 0 move group(s) are NOT READY': a complete all-clear manufactured from missing
    evidence, which is exactly what guardrail 3 forbids."""
    s = copy.deepcopy(sample)
    for k in ("cross_layer", "link_centrality", "migration_readiness"):
        s.pop(k, None)
    text = _runbook_text(tmp_path, s, "rb_noaxes.docx")
    gate = next(ln for ln in text.split("\n") if ln.startswith("Top gating decisions:"))
    assert "0 Critical cross-layer" not in gate and "0 bridge link(s)" not in gate, gate
    assert "0 of 0 move group(s)" not in gate, gate
    assert gate.count("[NOT OBSERVED") >= 2, gate
    assert "readiness was NOT computed" in gate, gate
    # the §1 register rows and §5's confidence line carry the same absent axes
    assert "Bridge links (single points of fabric partition)\n[NOT OBSERVED]" in text \
        or "Bridge links (single points of fabric partition) | [NOT OBSERVED]" in text \
        or any(ln.startswith("Bridge links") and "[NOT OBSERVED]" in ln for ln in text.split("\n")), \
        "the §1 metric register still prints a measured 0 bridge links"
    assert "Inferred-high. 0 inter-switch links" not in text


def test_runbook_gating_headline_unchanged_when_the_analyses_are_present(tmp_path, sample):
    """Refute-the-fix: on the unmodified sample fleet the counted sentence must survive verbatim."""
    text = _runbook_text(tmp_path, sample, "rb_ok.docx")
    gate = next(ln for ln in text.split("\n") if ln.startswith("Top gating decisions:"))
    assert "[NOT OBSERVED" not in gate, gate
    assert "Critical cross-layer single-point(s) of failure" in gate
    assert "move group(s)\nare NOT READY" in gate or "are NOT READY" in gate


def test_runbook_move_group_table_abstains_without_a_sequencing_record(tmp_path, sample):
    """§9 introduces the table with 'single-homed switches are hard cutovers (maintenance window;
    their endpoints take an outage)'. `ws_by_group.get(g, {})` turned a group with NO sequencing
    record into '0 make-before-break / 0 hard cutover / 0 at-risk endpoints' — i.e. a 19-switch
    group reading as needing no window and no outage."""
    s = copy.deepcopy(sample)
    s.pop("wave_sequencing", None)
    text = _runbook_text(tmp_path, s, "rb_nows.docx")
    rows = [ln for ln in text.split("\n") if ln.startswith("Group 1 |") and "NOT READY" in ln]
    assert rows, "the §9 move-group row did not render"
    assert "| 0 | 0 | 0" not in rows[0], \
        "an unsequenced group still reports zero hard cutovers and zero at-risk endpoints: " + rows[0]
    assert rows[0].count("[NOT OBSERVED]") == 3, rows[0]
    assert "no wave-sequencing record" in text and "NOT zero" in text


def test_runbook_move_group_table_keeps_observed_sequencing(tmp_path, sample):
    """Refute-the-fix: the unmodified fleet's Group 1 really is 2 make-before-break / 17 hard
    cutover / 45 at-risk endpoints and must still render those figures."""
    text = _runbook_text(tmp_path, sample, "rb_ws_ok.docx")
    row = next(ln for ln in text.split("\n") if ln.startswith("Group 1 |") and "NOT READY" in ln)
    assert "[NOT OBSERVED]" not in row, row
    assert row.rstrip().endswith("| 17 | 45"), row


# ---------------------------------------------------------------------------------------------
# ARCHREVIEW · a conformance GRADE must never be earned by silence
# ---------------------------------------------------------------------------------------------
def _verdict(snap, cid):
    ar = compute_architecture_review(snap)
    return next(c for c in ar["checks"] if c["id"] == cid)


def test_archreview_cap2_not_assessable_without_a_captured_port_utilisation(sample):
    """CAP-2 filtered non-numeric port_util rows away and then graded CONFORMS off whatever
    survived — so a capacity section whose utilisation column was never captured asserted 'All N
    device(s) hold port headroom below the 90% line' over N unmeasured closets. `port_util` is a
    known absent-prone field (runbook.py excludes rows whose active-port count was never observed
    for exactly this reason)."""
    s = copy.deepcopy(sample)
    for r in s["capacity"]:
        r["port_util"] = ""
    c = _verdict(s, "CAP-2")
    assert c["verdict"] == "not-assessable", c
    assert "no numeric port-utilisation figure" in c["observed"]


def test_archreview_cap2_not_assessable_on_partial_port_utilisation(sample):
    """Partial coverage is a blind spot, never health — the rule RES-3 already applies to PSU
    inventory. One unmeasured closet must not be carried to CONFORMS by its measured siblings."""
    s = copy.deepcopy(sample)
    s["capacity"][0]["port_util"] = ""
    for r in s["capacity"][1:]:
        r["port_util"] = 50.0
    assert _verdict(s, "CAP-2")["verdict"] == "not-assessable"


@pytest.mark.parametrize("mutate,label", [
    (lambda s: s.__setitem__("failure_impact", []), "present but empty"),
    (lambda s: s.__setitem__("failure_impact", "oops"), "truthy non-list (upload poison)"),
    (lambda s: [r.pop("stranded", None) for r in s["failure_impact"]], "rows carry no stranded figure"),
])
def test_archreview_res4_not_assessable_without_simulation_evidence(sample, mutate, label):
    """RES-4's guard was `snap.get("failure_impact") is None`, which catches ONLY a literally absent
    key. Every other way the evidence can be missing fell through to CONFORMS — 'the failure
    simulation strands no endpoints behind any single device / redundancy absorbs any one device
    loss' — and RAISED the conformance grade, on a fleet whose real simulation strands 45 endpoints
    behind core1."""
    s = copy.deepcopy(sample)
    mutate(s)
    assert _verdict(s, "RES-4")["verdict"] == "not-assessable", label


def test_archreview_res4_still_conforms_on_an_observed_zero(sample):
    """Refute-the-fix: `stranded: 0` is a REAL observed zero. The guard keys on the figure being
    absent, not on it being falsy, so a genuinely redundant fleet must still earn CONFORMS."""
    s = copy.deepcopy(sample)
    for r in s["failure_impact"]:
        r["stranded"] = 0
    c = _verdict(s, "RES-4")
    assert c["verdict"] == "conforms", c
    assert "simulated device(s) report a stranded-endpoint figure" in c["observed"]


def test_archreview_ops2_not_assessable_without_an_undefined_reference_count(sample):
    """Third instance of the same class in this file: `_as_int(<missing>)` coerced a hygiene record
    with no `summary.undefined` figure to 0, so a fleet whose hygiene analysis reported nothing
    graded CONFORMS — 'No undefined references across the fleet' — over silence. On the unmodified
    sample core1 really has 2 undefined references (an ACL and a route-map the configs reference but
    never define)."""
    s = copy.deepcopy(sample)
    for rec in s["config_hygiene"].values():
        rec.get("summary", {}).pop("undefined", None)
        rec.pop("undefined", None)
    c = _verdict(s, "OPS-2")
    assert c["verdict"] == "not-assessable", c
    assert _verdict(copy.deepcopy(sample), "OPS-2")["verdict"] == "advisory"  # real evidence survives


def test_archreview_ops2_falls_back_to_the_undefined_list(sample):
    """The abstention must not swallow evidence that IS present: dropping only the SUMMARY count
    must not lose a real finding, because the record's own `undefined` list still evidences it. The
    check stays ADVISORY rather than abstaining (an over-eager guard would hide core1's two
    undefined references, which is the opposite failure)."""
    s = copy.deepcopy(sample)
    for rec in s["config_hygiene"].values():
        rec.get("summary", {}).pop("undefined", None)
    c = _verdict(s, "OPS-2")
    assert c["verdict"] == "advisory", c
    assert "core1" in c["observed"]


def test_archreview_absent_evidence_never_raises_the_grade(sample):
    """The load-bearing property behind both guards: deleting evidence must never IMPROVE the
    architecture grade. Before the fix, blanking port_util took the sample fleet from 70% to 71%
    and emptying failure_impact did the same — a customer-facing grade, and the spend it justifies,
    moving in the wrong direction because a collection got shorter."""
    base = compute_architecture_review(copy.deepcopy(sample))["summary"]["score_pct"]
    for mutate in (lambda s: [r.__setitem__("port_util", "") for r in s["capacity"]],
                   lambda s: s.__setitem__("failure_impact", []),
                   lambda s: [r.pop("stranded", None) for r in s["failure_impact"]],
                   lambda s: [(r.get("summary", {}).pop("undefined", None), r.pop("undefined", None))
                              for r in s["config_hygiene"].values()]):
        s = copy.deepcopy(sample)
        mutate(s)
        got = compute_architecture_review(s)["summary"]["score_pct"]
        assert got <= base, f"deleting evidence raised the grade {base}% -> {got}%"


def test_archreview_golden_snapshot_grade_unchanged_by_the_guards(golden):
    """The golden snapshot is a fully-evidenced producer artifact: both guarded checks must keep
    reading their real evidence, so nothing here regrades a well-formed assessment."""
    ar = compute_architecture_review(copy.deepcopy(golden))
    ck = {c["id"]: c for c in ar["checks"]}
    assert ck["CAP-2"]["verdict"] != "not-assessable"
    assert ck["RES-4"]["verdict"] != "not-assessable"
    assert ck["OPS-2"]["verdict"] != "not-assessable"


# ---------------------------------------------------------------------------------------------
# whole-document smoke: the guards must not break a well-formed render
# ---------------------------------------------------------------------------------------------
@pytest.mark.parametrize("writer", ["mop", "runbook"])
def test_writers_still_render_the_golden_snapshot(tmp_path, golden, writer):
    text = (_mop_text(tmp_path, golden, "g_mop.docx") if writer == "mop"
            else _runbook_text(tmp_path, golden, "g_rb.docx"))
    assert len(text) > 2000
    assert "Traceback" not in text
