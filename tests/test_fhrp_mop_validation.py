"""FHRP validation rows are per-VLAN gates even when one command captures them all."""
import os

import pytest


pytest.importorskip("docx")
from docx import Document  # noqa: E402

from cisco_toolkit.mop import _join_group_records, write_mop_docx  # noqa: E402


def _fhrp_check(vlan):
    return {
        "device": "distA",
        "platform": "ios",
        "wave": "Group 1",
        "category": "FHRP",
        "subtype": "HSRP",
        "vlan": vlan,
        "severity": "High",
        "check": f"First-hop redundancy healthy for VLAN {vlan}",
        "command": "show standby brief",
        "expect": f"VLAN {vlan} shows exactly one Active + one Standby",
        "why": f"VLAN {vlan} must retain first-hop failover.",
    }


def _snapshot(checks=None):
    checks = [_fhrp_check(10), _fhrp_check(20)] if checks is None else checks
    return {
        "script_version": "V3.23.0",
        "devices": {"distA": {"platform": "ios"}},
        "move_groups": [{"switches": ["distA"], "endpoints": 8}],
        "migration_readiness": [{
            "group": "Group 1",
            "switches": ["distA"],
            "endpoints": 8,
            "readiness": "READY",
            "n_fail": 0,
            "n_warn": 0,
            "checks": [],
        }],
        "wave_sequencing": [{
            "group": "Group 1",
            "make_before_break": ["distA"],
            "hard_cutover": [],
            "sequence": "make-before-break",
        }],
        "validation_plan": {
            "items": checks,
            "by_wave": {"Group 1": checks},
            "summary": {"n_items": len(checks), "n_waves": 1, "n_high": len(checks),
                        "by_category": {"FHRP": len(checks)}},
        },
    }


def _tables_with_header(doc, headers):
    expected = list(headers)
    return [
        table for table in doc.tables
        if table.rows and [cell.text for cell in table.rows[0].cells] == expected
    ]


def test_join_retains_semantic_checks_but_dedupes_an_exact_duplicate():
    vlan10 = _fhrp_check(10)
    vlan20 = _fhrp_check(20)
    validation = {"Group 1": [vlan10, dict(vlan10), vlan20, "malformed-row"]}

    joined = _join_group_records(
        ["Group 1"], ["distA"], {}, {}, {}, validation)[3]

    assert [row["vlan"] for row in joined] == [10, 20]
    assert [row["check"] for row in joined] == [
        "First-hop redundancy healthy for VLAN 10",
        "First-hop redundancy healthy for VLAN 20",
    ]


def test_docx_keeps_two_vlan_gates_and_one_shared_baseline_capture(tmp_path):
    output = tmp_path / "fhrp-mop.docx"
    write_mop_docx(str(output), _snapshot(), "FHRP Retention Fixture")

    assert os.path.isfile(output) and output.stat().st_size > 0
    doc = Document(str(output))
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert "3.3 Pre-cutover baseline capture" in headings
    assert "3.6 Post-cutover validation (go/no-go)" in headings

    captures = _tables_with_header(doc, ["Device", "Command to capture"])
    assert len(captures) == 1
    capture_rows = [[cell.text for cell in row.cells] for row in captures[0].rows[1:]]
    assert capture_rows == [["distA", "show standby brief"]]

    validation = _tables_with_header(
        doc, ["Sev", "Category", "Device", "Check", "Command", "Observed baseline / acceptance"])
    assert len(validation) == 1
    rows = [[cell.text for cell in row.cells] for row in validation[0].rows[1:]]
    assert [row[3] for row in rows] == [
        "First-hop redundancy healthy for VLAN 10",
        "First-hop redundancy healthy for VLAN 20",
    ]
    assert [row[4] for row in rows] == ["show standby brief", "show standby brief"]
    assert [row[5] for row in rows] == [
        "VLAN 10 shows exactly one Active + one Standby",
        "VLAN 20 shows exactly one Active + one Standby",
    ]
    document_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "matching that degraded state after cutover is not success" in document_text


def test_docx_keeps_every_fhrp_vlan_gate_beyond_general_display_cap(tmp_path):
    checks = [_fhrp_check(vlan) for vlan in range(1, 43)]
    output = tmp_path / "fhrp-mop-over-cap.docx"
    write_mop_docx(str(output), _snapshot(checks), "FHRP Over-Cap Fixture")

    doc = Document(str(output))
    validation = _tables_with_header(
        doc, ["Sev", "Category", "Device", "Check", "Command", "Observed baseline / acceptance"])
    assert len(validation) == 1
    rows = [[cell.text for cell in row.cells] for row in validation[0].rows[1:]]
    assert len(rows) == 42
    assert "First-hop redundancy healthy for VLAN 1" in {row[3] for row in rows}
    assert "First-hop redundancy healthy for VLAN 42" in {row[3] for row in rows}

    captures = _tables_with_header(doc, ["Device", "Command to capture"])
    assert len(captures) == 1
    capture_rows = [[cell.text for cell in row.cells] for row in captures[0].rows[1:]]
    assert capture_rows == [["distA", "show standby brief"]]
