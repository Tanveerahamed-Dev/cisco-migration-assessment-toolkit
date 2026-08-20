"""NEW-V3.23.156: the Customer Requirements Document (CRD, DOCX) — the Plan-phase
requirements-capture instrument. python-docx is optional, so the module is skipped when it is
absent (the generator fails soft the same way). These tests pin the skeleton sections, the
evidence-primed environment summary, the discipline gating (no section without observed footprint),
the REQ-ID traceability skeleton, and the fail-soft path."""
import json
import os
import sys
from pathlib import Path

import pytest

docx = pytest.importorskip("docx")  # skip the file if the optional dep is absent
from docx import Document  # noqa: E402

from cisco_toolkit.crd import write_crd_docx  # noqa: E402


def _snap():
    """A compact snapshot exercising every evidence-gated CRD section."""
    return {
        "script_version": "V3.23.0",
        "devices": {"core1": {"model": "N9K-C93180YC"}, "acc1": {"model": "WS-C2960X-48FPD-L"}},
        "interfaces": {
            "acc1": {
                "Gi1/0/5": {"switchport_mode": "Access", "vlan": "20",
                            "end_host_mac": "bbbb.0000.0001", "dual_connection": "acc2:Gi1/0/5"},
                "Gi1/0/6": {"switchport_mode": "Access", "vlan": "30", "end_host_mac": "cccc.0000.0001"},
                "Gi1/0/1": {"switchport_mode": "Trunk", "cdp_neighbor": "core1"},
            },
            "core1": {"Vlan20": {"svi_ip": "10.0.20.1", "vrf": "PROD", "acl_in": "PROD_IN"}},
        },
        "routing_neighbors": {"core1": {"ospf": [{"neighbor": "10.0.0.2"}], "eigrp": [], "bgp": []}},
        "l3_forwarding": [{"switch": "core1", "vlan": "20", "svi_ip": "10.0.20.1", "fhrp": "HSRP active"}],
        "service_map": {
            "services": [{"service": "rtsp", "category": "CCTV", "port": "554", "proto": "tcp"}],
            "multicast": {"active_interfaces": 2, "active_switch_count": 1,
                          "classified_groups": [{"group": "224.0.1.129", "name": "PTP-primary"}],
                          "igmp_queriers": [], "ptp": {}},
        },
        "lifecycle_risk": {"summary": {"n_devices": 2, "n_past_eos": 1}},
        "collection_completeness": {"summary": {"complete": 2, "partial": 0, "not_collected": 0}},
        "punchlist": [{"severity": "Critical", "category": "L3 design",
                       "title": "VLAN 30 has a single gateway", "devices": ["core1"]}],
        # Two genuinely dual-homed endpoints (host MAC on two switches) — the canonical
        # redundancy-bearing count. Note the per-port dual_connection tally above is only 1, so a
        # CRD reading endpoint_dependencies (2) is provably distinct from the old per-port count (1).
        "endpoint_dependencies": {"dual_homed": [
            {"mac": "bbbb.0000.0001", "switches": ["acc1", "acc2"]},
            {"mac": "dddd.0000.0002", "switches": ["acc1", "core1"]},
        ]},
    }


def _all_text(doc):
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            parts.extend(c.text for c in row.cells)
    return "\n".join(parts)


def test_crd_has_skeleton_sections_and_furniture(tmp_path):
    out = str(tmp_path / "c.docx")
    write_crd_docx(out, _snap(), "Unit Test Fleet")
    d = Document(out)
    h1 = [p.text for p in d.paragraphs if p.style.name == "Heading 1"]
    for token in ("1. Engagement Context", "2. Current Environment Summary (evidence)",
                  "3. Business Requirements", "4. Technical Requirements",
                  "5. Operational Requirements", "6. Future Plans & Growth",
                  "7. Constraints & Assumptions", "8. Out of Scope",
                  "9. Requirements Traceability Matrix (RTM)"):
        assert any(t == token for t in h1), f"missing section: {token}; have {h1}"
    assert "Document Control" in h1 and "Document Acceptance" in h1
    text = _all_text(d)
    assert "Customer network owner" in text          # acceptance roles
    assert "Assessment workbook (.xlsx)" in text     # related-documents cross-reference


def test_crd_environment_summary_reconciles_to_evidence(tmp_path):
    out = str(tmp_path / "c.docx")
    write_crd_docx(out, _snap(), "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "OSPF" in text                                       # protocols observed
    assert "PROD" in text                                       # VRF observed
    assert "VLAN 30 has a single gateway" in text               # punch-list carried as known issue
    # endpoints = access ports with a host MAC (2); dual-homed is the canonical
    # endpoint_dependencies.dual_homed count, surfaced as its own honest stat line
    assert "Dual-homed endpoints (host MAC on two switches)" in text


def test_crd_dual_homed_reads_canonical_endpoint_dependencies(tmp_path):
    """A1 (SSOT): the CRD's dual-homed figure must be the canonical
    endpoint_dependencies.dual_homed count (host MAC on two switches) — the SAME source the HLD's
    preserve-dual-homed-endpoints design decision reads — not a looser per-port dual_connection
    tally. The fixture has 1 dual_connection port but 2 canonical dual-homed endpoints, so this
    discriminates: a regression to the per-port count would render 1, not 2."""
    snap = _snap()
    assert len(snap["endpoint_dependencies"]["dual_homed"]) == 2          # canonical truth
    out = str(tmp_path / "c.docx")
    write_crd_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "Dual-homed endpoints (host MAC on two switches)" in text       # honest label
    assert "for the 2 dual-homed endpoint" in text                         # REQ-T-LAN-002 cites 2, not 1
    assert "endpoint port(s)" not in text                                  # old misleading phrasing gone


def test_crd_future_plans_anchors_to_observed_scale(tmp_path):
    """N3: §6 Future Plans must anchor growth to the OBSERVED scale baseline (devices / VLANs /
    endpoints, canonical-first), so a growth target is measured against reality, not a blank prompt."""
    snap = _snap()
    snap["executive_brief"] = {"scale": {"n_devices": 2, "n_vlans": 5, "n_endpoints": 4242}}
    out = str(tmp_path / "c.docx")
    write_crd_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "Today (observed)" in text        # evidence-anchored scale baseline table
    assert "4242" in text                     # canonical n_endpoints, not a blank prompt


def test_crd_vlan_count_reads_published_scale_canonical_first(tmp_path):
    """SSOT: the 'VLANs in use' count must read the published executive_brief.scale.n_vlans (the ONE
    source the workbook / design / explorer / webapp reconcile to) canonical-first, not independently
    recompute vlan_inventory — mirroring the endpoint line one row below (n_endpoints) and design.py's
    'A5 SSOT fix' canonical-first read. Mutation-proof: pin a published scale.n_vlans that DISAGREES with
    vlan_inventory and assert the CRD renders the published value, so the read path cannot silently drift."""
    snap = _snap()
    # vlan_inventory(_snap()) is a small set; pin a disagreeing published canonical scalar (777).
    snap["executive_brief"] = {"scale": {"n_devices": 2, "n_vlans": 777, "n_endpoints": 4242}}
    out = str(tmp_path / "c.docx")
    write_crd_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "777" in text, "CRD must render the published scale.n_vlans (777), not a local vlan_inventory recount"


def test_crd_evidence_scope_says_collected_not_inventory_scope(tmp_path):
    """Coverage-honesty (inventoried vs collected): n_devices is the full INVENTORY (devices.json), not
    the collected count. The title page + §1 evidence-scope must phrase it 'C of N devices collected'
    (C = collection_completeness.complete), never the old 'N devices in evidence scope' which presented
    the inventory as if every device had been collected."""
    snap = _snap()
    snap["collection_completeness"] = {"summary": {"complete": 1, "inventory": 99}}
    out = str(tmp_path / "c.docx")
    write_crd_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "devices collected" in text and "1 of" in text     # collected framing + the collected count
    assert "devices in evidence scope" not in text            # the old inventory-as-scope wording is gone


def test_crd_requirement_tables_declare_verification_method(tmp_path):
    """N5: every requirement-capture table declares HOW each requirement is proven (a Verification
    column), so a requirement is testable — the technical rows reference the NRFU acceptance test."""
    out = str(tmp_path / "c.docx")
    write_crd_docx(out, _snap(), "Unit Test Fleet")
    d = Document(out)
    req_tabs = [t for t in d.tables if t.rows and any(c.text == "Verification" for c in t.rows[0].cells)]
    assert req_tabs, "no requirement table with a Verification column"
    text = _all_text(d)
    assert "NRFU technical acceptance test" in text     # REQ-T rows declare their verification method


def test_crd_has_rfc2119_classification_legend(tmp_path):
    """N1/N8: the CRD declares the BCP 14 (RFC 2119 / RFC 8174) normative-keyword convention, and the
    requirement tables classify strength with MUST/SHOULD/MAY rather than an ad-hoc H/M/L."""
    out = str(tmp_path / "c.docx")
    write_crd_docx(out, _snap(), "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "BCP 14" in text and "RFC 2119" in text       # classification legend cites the standard
    assert "Class (RFC 2119)" in text                     # requirement tables use the RFC-2119 class column
    assert "<MUST/SHOULD/MAY>" in text                    # the normative-keyword placeholder, not <H/M/L>


def test_crd_l3_no_fhrp_text_says_svi_instances_not_vlans(tmp_path):
    """A4: l3_forwarding row count (n_l3) is SVI INSTANCES, not distinct gateway VLANs. The
    no-FHRP REQ-T-L3-001 wording must not call the SVI-instance count 'gateway VLAN(s)'."""
    snap = _snap()
    # all-"none" FHRP so the else-branch (the mislabeled one) renders
    snap["l3_forwarding"] = [{"switch": "core1", "vlan": "20", "svi_ip": "10.0.20.1", "fhrp": "none"},
                             {"switch": "core1", "vlan": "30", "svi_ip": "10.0.30.1", "fhrp": "none"}]
    out = str(tmp_path / "c.docx")
    write_crd_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "gateway SVI instance(s)" in text
    assert "gateway VLAN(s) — every gateway is single-homed" not in text    # old mislabel gone


def test_crd_requirements_are_proposals_with_req_ids_and_traceability(tmp_path):
    out = str(tmp_path / "c.docx")
    write_crd_docx(out, _snap(), "Unit Test Fleet")
    text = _all_text(Document(out))
    # capture discipline: REQ-IDs, owners, testability guidance, confirm-or-strike framing
    for rid in ("REQ-B-001", "REQ-T-LAN-001", "REQ-T-L3-001", "REQ-T-MC-001",
                "REQ-T-SEC-001", "REQ-T-SVC-001", "REQ-O-001"):
        assert rid in text, rid
    assert "testable statement" in text.lower()
    assert "<owner>" in text and "<YES/AMEND>" in text
    # every seeded REQ-ID lands in the RTM as HONEST forward placeholders (never fabricated references)
    assert "to be traced in HLD" in text and "to be traced in NRFU" in text
    # the evidence-primed framing is explicit (proposals, not assumed requirements)
    assert "proposals" in text.lower() or "proposal" in text.lower()


def test_crd_discipline_sections_are_evidence_gated(tmp_path):
    """A pure-L2 fleet with no multicast / VRF / service evidence gets no §4.2-§4.5."""
    snap = _snap()
    snap["routing_neighbors"] = {}
    snap["l3_forwarding"] = []
    snap["service_map"] = {}
    snap["interfaces"]["core1"] = {}   # drop the VRF/ACL SVI
    out = str(tmp_path / "c.docx")
    write_crd_docx(out, snap, "Unit Test Fleet")
    d = Document(out)
    h2 = [p.text for p in d.paragraphs if p.style.name == "Heading 2"]
    assert any("4.1 Campus LAN" in t for t in h2)               # always present
    for absent in ("4.2 Layer 3", "4.3 Multicast", "4.4 Segmentation", "4.5 Services"):
        assert not any(absent in t for t in h2), absent
    text = _all_text(d)
    assert "none (pure L2 fleet)" in text


def test_crd_failsoft_without_python_docx(monkeypatch, tmp_path):
    import builtins, os
    real_import = builtins.__import__

    def _blocked(name, *a, **k):
        if name == "docx" or name.startswith("docx."):
            raise ImportError("simulated missing python-docx")
        return real_import(name, *a, **k)

    monkeypatch.setattr(builtins, "__import__", _blocked)
    out = str(tmp_path / "c.docx")
    write_crd_docx(out, _snap(), "Unit Test Fleet")   # must not raise
    assert not os.path.exists(out)


def test_crd_vlan_count_is_canonical_single_source(tmp_path):
    """Single source of truth: the CRD VLAN count must equal the canonical vlan_inventory (all
    distinct VLAN IDs, incl. SVI/L3 gateways), so it agrees with the design doc — not just the
    access-port VLANs (which omit L3-only VLANs)."""
    from cisco_toolkit.analyze import vlan_inventory
    from cisco_toolkit.crd import _evidence_facts
    snap = {
        "script_version": "V3.23.0",
        "devices": {"core1": {"model": "N9K"}, "acc1": {"model": "C2960"}},
        "interfaces": {
            "acc1": {"Gi1/0/1": {"switchport_mode": "Access", "vlan": "10",
                                 "end_host_mac": "aaaa.0000.0001"}},
            "core1": {"Vlan99": {"svi_ip": "10.0.99.1"}},   # SVI carries no 'vlan' field
        },
        # VLAN 50 has an L3 gateway but sits on NO access port -> only the canonical inventory sees it
        "l3_forwarding": [{"switch": "core1", "vlan": "50", "svi_ip": "10.0.50.1"},
                          {"switch": "core1", "vlan": "99", "svi_ip": "10.0.99.1"}],
    }
    canonical = {v for v, _ in vlan_inventory(snap)}
    assert canonical == {10, 50, 99}                        # access 10 + L3 50/99
    assert _evidence_facts(snap)["n_vlans"] == len(canonical) == 3


def test_crd_fhrp_count_excludes_none_sentinel_no_false_redundancy():
    """False-health guard (the 'none'-is-truthy bug class): l3_forwarding rows carry fhrp='none' when
    no FHRP is configured, and 'none'.strip() is truthy — so the CRD was counting EVERY gateway VLAN as
    'FHRP-protected', a customer-facing claim of redundancy that does not exist and that contradicts the
    workbook's FHRP Consistency sheet (0). Mirror the engine's canonical (fhrp or 'none') != 'none'
    check so only a real HSRP/VRRP/GLBP token counts."""
    from cisco_toolkit.crd import _evidence_facts
    snap = {"l3_forwarding": [
        {"switch": "c1", "vlan": "10", "fhrp": "none"},         # no FHRP -> must NOT count
        {"switch": "c1", "vlan": "20", "fhrp": ""},             # blank -> must NOT count
        {"switch": "c1", "vlan": "30", "fhrp": "HSRP active"},  # real FHRP -> counts
        {"switch": "c2", "vlan": "30", "fhrp": "HSRP standby"}, # same VLAN, real -> dedups to one
    ]}
    assert _evidence_facts(snap)["fhrp_vlans"] == ["30"]
    # a fleet with zero real FHRP must report an EMPTY set (not every gateway VLAN)
    none_snap = {"l3_forwarding": [{"switch": "c1", "vlan": str(v), "fhrp": "none"} for v in (10, 20, 30)]}
    assert _evidence_facts(none_snap)["fhrp_vlans"] == []


def test_crd_gains_design_driven_requirements_section(tmp_path):
    """A snapshot carrying the canonical design_blueprint renders §8 — the target-state design decisions
    read as requirement candidates (REQ-D-…), each citing a CCDE principle, plus the open design
    questions. It is data-gated: the other CRD tests (no design_blueprint) never see §8, proving it is
    never empty filler. One source of truth: the SAME design_blueprint behind the HLD/LLD and dashboards."""
    snap = _snap()
    snap["design_blueprint"] = {
        "summary": {"n_recommended": 1, "n_needs_requirement": 1, "n_critical": 1},
        "decisions": [
            {"title": "Introduce first-hop redundancy", "status": "recommended", "priority": "Critical",
             "recommended_action": "Deploy HSRP/VRRP across dual gateways", "driver": "Gateway resilience",
             "evidence": {"summary": "52 VLAN(s) without FHRP"}, "principle": {"citation": "CCDE In Depth — HA"}},
            {"title": "Right-size availability per tier", "status": "needs-requirement", "priority": "High",
             "evidence": {"summary": ""}, "principle": {"citation": "CCDE"},
             "requirements_needed": ["availability_tier"]},
        ],
        "coverage": {"caveat": "grounded only in collected evidence"},
    }
    out = str(tmp_path / "c.docx")
    write_crd_docx(out, snap, "Unit Test Fleet")
    d = Document(out)
    h1 = [p.text for p in d.paragraphs if p.style.name == "Heading 1"]
    assert any("10. Design-Driven Requirements" in t for t in h1), f"missing §10; have {h1}"
    text = _all_text(d)
    assert "REQ-D-001" in text
    assert "Introduce first-hop redundancy" in text and "CCDE In Depth — HA" in text
    assert "Right-size availability per tier" in text          # open design question surfaced
    # gated off when there is no blueprint (the other tests rely on this)
    d2 = Document(out)  # sanity: §10 only here
    assert any("10. Design-Driven Requirements" in p.text for p in d2.paragraphs)


def test_crd_design_requirements_enter_traceability_matrix(tmp_path):
    """REVIEW #8: the §8 design-driven REQ-D requirements must ALSO be registered in req_ids so they land
    in the §7 Requirement Traceability matrix (and the seeded-requirement count) — a requirement that
    traces to nothing is the exact 'orphan requirement' §7 forbids. The §7 row for a REQ-D pre-fills the
    KNOWN HLD §4 trace (the design blueprint)."""
    snap = _snap()
    snap["design_blueprint"] = {
        "summary": {"n_recommended": 1, "n_needs_requirement": 0, "n_critical": 1},
        "decisions": [
            {"title": "Introduce first-hop redundancy", "status": "recommended", "priority": "Critical",
             "recommended_action": "Deploy HSRP/VRRP", "driver": "Gateway resilience",
             "evidence": {"summary": "52 VLAN(s) without FHRP"}, "principle": {"citation": "CCDE — HA"}},
        ],
        "coverage": {},
    }
    out = str(tmp_path / "c.docx")
    write_crd_docx(out, snap, "Unit Test Fleet")
    d = Document(out)
    # locate the §7 traceability matrix by its header (distinct from the §2-6 capture tables and §8 detail)
    matrix = next((t for t in d.tables
                   if [c.text for c in t.rows[0].cells][:1] == ["REQ-ID"]
                   and "HLD section" in [c.text for c in t.rows[0].cells]), None)
    assert matrix is not None, "no §7 traceability matrix found"
    rids = [r.cells[0].text for r in matrix.rows[1:]]
    assert "REQ-D-001" in rids, f"REQ-D requirement is orphaned from the §7 matrix; have {rids}"
    row = next(r for r in matrix.rows[1:] if r.cells[0].text == "REQ-D-001")
    assert "§4" in " ".join(c.text for c in row.cells)          # traces to the known HLD §4 design blueprint


def test_crd_evidence_facts_tolerates_non_dict_executive_brief(tmp_path):
    """[multi-domain audit L4] crd.py's chained '(executive_brief or {}).get(scale)' crashed on a TRUTHY non-dict
    executive_brief (malformed/slimmed snapshot) -> the whole CRD silently aborted. The evidence-facts read and
    the doc builder must degrade, never raise."""
    from cisco_toolkit.crd import write_crd_docx, _evidence_facts
    ev = _evidence_facts({"executive_brief": "oops", "devices": {"sw1": {}}})
    assert ev["n_devices"] == 1
    out = str(tmp_path / "crd.docx")
    write_crd_docx(out, {"executive_brief": "oops", "devices": {"sw1": {}}}, "Test")   # must not raise


def test_crd_tolerates_non_dict_design_blueprint(tmp_path):
    """[audit-2 L5] my batch-2 L4 fix guarded design_blueprint in design.py but NOT crd.py -> write_crd_docx still
    crashed on a truthy non-dict design_blueprint. Now guarded."""
    out = str(tmp_path / "crd.docx")
    write_crd_docx(out, {"executive_brief": {}, "design_blueprint": "oops", "devices": {"s": {}}}, "x")  # no raise
    import os
    assert os.path.exists(out)


def test_crd_and_ops_survive_nondict_sections(tmp_path):
    """[audit-4 #10 totality] a malformed-but-ingest-valid uploaded snapshot (a TRUTHY non-dict section --
    executive_brief string, collection_completeness/interfaces/lifecycle_risk list-or-string) made write_crd_docx /
    write_ops_handbook_docx raise AttributeError -> HTTP 500 from the public webapp upload endpoint -- instead of
    degrading. Every section read must coerce."""
    from cisco_toolkit.crd import write_crd_docx
    from cisco_toolkit.ops import write_ops_handbook_docx
    cases = [
        {"devices": {}, "executive_brief": "corrupted"},
        {"devices": {}, "collection_completeness": [1]},
        {"devices": {}, "interfaces": [1, 2]},
        {"devices": {}, "routing_neighbors": ["x"], "l3_forwarding": {"k": "v"}, "service_map": "nope",
         "lifecycle_risk": "nope", "endpoint_dependencies": [1], "punchlist": [None, "x"]},
    ]
    for i, snap in enumerate(cases):
        write_crd_docx(str(tmp_path / f"crd{i}.docx"), snap, "L")            # must not raise
        write_ops_handbook_docx(str(tmp_path / f"ops{i}.docx"), snap, "L")   # must not raise


# --------------------------------------------------------------------------- Constraints / Out-of-Scope / RTM
def _snap_with_register():
    """The _snap() fleet PLUS a real requirements-overlaid design_blueprint — the SSOT read path the CRD
    uses for the Constraints / Out-of-Scope overlay. Built with the engine's own load_requirements +
    compute_design_blueprint (not a hand-mocked requirements_model), so the test exercises the actual echo
    of the register into design_blueprint.requirements_model.fields[*].value."""
    from cisco_toolkit.design_advisor import compute_design_blueprint
    snap = _snap()
    req = {
        "availability_tier": "gold",
        "fabric_operating_model": "nxos-evpn",
        "constraints": [
            "Standardize on Cisco Nexus 9000 / NX-OS, managed by Nexus Dashboard (NDFC + NDI)",
            "Out-of-band management on physically separate gear; it must not transit the fabric it manages",
        ],
    }
    snap["design_blueprint"] = compute_design_blueprint(snap, req)
    return snap


def test_crd_constraints_section_from_register_when_supplied(tmp_path):
    """Constraints (§7): when a requirements register is supplied (echoed into
    design_blueprint.requirements_model), its constraints are rendered as CONFIRMED rows — read off the
    canonical blueprint (SSOT), not re-loaded — and the availability tier + fabric operating model become
    confirmed constraints too. No OPEN QUESTION rows in the register-supplied path."""
    out = str(tmp_path / "c.docx")
    write_crd_docx(out, _snap_with_register(), "Unit Test Fleet")
    d = Document(out)
    h1 = [p.text for p in d.paragraphs if p.style.name == "Heading 1"]
    assert any(t == "7. Constraints & Assumptions" for t in h1), f"missing §7; have {h1}"
    text = _all_text(d)
    assert "CONFIRMED (requirements register)" in text                       # register-sourced, not invented
    assert "Nexus Dashboard (NDFC + NDI)" in text                            # a verbatim register constraint
    assert "must not transit the fabric" in text                             # the OOB register constraint
    assert "Standalone NX-OS VXLAN BGP-EVPN" in text                         # fabric_operating_model -> constraint
    assert "gold" in text                                                     # availability_tier -> constraint
    assert "CON-001" in text                                                  # constraints get CON-IDs
    # coverage-honesty: a supplied register does NOT surface the no-register OPEN QUESTIONS
    assert "no requirements register supplied" not in text.lower()


def test_crd_constraints_are_open_questions_without_register(tmp_path):
    """Constraints (§7), coverage-honesty: with NO requirements register the engagement constraints are
    surfaced as OPEN QUESTIONS (the fabric-model / OOB / overlapping-VLAN / installed-base / cutover asks),
    never asserted as if the customer had stated them. _snap() carries no design_blueprint -> no register."""
    out = str(tmp_path / "c.docx")
    write_crd_docx(out, _snap(), "Unit Test Fleet")     # _snap() has no design_blueprint
    text = _all_text(Document(out))
    assert "OPEN QUESTION — no requirements register supplied" in text
    assert "port-VLAN translation" in text                                   # overlapping-VLAN open question
    assert "Out-of-band management must NOT transit" in text                 # OOB open question
    assert "Installed-base reuse" in text                                    # installed-base open question
    # honesty: the no-register path must NOT print a CONFIRMED register row
    assert "CONFIRMED (requirements register)" not in text


def test_crd_constraints_surface_eol_forced_replacements_from_evidence(tmp_path):
    """Constraints (§7): EoL / past-LDoS hardware is an EVIDENCE-derived constraint (a forced replacement
    the migration cannot carry forward), read from lifecycle_risk — gated on observed state, distinct from
    the register-supplied constraints."""
    snap = _snap()
    snap["lifecycle_risk"] = {"summary": {
        "n_devices": 12, "n_past_ldos": 3, "n_past_eos": 4, "n_near": 5,
    }}
    out = str(tmp_path / "c.docx")
    write_crd_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "EVIDENCE (lifecycle risk)" in text
    assert "3 device(s) are past last-day-of-support" in text                # the observed LDoS count
    assert "forced replacements" in text
    assert "4 device(s) are past end-of-sale" in text
    assert "recorded window before LDoS" in text
    assert "date band does not establish support entitlement" in text
    assert "5 device(s) are within one year of LDoS" in text
    assert "schedule replacement before the recorded deadline" in text
    assert "vendor-support horizon" not in text


def test_crd_has_out_of_scope_boundary_section(tmp_path):
    """Out of Scope (§8): the explicit boundary statement — the standard CRD guardrail. Present with a
    default 'confirm with customer' exclusion list; not-collected devices are the coverage-honest anchor."""
    out = str(tmp_path / "c.docx")
    write_crd_docx(out, _snap(), "Unit Test Fleet")
    d = Document(out)
    h1 = [p.text for p in d.paragraphs if p.style.name == "Heading 1"]
    assert any(t == "8. Out of Scope" for t in h1), f"missing §8; have {h1}"
    text = _all_text(d)
    assert "confirm with customer" in text.lower()                           # exclusions are proposals to confirm
    assert "Not-collected devices" in text                                   # coverage-honest boundary
    assert "Application / server / storage changes" in text                  # a sensible default exclusion


def test_crd_rtm_traces_reqids_forward_to_four_downstream_columns(tmp_path):
    """RTM (§9): the Cisco-AS Requirements Traceability Matrix — every captured REQ-ID traced FORWARD to
    HLD / LLD / MOP / NRFU as HONEST placeholders ('to be traced in <deliverable>'), never a fabricated
    section number. The four downstream columns are all present."""
    out = str(tmp_path / "c.docx")
    write_crd_docx(out, _snap(), "Unit Test Fleet")
    d = Document(out)
    h1 = [p.text for p in d.paragraphs if p.style.name == "Heading 1"]
    assert any(t == "9. Requirements Traceability Matrix (RTM)" for t in h1), f"missing §9; have {h1}"
    # locate the RTM by its 5-column header (REQ-ID + the four forward deliverables)
    rtm = next((t for t in d.tables
                if [c.text for c in t.rows[0].cells] ==
                ["REQ-ID", "HLD section", "LLD object", "MOP step", "NRFU test case"]), None)
    assert rtm is not None, "no §9 RTM with the four forward columns"
    rids = [r.cells[0].text for r in rtm.rows[1:]]
    for seeded in ("REQ-B-001", "REQ-T-LAN-001", "REQ-O-001"):
        assert seeded in rids, f"{seeded} orphaned from the RTM; have {rids}"
    text = _all_text(d)
    # forward cells are honest placeholders, not invented references
    assert "to be traced in LLD" in text and "to be traced in MOP" in text and "to be traced in NRFU" in text


def test_crd_rtm_forward_cells_are_placeholders_not_fabricated_sections(tmp_path):
    """RTM honesty: a non-REQ-D requirement (no design blueprint behind it) must NOT trace to a fabricated
    downstream section number — its four forward cells stay 'to be traced in <deliverable>' placeholders.
    Guards against a regression that back-fills invented HLD/LLD/MOP/NRFU references."""
    import re
    out = str(tmp_path / "c.docx")
    write_crd_docx(out, _snap(), "Unit Test Fleet")     # no design_blueprint -> only seeded B/T/O rows
    d = Document(out)
    rtm = next(t for t in d.tables
               if [c.text for c in t.rows[0].cells] ==
               ["REQ-ID", "HLD section", "LLD object", "MOP step", "NRFU test case"])
    row = next(r for r in rtm.rows[1:] if r.cells[0].text == "REQ-B-001")
    for cell in row.cells[1:]:                                                # every forward cell
        assert cell.text.startswith("to be traced in"), f"fabricated forward ref: {cell.text!r}"
        assert not re.search(r"§\s*\d", cell.text), f"invented section number: {cell.text!r}"


def test_crd_reqd_row_traces_to_known_hld_but_placeholders_downstream(tmp_path):
    """RTM: a REQ-D requirement DOES trace to the KNOWN HLD §4 design blueprint (the same design_blueprint
    the CRD read — a real, not fabricated, reference), while LLD / MOP / NRFU stay honest placeholders."""
    snap = _snap()
    snap["design_blueprint"] = {
        "summary": {"n_recommended": 1, "n_needs_requirement": 0, "n_critical": 1},
        "decisions": [
            {"title": "Introduce first-hop redundancy", "status": "recommended", "priority": "Critical",
             "recommended_action": "Deploy HSRP/VRRP", "driver": "Gateway resilience",
             "evidence": {"summary": "52 VLAN(s) without FHRP"}, "principle": {"citation": "CCDE — HA"}},
        ],
        "coverage": {},
    }
    out = str(tmp_path / "c.docx")
    write_crd_docx(out, snap, "Unit Test Fleet")
    d = Document(out)
    rtm = next(t for t in d.tables
               if [c.text for c in t.rows[0].cells] ==
               ["REQ-ID", "HLD section", "LLD object", "MOP step", "NRFU test case"])
    row = next(r for r in rtm.rows[1:] if r.cells[0].text == "REQ-D-001")
    assert "§4" in row.cells[1].text                                         # KNOWN HLD trace (real)
    assert "to be traced in LLD" in row.cells[2].text                        # LLD placeholder
    assert "to be traced in MOP" in row.cells[3].text                        # MOP placeholder


def test_crd_requirements_overlay_reads_blueprint_ssot():
    """_requirements_overlay reads the register OFF design_blueprint.requirements_model (the ONE echo of
    the --requirements register into the snapshot) — provided flag + constraints/tier/fabric values —
    rather than re-loading the register file, so the CRD agrees with the HLD/deck/explorer."""
    from cisco_toolkit.crd import _requirements_overlay
    # provided register
    snap = _snap_with_register()
    reg = _requirements_overlay(snap["design_blueprint"])
    assert reg["provided"] is True
    assert reg["availability_tier"] == "gold"
    assert reg["fabric_operating_model"] == "nxos-evpn"
    assert any("Nexus Dashboard" in c for c in reg["constraints"])
    # no register / non-dict blueprint -> honest empties, provided False
    empty = _requirements_overlay({})
    assert empty["provided"] is False and empty["constraints"] == []
    assert empty["availability_tier"] is None and empty["fabric_operating_model"] is None


# ======================================================================================================
# Stored-DoS: a TRUTHY non-dict / non-list value ONE LEVEL INSIDE a well-formed container.
#
# The three regressions above (`..._non_dict_executive_brief`, `..._non_dict_design_blueprint`,
# `..._nondict_sections`) are all SECTION-level: they poison a whole top-level section, which the
# isinstance guards at crd.py:78/152 and the `_as_dict`/`_as_list` section reads already absorb. They
# cannot reach the class below, where every CONTAINER is well-formed and passes its upstream gate and
# only a ROW-INNER field (G2) or a PER-ELEMENT member (G3) is malformed. `x.get(k) or []` catches only
# FALSY values, so a truthy scalar survives and crashes the NEXT dereference: `len(5)`, `5.values()`,
# `5.get(...)`, `for d in 5`, `", ".join(5)`.
#
# Why it is a stored DoS, not a cosmetic defect: the snapshot is attacker-controllable — the webapp
# upload's only validation is `isinstance(snap, dict) and "devices" in snap`
# (webapp/backend/app.py:_parse_snapshot_bytes), the entry sanitizer `textutils.xml_safe_deep` passes
# int/float/bool/None through unchanged, and the POST is ACCEPTED (201) and the payload STORED before
# any render. `webapp/backend/deliverables.generate` re-raises after unlinking the temp file, which the
# read route turns into HTTP 500 — so one accepted upload 500s every later
# GET /api/snapshots/{id}/deliverable/crd.
# ======================================================================================================
def _snap_iface(port_value):
    """_snap() with ONE per-port interface value replaced; the rest of the fleet stays well-formed so
    the poisoned port is actually iterated (crd.py _evidence_facts `d = d or {}` -> d.get(...))."""
    s = _snap()
    s["interfaces"]["acc1"]["Gi1/0/5"] = port_value
    return s


def _snap_mcast(**over):
    """_snap() with ONE multicast field replaced. The sibling activity fields are untouched, so the
    multicast-active gate still fires and §4.3 still renders — poisoning the whole block at once would
    make the gate skip the section and the case would prove nothing."""
    s = _snap()
    s["service_map"]["multicast"].update(over)
    return s


def _snap_mcast_scaled(**over):
    """_snap_mcast(), plus the engine's canonical `executive_brief.scale` (which every real snapshot
    publishes). That short-circuits crd.py's `_eb_scale.get("n_vlans") or len(vlan_inventory(snap))`
    fallback, isolating THIS file's read from cisco_toolkit/analyze.py — `vlan_inventory` carries its own
    copy of the same `or []` bug on service_map.multicast.igmp_queriers (analyze.py:1632) and is another
    file's fix. See test_crd_route_survives_scalar_igmp_queriers below, which pins that residue openly
    instead of hiding it."""
    s = _snap_mcast(**over)
    s["executive_brief"] = {"scale": {"n_devices": 2, "n_vlans": 3, "n_endpoints": 2}}
    return s


def _snap_punch(severity):
    """_snap() with ONE punchlist row's `severity` replaced. An UNHASHABLE value makes the sort key's
    `_SEV_RANK.get(...)` raise TypeError before any coercion can help — dict.get hashes its argument."""
    s = _snap()
    s["punchlist"][0] = {**s["punchlist"][0], "severity": severity}
    return s


def _snap_bp(**over):
    """_snap_with_register() (a REAL compute_design_blueprint) with ONE blueprint CHILD replaced. The
    blueprint itself stays a dict, so the isinstance guard at crd.py:152 passes and the child is read."""
    s = _snap_with_register()
    s["design_blueprint"].update(over)
    return s


def _snap_bp_field(key):
    """_snap_with_register() with ONE requirements_model field's `key` replaced. crd.py builds
    `{f.get("key"): f.get("value") ...}`, so an UNHASHABLE key raises while the dict is being built."""
    s = _snap_with_register()
    fields = s["design_blueprint"]["requirements_model"]["fields"]
    fields[0] = {**fields[0], "key": key}
    return s


def _snap_bp_decision(status, **over):
    """_snap_with_register() with ONE field of the first REAL `status` decision poisoned. status/title
    stay intact so the row is still selected and rendered rather than filtered out upstream."""
    s = _snap_with_register()
    ds = s["design_blueprint"]["decisions"]
    i = next(i for i, d in enumerate(ds) if d.get("status") == status)
    ds[i] = {**ds[i], **over}
    return s


_INNER_POISON = {
    # crd.py:40 (G3) — a PER-PORT value that is a truthy scalar: `d = d or {}` passed it through to
    # d.get("switchport_mode") -> AttributeError.
    "iface_port_int": lambda: _snap_iface(5),
    "iface_port_str": lambda: _snap_iface("up"),
    # crd.py:69 (G2) — len(mc.get(...) or []) on a truthy scalar -> TypeError: object of type 'int' has no len()
    "mcast_classified_groups_int": lambda: _snap_mcast(classified_groups=5),
    "mcast_igmp_queriers_int": lambda: _snap_mcast_scaled(igmp_queriers=5),
    # crd.py:70 — (mc.get("ptp") or {}).values() on a scalar (G2), and (v or {}).get(...) per element (G3)
    "mcast_ptp_int": lambda: _snap_mcast(ptp=5),
    "mcast_ptp_element_int": lambda: _snap_mcast(ptp={"Gi1/0/1": 5}),
    # The SUBTLE one (crd.py:345-346, G3): [5] PASSES the crd.py:69 length gate (len == 1 -> mcast_active
    # True), so §4.3 renders and only THEN does groups[0].get("name") crash. A section-level or
    # activity-gate guard cannot catch this — the crash is two reads downstream of the check.
    "mcast_classified_groups_element_int": lambda: _snap_mcast(classified_groups=[5]),
    # crd.py:504 (G2) — `for d in (bp.get("decisions") or [])` over a scalar -> TypeError: not iterable
    "bp_decisions_int": lambda: _snap_bp(decisions=5),
    # crd.py:564-565 (G2) — cov = bp.get("coverage") or {} -> cov.get("caveat") on a scalar
    "bp_coverage_int": lambda: _snap_bp(coverage=5),
    # crd.py:551/552 (G2) — (d.get("evidence") or {}).get("summary") / (d.get("principle") or {}).get(...)
    "bp_decision_evidence_int": lambda: _snap_bp_decision("recommended", evidence=5),
    "bp_decision_principle_int": lambda: _snap_bp_decision("recommended", principle=5),
    # crd.py:563 (G2) — ", ".join(d.get("requirements_needed") or []); the list-of-ints variant is the
    # same reachable line one level deeper (join of a non-str element -> TypeError).
    "bp_decision_reqneeded_int": lambda: _snap_bp_decision("needs-requirement", requirements_needed=5),
    "bp_decision_reqneeded_list_of_ints": lambda: _snap_bp_decision("needs-requirement",
                                                                    requirements_needed=[7]),
    # crd.py:81 — the punchlist SORT KEY: `_SEV_RANK.get(i.get("severity"), 5)` on an UNHASHABLE
    # severity raises TypeError inside dict.get itself. Found by recursive-poison fuzzing the whole
    # snapshot, not by reading the `or []` sites — a different mechanism, same stored-DoS class.
    "punch_severity_list": lambda: _snap_punch(["x"]),
    "punch_severity_dict": lambda: _snap_punch({"a": 1}),
    # crd.py:123 — _requirements_overlay builds `{f.get("key"): ...}`; an UNHASHABLE field key raises
    # while the comprehension builds the dict. Sits at snapshot depth 5, past the depth cap of the
    # fuzz sweep — found by enumerating every HASHING site in the file instead.
    "reqmodel_field_key_list": lambda: _snap_bp_field(["availability_tier"]),
    "reqmodel_field_key_dict": lambda: _snap_bp_field({"a": 1}),
}


@pytest.mark.parametrize("name", sorted(_INNER_POISON))
def test_crd_survives_truthy_nondict_inner_value(tmp_path, name):
    """write_crd_docx must DEGRADE on a truthy non-dict/list inner value, never raise — and still write
    a readable .docx (a half-written file would be its own delivery defect)."""
    out = str(tmp_path / f"{name}.docx")
    write_crd_docx(out, _INNER_POISON[name](), "Poison Fleet")   # must not raise
    assert os.path.exists(out), f"{name}: no CRD written"
    Document(out)                                                # and it opens


def test_crd_evidence_facts_tolerates_non_str_routing_protocol_key(tmp_path):
    """crd.py:59 — `p.upper()` over routing_neighbors' PROTOCOL KEYS. Deliberately NOT in the route sweep
    above: JSON object keys are always strings, so this is unreachable from an upload. It is reachable on
    the CLI path, which hands write_crd_docx the IN-PROCESS snapshot dict (COLLECT_PARSE_V3_23_0:2886)
    without a JSON round-trip. Pinned here honestly at the level it can actually occur, rather than
    dressed up as a stored DoS it is not."""
    from cisco_toolkit.crd import _evidence_facts
    snap = _snap()
    snap["routing_neighbors"] = {"core1": {5: [{"neighbor": "10.0.0.2"}], "ospf": [{"neighbor": "10.0.0.3"}]}}
    ev = _evidence_facts(snap)                      # must not raise
    assert "OSPF" in ev["protos"], ev["protos"]     # the real protocol still renders
    assert "5" in ev["protos"], ev["protos"]        # and the odd key degrades to its string form
    write_crd_docx(str(tmp_path / "c.docx"), snap, "L")


def test_crd_wellformed_output_is_unchanged_by_the_inner_guards(tmp_path):
    """Non-vacuity anchor for the guards: on WELL-FORMED input every guarded read still yields its real
    value, so the coercers only ever swallow garbage. `as_list(x)` is identical to `(x or [])` for every
    list/None input and `as_dict(x)` to `(x or {})` for every dict/None input — this pins that in output
    terms at each of the guarded sites."""
    out = str(tmp_path / "good.docx")
    write_crd_docx(out, _snap_with_register(), "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "PTP-primary" in text              # :345-346 groups[0]["name"] still reaches §4.3 REQ-T-MC-001
    assert "1 classified group(s)" in text    # :348 len(groups) still counts the real group
    assert "2 dual-homed endpoint(s)" in text  # :54 canonical endpoint_dependencies read is intact
    # §10 rows still carry the REAL evidence summary + CCDE citation (:551/:552) and the real
    # requirements_needed join (:563) and coverage caveat (:564-565) — not "—" placeholders.
    assert "REQ-D-001" in text
    assert "growth_horizon" in text                                  # :563 join of a real needs list
    assert "Design decisions are grounded only in collected evidence" in text   # :564-565 coverage caveat
    # :551 / :552 — the REAL evidence summary and CCDE citation of the first recommended decision, read
    # verbatim off the blueprint so a guard that silently degraded them to "—" would fail here.
    bp = _snap_with_register()["design_blueprint"]
    first_rec = next(d for d in bp["decisions"] if d["status"] == "recommended")
    assert first_rec["evidence"]["summary"][:60] in text, "§10 lost the real evidence summary (:551)"
    assert first_rec["principle"]["citation"][:40] in text, "§10 lost the real CCDE citation (:552)"


# ------------------------------------------------------- E2E through the REAL upload -> render route
@pytest.fixture()
def hub_client(tmp_path):
    """The public route the stored-DoS actually travels. base_url=localhost so the default Host passes
    AssessHub's no-token DNS-rebinding allowlist; raise_server_exceptions=False so a render failure
    surfaces as a 500 STATUS CODE we can assert on, exactly as a real client experiences it."""
    pytest.importorskip("fastapi")
    from fastapi.testclient import TestClient
    webapp = str(Path(__file__).resolve().parents[1] / "webapp")
    if webapp not in sys.path:
        sys.path.insert(0, webapp)          # make `backend` importable from the tests/ suite
    from backend.app import create_app
    app = create_app(db_path=str(tmp_path / "hub.db"))
    with TestClient(app, base_url="http://localhost", raise_server_exceptions=False) as c:
        yield c


def _store_and_render_crd(client, snap):
    """POST the snapshot (the upload validates only `isinstance(snap, dict) and 'devices' in snap`, so
    the poison is ACCEPTED and STORED), then GET the CRD the stored snapshot renders."""
    cid = client.post("/api/campaigns", json={"name": "c"}).json()["id"]
    r = client.post(f"/api/campaigns/{cid}/snapshots",
                    files={"file": ("s.json", json.dumps(snap).encode(), "application/json")},
                    data={"label": "s"})
    assert r.status_code == 201, f"upload rejected ({r.status_code}) — the DoS premise needs a 201: {r.text[:200]}"
    return client.get(f"/api/snapshots/{r.json()['id']}/deliverable/crd")


@pytest.mark.parametrize("name", sorted(_INNER_POISON))
def test_crd_deliverable_route_survives_truthy_nondict_inner_value(hub_client, name):
    """End-to-end: the accepted-and-stored poison must render 200, not 500. Before the crd.py guards
    every case here returned 500 from this exact route."""
    r = _store_and_render_crd(hub_client, _INNER_POISON[name]())
    assert r.status_code == 200, f"{name}: expected 200, got {r.status_code}: {r.text[:300]}"


def test_crd_route_survives_scalar_igmp_queriers(hub_client):
    """Was xfail(strict=False) while the residue lived OUTSIDE this file, in
    cisco_toolkit/analyze.py::vlan_inventory (`(... igmp_queriers ...) or []` then `for q in 5`), which
    crd.py falls back to when executive_brief.scale.n_vlans is absent. analyze.py is guarded now, so the
    marker is dropped and this asserts for real — a regression in EITHER module fails here instead of
    being absorbed as an expected failure."""
    r = _store_and_render_crd(hub_client, _snap_mcast(igmp_queriers=5))
    assert r.status_code == 200, f"expected 200, got {r.status_code}: {r.text[:300]}"


def test_crd_deliverable_route_renders_wellformed_snapshot(hub_client):
    """Route-level non-vacuity: the same route on a well-formed snapshot returns a real .docx."""
    r = _store_and_render_crd(hub_client, _snap_with_register())
    assert r.status_code == 200, r.text[:300]
    assert r.content[:2] == b"PK", "expected a .docx (zip) body"
    assert len(r.content) > 10_000, f"suspiciously small CRD: {len(r.content)} bytes"


def _crd_cells(models, tmp_path):
    """Every table cell of a CRD built from a fleet with the given device models."""
    from cisco_toolkit.analyze import compute_lifecycle_risk
    from cisco_toolkit.crd import write_crd_docx
    snap = {"schema": "collect_parse_snapshot/1", "interfaces": {},
            "lifecycle_risk": compute_lifecycle_risk(
                {f"sw{i}": {"model": m, "sw_version": "x"} for i, m in enumerate(models)},
                asof="2026-06-07")}
    out = str(tmp_path / "crd.docx")
    write_crd_docx(out, snap, "Unit Test Fleet")
    return [c.text for t in Document(out).tables for r in t.rows for c in r.cells]


def test_crd_discloses_an_unbanded_fleet_in_BOTH_the_evidence_row_and_the_constraints(tmp_path):
    """Two exits carried the same false-clean reading; fixing one would have left the other.

    §2's evidence table printed `n_past_ldos` bare, so a fleet nothing could be banded on showed a
    plain "0" — identical to a fully lifecycle-banded fleet with no adverse date band. Separately,
    §4's constraints register
    emitted a hardware-lifecycle constraint only when `n_past_ldos > 0` or `n_past_eos > 0`; both read
    0 on an unbanded fleet, so NO constraint was recorded and the workshop wrote requirements as
    though refresh were a settled non-issue.

    The CRD is the instrument the requirements workshop is run from — both readings must be honest.
    """
    cells = _crd_cells(["TOTALLY-MADE-UP", "C9300-48T"], tmp_path)
    ldos_row = [c for c in cells if c.startswith("0 (+ 2 device(s) NOT ASSESSED")]
    assert ldos_row, f"§2 printed a bare LDoS count on an unbanded fleet: {cells!r}"
    assert "either no exact EoX row matched" in ldos_row[0]
    assert "source/date authority was withheld" in ldos_row[0]
    assert "no EoX bulletin matched them" not in ldos_row[0]
    assert any("could NOT be lifecycle-banded" in c for c in cells), \
        "§4 recorded no hardware-lifecycle constraint for a fleet whose lifecycle is undetermined"
    constraint = next(c for c in cells if "could NOT be lifecycle-banded" in c)
    assert "either no exact EoX row matched" in constraint
    assert "retained source/date authority was withheld or incomplete" in constraint
    assert "support entitlement are UNDETERMINED" in constraint
    assert any("COVERAGE GAP (lifecycle risk)" in c for c in cells), \
        "the constraint is not labelled as a coverage gap"


def test_crd_leaves_a_fully_banded_fleet_untouched(tmp_path):
    """Non-vacuity: neither disclosure may fire when every platform WAS banded, or they say nothing."""
    cells = _crd_cells(["WS-C4948E-F"], tmp_path)
    assert not any("NOT ASSESSED" in c for c in cells), \
        "a fully-banded fleet gained a spurious coverage disclosure in §2"
    assert not any("could NOT be lifecycle-banded" in c for c in cells)
    # ...and the real EoL constraint must still be the one recorded.
    assert any("past last-day-of-support (LDoS)" in c for c in cells)


# =====================================================================================================
# §4.3 REQ-T-MC-001 multicast classification AUTHORITY (review r9 EXIT D).
#
# The requirement named the classified-group count -- "Preserve multicast delivery for the N classified
# group(s) (e.g. PTP-primary) ..." -- with no authority qualification. The CRD is the instrument a
# requirements WORKSHOP is run from, so a requirement written against a CURATED classification and
# stated as though observed becomes a contractual commitment. analyze.compute_multicast_intelligence
# already publishes the basis (`summary.n_av_groups_authoritative`); on the SHIPPED port pack it is
# ZERO -- every multicast row is curated -- and the same on-air flag escalates a mac-alias finding from
# Medium to High (analyze.py:2774). The fix is DISCLOSURE inside the requirement, not a re-scoring.
# =====================================================================================================
_MC_BASIS_LEAD = "CLASSIFICATION BASIS: the group addresses and their multicast activity are OBSERVED"


def _real_multicast_crd_snap(groups=("224.0.1.129", "239.255.255.250"), authoritative=()):
    """A CRD snapshot whose multicast sections come from the REAL producers.

    `compute_service_map(igmp_groups=...)` classifies the addresses through the SHIPPED offline port
    registry, and `compute_multicast_intelligence` derives the authority census from that. Addresses
    named in `authoritative` have their registry semantics flipped BEFORE the intelligence pass, so the
    census is still the producer's own arithmetic -- the only way to obtain the authoritative fleet the
    shipped pack cannot produce (every multicast row in it is curated)."""
    from cisco_toolkit.analyze import compute_multicast_intelligence, compute_service_map
    sm = compute_service_map({}, {}, igmp_groups=list(groups))
    for g in sm["multicast"]["classified_groups"]:
        if g.get("group") in authoritative:
            g["semantics_authoritative"] = True
            g["semantics_source"] = "IANA multicast-addresses registry"
    sm["multicast"]["active_switch_count"], sm["multicast"]["active_interfaces"] = 1, 2
    snap = _snap()
    snap["service_map"] = sm
    snap["multicast_intelligence"] = compute_multicast_intelligence(sm)
    return snap


def _crd_text(snap, tmp_path, name="c.docx"):
    out = str(tmp_path / name)
    write_crd_docx(out, snap, "Multicast Fleet")
    return _all_text(Document(out))


def test_crd_multicast_requirement_states_its_curated_classification_basis(tmp_path):
    """Shipped-pack reality: NO multicast group is classified on-air by an authoritative source, so
    REQ-T-MC-001 must say so in the row the customer signs. Pre-fix the requirement ended at
    '...snooping/querier behaviour per VLAN.' and the word CURATED appeared nowhere in the CRD."""
    snap = _real_multicast_crd_snap()
    s = snap["multicast_intelligence"]["summary"]
    assert (s["n_av_groups"], s["n_av_groups_authoritative"]) == (1, 0)   # producer's verdict
    text = _crd_text(snap, tmp_path)
    assert "Preserve multicast delivery for the 2 classified group(s)" in text   # the ask is unchanged
    assert _MC_BASIS_LEAD in text, "REQ-T-MC-001 still states a curated classification as observed"
    assert ("NONE of them classified on-air by an authoritative source: the broadcast/AV label is a "
            "CURATED offline-registry classification, not a measurement") in text
    assert "Confirm the real use of each group in the workshop" in text


def test_crd_multicast_requirement_is_not_assessed_when_the_census_is_absent(tmp_path):
    """FAIL-CLOSED. A snapshot with classified groups but no `multicast_intelligence` at all must read
    NOT ASSESSED -- an absent census may never be presented as authority, and may never be silently
    dropped from a requirement that is about to become a commitment."""
    snap = _real_multicast_crd_snap()
    snap.pop("multicast_intelligence")
    text = _crd_text(snap, tmp_path)
    assert _MC_BASIS_LEAD in text
    assert ("on-air classification authority NOT ASSESSED in this snapshot; treat the broadcast/AV "
            "label as curated, not authoritative") in text


def test_crd_multicast_requirement_credits_an_authoritative_fleet(tmp_path):
    """NON-VACUITY 1. A fleet whose on-air classification IS authoritative must be reported as such --
    the qualifier is derived from the census, not an unconditional caveat on every CRD."""
    snap = _real_multicast_crd_snap(authoritative=("224.0.1.129",))
    s = snap["multicast_intelligence"]["summary"]
    assert (s["n_av_groups"], s["n_av_groups_authoritative"]) == (1, 1)   # producer's verdict
    text = _crd_text(snap, tmp_path)
    assert "1 of 1 on-air classification(s) rest on an authoritative source" in text
    assert "NONE of them classified on-air by an authoritative source" not in text
    assert "authority NOT ASSESSED in this snapshot" not in text


def test_crd_fleet_with_no_media_estate_gains_no_multicast_requirement_or_caveat(tmp_path):
    """NON-VACUITY 2. §4.3 is gated on observed multicast, so a fleet with no media estate must get
    NEITHER the requirement NOR its authority caveat -- the disclosure may not leak into a document
    that classifies nothing."""
    snap = _snap()
    snap["service_map"] = {"services": snap["service_map"]["services"]}
    text = _crd_text(snap, tmp_path)
    assert "REQ-T-MC-001" not in text
    for token in (_MC_BASIS_LEAD, "NONE of them classified on-air",
                  "on-air classification authority NOT ASSESSED"):
        assert token not in text, f"a non-media fabric acquired multicast authority text: {token!r}"


def test_crd_multicast_requirement_refuses_an_incoherent_census(tmp_path):
    """authoritative > total is incoherent (the counts are published independently). '7 of 3' is
    indistinguishable from a real ratio, so the incoherence is disclosed instead -- runbook.py carries
    the same check because an earlier fix printed a NEGATIVE count into a client workbook."""
    snap = _real_multicast_crd_snap()
    snap["multicast_intelligence"]["summary"].update(n_av_groups=3, n_av_groups_authoritative=7)
    text = _crd_text(snap, tmp_path)
    assert "census INCOHERENT: 7 classification(s) reported as authoritative out of 3 on-air group(s)" in text
    assert "7 of 3" not in text


@pytest.mark.parametrize("bad", [5, "x", [1, 2], {"k": 1}, None, {"summary": 5}])
def test_crd_multicast_authority_tolerates_a_malformed_census(tmp_path, bad):
    """A malformed `multicast_intelligence` must degrade to NOT ASSESSED and still emit an openable
    CRD -- never claim authority, never abort the deliverable (an HTTP 500 on the webapp route)."""
    import os
    snap = _real_multicast_crd_snap()
    snap["multicast_intelligence"] = bad
    out = str(tmp_path / "c.docx")
    write_crd_docx(out, snap, "Poison")
    assert os.path.exists(out)
    text = _all_text(Document(out))
    assert "authority NOT ASSESSED in this snapshot" in text
    assert "rest on an authoritative source" not in text


def test_crd_av_authority_mirrors_the_runbook_wording():
    """crd._av_authority is a deliberate MIRROR of runbook._av_authority (documented in both
    docstrings) so the CRD and the runbook cannot make two different authority claims about the same
    fact in one deliverable set. Drive both over one case table."""
    from cisco_toolkit.crd import _av_authority as crd_av
    from cisco_toolkit.runbook import _av_authority as runbook_av
    cases = [
        None, {}, 5, "x", {"summary": 5}, {"summary": {}},
        {"summary": {"n_av_groups": 3, "n_av_groups_authoritative": 0}},
        {"summary": {"n_av_groups": 3, "n_av_groups_authoritative": 2}},
        {"summary": {"n_av_groups": 3, "n_av_groups_authoritative": 3}},
        {"summary": {"n_av_groups": 3, "n_av_groups_authoritative": 7}},
        {"summary": {"n_av_groups": 3, "n_av_groups_authoritative": None}},
        {"summary": {"n_av_groups": 3, "n_av_groups_authoritative": "lots"}},
        {"summary": {"n_av_groups": 3, "n_av_groups_authoritative": -4}},
        {"summary": {"n_av_groups": None, "n_av_groups_authoritative": 2}},
    ]
    for case in cases:
        assert crd_av(case) == runbook_av(case), f"wording diverged on {case!r}"
    assert len({runbook_av(c) for c in cases}) >= 4    # the table is not vacuous
