"""NEW-V3.23.150: shared AS-style document furniture (document control + acceptance gate) for the
DOCX deliverable family. python-docx is optional, so the module is skipped when it is absent (the
helpers are only ever called by writers whose own docx import already succeeded)."""
import pytest

docx = pytest.importorskip("docx")  # skip the file if the optional dep is absent
from docx import Document  # noqa: E402

from cisco_toolkit.docmeta import (  # noqa: E402
    add_acceptance,
    add_document_control,
    add_excellence_front,
    add_glossary,
    add_inputs_required,
    add_toc,
    as_dict,
    as_list,
    related_rows,
)


def _all_text(doc):
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            parts.extend(c.text for c in row.cells)
    return "\n".join(parts)


def test_add_toc_marks_fields_to_rebuild_on_open():
    """DE-TOC: a field-code TOC renders as a placeholder until refreshed, so every deliverable
    must ship with <w:updateFields w:val="true"/> — the client (or a headless docx→pdf render)
    gets a built table of contents on open, not a 'press F9' stub. Guards the V3.23.172 fix."""
    from docx.oxml.ns import qn
    from cisco_toolkit.docmeta import enable_update_fields

    doc = Document()
    add_toc(doc)
    settings = doc.settings.element
    uf = settings.findall(qn("w:updateFields"))
    assert len(uf) == 1 and uf[0].get(qn("w:val")) == "true"
    # schema position: CT_Settings is an ordered sequence — updateFields must PRECEDE w:compat/w:rsids,
    # or a strict XSD validator rejects the deliverable (Word tolerates it; a client's toolchain may not).
    kids = list(settings)
    for later in ("w:compat", "w:rsids"):
        el = settings.find(qn(later))
        if el is not None:
            assert kids.index(uf[0]) < kids.index(el), "updateFields must precede %s (schema order)" % later
    # idempotent: re-invoking (a second TOC, or an explicit call) never duplicates the setting
    add_toc(doc)
    enable_update_fields(doc)
    uf2 = settings.findall(qn("w:updateFields"))
    assert len(uf2) == 1 and uf2[0].get(qn("w:val")) == "true"


def test_excellence_front_reads_canonical_facts_not_a_recount():
    # DE-01: the 'At a Glance' register must render the CANONICAL headline facts (single source of
    # truth), plus the ssot self-verification signal — proof it reads, not recounts.
    doc = Document()
    snap = {"executive_brief": {"scale": {"n_devices": 303, "n_collected": 253,
                                          "n_endpoints": 5127, "n_vlans": 202},
                                "posture": {"avg_health": 41, "n_critical": 117, "n_poor": 85}},
            "lifecycle_risk": {"summary": {"n_past_ldos": 152}}}
    add_excellence_front(doc, snap)
    txt = _all_text(doc)
    assert "At a Glance" in txt
    assert "253 of 303" in txt               # canonical scope, not a recount
    assert "152 device" in txt               # the marquee lifecycle fact
    assert "Single source of truth" in txt   # the Law-10 self-verified signal


def test_excellence_front_is_coverage_honest_on_absent_blocks():
    # The bite: with NO canonical blocks published, an unpublished fact must read [NOT OBSERVED],
    # never a fabricated healthy-looking 0 (the false-health class this whole layer exists to kill).
    doc = Document()
    add_excellence_front(doc, {})
    txt = _all_text(doc)
    assert "At a Glance" in txt
    assert "[NOT OBSERVED]" in txt
    assert "0 device(s) past last-date-of-support" not in txt


def test_glossary_always_carries_the_coverage_honesty_convention():
    doc = Document()
    add_glossary(doc, extra_terms=(("XYZ", "a doc-specific term"),))
    txt = _all_text(doc)
    assert "Glossary" in txt and "FHRP" in txt and "XYZ" in txt
    assert "[NOT OBSERVED]" in txt   # the not-observed-is-not-healthy row must always be present


def test_inputs_required_surfaces_not_collected_devices_coverage_honestly():
    # DE-01/Law 9: the register must name the not-collected blind spot as a concrete input to close,
    # never bury it. And it must NOT invent one when nothing is uncollected.
    doc = Document()
    add_inputs_required(doc, {"collection_completeness": {"summary": {"not_collected": 50}}})
    txt = _all_text(doc)
    assert "Inputs Required" in txt
    assert "50 inventoried device(s) were not collected" in txt and "[NOT OBSERVED]" in txt

    doc2 = Document()
    add_inputs_required(doc2, {"collection_completeness": {"summary": {"not_collected": 0}}})
    txt2 = _all_text(doc2)
    assert "Inputs Required" in txt2                       # heading + the standard gates still present
    assert "were not collected" not in txt2                # but no fabricated blind-spot row


def test_document_control_includes_a_distribution_list():
    doc = Document()
    add_document_control(doc, document="X.docx", label="Meridian", engine_version="9")
    txt = _all_text(doc)
    assert "Distribution list" in txt


def test_add_toc_renders_field_code_and_placeholder():
    # V3.23.171: the TOC field-code block was a verbatim copy in all 7 generators; the shared
    # helper must keep the contract every writer test relies on (the Right-click placeholder).
    doc = Document()
    add_toc(doc)
    assert any("Contents" in p.text for p in doc.paragraphs
               if p.style.name.startswith("Heading"))
    assert "Right-click → Update Field" in _all_text(doc)
    xml = doc.element.xml
    assert 'TOC \\o "1-2"' in xml and "fldChar" in xml      # the real Word field code
    deep = Document()
    add_toc(deep, depth="1-3")
    assert 'TOC \\o "1-3"' in deep.element.xml              # depth is configurable in ONE place


def test_shared_coercers_guard_truthy_non_containers():
    # the V3.23.159 review's bug class, now importable instead of copied per generator
    assert as_dict({"a": 1}) == {"a": 1}
    assert as_dict("truthy-non-dict") == {} and as_dict(None) == {} and as_dict(3) == {}
    assert as_list([1]) == [1]
    assert as_list("nope") == [] and as_list(None) == [] and as_list({"a": 1}) == []


def test_related_rows_excludes_the_document_being_written():
    rows = related_rows(exclude=("mop",))
    names = [n for n, _ in rows]
    assert not any("Method of Procedure" in n for n in names)
    assert any("Assessment workbook" in n for n in names)
    assert len(rows) == 11  # the 12-document family (incl. the ops handbook, V3.23.168) minus self
    # a bare-string exclude is one key, not an iterable of letters (set("mop") footgun)
    assert related_rows(exclude="mop") == rows


def test_related_rows_marks_the_web_only_kinds_the_engine_never_writes():
    """The table lists the whole FAMILY, but `cli_artifacts` (what a COMPLETE engine CLI/Atlas run
    leaves on disk) deliberately excludes `cutover` and `nrfu` — AssessHub renders those. Unmarked,
    an offline delivery advertised a Cutover Plan and an NRFU/ATP that are not in the folder and that
    no engine writer produces. Every OTHER row must stay unannotated (the note is not decoration)."""
    from cisco_toolkit.docmeta import (CLI_ARTIFACT_SUFFIX, FAMILY, WEB_ONLY_KINDS,
                                       WEB_ONLY_ROLE_NOTE)
    assert WEB_ONLY_KINDS and not (WEB_ONLY_KINDS & set(CLI_ARTIFACT_SUFFIX))
    by_name = dict(related_rows())
    for key, name, _role in FAMILY:
        role = by_name[name]
        if key in WEB_ONLY_KINDS:
            assert WEB_ONLY_ROLE_NOTE in role, f"{name} does not disclose its producer"
            assert "AssessHub" in role and "not produced by the offline engine run" in role
        else:
            assert WEB_ONLY_ROLE_NOTE not in role, f"{name} wrongly marked web-only"


def test_related_documents_intro_does_not_assert_a_reconciliation_it_never_ran():
    """The intro asserted "the set is internally consistent (every number reconciles)"
    UNCONDITIONALLY — in the same front matter where add_excellence_front's single-source-of-truth
    line may report that NONE of the figures could be reconciled, or that some disagreed. A claim of
    verification is not furniture: the reconcile verdict has one owner (ssot.summary, rendered under
    'At a Glance'), and this block must point at it rather than pre-empt it."""
    from docx import Document

    from cisco_toolkit.docmeta import add_related_documents
    doc = Document()
    add_related_documents(doc, exclude=("mop",), intro=True)
    text = _all_text(doc)
    assert "one of a set generated from the same assessment snapshot" in text   # the provenance fact
    assert "every number reconciles" not in text
    assert "internally consistent" not in text
    assert "single-source-of-truth line" in text and "does not assert it" in text


def test_document_control_renders_all_furniture():
    doc = Document()
    add_document_control(doc, document="Unit Doc", label="Fleet X", engine_version="V9",
                         generated_at="2026-06-10 09:00", audience="unit testers",
                         exclude=("design",), extra_assumptions=("extra unit caveat",))
    text = _all_text(doc)
    assert "Document Control" in text and "Revision history" in text
    assert "Fleet X" in text and "2026-06-10 09:00" in text
    assert "DRAFT — generated; not yet reviewed" in text
    assert "Related documents" in text and "Assessment workbook (.xlsx)" in text
    assert "As-Built Network Design Document (.docx)" not in text   # self excluded from the set
    assert "Assumptions & caveats" in text and "extra unit caveat" in text
    assert "unit testers" in text


def test_document_control_separates_collection_from_generation():
    """Provenance (wave R2-3-01): 'Snapshot captured' is the COLLECTION instant; 'Document generated'
    is the render time. On a --no-collect re-render these differ and must NOT be conflated."""
    doc = Document()
    add_document_control(doc, document="Unit Doc", label="Fleet X", engine_version="V9",
                         generated_at="2026-06-21T06:58:07", collected_at="2026-06-13T06:32:01",
                         audience="unit testers")
    rows = {r.cells[0].text: r.cells[1].text for tbl in doc.tables for r in tbl.rows}
    assert rows.get("Snapshot captured") == "2026-06-13T06:32:01"   # collection date headlines
    assert rows.get("Document generated") == "2026-06-21T06:58:07"  # render date is distinct, not conflated


def test_document_control_collected_at_falls_back_to_generated():
    """Back-compat: a snapshot predating collected_at still shows the generation date under
    'Snapshot captured' (never blank)."""
    doc = Document()
    add_document_control(doc, document="Unit Doc", label="Fleet X", generated_at="2026-06-10 09:00")
    rows = {r.cells[0].text: r.cells[1].text for tbl in doc.tables for r in tbl.rows}
    assert rows.get("Snapshot captured") == "2026-06-10 09:00"      # fallback preserved


def test_related_documents_block_renders_table_and_audience():
    """V3.23.152: the shared family cross-reference block (used by NRFU/PIR directly and by
    add_document_control with intro=True)."""
    from cisco_toolkit.docmeta import add_related_documents

    doc = Document()
    add_related_documents(doc, exclude=("nrfu",), audience="unit audience")
    text = _all_text(doc)
    assert "Related documents" in text
    assert "NRFU / Acceptance Test Plan (.docx)" not in text   # self excluded
    assert "Assessment workbook (.xlsx)" in text
    assert "Intended audience:" in text and "unit audience" in text
    assert "one of a set" not in text                          # intro only when intro=True


def test_acceptance_renders_roles_and_optional_heading():
    doc = Document()
    add_acceptance(doc, scope_note="unit scope note")
    text = _all_text(doc)
    assert "Document Acceptance" in text and "unit scope note" in text
    for role in ("Customer network owner", "Customer operations lead",
                 "Delivery engineer (author)", "Project / change manager"):
        assert role in text
    # heading=None appends the gate under an existing section without a new H1 (the MOP's use)
    doc2 = Document()
    add_acceptance(doc2, heading=None)
    h1 = [p.text for p in doc2.paragraphs if p.style.name == "Heading 1"]
    assert h1 == []
    assert "Customer network owner" in _all_text(doc2)
