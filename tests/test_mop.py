"""NEW-V3.23.149: the per-wave Method of Procedure (MOP, DOCX) deliverable. python-docx is optional, so
the module is skipped when it is absent (the generator fails soft the same way). These tests pin the
one-section-per-wave structure, the change-overview reconciliation, the reuse of the existing validation
plan as the post-cutover checks, the blocker surfacing, and the fail-soft path."""
import json
import os
import re
import sys

import pytest

docx = pytest.importorskip("docx")  # skip the file if the optional dep is absent
from docx import Document  # noqa: E402

from cisco_toolkit.mop import write_mop_docx  # noqa: E402


def _snap():
    """A two-wave migration snapshot exercising every part of a MOP wave section."""
    return {
        "script_version": "V3.23.0",
        "devices": {"distA": {"platform": "ios"}, "distB": {"platform": "nxos"}, "acc1": {"platform": "ios"}},
        "move_groups": [{"switches": ["distA", "distB"], "endpoints": 40},
                        {"switches": ["acc1"], "endpoints": 12}],
        "migration_readiness": [
            {"group": "Group 1", "switches": ["distA", "distB"], "endpoints": 40,
             "readiness": "CAUTION", "n_fail": 0, "n_warn": 2, "checks": []},
            {"group": "Group 2", "switches": ["acc1"], "endpoints": 12,
             "readiness": "NOT READY", "n_fail": 1, "n_warn": 0, "checks": []},
        ],
        "wave_sequencing": [
            {"group": "Group 1", "make_before_break": ["distA", "distB"], "hard_cutover": [],
             "hard_cutover_endpoints": 0, "sequence": "make-before-break"},
            {"group": "Group 2", "make_before_break": [], "hard_cutover": ["acc1"],
             "hard_cutover_endpoints": 12, "sequence": "hard-cutover"},
        ],
        "migration_scenarios": {
            "per_group": [
                {"group": "Group 1", "switches": 2, "endpoints": 40, "readiness": "CAUTION",
                 "recommended_scenario": "parallel-run", "rationale": "dual-homed — build beside",
                 "playbook": {"pre": "stage target beside legacy", "validate": "cut one leg, prove forwarding",
                              "rollback": "fail back to the legacy leg"}},
                {"group": "Group 2", "switches": 1, "endpoints": 12, "readiness": "NOT READY",
                 "recommended_scenario": "hard-cutover", "rationale": "single-homed access edge",
                 "playbook": {"pre": "schedule outage window", "validate": "ping the gateway",
                              "rollback": "re-cable to legacy port"}},
            ],
            "fleet_recommendation": "Cut the dual-homed core first, then the single-homed access edge.",
            "scenario_counts": {"parallel-run": 1, "hard-cutover": 1},
        },
        "validation_plan": {
            "items": [
                {"category": "Gateway", "device": "distA", "check": "VLAN 10 gateway up",
                 "command": "show ip interface brief", "expect": "10.0.10.1 up/up", "wave": "Group 1"},
                {"category": "FHRP", "device": "distA", "check": "HSRP role",
                 "command": "show standby brief", "expect": "Active/Standby", "wave": "Group 1"},
                {"category": "Reachability", "device": "acc1", "check": "gateway reachable",
                 "command": "ping 10.0.20.1", "expect": "100 percent", "wave": "Group 2"},
            ],
            "by_wave": {
                "Group 1": [
                    {"category": "Hygiene", "device": "distA", "check": "login banner present",
                     "command": "show banner login", "expect": "MOTD shown", "wave": "Group 1",
                     "severity": "Low", "why": "cosmetic / compliance"},
                    {"category": "Gateway", "device": "distA", "check": "VLAN 10 gateway up",
                     "command": "show ip interface brief", "expect": "10.0.10.1 up/up", "wave": "Group 1",
                     "severity": "High", "why": "VLAN 10 endpoints lose their gateway if this SVI is down"},
                    {"category": "FHRP", "device": "distA", "check": "HSRP role",
                     "command": "show standby brief", "expect": "Active/Standby", "wave": "Group 1",
                     "severity": "Medium"},
                ],
                "Group 2": [
                    {"category": "Reachability", "device": "acc1", "check": "gateway reachable",
                     "command": "ping 10.0.20.1", "expect": "100 percent", "wave": "Group 2"},
                ],
            },
            "summary": {"n_items": 3, "n_waves": 2, "n_high": 2, "by_category": {"Gateway": 1}},
            "banner": "3 checks across 2 waves",
        },
        "failure_impact": [{"host": "distA", "severity": "High", "stranded": 22, "vlans_impacted": 3},
                           {"host": "acc1", "severity": "Medium", "stranded": 6, "vlans_impacted": 1}],
        "remediation_plan": {"items": [
            {"source": "fhrp", "device": "acc1", "severity": "High", "title": "Add FHRP peer",
             "why": "single gateway strands the VLAN", "commands": ["interface Vlan20", " standby 20 ip 10.0.20.254"]},
            {"source": "hygiene-undefined", "device": "distA", "severity": "Low", "title": "unused ACL",
             "commands": ["! review only"]},
        ]},
        "punchlist": [{"severity": "Critical", "category": "L3 design", "devices": ["acc1"],
                       "title": "VLAN 20 single gateway", "detail": "x"},
                      {"severity": "Low", "category": "Hygiene", "devices": ["distA"],
                       "title": "cosmetic", "detail": "y"}],
        "executive_brief": {"top_gating": ["Resolve the VLAN 20 single-gateway exposure before wave 2."]},
    }


def _all_text(doc):
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            parts.extend(c.text for c in row.cells)
    return "\n".join(parts)


def _row_values(doc, label, col=1):
    """The `col`-th cell of every table row whose FIRST cell is exactly `label`, in document order.

    WHY THIS EXISTS (mutation-proved, 2026-07-28). A whole-document `assert "[NOT OBSERVED]" in
    text` cannot tell a rendered abstention ROW from the boilerplate paragraph that merely
    EXPLAINS the convention. mop.py quotes the marker verbatim in the BLUF intro, in the
    pre-implementation-checklist intro and in the terms-and-conventions section, so the token
    appears eleven times in a completely unmodified MOP — before any coverage-honesty fixture
    strips a single fact. Three tests here asserted exactly that and stayed green while the row
    each of them named silently degraded. Look up the row."""
    out = []
    for t in doc.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells]
            if cells and cells[0] == label and len(cells) > col:
                out.append(cells[col])
    return out


def test_mop_has_one_section_per_wave(tmp_path):
    out = str(tmp_path / "m.docx")
    write_mop_docx(out, _snap(), "Unit Test Fleet")
    d = Document(out)
    h1 = [p.text for p in d.paragraphs if p.style.name == "Heading 1"]
    assert any(t == "1. Change Overview" for t in h1)
    assert any(t.startswith("2. Global Prerequisites") for t in h1)
    # one MOP section per wave
    assert any(t == "3. MOP — Group 1" for t in h1), h1
    assert any(t == "4. MOP — Group 2" for t in h1), h1
    # final acceptance section closes it
    assert any("Post-Migration Acceptance" in t for t in h1)


def test_mop_reuses_validation_plan_as_checks(tmp_path):
    out = str(tmp_path / "m.docx")
    write_mop_docx(out, _snap(), "Unit Test Fleet")
    text = _all_text(Document(out))
    # the EXISTING validation plan items appear as the post-cutover go/no-go checks (one source of truth)
    assert "show standby brief" in text and "Active/Standby" in text          # Group 1 FHRP check
    assert "ping 10.0.20.1" in text and "100 percent" in text                 # Group 2 reachability check
    # strategy drives the procedure wording: make-before-break vs hard cutover
    assert "BESIDE" in text or "beside" in text
    assert "hard cutover" in text.lower()


def test_mop_mixed_homing_wave_is_not_classified_pure_make_before_break(tmp_path):
    """A wave with BOTH dual-homed (make_before_break) and single-homed (hard_cutover) members must not be
    classified pure make-before-break off the free-text 'sequence' — single-homed switches have no second
    path. The procedure sequences the dual-homed MBB first then takes the scheduled outage; the rollback
    flags the single-homed members' rollback as a brief outage, not 'non-disruptive'."""
    snap = _snap()
    g1 = snap["wave_sequencing"][0]                     # was pure MBB (make_before_break=[distA, distB])
    g1["hard_cutover"] = ["accSingle"]                  # now MIXED
    g1["sequence"] = "2 make-before-break (dual-homed) + 1 hard cutover (single-homed)"
    out = str(tmp_path / "mop_mixed.docx")
    write_mop_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "make-before-break the 2 dual-homed" in text          # dual-homed MBB sequenced FIRST
    assert "scheduled outage" in text.lower()                    # then the single-homed outage (not pure MBB)
    assert "brief outage" in text.lower()                        # rollback honesty for the single-homed members


def test_mop_validation_table_surfaces_severity_high_first(tmp_path):
    """N23: the §N.6 go/no-go table must surface each check's severity (already on every
    validation_plan item) and order High-first, so the most critical checks survive the per-wave
    display cap. The fixture lists a Low check BEFORE a High one, so a regression that keeps source
    order renders Low first."""
    out = str(tmp_path / "m.docx")
    write_mop_docx(out, _snap(), "Unit Test Fleet")
    d = Document(out)
    vtabs = [t for t in d.tables if t.rows and t.rows[0].cells[0].text == "Sev"]
    assert vtabs, "no §N.6 validation table with a leading 'Sev' column"
    w1 = next((t for t in vtabs if any("VLAN 10 gateway up" in c.text for r in t.rows for c in r.cells)), None)
    assert w1 is not None, "wave-1 validation table not found"
    sevs = [r.cells[0].text for r in w1.rows[1:]]
    assert "High" in sevs and "Low" in sevs and sevs.index("High") < sevs.index("Low"), sevs


def test_mop_port_mapping_surfaces_portchannel_redundancy(tmp_path):
    """N27: the §x.4 port-mapping must surface uplink port-channel membership so redundant uplink
    pairs are explicit — two uplink ports on one switch sharing a Po bundle read as a redundant pair."""
    snap = _snap()
    snap["interfaces"] = {
        "distA": {
            "Te1/1": {"switchport_mode": "Trunk", "cdp_neighbor": "core1", "port_channel": "Po40"},
            "Te1/2": {"switchport_mode": "Trunk", "cdp_neighbor": "core1", "port_channel": "Po40"},
        },
    }
    out = str(tmp_path / "m.docx")
    write_mop_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    assert text.count("[Po40]") >= 2          # both uplink legs tagged with their bundle


def test_mop_software_standardization_table(tmp_path):
    """N24: the MOP §2 states the per-platform software standardization (current train + disposition)
    from the software-advisory screening, so the cutover team knows which platforms need an upgrade."""
    snap = _snap()
    snap["software_risk"] = {"per_device": [
        {"host": "a", "platform": "nxos", "sw_version": "6.0", "train": "NX-OS 6.x", "train_band": "Replace/Upgrade"},
        {"host": "b", "platform": "ios", "sw_version": "15.2", "train": "IOS 15.2", "train_band": "Acceptable"}]}
    out = str(tmp_path / "m.docx")
    write_mop_docx(out, snap, "Unit Test Fleet")
    d = Document(out)
    heads = [p.text for p in d.paragraphs if p.style.name.startswith("Heading")]
    text = _all_text(d)
    assert any("Software" in h and ("standardization" in h.lower() or "image" in h.lower()) for h in heads), heads
    assert "NX-OS 6.x" in text and "Replace/Upgrade" in text     # current train + disposition surfaced


def test_mop_surfaces_blockers_and_rollback(tmp_path):
    out = str(tmp_path / "m.docx")
    write_mop_docx(out, _snap(), "Unit Test Fleet")
    text = _all_text(Document(out))
    # the high-severity remediation for acc1 is surfaced as a wave-2 blocker (Low items are not)
    assert "Add FHRP peer" in text
    assert "standby 20 ip 10.0.20.254" in text
    # the critical punch-list item touching acc1 surfaces; rollback + sign-off scaffolding present
    assert "VLAN 20 single gateway" in text
    assert "Rollback" in text and "fail back to the legacy leg" in text
    assert "Implementing engineer" in text
    # the change overview reconciles the wave count
    assert "2 wave(s)" in text or "sequenced into 2 wave(s)" in text


def test_mop_carries_document_furniture(tmp_path):
    """V3.23.150: AS-style front matter (Document Control with the 'Ready for service' caveat from
    the Cisco migration-service MOP definition) + the closing signature gate under the existing
    Post-Migration Acceptance section (no new H1)."""
    out = str(tmp_path / "m.docx")
    write_mop_docx(out, _snap(), "Unit Test Fleet")
    d = Document(out)
    h1 = [p.text for p in d.paragraphs if p.style.name == "Heading 1"]
    assert "Document Control" in h1
    assert "Document Acceptance" not in h1            # gate renders under Post-Migration Acceptance
    text = _all_text(d)
    assert "Ready for service" in text                # MOP-specific acceptance-criteria caveat
    assert "Customer network owner" in text           # closing signature roles
    assert "Assessment workbook (.xlsx)" in text      # related-documents cross-reference


def test_mop_port_mapping_raci_and_success_criteria(tmp_path):
    """V3.23.155: the migration-service MOP shape — §2.1 RACI, per-wave §x.4 port/interface mapping
    with <target> columns + evidence-derived staged config stubs, and a success criterion per step."""
    snap = _snap()
    snap["interfaces"] = {
        "acc1": {
            "Gi1/0/5": {"switchport_mode": "Access", "vlan": "20", "end_host_mac": "bbbb.0000.0001",
                        "description": "CAM-12"},
            "Gi1/0/1": {"switchport_mode": "Trunk", "cdp_neighbor": "distA"},
            "Gi1/0/9": {"switchport_mode": "Access", "vlan": "30"},   # no endpoint evidence
        },
    }
    out = str(tmp_path / "m.docx")
    write_mop_docx(out, snap, "Unit Test Fleet")
    d = Document(out)
    text = _all_text(d)
    h2 = [p.text for p in d.paragraphs if p.style.name == "Heading 2"]
    # RACI up front, per the AS migration-service ownership split
    assert "2.1 Responsibilities (RACI)" in h2
    assert "In-window execution of each step" in text and "Customer operations" in text
    # wave 2 (acc1) = section 4: mapping table reads the evidence, target columns stay placeholders
    assert "4.4 Port / interface mapping & staged configuration" in h2, h2
    assert "Gi1/0/5" in text and "<target device>" in text
    assert "Uplink → distA" in text
    assert "Gi1/0/9" not in text                        # no endpoint evidence → excluded
    assert "switchport access vlan 20" in text          # evidence-derived staged stub
    assert "description CAM-12" in text
    # renumbered chain + per-step success criteria
    assert "4.5 Cutover procedure" in h2 and "4.8 Sign-off" in h2
    assert "Success:" in text
    # wave 1 (distA/distB) has no interface evidence → the fresh-collection fallback renders
    assert "build the mapping from a fresh interface collection" in text


def test_mop_failsoft_without_python_docx(monkeypatch, tmp_path):
    import builtins, os
    real_import = builtins.__import__

    def _blocked(name, *a, **k):
        if name == "docx" or name.startswith("docx."):
            raise ImportError("simulated missing python-docx")
        return real_import(name, *a, **k)

    monkeypatch.setattr(builtins, "__import__", _blocked)
    out = str(tmp_path / "m.docx")
    write_mop_docx(out, _snap(), "Unit Test Fleet")   # must not raise
    assert not os.path.exists(out)


def test_mop_join_apportions_endpoints_to_subwave():
    """AUDIT FIX: a coupled sub-wave (a SLICE of a larger move-group) must NOT inherit the parent group's
    full endpoint total — apportion by the sub-wave's switch share. A full-coverage wave keeps the total."""
    from cisco_toolkit.mop import _join_group_records
    rbg = {"Group 1": {"group": "Group 1", "switches": ["a", "b", "c", "d"], "endpoints": 400,
                       "n_fail": 0, "n_warn": 2, "readiness": "CAUTION"}}
    r1, *_ = _join_group_records(["Group 1"], ["a"], rbg, {}, {}, {})       # 1 of 4 switches
    assert r1["endpoints"] == 100                                          # 400 * 1/4, NOT 400
    r4, *_ = _join_group_records(["Group 1"], ["a", "b", "c", "d"], rbg, {}, {}, {})
    assert r4["endpoints"] == 400                                          # full coverage -> full total


def test_mop_join_unions_cutover_split_across_groups():
    """A batch wave spans MULTIPLE groups; the cutover host split (make_before_break / hard_cutover)
    must UNION across all of them, not just the first/representative group. Else a wave mixing a
    dual-homed group (make-before-break) with a single-homed group (hard cutover) reads only the first
    group's split — taking the wrong procedure branch and silently dropping the other group's switches."""
    from cisco_toolkit.mop import _join_group_records
    seq_by_group = {"G1": {"make_before_break": ["a", "b"], "hard_cutover": []},
                    "G2": {"make_before_break": [], "hard_cutover": ["c"]}}
    rbg = {"G1": {"switches": ["a", "b"]}, "G2": {"switches": ["c"]}}
    _r, seq, _scen, _val = _join_group_records(["G1", "G2"], ["a", "b", "c"], rbg, seq_by_group, {}, {})
    assert set(seq.get("make_before_break") or []) == {"a", "b"}
    assert set(seq.get("hard_cutover") or []) == {"c"}     # the single-homed group is NOT dropped -> branch is correct


def _snap_waves():
    """_snap() + a canonical wave_plan: Group 1 (distA,distB) sliced into 2 coupled-subwaves; Group 2
    (acc1) as an independent-batch — 3 waves."""
    snap = _snap()
    snap["design_blueprint"] = {"target_state": {"wave_plan": {
        "waves": [
            {"wave": 1, "kind": "coupled-subwave", "n_switches": 1, "switches": ["distA"], "source_groups": [0]},
            {"wave": 2, "kind": "coupled-subwave", "n_switches": 1, "switches": ["distB"], "source_groups": [0]},
            {"wave": 3, "kind": "independent-batch", "n_switches": 1, "switches": ["acc1"], "source_groups": [1]},
        ],
        "n_waves": 3, "wave_cap": 40, "n_move_groups": 2, "largest_group": 2, "n_subdivided_groups": 1,
        "note": "candidate waves"}}}
    return snap


def test_mop_sections_keyed_to_wave_plan(tmp_path):
    """F2: with a canonical wave_plan the MOP emits one section per WAVE (wave-keyed title), not one per
    move-group/readiness row."""
    out = str(tmp_path / "m.docx")
    write_mop_docx(out, _snap_waves(), "Unit Test Fleet")
    h1 = [p.text for p in Document(out).paragraphs if p.style.name == "Heading 1"]
    assert any(t == "3. MOP — Wave 1 (coupled-subwave)" for t in h1), h1
    assert any(t == "4. MOP — Wave 2 (coupled-subwave)" for t in h1), h1
    assert any(t == "5. MOP — Wave 3 (independent-batch)" for t in h1), h1
    assert not any("MOP — Group 1" in t for t in h1)             # NOT keyed on move-groups when wave_plan present


def test_mop_coupled_subwave_surfaces_shared_vlan_caveat(tmp_path):
    """F2 honesty: coupled-subwave sections carry the shared-L2/VLAN SEQUENCE caveat; an independent-batch
    wave does not."""
    out = str(tmp_path / "m.docx")
    write_mop_docx(out, _snap_waves(), "Unit Test Fleet")
    d = Document(out)
    paras = [p.text for p in d.paragraphs]
    # the caveat appears (coupled sub-wave shares VLANs, is a SEQUENCE)
    assert any("coupled sub-wave" in t.lower() and "SEQUENCE" in t for t in paras), paras[:0] or "no caveat"
    assert sum(1 for t in paras if "Shared-L2 coordination" in t) >= 2   # one per coupled-subwave (waves 1 & 2)


def test_mop_reads_wave_plan_not_recompute(tmp_path):
    """F2 REFUTATION / SSOT mutation guard: a wave_plan grouping move_groups would NEVER produce
    (distA + acc1 together) must render verbatim — proving the MOP reads wave_plan, not a re-slice of
    move_groups."""
    snap = _snap()
    snap["design_blueprint"] = {"target_state": {"wave_plan": {
        "waves": [
            {"wave": 1, "kind": "independent-batch", "n_switches": 2, "switches": ["distA", "acc1"], "source_groups": [0, 1]},
            {"wave": 2, "kind": "coupled-subwave", "n_switches": 1, "switches": ["distB"], "source_groups": [0]},
        ],
        "n_waves": 2, "wave_cap": 40, "n_move_groups": 2, "largest_group": 2, "n_subdivided_groups": 0, "note": "x"}}}
    out = str(tmp_path / "m.docx")
    write_mop_docx(out, snap, "Unit Test Fleet")
    d = Document(out)
    h1 = [p.text for p in d.paragraphs if p.style.name == "Heading 1"]
    assert any("Wave 1 (independent-batch)" in t for t in h1) and any("Wave 2 (coupled-subwave)" in t for t in h1)
    assert not any("MOP — Group 1" in t for t in h1)             # not recomputed from move_groups
    assert "distA, acc1" in _all_text(d)                          # wave 1 co-locates distA+acc1 (impossible from a move-group recompute)


def test_mop_join_worst_readiness_uses_real_vocabulary():
    """REVIEW #2: the worst-verdict reducer must rank the engine's ACTUAL readiness values
    (NOT READY worse than CAUTION worse than READY). The old rank map keyed on a fictional
    vocabulary so 'NOT READY' fell to the default and lost to 'READY'/'CAUTION'."""
    from cisco_toolkit.mop import _join_group_records
    base = {"switches": ["x"], "endpoints": 10, "n_fail": 0, "n_warn": 0}
    rbg = {"G1": {**base, "group": "G1", "switches": ["a"], "readiness": "NOT READY", "n_fail": 1},
           "G2": {**base, "group": "G2", "switches": ["b"], "readiness": "READY"}}
    r, *_ = _join_group_records(["G1", "G2"], ["a", "b"], rbg, {}, {}, {})
    assert r["readiness"] == "NOT READY", r["readiness"]            # NOT 'READY'
    rbg2 = {"G1": rbg["G1"], "G2": {**rbg["G2"], "readiness": "CAUTION"}}
    r2, *_ = _join_group_records(["G1", "G2"], ["a", "b"], rbg2, {}, {}, {})
    assert r2["readiness"] == "NOT READY", r2["readiness"]          # NOT READY beats CAUTION
    rbg3 = {"G1": {**rbg["G1"], "readiness": "READY", "n_fail": 0}, "G2": {**rbg["G2"], "readiness": "CAUTION"}}
    r3, *_ = _join_group_records(["G1", "G2"], ["a", "b"], rbg3, {}, {}, {})
    assert r3["readiness"] == "CAUTION", r3["readiness"]            # CAUTION beats READY
    rbg4 = {"G1": {**rbg["G1"], "readiness": "READY", "n_fail": 0}, "G2": {**rbg["G2"], "readiness": "READY"}}
    r4, *_ = _join_group_records(["G1", "G2"], ["a", "b"], rbg4, {}, {}, {})
    assert r4["readiness"] == "READY", r4["readiness"]              # READY only when ALL ready


def test_mop_join_apportions_fail_warn_to_subwave():
    """REVIEW #1: n_fail/n_warn, like endpoints, must apportion to a coupled sub-wave's switch share —
    a 1-of-4 slice must not reprint the parent group's full blocking/warning counts."""
    from cisco_toolkit.mop import _join_group_records
    rbg = {"G1": {"group": "G1", "switches": ["a", "b", "c", "d"], "endpoints": 400,
                  "n_fail": 4, "n_warn": 8, "readiness": "NOT READY"}}
    r1, *_ = _join_group_records(["G1"], ["a"], rbg, {}, {}, {})            # 1 of 4
    assert (r1["n_fail"], r1["n_warn"]) == (1, 2), (r1["n_fail"], r1["n_warn"])
    r4, *_ = _join_group_records(["G1"], ["a", "b", "c", "d"], rbg, {}, {}, {})
    assert (r4["n_fail"], r4["n_warn"]) == (4, 8), (r4["n_fail"], r4["n_warn"])   # full coverage -> full


def test_mop_join_never_masks_a_blocker_via_rounding():
    """Coverage-honesty: a SINGLE blocking check on a large group must not round to 0 in a small coupled
    sub-wave (1*0.25=0.25 -> round 0) — that would show a NOT-READY wave with '0 blockers'. A positive
    share of a positive count contributes at least 1, so the gate can never be silently masked."""
    from cisco_toolkit.mop import _join_group_records
    rbg = {"G1": {"group": "G1", "switches": ["a", "b", "c", "d"], "n_fail": 1, "n_warn": 1,
                  "readiness": "NOT READY"}}
    r1, *_ = _join_group_records(["G1"], ["a"], rbg, {}, {}, {})            # 1 of 4 -> 0.25 must NOT mask to 0
    assert r1["n_fail"] == 1 and r1["n_warn"] == 1, (r1["n_fail"], r1["n_warn"])


def test_mop_batch_wave_surfaces_all_group_strategies_and_rollbacks(tmp_path):
    """REVIEW #4: an independent-batch wave bundling several move-groups must surface EVERY group's
    cutover strategy + rationale + rollback, not only the first non-empty group's (the old joiner kept
    only the first, silently dropping the rest)."""
    snap = _snap()
    snap["design_blueprint"] = {"target_state": {"wave_plan": {
        "waves": [{"wave": 1, "kind": "independent-batch", "n_switches": 3,
                   "switches": ["distA", "distB", "acc1"], "source_groups": [0, 1]}],
        "n_waves": 1, "wave_cap": 40, "n_move_groups": 2, "largest_group": 2,
        "n_subdivided_groups": 0, "note": "x"}}}
    out = str(tmp_path / "m.docx")
    write_mop_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    # Group 2's rationale + rollback (previously DROPPED because Group 1 was the first non-empty record)
    assert "single-homed access edge" in text                 # Group 2 rationale
    assert "re-cable to legacy port" in text                  # Group 2 rollback
    # Group 1's are still present
    assert "dual-homed — build beside" in text                # Group 1 rationale
    assert "fail back to the legacy leg" in text              # Group 1 rollback


def test_mop_endpoint_figure_labeled_as_mac_sum(tmp_path):
    """REVIEW #14: the per-wave endpoint figure is an apportioned per-switch-MAC sum; label it so
    (consistent with excel.py's B3 relabel), not the bare 'Endpoints'."""
    out = str(tmp_path / "m.docx")
    write_mop_docx(out, _snap(), "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "Endpoint MACs" in text


def test_mop_robust_to_malformed_uploaded_snapshot(tmp_path):
    """Robustness (webapp raw-upload path): a hand-edited / hostile snapshot can pass the webapp
    validator (which only checks a top-level 'devices' key) with sections that are the WRONG truthy type
    -- migration_readiness as an object, punchlist devices as a string, etc. The old `or {}`/`or []`
    idioms do NOT guard a truthy non-dict/non-list, so write_mop_docx char-iterated or raised
    AttributeError mid-render. With the shared docmeta coercers it must degrade and still produce a doc
    (matching engagement/ops/archreview + the webapp cutover hardening)."""
    out = str(tmp_path / "mop_malformed.docx")
    bad = {
        "devices": {"SW1": {"platform": "ios"}},
        "migration_readiness": {"group": "G1", "switches": ["SW1"]},   # object, not a list
        "move_groups": "oops",                                          # string, not a list
        "wave_sequencing": {"k": "v"},
        "migration_scenarios": "nope",
        "validation_plan": [1, 2, 3],
        "failure_impact": "str",
        "remediation_plan": {"items": "notalist"},
        "punchlist": [{"severity": "High", "category": "X", "title": "t", "devices": "SW1"}],
        "software_risk": "x",
        "executive_brief": 5,
        "interfaces": "bad",
    }
    write_mop_docx(out, bad, "Malformed Fleet")   # must NOT raise
    assert os.path.isfile(out) and os.path.getsize(out) > 0


def _evpn_mig(applicable=True):
    return {"applicable": applicable,
            "model_basis": "requirement-confirmed (fabric_operating_model = nxos-evpn)",
            "summary": "5 EVPN-migration guardrail(s)",
            "guardrails": ([] if not applicable else [
                {"id": "evpn-cut-single-active-l2-interconnect", "phase": "cutover-gate", "severity": "Critical",
                 "title": "Exactly ONE active L2 interconnect (the overlay will NOT break a loop)",
                 "basis": "2 vPC domain(s); STP present on 3 device(s)",
                 "detail": "Permit exactly ONE active Layer-2 connection ...", "source": "BRKDCN-2951"},
                {"id": "evpn-pre-nxos-1023-gateway-coexistence", "phase": "pre-cutover", "severity": "High",
                 "title": "NX-OS 10.2(3) HSRP↔Anycast-Gateway coexistence gate",
                 "basis": "1 of 2 NX-OS device(s) run a release below 10.2(3)",
                 "detail": "Any node that must hold both ... needs NX-OS >= 10.2(3).", "source": "BRKDCN-2951"},
            ])}


def test_mop_renders_evpn_guardrails_when_applicable(tmp_path):
    """The MOP's §2.2 surfaces the gated EVPN-migration guardrails (heading, the model basis, each guardrail +
    its primary source); the software subsection then renumbers to 2.3."""
    snap = _snap()
    snap["design_blueprint"] = {"evpn_migration": _evpn_mig()}
    snap["software_risk"] = {"per_device": [{"host": "distB", "train": "9.3", "disposition": "review"}]}
    out = str(tmp_path / "m.docx")
    write_mop_docx(out, snap, "Unit Test Fleet")
    d = Document(out)
    h2 = [p.text for p in d.paragraphs if p.style.name == "Heading 2"]
    assert any(t.startswith("2.2 EVPN-migration guardrails") for t in h2), h2
    assert any(t.startswith("2.3 Software") for t in h2), h2          # software renumbered below the EVPN block
    text = _all_text(d)
    assert "Exactly ONE active L2 interconnect" in text and "BRKDCN-2951" in text
    assert "requirement-confirmed" in text


def test_mop_silent_on_evpn_when_not_applicable(tmp_path):
    snap = _snap()
    snap["design_blueprint"] = {"evpn_migration": _evpn_mig(applicable=False)}
    out = str(tmp_path / "m.docx")
    write_mop_docx(out, snap, "Unit Test Fleet")
    h2 = [p.text for p in Document(out).paragraphs if p.style.name == "Heading 2"]
    assert not any("EVPN-migration guardrails" in t for t in h2)


def test_mop_check_counts_sum_to_group_ssot_across_subwaves():
    """[audit-2 #3] a coupled group's fail/warn check counts must SUM across its sub-waves to the group SSOT, not
    be share-with-floor double-counted (Meridian Group 1's true 3/4 rendered 7/7 across 7 sub-waves). Each check is
    attributed to the ONE sub-wave owning its canonical member switch."""
    from cisco_toolkit.mop import _join_group_records
    readiness = {"G1": {"group": "G1", "switches": ["swA", "swB", "swC", "swD"], "endpoints": 40,
                        "n_fail": 1, "n_warn": 1, "readiness": "NOT READY",
                        "checks": [{"status": "fail", "note": "sole gateway on swA, swB (no FHRP)"},
                                   {"status": "warn", "note": "err-disabled on swC, swD"},
                                   {"status": "pass", "note": "no inconsistent ports"}]}}
    w1 = _join_group_records(["G1"], ["swA", "swB"], readiness, {}, {}, {})[0]
    w2 = _join_group_records(["G1"], ["swC", "swD"], readiness, {}, {}, {})[0]
    assert w1["n_fail"] + w2["n_fail"] == 1 and w1["n_warn"] + w2["n_warn"] == 1   # each check counted ONCE
    full = _join_group_records(["G1"], ["swA", "swB", "swC", "swD"], readiness, {}, {}, {})[0]
    assert full["n_fail"] == 1 and full["n_warn"] == 1                              # whole group -> full SSOT


# ---- DE-01 P1: Executive-Summary BLUF (answer-first, front matter) ----------------------------
def test_mop_bluf_present_and_carries_wave_and_gate(tmp_path):
    """DE-01 / Cisco-AS: a Bottom-Line-Up-Front executive summary sits at the TOP of the MOP (before the
    numbered sections), and it carries the load-bearing decision facts pulled from the snapshot — the wave
    count, the go/no-go gate (worst readiness across the waves), the one-line rollback, and the window
    estimate. None invented: the fleet recommendation + top-gating are the snapshot's own."""
    out = str(tmp_path / "m.docx")
    write_mop_docx(out, _snap(), "Unit Test Fleet")
    d = Document(out)
    h1 = [p.text for p in d.paragraphs if p.style.name == "Heading 1"]
    assert any("Bottom Line Up Front" in t or "Bottom-Line-Up-Front" in t for t in h1), h1
    text = _all_text(d)
    # the BLUF names the wave count
    assert "2 wave(s)" in text
    # STRUCTURAL (non-vacuous): the BLUF go/no-go GATE ROW itself must carry the worst readiness — not just
    # "NOT READY" appearing somewhere in the doc (it also renders in every wave's scope table, which would
    # make a whole-doc substring check vacuous — adversarial-review finding, 2026-07-05).
    gate_rows = [[c.text.strip() for c in row.cells]
                 for tbl in d.tables for row in tbl.rows
                 if any("no-go gate" in c.text.lower() or "worst readiness" in c.text.lower()
                        for c in row.cells)]
    assert gate_rows, "BLUF go/no-go gate row not found in any table"
    assert any("NOT READY" in " ".join(cells) for cells in gate_rows), gate_rows
    # one-line rollback + a maintenance-window estimate line, both answer-first
    assert "Rollback in one line" in text or "Rollback (one line)" in text
    assert "Maintenance window" in text
    # the BLUF must appear BEFORE the numbered §1 Change Overview (it is genuinely up front)
    paras = [p.text for p in d.paragraphs]
    bluf_i = next(i for i, t in enumerate(paras)
                  if "Bottom Line Up Front" in t or "Bottom-Line-Up-Front" in t)
    co_i = next(i for i, t in enumerate(paras) if t == "1. Change Overview")
    assert bluf_i < co_i, (bluf_i, co_i)


def test_bluf_blocker_count_matches_the_section_it_cross_references(tmp_path):
    """One label, two numbers: the BLUF's 'Open blockers gating the change' counted migration-
    readiness FAIL checks while §1's Blockers column and each wave's §x.2 'Blockers to clear' are
    built from remediation_plan + Critical/High punch-list — and the BLUF pointed the approver at
    §x.2, the section that contradicted it. The discriminating fixture has 2 blockers and 1 failed
    readiness check, so a regression to the old source renders 1 where §1 totals 2."""
    out = str(tmp_path / "m_blockers.docx")
    write_mop_docx(out, _snap(), "Unit Test Fleet")
    d = Document(out)
    # the BLUF row (structural: the row itself, not a substring anywhere in the document)
    row = next(r for t in d.tables for r in t.rows
               if r.cells[0].text.strip().startswith("Open blockers"))
    detail = row.cells[1].text
    n_bluf = int(re.match(r"\s*(\d+)\s+Critical/High blocker", detail).group(1))
    # §1's overview table, the other renderer of the same set
    ov = next(t for t in d.tables if "Blockers" in [c.text.strip() for c in t.rows[0].cells])
    col = [c.text.strip() for c in ov.rows[0].cells].index("Blockers")
    n_overview = sum(int(r.cells[col].text) for r in ov.rows[1:])
    assert n_bluf == n_overview == 2, (detail, n_overview)
    # the readiness-FAIL count is still reported — under its OWN label, never as "blockers"
    assert "1 migration-readiness check(s) FAIL" in detail
    assert n_bluf != 1, "the fixture must discriminate the two sets"
    # and the cross-reference now points at a section that agrees with it
    assert "Blockers-to-clear subsection" in detail or "Blockers-to-clear" in detail


def test_mop_bluf_window_estimate_is_coverage_honest(tmp_path):
    """The window estimate is a derived planning figure, not a collected fact — it must be labelled as an
    estimate to confirm with the change owner, never presented as an authoritative collected number.

    NON-VACUITY (mutation-proved): the old whole-document form found "estimate" in the row's own
    LABEL ("Maintenance window (estimate)") and "confirm with the change owner" in the per-wave
    rollback-trigger table, so gutting `_window_estimate()` down to a bare "2-4 hours per
    maintenance window." — the derived figure stated as an authoritative number, which is the one
    thing this test names — left it GREEN. Assert the row's VALUE."""
    out = str(tmp_path / "m.docx")
    write_mop_docx(out, _snap(), "Unit Test Fleet")
    est = _row_values(Document(out), "Maintenance window (estimate)")
    assert len(est) == 1, f"the BLUF window row moved or duplicated: {est}"
    low = est[0].lower()
    assert "estimate" in low and "confirm with the change owner" in low, (
        f"the window figure reads as a collected number, not a planning estimate: {est[0]!r}")


# ---- DE-01 P1: per-step quantified rollback triggers -----------------------------------------
def test_mop_each_wave_has_quantified_rollback_trigger(tmp_path):
    """Cisco-AS standard: every wave carries an EXPLICIT, QUANTIFIED rollback trigger — not 'if something
    breaks' but a boolean condition (e.g. '>0.1% validation failures', 'no convergence within N minutes',
    'any Critical check fails'). The §x.7 rollback-triggers block must render a quantified condition per
    wave, ahead of the rollback procedure."""
    out = str(tmp_path / "m.docx")
    write_mop_docx(out, _snap(), "Unit Test Fleet")
    d = Document(out)
    h2 = [p.text for p in d.paragraphs if p.style.name == "Heading 2"]
    # a dedicated rollback-triggers heading per wave (waves are §3 and §4)
    assert any("Rollback triggers" in t for t in h2), h2
    text = _all_text(d)
    # QUANTIFIED default conditions (Cisco-AS), with the confirm-with-owner caveat where the snapshot
    # has no device-specific threshold
    assert "any Critical" in text and "validation" in text.lower()
    assert "convergence" in text.lower()
    assert "confirm with the change owner" in text.lower()


def test_mop_rollback_trigger_is_a_boolean_per_step(tmp_path):
    """Each rollback trigger reads as an explicit boolean 'Roll back if:' condition, so the in-window
    engineer has an unambiguous abort test rather than prose judgement."""
    out = str(tmp_path / "m.docx")
    write_mop_docx(out, _snap(), "Unit Test Fleet")
    text = _all_text(Document(out))
    assert text.count("Roll back if") >= 2      # at least one per wave (two waves in the fixture)


# ---- DE-01 P1: pre-implementation checklist (evidence-gated vs attested) ----------------------
def test_mop_pre_implementation_checklist_present(tmp_path):
    """Cisco-AS: a per-window pre-implementation checklist gates the window — config backup taken, OOB
    reachable, window confirmed, CAB approved, rollback tested, TAC pre-opened for high-risk. The
    snapshot-checkable preconditions (OOB port evidence, backup) are evidence-gated; the human-attested
    ones are clearly distinguished."""
    out = str(tmp_path / "m.docx")
    write_mop_docx(out, _snap(), "Unit Test Fleet")
    d = Document(out)
    h2 = [p.text for p in d.paragraphs if p.style.name == "Heading 2"]
    assert any("Pre-implementation checklist" in t for t in h2), h2
    text = _all_text(d)
    # the standard AS preconditions
    for token in ("Config backup", "Out-of-band", "CAB", "Rollback tested", "maintenance window"):
        assert token in text or token.lower() in text.lower(), token
    # the evidence-gated vs human-attested distinction is explicit
    assert "Evidence" in text and "Attested" in text


def test_mop_checklist_evidence_gates_oob_from_snapshot(tmp_path):
    """The OOB-reachability precondition is EVIDENCE-GATED: when the snapshot carries a management-port for
    a wave device it reads as evidenced (not merely attested); with no mgmt-port evidence it is [NOT
    OBSERVED] — coverage-honest, never assumed reachable."""
    snap = _snap()
    snap["interfaces"] = {
        "distA": {"mgmt0": {"switchport_mode": "", "mgmt_ip": "10.99.0.5"}},
    }
    out = str(tmp_path / "m.docx")
    write_mop_docx(out, snap, "Unit Test Fleet")
    d = Document(out)
    # NON-VACUITY (mutation-proved): `"[NOT OBSERVED]" in _all_text(doc)` is true of EVERY MOP —
    # the checklist's own intro quotes the convention verbatim — so replacing this row's abstention
    # with "Confirm OOB reachability before the window." (i.e. assume it) left the test GREEN while
    # the precondition stopped declaring the blind spot. Assert the ROW, per wave.
    oob = _row_values(d, "Out-of-band / console access to every device in scope", col=2)
    assert len(oob) == 2, f"expected one OOB precondition per wave, got {len(oob)}: {oob}"
    # wave 1 (distA) has mgmt-port evidence -> the OOB precondition cites it
    assert "mgmt0" in oob[0] or "10.99.0.5" in oob[0], oob[0]
    assert "[NOT OBSERVED]" not in oob[0], "an evidenced wave must not abstain"
    # wave 2 (acc1) has NO mgmt-port evidence -> coverage-honest not-observed marker
    assert oob[1].startswith("[NOT OBSERVED]"), (
        f"wave 2 has no management-port evidence but its OOB precondition does not abstain — "
        f"the window would open on an assumed-reachable console path: {oob[1]!r}")
    assert "do NOT assume" in oob[1]


def test_mop_checklist_tac_preopen_for_high_risk_wave(tmp_path):
    """A high-risk wave (NOT-READY / has Critical blockers) surfaces the 'pre-open a Cisco TAC case'
    precondition; a low-risk wave does not force it. In the fixture Group 2 (acc1) is NOT READY with a
    Critical punch-list item -> TAC pre-open surfaces."""
    out = str(tmp_path / "m.docx")
    write_mop_docx(out, _snap(), "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "TAC" in text and "pre-open" in text.lower()


# ---- DE-01 P1: communications / escalation plan ----------------------------------------------
def test_mop_comms_and_escalation_plan_present(tmp_path):
    """Cisco-AS: a communications / escalation plan — roles (change owner, executor, verifier,
    escalation), a T-minus comms cadence, and the escalation matrix (incl. Cisco TAC pre-open for
    high-risk windows)."""
    out = str(tmp_path / "m.docx")
    write_mop_docx(out, _snap(), "Unit Test Fleet")
    d = Document(out)
    h2 = [p.text for p in d.paragraphs if p.style.name == "Heading 2"]
    assert any("Communications" in t and "escalation" in t.lower() for t in h2), h2
    text = _all_text(d)
    # T-minus cadence
    assert "T-minus" in text or "T-" in text
    # escalation matrix roles + the TAC escalation tier
    assert "Change owner" in text and "Escalation" in text
    assert "Cisco TAC" in text


# ---- DE-01 P1: PRE / DURING / POST phase labels (roadmap C2) + verify AND rollback (C3) -------
def test_mop_procedure_steps_carry_phase_labels(tmp_path):
    """Roadmap C2: the cutover procedure labels its steps PRE / DURING / POST so the reader can see the
    window structure at a glance."""
    out = str(tmp_path / "m.docx")
    write_mop_docx(out, _snap(), "Unit Test Fleet")
    text = _all_text(Document(out))
    assert "[PRE]" in text and "[DURING]" in text and "[POST]" in text


def test_mop_every_wave_has_verify_and_rollback(tmp_path):
    """Roadmap C3: every wave has at least one VERIFY (post-cutover validation) AND one ROLLBACK section.
    Guards against a wave that renders a procedure with no go/no-go or no abort path."""
    out = str(tmp_path / "m.docx")
    write_mop_docx(out, _snap(), "Unit Test Fleet")
    h2 = [p.text for p in Document(out).paragraphs if p.style.name == "Heading 2"]
    # two waves -> two validation (verify) + two rollback headings
    assert sum(1 for t in h2 if "Post-cutover validation" in t) >= 2
    assert sum(1 for t in h2 if t.endswith("Rollback")) >= 2


# ---- coverage-honesty: not-assessable when evidence absent -----------------------------------
def test_mop_bluf_not_assessable_when_evidence_absent(tmp_path):
    """Coverage-honesty: with no fleet recommendation and no top-gating in the snapshot, the BLUF must
    DECLARE the missing facts not-assessable ([NOT OBSERVED]) rather than fabricate a recommendation."""
    snap = _snap()
    snap["migration_scenarios"] = {"per_group": snap["migration_scenarios"]["per_group"]}   # drop fleet_recommendation
    snap["executive_brief"] = {}                                                             # drop top_gating
    out = str(tmp_path / "m.docx")
    write_mop_docx(out, snap, "Unit Test Fleet")
    d = Document(out)
    # NON-VACUITY (mutation-proved): `"[NOT OBSERVED]" in _all_text(doc)` is satisfied by the
    # BLUF's own intro paragraph, which quotes the convention, in EVERY MOP — the token is present
    # 11 times before this test strips anything. Changing _write_bluf's `_v` fallback to a silent
    # em-dash (the classic false-health blank) left it GREEN. Assert the two rows by name.
    for label in ("Fleet-level recommendation", "Top gating item to clear first"):
        assert _row_values(d, label) == ["[NOT OBSERVED]"], (
            f"the BLUF's {label!r} row does not declare itself unassessable; got "
            f"{_row_values(d, label)!r}")
    # POSITIVE CONTROL: when the snapshot DOES publish them, the same rows carry the fact — so the
    # assertion above is about abstention, not about the rows being empty in every document.
    full = str(tmp_path / "full.docx")
    write_mop_docx(full, _snap(), "Unit Test Fleet")
    rec, = _row_values(Document(full), "Fleet-level recommendation")
    assert rec == "Cut the dual-homed core first, then the single-homed access edge."
    # even with no recommendation the structural gate facts (wave count) still render
    assert "2 wave(s)" in _all_text(d)


# ---- EVPN gating: rollback triggers cite the evpn abort conditions only when EVPN is the target
def test_mop_rollback_triggers_surface_evpn_abort_conditions_when_applicable(tmp_path):
    """Reference-profile specificity: when the target fabric is NX-OS VXLAN-EVPN, the quantified rollback triggers
    surface the primary-source EVPN abort conditions (gateway-MAC black-hole, L2 loop, control-plane not
    Established). Gated: a non-EVPN MOP does not."""
    snap = _snap()
    snap["design_blueprint"] = {"evpn_migration": _evpn_mig()}
    out = str(tmp_path / "m.docx")
    write_mop_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    # the EVPN rollback guardrail's quantified abort conditions appear in the triggers block
    assert "control plane" in text.lower() and "Established" in text
    assert "black-hole" in text.lower() or "loop" in text.lower()


def test_mop_rollback_triggers_silent_on_evpn_when_not_the_target(tmp_path):
    """Refutation / gate: a non-EVPN MOP's rollback triggers do NOT reference the EVPN-specific abort
    conditions (they are gated on the target fabric).

    NON-VACUITY (mutation-proved, 2026-07-28). This refuted two tokens the rollback-trigger block
    never emits: "DAG MAC" appears nowhere in any rendered MOP (it lives in evpn_migration.py), and
    "Anycast" comes from a DIFFERENT gated block (the §2.2 guardrail section). Ungating the trigger
    block itself — mop.py's `if evpn_on:` inside the rollback-trigger builder, set to `if True:` —
    added a fabricated "EVPN cutover abort (target = NX-OS VXLAN-EVPN)" trigger to every non-EVPN
    MOP and this test stayed GREEN. Refute the strings the block ACTUALLY renders."""
    out = str(tmp_path / "m.docx")
    write_mop_docx(out, _snap(), "Unit Test Fleet")   # _snap() has no evpn_migration
    text = _all_text(Document(out))
    for token in ("EVPN cutover abort", "target = NX-OS VXLAN-EVPN", "BRKDCN-2951",
                  "gateway-MAC mismatch", "Anycast"):
        assert token not in text, (
            f"a non-EVPN MOP surfaced the EVPN-gated string {token!r} — the trigger set is no "
            f"longer gated on the target fabric")
    # ...and the same tokens ARE present once EVPN is the target, so the refutation above is
    # discriminating rather than naming strings nothing ever emits (which is what it used to do).
    ev = _snap()
    ev["design_blueprint"] = {"evpn_migration": _evpn_mig()}
    out2 = str(tmp_path / "evpn.docx")
    write_mop_docx(out2, ev, "Unit Test Fleet")
    evpn_text = _all_text(Document(out2))
    for token in ("EVPN cutover abort", "target = NX-OS VXLAN-EVPN", "BRKDCN-2951"):
        assert token in evpn_text, f"{token!r} is refuted above but nothing ever emits it"


def test_mop_subwave_strategy_and_split_are_slice_scoped(tmp_path):
    """QA row-13: every ≤40-switch sub-wave row of the 251-switch parent repeated the PARENT's
    '107 hard cutover + 144 make-before-break' (impossible on a 40-switch row), and §wi.5 sequenced
    parent-scale counts. The joined host split must be SLICED to this wave's members: the §1 Strategy
    cell derives per-row counts, the slices sum to the parent split, and the parent-scale text is gone."""
    snap = _snap()
    sw = [f"s{i}" for i in range(1, 7)]     # one parent group of 6: 4 dual-homed + 2 single-homed
    snap["devices"] = {h: {"platform": "ios"} for h in sw}
    snap["move_groups"] = [{"switches": sw, "endpoints": 60}]
    snap["migration_readiness"] = [{"group": "Group 1", "switches": sw, "endpoints": 60,
                                    "readiness": "CAUTION", "n_fail": 0, "n_warn": 1, "checks": []}]
    snap["wave_sequencing"] = [{
        "group": "Group 1", "make_before_break": ["s1", "s2", "s4", "s5"],
        "hard_cutover": ["s3", "s6"], "hard_cutover_endpoints": 20,
        "sequence": "2 hard cutover (single-homed) + 4 make-before-break (dual-homed)"}]
    snap["migration_scenarios"]["per_group"] = [
        {"group": "Group 1", "switches": 6, "endpoints": 60, "readiness": "CAUTION",
         "recommended_scenario": "hybrid", "rationale": "mixed homing",
         "playbook": {"pre": "stage", "validate": "prove", "rollback": "fail back"}}]
    snap["validation_plan"]["by_wave"] = {"Group 1": snap["validation_plan"]["by_wave"]["Group 1"]}
    snap["design_blueprint"] = {"target_state": {"wave_plan": {
        "n_waves": 2, "n_move_groups": 1, "largest_group": 6, "wave_cap": 3,
        "waves": [{"wave": 1, "kind": "coupled-subwave", "switches": ["s1", "s2", "s3"]},
                  {"wave": 2, "kind": "coupled-subwave", "switches": ["s4", "s5", "s6"]}]}}}
    out = str(tmp_path / "m.docx")
    write_mop_docx(out, snap, "Unit Test Fleet")
    text = _all_text(Document(out))
    # per-slice derived split (1 hard + 2 mbb), on BOTH the §1 row and each §wi.1 scope table
    assert text.count("1 hard cutover (single-homed) + 2 make-before-break (dual-homed)") >= 2
    # the parent-scale free text must appear NOWHERE (that was the impossible-per-row claim)
    assert "2 hard cutover (single-homed) + 4 make-before-break (dual-homed)" not in text
    # the mixed-homing procedure sequences THIS slice's counts, not the parent's
    assert "make-before-break the 2 dual-homed" in text
    assert "the 1 single-homed switch(es)" in text
    assert "the 4 dual-homed" not in text


# ---- stored-DoS: INNER (row / per-element) poison, not just top-level sections ---------------------
# test_mop_robust_to_malformed_uploaded_snapshot (above) poisons whole SECTIONS, which the module already
# routes through _as_dict/_as_list at :329-338. It cannot reach the residue: a WELL-FORMED section whose
# inner ROW field (G2) or per-element MAP value (G3) is a truthy scalar. `x.get("k") or []` catches only
# FALSY values, so `5` survives and detonates on the next dereference (len()/iteration/slice/set()/.get()/
# .items()). The webapp stores an uploaded snapshot before rendering it (its only validation is
# `isinstance(snap, dict) and "devices" in snap`), so each of these is a STORED availability DoS: every
# later GET /api/snapshots/{id}/deliverable/mop 500s. One value poisoned per case -- poisoning several at
# once lets an upstream guard drop the row and makes the case vacuous.

def _poison(path, value, base=None):
    """_snap() (or `base`) with exactly ONE inner value replaced. `path` is a key/index chain."""
    snap = base if base is not None else _snap()
    node = snap
    for key in path[:-1]:
        node = node[key]
    node[path[-1]] = value
    return snap


def _batch_wave_snap():
    """_snap() + an independent-batch wave that bundles BOTH move-groups into one wave, so the per-group
    rollback loop (which reads each sibling group's playbook) actually runs."""
    snap = _snap()
    snap["design_blueprint"] = {"target_state": {"wave_plan": {
        "waves": [{"wave": 1, "kind": "independent-batch", "n_switches": 3,
                   "switches": ["distA", "distB", "acc1"], "source_groups": [0, 1]}],
        "n_waves": 1, "wave_cap": 40, "n_move_groups": 2, "largest_group": 2,
        "n_subdivided_groups": 0, "note": "x"}}}
    return snap


# name -> snapshot. Each entry names the raw read it detonates on in cisco_toolkit/mop.py.
INNER_POISON = {
    # G2 -- a row of a well-formed list-of-dicts section
    "readiness_switches_scalar":  _poison(["migration_readiness", 0, "switches"], 5),          # members_all.update(5)/len(5)
    "readiness_checks_scalar":    _poison(["migration_readiness", 0, "checks"], 5),            # for c in 5
    "seq_make_before_break_scalar": _poison(["wave_sequencing", 0, "make_before_break"], 5),   # for _h in 5
    "seq_hard_cutover_scalar":    _poison(["wave_sequencing", 0, "hard_cutover"], 5),          # for _h in 5
    "punchlist_devices_scalar":   _poison(["punchlist", 0, "devices"], 5),                     # set(5)
    "remediation_commands_scalar": _poison(["remediation_plan", "items", 0, "commands"], 5),   # 5[:8]
    "scenario_playbook_scalar":   _poison(["migration_scenarios", "per_group", 0, "playbook"], 5),  # 5.get("pre")
    # G3 -- a per-element value of a well-formed core map
    "val_by_wave_scalar":         _poison(["validation_plan", "by_wave", "Group 1"], 5),       # for v in 5
    "interfaces_host_scalar":     _poison(["interfaces"], {"distA": 5}),                       # 5.items()
    "interfaces_port_scalar":     _poison(["interfaces"], {"distA": {"Gi1/0/1": 5}}),          # 5.get("switchport_mode")
    # the multi-group rollback loop: poison the SIBLING group's playbook so the representative one stays
    # well-formed and the crash isolates to the per-group read, not the shared one.
    "batch_sibling_playbook_scalar": _poison(["migration_scenarios", "per_group", 1, "playbook"], 5,
                                             base=_batch_wave_snap()),                         # 5.get("rollback")

    # G4 -- a well-formed CONTAINER whose ELEMENT is a scalar. `as_list` guards the container's TYPE and
    # is structurally blind to what is inside it, so these all survive the G2/G3 coercers above and
    # detonate on the first join / add_paragraph / .get() of the element.
    "gating_element_scalar":       _poison(["executive_brief", "top_gating", 0], 5),           # add_paragraph(5)
    "readiness_switch_element_scalar": _poison(["migration_readiness", 0, "switches", 0], 5),  # ", ".join([5,...])
    # the OTHER switch producer: _wave_sections() reads wave_plan.waves[].switches, _waves() reads
    # migration_readiness[].switches -- one coerced producer would leave the other open.
    "wave_plan_switch_element_scalar": _poison(
        ["design_blueprint", "target_state", "wave_plan", "waves", 0, "switches", 0], 5,
        base=_batch_wave_snap()),                                                              # ", ".join([5,...])
    "remediation_command_element_scalar": _poison(
        ["remediation_plan", "items", 0, "commands", 0], 5),                                   # add_paragraph(5)
    "val_by_wave_element_scalar":  _poison(["validation_plan", "by_wave", "Group 1", 0], 5),   # 5.get("device")
    # a STRING element is the nastier half of the same site: it is truthy AND joinable, so every
    # coercion that only rescues numbers still leaves `"x".get("device")` raising.
    "val_by_wave_element_string":  _poison(["validation_plan", "by_wave", "Group 1", 0], "x"),  # "x".get("device")

    # G5 -- a string-expected LEAF holding a truthy scalar. The bare `if playbook.get(k):` truthiness gate
    # admits it, then a string METHOD is called on it (`.rstrip(". ")`) -- as_dict guarded the playbook,
    # not the leaf.
    "playbook_pre_leaf_scalar":    _poison(["migration_scenarios", "per_group", 0, "playbook", "pre"], 5),
    "playbook_validate_leaf_scalar": _poison(["migration_scenarios", "per_group", 0, "playbook", "validate"], 5),
    "playbook_rollback_leaf_scalar": _poison(["migration_scenarios", "per_group", 0, "playbook", "rollback"], 5),
    # the multi-group rollback loop reads each SIBLING group's playbook.rollback through a DIFFERENT
    # expression than the single-group branch above -- poison the sibling so the representative stays good.
    "batch_sibling_rollback_leaf_scalar": _poison(
        ["migration_scenarios", "per_group", 1, "playbook", "rollback"], 5, base=_batch_wave_snap()),
}


@pytest.mark.parametrize("name", list(INNER_POISON))
def test_mop_robust_to_inner_scalar_poison(tmp_path, name):
    """Every one of these raises against the pre-guard module (verified by importing
    `git show origin/main:cisco_toolkit/mop.py`); with the docmeta coercers each must degrade and still
    produce a document."""
    out = str(tmp_path / f"mop_{name}.docx")
    write_mop_docx(out, INNER_POISON[name], "Poisoned Fleet")   # must NOT raise
    assert os.path.isfile(out) and os.path.getsize(out) > 0
    # non-degenerate: the wave sections still render (the guard degrades the poisoned value, not the doc)
    h1 = [p.text for p in Document(out).paragraphs if p.style.name == "Heading 1"]
    assert any(t.startswith("3. MOP —") for t in h1), h1


def test_mop_wellformed_render_unchanged_by_the_guards(tmp_path):
    """Behaviour-preservation anchor: `as_list(x)` is identical to `(x or [])` for every list/None input,
    so a well-formed snapshot must render EXACTLY as before. Pins the full text of both the default and
    the wave_plan snapshots so a coercer that silently swallowed good data would fail here."""
    for tag, snap in (("default", _snap()), ("waves", _snap_waves()), ("batch", _batch_wave_snap())):
        out = str(tmp_path / f"m_{tag}.docx")
        write_mop_docx(out, snap, "Unit Test Fleet")
        doc = Document(out)
        text = _all_text(doc)
        # the §x.1 "Devices in scope" join is the element-coerced site -- every wave's switch list must
        # still reach it verbatim (a coercer that dropped or mangled an element would shrink this union)
        scope = [r.cells[1].text for t in doc.tables for r in t.rows
                 if r.cells[0].text == "Devices in scope"]
        assert scope, f"{tag}: no 'Devices in scope' row"
        assert {s for cell in scope for s in cell.split(", ")} == {"distA", "distB", "acc1"}, (tag, scope)
        # every inner value the guards now route through a coercer still reaches the document
        assert "distA" in text and "acc1" in text                       # readiness/sequencing switches
        assert "stage target beside legacy" in text                     # playbook.pre
        assert "cut one leg, prove forwarding" in text                  # playbook.validate
        assert "fail back to the legacy leg" in text                    # playbook.rollback
        assert "standby 20 ip 10.0.20.254" in text                      # remediation_plan.items[].commands
        assert "VLAN 20 single gateway" in text                         # punchlist row matched via devices[]
        assert "show banner login" in text                              # validation_plan.by_wave[]
        assert "Resolve the VLAN 20 single-gateway exposure" in text    # executive_brief.top_gating[]


# ---- E2E through the real AssessHub route (the attacker's actual path) -----------------------------

@pytest.fixture()
def hub_client(tmp_path):
    """AssessHub TestClient. base_url=localhost so the default Host passes the no-token DNS-rebinding
    guard; raise_server_exceptions=False so a stored DoS surfaces as a 500 RESPONSE (what a real client
    sees) instead of a raised exception."""
    pytest.importorskip("fastapi")
    pytest.importorskip("httpx")
    from fastapi.testclient import TestClient
    root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    if root not in sys.path:                      # webapp is a namespace package off the repo root
        sys.path.insert(0, root)
    from webapp.backend.app import create_app
    with TestClient(create_app(db_path=str(tmp_path / "hub.db")), base_url="http://localhost",
                    raise_server_exceptions=False) as c:
        yield c


def _upload_and_get_mop(client, snap):
    """POST the snapshot (the webapp validates only `isinstance(snap, dict) and 'devices' in snap`, so
    the poison is ACCEPTED and STORED), then GET the MOP deliverable rendered from it."""
    cid = client.post("/api/campaigns", json={"name": "c"}).json()["id"]
    r = client.post(f"/api/campaigns/{cid}/snapshots",
                    files={"file": ("s.json", json.dumps(snap).encode(), "application/json")},
                    data={"label": "s"})
    assert r.status_code == 201, r.text            # stored first -> this is a STORED availability DoS
    sid = r.json()["id"]
    return client.get(f"/api/snapshots/{sid}/deliverable/mop")


@pytest.mark.parametrize("name", list(INNER_POISON))
def test_mop_deliverable_route_survives_inner_scalar_poison(hub_client, name):
    """The route re-raises after unlinking its temp file, so a raise inside write_mop_docx is an HTTP
    500 on EVERY later GET for that stored snapshot. Each case 500s against the pre-guard module."""
    r = _upload_and_get_mop(hub_client, INNER_POISON[name])
    assert r.status_code == 200, f"{name} -> {r.status_code}: {r.text[:300]}"
    assert len(r.content) > 0


def test_mop_deliverable_route_still_serves_a_wellformed_snapshot(hub_client):
    """Non-vacuity anchor for the route itself: a good snapshot downloads a real .docx."""
    r = _upload_and_get_mop(hub_client, _snap())
    assert r.status_code == 200, r.text
    assert r.content[:2] == b"PK" and len(r.content) > 10_000     # a populated OOXML package
