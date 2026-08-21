"""Focused regressions for exact current-run custody and fail-closed pipeline inputs."""
import copy
import hashlib
from importlib.resources import as_file, files
import json
import logging
import re
import sys
from types import SimpleNamespace

import pytest
from openpyxl import Workbook

import COLLECT_PARSE_V3_23_0 as cp
from cisco_toolkit import (
    assertions,
    cmdio,
    gate_state,
    input_custody,
    manifest,
    traffic_assurance,
)

def _args(**changes):
    base = {
        "requirements": "",
        "golden_config": None,
        "import_inventory": None,
        "scenario": None,
        "path_intents": None,
        "traffic_intents": None,
        "assert_pack": None,
        "flow_src": None,
        "flow_dst": None,
    }
    base.update(changes)
    return SimpleNamespace(**base)


def _finalize_ctx(tmp_path):
    """A minimal healthy producer boundary; callers inject one failure at a time."""
    cp._PHASE_TIMINGS.clear()
    gate_state.reset_verdicts()
    out = tmp_path / "Assessment.xlsx"
    wb = Workbook()
    wb.active.append(["Hostname", "Port", "Status"])
    cp._start_run_custody(str(out))
    cp._RUN_CUSTODY["generated_at"] = "2026-01-01T00:00:00"
    ctx = SimpleNamespace(
        out_xlsx=str(out),
        snap_path=str(tmp_path / "Assessment.snapshot.json"),
        snap_dict={"collected_at": "2026-01-01T00:00:00"},
        root_dir=str(tmp_path),
        args=SimpleNamespace(redact_collection=False),
        all_devices_meta=[],
        all_cmd_to_files={},
        workers=1,
        wb=wb,
    )
    return ctx


def _receipt_finalize_ctx(tmp_path):
    ctx = _finalize_ctx(tmp_path)
    ctx.snap_dict = {
        "script_version": "custody-receipt-test/1",
        "collected_at": "2026-01-01T00:00:00",
        "devices": {"sw1": {}},
        "interfaces": {"sw1": {}},
    }
    ctx.protocol_assurance_receipt_enabled = True
    ctx.protocol_assurance_export_path = str(
        tmp_path / "Assessment.protocol-assurance.json")
    ctx.html_output_path = ""
    ctx.runbook_output_path = ""
    ctx.mop_output_path = ""
    ctx.flow_paths = {}
    return ctx


def _receipt_authority(tmp_path):
    from cisco_toolkit.protocol_assurance import bind_snapshot_json_bytes
    from webapp.backend.protocol_portfolio import (
        build_protocol_single_snapshot_bundle,
        canonical_export_bytes,
        emitted_snapshot_binding,
    )

    cp._start_run_custody(str(tmp_path / "Assessment.xlsx"))
    snapshot_path = tmp_path / "Assessment.snapshot.json"
    snapshot_bytes = json.dumps({
        "script_version": "custody-receipt-test/1",
        "devices": {"sw1": {}},
    }, separators=(",", ":")).encode("utf-8")
    snapshot_path.write_bytes(snapshot_bytes)
    bound = bind_snapshot_json_bytes(snapshot_bytes)
    binding = emitted_snapshot_binding(bound, label=snapshot_path.name)
    bundle = build_protocol_single_snapshot_bundle(bound, binding, subject_cap=1)
    export_path = tmp_path / "Assessment.protocol-assurance.json"
    export_path.write_bytes(canonical_export_bytes(bundle["complete_export"]))
    cp._RUN_CUSTODY["snapshot_authority"] = {
        "source": binding["source"],
        "name": snapshot_path.name,
        "sha256": binding["sha256"],
        "bytes": binding["bytes"],
    }
    cp._ACTIVE_INTEGRITY_SNAPSHOT = None
    return snapshot_path, export_path, binding, bundle


def test_manifest_uses_exact_current_run_registry_and_excludes_stale_siblings(tmp_path):
    workbook = tmp_path / "Assessment.xlsx"
    Workbook().save(workbook)
    stale = tmp_path / "Assessment_design.docx"
    stale.write_bytes(b"older run; deliberately preserved")
    unrelated = tmp_path / "Assessment_operator_notes.txt"
    unrelated.write_text("not a producer-owned output", encoding="utf-8")

    cp._start_run_custody(str(workbook))
    cp._register_artifact(str(workbook), kind="xlsx", source="current workbook writer")
    gate_state.reset_verdicts()
    run = cp.build_run_manifest(str(workbook), {"collected_at": "2026-01-01T00:00:00"})

    assert [row["name"] for row in run["artifacts"]] == ["Assessment.xlsx"]
    excluded = run["metadata"]["excluded_stale_outputs"]
    assert excluded == [{
        "name": stale.name,
        "reason": "pre-existing or unregistered output; excluded from this run's seal",
        "existed_before_run": True,
    }]
    assert unrelated.name not in json.dumps(run), "a prefix sibling leaked into exact custody"
    assert stale.exists() and unrelated.exists(), "custody must preserve, never clean, older files"


def test_partial_writer_output_is_left_unsealed(tmp_path):
    workbook = tmp_path / "Assessment.xlsx"
    Workbook().save(workbook)
    partial = tmp_path / "Assessment_runbook.docx"
    partial.write_bytes(b"PK\x03\x04partial")
    cp._start_run_custody(str(workbook))

    with pytest.raises(ValueError, match="structural validation"):
        cp._register_artifact(str(partial), kind="docx", source="failed writer")
    assert partial.exists(), "invalid output is forensic evidence and must not be deleted"
    assert cp._RUN_CUSTODY["artifacts"] == {}


def test_writer_noop_cannot_admit_a_preexisting_artifact(tmp_path):
    workbook = tmp_path / "Assessment.xlsx"
    Workbook().save(workbook)
    cp._start_run_custody(str(workbook))

    assert cp._emit_artifact(
        "No-op workbook writer", str(workbook), "xlsx", lambda: None) is False
    assert cp._RUN_CUSTODY["artifacts"] == {}
    assert workbook.exists(), "the stale file must be preserved, not cleaned"


def test_raw_evidence_records_hash_exact_consumed_files(tmp_path):
    root = tmp_path / "collection"
    device = root / "switch"
    device.mkdir(parents=True)
    capture = device / "show_version.txt"
    capture.write_text("version evidence", encoding="utf-8")
    sidecar = device / "_capture_meta.json"
    sidecar.write_text('{"show version":"prompt_verified"}', encoding="utf-8")

    records = cp._evidence_records(
        {"switch": {"show version": str(capture)}}, str(root))
    assert records["n_files"] == 2
    assert {row["path"] for row in records["files"]} == {
        "switch/show_version.txt", "switch/_capture_meta.json"}
    assert all(len(row["sha256"]) == 64 for row in records["files"])
    assert len(records["root_sha256"]) == 64


def test_raw_evidence_includes_sidecar_for_device_with_no_command_files(tmp_path):
    root = tmp_path / "collection"
    device = root / "switch"
    device.mkdir(parents=True)
    sidecar = device / "_capture_meta.json"
    sidecar.write_text('{"collection":"failed before first command"}', encoding="utf-8")

    records = cp._evidence_records({"switch": {}}, str(root))

    assert records["n_files"] == 1
    assert records["files"][0]["path"] == "switch/_capture_meta.json"
    assert records["files"][0]["commands"] == ["<capture-metadata>"]


def test_raw_evidence_mutation_after_analysis_binding_forces_incomplete_run(tmp_path):
    root = tmp_path / "collection"
    device = root / "switch"
    device.mkdir(parents=True)
    capture = device / "show_version.txt"
    capture.write_text("bytes actually parsed", encoding="utf-8")
    mapping = {"switch": {"show version": str(capture)}}
    ctx = _finalize_ctx(tmp_path)
    ctx.root_dir = str(root)
    ctx.all_cmd_to_files = mapping
    cp._RUN_CUSTODY["evidence"]["analysis_input"] = cp._evidence_records(
        mapping, str(root))

    capture.write_text("different bytes present at sealing", encoding="utf-8")
    result = cp._stage_finalize(ctx)

    assert result.complete is False and result.exit_code == 1
    assert "Raw evidence custody" in result.failed_mandatory
    marker = json.loads((tmp_path / "Assessment.incomplete.json").read_text(
        encoding="utf-8"))
    assert any(row["step"] == "Raw evidence custody"
               for row in marker["failed_mandatory"])


def test_transient_raw_mutation_cannot_be_hidden_by_restoring_file_before_seal(
        tmp_path):
    root = tmp_path / "collection"
    device = root / "switch"
    device.mkdir(parents=True)
    capture = device / "show_version.txt"
    original = "version evidence that the parser is allowed to consume"
    capture.write_text(original, encoding="utf-8")
    mapping = {"switch": {"show version": str(capture)}}
    ctx = _finalize_ctx(tmp_path)
    ctx.root_dir = str(root)
    ctx.all_cmd_to_files = mapping
    bindings = []
    cp._RUN_CUSTODY["evidence"]["analysis_input"] = cp._evidence_records(
        mapping, str(root), _bindings_out=bindings)
    input_custody.bind_files(bindings)

    capture.write_text("transient attacker-controlled bytes", encoding="utf-8")
    assert cmdio._load_cmd_output(mapping["switch"], "show version") == ""
    capture.write_text(original, encoding="utf-8")
    assert cp._evidence_records(mapping, str(root)) == \
        cp._RUN_CUSTODY["evidence"]["analysis_input"], "precondition: bytes were restored"

    result = cp._stage_finalize(ctx)

    assert result.complete is False and result.exit_code == 1
    assert "Raw evidence custody" in result.failed_mandatory
    assert "outside the pre-analysis binding" in json.dumps(
        json.loads((tmp_path / "Assessment.incomplete.json").read_text(
            encoding="utf-8")))


@pytest.mark.parametrize(
    "field,content",
    [
        ("scenario", '{"scenarios": []}'),
        ("path_intents", '{"intents":[{"id":"x","src":"a"}]}'),
        ("traffic_intents", '{"intents":[{"src":"192.0.2.1","dst":"198.51.100.1"}]}'),
        ("assert_pack", '{"assertions":[{"id":"x","subject":"health"}]}'),
        ("requirements", '{"unknown_only":"value"}'),
    ],
)
def test_supplied_invalid_optional_inputs_fail_closed(tmp_path, field, content):
    path = tmp_path / f"{field}.json"
    path.write_text(content, encoding="utf-8")
    args = _args(**{field: str(path)})
    with pytest.raises(ValueError):
        cp._preflight_optional_inputs(args)


@pytest.mark.parametrize("mode_args", [
    ["--compare", "old.json", "new.json"],
    ["--trend", "old.json", "new.json"],
])
def test_traffic_intents_refuse_compare_trend_path_engine_reruns(monkeypatch, capsys, mode_args):
    monkeypatch.setattr(sys, "argv", [
        "cisco-assess", *mode_args, "--traffic-intents", "must-not-be-read.json",
    ])

    with pytest.raises(SystemExit) as exc:
        cp.main()

    assert exc.value.code == 2
    error = capsys.readouterr().err
    assert "project any stored traffic_assurance receipt" in error
    assert "never rerun the path engine with a new intent catalog" in error


def test_assert_pack_preserves_and_validates_complete_object_grammar(tmp_path):
    pack = {
        "assertions": [
            {
                "id": "ordinary", "subject": "collection_completeness",
                "all_of": [{"type": "contains", "value": "summary"}],
            },
            {
                "id": "inline-objects", "collection": "interfaces",
                "field_rules": [{"field": "mtu", "op": "max", "value": 9216}],
            },
        ],
        "for_each": [{
            "id": "unique-addresses", "for_each": "interfaces",
            "unique_by": "svi_ip",
        }],
        "pack_metadata": {"owner": "change-control"},
    }
    path = tmp_path / "assertions.json"
    path.write_text(json.dumps(pack), encoding="utf-8")

    loaded = cp._preflight_optional_inputs(_args(assert_pack=str(path)))

    assert loaded["assert_pack"] == pack
    assert loaded["assert_pack"]["for_each"][0]["unique_by"] == "svi_ip"
    assert loaded["assert_pack"]["pack_metadata"] == {"owner": "change-control"}
    evaluated = assertions.evaluate_pack({
        "collection_completeness": {"summary": {"complete": 1}},
        "interfaces": [
            {"host": "a", "ifname": "Gi1", "mtu": 1500, "svi_ip": "192.0.2.1"},
            {"host": "b", "ifname": "Gi1", "mtu": 9000, "svi_ip": "192.0.2.2"},
        ],
    }, loaded["assert_pack"])
    assert {row["id"] for row in evaluated["results"]} == {
        "ordinary", "inline-objects", "unique-addresses"}


def test_optional_input_hash_and_parser_share_one_bound_byte_string(tmp_path):
    path = tmp_path / "scenario.json"
    original = b'{"scenarios":[{"name":"loss","failures":[{"type":"node","id":"sw1"}]}]}'
    path.write_bytes(original)

    loaded = cp._preflight_optional_inputs(_args(scenario=str(path)))

    record = loaded["records"][0]
    assert record["sha256"] == hashlib.sha256(original).hexdigest()
    assert loaded["scenario"][0]["name"] == "loss"
    path.write_text(
        '{"scenarios":[{"name":"mutated","failures":[{"type":"node","id":"sw2"}]}]}',
        encoding="utf-8",
    )
    with pytest.raises(ValueError, match="mutation detected"):
        cp._require_current_bindings(loaded["bindings"], "optional input")


def test_traffic_intents_preflight_preserves_bound_rows_and_rejects_identity_ambiguity(tmp_path):
    path = tmp_path / "traffic-intents.json"
    catalog = {
        "intents": [{
            "id": "client-to-api", "src": "192.0.2.10", "dst": "198.51.100.20",
            "protocol": "tcp", "src_port": 49152, "dst_port": 443,
            "return_required": False,
            "failure": {"action": "fail_node", "id": "edge-2"},
        }]
    }
    path.write_text(json.dumps(catalog), encoding="utf-8")

    loaded = cp._preflight_optional_inputs(_args(traffic_intents=str(path)))

    assert loaded["traffic_intents"] == catalog["intents"]
    assert loaded["records"][0]["role"] == "traffic_intents"
    assert loaded["bindings"][0]["role"] == "traffic_intents"

    for rows, error in (
        ([{"id": "same"}, {"id": " same "}], "IDs must be unique"),
        ([{"id": "x", "failure": "fail_node"}], "failure must be a JSON object"),
    ):
        path.write_text(json.dumps({"intents": rows}), encoding="utf-8")
        with pytest.raises(ValueError, match=error):
            cp._preflight_optional_inputs(_args(traffic_intents=str(path)))


def test_tracked_traffic_intents_example_is_bound_and_semantically_valid():
    resource = files("cisco_toolkit").joinpath("data/traffic-intents.example.json")
    with as_file(resource) as path:
        loaded = cp._preflight_optional_inputs(_args(traffic_intents=str(path)))
    result = traffic_assurance.assess_flows({}, loaded["traffic_intents"])

    assert loaded["records"][0]["role"] == "traffic_intents"
    assert loaded["bindings"][0]["role"] == "traffic_intents"
    assert result["summary"]["n"] == len(loaded["traffic_intents"])
    assert result["summary"]["invalid"] == 0
    assert all(row["valid"] for row in result["results"])
    assert all(row["supported"] for row in result["results"])
    assert all(row["verdict"] == "indeterminate" for row in result["results"])
    failure = result["results"][1]["failure"]
    assert failure["requested"] is True
    assert failure["action"] == "fail_node"
    assert failure["mutation"]["valid"] is True


def test_devices_are_parsed_from_the_same_bytes_recorded_for_custody(tmp_path):
    path = tmp_path / "devices.json"
    original = json.dumps([{
        "ip": "192.0.2.1", "hostname": "switch-a", "username": "reader",
    }]).encode("utf-8")
    path.write_bytes(original)
    data, record, binding = cp._bind_input(path, role="devices_file")

    devices = cp.load_devices(str(path), allow_prompt=False, _bound_bytes=data)
    path.write_text(json.dumps([{
        "ip": "192.0.2.2", "hostname": "switch-b", "username": "reader",
    }]), encoding="utf-8")

    assert devices[0]["hostname"] == "switch-a"
    assert record["sha256"] == hashlib.sha256(original).hexdigest()
    with pytest.raises(ValueError, match="bytes changed after parsing"):
        cp._require_current_bindings([binding], "devices input")


@pytest.mark.parametrize("bad_object", [
    {"id": "missing-path", "field_rules": [{"field": "mtu", "op": "max"}]},
    {"id": "bad-rules", "collection": "interfaces", "field_rules": "not-a-list"},
    {"id": "no-constraint", "collection": "interfaces"},
    {"id": "bad-unique", "collection": "interfaces", "unique_by": ""},
])
def test_assert_pack_object_grammar_fails_closed(tmp_path, bad_object):
    path = tmp_path / "assertions.json"
    path.write_text(json.dumps({"for_each": [bad_object]}), encoding="utf-8")
    with pytest.raises(ValueError):
        cp._preflight_optional_inputs(_args(assert_pack=str(path)))


def test_devices_loader_rejects_valid_prefix_with_malformed_suffix(tmp_path):
    path = tmp_path / "devices.json"
    path.write_text(
        '[{"ip":"192.0.2.1","hostname":"switch","username":"reader"}] trailing-garbage',
        encoding="utf-8",
    )
    with pytest.raises(ValueError, match="valid JSON"):
        cp.load_devices(str(path), allow_prompt=False)


def test_late_phase_failure_updates_active_snapshot_immediately():
    snap = {}
    cp._ACTIVE_INTEGRITY_SNAPSHOT = snap
    cp._PHASE_TIMINGS.clear()

    def fail():
        raise RuntimeError("late writer failed")

    assert cp._run_phase("Late writer", fail, _default=None) is None
    assert snap["assessment_integrity"]["failed_phases"] == ["Late writer"]
    assert "RuntimeError" in snap["assessment_integrity"]["phase_errors"]["Late writer"]


def test_normal_finalizer_binds_exact_snapshot_export_and_workbook_bytes(tmp_path):
    from openpyxl import load_workbook
    from webapp.backend.protocol_portfolio import canonical_export_bytes

    ctx = _receipt_finalize_ctx(tmp_path)
    result = cp._stage_finalize(ctx)

    assert result.complete is True
    snapshot_bytes = (tmp_path / "Assessment.snapshot.json").read_bytes()
    bundle = ctx.protocol_assurance_bundle
    source = bundle["receipt"]["source_binding"]
    assert source["sha256"] == "sha256:" + hashlib.sha256(snapshot_bytes).hexdigest()
    assert source["bytes"] == len(snapshot_bytes)
    assert not ({"snapshot_id", "campaign_id", "engagement_id"} & set(source))
    export_bytes = (tmp_path / "Assessment.protocol-assurance.json").read_bytes()
    assert export_bytes == canonical_export_bytes(bundle["complete_export"])
    assert bundle["receipt"]["complete_export"]["sha256"] == (
        "sha256:" + hashlib.sha256(export_bytes).hexdigest())
    workbook = load_workbook(tmp_path / "Assessment.xlsx", read_only=True, data_only=True)
    text = "\n".join(
        str(cell.value) for row in workbook["Protocol Assurance"].iter_rows()
        for cell in row if cell.value is not None)
    workbook.close()
    assert "BOUND" in text and source["sha256"] in text
    assert "Assessment.protocol-assurance.json" in text
    assert "/api/snapshots/" not in text


def test_direct_legacy_finalizer_remains_not_verified_and_mints_no_local_receipt(tmp_path):
    from openpyxl import load_workbook

    ctx = _finalize_ctx(tmp_path)
    result = cp._stage_finalize(ctx)

    assert result.complete is True
    assert not (tmp_path / "Assessment.protocol-assurance.json").exists()
    workbook = load_workbook(tmp_path / "Assessment.xlsx", read_only=True, data_only=True)
    text = "\n".join(
        str(cell.value) for row in workbook["Protocol Assurance"].iter_rows()
        for cell in row if cell.value is not None)
    workbook.close()
    assert "NOT VERIFIED" in text
    assert "This renderer did not hash the parsed snapshot or mint source custody" in text


def test_late_optional_bound_refresh_failure_preserves_frozen_snapshot_and_not_verified_artifact(
        tmp_path, monkeypatch):
    ctx = _receipt_finalize_ctx(tmp_path)
    explorer = tmp_path / "Assessment_explorer.html"
    not_verified = b"<html><body>NOT VERIFIED current-run pre-pass</body></html>"
    explorer.write_bytes(not_verified)
    cp._register_artifact(str(explorer), kind="html", source="HTML Explorer pre-pass")
    ctx.html_output_path = str(explorer)

    def fail_refresh(*_args, **_kwargs):
        raise RuntimeError("injected late renderer failure")

    monkeypatch.setattr(cp, "write_html_explorer", fail_refresh)
    result = cp._stage_finalize(ctx)

    assert result.complete is True, "an optional presentation refresh stays fail-soft"
    snapshot_bytes = (tmp_path / "Assessment.snapshot.json").read_bytes()
    assert ctx.protocol_assurance_bundle["receipt"]["source_binding"]["sha256"] == (
        "sha256:" + hashlib.sha256(snapshot_bytes).hexdigest())
    snapshot = json.loads(snapshot_bytes)
    assert "HTML Explorer BOUND receipt" not in (
        (snapshot.get("assessment_integrity") or {}).get("failed_phases") or [])
    assert explorer.read_bytes() == not_verified
    run_manifest = json.loads((tmp_path / "Assessment.run_manifest.json").read_text(
        encoding="utf-8"))
    assert "HTML Explorer BOUND receipt" in (
        run_manifest["metadata"]["producer_finalization"]["post_snapshot_failures"])
    assert explorer.name not in {row["name"] for row in run_manifest["artifacts"]}
    assert explorer.name in {
        row["name"] for row in run_manifest["metadata"]["excluded_stale_outputs"]}


def test_snapshot_hash_drift_withholds_bound_replacement_and_forces_incomplete(
        tmp_path, monkeypatch):
    ctx = _receipt_finalize_ctx(tmp_path)
    explorer = tmp_path / "Assessment_explorer.html"
    not_verified = b"<html><body>NOT VERIFIED current-run pre-pass</body></html>"
    explorer.write_bytes(not_verified)
    cp._register_artifact(str(explorer), kind="html", source="HTML Explorer pre-pass")
    ctx.html_output_path = str(explorer)

    def drift_source(staged_path, *_args, **_kwargs):
        with open(staged_path, "w", encoding="utf-8") as fh:
            fh.write("<html><body>BOUND staged</body></html>")
        with open(ctx.snap_path, "w", encoding="utf-8") as fh:
            json.dump({"script_version": "mutated-after-freeze"}, fh)

    monkeypatch.setattr(cp, "write_html_explorer", drift_source)
    result = cp._stage_finalize(ctx)

    assert result.complete is False
    assert explorer.read_bytes() == not_verified
    assert "Protocol Assurance pre-manifest authority" in result.failed_mandatory
    assert not any("BOUND staged" in path.read_text(encoding="utf-8", errors="ignore")
                   for path in tmp_path.glob(".*.receipt-*.tmp.html"))


def test_export_byte_drift_is_rechecked_before_bound_replace_and_manifest(
        tmp_path, monkeypatch):
    ctx = _receipt_finalize_ctx(tmp_path)
    explorer = tmp_path / "Assessment_explorer.html"
    not_verified = b"<html><body>NOT VERIFIED current-run pre-pass</body></html>"
    explorer.write_bytes(not_verified)
    cp._register_artifact(str(explorer), kind="html", source="HTML Explorer pre-pass")
    ctx.html_output_path = str(explorer)

    def drift_export(staged_path, *_args, **_kwargs):
        with open(staged_path, "w", encoding="utf-8") as fh:
            fh.write("<html><body>BOUND staged</body></html>")
        with open(ctx.protocol_assurance_export_path, "wb") as fh:
            fh.write(b"{}")

    monkeypatch.setattr(cp, "write_html_explorer", drift_export)
    result = cp._stage_finalize(ctx)

    assert result.complete is False
    assert explorer.read_bytes() == not_verified
    assert "Protocol Assurance pre-manifest authority" in result.failed_mandatory
    assert (tmp_path / "Assessment.protocol-assurance.json").read_bytes() == b"{}"


def test_rehashed_family_mutation_is_rejected_against_exact_snapshot_owner(tmp_path):
    from cisco_toolkit.docmeta import protocol_assurance_receipt_view
    from cisco_toolkit.protocol_assurance import canonical_sha256
    from webapp.backend.protocol_portfolio import canonical_export_bytes

    snapshot_path, export_path, binding, bundle = _receipt_authority(tmp_path)
    forged = copy.deepcopy(bundle)
    forged["complete_export"]["families"][0]["status_reason"] = "forged but rehashed"
    forged["receipt"]["families"][0]["status_reason"] = "forged but rehashed"
    forged["receipt"]["complete_export"]["sha256"] = canonical_sha256(
        forged["complete_export"])
    unsigned = dict(forged["receipt"])
    unsigned.pop("receipt_sha256")
    forged["receipt"]["receipt_sha256"] = canonical_sha256(unsigned)
    export_path.write_bytes(canonical_export_bytes(forged["complete_export"]))

    assert protocol_assurance_receipt_view(forged)["valid"] is True, (
        "precondition: detached hashes and duplicated fields were consistently forged")
    with pytest.raises(ValueError, match="does not match the exact frozen snapshot"):
        cp._verify_protocol_receipt_authority(
            str(snapshot_path), binding, forged, str(export_path))


def test_bound_refresh_interrupt_restores_not_verified_artifact_and_cleans_stage(tmp_path):
    snapshot_path, export_path, binding, bundle = _receipt_authority(tmp_path)
    explorer = tmp_path / "Assessment_explorer.html"
    not_verified = b"<html><body>NOT VERIFIED current-run pre-pass</body></html>"
    explorer.write_bytes(not_verified)
    cp._register_artifact(str(explorer), kind="html", source="HTML Explorer pre-pass")

    def interrupt(*_args, **_kwargs):
        raise KeyboardInterrupt()

    with pytest.raises(KeyboardInterrupt):
        cp._atomic_receipt_refresh(
            "HTML Explorer BOUND receipt", str(explorer), "html", interrupt,
            snapshot_path=str(snapshot_path), source_binding=binding,
            bundle=bundle, export_path=str(export_path),
        )

    assert explorer.read_bytes() == not_verified
    assert not list(tmp_path.glob(".*.receipt-*"))
    assert cp._RUN_CUSTODY["artifacts"][str(explorer.resolve())]["sha256"] == (
        hashlib.sha256(not_verified).hexdigest())


@pytest.mark.parametrize(
    "kind,suffix",
    [("html", ".html"), ("xlsx", ".xlsx"), ("docx", ".docx")],
)
def test_structurally_valid_receiptless_writer_cannot_be_registered_as_bound(
        tmp_path, kind, suffix):
    snapshot_path, export_path, binding, bundle = _receipt_authority(tmp_path)
    target = tmp_path / f"receiptless{suffix}"

    def receiptless_writer(staged_path):
        if kind == "html":
            with open(staged_path, "w", encoding="utf-8") as fh:
                fh.write("<html><body>receipt accidentally omitted</body></html>")
        elif kind == "xlsx":
            Workbook().save(staged_path)
        else:
            from docx import Document
            Document().save(staged_path)

    assert cp._atomic_receipt_refresh(
        f"{kind} BOUND receipt", str(target), kind, receiptless_writer,
        snapshot_path=str(snapshot_path), source_binding=binding,
        bundle=bundle, export_path=str(export_path),
    ) is False

    assert not target.exists()
    assert str(target.resolve()) not in cp._RUN_CUSTODY["artifacts"]
    assert f"{kind} BOUND receipt" in cp._RUN_CUSTODY["post_snapshot_failures"]


@pytest.mark.parametrize("duplicate", [False, True])
def test_html_bound_refresh_rejects_commented_or_duplicate_receipt_assignment(
        tmp_path, duplicate):
    from cisco_toolkit.protocol_assurance import bind_snapshot_json_bytes
    from cisco_toolkit.protocol_receipt_surfaces import protocol_assurance_surface_payload

    snapshot_path, export_path, binding, bundle = _receipt_authority(tmp_path)
    target = tmp_path / "decoy.html"
    surface = protocol_assurance_surface_payload(
        bundle, complete_export_reference=export_path.name)
    encoded = json.dumps(surface, ensure_ascii=False, separators=(",", ":"))

    def decoy_writer(staged_path):
        if duplicate:
            cp.write_html_explorer(
                staged_path,
                bind_snapshot_json_bytes(snapshot_path.read_bytes()),
                "receipt-test",
                protocol_assurance_bundle=bundle,
                complete_export_reference=export_path.name,
            )
            with open(staged_path, "a", encoding="utf-8") as fh:
                fh.write(
                    "<!-- const EMBEDDED_PROTOCOL_ASSURANCE=" + encoded
                    + ";\nconst EMBEDDED_SNAPSHOT={}; duplicate -->"
                )
        else:
            with open(staged_path, "w", encoding="utf-8") as fh:
                fh.write(
                    "<html><body>Protocol Assurance receipt NOT VERIFIED</body><!-- "
                    "const EMBEDDED_PROTOCOL_ASSURANCE=" + encoded
                    + ";\nconst EMBEDDED_SNAPSHOT={};\n"
                    "load(EMBEDDED_SNAPSHOT,\"decoy\",true,EMBEDDED_PROTOCOL_ASSURANCE);"
                    " --></html>"
                )

    assert cp._atomic_receipt_refresh(
        "HTML BOUND receipt", str(target), "html", decoy_writer,
        snapshot_path=str(snapshot_path), source_binding=binding,
        bundle=bundle, export_path=str(export_path),
        expected_html_label="receipt-test" if duplicate else "decoy",
    ) is False
    assert not target.exists()
    assert str(target.resolve()) not in cp._RUN_CUSTODY["artifacts"]


@pytest.mark.parametrize("mutation", ["wrong_snapshot", "dead_branch", "changed_label"])
def test_html_bound_refresh_requires_exact_canonical_executable_projection(
        tmp_path, mutation):
    from cisco_toolkit.protocol_assurance import bind_snapshot_json_bytes

    snapshot_path, export_path, binding, bundle = _receipt_authority(tmp_path)
    target = tmp_path / f"{mutation}.html"

    def mutated_writer(staged_path):
        cp.write_html_explorer(
            staged_path,
            bind_snapshot_json_bytes(snapshot_path.read_bytes()),
            "receipt-test",
            protocol_assurance_bundle=bundle,
            complete_export_reference=export_path.name,
        )
        document = open(staged_path, encoding="utf-8").read()
        if mutation == "wrong_snapshot":
            document = re.sub(
                r"(?m)^const EMBEDDED_SNAPSHOT=\{[^\r\n]*\};$",
                "const EMBEDDED_SNAPSHOT={};",
                document,
                count=1,
            )
        elif mutation == "dead_branch":
            document = document.replace(
                "const EMBEDDED_PROTOCOL_ASSURANCE=",
                "if(false){\nconst EMBEDDED_PROTOCOL_ASSURANCE=",
                1,
            ).replace(
                "load(EMBEDDED_SNAPSHOT,\"receipt-test\",true,"
                "EMBEDDED_PROTOCOL_ASSURANCE);",
                "load(EMBEDDED_SNAPSHOT,\"receipt-test\",true,"
                "EMBEDDED_PROTOCOL_ASSURANCE);\n}",
                1,
            )
        with open(staged_path, "w", encoding="utf-8") as fh:
            fh.write(document)

    assert cp._atomic_receipt_refresh(
        "HTML BOUND receipt", str(target), "html", mutated_writer,
        snapshot_path=str(snapshot_path), source_binding=binding,
        bundle=bundle, export_path=str(export_path),
        expected_html_label=("expected-run-label" if mutation == "changed_label"
                             else "receipt-test"),
    ) is False
    assert not target.exists()
    assert str(target.resolve()) not in cp._RUN_CUSTODY["artifacts"]


@pytest.mark.parametrize("kind,suffix", [("xlsx", ".xlsx"), ("docx", ".docx")])
def test_office_bound_refresh_rejects_not_verified_decoy_token_bag(
        tmp_path, kind, suffix):
    from cisco_toolkit.protocol_receipt_surfaces import protocol_assurance_surface_payload

    snapshot_path, export_path, binding, bundle = _receipt_authority(tmp_path)
    target = tmp_path / f"decoy{suffix}"
    surface = protocol_assurance_surface_payload(
        bundle, complete_export_reference=export_path.name)
    tokens = [
        surface["schema"], "BOUND", binding["sha256"], surface["receipt_sha256"],
        surface["complete_export"]["sha256"], export_path.name,
    ]

    def decoy_writer(staged_path):
        if kind == "xlsx":
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = "Protocol Assurance"
            sheet["A1"] = "Protocol Assurance · single-snapshot receipt"
            sheet["A3"] = "Receipt status"
            sheet["B3"] = "NOT VERIFIED"
            for row_number, token in enumerate(tokens, start=20):
                sheet.cell(row_number, 1, token)
            workbook.save(staged_path)
        else:
            from docx import Document

            document = Document()
            document.add_heading("Source-bound Protocol Assurance receipt", level=2)
            document.add_paragraph("Status: NOT VERIFIED; receipt was not rendered.")
            for token in tokens:
                document.add_paragraph(token)
            document.save(staged_path)

    assert cp._atomic_receipt_refresh(
        f"{kind} BOUND receipt", str(target), kind, decoy_writer,
        snapshot_path=str(snapshot_path), source_binding=binding,
        bundle=bundle, export_path=str(export_path),
    ) is False
    assert not target.exists()
    assert str(target.resolve()) not in cp._RUN_CUSTODY["artifacts"]


def test_clean_finalization_returns_structured_complete_result_and_clears_old_marker(
        tmp_path):
    marker = tmp_path / "Assessment.incomplete.json"
    marker.write_text('{"status":"incomplete","run_id":"older"}', encoding="utf-8")

    result = cp._stage_finalize(_finalize_ctx(tmp_path))

    assert isinstance(result, cp.FinalizationResult)
    assert result.complete is True and result.exit_code == 0
    assert result.manifest_sealed is True
    assert result.failed_mandatory == ()
    assert not marker.exists(), "only a fully verified seal may clear an older marker"


def test_workbook_writer_failure_is_nonzero_and_persists_incomplete_marker(
        tmp_path, monkeypatch, caplog):
    ctx = _finalize_ctx(tmp_path)

    def fail_save(_path):
        raise OSError("injected workbook writer failure")

    monkeypatch.setattr(ctx.wb, "save", fail_save)
    with caplog.at_level(logging.INFO):
        result = cp._stage_finalize(ctx)

    assert result.complete is False and result.exit_code == 1
    assert "Assessment workbook" in result.failed_mandatory
    assert result.marker_written is True
    marker = json.loads((tmp_path / "Assessment.incomplete.json").read_text(
        encoding="utf-8"))
    assert marker["status"] == "incomplete"
    assert marker["run_id"] == result.run_id
    assert any(row["step"] == "Assessment workbook"
               for row in marker["failed_mandatory"])
    assert "[COMPLETE]" not in caplog.text


def test_manifest_writer_failure_is_nonzero_and_never_emits_completion(
        tmp_path, monkeypatch, caplog):
    ctx = _finalize_ctx(tmp_path)
    real_write = cp.write_json_file

    def fail_manifest(path, data, **kwargs):
        if str(path).endswith(".run_manifest.json"):
            raise OSError("injected manifest writer failure")
        return real_write(path, data, **kwargs)

    monkeypatch.setattr(cp, "write_json_file", fail_manifest)
    with caplog.at_level(logging.INFO):
        result = cp._stage_finalize(ctx)

    assert result.complete is False and result.exit_code == 1
    assert result.manifest_sealed is False
    assert "Run manifest" in result.failed_mandatory
    assert json.loads((tmp_path / "Assessment.incomplete.json").read_text(
        encoding="utf-8"))["output"]["manifest_sealed"] is False
    assert "[COMPLETE]" not in caplog.text


def test_manifest_self_verification_failure_is_nonzero_and_durable(
        tmp_path, monkeypatch, caplog):
    ctx = _finalize_ctx(tmp_path)
    monkeypatch.setattr(
        manifest, "verify_file",
        lambda *_args, **_kwargs: {"ok": False, "reason": "injected verification failure"})

    with caplog.at_level(logging.INFO):
        result = cp._stage_finalize(ctx)

    assert result.complete is False and result.exit_code == 1
    assert result.manifest_sealed is False
    assert "Run manifest" in result.failed_mandatory
    marker = json.loads((tmp_path / "Assessment.incomplete.json").read_text(
        encoding="utf-8"))
    assert marker["status"] == "incomplete"
    assert "injected verification failure" in json.dumps(marker)
    assert "[COMPLETE]" not in caplog.text


def test_incomplete_marker_write_failure_still_returns_nonzero_and_logs_loudly(
        tmp_path, monkeypatch, caplog):
    ctx = _finalize_ctx(tmp_path)
    monkeypatch.setattr(ctx.wb, "save", lambda _path: (_ for _ in ()).throw(
        OSError("injected workbook failure")))
    monkeypatch.setattr(
        cp, "_write_incomplete_marker",
        lambda *_args, **_kwargs: (_ for _ in ()).throw(
            OSError("injected marker failure")))

    with caplog.at_level(logging.INFO):
        result = cp._stage_finalize(ctx)

    assert result.complete is False and result.exit_code == 1
    assert result.marker_written is False
    assert "Incomplete marker" in result.failed_mandatory
    assert "durable marker write also FAILED" in caplog.text
    assert "[COMPLETE]" not in caplog.text


def test_manifest_seals_rich_artifact_and_metadata_records(tmp_path):
    workbook = tmp_path / "Assessment.xlsx"
    Workbook().save(workbook)
    cp._start_run_custody(str(workbook), [{
        "role": "devices_file", "name": "devices.json", "size": 10, "sha256": "a" * 64}])
    cp._register_artifact(str(workbook), kind="xlsx", source="current workbook writer")
    run = cp.build_run_manifest(str(workbook), {
        "collected_at": "2026-01-01T00:00:00",
        "assessment_integrity": {"failed_phases": ["Synthetic"]},
    })
    assert manifest._sealed_metadata(run["chain"]) == run["metadata"]
    assert manifest._sealed_artifacts(run["chain"]) == run["artifacts"]
    assert run["devices_file_sha256"] == "a" * 64
    assert run["metadata"]["artifact_registry"] == {
        "mode": "exact-current-run", "n_artifacts": 1}


@pytest.mark.parametrize("supply_missing_dir", [False, True])
def test_no_collect_refuses_absent_evidence_source(
        tmp_path, monkeypatch, supply_missing_dir):
    devices = tmp_path / "devices.json"
    devices.write_text(json.dumps([{
        "hostname": "switch", "ip": "192.0.2.1", "username": "reader",
        "platform": "ios",
    }]), encoding="utf-8")
    template = tmp_path / "template.xlsx"
    wb = Workbook()
    wb.active.append(["Hostname", "Port", "Status"])
    wb.save(template)
    missing = tmp_path / "missing-collection"

    argv = [
        "cisco-assess", "--no-collect",
        "--devices-file", str(devices), "--template", str(template),
        "--output", str(tmp_path / "out.xlsx"),
    ]
    if supply_missing_dir:
        argv.extend(["--collection-dir", str(missing)])
    monkeypatch.chdir(tmp_path)
    monkeypatch.setattr(sys, "argv", argv)

    with pytest.raises(SystemExit) as exc:
        cp.main()
    assert exc.value.code == 2
    assert not missing.exists(), "offline mode must not synthesize an empty evidence source"


def test_compare_cli_wires_exact_file_bindings_and_explicit_identity_context(
        tmp_path, monkeypatch):
    old = tmp_path / "old.json"
    new = tmp_path / "new.json"
    old.write_text('{"script_version":"V1","devices":{}}', encoding="utf-8")
    new.write_text('{"script_version":"V2","devices":{}}', encoding="utf-8")
    context = tmp_path / "context.json"
    context.write_text(json.dumps({
        "schema": "offline_comparison_context/1",
        "engagement_id": "ENG-CUSTODY",
        "campaign_id": 9,
        "before_snapshot_id": 90,
        "after_snapshot_id": 91,
        "change_intent": {"expected_changes": [], "note": "custody test"},
    }), encoding="utf-8")
    captured = {}

    fake_comparison = {
        "precert": {"verdict": "INDETERMINATE"},
        "cutover_gate": {"verdict": "INDETERMINATE", "note": "schema mismatch"},
    }

    def fake_compare(before, after, **kwargs):
        captured["compare"] = kwargs
        return fake_comparison

    def fake_diff(before, after, output, **kwargs):
        captured["diff"] = kwargs

    monkeypatch.setattr(cp, "compare_bound_pair", fake_compare)
    monkeypatch.setattr(cp, "write_diff_workbook", fake_diff)
    monkeypatch.setattr(sys, "argv", [
        "cisco-assess", "--compare", str(old), str(new),
        "--comparison-context", str(context), "--allow-schema-mismatch",
        "--output", str(tmp_path / "diff.xlsx"),
    ])

    cp.main()

    before_binding = captured["compare"]["before_binding"]
    after_binding = captured["compare"]["after_binding"]
    assert before_binding == {
        "source": cp.OFFLINE_FILE_SOURCE,
        "sha256": "sha256:" + cp.file_sha256(str(old)),
        "bytes": old.stat().st_size,
        "snapshot_id": 90,
        "campaign_id": 9,
        "engagement_id": "ENG-CUSTODY",
        "label": old.name,
        "script_version": "V1",
    }
    assert after_binding == {
        "source": cp.OFFLINE_FILE_SOURCE,
        "sha256": "sha256:" + cp.file_sha256(str(new)),
        "bytes": new.stat().st_size,
        "snapshot_id": 91,
        "campaign_id": 9,
        "engagement_id": "ENG-CUSTODY",
        "label": new.name,
        "script_version": "V2",
    }
    assert captured["compare"]["change_intent"] == {
        "expected_changes": [], "note": "custody test",
    }
    assert captured["diff"]["comparison"] is fake_comparison


def test_compare_rejects_source_mutation_between_parse_and_publication(
        tmp_path, monkeypatch):
    old = tmp_path / "old.json"
    new = tmp_path / "new.json"
    old.write_text('{"script_version":"V1","devices":{"a":{}}}', encoding="utf-8")
    new.write_text('{"script_version":"V1","devices":{"a":{}}}', encoding="utf-8")
    output = tmp_path / "diff.xlsx"
    real_schema_status = cp.schema_compat_status

    def mutate_after_parse(snapshots, **kwargs):
        result = real_schema_status(snapshots, **kwargs)
        old.write_text(
            '{"script_version":"V1","devices":{"mutated":{}}}', encoding="utf-8")
        return result

    monkeypatch.setattr(cp, "schema_compat_status", mutate_after_parse)
    monkeypatch.setattr(sys, "argv", [
        "cisco-assess", "--compare", str(old), str(new),
        "--output", str(output),
    ])

    with pytest.raises(SystemExit) as exc:
        cp.main()

    assert exc.value.code == 2
    assert not output.exists(), "a certificate was emitted after its source path mutated"


def test_trend_cli_wires_each_exact_hash_and_schema_status(tmp_path, monkeypatch):
    paths = []
    for index in range(3):
        path = tmp_path / f"s{index}.json"
        path.write_text(
            json.dumps({"script_version": "V1", "devices": {}, "sequence": index}),
            encoding="utf-8",
        )
        paths.append(path)
    captured = {}

    def fake_campaign(snapshots, output, **kwargs):
        captured.update(kwargs)

    monkeypatch.setattr(cp, "write_campaign_workbook", fake_campaign)
    monkeypatch.setattr(sys, "argv", [
        "cisco-assess", "--trend", *(str(path) for path in paths),
        "--output", str(tmp_path / "trend.xlsx"),
    ])

    cp.main()

    assert captured["source_bindings"] == [
        cp.file_sha256(str(path)) for path in paths]
    assert captured["schema_status"] == {
        "status": "ok", "message": "", "override": False}


def test_evidence_seals_when_the_root_is_a_different_spelling_of_the_same_directory(tmp_path):
    """Custody's containment check compared a REALPATH'D file against a RAW root.

    Two spellings of one directory then read as different directories, and custody refused every
    evidence file of a healthy run: the first windows-latest CI leg ever to complete had %TEMP%
    in DOS 8.3 form (`RUNNER~1`), realpath expanded the files to the long form (`runneradmin`),
    and the engine exited 1 with "cannot seal raw evidence ...: evidence file resolves outside
    collection root" — a FALSE refusal, fail-closed on nothing.

    8.3 names cannot be minted portably, so the same divergence is reproduced with a directory
    LINK (junction on Windows, symlink on POSIX): the root is addressed through the link while
    realpath resolves the files to the target — different spelling, same directory, exactly the
    8.3 shape. Both sides must be resolved before comparison.
    """
    import os
    import subprocess

    import COLLECT_PARSE_V3_23_0 as engine

    target = tmp_path / "real_collection"
    (target / "CORE-1").mkdir(parents=True)
    capture = target / "CORE-1" / "show_version.txt"
    capture.write_text("Cisco IOS Software, Version 15.2\n", encoding="utf-8")

    link = tmp_path / "spelled_differently"
    try:
        os.symlink(str(target), str(link), target_is_directory=True)
    except OSError:
        # Windows without developer mode: a JUNCTION needs no privilege at all.
        proc = subprocess.run(["cmd", "/c", "mklink", "/J", str(link), str(target)],
                              capture_output=True)
        if proc.returncode != 0:
            pytest.skip("neither symlink nor junction can be created here: "
                        f"{proc.stderr.decode(errors='replace')[:120]}")

    # evidence addressed THROUGH the link, root given AS the link -- realpath diverges the file,
    # and before the fix the raw root then failed commonpath containment.
    mapping = {"CORE-1": {"show version": str(link / "CORE-1" / "show_version.txt")}}
    records = engine._evidence_records(mapping, str(link))
    paths = [r["path"] for r in records["files"]] if isinstance(records, dict) and "files" in records \
        else [r["path"] for r in records.get("rows", [])] if isinstance(records, dict) \
        else [r["path"] for r in records]
    assert any("show_version.txt" in p for p in paths), (
        f"the evidence file was not sealed: {records!r}")

    # NON-VACUITY: a file genuinely OUTSIDE the root must still be refused after both sides
    # resolve -- the fix must not have widened containment into acceptance.
    outside = tmp_path / "elsewhere.txt"
    outside.write_text("outside\n", encoding="utf-8")
    with pytest.raises(RuntimeError, match="outside collection root"):
        engine._evidence_records({"CORE-1": {"show version": str(outside)}}, str(link))
