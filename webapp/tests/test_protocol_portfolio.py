"""Source-bound single-snapshot Protocol Assurance receipt and export contracts."""

from __future__ import annotations

import dataclasses
import hashlib
import io
import json
import sys
from pathlib import Path

import pytest
from fastapi.testclient import TestClient

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from backend.app import create_app  # noqa: E402
from backend.protocol_portfolio import (  # noqa: E402
    SUBJECT_RENDER_CAP,
    build_protocol_single_snapshot_bundle,
)
from cisco_toolkit.protocol_assurance import canonical_sha256  # noqa: E402
from tests.test_etherchannel_operational_evidence import (  # noqa: E402
    _copy_paths as _copy_etherchannel_paths,
    _snapshot_from_paths as _etherchannel_snapshot,
)
from tests.test_multichassis_snapshot_reconciliation import _persisted_pair  # noqa: E402


@pytest.fixture()
def client(tmp_path):
    app = create_app(db_path=str(tmp_path / "portfolio.db"))
    with TestClient(app, base_url="http://localhost") as value:
        yield value


def _campaign(client: TestClient) -> int:
    response = client.post("/api/campaigns", json={"name": "protocol portfolio"})
    assert response.status_code == 201, response.text
    return response.json()["id"]


def _upload(client: TestClient, snapshot: dict, *, pretty: bool = False) -> tuple[int, bytes]:
    campaign_id = _campaign(client)
    raw = json.dumps(snapshot, indent=2 if pretty else None).encode("utf-8")
    response = client.post(
        f"/api/campaigns/{campaign_id}/snapshots",
        files={"file": ("snapshot.json", raw, "application/json")},
        data={"label": "portfolio source"},
    )
    assert response.status_code == 201, response.text
    return response.json()["id"], raw


def _minimal(**extra) -> dict:
    return {"script_version": "portfolio-test/1", "devices": {"sw1": {}}, **extra}


def _docx_text(payload: bytes) -> str:
    from docx import Document

    document = Document(io.BytesIO(payload))
    parts = [paragraph.text for paragraph in document.paragraphs]
    parts.extend(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    return "\n".join(parts)


def _embedded_protocol_assurance(html: str) -> dict:
    encoded = html.split("const EMBEDDED_PROTOCOL_ASSURANCE=", 1)[1].split(
        ";\nconst EMBEDDED_SNAPSHOT=", 1
    )[0]
    return json.loads(encoded)


def test_single_snapshot_receipt_binds_exact_persisted_blob_and_script_owner(client):
    snapshot_id, upload_bytes = _upload(client, _minimal(), pretty=True)
    response = client.get(f"/api/snapshots/{snapshot_id}/section/protocol_assurance")
    assert response.status_code == 200, response.text
    receipt = response.json()["data"]["receipt"]

    store = client.app.state.store
    with store._lock:
        row = store._conn.execute(
            "SELECT CAST(snapshot_json AS BLOB) AS blob FROM snapshots WHERE id=?",
            (snapshot_id,),
        ).fetchone()
    persisted = bytes(row["blob"])

    source = receipt["source_binding"]
    assert receipt["schema"] == "protocol_single_snapshot_receipt/1"
    assert receipt["owns_score"] is False and receipt["owns_verdict"] is False
    assert receipt["custody_status"] == "bound"
    assert source["source"] == "persisted snapshots.snapshot_json blob"
    assert source["bytes"] == len(persisted)
    assert source["sha256"] == "sha256:" + hashlib.sha256(persisted).hexdigest()
    assert persisted != upload_bytes  # upload bytes are normalized/stamped and are not retained
    assert receipt["script_owner"] == {
        "source": "snapshot.script_version + snapshots.script_version column",
        "snapshot_value": "portfolio-test/1",
        "stored_value": "portfolio-test/1",
        "status": "bound",
    }
    assert "does not retain or claim the original upload bytes" in receipt["custody_note"]
    unsigned = dict(receipt)
    claimed = unsigned.pop("receipt_sha256")
    assert claimed == canonical_sha256(unsigned)

    meta = client.get(f"/api/snapshots/{snapshot_id}").json()
    section = next(row for row in meta["summary"]["sections"] if row["key"] == "protocol_assurance")
    assert section["label"] == "Protocol Assurance"
    assert section["count"] == receipt["summary"]["n_families"] == len(receipt["support_profiles"])


def test_offline_explorer_embeds_the_same_bound_receipt_and_names_complete_export(client):
    snapshot_id, _ = _upload(client, _minimal(), pretty=True)
    receipt = client.get(
        f"/api/snapshots/{snapshot_id}/section/protocol_assurance"
    ).json()["data"]["receipt"]

    response = client.get(f"/api/snapshots/{snapshot_id}/explorer")
    assert response.status_code == 200, response.text
    surface = _embedded_protocol_assurance(response.text)

    assert surface["receipt_valid"] is True
    assert surface["status"] == "BOUND"
    assert surface["source_binding"] == receipt["source_binding"]
    assert surface["receipt_sha256"] == receipt["receipt_sha256"]
    assert surface["complete_export"]["sha256"] == receipt["complete_export"]["sha256"]
    assert surface["complete_export"]["reference"] == (
        f"/api/snapshots/{snapshot_id}/protocol-assurance/export"
    )
    assert surface["subject_cap"] == {
        field: sum(family["subjects"][field] for family in receipt["families"])
        for field in ("total", "rendered", "omitted")
    }


def test_missing_and_malformed_family_evidence_remain_neutral_not_verified(client):
    snapshot_id, _ = _upload(
        client,
        _minimal(bgp_configured_peer_baseline="malformed", protocol_assessability=5),
    )
    response = client.get(f"/api/snapshots/{snapshot_id}/section/protocol_assurance")
    assert response.status_code == 200, response.text
    receipt = response.json()["data"]["receipt"]
    families = {row["family"]: row for row in receipt["families"]}

    bgp = families["bgp_configured_peer"]
    assert bgp["evidence_status"] == "not_verified"
    assert bgp["subject_total"] == 0
    assert bgp["subjects"] == {"total": 0, "rendered": 0, "omitted": 0, "rows": []}
    assert "malformed" in bgp["status_reason"].lower() or "invalid" in bgp["status_reason"].lower()

    # No family is upgraded merely because its executable support profile exists.
    assert all(row["evidence_status"] != "observed" for row in receipt["families"])
    assert receipt["summary"]["by_evidence_status"]["not_verified"] >= 1
    assert all(profile["implementation_state"] == "implemented" for profile in receipt["support_profiles"])


def test_etherchannel_portfolio_projects_typed_decision_depth_without_recomputation(
        client, tmp_path):
    snapshot = _etherchannel_snapshot(
        _copy_etherchannel_paths(tmp_path, "ios", "sw1", {}),
        "ios",
        "sw1",
    )
    snapshot_id, _ = _upload(client, snapshot)

    receipt = client.get(
        f"/api/snapshots/{snapshot_id}/section/protocol_assurance"
    ).json()["data"]["receipt"]
    family = next(
        row for row in receipt["families"] if row["family"] == "etherchannel"
    )

    assert family["evidence_status"] == "observed"
    assert family["subject_total"] == 1
    subject = family["subjects"]["rows"][0]
    assert subject["subject"] == "sw1|Po10"
    assert subject["kind"] == "single_chassis_local_group"
    assert subject["source_contract"] == "etherchannel_operational_evidence/1"
    detail = subject["detail"]
    assert [row["mode"] for row in detail["configured_members"]] == [
        "active", "passive",
    ]
    assert detail["partner"]["system_id"] == "0011.2233.4455"
    assert detail["partner"]["aggregation_id"] == "2"
    assert detail["min_links"]["value"] == 1
    assert detail["capacity"]["forwarding_member_count"] == 2
    assert detail["capacity"]["forwarding_bandwidth_mbps"] == 2000
    assert detail["hashing"]["algorithm"] == "src-dst-ip"
    assert detail["counter_evidence"]["fault_total"] == 0
    assert detail["member_failure_rehearsal"]["status"] == "pass"
    assert detail["member_failure_rehearsal"][
        "service_path_survival"] == "not_verified"


def test_vtp_extended_subject_is_source_bound_capped_and_complete_in_export(
        client, tmp_path):
    from cisco_toolkit.vtp_extended import embedded_vtp_extended_evidence
    from cisco_toolkit.vtp_safety import embedded_vtp_safety_baseline
    from tests.test_vtp_extended_evidence import _sources, _spec

    protected, extended, _paths, _integrity = _sources(
        tmp_path / "vtp-portfolio", {"sw1": _spec(password="portfolio-secret")})
    snapshot_id, _ = _upload(client, _minimal(
        vtp_safety_baseline=embedded_vtp_safety_baseline(protected),
        vtp_extended_evidence=embedded_vtp_extended_evidence(extended),
    ))
    receipt = client.get(
        f"/api/snapshots/{snapshot_id}/section/protocol_assurance"
    ).json()["data"]["receipt"]
    vtp = next(row for row in receipt["families"] if row["family"] == "vtp_safety")

    assert vtp["evidence_contracts"] == [
        "vtp_safety_baseline/1", "vtp_extended_evidence/1",
    ]
    assert vtp["evidence_status"] == "observed"
    assert vtp["subjects"]["total"] == vtp["subjects"]["rendered"] == 1
    subject = vtp["subjects"]["rows"][0]
    assert subject["kind"] == "local_vtp_vlan_database"
    assert subject["source_contract"] == (
        "vtp_safety_baseline/1 + vtp_extended_evidence/1")
    assert subject["detail"] == {
        "switch": "sw1",
        "mode": "server",
        "domain": "CAMPUS",
        "version": "2",
        "revision": 7,
        "database_identity": "domain=CAMPUS;version=2",
        "vlan_database_digest": subject["detail"]["vlan_database_digest"],
        "vlan_count": 3,
        "pruning_state": "not_configured",
        "authentication_configured": True,
    }
    assert subject["detail"]["vlan_database_digest"].startswith("sha256:")
    assert "portfolio-secret" not in json.dumps(receipt)

    bound, binding = client.app.state.store.get_bound_snapshot(snapshot_id)
    bundle = build_protocol_single_snapshot_bundle(bound, binding, subject_cap=0)
    capped = next(
        row for row in bundle["receipt"]["families"] if row["family"] == "vtp_safety")
    complete = next(
        row for row in bundle["complete_export"]["families"]
        if row["family"] == "vtp_safety")
    assert capped["subjects"] == {"total": 1, "rendered": 0, "omitted": 1, "rows": []}
    assert complete["subject_total"] == 1 and len(complete["subjects"]) == 1
    assert complete["subjects"][0]["detail"] == subject["detail"]


@pytest.mark.parametrize(
    "family",
    [
        "ipv4_routing_adjacency",
        "ipv6_routing_adjacency",
        "bgp_configured_peer",
        "stp_consistency",
        "stp_topology",
        "etherchannel",
        "vtp_safety",
        "fhrp_configured_group",
        "fhrp_redundancy_domain",
    ],
)
@pytest.mark.parametrize(
    "device_roster",
    [{}, {"other-site-reused-domain": {}}],
    ids=("empty-roster", "cross-site-roster"),
)
def test_exact_bound_portfolio_withholds_native_family_on_roster_mismatch(
        client, family, device_roster):
    golden = Path(__file__).resolve().parents[2] / "tests" / "golden" / "snapshot.json"
    snapshot = json.loads(golden.read_text(encoding="utf-8"))
    snapshot["devices"] = device_roster
    snapshot_id, _ = _upload(client, snapshot)

    receipt = client.get(
        f"/api/snapshots/{snapshot_id}/section/protocol_assurance"
    ).json()["data"]["receipt"]
    projected = next(row for row in receipt["families"] if row["family"] == family)

    # Exact storage custody is real, but it cannot repair an owner denominator from another estate.
    assert receipt["custody_status"] == "bound"
    assert projected["evidence_status"] == "not_verified"
    assert projected["subject_total"] == 0
    assert projected["subjects"] == {
        "total": 0, "rendered": 0, "omitted": 0, "rows": [],
    }
    assert "snapshot devices" in projected["status_reason"].lower()


def test_portfolio_rejects_stale_multichassis_baseline_over_truncated_typed_rows(client):
    snapshot = _persisted_pair()
    snapshot["multichassis_lag_typed_observations"]["observations"].pop()
    snapshot_id, _ = _upload(client, snapshot)

    receipt = client.get(
        f"/api/snapshots/{snapshot_id}/section/protocol_assurance"
    ).json()["data"]["receipt"]
    multichassis = next(
        row for row in receipt["families"] if row["family"] == "multichassis_lag"
    )

    assert multichassis["evidence_status"] == "not_verified"
    assert multichassis["subject_total"] == 0
    assert multichassis["subjects"] == {
        "total": 0, "rendered": 0, "omitted": 0, "rows": [],
    }
    assert "typed observation count does not reconcile" in multichassis[
        "status_reason"
    ].lower()


def test_portfolio_rejects_typed_multichassis_baseline_that_omits_legacy_local(client):
    snapshot = _persisted_pair()
    snapshot["devices"]["leaf-c"] = {}
    snapshot["vpc"] = {
        "leaf-a": {"domain_id": "10", "peer_status": "peer adjacency formed ok"},
        "leaf-b": {"domain_id": "10", "peer_status": "peer adjacency formed ok"},
        "leaf-c": {"domain_id": "99", "peer_status": "peer adjacency not formed"},
    }
    snapshot_id, _ = _upload(client, snapshot)

    receipt = client.get(
        f"/api/snapshots/{snapshot_id}/section/protocol_assurance"
    ).json()["data"]["receipt"]
    multichassis = next(
        row for row in receipt["families"] if row["family"] == "multichassis_lag"
    )

    assert multichassis["evidence_status"] == "not_verified"
    assert multichassis["subject_total"] == 0
    assert "legacy local subjects do not reconcile" in multichassis[
        "status_reason"
    ].lower()


def test_portfolio_rejects_fhrp_domain_receipt_from_different_svi_projection(
        client, tmp_path):
    from cisco_toolkit.fhrp_intent import embedded_fhrp_configured_group_baseline
    from cisco_toolkit.fhrp_redundancy import (
        compute_fhrp_redundancy_domain_baseline,
        embedded_fhrp_redundancy_domain_baseline,
    )
    from tests.test_fhrp_redundancy_domain_baseline import _owner

    source_dir = tmp_path / "fhrp"
    source_dir.mkdir()
    configured, interfaces = _owner(source_dir, {
        "edge-a": {"ip": "10.0.10.2", "role": "Active"},
        "edge-b": {"ip": "10.0.10.3", "role": "", "group": False},
    })
    actual = compute_fhrp_redundancy_domain_baseline(interfaces, configured)
    assert actual["domains"][0]["status"] == "review"
    graft = compute_fhrp_redundancy_domain_baseline({}, configured)
    snapshot = {
        "script_version": "portfolio-fhrp-svi-graft/1",
        "devices": {"edge-a": {}, "edge-b": {}},
        "interfaces": json.loads(json.dumps(
            interfaces,
            default=lambda item: dataclasses.asdict(item)
            if dataclasses.is_dataclass(item) else str(item),
        )),
        "fhrp_configured_group_baseline":
            embedded_fhrp_configured_group_baseline(configured),
        "fhrp_redundancy_domain_baseline":
            embedded_fhrp_redundancy_domain_baseline(
                graft, configured_group_baseline=configured),
    }
    snapshot_id, _ = _upload(client, snapshot)

    receipt = client.get(
        f"/api/snapshots/{snapshot_id}/section/protocol_assurance"
    ).json()["data"]["receipt"]
    domain = next(
        row for row in receipt["families"]
        if row["family"] == "fhrp_redundancy_domain"
    )

    assert domain["evidence_status"] == "not_verified"
    assert domain["subject_total"] == 0
    assert "svi projection digest mismatch" in domain["status_reason"].lower()


def test_portfolio_owner_withholds_custody_from_a_detached_parsed_snapshot(client):
    snapshot_id, _ = _upload(client, _minimal())
    bound, binding = client.app.state.store.get_bound_snapshot(snapshot_id)
    assert build_protocol_single_snapshot_bundle(bound, binding)["receipt"]["custody_status"] == "bound"

    detached = build_protocol_single_snapshot_bundle(dict(bound), binding)["receipt"]
    assert detached["custody_status"] == "not_verified"
    assert "exact persisted snapshot byte authority is unavailable" in detached["custody_failures"]
    assert all(family["evidence_status"] == "not_verified" for family in detached["families"])


def test_subject_cap_is_disclosed_and_complete_export_is_uncapped(client):
    golden = Path(__file__).resolve().parents[2] / "tests" / "golden" / "snapshot.json"
    snapshot = json.loads(golden.read_text(encoding="utf-8"))
    snapshot["routing_neighbors"] = {
        "core1": {
            "ospf": [
                {
                    "neighbor": f"10.200.{index // 254}.{index % 254 + 1}",
                    "address": f"10.200.{index // 254}.{index % 254 + 1}",
                    "interface": f"Vlan{index + 1}",
                    "state": "FULL/DR",
                }
                for index in range(SUBJECT_RENDER_CAP + 5)
            ],
            "bgp": [],
            "eigrp": [],
        }
    }
    snapshot_id, _ = _upload(client, snapshot)

    section = client.get(f"/api/snapshots/{snapshot_id}/section/protocol_assurance").json()["data"]
    receipt = section["receipt"]
    ipv4 = next(row for row in receipt["families"] if row["family"] == "ipv4_routing_adjacency")
    assert ipv4["subjects"]["total"] == SUBJECT_RENDER_CAP + 5
    assert ipv4["subjects"]["rendered"] == SUBJECT_RENDER_CAP
    assert ipv4["subjects"]["omitted"] == 5
    assert len(ipv4["subjects"]["rows"]) == SUBJECT_RENDER_CAP
    assert section["complete_export"]["url"].endswith(
        f"/snapshots/{snapshot_id}/protocol-assurance/export"
    )

    exported = client.get(section["complete_export"]["url"])
    assert exported.status_code == 200, exported.text
    assert exported.headers["content-disposition"].endswith(
        f'protocol-assurance-snapshot-{snapshot_id}.json"'
    )
    payload = exported.json()
    assert payload["schema"] == "protocol_single_snapshot_export/1"
    complete_ipv4 = next(row for row in payload["families"] if row["family"] == "ipv4_routing_adjacency")
    assert complete_ipv4["subject_total"] == SUBJECT_RENDER_CAP + 5
    assert len(complete_ipv4["subjects"]) == SUBJECT_RENDER_CAP + 5
    assert canonical_sha256(payload) == receipt["complete_export"]["sha256"]
    assert exported.headers["x-atlas-content-sha256"] == receipt["complete_export"]["sha256"]

    document = client.get(f"/api/snapshots/{snapshot_id}/deliverable/nrfu")
    assert document.status_code == 200, document.text
    text = _docx_text(document.content)
    assert f"{SUBJECT_RENDER_CAP} / {SUBJECT_RENDER_CAP + 5} / 5" in text
    totals = {
        field: sum(family["subjects"][field] for family in receipt["families"])
        for field in ("total", "rendered", "omitted")
    }
    assert (
        "Signed receipt payload cap (rendered / total / omitted): "
        f"{totals['rendered']} / {totals['total']} / {totals['omitted']}"
    ) in text
    assert f"subject-detail display is 0 / {totals['total']} / {totals['total']}" in text
    assert receipt["complete_export"]["sha256"] in text
    assert f"all {totals['total']} subject row(s), uncapped" in text


@pytest.mark.parametrize("kind", ["runbook", "mop", "nrfu"])
def test_operator_documents_project_one_source_bound_portfolio_receipt(client, kind):
    snapshot_id, _ = _upload(client, _minimal(generated_at="2026-08-20T00:00:00Z"))
    receipt = client.get(
        f"/api/snapshots/{snapshot_id}/section/protocol_assurance"
    ).json()["data"]["receipt"]

    response = client.get(f"/api/snapshots/{snapshot_id}/deliverable/{kind}")
    assert response.status_code == 200, response.text
    text = _docx_text(response.content)
    totals = {
        field: sum(family["subjects"][field] for family in receipt["families"])
        for field in ("total", "rendered", "omitted")
    }

    assert "protocol_single_snapshot_receipt/1" in text
    assert "owns no score and no verdict" in text
    assert receipt["source_binding"]["sha256"] in text
    assert receipt["receipt_sha256"] in text
    assert receipt["complete_export"]["sha256"] in text
    assert (
        "Signed receipt payload cap (rendered / total / omitted): "
        f"{totals['rendered']} / {totals['total']} / {totals['omitted']}"
    ) in text
    assert f"subject-detail display is 0 / {totals['total']} / {totals['total']}" in text
    assert (
        f"/api/snapshots/{snapshot_id}/protocol-assurance/export"
    ) in text
    assert "This renderer did not hash the parsed snapshot or mint source custody" not in text


@pytest.mark.parametrize("kind", ["runbook", "mop", "nrfu"])
def test_direct_document_renderer_never_mints_receipt_from_snapshot_claims(tmp_path, kind):
    fake_digest = "sha256:" + "f" * 64
    snapshot = _minimal(
        protocol_assurance={"schema": "protocol_single_snapshot_receipt/1", "sha256": fake_digest},
        protocol_single_snapshot_receipt={
            "schema": "protocol_single_snapshot_receipt/1",
            "receipt_sha256": fake_digest,
        },
    )
    output = tmp_path / f"{kind}.docx"
    if kind == "runbook":
        from cisco_toolkit.runbook import write_runbook_docx as writer
    elif kind == "mop":
        from cisco_toolkit.mop import write_mop_docx as writer
    else:
        from backend.nrfu_docx import write_nrfu_docx as writer

    writer(str(output), snapshot, "Detached receipt fixture")
    text = _docx_text(output.read_bytes())
    assert "Source-bound protocol_single_snapshot_receipt/1 unavailable" in text
    assert "This renderer did not hash the parsed snapshot or mint source custody" in text
    assert fake_digest not in text


def test_document_renderer_rejects_a_tampered_portfolio_sidecar(client, tmp_path):
    snapshot_id, _ = _upload(client, _minimal())
    bound, binding = client.app.state.store.get_bound_snapshot(snapshot_id)
    bundle = build_protocol_single_snapshot_bundle(bound, binding)
    original_digest = bundle["receipt"]["receipt_sha256"]
    bundle["receipt"]["source_binding"]["bytes"] += 1

    from backend.nrfu_docx import write_nrfu_docx

    output = tmp_path / "tampered-nrfu.docx"
    write_nrfu_docx(
        str(output),
        bound,
        "Tampered receipt fixture",
        protocol_assurance_bundle=bundle,
    )
    text = _docx_text(output.read_bytes())
    assert "Status: NOT VERIFIED (receipt digest does not reconcile)" in text
    assert original_digest not in text


def test_protocol_assurance_export_404s_for_unknown_snapshot(client):
    assert client.get("/api/snapshots/999999/section/protocol_assurance").status_code == 404
    assert client.get("/api/snapshots/999999/protocol-assurance/export").status_code == 404
