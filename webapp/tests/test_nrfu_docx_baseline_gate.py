"""NRFU must retain typed validation authority and refuse blocker-as-acceptance wording."""

from __future__ import annotations

import sys
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

pytest.importorskip("docx", reason="python-docx not installed on this runner")


def _review_snapshot() -> dict:
    row = {
        "device": "gw2",
        "platform": "ios",
        "wave": "Gateway wave",
        "category": "Routing",
        "severity": "High",
        "check": "Verify source-bound review evidence",
        "command": "show ip ospf neighbor",
        "expect": (
            "PRE-CUTOVER REVIEW — BLOCKER: the source-bound observation requires simultaneous "
            "verification or explicit disposition."
        ),
        "why": "The evidence does not establish an acceptable current state.",
        "evidence_state": "review",
        "projection_custody": "source_bound_current_run",
        "source_key": "routing_baseline_receipt/subjects/gw2/OSPF",
    }
    return {
        "devices": {"gw2": {"platform": "ios"}},
        "validation_plan": {
            "items": [row],
            "by_wave": {"Gateway wave": [dict(row)]},
            "summary": {
                "n_items": 1,
                "n_waves": 1,
                "by_category": {"Routing": 1},
                "n_high": 1,
            },
        },
    }


def test_nrfu_renders_typed_authority_and_indeterminate_blocker_banner(tmp_path):
    from docx import Document

    from backend.nrfu_docx import write_nrfu_docx

    out = tmp_path / "typed-baseline-nrfu.docx"
    write_nrfu_docx(str(out), _review_snapshot(), "Typed baseline")

    doc = Document(out)
    paragraphs = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    cells = [cell.text for table in doc.tables for row in table.rows for cell in row.cells]

    assert "CURRENT BASELINE INDETERMINATE" in paragraphs
    assert "Matching a blocker observation is not acceptance" in paragraphs
    assert "Evidence state" in cells
    assert "Custody / source" in cells
    assert "review" in cells
    assert any(
        "source_bound_current_run" in cell
        and "routing_baseline_receipt/subjects/" in cell
        for cell in cells
    )
    assert any("simultaneous verification or explicit disposition" in cell for cell in cells)
