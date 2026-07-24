"""Stored-DoS guard: a malformed snapshot section must not 500 ``GET /api/snapshots/{id}/deliverable/nrfu``.

The NRFU / Acceptance Test Plan is an AssessHub *synthesis* — ``nrfu_docx.write_nrfu_docx`` reads a
dozen sections straight out of the stored snapshot. The upload path validates only
``isinstance(snap, dict) and "devices" in snap`` (``app.py::_parse_snapshot_bytes``) and stores the JSON
verbatim, and ``deliverables.generate`` RE-RAISES after unlinking its temp file, which ``app.py`` turns
into ``HTTPException(500, "Failed to generate nrfu: ...")``. So any value that raises inside the writer
is a **stored availability DoS**: the POST is accepted (201) and *every later GET of the NRFU 500s*.

The bug class is the falsy-only guard ``X.get("k") or {}`` / ``or []``: ``or`` catches None/{}/[]/0 but a
TRUTHY non-dict / non-list (``5``, ``"x"``, ``[1, 2]``) survives and detonates on the next dereference
(``.get`` / ``.items`` / ``len`` / iteration). Several sites do not detonate at the read — they are
straight-line prologue reads whose value blows up much later:

  * ``devices`` -> ``len(devices)`` in the "Devices in scope" row (only on the non-int-``scale`` branch)
    AND in the *eagerly evaluated* default of ``coll.get('inventory', len(devices))``, which fires even
    when the ``inventory`` key IS present — two independent triggers from one guard;
  * ``collection_completeness.summary`` -> ``coll.get("not_collected")`` in the scope-limits block;
  * ``software_risk.summary`` -> ``swrisk.get("n_config_not_assessable")`` in the same block;
  * ``executive_brief.scale`` -> ``scale.get("n_devices")`` in the document-control table.

Because the prologue is straight-line, the FIRST poisoned section wins — so every case here poisons
exactly ONE section, otherwise the later sites would never be reached and would stay untested.

NON-VACUITY (verified at development time by loading ``git show origin/main:webapp/backend/nrfu_docx.py``
as ``backend.nrfu_docx`` via importlib and re-running the route): every case in ``_DEEP_CASES`` returns
**500** against the pre-guard writer and **200** after it. The ``_SECTIONS`` sweep is the whole-class net
around those anchors (a few of its cells — e.g. a *string* ``devices``, whose ``len()`` happens to work —
were already safe; they are kept as regression anchors, not claimed as pre-fix failures).

The fix routes every read through ``summary._as_list`` / a local ``_as_dict`` (the same coercers
``webapp/backend/cutover.py`` uses, mirroring ``cisco_toolkit.docmeta.as_dict``/``as_list``). For every
well-formed value those are IDENTICAL to the ``or {}`` / ``or []`` they replace, so a good snapshot
renders the same document: rendering ``_WELLFORMED`` with the pre-fix and post-fix writer produced 199
identical text segments (timestamps normalised). That is pinned here by
``test_wellformed_nrfu_render_unchanged`` — which PASSES against the pre-fix writer too, so it is a real
baseline and not a post-hoc restatement — and by the algebraic equivalence check at the bottom.

OUT OF SCOPE (deliberately not fixed here, different bug class — see the PR body): ``len({(s.get("service"),
s.get("category")) for s in services})`` raises ``TypeError: unhashable type`` when a ``service_map.services``
row carries a list/dict leaf. ``as_dict``/``as_list`` cannot fix that; it needs the unhashable-key treatment.
"""
import io
import json
import sys
from pathlib import Path

import pytest
from fastapi.testclient import TestClient

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))  # make `backend` importable

from backend.app import create_app  # noqa: E402


@pytest.fixture()
def client(tmp_path):
    app = create_app(db_path=str(tmp_path / "test.db"))
    # base_url=localhost so the default Host passes the no-token DNS-rebinding guard;
    # raise_server_exceptions=False so a stored DoS surfaces as a 500 RESPONSE we can assert on,
    # exactly as a real client experiences it.
    with TestClient(app, base_url="http://localhost", raise_server_exceptions=False) as c:
        yield c


def _upload(client, snap):
    cid = client.post("/api/campaigns", json={"name": "c"}).json()["id"]
    r = client.post(
        f"/api/campaigns/{cid}/snapshots",
        files={"file": ("s.json", json.dumps(snap).encode(), "application/json")},
        data={"label": "s"},
    )
    assert r.status_code == 201, r.text          # the POISON IS ACCEPTED -- that is what makes it *stored*
    return r.json()["id"]


def _get_nrfu(client, snap):
    """POST the poison (accepted 201 -> stored), then GET the NRFU deliverable for it."""
    return client.get(f"/api/snapshots/{_upload(client, snap)}/deliverable/nrfu")


def _renders(r, ctx=""):
    """Assert the deliverable actually rendered. 503 = python-docx absent on this runner (not a DoS)."""
    if r.status_code == 503:
        pytest.skip("python-docx not installed on this runner")
    assert r.status_code == 200, f"{ctx} -> {r.status_code}: {r.text[:300]}"


def _base(**over):
    """Minimal well-formed snapshot; `devices` present so the upload validator accepts it."""
    snap = {"devices": {"sw1": {}, "sw2": {}}}
    snap.update(over)
    return snap


# ---- whole-class sweep: every NRFU-consumed TOP-LEVEL section x {int, str, list} ----------------------
_SECTIONS = [
    "devices", "lifecycle_risk", "collection_completeness", "executive_brief", "validation_plan",
    "service_map", "application_intelligence", "design_blueprint", "design_nrfu", "software_risk",
]


@pytest.mark.parametrize("section", _SECTIONS)
@pytest.mark.parametrize("value", [5, "x", [1, 2]], ids=["int", "str", "list"])
def test_nrfu_survives_toplevel_section_poison(client, section, value):
    r = _get_nrfu(client, _base(**{section: value}))
    _renders(r, f"section={section} value={value!r}")


# ---- targeted anchors: one per guarded read, incl. the DEFERRED detonations --------------------------
# The trailing comment on each is the site that actually raised against the pre-guard writer, read off a
# real traceback (not inferred) -- named, not line-numbered, so the guard can move without rotting this.
_DEEP_CASES = {
    # `devices` -- raises in the `else len(devices)` branch of the Devices-in-scope row...
    "devices_scalar": {"devices": 5},
    # ...and independently in `coll.get('inventory', len(devices))`, an EAGER default evaluated even
    # though `inventory` IS present and the canonical int scale takes the other branch above. Two
    # separate triggers from one guard, so both need their own case.
    "devices_scalar_eager_default": {
        "devices": 5,
        "executive_brief": {"scale": {"n_devices": 2}},
        "collection_completeness": {"summary": {"inventory": 2, "complete": 2, "partial": 0,
                                                "not_collected": 0}},
    },
    # `lifecycle_risk` / `.per_device` -> `.get` on an int / `for d in 5` (raises at the read)
    "lifecycle_risk_scalar": {**_base(), "lifecycle_risk": 5},
    "lifecycle_per_device_scalar": {**_base(), "lifecycle_risk": {"per_device": 5}},
    # `collection_completeness` -> at the read; `.summary` -> DEFERRED to `coll.get("not_collected")`
    "collection_completeness_scalar": {**_base(), "collection_completeness": 5},
    "collection_summary_scalar": {**_base(), "collection_completeness": {"summary": 5}},
    # `executive_brief` -> at the read; `.scale` -> DEFERRED to `scale.get("n_devices")`
    "executive_brief_scalar": {**_base(), "executive_brief": 5},
    "exec_brief_scale_scalar": {**_base(), "executive_brief": {"scale": 5}},
    # `validation_plan` / `.items` (Phase II rows)
    "validation_plan_scalar": {**_base(), "validation_plan": 5},
    "validation_items_scalar": {**_base(), "validation_plan": {"items": 5}},
    # `service_map` / `.services` (Phase III rows)
    "service_map_scalar": {**_base(), "service_map": 5},
    "service_map_services_scalar": {**_base(), "service_map": {"services": 5}},
    # `application_intelligence` / `.domains` (Phase III rows)
    "app_intel_scalar": {**_base(), "application_intelligence": 5},
    "app_intel_domains_scalar": {**_base(), "application_intelligence": {"domains": 5}},
    # `design_blueprint` / `.decisions` (§2.1 traceability)
    "design_blueprint_scalar": {**_base(), "design_blueprint": 5},
    "design_decisions_scalar": {**_base(), "design_blueprint": {"decisions": 5}},
    # `design_nrfu` / `.items` (§2.1 phase lookup)
    "design_nrfu_scalar": {**_base(), "design_nrfu": 5},
    "design_nrfu_items_scalar": {**_base(), "design_nrfu": {"items": 5}},
    # `software_risk` -> at the read; `.summary` -> DEFERRED to `swrisk.get("n_config_not_assessable")`
    "software_risk_scalar": {**_base(), "software_risk": 5},
    "software_risk_summary_scalar": {**_base(), "software_risk": {"summary": 5}},
    # ROW-INNER: a well-formed list-of-dicts whose ROW holds a scalar -> `len(sw)` in the Phase III
    # "Endpoints across the N switch(es)" cell. The isinstance(dict) row filter does not catch this.
    "domain_switches_scalar": {
        **_base(), "application_intelligence": {"domains": [{"domain": "d1", "switches": 5}]}},
}


@pytest.mark.parametrize("snap", list(_DEEP_CASES.values()), ids=list(_DEEP_CASES.keys()))
def test_nrfu_survives_deep_field_poison(client, snap):
    """All 21 of these return 500 against the pre-guard writer (see the module docstring)."""
    _renders(_get_nrfu(client, snap))


# ---- regression anchors: shapes that were ALREADY safe and must stay so -----------------------------
@pytest.mark.parametrize("value", [5, "x", [1, 2]], ids=["int", "str", "list"])
def test_nrfu_multicast_intelligence_poison_was_already_safe(client, value):
    """`mcast` is never dereferenced (only truth-tested), so it needs NO coercion. Pinning that here so a
    later 'tidy-up' cannot quietly turn it into a dereference without a red test."""
    _renders(_get_nrfu(client, _base(multicast_intelligence=value)), f"mcast={value!r}")


@pytest.mark.parametrize("snap", [
    {**_base(), "lifecycle_risk": {"per_device": [5, "x", None]}},
    {**_base(), "validation_plan": {"items": [5, "x"]}},
    {**_base(), "service_map": {"services": [5, "x"]}},
    {**_base(), "application_intelligence": {"domains": [5, "x"]}},
], ids=["per_device", "val_items", "services", "domains"])
def test_nrfu_survives_scalar_rows_inside_wellformed_lists(client, snap):
    """The `if isinstance(..., dict)` row filters already handled these; regression anchor only."""
    _renders(_get_nrfu(client, snap))


# ---- behaviour preservation: a well-formed snapshot renders exactly what it did before ----------------
_WELLFORMED = {
    "script_version": "V3.23.0",
    "devices": {"a": {}, "b": {}},
    "generated_at": "2026-01-02 03:04",
    "executive_brief": {"scale": {"n_devices": 303, "n_vlans": 202, "n_endpoints": 5127}},
    "collection_completeness": {"summary": {"inventory": 303, "complete": 250, "partial": 0,
                                            "not_collected": 53}},
    "software_risk": {"summary": {"n_config_not_assessable": 7}},
    "lifecycle_risk": {"per_device": [
        {"host": "sw1", "model": "WS-C3750X", "sw_version": "15.2(4)E", "band": "Past end-of-support"}]},
    "validation_plan": {"items": [
        {"device": "sw1", "category": "reachability", "severity": "High", "check": "gateway reachable",
         "command": "ping 10.0.0.1", "expect": "5/5 replies"}]},
    "service_map": {"services": [{"service": "syslog", "category": "management", "port": 514,
                                  "proto": "udp"}]},
    "application_intelligence": {"domains": [{"domain": "voice", "switches": ["sw1", "sw2"]}]},
    "multicast_intelligence": {"pim": True},
    "design_blueprint": {"decisions": [
        {"id": "D1", "title": "Collapsed-core to spine/leaf", "status": "recommended", "priority": "High"},
        {"id": "D2", "title": "Wireless controller placement", "status": "needs-requirement",
         "priority": "Medium"}]},
    "design_nrfu": {"items": [{"decision_id": "D1", "phase": "post-cutover-routing"}]},
}


def _nrfu_cells(client, snap):
    r = _get_nrfu(client, snap)
    _renders(r)
    from docx import Document
    doc = Document(io.BytesIO(r.content))
    paras = [p.text for p in doc.paragraphs]
    cells = [c.text for t in doc.tables for row in t.rows for c in row.cells]
    return paras, cells


def test_wellformed_nrfu_render_unchanged(client):
    """Every value the guarded reads feed, pinned. These are exactly the strings the PRE-guard writer
    produced for this fixture -- this test passes against both writers (the full 199-segment normalised
    text diff was identical), so the coercion cannot have silently changed a well-formed render."""
    paras, cells = _nrfu_cells(client, _WELLFORMED)
    text = "\n".join(paras)

    # :43 devices + :50 executive_brief.scale -- canonical scale wins, len(devices)=2 is only a fallback
    i = next(k for k, c in enumerate(cells) if "Devices in scope" in c)
    assert cells[i + 1] == "303"
    # :46 collection_completeness.summary -- incl. the `inventory` key that shadows the eager len() default
    assert "250 of 303 device(s) fully collected, 0 partial, 53 not collected" in text
    # :46 + :60 -- the two scope-limit sentences read from coll / swrisk
    assert "53 device(s) were not collected at assessment" in text
    assert "7 device(s) had configuration that could not be assessed" in text
    # :57 design_blueprint.decisions + :59 design_nrfu.items -- the traceability rows and their phase
    assert any("Collapsed-core to spine/leaf" in c for c in cells)
    assert any(c == "Tested — post-cutover-routing" for c in cells)
    assert any("Wireless controller placement" in c for c in cells)
    assert "1 target-state design area(s) still need a requirement" in text
    # :44 lifecycle_risk.per_device -- Phase I row
    assert any("NRFU-I-001" in c for c in cells)
    assert any("Model WS-C3750X; IOS 15.2(4)E" in c for c in cells)
    # :51 validation_plan.items -- Phase II row
    assert any("NRFU-II-001" in c for c in cells)
    assert any("ping 10.0.0.1" in c for c in cells)
    # :52 service_map.services + :53/:237 domains & their switch count + mcast row
    assert any("syslog reachable end-to-end on udp/514" in c for c in cells)
    assert any("App domain: voice" in c for c in cells)
    assert any("Endpoints across the 2 switch(es) in this domain" in c for c in cells)
    assert any("Multicast / timing" in c for c in cells)
    # test-summary totals: 1 device + 1 validation item + (1 service + 1 domain + 1 multicast) = 5
    j = next(k for k, c in enumerate(cells) if c == "Total")
    assert cells[j + 2] == "5"


def test_guarded_read_matches_or_form_on_every_wellformed_value():
    """Algebraic proof that the swap is behaviour-preserving: for every value the OLD `or {}` / `or []`
    handled correctly, the coercer returns the identical object. Only truthy non-dict/non-list inputs --
    which the old form let through to crash -- differ."""
    from backend import summary
    from backend.nrfu_docx import _as_dict

    for v in [None, {}, {"a": 1}, {"a": {"b": 2}}]:
        assert _as_dict(v) == (v or {})
    for v in [None, [], [1, 2], [{"a": 1}]]:
        assert summary._as_list(v) == (v or [])
    # and the values the old form could NOT handle now degrade instead of raising
    for bad in [5, "x", [1, 2], 3.5, True]:
        assert _as_dict(bad) == {}
    for bad in [5, "x", {"a": 1}, 3.5, True]:
        assert summary._as_list(bad) == []
