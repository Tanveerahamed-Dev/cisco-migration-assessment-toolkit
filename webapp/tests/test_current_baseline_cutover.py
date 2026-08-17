"""AssessHub must carry the engine-owned current-baseline gate into the cutover workflow."""

from __future__ import annotations

import sys
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))


def _row(index: int, *, state: str = "assessed", blocker_text: str = "") -> dict:
    expect = blocker_text or f"Observed baseline {index} remains stable"
    return {
        "device": "sw1",
        "platform": "ios",
        "wave": "G1",
        "category": "Routing",
        "severity": "High",
        "check": f"validation-{index}",
        "command": f"show validation {index}",
        "expect": expect,
        "why": f"reason-{index}",
        "evidence_state": state,
        "projection_custody": "embedded_unverified",
        "source_key": f"routing_neighbors.sw1.{index}",
    }


def _validation_plan(rows: list[dict]) -> dict:
    return {
        "items": rows,
        "by_wave": {"G1": list(rows)},
        "summary": {
            "n_items": len(rows),
            "n_waves": 1,
            "by_category": {"Routing": len(rows)},
            "n_high": len(rows),
        },
    }


def _snapshot(rows: list[dict]) -> dict:
    return {
        "devices": {"sw1": {"platform": "ios"}},
        "wave_sequencing": [{
            "group": "G1", "make_before_break": ["sw1"], "hard_cutover": [],
            "hard_cutover_endpoints": 0,
        }],
        "migration_readiness": [{
            "group": "G1", "switches": ["sw1"], "readiness": "READY",
            "n_fail": 0, "n_warn": 0, "checks": [],
        }],
        "move_groups": [{"group": "G1", "switches": ["sw1"], "endpoints": 1}],
        "validation_plan": _validation_plan(rows),
    }


def _degraded_row(index: int = 1) -> dict:
    return _row(
        index,
        state="degraded",
        blocker_text=(
            "PRE-CUTOVER DEGRADED — BLOCKER: observed OSPF peer remains EXSTART/DR; "
            "matching this state is not acceptance."
        ),
    )


def test_degraded_current_baseline_forces_wave_and_fleet_no_go():
    from backend import cutover

    plan = cutover.build_plan(_snapshot([_degraded_row()]))

    assert plan["summary"]["verdict"] == cutover.GATE_NOGO
    assert plan["summary"]["current_baseline"]["verdict"] == cutover.BASELINE_BLOCKED
    assert plan["summary"]["n_baseline_blockers"] == 1
    wave = plan["waves"][0]
    assert wave["gate"] == cutover.GATE_NOGO
    assert wave["current_baseline"]["verdict"] == cutover.BASELINE_BLOCKED
    assert wave["baseline_blockers"][0]["source_key"] == "routing_neighbors.sw1.1"
    assert wave["validation"][0]["baseline_blocker"] is True
    assert "Current baseline: BLOCKED" in plan["summary"]["statement"]


def test_vtp_review_row_generically_forces_no_go_and_preserves_authority():
    from backend import cutover

    vtp = _row(
        1,
        state="review",
        blocker_text=(
            "PRE-CUTOVER REVIEW — BLOCKER: VTP Server mode with configuration revision "
            "150 was observed; matching revision 150 is NOT ACCEPTANCE."
        ),
    )
    vtp.update({
        "category": "VTP",
        "check": "VTP high-revision cutover exposure review",
        "command": "show vtp status",
        "source_key": "show vtp status#mode/domain/revision/version",
        "projection_custody": "current_run_source_bound",
    })
    snap = _snapshot([vtp])
    snap["validation_plan"]["summary"]["by_category"] = {"VTP": 1}

    plan = cutover.build_plan(snap)
    assert plan["summary"]["verdict"] == cutover.GATE_NOGO
    assert plan["summary"]["current_baseline"]["verdict"] == cutover.BASELINE_INDETERMINATE
    assert plan["waves"][0]["gate"] == cutover.GATE_NOGO
    blocker = plan["waves"][0]["baseline_blockers"][0]
    assert blocker["category"] == "VTP"
    assert blocker["command"] == "show vtp status"
    assert blocker["projection_custody"] == "current_run_source_bound"
    assert blocker["source_key"] == "show vtp status#mode/domain/revision/version"


def test_baseline_blocker_only_stops_its_affected_wave():
    from backend import cutover

    blocked = _degraded_row()
    clear = {**_row(2), "device": "sw2", "wave": "G2"}
    snap = _snapshot([blocked])
    snap["devices"]["sw2"] = {"platform": "ios"}
    snap["wave_sequencing"].append({
        "group": "G2", "make_before_break": ["sw2"], "hard_cutover": [],
        "hard_cutover_endpoints": 0,
    })
    snap["migration_readiness"].append({
        "group": "G2", "switches": ["sw2"], "readiness": "READY",
        "n_fail": 0, "n_warn": 0, "checks": [],
    })
    snap["move_groups"].append({"group": "G2", "switches": ["sw2"], "endpoints": 1})
    snap["validation_plan"] = {
        "items": [blocked, clear],
        "by_wave": {"G1": [blocked], "G2": [clear]},
        "summary": {"n_items": 2, "n_waves": 2, "by_category": {"Routing": 2}, "n_high": 2},
    }

    plan = cutover.build_plan(snap)
    by_group = {wave["group"]: wave for wave in plan["waves"]}
    assert by_group["G1"]["gate"] == cutover.GATE_NOGO
    assert by_group["G2"]["gate"] == cutover.GATE_GO
    assert by_group["G2"]["current_baseline"]["verdict"] == cutover.BASELINE_CLEAR
    assert plan["summary"]["verdict"] == cutover.GATE_NOGO


@pytest.mark.parametrize("state, marker", [
    ("review", "PRE-CUTOVER REVIEW — BLOCKER: ambiguous peer identity"),
    ("not_verified", "ROUTING BASELINE NOT VERIFIED — BLOCKER: receipt mismatch"),
])
def test_uncertain_blocker_does_not_make_an_unrelated_clear_wave_indeterminate(state, marker):
    from backend import cutover

    uncertain = _row(1, state=state, blocker_text=marker)
    clear = {**_row(2), "device": "sw2", "wave": "G2"}
    snap = _snapshot([uncertain])
    snap["devices"]["sw2"] = {"platform": "ios"}
    snap["wave_sequencing"].append({
        "group": "G2", "make_before_break": ["sw2"], "hard_cutover": [],
        "hard_cutover_endpoints": 0,
    })
    snap["migration_readiness"].append({
        "group": "G2", "switches": ["sw2"], "readiness": "READY",
        "n_fail": 0, "n_warn": 0, "checks": [],
    })
    snap["move_groups"].append({"group": "G2", "switches": ["sw2"], "endpoints": 1})
    snap["validation_plan"] = {
        "items": [uncertain, clear],
        "by_wave": {"G1": [uncertain], "G2": [clear]},
        "summary": {"n_items": 2, "n_waves": 2, "by_category": {"Routing": 2}, "n_high": 2},
    }

    by_group = {wave["group"]: wave for wave in cutover.build_plan(snap)["waves"]}
    assert by_group["G1"]["current_baseline"]["verdict"] == cutover.BASELINE_INDETERMINATE
    assert by_group["G1"]["gate"] == cutover.GATE_NOGO
    assert by_group["G2"]["current_baseline"]["verdict"] == cutover.BASELINE_CLEAR
    assert by_group["G2"]["gate"] == cutover.GATE_GO


@pytest.mark.parametrize("validation_plan, verdict", [
    ({}, "NOT_ASSESSED"),
    (5, "INDETERMINATE"),
    ({"items": "invalid", "by_wave": {}, "summary": {}}, "INDETERMINATE"),
])
def test_only_clear_baseline_can_produce_go(validation_plan, verdict):
    from backend import cutover

    snap = _snapshot([_row(1)])
    snap["validation_plan"] = validation_plan
    plan = cutover.build_plan(snap)
    assert plan["summary"]["current_baseline"]["verdict"] == verdict
    assert plan["waves"][0]["gate"] == cutover.GATE_COND

    clear = cutover.build_plan(_snapshot([_row(1)]))
    assert clear["summary"]["current_baseline"]["verdict"] == cutover.BASELINE_CLEAR
    assert clear["waves"][0]["gate"] == cutover.GATE_GO


def test_invalid_plan_with_hostile_typed_row_abstains_without_echoing_it():
    from backend import cutover

    hostile = _degraded_row()
    hostile["expect"] += " HOSTILE-PAYLOAD-MUST-NOT-SURFACE"
    snap = _snapshot([hostile])
    snap["validation_plan"]["summary"]["n_items"] = 99  # break the receipt reconciliation

    plan = cutover.build_plan(snap)

    assert plan["summary"]["current_baseline"]["verdict"] == cutover.BASELINE_INDETERMINATE
    assert plan["summary"]["current_baseline"]["blockers"] == []
    assert plan["baseline_blockers"] == []
    assert plan["waves"][0]["validation"] == []
    assert plan["waves"][0]["gate"] == cutover.GATE_COND
    assert "HOSTILE-PAYLOAD-MUST-NOT-SURFACE" not in repr(plan)


def test_cutover_docx_retains_blocker_after_ordinary_validation_cap(tmp_path):
    pytest.importorskip("docx")
    from docx import Document

    from backend.cutover_docx import write_cutover_docx

    rows = [_row(i) for i in range(1, 32)]
    blocker = _degraded_row(32)
    blocker["check"] = "BLOCKER-AT-ROW-32"
    rows.append(blocker)
    out = tmp_path / "baseline-gate.docx"

    write_cutover_docx(str(out), _snapshot(rows), "Baseline gate test")

    cells = [cell.text for table in Document(out).tables for row in table.rows for cell in row.cells]
    assert any("BLOCKER-AT-ROW-32" in cell for cell in cells)
    assert any("Command: show validation 32" in cell for cell in cells)
    assert any("PRE-CUTOVER DEGRADED — BLOCKER" in cell for cell in cells)
    assert any("embedded_unverified" in cell and "routing_neighbors.sw1.32" in cell for cell in cells)
    assert any("31 ordinary; 1 gated above" in p.text for p in Document(out).paragraphs)


def test_genuine_duplicate_blocker_occurrences_are_not_collapsed():
    from backend import cutover

    first = _degraded_row()
    duplicate = dict(first)
    plan = cutover.build_plan(_snapshot([first, duplicate]))

    assert plan["summary"]["current_baseline"]["summary"]["n_blockers"] == 2
    assert plan["summary"]["n_baseline_blockers"] == 2
    assert len(plan["baseline_blockers"]) == 2
    assert len(plan["waves"][0]["baseline_blockers"]) == 2


def test_execution_freezes_blocker_and_cannot_turn_it_into_clean_success():
    from backend import execution

    state = execution.start_run(_snapshot([_degraded_row()]), "blocked run", "lead")
    wave = state["waves"][0]
    check = wave["checks"][0]
    assert check["baseline_blocker"] is True
    assert check["baseline_state"] == "degraded"
    assert check["projection_custody"] == "embedded_unverified"
    assert check["source_key"] == "routing_neighbors.sw1.1"

    with pytest.raises(ValueError, match="cannot be recorded as a plain PASS"):
        execution.apply_check(state, wave["group"], 0, "pass", "still EXSTART", "lead")
    assert check["result"] == "pending"

    # Even a legacy writer or direct state mutation cannot promote this new, frozen non-CLEAR run
    # to clean success; a fresh clear snapshot must start a new execution record.
    for step in wave["steps"]:
        step["status"] = "done"
    check["result"] = "pass"
    wave["closeout"]["decision"] = "COMPLETE"
    assert execution._derive_outcome(state, "completed") == execution.OUTCOME_PARTIAL


def test_pir_discloses_frozen_baseline_gate_and_evidence_authority(tmp_path):
    pytest.importorskip("docx")
    from docx import Document

    from backend import execution
    from backend.pir_docx import write_pir_docx

    state = execution.start_run(_snapshot([_degraded_row()]), "blocked run", "lead")
    out = tmp_path / "blocked-pir.docx"
    write_pir_docx(str(out), state, "Blocked snapshot")

    doc = Document(out)
    paragraphs = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    cells = [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    assert "Frozen current-baseline acceptance gate" in paragraphs
    assert "Verdict at execution start: BLOCKED" in paragraphs
    assert any("show validation 1" in cell for cell in cells)
    assert any("embedded_unverified" in cell and "routing_neighbors.sw1.1" in cell for cell in cells)
    assert any("PRE-CUTOVER DEGRADED — BLOCKER" in cell for cell in cells)


def _snapshot_with_unbound_review_rows(n_rows: int) -> dict:
    rows = []
    for index in range(1, n_rows + 1):
        row = _row(index, state="review")
        row.update({
            "device": f"orphan-{index}",
            "wave": "Unscheduled",
            "check": f"UNBOUND-BLOCKER-{index:03d}",
            "expect": (
                "PRE-CUTOVER REVIEW — BLOCKER: source-bound current-state evidence requires "
                "simultaneous verification or explicit disposition."
            ),
        })
        rows.append(row)
    snap = _snapshot([])
    snap["validation_plan"] = {
        "items": rows,
        "by_wave": {"Unscheduled": list(rows)},
        "summary": {
            "n_items": len(rows),
            "n_waves": 1,
            "by_category": {"Routing": len(rows)},
            "n_high": len(rows),
        },
    }
    return snap


def test_execution_freezes_every_unbound_fleet_blocker_outside_gate_sample_cap():
    from backend import execution

    state = execution.start_run(
        _snapshot_with_unbound_review_rows(52), "unbound blocker run", "lead"
    )

    assert state["plan_summary"]["current_baseline"]["summary"]["blockers_capped"] is True
    assert state["plan_summary"]["n_unbound_baseline_blockers"] == 52
    assert len(state["baseline_blockers"]) == 52
    assert len(state["unbound_baseline_blockers"]) == 52
    assert state["baseline_blockers"][-1]["check"] == "UNBOUND-BLOCKER-052"
    assert state["waves"][0]["baseline_blockers"] == []
    assert state["waves"][0]["checks"] == []


def test_pir_renders_every_unbound_fleet_blocker_outside_gate_sample_cap(tmp_path):
    pytest.importorskip("docx")
    from docx import Document

    from backend import execution
    from backend.pir_docx import write_pir_docx

    state = execution.start_run(
        _snapshot_with_unbound_review_rows(52), "unbound blocker run", "lead"
    )
    out = tmp_path / "unbound-blockers-pir.docx"
    write_pir_docx(str(out), state, "Unbound blocker snapshot")

    doc = Document(out)
    paragraphs = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    cells = [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    assert "Fleet-level / unbound blockers:" in paragraphs
    assert "52 occurrence(s) were not bound to a scheduled wave" in paragraphs
    assert sum("Fleet-level / unbound" in cell for cell in cells) == 52
    assert any("UNBOUND-BLOCKER-052" in cell for cell in cells)
    assert any("routing_neighbors.sw1.52" in cell for cell in cells)
