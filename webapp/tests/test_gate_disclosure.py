"""AssessHub's on-demand design/MOP download vs the PPDIOO document gates (decision 2026-07-22).

The recorded decision is DISCLOSE-not-block: this surface consults `cisco_toolkit.gate_state` and
reports an unsatisfied document gate on the response, but always delivers the file. The reasoning
(an AssessHub campaign is a DB row with no filesystem engagement root, so a server-root ledger
would wrongly refuse OTHER campaigns; and there is no --override-gate here) lives in
`backend.deliverables.generate`'s docstring; the source-level half of the pin is
`tests/test_gate_state.py::test_assesshub_deliverable_download_discloses_gates`.

Both halves matter and they fail in opposite directions: a regression to silence loses the
disclosure, a regression to refusal loses the deliverable. Every test here asserts BOTH.
"""

import sys
from pathlib import Path

import pytest
from fastapi.testclient import TestClient
from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))  # make `backend` importable

from backend import deliverables  # noqa: E402
from backend.app import create_app  # noqa: E402

from cisco_toolkit import gate_state  # noqa: E402  (backend import bootstraps sys.path)

_GATED = ("design", "mop")


def _seed(client) -> tuple[int, str]:
    r = client.post("/api/demo/seed")
    assert r.status_code == 200, r.text
    body = r.json()
    return body["snapshot"]["id"], body["campaign"]["engagement_id"]


@pytest.fixture()
def engagement(tmp_path, monkeypatch):
    """Run the server with its cwd AT an engagement root, which is how gate_state resolves the
    ledger (root='.'). Yields the root so a test can record decisions into it."""
    root = tmp_path / "engagement"
    root.mkdir()
    monkeypatch.chdir(root)
    return root


def _client(tmp_path):
    app = create_app(db_path=str(tmp_path / "gate.db"))
    # base_url=localhost so the default Host passes the no-token DNS-rebinding guard (#384).
    return TestClient(app, base_url="http://localhost")


def _download(client, sid, kind):
    r = client.get(f"/api/snapshots/{sid}/deliverable/{kind}",
                   headers={"sec-fetch-site": "same-origin"})
    return r


def _docx_text(response, tmp_path, name):
    path = tmp_path / name
    path.write_bytes(response.content)
    document = Document(path)
    text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
        + [
            paragraph.text
            for section in document.sections
            for paragraph in section.header.paragraphs
        ]
    )
    return document, text


@pytest.mark.parametrize("kind", _GATED)
def test_unapproved_gate_discloses_but_still_delivers(tmp_path, engagement, kind):
    """The load-bearing case: upstream approval REVOKED (a peer review that actively said no).

    A revoked ledger must produce a disclosure header AND a real document — that pairing IS the
    decision. Asserting only the header would pass if the route had started refusing.

    The token is `pending`, the one `gate_state.enforce`/`pending_approvals` report for a LOCATED
    ledger with missing approvals. It read `ungated` until 2026-07-25 — a token `UNEVALUATED`
    reserves for approvals that were never evaluated at all, so the header told an operator whose
    LLD had been REVOKED that no gates were tracked. See
    tests/test_gate_state.py::test_gate_disclosure_token_matches_the_owner_of_the_posture_fact."""
    with _client(tmp_path) as c:
        sid, engagement_id = _seed(c)
        gate_state.bind_engagement(engagement_id, root=str(engagement), by="reviewer")
        # Revoke design's upstream and both of the MOP's for this exact campaign identity.
        for g in ("assessment_approved", "lld_approved", "baseline_captured"):
            gate_state.record_decision(g, "revoked", root=str(engagement), by="reviewer")
        if not deliverables.availability().get(kind):
            pytest.skip("python-docx not installed — the availability 503 pre-empts the gate path")
        r = _download(c, sid, kind)

    assert r.status_code == 200, f"the gate BLOCKED the download (decision is disclose): {r.text[:200]}"
    gate_hdr = r.headers.get("X-Gate-Status", "")
    assert gate_hdr.startswith("pending:"), \
        f"an unapproved gate was not disclosed as pending: {gate_hdr!r}"
    assert not gate_hdr.startswith("ungated:"), \
        f"a REVOKED approval was disclosed as 'ungated', which reads as 'no gates here': {gate_hdr!r}"
    # DOCX is a zip — 'PK' proves a real document was streamed, not an empty/stub refusal body.
    assert r.content[:2] == b"PK" and len(r.content) > 5000, \
        f"disclosed but delivered nothing usable ({len(r.content)} bytes)"
    document, text = _docx_text(r, tmp_path, f"{kind}-pending.docx")
    assert "UNAPPROVED DRAFT" in text
    if kind == "design":
        assert "assessment_approved" in text
    else:
        assert "lld_approved" in text and "baseline_captured" in text
    assert document.core_properties.content_status == "UNAPPROVED DRAFT"
    assert "atlas_gate_status=pending" in (document.core_properties.keywords or "")


@pytest.mark.parametrize("kind", _GATED)
def test_unbound_brownfield_is_marked_and_exact_bound_approval_is_silent(tmp_path, engagement, kind):
    """Only approvals bound to the exact AssessHub campaign may remove the draft marker."""
    with _client(tmp_path) as c:
        sid, engagement_id = _seed(c)
        if not deliverables.availability().get(kind):
            pytest.skip("python-docx not installed — the availability 503 pre-empts the gate path")

        # (a) no ledger can confirm ownership for this campaign: disclose and mark the file.
        assert not (engagement / "docs" / "engagement-state.json").exists()
        r = _download(c, sid, kind)
        assert r.status_code == 200
        assert r.headers.get("X-Gate-Status", "").startswith("ownership_unbound:")
        unbound, text = _docx_text(r, tmp_path, f"{kind}-unbound.docx")
        assert "UNAPPROVED DRAFT" in text
        assert unbound.core_properties.content_status == "UNAPPROVED DRAFT"

        # (b) the ledger is bound to this exact campaign and every upstream gate is approved.
        gate_state.bind_engagement(engagement_id, root=str(engagement), by="lead")
        for g in ("assessment_approved", "lld_approved", "baseline_captured"):
            gate_state.record_decision(g, "approved", root=str(engagement), by="lead")
        r = _download(c, sid, kind)
        assert r.status_code == 200 and "X-Gate-Status" not in r.headers, \
            f"an approved engagement was flagged: {r.headers.get('X-Gate-Status')!r}"
        approved, text = _docx_text(r, tmp_path, f"{kind}-approved.docx")
        assert "UNAPPROVED DRAFT" not in text
        assert approved.core_properties.content_status != "UNAPPROVED DRAFT"


def test_ungated_deliverable_never_carries_the_header(tmp_path, engagement):
    """Only design and mop are gated (gate_state.GENERATOR_REQUIRES). A revoked ledger must not
    leak a disclosure onto the runbook, which no document gate governs."""
    gate_state.record_decision("assessment_approved", "revoked", root=str(engagement), by="reviewer")
    with _client(tmp_path) as c:
        sid, _ = _seed(c)
        if not deliverables.availability().get("runbook"):
            pytest.skip("python-docx not installed")
        r = _download(c, sid, "runbook")
    assert r.status_code == 200 and "X-Gate-Status" not in r.headers, \
        f"an ungated deliverable was gate-flagged: {r.headers.get('X-Gate-Status')!r}"


def test_corrupt_ledger_discloses_and_still_delivers(tmp_path, engagement):
    """The asymmetry this surface deliberately inverts: the CLI fails CLOSED on an unreadable
    store, AssessHub cannot (a broken ledger must never withhold a deliverable) — so it must fail
    OPEN AND SAY SO. Silence here would be the worst outcome: an operator whose ledger is corrupt
    would read it as 'gates fine'."""
    docs = engagement / "docs"
    docs.mkdir(parents=True)
    (docs / "engagement-state.json").write_text("{ this is not json", encoding="utf-8")

    with _client(tmp_path) as c:
        sid, _ = _seed(c)
        if not deliverables.availability().get("design"):
            pytest.skip("python-docx not installed")
        r = _download(c, sid, "design")

    assert r.status_code == 200, f"a corrupt ledger blocked the download: {r.text[:200]}"
    assert r.headers.get("X-Gate-Status", "").startswith("unreadable:"), \
        f"a corrupt ledger was not disclosed: {r.headers.get('X-Gate-Status')!r}"
    assert r.content[:2] == b"PK"
    document, text = _docx_text(r, tmp_path, "design-unreadable.docx")
    assert "UNAPPROVED DRAFT" in text and "UNREADABLE" in text
    assert str(engagement) not in text
    assert document.core_properties.content_status == "UNAPPROVED DRAFT"


def test_other_campaigns_approved_ledger_cannot_silence_this_document(tmp_path, engagement):
    gate_state.bind_engagement("urn:uuid:another-campaign", root=str(engagement), by="other-lead")
    gate_state.record_decision(
        "assessment_approved", "approved", root=str(engagement), by="other-lead"
    )
    with _client(tmp_path) as client:
        sid, this_engagement = _seed(client)
        if not deliverables.availability().get("design"):
            pytest.skip("python-docx not installed")
        response = _download(client, sid, "design")
    assert this_engagement != "urn:uuid:another-campaign"
    assert response.headers.get("X-Gate-Status", "").startswith("ownership_mismatch:")
    document, text = _docx_text(response, tmp_path, "ownership-mismatch.docx")
    assert "UNAPPROVED DRAFT" in text and "OWNERSHIP_MISMATCH" in text
    assert document.core_properties.content_status == "UNAPPROVED DRAFT"


def test_one_gate_read_drives_header_and_document(tmp_path, engagement, monkeypatch):
    calls = []
    original = deliverables.gate_disclosure

    def changing_gate(kind, gate_root=".", engagement=None):
        calls.append((kind, gate_root, engagement))
        if len(calls) == 1:
            return {
                "status": "pending",
                "missing": ["assessment_approved"],
                "revoked": [],
            }
        return None

    monkeypatch.setattr(deliverables, "gate_disclosure", changing_gate)
    try:
        with _client(tmp_path) as client:
            sid, engagement_id = _seed(client)
            if not deliverables.availability().get("design"):
                pytest.skip("python-docx not installed")
            response = _download(client, sid, "design")
    finally:
        monkeypatch.setattr(deliverables, "gate_disclosure", original)

    assert len(calls) == 1
    assert calls[0][2] == engagement_id
    assert response.headers["X-Gate-Status"].startswith("pending:")
    document, text = _docx_text(response, tmp_path, "single-read.docx")
    assert "UNAPPROVED DRAFT" in text
    assert document.core_properties.content_status == "UNAPPROVED DRAFT"


def test_gate_evaluation_exception_marks_document_instead_of_becoming_silence(
    tmp_path, engagement, monkeypatch, caplog,
):
    monkeypatch.setattr(
        gate_state,
        "load_store",
        lambda *args, **kwargs: (_ for _ in ()).throw(
            RuntimeError("private path detail\nFORGED LOG RECORD")
        ),
    )
    with _client(tmp_path) as client:
        sid, _ = _seed(client)
        response = _download(client, sid, "design")
    assert response.status_code == 200
    assert response.headers.get("X-Gate-Status", "").startswith("unreadable:")
    _, text = _docx_text(response, tmp_path, "gate-exception.docx")
    assert "UNAPPROVED DRAFT" in text
    assert "private path detail" not in text
    assert "private path detail" not in caplog.text
    assert "FORGED LOG RECORD" not in caplog.text
    assert caplog.text.count("gated document disclosure evaluation failed") == 1
