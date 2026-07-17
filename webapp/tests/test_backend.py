"""End-to-end backend tests for AssessHub against an isolated (temp) SQLite store.

Exercises the full flow the frontend depends on: demo seed, campaign + snapshot CRUD, summary
projection, section slicing, explorer render, compare, and trend — all without touching the real DB
or the engine's golden contract.
"""

import sys
from pathlib import Path

import pytest
from fastapi.testclient import TestClient

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))  # make `backend` importable
_ENGINE_TESTS = str(Path(__file__).resolve().parents[2] / "tests")  # engine fixtures (synthetic collection)
if _ENGINE_TESTS not in sys.path:
    sys.path.append(_ENGINE_TESTS)  # append (not insert) so webapp modules keep import priority

from backend.app import create_app  # noqa: E402


@pytest.fixture()
def client(tmp_path):
    app = create_app(db_path=str(tmp_path / "test.db"))
    # base_url=localhost so the default Host passes the no-token DNS-rebinding guard (app.py
    # _request_host_allowed) — the dev server this emulates is reached over loopback.
    with TestClient(app, base_url="http://localhost") as c:
        yield c


def test_health(client):
    r = client.get("/api/health")
    assert r.status_code == 200
    assert r.json()["status"] == "ok"
    assert r.json()["sample_available"] is True


def test_demo_seed_then_summary(client):
    r = client.post("/api/demo/seed")
    assert r.status_code == 200, r.text
    body = r.json()
    snap_id = body["snapshot"]["id"]

    # campaign listing carries a latest_summary used on dashboard cards
    cl = client.get("/api/campaigns").json()
    assert len(cl) == 1
    assert cl[0]["n_snapshots"] == 1
    assert cl[0]["latest_summary"]["n_switches"] >= 1

    # snapshot meta + derived summary
    meta = client.get(f"/api/snapshots/{snap_id}").json()
    s = meta["summary"]
    assert s["n_switches"] >= 1
    assert s["punchlist"]["total"] >= 1
    assert set(s["bands"]).issubset(set(["Excellent", "Good", "Fair", "Poor", "Critical"]))
    assert isinstance(s["keystones"], list) and s["keystones"]
    # at least the punch-list section is reported present
    assert any(sec["key"] == "punchlist" for sec in s["sections"])


def test_section_slice_and_guard(client):
    snap_id = client.post("/api/demo/seed").json()["snapshot"]["id"]
    r = client.get(f"/api/snapshots/{snap_id}/section/punchlist")
    assert r.status_code == 200
    assert r.json()["section"] == "punchlist"
    assert isinstance(r.json()["data"], list)
    # unknown section is rejected, not silently empty
    assert client.get(f"/api/snapshots/{snap_id}/section/not_a_section").status_code == 400


def test_architecture_coverage_endpoint(client):
    """The architecture-coverage SSOT is served (computed server-side with the SAME engine function the
    explorer/CLI use -- the dashboard never re-derives coverage). The class count is locked to the engine's
    OWN registry (drift-proof: when the engine adds an architecture class the webapp contract moves with it,
    so the lock can never silently fall stale the way a hardcoded number does)."""
    import cisco_toolkit.design_advisor as da
    n_ssh = sum(1 for _axis, _label, ch, _pids in da._ARCH_COVERAGE_REGISTRY if ch == "ssh")
    n_json = sum(1 for _axis, _label, ch, _pids in da._ARCH_COVERAGE_REGISTRY if ch == "json")
    snap_id = client.post("/api/demo/seed").json()["snapshot"]["id"]
    r = client.get(f"/api/snapshots/{snap_id}/architecture_coverage")
    assert r.status_code == 200
    cov = r.json()
    # Cross-surface SSOT: the served count == the engine's canonical registry (and is internally consistent).
    assert isinstance(cov.get("classes"), list)
    assert cov["summary"]["n_classes"] == len(da._ARCH_COVERAGE_REGISTRY) == len(cov["classes"])
    assert cov["summary"]["by_channel"]["json"] == n_json and cov["summary"]["by_channel"]["ssh"] == n_ssh
    by = {c["key"]: c for c in cov["classes"]}
    assert by["aci"]["channel"] == "json" and by["sdwan"]["channel"] == "json"
    # coverage-honest: every class is observed-and-status or not-observed (never silently 'healthy')
    assert all(c["status"] in ("finding", "clean", "not-observed") for c in cov["classes"])
    assert client.get("/api/snapshots/999999/architecture_coverage").status_code == 404


def test_domain_packs_endpoint(client):
    """The domain skill-packs a snapshot engages (Phase-3 / D6) are served, selected by the engine SSOT
    (domain_packs.select_packs) -- never re-derived in JS. Cross-endpoint SSOT: a pack loads IFF one of its
    architecture classes is OBSERVED in the SAME /architecture_coverage map, so the chip strip can never
    disagree with the coverage grid beside it. Coverage-honest: an empty selection is stated, not silent."""
    from cisco_toolkit.domain_packs import PACKS
    snap_id = client.post("/api/demo/seed").json()["snapshot"]["id"]
    r = client.get(f"/api/snapshots/{snap_id}/domain_packs")
    assert r.status_code == 200
    dp = r.json()
    assert isinstance(dp.get("selected"), list) and isinstance(dp.get("loaded"), list) and dp.get("note")
    assert {s["pack"] for s in dp["selected"]} == set(dp["loaded"]) <= set(PACKS)
    # the invariant that keeps the two panels honest: loaded IFF an observed class in that pack
    cov = client.get(f"/api/snapshots/{snap_id}/architecture_coverage").json()
    observed = {c["key"] for c in cov["classes"] if c.get("observed")}
    for pid, spec in PACKS.items():
        assert (pid in dp["loaded"]) == bool(observed & spec["classes"]), pid
    # every loaded pack cites the OBSERVED class(es) that triggered it (never an empty/ungrounded trigger)
    for s in dp["selected"]:
        assert s["triggered_by"] and set(s["triggered_by"]) <= observed
    assert client.get("/api/snapshots/999999/domain_packs").status_code == 404


def test_nos_quartet_sections_reachable(client):
    """NEW-V3.23.176: the syslog / QoS / software-risk / platform-health axes (V3.23.164-.167)
    are tabbed AND fetchable -- the one-source-of-truth audit found them unreachable from the
    web platform (they landed after SECTION_LABELS was authored)."""
    snap_id = client.post("/api/demo/seed").json()["snapshot"]["id"]
    meta = client.get(f"/api/snapshots/{snap_id}").json()
    keys = {sec["key"] for sec in meta["summary"]["sections"]}
    quartet = {"syslog_intelligence", "qos_audit", "software_risk", "platform_health"}
    assert quartet <= keys, f"missing tabs: {quartet - keys}"
    for name in sorted(quartet):
        r = client.get(f"/api/snapshots/{snap_id}/section/{name}")
        assert r.status_code == 200, f"{name}: {r.text}"
        assert isinstance(r.json()["data"], dict)


def test_device_risk_register_section(client):
    """NEW-V3.23.174: the Device Risk Register is a whitelisted section and the demo sample
    (regenerated through the real pipeline) carries it, pre-ranked riskiest-first."""
    snap_id = client.post("/api/demo/seed").json()["snapshot"]["id"]
    meta = client.get(f"/api/snapshots/{snap_id}").json()
    assert any(sec["key"] == "device_dossiers" for sec in meta["summary"]["sections"])
    r = client.get(f"/api/snapshots/{snap_id}/section/device_dossiers")
    assert r.status_code == 200
    dd = r.json()["data"]
    rows = dd["per_device"]
    assert rows and {"host", "risk_band", "risk_index", "verdict"} <= set(rows[0])
    ranks = {"Severe": 0, "Elevated": 1, "Guarded": 2, "Low": 3}
    assert [ranks[d["risk_band"]] for d in rows] == sorted(ranks[d["risk_band"]] for d in rows)


def test_orchestration_peer_engine_sections_served(client):
    """NEW (orchestration-peer wave): the three always-on engines (G1 acl_line_reachability,
    I2 feature_compliance, K1 capture_integrity) are whitelisted tabs, indexed for visibility, and
    served with their findings + summary intact -- the marquee different-action shadow survives the
    round-trip, so the webapp reaches parity with the explorer's coverage-honest cards."""
    import json
    cid = client.post("/api/campaigns", json={"name": "engines"}).json()["id"]
    snap = {
        "script_version": "test", "devices": {"R1": {}}, "health_scores": [],
        "acl_line_reachability": {
            "findings": [{"host": "R1", "acl": "GUEST", "line_index": 6, "action": "permit",
                          "reason": "BLOCKING_LINES", "blocking_lines": [3], "different_action": True,
                          "detail": "permit hidden behind an earlier deny"}],
            "summary": {"n_findings": 1, "n_shadowed": 1, "n_different_action": 1,
                        "n_unmatchable": 0, "n_indeterminate": 0, "n_bad_reference": 0}},
        "feature_compliance": {
            "features": [{"feature": "aaa", "n_baseline": 3, "n_drifting": 1}],
            "per_device_feature": [{"host": "R1", "feature": "aaa", "n_missing": 1, "status": "drift"}],
            "summary": {"n_features": 1, "n_rows": 1, "n_drift_rows": 1}},
        "capture_integrity": {
            "findings": [{"host": "R1", "command": "show running-config", "status": "incomplete",
                          "reason": "no terminating 'end' near the tail"}],
            "summary": {"n_findings": 1, "n_incomplete": 1, "n_error": 0, "n_empty": 0, "n_hosts_affected": 1}},
    }
    up = client.post(f"/api/campaigns/{cid}/snapshots",
                     files={"file": ("snap.json", json.dumps(snap).encode(), "application/json")},
                     data={"label": "engines"})
    assert up.status_code == 201, up.text
    sid = up.json()["id"]

    keys = {sec["key"] for sec in client.get(f"/api/snapshots/{sid}").json()["summary"]["sections"]}
    trio = {"acl_line_reachability", "feature_compliance", "capture_integrity"}
    assert trio <= keys, f"missing tabs: {trio - keys}"
    for name in sorted(trio):
        r = client.get(f"/api/snapshots/{sid}/section/{name}")
        assert r.status_code == 200, f"{name}: {r.text}"
        assert r.json()["section"] == name
        assert isinstance(r.json()["data"].get("summary"), dict)

    # G1 marquee: the dangerous DIFFERENT-action shadow is preserved through store + serve
    acl = client.get(f"/api/snapshots/{sid}/section/acl_line_reachability").json()["data"]
    assert acl["findings"][0]["different_action"] is True
    assert acl["summary"]["n_different_action"] == 1

    # a clean capture estate hides its own tab (zero findings -> not indexed), never a false-green
    clean = dict(snap, capture_integrity={"findings": [],
                 "summary": {"n_findings": 0, "n_incomplete": 0, "n_error": 0, "n_empty": 0, "n_hosts_affected": 0}})
    up2 = client.post(f"/api/campaigns/{cid}/snapshots",
                      files={"file": ("snap2.json", json.dumps(clean).encode(), "application/json")},
                      data={"label": "clean"})
    keys2 = {sec["key"] for sec in client.get(f"/api/snapshots/{up2.json()['id']}").json()["summary"]["sections"]}
    assert "capture_integrity" not in keys2


def test_orchestration_peer_optin_sections_served(client):
    """NEW (orchestration-peer wave): the four OPT-IN engines (A1 state_assertions, G3 path_intents,
    B external_reconcile, G4 whatif) are whitelisted tabs and served when their key is present -- so a
    snapshot produced with --assert-pack/--path-intents/--import-inventory/--scenario reaches the webapp.
    Coverage-honesty survives the round-trip: a not_observed verdict and a lost_path (not a fabricated block)."""
    import json
    cid = client.post("/api/campaigns", json={"name": "optin"}).json()["id"]
    snap = {
        "script_version": "test", "devices": {"R1": {}}, "health_scores": [],
        "state_assertions": {
            "results": [{"id": "grade", "title": "grade A/B", "status": "fail", "subject": "security_grade"},
                        {"id": "vpc", "title": "vPC up", "status": "not_observed", "subject": "vpc", "abstention": "not_collected"}],
            "summary": {"n_pass": 0, "n_fail": 1, "n_not_observed": 1, "n_assessed": 1, "grade": "fail"}},
        "path_intents": {
            "results": [{"id": "pci-iso", "src": "10.0.1.1", "dst": "10.0.2.5", "expect": "ISOLATED",
                         "verdict": "fail", "status": "computed:reached"}],
            "summary": {"pass": 0, "fail": 1, "not_observed": 0}},
        "external_reconcile": {
            "rows": [{"type": "UNVERIFIABLE", "host": "ACC-5", "detail": "declared but never collected"}],
            "summary": {"MISSING_DEVICE": 0, "UNDOCUMENTED_DEVICE": 0, "MODEL_MISMATCH": 0,
                        "IP_DRIFT": 0, "UNVERIFIABLE": 1, "n_declared": 1, "n_observed": 0, "n_rows": 1}},
        "whatif": [{"name": "Lose DIST-1", "removed_hosts": ["DIST-1"],
                    "summary": {"blocked": 0, "lost_path": 12, "preserved": 8, "inconclusive_other": 48, "other": 12}}],
    }
    up = client.post(f"/api/campaigns/{cid}/snapshots",
                     files={"file": ("snap.json", json.dumps(snap).encode(), "application/json")},
                     data={"label": "optin"})
    assert up.status_code == 201, up.text
    sid = up.json()["id"]

    keys = {sec["key"] for sec in client.get(f"/api/snapshots/{sid}").json()["summary"]["sections"]}
    quartet = {"state_assertions", "path_intents", "external_reconcile", "whatif"}
    assert quartet <= keys, f"missing tabs: {quartet - keys}"
    for name in sorted(quartet):
        assert client.get(f"/api/snapshots/{sid}/section/{name}").status_code == 200

    # coverage-honesty preserved through the round-trip
    sa = client.get(f"/api/snapshots/{sid}/section/state_assertions").json()["data"]
    assert any(r["status"] == "not_observed" for r in sa["results"])   # a blind spot, not a silent pass
    wi = client.get(f"/api/snapshots/{sid}/section/whatif").json()["data"]
    assert isinstance(wi, list) and wi[0]["summary"]["blocked"] == 0 and wi[0]["summary"]["lost_path"] == 12


def test_parse_yield_section_served(client):
    """Plan A / Tier-1 #3 surfacing: snap['parse_yield'] (the zero-parse yield ledger the workbook's
    Collection Completeness sheet already renders) is a whitelisted, indexed tab and is served intact --
    the summary counts (zero_yield_suspect / expected / errors), the events table, and the engine's
    coverage-honest note VERBATIM (collected-but-unparsed evidence = a possible parser format gap,
    never a device verdict). A run where every content-bearing command parsed (no events) hides its
    own tab -- the platform convention -- so the telemetry never reads like a device finding."""
    import json

    from cisco_toolkit import cmdio
    note = cmdio.parse_yield_report()["summary"]["note"]   # the engine's own wording (SSOT, no paraphrase)

    cid = client.post("/api/campaigns", json={"name": "parse-yield"}).json()["id"]
    snap = {
        "script_version": "test", "devices": {"N9K-1": {}}, "health_scores": [],
        "parse_yield": {
            "summary": {"parsers_called": 3, "zero_yield_suspect": 1, "zero_yield_expected": 1,
                        "parse_errors": 1, "note": note},
            "per_parser": {
                # the marquee class: a real NX-OS RIB printed 220 lines yet parsed to 0 routes
                "parse_ip_routes": {"calls": 1, "with_content": 1, "zero_yield": 1, "errors": 0,
                                    "may_be_empty": False},
                "parse_acl_hitcounts": {"calls": 1, "with_content": 1, "zero_yield": 1, "errors": 0,
                                        "may_be_empty": True},
                "parse_vpc": {"calls": 1, "with_content": 1, "zero_yield": 0, "errors": 1,
                              "may_be_empty": False},
            },
            "events": [
                {"parser": "parse_ip_routes", "device": "N9K-1", "cmd": "show ip route vrf all",
                 "file": "show_ip_route_vrf_all.txt", "lines_in": 220, "error": False},
                {"parser": "parse_acl_hitcounts", "device": "N9K-1", "cmd": "show access-lists",
                 "file": "show_access-lists.txt", "lines_in": 12, "error": False},
                {"parser": "parse_vpc", "device": "N9K-1", "cmd": "show vpc",
                 "file": "show_vpc.txt", "lines_in": 40, "error": True},
            ],
            "events_truncated": False,
        },
    }
    up = client.post(f"/api/campaigns/{cid}/snapshots",
                     files={"file": ("snap.json", json.dumps(snap).encode(), "application/json")},
                     data={"label": "py"})
    assert up.status_code == 201, up.text
    sid = up.json()["id"]

    # tabbed (indexed by the events list -- the section's meaningful inner list) ...
    secs = {s["key"]: s for s in client.get(f"/api/snapshots/{sid}").json()["summary"]["sections"]}
    assert "parse_yield" in secs, f"parse_yield tab missing: {sorted(secs)}"
    assert secs["parse_yield"]["count"] == 3
    assert secs["parse_yield"]["label"] == "Parse yield"
    # ... and fetchable, with the ledger intact
    r = client.get(f"/api/snapshots/{sid}/section/parse_yield")
    assert r.status_code == 200, r.text
    d = r.json()["data"]
    s = d["summary"]
    assert (s["zero_yield_suspect"], s["zero_yield_expected"], s["parse_errors"]) == (1, 1, 1)
    assert s["note"] == note                          # verbatim engine wording survives the round-trip
    assert "never a device" in s["note"]
    assert d["per_parser"]["parse_acl_hitcounts"]["may_be_empty"] is True   # drives the class column
    assert len(d["events"]) == 3 and d["events"][0]["parser"]

    # the /api/meta contract carries the label the tab bar renders
    labels = {sl["key"]: sl["label"] for sl in client.get("/api/meta").json()["section_labels"]}
    assert labels.get("parse_yield") == "Parse yield"

    # a fully-parsed run (no events) hides its own tab -- telemetry, never a device verdict
    clean = dict(snap, parse_yield={
        "summary": {"parsers_called": 90, "zero_yield_suspect": 0, "zero_yield_expected": 0,
                    "parse_errors": 0, "note": note},
        "per_parser": {}, "events": [], "events_truncated": False})
    up2 = client.post(f"/api/campaigns/{cid}/snapshots",
                      files={"file": ("clean.json", json.dumps(clean).encode(), "application/json")},
                      data={"label": "clean"})
    keys2 = {s["key"] for s in client.get(f"/api/snapshots/{up2.json()['id']}").json()["summary"]["sections"]}
    assert "parse_yield" not in keys2


def test_explorer_render_embeds_snapshot(client):
    snap_id = client.post("/api/demo/seed").json()["snapshot"]["id"]
    r = client.get(f"/api/snapshots/{snap_id}/explorer")
    assert r.status_code == 200
    assert "text/html" in r.headers["content-type"]
    assert "EMBEDDED_SNAPSHOT" in r.text  # the live snapshot got baked into the template


def test_upload_and_compare(client):
    c = client.post("/api/campaigns", json={"name": "Wave test"}).json()
    cid = c["id"]
    sample = (Path(__file__).resolve().parents[2] / "tests" / "golden" / "snapshot.json")
    raw = sample.read_bytes()

    def upload(label):
        return client.post(f"/api/campaigns/{cid}/snapshots",
                           files={"file": ("snap.json", raw, "application/json")},
                           data={"label": label})

    a = upload("wave0")
    b = upload("wave1")
    assert a.status_code == 201 and b.status_code == 201

    cmp = client.post("/api/compare", json={"old_id": a.json()["id"], "new_id": b.json()["id"]})
    assert cmp.status_code == 200
    assert "verdict" in cmp.json()
    assert "schema_compat" in cmp.json()          # P3-E2: the diff surfaces its schema-compat verdict

    trend = client.get(f"/api/campaigns/{cid}/trend")
    assert trend.status_code == 200
    assert len(trend.json()["timeline"]) == 2
    assert "schema_compat" in trend.json()        # P3-E2: the trend surfaces its schema-compat verdict


def test_engine_diff_surfaces_schema_compat_status():
    """P3-E2 (webapp surface): snapshot_delta / campaign_trend surface the pair/series schema-compat
    verdict (ok | unverifiable | mismatch) so the interactive UI never presents a diff across an engine
    schema change as a real network change. Additive + non-breaking; ASCII message (cp1252 safe)."""
    from webapp.backend import engine
    mism = engine.snapshot_delta({"script_version": "V3.23.0", "devices": {}},
                                 {"script_version": "V9.9.9", "devices": {}})
    assert mism["schema_compat"]["status"] == "mismatch"
    assert mism["schema_compat"]["message"].isascii()
    ok = engine.snapshot_delta({"script_version": "V3.23.0"}, {"script_version": "V3.23.0"})
    assert ok["schema_compat"]["status"] == "ok"
    trend = engine.campaign_trend([{"script_version": "V3.23.0"}, {"script_version": "V9.9.9"}])
    assert trend["schema_compat"]["status"] == "mismatch"


def test_graph_endpoint(client):
    snap_id = client.post("/api/demo/seed").json()["snapshot"]["id"]
    r = client.get(f"/api/snapshots/{snap_id}/graph")
    assert r.status_code == 200
    g = r.json()
    assert g["nodes"] and g["edges"], "graph should have nodes and edges"
    ids = {n["id"] for n in g["nodes"]}
    # every edge connects two known nodes (no dangling APs/phones)
    for e in g["edges"]:
        assert e["source"] in ids and e["target"] in ids
        assert "is_bridge" in e
    # nodes carry the health band used to colour them
    assert any(n.get("band") for n in g["nodes"])
    # 404 for a missing snapshot
    assert client.get("/api/snapshots/999999/graph").status_code == 404


def test_cable_map_endpoint(client):
    """EDA-style cable map: role-tiered node/port/cable model, op-status coloured, coverage-honest.
    (Exercises the rehydration fallback — the demo snapshot predates the cable-map engine.)"""
    snap_id = client.post("/api/demo/seed").json()["snapshot"]["id"]
    r = client.get(f"/api/snapshots/{snap_id}/cable_map")
    assert r.status_code == 200
    cm = r.json()
    assert cm["nodes"] and cm["cables"] and cm["tiers"], "cable map should have nodes/cables/tiers"
    hosts = {n["host"] for n in cm["nodes"]}
    valid = {"up", "down", "unknown"}
    for n in cm["nodes"]:
        assert n["op_status"] in valid
        assert isinstance(n["tier"], int) and isinstance(n["collected"], bool)
        assert isinstance(n.get("kind"), str) and n["kind"]      # evidence-based kind (fabric-only filter input)
        # an uncollected device is [NOT OBSERVED] neutral, never a fake 'up'
        if not n["collected"]:
            assert n["op_status"] == "unknown" and "uncollected" in n["badges"]
    for c in cm["cables"]:
        assert c["op_status"] in valid
        assert c["a"] in hosts and c["b"] in hosts
        assert "speed" in c                                       # link speed surfaces on every cable
    # the summary op rollup accounts for every cable
    op = cm["summary"]["op"]
    assert op["up"] + op["down"] + op["unknown"] == len(cm["cables"])
    # 404 for a missing snapshot
    assert client.get("/api/snapshots/999999/cable_map").status_code == 404


def test_cutover_plan(client):
    snap_id = client.post("/api/demo/seed").json()["snapshot"]["id"]
    r = client.get(f"/api/snapshots/{snap_id}/cutover")
    assert r.status_code == 200, r.text
    plan = r.json()

    s = plan["summary"]
    assert s["verdict"] in ("GO", "CONDITIONAL GO", "NO-GO")
    assert s["n_waves"] == len(plan["waves"]) >= 1
    # n_devices is corroborated against per-wave bucket membership (an independent recount, not its own
    # definition) — this would fail if a switch were double-counted across buckets/waves.
    bucket_devices = {sw for w in plan["waves"] for sw in (w["make_before_break"] + w["hard_cutover"])}
    assert s["n_devices"] == len(bucket_devices)

    waves = plan["waves"]
    # pilot-first sequencing: order is monotonic and never schedules a worse gate before a better one
    rank = {"GO": 0, "CONDITIONAL GO": 1, "NO-GO": 2}
    assert [w["order"] for w in waves] == list(range(1, len(waves) + 1))
    assert [rank[w["gate"]] for w in waves] == sorted(rank[w["gate"]] for w in waves)

    for w in waves:
        assert w["gate"] in rank
        assert w["strategy"] in ("make-before-break", "hard-cutover", "mixed")
        # a make-before-break-only wave is zero-outage; a hard-cutover wave needs a window
        if not w["hard_cutover"]:
            assert w["est_window_minutes"] == 0
        else:
            assert w["est_window_minutes"] > 0
        # every wave carries a PPDIOO run-of-show ending at the rollback gate
        phases = [step["phase"] for step in w["run_of_show"]]
        assert phases[0] == "Baseline capture" and phases[-1] == "Rollback gate"
        # a NO-GO wave is gated by a failing readiness check or a Critical cross-layer hit
        if w["gate"] == "NO-GO":
            assert w["n_fail"] > 0 or w["critical_crosslayer"]

    assert client.get("/api/snapshots/999999/cutover").status_code == 404


def test_cutover_gate_critical_crosslayer_only():
    """A wave that passes every readiness check but is hit by a Critical cross-layer must still be
    NO-GO — the cross-layer-only gating path the API fixture (which fails on readiness) doesn't reach."""
    from backend import cutover

    snap = {
        "devices": {"sw1": {}, "sw2": {}},
        "wave_sequencing": [{"group": "G1", "make_before_break": ["sw1", "sw2"],
                             "hard_cutover": [], "hard_cutover_endpoints": 0}],
        "migration_readiness": [{"group": "G1", "switches": ["sw1", "sw2"], "readiness": "READY",
                                 "n_fail": 0, "n_warn": 0, "checks": []}],
        "move_groups": [{"switches": ["sw1", "sw2"], "endpoints": 10}],
        "cross_layer": [{"id": "CL-1", "severity": "Critical", "title": "x", "layers": "L1+L3",
                         "recommendation": "fix", "hosts": ["sw1"]}],
    }
    plan = cutover.build_plan(snap)
    wave = plan["waves"][0]
    assert wave["n_fail"] == 0 and wave["critical_crosslayer"]      # readiness clean, cross-layer hit
    assert wave["gate"] == "NO-GO"
    assert plan["summary"]["verdict"] == "NO-GO"


def test_cutover_blind_devices_block_go_and_are_disclosed():
    """[audit-5 #15] A wave whose readiness is otherwise clean but that contains a device the collection NEVER
    reached (health band 'Insufficient Data' / data_quality 0) must NOT be gated a confident GO -- an un-assessed
    device cannot be certified ready -- and the fleet statement must DISCLOSE those devices rather than asserting a
    confident make-before-break / window posture over them (a blind subsystem must not read like an assessed one)."""
    from backend import cutover

    base = {
        "devices": {"sw1": {}, "sw2": {}},
        "wave_sequencing": [{"group": "G1", "make_before_break": ["sw1", "sw2"],
                             "hard_cutover": [], "hard_cutover_endpoints": 0}],
        "migration_readiness": [{"group": "G1", "switches": ["sw1", "sw2"], "readiness": "READY",
                                 "n_fail": 0, "n_warn": 0, "checks": []}],
        "move_groups": [{"switches": ["sw1", "sw2"], "endpoints": 10}],
    }
    # sw2 was never collected -> banded 'Insufficient Data', data_quality 0.
    blind = dict(base, health_scores=[{"switch": "sw1", "band": "Good", "score": 80, "data_quality": 1.0},
                                      {"switch": "sw2", "band": "Insufficient Data", "score": None, "data_quality": 0.0}])
    plan = cutover.build_plan(blind)
    wave = plan["waves"][0]
    assert wave["n_fail"] == 0                                   # readiness is clean...
    assert wave["n_blind"] == 1 and "sw2" in wave.get("blind_switches", [])
    assert wave["gate"] != "GO"                                 # ...but an un-assessed device blocks a confident GO
    assert plan["summary"]["n_not_assessed"] == 1
    assert "assess" in plan["summary"]["statement"].lower()     # the statement discloses the coverage gap
    # control: a fully-collected wave is unaffected.
    seen = dict(base, health_scores=[{"switch": "sw1", "band": "Good", "score": 80, "data_quality": 1.0},
                                     {"switch": "sw2", "band": "Good", "score": 75, "data_quality": 1.0}])
    plan2 = cutover.build_plan(seen)
    assert plan2["waves"][0]["n_blind"] == 0 and plan2["waves"][0]["gate"] == "GO"
    assert plan2["summary"]["n_not_assessed"] == 0


def test_cutover_robust_to_malformed_snapshot():
    """build_plan must degrade gracefully on a malformed (e.g. uploaded) snapshot: a string `hosts`
    field still matches (not split into characters), and non-numeric counters don't crash."""
    from backend import cutover

    snap = {
        "devices": {"sw1": {}},
        "wave_sequencing": [{"group": "G1", "make_before_break": ["sw1"], "hard_cutover": [],
                             "hard_cutover_endpoints": None}],
        "migration_readiness": [{"group": "G1", "switches": ["sw1"], "readiness": "READY",
                                 "n_fail": None, "n_warn": None, "checks": []}],
        "cross_layer": [{"id": "CL-1", "severity": "Critical", "title": "x", "hosts": "sw1"}],  # string
        "failure_impact": [{"host": "sw1", "severity": "High", "stranded": None, "vlans_impacted": "3"}],
    }
    plan = cutover.build_plan(snap)              # must not raise
    wave = plan["waves"][0]
    assert wave["gate"] == "NO-GO"               # the string-host Critical cross-layer still gates
    assert wave["critical_crosslayer"]
    assert wave["blast_radius"]["stranded"] == 0          # None coerced, not crashed
    assert wave["blast_radius"]["vlans_impacted"] == 3    # "3" coerced


def test_deliverables(client):
    snap_id = client.post("/api/demo/seed").json()["snapshot"]["id"]
    cat = client.get("/api/meta").json()["deliverables"]
    assert {d["key"] for d in cat} == {"engagement", "crd", "runbook", "design", "archreview",
                                       "mop", "cutover", "nrfu", "opshandbook", "deck"}
    for d in cat:
        r = client.get(f"/api/snapshots/{snap_id}/deliverable/{d['key']}")
        if d["available"]:
            assert r.status_code == 200, r.text
            assert r.content[:2] == b"PK"          # docx/pptx are ZIP containers
            assert len(r.content) > 1000
            assert d["key"] in r.headers.get("content-disposition", "")
        else:
            assert r.status_code == 503
    assert client.get(f"/api/snapshots/{snap_id}/deliverable/nope").status_code == 400


def test_gate_board_roundtrip_and_plan_of_record_feedback(client):
    """V3.23.158: the gate board records T-minus sign-offs per (wave, gate); the engagement plan of
    record carries them back as §4.3 'Gate record (as signed)'. Pending clears the row, so an
    untouched board leaves the document without an as-signed section (no invented project state)."""
    seeded = client.post("/api/demo/seed").json()
    cid, snap_id = seeded["campaign"]["id"], seeded["snapshot"]["id"]

    board = client.get(f"/api/campaigns/{cid}/gates").json()
    assert [g["key"] for g in board["cadence"]] == [
        "commit", "checkpoint", "readiness", "go_no_go", "window", "hypercare_exit"]
    assert board["records"] == []                        # nothing signed yet
    # V3.23.159: no defensive fallback — if the demo fleet stops deriving waves, FAIL here
    # (a phantom fallback wave would keep this test green while the board/loop is broken).
    assert board["waves"], "demo fleet must derive at least one wave"
    wave = board["waves"][0]

    # sign, then upsert the same gate (one row, latest decision wins)
    r = client.post(f"/api/campaigns/{cid}/gates",
                    json={"wave": wave, "gate": "go_no_go", "decision": "no-go",
                          "signed_by": "A. Engineer", "note": "blocker open"})
    assert r.status_code == 200
    r = client.post(f"/api/campaigns/{cid}/gates",
                    json={"wave": wave, "gate": "go_no_go", "decision": "go",
                          "signed_by": "A. Engineer", "note": "checks green"})
    recs = r.json()["records"]
    assert len(recs) == 1 and recs[0]["decision"] == "go" and recs[0]["signed_by"] == "A. Engineer"

    # validation: unknown gate / unknown decision / unknown wave / size caps / missing campaign
    bad_gate = client.post(f"/api/campaigns/{cid}/gates",
                           json={"wave": wave, "gate": "bogus", "decision": "go"})
    assert bad_gate.status_code == 400
    bad_dec = client.post(f"/api/campaigns/{cid}/gates",
                          json={"wave": wave, "gate": "commit", "decision": "maybe"})
    assert bad_dec.status_code == 400
    # V3.23.159: a typo'd wave label must not mint a permanent governance row
    bad_wave = client.post(f"/api/campaigns/{cid}/gates",
                           json={"wave": "Goup 1 (typo)", "gate": "commit", "decision": "go"})
    assert bad_wave.status_code == 400 and "Unknown wave" in bad_wave.text
    # V3.23.159: unbounded strings are rejected by the model, not stored/echoed/rendered
    too_big = client.post(f"/api/campaigns/{cid}/gates",
                          json={"wave": wave, "gate": "commit", "decision": "go",
                                "note": "x" * 10_000})
    assert too_big.status_code == 422
    assert client.get("/api/campaigns/999999/gates").status_code == 404
    assert client.post("/api/campaigns/999999/gates",
                       json={"wave": wave, "gate": "commit", "decision": "go"}).status_code == 404

    # the feedback loop: the signed decision lands in the engagement DOCX
    r = client.get(f"/api/snapshots/{snap_id}/deliverable/engagement")
    if r.status_code == 503:
        pytest.skip("python-docx not installed on this runner")
    assert r.status_code == 200, r.text

    import io

    from docx import Document

    def _text(resp):
        doc = Document(io.BytesIO(resp.content))
        parts = [p.text for p in doc.paragraphs]
        parts += [c.text for t in doc.tables for row in t.rows for c in row.cells]
        return "\n".join(parts)

    text = _text(r)
    assert "Gate record (as signed)" in text
    assert "GO" in text and "A. Engineer" in text and "checks green" in text

    # clearing back to pending removes the row AND the as-signed section — and is idempotent
    r = client.post(f"/api/campaigns/{cid}/gates",
                    json={"wave": wave, "gate": "go_no_go", "decision": "pending"})
    assert r.json()["records"] == []
    r = client.post(f"/api/campaigns/{cid}/gates",
                    json={"wave": wave, "gate": "go_no_go", "decision": "pending"})
    assert r.status_code == 200 and r.json()["records"] == []   # clearing a clear cell is a no-op
    r = client.get(f"/api/snapshots/{snap_id}/deliverable/engagement")
    assert "Gate record (as signed)" not in _text(r)


def test_archreview_endpoint(client):
    """V3.23.163: the Ask-the-Engineer panel's data — the senior-engineer design review. The server
    prefers the snapshot's stored architecture_review section and computes with the SAME engine
    function otherwise, so any stored snapshot (old or new) gets a review; absent evidence renders
    not-assessable verdicts, never a fabricated grade."""
    snap_id = client.post("/api/demo/seed").json()["snapshot"]["id"]
    r = client.get(f"/api/snapshots/{snap_id}/archreview")
    assert r.status_code == 200, r.text
    ar = r.json()
    assert len(ar["domains"]) == 8
    assert ar["summary"]["n_checks"] == len(ar["checks"]) >= 20
    assert ar["summary"]["grade"] in ("A", "B", "C", "D", "F", "N/A")
    allowed = {"conforms", "advisory", "deviation", "critical", "not-assessable"}
    assert {c["verdict"] for c in ar["checks"]} <= allowed
    # every check carries the full senior-engineer block the panel renders
    for c in ar["checks"]:
        for k in ("observed", "implication", "recommendation", "reference", "evidence"):
            assert k in c, (c.get("id"), k)
    assert client.get("/api/snapshots/999999/archreview").status_code == 404


def test_design_blueprint_endpoint_and_requirements_overlay(client):
    """The CCDE-grounded design blueprint endpoint — the SAME compute_design_blueprint object the HLD/LLD
    DOCX and the explorer Design mode read. GET returns the evidence-grounded baseline (every decision
    cites a CCDE principle); POSTing a requirements register right-sizes (re-scores) every decision. The
    right-sizing logic lives only in Python — one source of truth across script and dashboard."""
    snap_id = client.post("/api/demo/seed").json()["snapshot"]["id"]
    r = client.get(f"/api/snapshots/{snap_id}/design")
    assert r.status_code == 200, r.text
    bp = r.json()
    assert isinstance(bp["decisions"], list)
    assert len(bp["tradeoff_scorecard"]) >= 1
    for k in ("summary", "requirements_model", "coverage", "axes"):
        assert k in bp, k
    for d in bp["decisions"]:
        for k in ("id", "title", "priority", "status", "evidence", "principle", "axes"):
            assert k in d, (d.get("id"), k)
        assert d["principle"]["citation"], d.get("id")    # every decision cites a CCDE source
    # NEW (target_state surfaces the dashboards render — SAME object the HLD §5 reads, no client recompute)
    ts = bp["target_state"]
    for k in ("dimensions", "replacement_bom", "addressing_plan", "wave_plan", "segmentation_plan"):
        assert k in ts, k
    assert isinstance(ts["dimensions"], list)
    assert {"n_replace", "n_refresh", "replace_now", "refresh_soon"} <= set(ts["replacement_bom"])
    assert ts["addressing_plan"].get("status") in ("candidate", "needs-requirement")
    assert {"waves", "n_waves", "wave_cap"} <= set(ts["wave_plan"])
    # segmentation_plan is the field the dashboards' "Target segmentation" block reads (SSOT, no client recompute)
    assert ts["segmentation_plan"].get("status") in ("candidate", "needs-requirement")
    assert ts["segmentation_plan"].get("observed") and ts["segmentation_plan"].get("target")
    # requirements overlay: supplying a register re-scores (effective_priority) every decision
    r2 = client.post(f"/api/snapshots/{snap_id}/design",
                     json={"availability_tier": "gold", "critical_apps": ["voice"], "growth_horizon": "3y"})
    assert r2.status_code == 200, r2.text
    bp2 = r2.json()
    assert all("effective_priority" in d for d in bp2["decisions"])
    assert "target_state" in bp2                          # the overlay (server-computed) carries target_state too
    # the stored section is also reachable through the generic section reader (SSOT)
    assert client.get(f"/api/snapshots/{snap_id}/section/design_blueprint").status_code in (200, 404)
    assert client.get("/api/snapshots/999999/design").status_code == 404


def test_design_overlay_accepts_interview_answers(client):
    """C1 (audit fix): the requirements loop closes from the engagement INTERVIEW too. POSTing
    {"interview_answers": {...}} maps the typed answers through the SAME requirements_from_interview
    bridge the CLI references, then recomputes server-side — so interview output is no longer a dead path
    (the bridge previously had no production caller). One right-sizing source: Python."""
    def _status(bp, did):
        return next((d["status"] for d in bp["decisions"] if d["id"] == did), None)
    snap_id = client.post("/api/demo/seed").json()["snapshot"]["id"]
    base = client.get(f"/api/snapshots/{snap_id}/design").json()
    # the defense-in-depth decision is an OPEN question until a data_classification requirement is supplied
    assert _status(base, "security-defense-in-depth-segmentation") == "needs-requirement"
    r = client.post(f"/api/snapshots/{snap_id}/design", json={"interview_answers": {
        "availability_tier": "Gold", "critical_apps": "voice, video", "growth_horizon": "double in 3y",
        "data_classification": ["restricted", "internal"], "convergence_budget_ms": "200"}})
    assert r.status_code == 200, r.text
    bp = r.json()
    assert all("effective_priority" in d for d in bp["decisions"]), "answers must be applied (re-scored)"
    # discriminating check: this only flips if data_classification was genuinely EXTRACTED from the
    # interview answers (a raw {"interview_answers": {...}} register would leave it needs-requirement)
    assert _status(bp, "security-defense-in-depth-segmentation") == "recommended", \
        "interview answers must be mapped via requirements_from_interview, not treated as a raw register"


def test_design_overlay_resolves_fabric_operating_model_choice(client):
    """SSOT interactivity: the DC fabric operating-model CHOICE (Cisco ACI vs standalone NX-OS VXLAN-EVPN)
    is requirement-gated. POSTing fabric_operating_model flips it from open-question to recommended through
    the SAME compute_design_blueprint the CLI/HLD run — the dashboard never re-derives the choice (and free
    text like 'ACI' is canonicalised server-side, not in the browser)."""
    def _status(bp, did):
        return next((d["status"] for d in bp["decisions"] if d["id"] == did), None)
    pid = "dc-fabric-aci-vs-nxos-evpn-operating-model"
    snap_id = client.post("/api/demo/seed").json()["snapshot"]["id"]
    base = client.get(f"/api/snapshots/{snap_id}/design").json()
    assert _status(base, pid) == "needs-requirement"          # open question until the WHY is supplied
    r = client.post(f"/api/snapshots/{snap_id}/design", json={"fabric_operating_model": "ACI"})
    assert r.status_code == 200, r.text
    assert _status(r.json(), pid) == "recommended"            # resolved server-side via the same engine


def test_design_nrfu_endpoint(client):
    """Design-driven NRFU/ATP endpoint: GET /api/snapshots/{id}/design/nrfu returns a structured
    acceptance-test checklist derived from the recommended design decisions — one item per decision,
    phased across pre-cutover / post-cutover-functional / post-cutover-operational. Every item must
    carry decision_id, title, priority, phase, description, pass_criteria, devices, principle_citation
    and trace back to a recommended decision from the baseline blueprint (SSOT: Python only)."""
    snap_id = client.post("/api/demo/seed").json()["snapshot"]["id"]
    r = client.get(f"/api/snapshots/{snap_id}/design/nrfu")
    assert r.status_code == 200, r.text
    result = r.json()
    assert "items" in result and "n_items" in result and "note" in result
    assert result["n_items"] == len(result["items"])
    required = {"decision_id", "title", "priority", "phase", "description",
                "pass_criteria", "devices", "principle_citation"}
    for item in result["items"]:
        missing = required - set(item)
        assert not missing, f"NRFU item {item.get('decision_id')} missing keys: {missing}"
    # all items trace to recommended decisions from the same blueprint
    bp = client.get(f"/api/snapshots/{snap_id}/design").json()
    rec_ids = {d["id"] for d in bp["decisions"] if d["status"] == "recommended"}
    for item in result["items"]:
        assert item["decision_id"] in rec_ids, (
            f"NRFU item {item['decision_id']} does not trace to a recommended decision"
        )
    # phases are the expected values
    valid_phases = {"pre-cutover", "post-cutover-functional", "post-cutover-operational"}
    for item in result["items"]:
        assert item["phase"] in valid_phases, f"unknown phase: {item['phase']}"
    # 404 on unknown snapshot
    assert client.get("/api/snapshots/999999/design/nrfu").status_code == 404


def test_design_nrfu_overlay_reflects_requirements(client):
    """REVIEW #13: POST /api/snapshots/{id}/design/nrfu right-sizes the NRFU checklist to a requirements
    register — a design decision that flips needs-requirement -> recommended under the overlay appears as a
    NEW NRFU item, so the dashboard NRFU tab reflects right-sizing (not the stale baseline). SSOT: the NRFU
    is derived (server-side) from the SAME overlay blueprint POST /design returns, never re-derived in JS."""
    snap_id = client.post("/api/demo/seed").json()["snapshot"]["id"]
    base_nrfu = client.get(f"/api/snapshots/{snap_id}/design/nrfu").json()
    base_ids = {it["decision_id"] for it in base_nrfu["items"]}
    reg = {"availability_tier": "gold", "growth_horizon": "double the campus in 3 years",
           "data_classification": ["PCI", "corp"], "critical_apps": ["voice"]}
    r = client.post(f"/api/snapshots/{snap_id}/design/nrfu", json=reg)
    assert r.status_code == 200, r.text
    over = r.json()
    assert "items" in over and over["n_items"] == len(over["items"])
    over_ids = {it["decision_id"] for it in over["items"]}
    # SSOT: every overlay NRFU item traces to a recommended decision of the OVERLAY blueprint
    rec_over = {d["id"] for d in client.post(f"/api/snapshots/{snap_id}/design", json=reg).json()["decisions"]
                if d["status"] == "recommended"}
    assert over_ids <= rec_over
    # right-sizing had an effect: the overlay flipped at least one decision, and the NRFU reflects it
    rec_base = {d["id"] for d in client.get(f"/api/snapshots/{snap_id}/design").json()["decisions"]
                if d["status"] == "recommended"}
    assert rec_over != rec_base, "requirements should flip at least one open design question"
    assert over_ids != base_ids, "the NRFU checklist must reflect right-sizing, not the baseline"
    # also accepts the interview-answers wrapper (same bridge as POST /design)
    assert client.post(f"/api/snapshots/{snap_id}/design/nrfu",
                       json={"interview_answers": {"availability_tier": "gold"}}).status_code == 200
    assert client.post("/api/snapshots/999999/design/nrfu", json=reg).status_code == 404


def test_design_overlay_address_space_unlocks_ip_plan(client):
    """SSOT fix: POSTing address_space to /design recomputes the blueprint with a CANDIDATE
    addressing_plan (status='candidate') — the IP plan must go from needs-requirement to a live
    candidate allocation. Verifies the webapp requirements form can now supply address_space
    (previously missing from the DesignBlueprint.tsx form, so the IP plan stayed permanently gated).
    Also confirms n_census_vlans and n_unsizable are present on both the needs-requirement and
    candidate paths."""
    snap_id = client.post("/api/demo/seed").json()["snapshot"]["id"]
    # baseline: no address_space -> needs-requirement
    base = client.get(f"/api/snapshots/{snap_id}/design").json()
    ap_base = base["target_state"]["addressing_plan"]
    assert ap_base["status"] == "needs-requirement"
    assert "n_census_vlans" in ap_base and "n_unsizable" in ap_base, (
        "n_census_vlans / n_unsizable must be present even when needs-requirement"
    )
    # with address_space -> candidate
    r = client.post(f"/api/snapshots/{snap_id}/design",
                    json={"address_space": "10.0.0.0/16"})
    assert r.status_code == 200, r.text
    bp2 = r.json()
    ap2 = bp2["target_state"]["addressing_plan"]
    assert ap2["status"] == "candidate", "address_space must unlock the IP plan to candidate"
    assert isinstance(ap2.get("subnets"), list) and len(ap2["subnets"]) > 0, "candidate plan must have subnets"
    assert "n_census_vlans" in ap2 and "n_unsizable" in ap2


def test_cutover_deliverable_content(client):
    """The Cutover Plan DOCX is the one deliverable with no engine-side test — validate it actually
    renders the plan (headings + tables), not just that it's a >1KB zip."""
    snap_id = client.post("/api/demo/seed").json()["snapshot"]["id"]
    r = client.get(f"/api/snapshots/{snap_id}/deliverable/cutover")
    if r.status_code == 503:
        pytest.skip("python-docx not installed on this runner")
    assert r.status_code == 200, r.text

    import io

    from docx import Document

    doc = Document(io.BytesIO(r.content))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Migration Cutover Plan" in text            # title page
    assert "Recommended wave sequence" in text         # §2 wave table section
    assert "Methodology" in text                       # grounding section
    assert "Run-of-show" in text                       # per-wave run-of-show heading
    assert "Document Control" in text                  # AS-style front matter (V3.23.150)
    assert "Document Acceptance" in text               # closing signature gate
    all_rows = [c.text for t in doc.tables for row in t.rows for c in row.cells]
    assert any("Customer network owner" in c for c in all_rows)   # acceptance roles
    assert len(doc.tables) >= 2                         # summary + sequence tables at minimum
    # V3.23.154: manual step numbering restarts per wave — Word's "List Number" style would
    # continue counting across waves (same defect class as the MOP's e263b6d fix)
    n_ros = sum(1 for p in doc.paragraphs if p.text == "Run-of-show")
    n_first = sum(1 for p in doc.paragraphs if p.text.startswith("1. [Baseline capture]"))
    assert n_ros >= 1 and n_first == n_ros


def test_run_of_show_carries_impact_on_cutover_steps():
    """V3.23.154: the outage/impact callout rides the step that causes it (AS/Barstow convention) —
    an additive `impact` field on the two cutover steps only."""
    from backend.cutover import _run_of_show

    steps = _run_of_show(mbb=["a"], hard=["b"], hard_ep=7, window=45, n_val=3, n_rem=0, blockers=[])
    by_phase = {s["phase"]: s for s in steps}
    assert by_phase["Cutover · make-before-break"]["impact"].startswith("No outage")
    hard = by_phase["Cutover · hard cutover"]["impact"]
    assert hard.startswith("OUTAGE") and "7 endpoint(s)" in hard
    assert "impact" not in by_phase["Baseline capture"]
    assert "impact" not in by_phase["Validation"]


def test_cutover_no_waves_still_carries_acceptance(tmp_path):
    """A degenerate snapshot (no derivable waves) takes the early-return path — the document must
    still carry the full furniture, including the closing acceptance gate (review fix, V3.23.151)."""
    pytest.importorskip("docx")
    from docx import Document

    from backend.cutover_docx import write_cutover_docx

    out = str(tmp_path / "cutover_empty.docx")
    write_cutover_docx(out, {"devices": {}}, "Empty Fleet")
    text = "\n".join(p.text for p in Document(out).paragraphs)
    assert "No migration waves were derived" in text
    assert "Document Control" in text
    assert "Document Acceptance" in text


def test_nrfu_deliverable_content(client):
    """The NRFU / Acceptance Test Plan is a web-layer synthesis with no engine test — validate it
    renders the document-control front matter and all three test phases, not just a valid zip."""
    snap_id = client.post("/api/demo/seed").json()["snapshot"]["id"]
    r = client.get(f"/api/snapshots/{snap_id}/deliverable/nrfu")
    if r.status_code == 503:
        pytest.skip("python-docx not installed on this runner")
    assert r.status_code == 200, r.text
    assert "nrfu" in r.headers.get("content-disposition", "")

    import io

    from docx import Document

    doc = Document(io.BytesIO(r.content))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Network Ready-For-Use" in text
    assert "Document control" in text and "Sign-off" in text     # front matter
    assert "Related documents" in text                           # family cross-reference (V3.23.150)
    assert "Phase I" in text and "Phase II" in text and "Phase III" in text
    assert "Entry criteria" in text and "Exit criteria" in text
    # Phase II reuses the engine's validation_plan — its test rows carry a command + expected baseline
    all_rows = [c.text for t in doc.tables for row in t.rows for c in row.cells]
    assert any("NRFU-II-" in c for c in all_rows)                # generated test IDs
    assert any("show " in c for c in all_rows)                   # runnable commands


def test_nrfu_devices_in_scope_reads_canonical_scale(tmp_path):
    """C9 (SSOT): the NRFU 'Devices in scope' header must read the canonical executive_brief.scale —
    the published single source the explorer/deck/HLD read — not a local len(devices) recompute. The
    web-layer NRFU writer was the last surface recounting fleet scale. Discriminating fixture: scale
    says 303 while the raw devices array holds 2, so a recompute regression renders 2, not 303."""
    pytest.importorskip("docx")
    from docx import Document

    from backend.nrfu_docx import write_nrfu_docx
    snap = {
        "script_version": "V3.23.0",
        "devices": {"a": {}, "b": {}},
        "executive_brief": {"scale": {"n_devices": 303, "n_vlans": 202, "n_endpoints": 5127}},
        "collection_completeness": {"summary": {"inventory": 303, "complete": 250, "not_collected": 53}},
        "lifecycle_risk": {"per_device": []}, "validation_plan": {"items": []},
        "service_map": {"services": []}, "application_intelligence": {"domains": []},
        "multicast_intelligence": {}, "design_blueprint": {"decisions": [], "design_nrfu": {"items": []}},
    }
    out = str(tmp_path / "nrfu.docx")
    write_nrfu_docx(out, snap, "Unit Test Fleet")
    rows = [c.text for t in Document(out).tables for row in t.rows for c in row.cells]
    i = next(k for k, c in enumerate(rows) if "Devices in scope" in c)
    assert rows[i + 1] == "303"    # canonical scale.n_devices, not len(devices)=2

    # WEBAP-03: the writer used `scale.get("n_devices") or len(devices)`, so a canonical 0 (a legitimate
    # all-not-yet-collected inventory, which storage.py preserves with an `is None` check) was masked by the
    # raw-array recount. A canonical 0 with a non-empty devices map must render 0, not 3.
    snap0 = dict(snap, devices={"a": {}, "b": {}, "c": {}},
                 executive_brief={"scale": {"n_devices": 0, "n_vlans": 0, "n_endpoints": 0}})
    out0 = str(tmp_path / "nrfu0.docx")
    write_nrfu_docx(out0, snap0, "Unit Test Fleet")
    rows0 = [c.text for t in Document(out0).tables for row in t.rows for c in row.cells]
    j = next(k for k, c in enumerate(rows0) if "Devices in scope" in c)
    assert rows0[j + 1] == "0"     # canonical 0 honoured, NOT len(devices)=3


def test_snapshot_meta_n_devices_reads_canonical_scale(client):
    """SSOT (python<->dashboard): the stored snapshot-meta n_devices (shown in the dashboard's snapshot
    list) must read the canonical executive_brief.scale.n_devices, not a server-side len(devices) recount.
    Discriminating fixture: scale says 303 while the raw devices array holds only 2."""
    import json
    cid = client.post("/api/campaigns", json={"name": "scale"}).json()["id"]
    snap = {"script_version": "V3.23.0", "devices": {"a": {}, "b": {}},
            "health_scores": [{"switch": "a", "band": "Good", "score": 80}],
            "executive_brief": {"scale": {"n_devices": 303, "n_vlans": 202, "n_endpoints": 5127},
                                "posture": {"avg_health": 80, "n_critical": 0}}}
    r = client.post(f"/api/campaigns/{cid}/snapshots",
                    files={"file": ("snap.json", json.dumps(snap).encode(), "application/json")},
                    data={"label": "scale-test"})
    assert r.status_code == 201, r.text
    assert r.json()["n_devices"] == 303          # canonical scale.n_devices, NOT len(devices)=2


def test_upload_robust_to_truthy_non_list_sections(client):
    """Robustness (malformed-upload DoS): summarize() runs on every upload and reads health_scores /
    punchlist / migration_readiness. `(snap.get(k) or [])` only coerces FALSY values, so a truthy
    NON-list (e.g. an int) flowed into a list-comprehension and raised TypeError -> an unhandled HTTP 500
    on a structurally-valid-but-hostile upload. It must now degrade to 201 (empty section)."""
    import json
    cid = client.post("/api/campaigns", json={"name": "robust-list"}).json()["id"]
    snap = {"devices": {"sw1": {}}, "health_scores": 5, "punchlist": "oops",
            "migration_readiness": {"not": "a list"}}
    r = client.post(f"/api/campaigns/{cid}/snapshots",
                    files={"file": ("snap.json", json.dumps(snap).encode(), "application/json")},
                    data={"label": "hostile"})
    assert r.status_code == 201, r.text          # was 500 before the _as_list coercion


def test_add_snapshot_robust_to_truthy_non_dict_executive_brief(client):
    """Robustness: a snapshot whose health_scores is a valid list (so summarize() succeeds) but whose
    executive_brief is a truthy NON-dict ('CORRUPT') reaches Store.add_snapshot, whose canonical
    n_devices read chained .get() through `(executive_brief or {})` -- which does NOT guard a truthy
    non-dict -> AttributeError -> HTTP 500. It must now degrade to 201."""
    import json
    cid = client.post("/api/campaigns", json={"name": "robust-eb"}).json()["id"]
    snap = {"devices": {"sw1": {}}, "health_scores": [{"switch": "sw1", "band": "Good", "score": 80}],
            "executive_brief": "CORRUPT"}
    r = client.post(f"/api/campaigns/{cid}/snapshots",
                    files={"file": ("snap.json", json.dumps(snap).encode(), "application/json")},
                    data={"label": "corrupt-eb"})
    assert r.status_code == 201, r.text          # was 500 before the isinstance guard
    assert r.json()["n_devices"] == 1            # falls back to len(devices) when scale is unreadable


def test_snapshot_meta_n_devices_canonical_zero_not_recounted(client):
    """SSOT `or`-masks-zero: a snapshot canonically publishing n_devices == 0 (an empty-inventory
    collection) that nonetheless carries a non-empty raw devices map must record the CANONICAL 0, not
    fall through `0 or len(devices)` to the client-side recount. Requires the `is not None` fix."""
    import json
    cid = client.post("/api/campaigns", json={"name": "zero"}).json()["id"]
    snap = {"devices": {"sw1": {}, "sw2": {}, "sw3": {}},
            "executive_brief": {"scale": {"n_devices": 0}}}
    r = client.post(f"/api/campaigns/{cid}/snapshots",
                    files={"file": ("snap.json", json.dumps(snap).encode(), "application/json")},
                    data={"label": "zero-canonical"})
    assert r.status_code == 201, r.text
    assert r.json()["n_devices"] == 0            # canonical 0, NOT len(devices)=3


def test_nrfu_carries_design_traceability_and_scope_limits(tmp_path):
    """N29+N30: the NRFU/ATP must trace its coverage back to the target-state design decisions, and
    state its SCOPE LIMITS (what it does NOT validate). A needs-requirement design area + not-collected
    devices are explicit coverage boundaries, not silent gaps."""
    pytest.importorskip("docx")
    from docx import Document

    from backend.nrfu_docx import write_nrfu_docx
    snap = {
        "script_version": "V3.23.0", "devices": {"a": {}, "b": {}},
        "executive_brief": {"scale": {"n_devices": 303}},
        "collection_completeness": {"summary": {"inventory": 303, "complete": 250, "not_collected": 50}},
        "lifecycle_risk": {"per_device": []}, "validation_plan": {"items": []},
        "service_map": {"services": []}, "application_intelligence": {"domains": []},
        "multicast_intelligence": {}, "software_risk": {"summary": {"n_config_not_assessable": 50}},
        "design_blueprint": {"decisions": [
            {"id": "fhrp-first-hop-gateway-redundancy", "title": "Introduce first-hop gateway redundancy",
             "domain": "availability", "priority": "Critical", "status": "recommended"},
            {"id": "dc-three-tier-vs-collapsed-core", "title": "Three-tier vs collapsed core",
             "domain": "dc-fabric", "priority": "High", "status": "needs-requirement"}]},
        "design_nrfu": {"items": [
            {"decision_id": "fhrp-first-hop-gateway-redundancy", "phase": "post-cutover-functional"}]},
    }
    out = str(tmp_path / "nrfu.docx")
    write_nrfu_docx(out, snap, "Unit Test Fleet")
    d = Document(out)
    heads = [p.text for p in d.paragraphs if p.style.name.startswith("Heading")]
    text = "\n".join(p.text for p in d.paragraphs)
    rows = [c.text for t in d.tables for r in t.rows for c in r.cells]
    assert any("coverage" in h.lower() and "scope" in h.lower() for h in heads), heads
    assert "Introduce first-hop gateway redundancy" in rows                 # recommended decision traced
    assert any("not testable" in r.lower() for r in rows)                   # needs-requirement flagged (N30)
    assert "50" in text and ("not validated" in text.lower() or "not collected" in text.lower())  # N29 scope limit


def test_execution_run_lifecycle(client):
    """The war-room flow end to end: start a run from the cutover plan, check off a step, record
    validation results, scribe a deviation, close out a wave, finish — then the run is read-only."""
    snap_id = client.post("/api/demo/seed").json()["snapshot"]["id"]

    r = client.post(f"/api/snapshots/{snap_id}/executions", json={"label": "", "operator": "lead"})
    assert r.status_code == 201, r.text
    ex = r.json()
    eid = ex["id"]
    assert ex["label"] == "Cutover run 1"            # auto-labelled
    assert ex["status"] == "in_progress" and ex["outcome"] is None

    # the run froze the plan: same wave groups, pilot-first order, steps/checks all pending
    plan = client.get(f"/api/snapshots/{snap_id}/cutover").json()
    assert [w["group"] for w in ex["waves"]] == [w["group"] for w in plan["waves"]]
    w0 = ex["waves"][0]
    assert [s["phase"] for s in w0["steps"]] == [s["phase"] for s in plan["waves"][0]["run_of_show"]]
    assert all(s["status"] == "pending" for w in ex["waves"] for s in w["steps"])
    assert len(w0["checks"]) == len(plan["waves"][0]["validation"])
    assert ex["progress"]["pct"] == 0

    g = w0["group"]
    # step check-off is timestamped and attributed
    ex = client.post(f"/api/executions/{eid}/step",
                     json={"wave": g, "index": 0, "status": "done", "operator": "lead"}).json()
    s0 = ex["waves"][0]["steps"][0]
    assert s0["status"] == "done" and s0["at"] and s0["by"] == "lead"
    assert ex["progress"]["n_steps_done"] == 1 and ex["progress"]["pct"] > 0

    # a failing validation check is recorded AND auto-scribed as a deviation
    ex = client.post(f"/api/executions/{eid}/check",
                     json={"wave": g, "index": 0, "result": "fail",
                           "observed": "neighbor missing", "operator": "lead"}).json()
    assert ex["waves"][0]["checks"][0]["result"] == "fail"
    assert ex["progress"]["checks"]["fail"] == 1
    assert any(e["kind"] == "deviation" and "neighbor missing" in e["text"] for e in ex["events"])

    # explicit deviation + wave closeout
    client.post(f"/api/executions/{eid}/event",
                json={"kind": "deviation", "text": "re-seated SFP", "wave": g})
    ex = client.post(f"/api/executions/{eid}/closeout",
                     json={"wave": g, "decision": "COMPLETE", "note": "with workaround"}).json()
    assert ex["waves"][0]["closeout"]["decision"] == "COMPLETE"
    assert ex["progress"]["waves"][0]["state"] == "complete"

    # finish: outcome derives PARTIAL (other waves never closed), then the run is immutable
    ex = client.post(f"/api/executions/{eid}/finish", json={"status": "completed"}).json()
    assert ex["status"] == "completed" and ex["outcome"] == "PARTIALLY IMPLEMENTED"
    r = client.post(f"/api/executions/{eid}/step", json={"wave": g, "index": 1, "status": "done"})
    assert r.status_code == 409

    # listed under the snapshot; bad inputs rejected; deletable
    runs = client.get(f"/api/snapshots/{snap_id}/executions").json()
    assert [(x["id"], x["status"]) for x in runs] == [(eid, "completed")]
    assert client.post(f"/api/executions/{eid + 99}/step",
                       json={"wave": g, "index": 0, "status": "done"}).status_code == 404
    assert client.delete(f"/api/executions/{eid}").status_code == 204
    assert client.get(f"/api/executions/{eid}").status_code == 404


def test_execution_outcome_vocabulary(client):
    """Outcome derivation follows the PIR vocabulary: a clean run is SUCCESSFUL, a rolled-back wave
    dominates, and an abort is ABORTED regardless of progress."""
    snap_id = client.post("/api/demo/seed").json()["snapshot"]["id"]

    def run():
        return client.post(f"/api/snapshots/{snap_id}/executions", json={}).json()

    def close_all(eid, ex, decision):
        for w in ex["waves"]:
            ex = client.post(f"/api/executions/{eid}/closeout",
                             json={"wave": w["group"], "decision": decision}).json()
        return ex

    # every wave completed, nothing skipped/failed -> SUCCESSFUL
    ex = run()
    ex = close_all(ex["id"], ex, "COMPLETE")
    ex = client.post(f"/api/executions/{ex['id']}/finish", json={"status": "completed"}).json()
    assert ex["outcome"] == "SUCCESSFUL"

    # a rolled-back wave dominates the verdict, even when every other wave completed
    ex = run()
    eid = ex["id"]
    ex = client.post(f"/api/executions/{eid}/closeout",
                     json={"wave": ex["waves"][0]["group"], "decision": "ROLLED BACK"}).json()
    for w in ex["waves"]:
        if not w["closeout"]["decision"]:
            ex = client.post(f"/api/executions/{eid}/closeout",
                             json={"wave": w["group"], "decision": "COMPLETE"}).json()
    ex = client.post(f"/api/executions/{eid}/finish", json={"status": "completed"}).json()
    assert ex["outcome"] == "ROLLED BACK"

    # an aborted run is ABORTED
    ex = run()
    ex = client.post(f"/api/executions/{ex['id']}/finish", json={"status": "aborted"}).json()
    assert ex["outcome"] == "ABORTED"


def test_execution_record_integrity_guards(client):
    """The as-executed record is an audit artifact: negative indexes must not address from the end,
    and a closed-out wave's steps/checks are part of the signed record (409, like a finished run)."""
    snap_id = client.post("/api/demo/seed").json()["snapshot"]["id"]
    ex = client.post(f"/api/snapshots/{snap_id}/executions", json={}).json()
    eid, g = ex["id"], ex["waves"][0]["group"]

    # negative index -> 400, and no step was touched
    r = client.post(f"/api/executions/{eid}/step", json={"wave": g, "index": -1, "status": "done"})
    assert r.status_code == 400
    ex = client.get(f"/api/executions/{eid}").json()
    assert all(s["status"] == "pending" for s in ex["waves"][0]["steps"])

    # close the wave out, then late mutations are conflicts and the record is unchanged
    client.post(f"/api/executions/{eid}/closeout", json={"wave": g, "decision": "COMPLETE"})
    r = client.post(f"/api/executions/{eid}/step", json={"wave": g, "index": 0, "status": "done"})
    assert r.status_code == 409
    r = client.post(f"/api/executions/{eid}/check", json={"wave": g, "index": 0, "result": "pass"})
    assert r.status_code == 409
    r = client.post(f"/api/executions/{eid}/closeout", json={"wave": g, "decision": "ROLLED BACK"})
    assert r.status_code == 409
    ex = client.get(f"/api/executions/{eid}").json()
    assert ex["waves"][0]["closeout"]["decision"] == "COMPLETE"
    assert all(s["status"] == "pending" for s in ex["waves"][0]["steps"])


def test_execution_duplicate_wave_groups_disambiguated():
    """Uploaded snapshots can carry duplicate/missing wave group names — the run must give every
    wave a unique address or mutations for the second wave land on the first."""
    from backend import execution

    snap = {
        "devices": {"sw1": {}, "sw2": {}},
        "wave_sequencing": [{"make_before_break": ["sw1"], "hard_cutover": []},
                            {"make_before_break": ["sw2"], "hard_cutover": []}],
    }
    state = execution.start_run(snap, "run", "")
    groups = [w["group"] for w in state["waves"]]
    assert len(set(groups)) == len(groups) and all(groups)
    execution.apply_step(state, groups[1], 0, "done", "", "op")
    assert state["waves"][1]["steps"][0]["status"] == "done"
    assert state["waves"][0]["steps"][0]["status"] == "pending"


def test_execution_pir_report(client):
    """The PIR / as-executed record renders the run: outcome, planned-vs-actual, the timestamped
    deviation log, and the per-wave validation results."""
    snap_id = client.post("/api/demo/seed").json()["snapshot"]["id"]
    ex = client.post(f"/api/snapshots/{snap_id}/executions",
                     json={"label": "Window 7", "operator": "tanveer"}).json()
    eid = ex["id"]
    g = ex["waves"][0]["group"]
    client.post(f"/api/executions/{eid}/step", json={"wave": g, "index": 0, "status": "done",
                                                     "operator": "tanveer"})
    client.post(f"/api/executions/{eid}/check", json={"wave": g, "index": 0, "result": "pass",
                                                      "operator": "tanveer"})
    client.post(f"/api/executions/{eid}/event",
                json={"kind": "deviation", "text": "uplink LED amber, re-seated SFP", "wave": g})
    client.post(f"/api/executions/{eid}/closeout", json={"wave": g, "decision": "COMPLETE"})
    client.post(f"/api/executions/{eid}/finish", json={"status": "completed"})

    r = client.get(f"/api/executions/{eid}/report")
    if r.status_code == 503:
        pytest.skip("python-docx not installed on this runner")
    assert r.status_code == 200, r.text
    assert r.content[:2] == b"PK"
    assert "pir" in r.headers.get("content-disposition", "")

    import io

    from docx import Document

    doc = Document(io.BytesIO(r.content))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Post-Implementation Review" in text
    assert "Window 7" in text
    assert "Planned vs actual" in text
    assert "Timeline & deviation log" in text
    assert "Related documents" in text                            # family cross-reference (V3.23.150)
    all_rows = [c.text for t in doc.tables for row in t.rows for c in row.cells]
    assert any("re-seated SFP" in c for c in all_rows)         # the scribed deviation is in the log
    assert any("PASS" in c for c in all_rows)                  # the recorded validation result
    assert any("tanveer" in c for c in all_rows)               # actions are attributed


def _fixture_collection_zip(tmp_path, wrap="export/fleet", include_devices_json=False) -> bytes:
    """A real offline-collection ZIP built from the engine test-suite's synthetic fixtures."""
    import io
    import zipfile

    import synthetic_fixtures as fx

    root = tmp_path / "zipsrc"
    coll = root / wrap if wrap else root
    fx.write_collection(str(coll))
    if include_devices_json:
        import json
        (coll / "devices.json").write_text(json.dumps(fx.DEVICES), encoding="utf-8")
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        for p in root.rglob("*"):
            if p.is_file():
                zf.write(p, p.relative_to(root).as_posix())
    return buf.getvalue()


def test_ingest_collection_runs_real_engine(client, tmp_path):
    """The flagship ingest path: a ZIP of raw show outputs (wrapped in a folder, NO devices.json —
    platforms autodetect) is run through the real engine and stored as a first-class snapshot."""
    raw = _fixture_collection_zip(tmp_path)
    cid = client.post("/api/campaigns", json={"name": "ingest"}).json()["id"]
    r = client.post(f"/api/campaigns/{cid}/ingest",
                    files={"file": ("fleet.zip", raw, "application/zip")},
                    data={"label": "Ingested wave"})
    assert r.status_code == 201, r.text
    meta = r.json()
    assert meta["label"] == "Ingested wave"
    assert meta["n_devices"] == 3                       # core1 / core2 / access1
    assert meta["ingest"]["devices_json"] == "synthesized"
    assert sorted(meta["ingest"]["devices"]) == ["access1", "core1", "core2"]
    # WEBAP-02: the headline directory count must equal the fleet actually assessed -- skipped non-round-trippable
    # folders excluded -- never over-report (here nothing is skipped, so it pins the formula).
    assert meta["ingest"]["n_device_dirs"] == 3
    assert (meta["ingest"]["n_device_dirs"]
            == len(meta["ingest"]["devices"]) - len(meta["ingest"]["skipped_dirs"]))
    # WEBAP-01: the engine log tail surfaced to the (remote) client must not leak the server working-dir
    # absolute path; the engine references it (writes its outputs there), so the scrub placeholder must appear.
    tail = meta["ingest"].get("engine_log_tail", "")
    assert "<workdir>" in tail and tmp_path.name not in tail

    # the stored snapshot is the engine's own: summary derives and the cutover plan synthesizes
    s = meta["summary"]
    assert s["n_switches"] == 3 and s["punchlist"]["total"] >= 1
    plan = client.get(f"/api/snapshots/{meta['id']}/cutover").json()
    assert plan["summary"]["n_waves"] >= 1


def test_ingest_mismatched_devices_json_falls_back(client, tmp_path):
    """A bundled devices.json whose hostnames match no folder must not be trusted blindly: the run
    falls back to folder-name synthesis (assessing what is actually in the archive) instead of
    pointing the engine at non-existent directories and storing an empty snapshot."""
    import io
    import json
    import zipfile

    import synthetic_fixtures as fx

    root = tmp_path / "zipsrc" / "fleet"
    fx.write_collection(str(root))
    # FQDN hostnames that match no folder
    (root / "devices.json").write_text(json.dumps(
        [{"hostname": f"{h}.example.net", "platform": "ios"} for h in ("core1", "core2", "access1")]),
        encoding="utf-8")
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        for p in (tmp_path / "zipsrc").rglob("*"):
            if p.is_file():
                zf.write(p, p.relative_to(tmp_path / "zipsrc").as_posix())

    cid = client.post("/api/campaigns", json={"name": "mismatch"}).json()["id"]
    r = client.post(f"/api/campaigns/{cid}/ingest",
                    files={"file": ("fleet.zip", buf.getvalue(), "application/zip")})
    assert r.status_code == 201, r.text
    meta = r.json()
    assert meta["ingest"]["devices_json"] == "synthesized"      # the useless file was not trusted
    assert meta["n_devices"] == 3
    # the snapshot carries REAL parsed data (the empty-snapshot failure mode)
    assert meta["summary"]["punchlist"]["total"] >= 1


def test_ingest_rejects_bad_archives(client, tmp_path):
    import io
    import zipfile

    cid = client.post("/api/campaigns", json={"name": "x"}).json()["id"]

    def post(content, name="c.zip"):
        return client.post(f"/api/campaigns/{cid}/ingest",
                           files={"file": (name, content, "application/zip")})

    # not a zip at all
    assert post(b"definitely not a zip").status_code == 400
    # path traversal entry is refused before anything runs
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("../evil.txt", "x")
    r = post(buf.getvalue())
    assert r.status_code == 400 and "traversal" in r.json()["detail"]
    # show outputs at the archive root (no per-device folders) get an actionable message
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("show_interface_status.txt", "Port Name Status")
    r = post(buf.getvalue())
    assert r.status_code == 400 and "own folder" in r.json()["detail"]
    # a zip with no device outputs at all
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("readme.md", "hello")
    r = post(buf.getvalue())
    assert r.status_code == 400 and "No device outputs" in r.json()["detail"]
    # unknown campaign
    raw = _fixture_collection_zip(tmp_path)
    assert client.post("/api/campaigns/999999/ingest",
                       files={"file": ("c.zip", raw, "application/zip")}).status_code == 404


def test_bad_upload_rejected(client):
    cid = client.post("/api/campaigns", json={"name": "x"}).json()["id"]
    r = client.post(f"/api/campaigns/{cid}/snapshots",
                    files={"file": ("bad.json", b"not json", "application/json")},
                    data={"label": "bad"})
    assert r.status_code == 400


# ---------------------------------------------------------------------------
# Single-source-of-truth: the dashboard reader must trust the engine's canonical
# executive_brief.scale / .posture, NOT re-derive the headline numbers from the raw
# arrays. These fixtures make canonical DISAGREE with what a recompute would yield,
# so the asserted values can only come from the canonical block — any regression to
# a client-side recount turns them red. (Engine side is locked by
# tests/test_pipeline_inprocess.py; this is the webapp half of the contract.)
# ---------------------------------------------------------------------------

def _divergent_snap():
    """A snapshot whose canonical brief deliberately disagrees with the raw arrays:
    len(devices)=3 / avg(scores)=80.0 / Critical-band tally=1, but the brief says
    42 / 63.4 / 5. A canonical-first reader returns the brief's numbers."""
    return {
        "script_version": "V3.23.0",
        "generated_at": "2026-06-18T00:00:00",
        "devices": {"a": {}, "b": {}, "c": {}},                       # fallback n_switches == 3
        "health_scores": [                                            # fallback avg == 80.0, Critical tally == 1
            {"host": "a", "score": 90, "band": "Good"},
            {"host": "b", "score": 80, "band": "Fair"},
            {"host": "c", "score": 70, "band": "Critical"},
        ],
        "punchlist": [{"severity": "High", "category": "L2"},
                      {"severity": "Low", "category": "Hygiene"}],
        "executive_brief": {                                         # the canonical source of truth
            "scale": {"n_devices": 42, "n_domains": 9, "n_endpoints": 5127, "n_vlans": 172},
            "posture": {"avg_health": 63.4, "n_critical": 5, "n_poor": 2, "worst_band": "Poor"},
            "axes": [], "top_gating": [], "posture_statement": "—",
        },
    }


def test_trend_point_honors_canonical_scale_posture_over_recompute():
    """SSOT: trend_point must read executive_brief.scale/.posture, not recount the raw arrays."""
    from backend import engine
    tp = engine.trend_point(_divergent_snap())
    assert tp["n_switches"] == 42, "n_switches must be scale.n_devices (canonical), not len(devices)=3"
    assert tp["avg_health"] == 63.4, "avg_health must be posture.avg_health, not avg(scores)=80.0"
    assert tp["n_critical"] == 5, "n_critical must be posture.n_critical, not the Critical-band tally=1"


def test_summarize_projects_the_canonical_headline():
    """SSOT: the API-facing summary projection inherits the canonical headline (it re-uses trend_point)."""
    from backend.summary import summarize
    s = summarize(_divergent_snap())
    assert s["n_switches"] == 42 and s["avg_health"] == 63.4 and s["n_critical"] == 5


def test_summarize_near_eos_reads_canonical_n_near():
    """Coverage-honesty: the dashboard summary's near-end-of-support figure must read the canonical
    lifecycle field n_near (Near-LDoS). It previously read a non-existent n_near_eos, silently
    blanking the count while past_eos/past_ldos rendered — a half-shown lifecycle posture."""
    from backend.summary import summarize
    life = summarize({"lifecycle_risk": {"summary": {"n_past_ldos": 152, "n_past_eos": 0, "n_near": 61}}})["lifecycle"]
    assert life["near_eos"] == 61, life
    assert life["past_ldos"] == 152 and life["past_eos"] == 0


def test_trend_point_falls_back_when_brief_absent():
    """Back-compat: a pre-brief snapshot (no executive_brief) still resolves via the local recompute,
    so the canonical-first read degrades gracefully instead of returning blanks."""
    from backend import engine
    snap = _divergent_snap()
    snap.pop("executive_brief")
    tp = engine.trend_point(snap)
    assert tp["n_switches"] == 3 and tp["avg_health"] == 80.0 and tp["n_critical"] == 1


def test_causal_flows_endpoint(client):
    """The unified Causal Flow model (engine compute_causal_flows) — the SAME normalization the explorer's
    Causal Flow mode renders, served so the dashboard never re-derives causal intent. Cross-layer once."""
    snap_id = client.post("/api/demo/seed").json()["snapshot"]["id"]
    r = client.get(f"/api/snapshots/{snap_id}/causal_flows")
    assert r.status_code == 200
    body = r.json()
    assert body["summary"]["n_flows"] > 0
    fams = {f["key"] for f in body["families"]}
    assert "struct" in fams and "xlayer" in fams              # structural SPOFs + cross-layer always present
    # every flow carries the four narrative stages + the magnitude / shape contract the UI renders
    f0 = body["flows"][0]
    for k in ("trigger", "mechanism", "impact", "mitigation", "severity",
              "blast", "blast_unit", "shape", "family_label", "icon"):
        assert k in f0, f"flow missing {k}"
    # cross-layer compounds promote to a bowtie
    assert any(f["shape"] == "bowtie" for f in body["flows"])
    # keys unique END-TO-END (cross_layer CL-xx ids repeat — locks the index-based-key fix at the wire)
    keys = [f["key"] for f in body["flows"]]
    assert len(keys) == len(set(keys)), "flow keys must be unique at the endpoint"
    # the cross-layer VLAN magnitude reaches the response ("N VLANs", not "1 device")
    assert any(f["blast_unit"] in ("VLAN", "VLANs") for f in body["flows"]), "cross-layer VLAN magnitude must surface"
    # de-dup: the punch-list 'Cross-layer' rows are NOT re-emitted (only the cross_layer array feeds xlayer)
    assert not any(f["family"] == "Cross-layer" for f in body["flows"])
    n_xlayer = sum(1 for f in body["flows"] if f["family"] == "xlayer")
    cl = client.get(f"/api/snapshots/{snap_id}/section/cross_layer")
    if cl.status_code == 200:
        assert n_xlayer == len(cl.json()["data"]), "xlayer count must equal cross_layer length (no double-count)"
    # 404 for a missing snapshot
    assert client.get("/api/snapshots/999999/causal_flows").status_code == 404


def test_causal_flows_total_over_malformed_snapshot():
    """The engine fn the /causal_flows route calls must be TOTAL over any dict — a malformed-but-truthy
    container field (a string where a list is expected, an unhashable severity) must NOT raise, else the
    route 500s. Mirrors the JS Array.isArray defensiveness; the route also has a try/except backstop."""
    from cisco_toolkit.causal import compute_causal_flows
    for bad in [
        {"causality": "notalist"}, {"cross_layer": "x"}, {"punchlist": "x"},
        {"design_blueprint": {"decisions": "x"}}, {"devices": "x"},
        {"causality": [{"severity": {"a": 1}, "hosts": ["a"]}]},   # unhashable severity
        {"punchlist": [{"category": "X", "devices": "notalist", "detail": 12345}]},
    ]:
        r = compute_causal_flows(bad)
        assert r["summary"]["n_flows"] >= 0 and isinstance(r["flows"], list)


def test_causal_flows_computes_design_family_like_design_endpoint(client):
    """Internal consistency: the /causal_flows route computes design_blueprint on the fly when the snapshot
    didn't store one (exactly as /design does), so the design-decision family appears. The sample carries NO
    stored design_blueprint but yields recommended decisions — and the count MUST equal /design's recommended
    decisions (cross-endpoint single-source-of-truth)."""
    snap_id = client.post("/api/demo/seed").json()["snapshot"]["id"]
    body = client.get(f"/api/snapshots/{snap_id}/causal_flows").json()
    fams = {f["key"] for f in body["families"]}
    assert "design" in fams, "design family must be present (computed when not stored, like /design)"
    n_design = sum(1 for f in body["flows"] if f["family"] == "design")
    assert n_design > 0
    bp = client.get(f"/api/snapshots/{snap_id}/design").json()
    rec = sum(1 for d in bp["decisions"] if d.get("status") == "recommended")
    assert n_design == rec, f"causal-flow design family ({n_design}) must equal /design recommended ({rec})"


def test_build_graph_robust_to_non_dict_interfaces_devices():
    """WEBAP-02: build_graph used `snap.get("interfaces") or {}` / `(snap.get("devices") or {})` -- a TRUTHY
    non-dict (a JSON string/list in a malformed upload) slipped through and .keys() raised AttributeError -> an
    unhandled 500 on GET /graph. It must degrade instead."""
    from webapp.backend.graph import build_graph
    g = build_graph({"interfaces": "oops", "devices": [1, 2, 3], "health_scores": "nope"})   # must not raise
    assert isinstance(g, dict) and "nodes" in g and "edges" in g


def test_get_snapshot_section_robust_to_scalar_section(tmp_path):
    """WEBAP-01: get_snapshot_section json.loads()'d sqlite json_extract output, but a JSON SCALAR section
    (string/number) comes back as the native value -- json.loads(int) raised TypeError and json.loads(a bare
    string) raised JSONDecodeError, neither caught -> 500. A scalar section must be returned, not raise."""
    from webapp.backend.storage import Store
    st = Store(str(tmp_path / "t.db"))
    cid = st.create_campaign("c", "cust")["id"]
    snap = {"devices": {"sw1": {}}, "design_blueprint": 5, "architecture_coverage": "weird-scalar"}
    sid = st.add_snapshot(cid, "s", snap, {"n_switches": 1})["id"]
    assert st.get_snapshot_section(sid, "design_blueprint") == 5           # native int, not a 500
    assert st.get_snapshot_section(sid, "architecture_coverage") == "weird-scalar"
    assert st.get_snapshot_section(sid, "devices") == {"sw1": {}}          # objects still decode


def test_reconcile_gate_flags_a_drifting_snapshot(caplog):
    """W3-5: deliverables.generate runs a fail-soft SSOT pre-emission check — a snapshot whose published facts
    disagree with the raw evidence is loudly logged before the artifact is written (never silently emitted),
    but a single drift never blocks the deliverable. Total/fail-open on bad input."""
    import logging
    from backend.deliverables import _reconcile_gate
    assert _reconcile_gate({}, "mop") == []                                # nothing published -> clean
    assert _reconcile_gate(None, "mop") == []                              # total on bad input, no crash
    drift = {"executive_brief": {"scale": {"n_devices": 999}},
             "health_scores": [{"switch": "a"}, {"switch": "b"}]}
    with caplog.at_level(logging.WARNING):
        viol = _reconcile_gate(drift, "mop")
    assert viol and any("n_devices" in v for v in viol)                    # the drift is returned
    assert any("unreconciled" in r.getMessage() for r in caplog.records)   # ...and loudly logged


def test_build_graph_tolerates_health_row_without_switch_key():
    """[multi-domain audit #10] a health_scores row lacking a string 'switch' key injected a None node id and made
    sorted(node_ids) raise TypeError (str vs None) -> an unhandled 500 on /graph."""
    from webapp.backend.graph import build_graph
    g = build_graph({"devices": {"sw1": {}}, "interfaces": {}, "health_scores": [{"band": "Good", "score": 90}]})
    assert "nodes" in g and all(n["id"] is not None for n in g["nodes"])


def test_summarize_survives_nondict_lifecycle_summary():
    """[audit-3 #13 totality] a snapshot whose lifecycle_risk.summary is a truthy NON-dict (an older engine's
    'lifecycle not computed' string) survived the `or {}` guard and reached lr.get(...) -> AttributeError -> the
    unauthenticated upload endpoint 500'd. summarize() must degrade on every field."""
    from webapp.backend import summary
    snap = {"devices": {"sw1": {"model": "WS-C3850-48P"}},
            "lifecycle_risk": {"summary": "lifecycle not computed (older engine)", "per_device": []}}
    out = summary.summarize(snap)        # must not raise
    assert isinstance(out, dict) and "lifecycle" in out


def test_section_device_dossiers_recomputes_stale_unassessed(client):
    """[audit-4 #20 false-health] a pre-V3.23.174 snapshot bands the uncollected fleet 'Low / routine migration
    handling' instead of 'Unassessed'; /section/device_dossiers served it verbatim (no recompute fallback unlike
    the sibling sections), so the live AssessHub Risk Register read blind devices as routine. It must recompute
    server-side when the stored section is stale."""
    import json
    cid = client.post("/api/campaigns", json={"name": "h", "client": "x"}).json()["id"]
    snap = {
        "devices": {"blind1": {}, "good1": {"model": "C9300", "sw_version": "17.9"}},
        "health_scores": [{"switch": "blind1", "band": "Insufficient Data", "score": 90},
                          {"switch": "good1", "band": "Good", "score": 85}],
        "failure_impact": [], "punchlist": [],
        "device_dossiers": {"summary": {"bands": {"Low": 1, "Guarded": 0, "Elevated": 0, "Severe": 0}},
                            "per_device": [{"host": "blind1", "risk_band": "Low", "health_band": "Insufficient Data",
                                            "verdict": "No stacked risk — routine migration handling."}]},
    }
    sid = client.post(f"/api/campaigns/{cid}/snapshots",
                      files={"file": ("s.json", json.dumps(snap).encode(), "application/json")},
                      data={"label": "stale"}).json()["id"]
    data = client.get(f"/api/snapshots/{sid}/section/device_dossiers").json()["data"]
    assert data["summary"]["bands"].get("Unassessed", 0) >= 1            # blind device surfaced, not silently 'Low'
    blind = next(d for d in data["per_device"] if d["host"] == "blind1")
    assert blind["risk_band"] == "Unassessed" and "routine" not in blind["verdict"].lower()
