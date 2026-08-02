"""Shared python-docx conventions for AssessHub's web-layer deliverable writers.

cutover_docx / nrfu_docx / pir_docx all render in the engine writers' house style (Calibri 10.5pt
body, navy headings, "Light Grid Accent 1" banded tables). The helpers lived as three diverging
copies — pir_docx's column-width fix never reached the other two — so the style layer now lives here
once. python-docx stays an optional dependency: nothing here imports it at module load; helpers
import it inside the call, exactly like the writers themselves.

Every string these helpers write is routed through ``cisco_toolkit.textutils.xml_safe``, exactly like
the engine twin ``cisco_toolkit.docmeta.add_table`` / ``_kv``. Device free-text is collected with
``errors='ignore'``, which passes valid-UTF-8 C0 control chars, U+FFFE/U+FFFF noncharacters and lone
surrogates straight through — and ONE of them anywhere in a cell makes ``doc.save()`` raise
``ValueError: All strings must be XML compatible``. On this surface that is not a crash, it is a STORED
denial of service: the snapshot is persisted verbatim, so ``GET /api/snapshots/{id}/deliverable/nrfu``
(and /cutover, and the PIR export) then returns HTTP 500 for that snapshot permanently.
"""

from __future__ import annotations

from typing import Any, List

NAVY = (0x1F, 0x38, 0x64)
GREY = (0x59, 0x59, 0x59)


def ink(rgb: tuple) -> Any:
    """RGBColor from an (r, g, b) tuple."""
    from docx.shared import RGBColor

    return RGBColor(*rgb)


def new_document() -> Any:
    """A Document with the house body style applied."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    return doc


def safe(value: Any) -> Any:
    """``cisco_toolkit.textutils.xml_safe``, imported inside the call like the docx imports above
    (the engine's sys.path bootstrap runs in `backend.engine`, which the writers import first)."""
    from cisco_toolkit.textutils import xml_safe

    return xml_safe(value)


def kv(p: Any, label_text: str, value: str, color: tuple = NAVY) -> None:
    """Bold coloured label + plain value on an existing paragraph."""
    r = p.add_run(safe(label_text))
    r.bold = True
    r.font.color.rgb = ink(color)
    p.add_run(" " + (safe(value) if value else "—"))


def add_table(doc: Any, headers: List[str], rows: List[List[Any]],
              widths: List[float] | None = None) -> Any:
    """Banded house-style table with a bold header row and explicit column widths.

    Explicit widths only render once autofit is off, and renderers respect them most reliably when
    set on both the column and every cell — without this, multi-column tables silently fall back to
    equal widths (the bug this shared helper fixed in pir_docx and backports to its siblings).
    """
    from docx.shared import Inches

    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, hd in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = safe(str(hd))     # sanitize: one XML-illegal char in any cell aborts the whole .docx
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = "" if v is None else safe(str(v))
    if widths:
        t.autofit = False
        for i, w in enumerate(widths):
            t.columns[i].width = Inches(w)
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t
