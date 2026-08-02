"""Render a finished (or in-flight) execution run as a **Post-Implementation Review (PIR) /
As-Executed Cutover Record** .docx.

This is the closure artifact of change management: the ITIL post-implementation review is written
from the change window's timestamped record — what was planned vs what actually happened, which
validation checks passed against their captured baselines, what deviated, and whether the change met
its objectives, met them with deviations, was partially implemented, or was backed out.

Like ``cutover_docx`` / ``nrfu_docx`` this is an AssessHub web-layer synthesis (the execution run only
exists in AssessHub), matching the engine writers' python-docx conventions (Calibri 10.5pt body, navy
headings, "Light Grid Accent 1" tables) so it reads as one family with the rest of the deliverable set.
"""

from __future__ import annotations

from datetime import datetime
from typing import Any, Dict, List

from . import engine, execution
from .docx_style import GREY as _GREY
from .docx_style import NAVY as _NAVY
from .docx_style import add_table, ink, new_document

_GREEN = (0x1A, 0x7F, 0x37)
_AMBER = (0x9A, 0x67, 0x00)
_RED = (0xCF, 0x22, 0x2E)

_OUTCOME_INK = {
    execution.OUTCOME_SUCCESS: _GREEN,
    execution.OUTCOME_DEVIATIONS: _AMBER,
    execution.OUTCOME_PARTIAL: _AMBER,
    execution.OUTCOME_ROLLED_BACK: _RED,
    execution.OUTCOME_ABORTED: _RED,
}

_DECISION_LABEL = {None: "not closed out", "COMPLETE": "Complete",
                   "ROLLED BACK": "ROLLED BACK", "DEFERRED": "Deferred"}


def _fmt_ts(ts: Any) -> str:
    if not ts:
        return "—"
    try:
        return datetime.fromisoformat(str(ts)).strftime("%Y-%m-%d %H:%M:%S")
    except ValueError:
        return str(ts)


def _fmt_min(m: Any) -> str:
    try:
        m = int(m)
    except (TypeError, ValueError, OverflowError):   # OverflowError: a JSON Infinity duration
        return "—"
    if m <= 0:
        return "0 min"
    h, mm = divmod(m, 60)
    return f"{h}h{mm:02d}m" if h else f"{mm} min"


def write_pir_docx(output_path: str, state: Dict[str, Any], snapshot_label: str) -> None:
    """Write the as-executed record / PIR for one execution run to ``output_path``."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    from cisco_toolkit.textutils import xml_safe, xml_safe_deep

    # Same reasoning as cutover_docx: §3 emits the move-group name through a DIRECT
    # `doc.add_heading(f"Wave {order} — {group}")` that bypasses the sanitizing table/kv helpers, and
    # the run state carries operator-typed notes and device-derived labels verbatim. One XML-illegal
    # character aborts `doc.save()`, and the record it is stored in does not change — so the PIR
    # export 500s permanently for that run.
    state = xml_safe_deep(state if isinstance(state, dict) else {})
    snapshot_label = (xml_safe(snapshot_label) if isinstance(snapshot_label, str)
                      else (str(snapshot_label) if snapshot_label is not None else ""))

    doc = new_document()

    def table(headers: List[str], rows: List[List[Any]], widths: List[float] | None = None):
        return add_table(doc, headers, rows, widths)

    waves: List[Dict[str, Any]] = state.get("waves") or []
    events: List[Dict[str, Any]] = state.get("events") or []
    prog = execution.progress(state)
    outcome = state.get("outcome") or "IN PROGRESS"
    finished = state.get("status") in ("completed", "aborted")

    # ---- title page ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Post-Implementation Review")
    tr.bold = True
    tr.font.size = Pt(26)
    tr.font.color.rgb = ink(_NAVY)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(f"As-Executed Cutover Record · {state.get('label', 'Cutover run')}")
    sr.font.size = Pt(13)
    sr.font.color.rgb = ink(_GREY)
    verdict = doc.add_paragraph()
    verdict.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vr = verdict.add_run(f"Outcome: {outcome}")
    vr.bold = True
    vr.font.size = Pt(16)
    vr.font.color.rgb = ink(_OUTCOME_INK.get(outcome, _NAVY))
    if not finished:
        note = doc.add_paragraph()
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        nr = note.add_run("INTERIM — this run is still in progress; the record below is a point-in-time export.")
        nr.italic = True
        nr.font.color.rgb = ink(_GREY)

    # ---- document control ----
    doc.add_heading("Document control", level=1)
    elapsed_min = prog["elapsed_seconds"] // 60
    table(["Field", "Value"], [
        ["Document", "Post-Implementation Review / As-Executed Cutover Record"],
        ["Change / run", state.get("label", "")],
        ["Assessment snapshot", snapshot_label],
        ["Run lead", state.get("operator") or "—"],
        ["Started", _fmt_ts(state.get("started_at"))],
        ["Ended", _fmt_ts(state.get("ended_at"))],
        ["Actual duration", _fmt_min(elapsed_min)],
        ["Planned window (hard-cutover)", _fmt_min(prog["planned_window_minutes"])],
        ["Status / outcome", f"{state.get('status', '')} · {outcome}"],
        ["Generated", datetime.now().strftime("%Y-%m-%d %H:%M")],
        ["Engine", engine.ENGINE_SCHEMA_VERSION],
    ], widths=[2.2, 4.3])
    # Shared family cross-reference (imported here, after the engine sys.path bootstrap).
    # V3.23.159 review fix: the PIR has no FAMILY key to self-exclude, so it lists the FULL set —
    # the old exclude=("deck",) was a copy-paste of a sibling's self-exclusion that silently
    # dropped the Executive Presentation Deck from the as-executed record's cross-references.
    from cisco_toolkit.docmeta import add_related_documents
    add_related_documents(
        doc, exclude=(),
        audience="change management, the service owner, and the customer operations team reviewing "
                 "the executed change.")

    # ---- 1. Execution summary ----
    doc.add_heading("1. Execution summary", level=1)
    checks = prog["checks"]
    n_complete = sum(1 for w in waves if w["closeout"]["decision"] == "COMPLETE")
    n_rolled = sum(1 for w in waves if w["closeout"]["decision"] == "ROLLED BACK")
    doc.add_paragraph(
        f"{n_complete} of {len(waves)} wave(s) completed"
        + (f", {n_rolled} rolled back" if n_rolled else "")
        + f". {prog['n_steps_done']} of {prog['n_steps']} run-of-show step(s) completed"
        + (f" ({prog['n_steps_skipped']} skipped)" if prog['n_steps_skipped'] else "")
        + f"; validation: {checks['pass']} pass / {checks['fail']} fail / {checks['na']} n-a "
        f"/ {checks['pending']} not run; {prog['n_deviations']} deviation(s) logged."
    )
    table(["Metric", "Planned", "Actual"], [
        ["Waves", len(waves), f"{n_complete} complete · {n_rolled} rolled back"],
        ["Run-of-show steps", prog["n_steps"],
         f"{prog['n_steps_done']} done · {prog['n_steps_skipped']} skipped"],
        ["Validation checks", sum(len(w.get('checks') or []) for w in waves),
         f"{checks['pass']} pass · {checks['fail']} fail · {checks['na']} n/a"],
        ["Maintenance window", _fmt_min(prog["planned_window_minutes"]), _fmt_min(elapsed_min)],
        ["Deviations", 0, prog["n_deviations"]],
    ], widths=[2.2, 2.1, 2.2])

    # ---- 2. Planned vs actual, per wave ----
    doc.add_heading("2. Planned vs actual, per wave", level=1)
    rows = []
    for w in waves:
        actual = execution.wave_elapsed_minutes(w)
        n_pass = sum(1 for c in w["checks"] if c["result"] == "pass")
        n_fail = sum(1 for c in w["checks"] if c["result"] == "fail")
        rows.append([w["group"], w["gate"], _DECISION_LABEL.get(w["closeout"]["decision"], "—"),
                     w["est_window_label"], _fmt_min(actual) if actual is not None else "—",
                     f"{n_pass}✓ / {n_fail}✗ of {len(w['checks'])}"])
    table(["Wave", "Gate at plan", "Closeout", "Planned window", "Actual duration", "Validation"],
          rows, widths=[1.0, 1.1, 1.1, 1.1, 1.1, 1.1])

    # ---- 3. Per-wave as-executed log ----
    doc.add_heading("3. Per-wave as-executed log", level=1)
    for w in waves:
        doc.add_heading(f"Wave {w['order']} — {w['group']}", level=2)
        meta = doc.add_paragraph()
        r = meta.add_run("Closeout: ")
        r.bold = True
        r.font.color.rgb = ink(_NAVY)
        decision = w["closeout"]["decision"]
        meta.add_run(f"{_DECISION_LABEL.get(decision, '—')}"
                     + (f" at {_fmt_ts(w['closeout']['at'])}" if w["closeout"]["at"] else "")
                     + (f" by {w['closeout']['by']}" if w["closeout"]["by"] else "")
                     + (f" — {w['closeout']['note']}" if w["closeout"]["note"] else ""))

        table(["#", "Phase", "Step", "Status", "Time", "By / note"],
              [[i + 1, s["phase"],
                (s["action"][:180] + "…") if len(s["action"]) > 180 else s["action"],
                s["status"].upper() if s["status"] != "pending" else "—",
                _fmt_ts(s["at"]) if s["at"] else "—",
                " · ".join(x for x in (s["by"], s["note"]) if x) or "—"]
               for i, s in enumerate(w["steps"])],
              widths=[0.3, 1.1, 2.6, 0.7, 1.0, 0.8])

        if w["checks"]:
            doc.add_heading("Validation results", level=3)
            table(["Result", "Check", "Command", "Expected", "Observed"],
                  [[(c["result"].upper() if c["result"] != "pending" else "NOT RUN"),
                    c["check"], c["command"], c["expect"], c["observed"] or "—"]
                   for c in w["checks"]],
                  widths=[0.7, 1.6, 1.5, 1.5, 1.2])

    # ---- 4. Timeline / deviation log ----
    doc.add_heading("4. Timeline & deviation log", level=1)
    doc.add_paragraph("Every action recorded during the run, in order. Deviations (validation "
                      "failures, rollbacks, and explicitly logged deviations) are the inputs to the "
                      "review verdict below.")
    table(["Time", "Kind", "Wave", "Entry", "By"],
          [[_fmt_ts(e["at"]), e["kind"], e.get("wave") or "—", e["text"], e.get("by") or "—"]
           for e in events],
          widths=[1.1, 0.7, 0.8, 3.0, 0.9])

    # ---- 5. Review & sign-off ----
    doc.add_heading("5. Review & sign-off", level=1)
    doc.add_paragraph(
        "Per standard change-management practice, a post-implementation review concludes one of: the "
        "change met its objectives; it met them with deviations to feed back into the next wave; it "
        "was partially implemented; or it was backed out. This run's derived verdict is: "
    ).add_run(outcome).bold = True
    doc.add_paragraph("Lessons learned / actions for the next wave:")
    for _ in range(3):
        doc.add_paragraph("•  ")
    table(["Role", "Name", "Signature", "Date"], [
        ["Implementation lead", state.get("operator") or "", "", ""],
        ["Network engineer", "", "", ""],
        ["Change manager", "", "", ""],
        ["Customer acceptance", "", "", ""],
    ], widths=[1.8, 2.0, 1.7, 1.0])

    doc.save(output_path)
