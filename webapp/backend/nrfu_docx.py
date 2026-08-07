"""Render a **Network Ready-For-Use (NRFU) / Acceptance Test Plan (ATP)** as a .docx deliverable.

This is an AssessHub synthesis (no engine writer), written in the web layer like ``cutover_docx`` and
matching the engine writers' python-docx conventions (Calibri body, navy headings, "Light Grid Accent
1" tables) so it reads as one family with the rest of the deliverable set.

It follows the standard Cisco NRFU/ATP shape: document-control + sign-off front matter, a two-part /
three-phase test structure, and a per-test row carrying an ID, the scope, the command to run, the
EXPECTED result, and a blank Result cell for the tester. The three phases map to Cisco's NRFU model:

  * **Phase I — Device readiness** (per-device standalone): inventory / software / lifecycle, built
    from ``lifecycle_risk.per_device`` and ``collection_completeness``.
  * **Phase II — Logical config & connectivity**: the engine's own ``validation_plan`` checks (FHRP,
    gateway, reachability, routing, STP, link) — each a ready-to-run command with an 'expect' baseline.
  * **Phase III — Service & traffic verification** (end-to-end): the detected services and application
    domains from ``service_map`` / ``application_intelligence`` / multicast.

Signature mirrors the engine writers — ``write_nrfu_docx(output_path, snap_dict, label)`` — so it slots
into ``deliverables.generate`` exactly like the others.
"""

from __future__ import annotations

from datetime import datetime
from typing import Any, Dict, List

from . import engine, summary
from .docx_style import GREY as _GREY
from .docx_style import NAVY as _NAVY
from .docx_style import add_table, ink, kv, new_document

# List coercion is `summary._as_list` (the web layer's single copy, as `cutover.py` uses it); the dict
# twin lives here for the same reason it lives in `cutover.py` -- see `cisco_toolkit.docmeta.as_dict`
# for the engine-side equivalent the writer family shares.


def _as_dict(v: Any) -> Dict[str, Any]:
    """Coerce a snapshot section/sub-section to a dict. `... or {}` only guards FALSY values; a truthy
    NON-dict (an int/str/list in a malformed or hostile upload) survives it and raises on the very next
    `.get`. This writer is reached by a plain `GET /api/snapshots/{id}/deliverable/nrfu` over a snapshot
    that the upload path stores verbatim, so that raise escapes as an HTTP 500 on every later read -- a
    STORED denial of service. Returning {} degrades gracefully; for a well-formed value (dict or None)
    the result is identical to the `or {}` it replaces."""
    return v if isinstance(v, dict) else {}


def write_nrfu_docx(output_path: str, snap_dict: Dict[str, Any], label: str) -> None:
    """Write the NRFU / Acceptance Test Plan to ``output_path`` as a .docx."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    from cisco_toolkit.textutils import xml_safe, xml_safe_deep

    # Deep-sanitize ALL device-derived text up front, exactly like the engine writers (design/crd/
    # archreview/engagement). `add_table` sanitizes every cell, but this writer also emits device text on
    # DIRECT paragraph paths that bypass the table helper (the §2 "Phase II tests by category" join reads
    # `validation_plan` categories straight from the snapshot), and one U+FFFE/U+FFFF/lone surrogate/C0
    # control char there still aborts `doc.save()` with "All strings must be XML compatible" — a 500 that
    # repeats on every later read of a snapshot the upload path stores verbatim (stored DoS, see docx_style).
    snap_dict = xml_safe_deep(snap_dict if isinstance(snap_dict, dict) else {})
    label = xml_safe(label) if isinstance(label, str) else (str(label) if label is not None else "")

    doc = new_document()

    def table(headers: List[str], rows: List[List[Any]], widths: List[float] | None = None):
        return add_table(doc, headers, rows, widths)

    # _as_dict / summary._as_list at EVERY level, never `... or {}` / `or []`: the snapshot is stored
    # verbatim from an upload whose only validation is "is a dict and has 'devices'", so a truthy
    # non-dict/non-list here is attacker-reachable and would 500 this read route for good (see _as_dict).
    # Several of these do not detonate at the read but much later -- `devices` at the len(devices)
    # fallback AND at the eagerly-evaluated `coll.get('inventory', len(devices))` default; `coll`,
    # `scale` and `swrisk` at their .get()s in the document-control / scope-limit blocks below.
    devices = _as_dict(snap_dict.get("devices"))
    per_device = [d for d in summary._as_list(_as_dict(snap_dict.get("lifecycle_risk")).get("per_device"))
                  if isinstance(d, dict)]
    coll = _as_dict(_as_dict(snap_dict.get("collection_completeness")).get("summary"))
    # SSOT: the fleet-scale header reads the canonical executive_brief.scale (the published single
    # source the explorer/deck/HLD read), with len(devices) only as a pre-brief fallback (C9 fix —
    # the web-layer NRFU writer was the last surface recomputing fleet scale from the raw array).
    scale = _as_dict(_as_dict(snap_dict.get("executive_brief")).get("scale"))
    val_items = [i for i in summary._as_list(_as_dict(snap_dict.get("validation_plan")).get("items"))
                 if isinstance(i, dict)]
    services = [s for s in summary._as_list(_as_dict(snap_dict.get("service_map")).get("services"))
                if isinstance(s, dict)]
    domains = [d for d in summary._as_list(_as_dict(snap_dict.get("application_intelligence")).get("domains"))
               if isinstance(d, dict)]
    # NOT coerced on purpose: `mcast` is never dereferenced, only truth-tested (§2 test count and the
    # Phase III multicast row), so any value is safe and a guard here would be churn.
    mcast = (snap_dict.get("multicast_intelligence") or {})
    # design-decision coverage + scope limits (N29/N30): trace the ATP back to the target-state design
    bp_decisions = [d for d in summary._as_list(_as_dict(snap_dict.get("design_blueprint")).get("decisions"))
                    if isinstance(d, dict)]
    nrfu_items = [i for i in summary._as_list(_as_dict(snap_dict.get("design_nrfu")).get("items"))
                  if isinstance(i, dict)]
    swrisk = _as_dict(_as_dict(snap_dict.get("software_risk")).get("summary"))

    # ---- title page ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Network Ready-For-Use (NRFU)")
    tr.bold = True
    tr.font.size = Pt(26)
    tr.font.color.rgb = ink(_NAVY)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(f"Acceptance Test Plan · {label}")
    sr.font.size = Pt(13)
    sr.font.color.rgb = ink(_GREY)
    stamp = doc.add_paragraph()
    stamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st = stamp.add_run("DRAFT — test cases derived from the assessment snapshot; review and time each "
                       "case against your change-control before execution.")
    st.italic = True
    st.font.color.rgb = ink(_GREY)

    # ---- document control + sign-off ----
    doc.add_heading("Document control", level=1)
    gen = snap_dict.get("generated_at") or datetime.now().strftime("%Y-%m-%d %H:%M")
    table(["Field", "Value"], [
        ["Document", "Network Ready-For-Use / Acceptance Test Plan"],
        ["Subject", label],
        ["Version", "0.1 (DRAFT)"],
        ["Generated", datetime.now().strftime("%Y-%m-%d %H:%M")],
        ["Snapshot captured", gen],
        ["Engine", engine.ENGINE_SCHEMA_VERSION],
        ["Devices in scope", scale.get("n_devices") if isinstance(scale.get("n_devices"), int) else len(devices)],
    ], widths=[2.2, 4.3])
    doc.add_heading("Sign-off", level=2)
    table(["Role", "Name", "Signature", "Date"], [
        ["Test lead", "", "", ""],
        ["Network engineer", "", "", ""],
        ["Change manager", "", "", ""],
        ["Customer acceptance", "", "", ""],
    ], widths=[1.8, 2.0, 1.7, 1.0])
    # Shared family cross-reference (imported here, after the engine sys.path bootstrap).
    from cisco_toolkit.docmeta import add_related_documents
    add_related_documents(
        doc, exclude=("nrfu",),
        audience="the test lead and network engineers executing the acceptance tests, and the "
                 "customer signatories accepting the network for production.")

    # ---- 1. introduction & scope ----
    doc.add_heading("1. Introduction, scope & test approach", level=1)
    doc.add_paragraph(
        "This Network Ready-For-Use (NRFU) plan — also an Acceptance Test Plan (ATP) — verifies that the "
        "migrated fleet is ready for production. It is organised in two parts (per-device standalone "
        "checks and end-to-end service checks) across the three standard NRFU phases. Each test states "
        "the command to run and the EXPECTED result captured from the assessment baseline; a deviation "
        "is a regression to investigate before acceptance.")
    p = doc.add_paragraph()
    kv(p, "Entry criteria:", "the wave/site cutover is complete, target configuration is applied, and "
                             "management plane is reachable on every in-scope device.")
    p = doc.add_paragraph()
    kv(p, "Exit criteria:", "every High-severity test passes, zero Critical tests fail, and any "
                            "deviation is logged and dispositioned (accepted or remediated) before sign-off.")

    # ---- 2. test summary ----
    val_by_cat: Dict[str, int] = {}
    for it in val_items:
        c = str(it.get("category", "") or "—")
        val_by_cat[c] = val_by_cat.get(c, 0) + 1
    n_p1, n_p2 = len(per_device), len(val_items)
    # str(): an unhashable service/category (a list/dict in a malformed upload) would 500 the set build
    n_p3 = (len({(str(s.get("service")), str(s.get("category"))) for s in services})
            + len(domains) + (1 if mcast else 0))
    doc.add_heading("2. Test summary", level=1)
    table(["Phase", "Coverage", "Tests"], [
        ["Phase I — Device readiness", "Per-device inventory, software & lifecycle", n_p1],
        ["Phase II — Logical config & connectivity", "FHRP / gateway / reachability / routing / STP / link", n_p2],
        ["Phase III — Service & traffic", "Detected services, application domains, multicast/timing", n_p3],
        ["Total", "", n_p1 + n_p2 + n_p3],
    ], widths=[3.0, 2.7, 0.8])
    if val_by_cat:
        doc.add_paragraph("Phase II tests by category: "
                          + " · ".join(f"{k} {v}" for k, v in sorted(val_by_cat.items(), key=lambda kv: -kv[1])))

    # ---- 2.1 design-decision coverage & scope limits (N30 traceability + N29 coverage-honesty) ----
    rec = [d for d in bp_decisions if d.get("status") == "recommended"]
    needs = [d for d in bp_decisions if d.get("status") == "needs-requirement"]
    nc = coll.get("not_collected") or 0
    n_na = swrisk.get("n_config_not_assessable") or 0
    limits = []
    if nc:
        limits.append(f"{nc} device(s) were not collected at assessment — their post-cutover state is NOT "
                      "validated here; re-collect and run Phase I against them before sign-off.")
    if n_na:
        limits.append(f"{n_na} device(s) had configuration that could not be assessed for software-advisory "
                      "exposure — NOT validated here.")
    if needs:
        limits.append(f"{len(needs)} target-state design area(s) still need a requirement before they can be "
                      "designed and acceptance-tested — see the design blueprint's open questions.")
    # EX-nrfu-004: which RECOMMENDED decisions actually have a derived acceptance test. `phase_by` is keyed
    # on presence, so a decision absent from design_nrfu is UNCOVERED — it must never render as "Tested".
    # str() on both sides together (an unhashable decision_id/id in a malformed upload would 500 the dict
    # build and the lookup alike; stringifying only one side would stop well-formed lookups resolving).
    phase_by = {str(i.get("decision_id")): i.get("phase", "") for i in nrfu_items}
    uncovered = [d for d in rec if str(d.get("id")) not in phase_by]
    if uncovered:
        limits.append(f"{len(uncovered)} recommended design decision(s) have NO acceptance test derived in "
                      "the design NRFU — they are NOT validated by this plan; derive test cases (or record "
                      "an explicit waiver) before sign-off.")
    if rec or needs or limits:
        doc.add_heading("2.1 Design-decision coverage & scope limits", level=2)
        doc.add_paragraph(
            "Traceability back to the assessment's target-state design: a RECOMMENDED design decision reads "
            "'Tested' ONLY where the design NRFU actually derived an acceptance test for it, in the phase "
            "shown; a decision with no derived test, and one that still needs a requirement, are explicit "
            "coverage boundaries listed below — never a silent gap.")
        if rec or needs:
            # EX-nrfu-004: the coverage column is driven by ACTUAL PRESENCE in design_nrfu. It previously
            # printed "Tested — " + (phase or "post-cutover-functional") UNCONDITIONALLY, so a decision with
            # no design_nrfu item read as covered and a test lead signed the ATP believing it had test cases
            # it does not — the exact silent gap the paragraph above promises never happens. `phase_by` may
            # legitimately hold "" (an item that states no phase); that is still COVERED, just unphased, so
            # the branch keys on `in phase_by`, never on the phase's truthiness.
            def _coverage(d: Dict[str, Any]) -> str:
                key = str(d.get("id"))
                if key not in phase_by:
                    return "NOT COVERED — no acceptance test derived for this decision"
                return "Tested — " + (str(phase_by[key]).strip() or "phase not stated in the design NRFU")

            trace = [(d.get("title", d.get("id", "")), d.get("priority", ""), _coverage(d)) for d in rec]
            trace += [(d.get("title", d.get("id", "")), d.get("priority", ""),
                       "Not testable until the requirement is supplied") for d in needs]
            table(["Target-state design decision", "Priority", "NRFU coverage"], trace[:40],
                  widths=[3.4, 0.9, 2.2])
            if len(trace) > 40:
                doc.add_paragraph(f"… and {len(trace) - 40} more decision(s); full set in the design blueprint.")
        if limits:
            doc.add_paragraph("Scope limits — NOT validated in this NRFU:")
            for lm in limits:
                doc.add_paragraph(lm, style="List Bullet")

    # ---- 3. Phase I — device readiness ----
    doc.add_heading("3. Phase I — Device readiness (per-device standalone)", level=1)
    if coll:
        doc.add_paragraph(
            f"Collection completeness at assessment: {coll.get('complete', 0)} of "
            f"{coll.get('inventory', len(devices))} device(s) fully collected, "
            f"{coll.get('partial', 0)} partial, {coll.get('not_collected', 0)} not collected. "
            "Re-run the inventory check below on every device post-cutover.")
    rows1: List[List[Any]] = []
    for n, d in enumerate(per_device, start=1):
        host = d.get("host", "")
        model = d.get("model", "") or "—"
        ver = d.get("sw_version", "") or "—"
        band = d.get("band", "")
        if band == "Past-LDoS":
            note = "Past last-day-of-support — replacement, not just verification"
        elif band == "Past-EoS":
            note = ("Past end-of-sale with LDoS still future — plan refresh; this date band does not "
                    "establish support entitlement")
        elif band == "Near-LDoS":
            note = "Within one year of LDoS — schedule refresh before the recorded deadline"
        elif band == "Active":
            note = "Pre-EoS date band (schema: Active); support entitlement not assessed"
        elif band == "Unknown":
            note = ("NOT ASSESSED — either no exact EoX row matched or the matched row's retained "
                    "source/date authority was withheld or incomplete")
        else:
            note = band or ""
        rows1.append([f"NRFU-I-{n:03d}", host,
                      "Device reachable; model & software match the as-built baseline",
                      "show version | i Model number|Version|System image",
                      f"Model {model}; IOS {ver}", "", note])
    if rows1:
        table(["ID", "Device", "Check", "Command", "Expected", "Result", "Notes"], rows1,
              widths=[0.7, 0.8, 1.6, 1.7, 1.4, 0.5, 1.0])
    else:
        doc.add_paragraph("No per-device lifecycle data in this snapshot.")

    # ---- 4. Phase II — logical config & connectivity ----
    doc.add_heading("4. Phase II — Logical configuration & connectivity", level=1)
    doc.add_paragraph("These are the engine's own post-cutover validation checks. Run each and compare "
                      "to the Expected baseline captured pre-cutover.")
    rows2: List[List[Any]] = []
    for n, it in enumerate(val_items, start=1):
        rows2.append([f"NRFU-II-{n:03d}", it.get("device", ""), it.get("category", ""),
                      it.get("severity", ""), it.get("check", ""), it.get("command", ""),
                      it.get("expect", ""), ""])
    if rows2:
        table(["ID", "Device", "Category", "Sev", "Check", "Command", "Expected", "Result"], rows2,
              widths=[0.7, 0.7, 0.8, 0.5, 1.3, 1.2, 1.6, 0.5])
    else:
        doc.add_paragraph("No validation_plan in this snapshot.")

    # ---- 5. Phase III — service & traffic ----
    doc.add_heading("5. Phase III — Service & traffic verification (end-to-end)", level=1)
    rows3: List[List[Any]] = []
    seen: set = set()
    n3 = 0
    for s in services:
        # str(): an unhashable service/category/port would 500 both the `in seen` test and seen.add()
        key = (str(s.get("service")), str(s.get("category")), str(s.get("port")))
        if key in seen:
            continue
        seen.add(key)
        n3 += 1
        svc = s.get("service", "") or "service"
        port = s.get("port", "")
        proto = s.get("proto", "")
        rows3.append([f"NRFU-III-{n3:03d}", f"{svc} ({s.get('category', '') or '—'})",
                      f"{svc} reachable end-to-end on {proto}/{port}",
                      f"telnet <server> {port}  /  show ip access-list (verify {svc} permitted)",
                      f"{svc} session establishes; permitted by ACL", ""])
    for d in domains:
        n3 += 1
        dom = d.get("domain", "") or d.get("id", "")
        # _as_list, not `or []`: a well-formed domains list whose ROW carries a scalar `switches`
        # (a truthy non-list) survives the row's isinstance(dict) filter and 500s the len() below.
        sw = summary._as_list(d.get("switches"))
        rows3.append([f"NRFU-III-{n3:03d}", f"App domain: {dom}",
                      "Intra-domain reachability after cutover",
                      "ping / traceroute between two endpoints in the domain",
                      f"Endpoints across the {len(sw)} switch(es) in this domain reach each other", ""])
    if mcast:
        n3 += 1
        rows3.append([f"NRFU-III-{n3:03d}", "Multicast / timing",
                      "Multicast forwarding & clock health (if in use)",
                      "show ip mroute / show ip pim neighbor / show ptp clock",
                      "PIM neighbors up; mroutes present; clock locked (if PTP/NTP in use)", ""])
    if rows3:
        table(["ID", "Scope", "Check", "Command", "Expected", "Result"], rows3,
              widths=[0.7, 1.5, 1.5, 1.6, 1.5, 0.5])
    else:
        doc.add_paragraph("No service-map / application-domain data in this snapshot.")

    # ---- 6. results & acceptance ----
    doc.add_heading("6. Results summary & acceptance", level=1)
    doc.add_paragraph("Record the outcome of each phase. The network is accepted for production only "
                      "when the exit criteria in §1 are met.")
    table(["Phase", "Total", "Pass", "Fail", "Deviations / notes"], [
        ["Phase I — Device readiness", n_p1, "", "", ""],
        ["Phase II — Logical config & connectivity", n_p2, "", "", ""],
        ["Phase III — Service & traffic", n_p3, "", "", ""],
        ["Overall", n_p1 + n_p2 + n_p3, "", "", ""],
    ], widths=[3.0, 0.7, 0.7, 0.7, 1.7])
    acc = doc.add_paragraph()
    kv(acc, "Acceptance:", "The undersigned confirm the exit criteria are met and accept the network "
                           "for production use. (See the Sign-off table on page 1.)")

    doc.save(output_path)
