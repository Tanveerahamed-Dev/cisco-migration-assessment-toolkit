"""Generate registry-declared snapshot deliverables for download from AssessHub.

Engine-backed entries reuse the engine writer. Cutover and NRFU are explicitly AssessHub-owned
syntheses; the conditional PIR has a separate execution-report route and is intentionally absent
from this snapshot catalogue. Optional dependencies remain fail-soft: an unavailable renderer is
reported as unavailable instead of becoming a generation failure.
"""

from __future__ import annotations

import importlib
import importlib.util
import logging
import os
import re
import tempfile
from collections.abc import Mapping
from typing import Dict

from . import engine  # noqa: F401  (import bootstraps sys.path so cisco_toolkit.* resolves)
from cisco_toolkit.docmeta import (  # noqa: E402  (engine establishes package path in portable mode)
    ASSESSHUB_DOWNLOAD_SPECS,
    ArtifactSpec,
)

logger = logging.getLogger(__name__)

_GATE_NOT_SUPPLIED = object()
_DISCLOSED_GATE_STATUSES = frozenset(
    {"bad_root", "unreadable", "ownership_mismatch", "ownership_unbound", "pending"}
)
_GATED_DOCUMENT_KINDS = frozenset({"design", "mop"})
_APPROVAL_KEY = re.compile(r"^[a-z][a-z0-9_]{0,63}$")


def _document_gate_projection(disclosure: object) -> dict[str, object] | None:
    """Return the bounded gate facts safe to place inside a forwarded document.

    Filesystem paths, engagement identifiers and free-form ledger/error text are deliberately
    excluded.  The HTTP header and DOCX are projections of this one immutable read.
    """
    if not isinstance(disclosure, Mapping):
        return None
    status = str(disclosure.get("status") or "")
    if status not in _DISCLOSED_GATE_STATUSES:
        return None

    def keys(name: str) -> list[str]:
        value = disclosure.get(name)
        if not isinstance(value, list):
            return []
        return sorted(
            {str(item) for item in value if isinstance(item, str) and _APPROVAL_KEY.fullmatch(item)}
        )

    return {"status": status, "missing": keys("missing"), "revoked": keys("revoked")}


def _stamp_unapproved_draft(path: str, kind: str, disclosure: object) -> None:
    """Embed an unsatisfied AssessHub gate in visible DOCX text and core metadata."""
    gate = _document_gate_projection(disclosure)
    if gate is None or kind not in {"design", "mop"}:
        return

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    status = str(gate["status"]).upper()
    missing = ", ".join(gate["missing"]) if gate["missing"] else "UNCONFIRMED"
    revoked = ", ".join(gate["revoked"]) if gate["revoked"] else "NONE RECORDED"
    marker = (
        f"UNAPPROVED DRAFT | gate status: {status} | missing approvals: {missing} | "
        f"revoked approvals: {revoked}"
    )

    doc = Document(path)
    banner = doc.paragraphs[0].insert_paragraph_before() if doc.paragraphs else doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = banner.add_run(marker)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(192, 0, 0)

    # Repeat the state in every section header so pages separated from the cover remain honest.
    for section in doc.sections:
        paragraph = section.header.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = paragraph.add_run(f"UNAPPROVED DRAFT | {status}")
        header_run.bold = True
        header_run.font.color.rgb = RGBColor(192, 0, 0)

    # Replace the generic draft row in the existing Document Control table when present.
    replaced = False
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) >= 2 and row.cells[0].text.strip().casefold() == "status":
                row.cells[1].text = marker
                replaced = True
                break
        if replaced:
            break

    # Standard OOXML core properties are independently readable from docProps/core.xml.
    doc.core_properties.content_status = "UNAPPROVED DRAFT"
    doc.core_properties.keywords = (
        f"atlas_gate_status={gate['status']};"
        f"missing={','.join(gate['missing'])};revoked={','.join(gate['revoked'])}"
    )[:255]
    doc.save(path)


def _reconcile_gate(snap: dict, kind: str) -> list:
    """Fail-soft SSOT pre-emission check (universal-best roadmap W3-5): never SILENTLY emit a deliverable from
    a snapshot whose published facts disagree with the raw evidence. Logs the violations (the snapshot's
    assessment_integrity already carries the machine-readable disclosure that the deliverables render); it never
    blocks the emit -- a single benign drift must not wedge the whole deliverable set. Total/fail-open."""
    try:
        from cisco_toolkit import ssot
        violations = ssot.reconcile(snap if isinstance(snap, dict) else {})
    except Exception:
        return []
    if violations:
        logger.warning("[SSOT] %s deliverable generated from a snapshot with %d unreconciled fact(s): %s",
                       kind, len(violations), violations[:3])
    return violations


def gate_disclosure(kind: str, gate_root: str = ".", engagement: str | None = None) -> dict | None:
    """The PPDIOO document-gate verdict for `kind`, as a DISCLOSURE — never a refusal.

    None means "nothing to say": an ungated deliverable, a legacy direct caller that supplied no
    engagement identity and has no store, or approvals verified for the declared engagement. The
    AssessHub route always supplies its campaign engagement ID, so no store becomes
    `ownership_unbound` rather than silent brownfield. A dict means the ledger says this document's
    upstream approvals are missing/revoked, or that gate state/ownership could not be determined;
    see `generate` for why AssessHub discloses where the CLI refuses.

    The status token is NOT computed here. `gate_state.pending_approvals` owns the gate-posture
    fact (SSOT Law 1) and is the read-only counterpart of the deciding path `enforce()` — same
    `STATUSES` vocabulary, never writes, never raises, which is exactly this surface's contract.
    Re-deriving the token locally is what inverted it: a LOCATED ledger with missing/revoked
    approvals was labelled `ungated`, a token `gate_state.UNEVALUATED` reserves for "the approvals
    were NEVER EVALUATED". So an AssessHub MOP download for an engagement whose LLD a peer review
    had actively REVOKED disclosed `X-Gate-Status: ungated` — the one word that reads as "no gates
    are tracked here, nothing to worry about" — while `enforce()` and `pending_approvals()` both
    called the same ledger `pending`. Three computations of one fact, and the third disagreed.

    Not `enforce()`, for two reasons that both still hold: it is the BLOCKING api, and its override
    arm APPENDS an audit line, which a read for a download header must never do. Total/fail-open
    like `_reconcile_gate` — a gate ledger problem must never be able to withhold a deliverable,
    which is the whole point of this mode."""
    if kind not in _GATED_DOCUMENT_KINDS:
        return None
    try:
        from cisco_toolkit import gate_state
        if kind not in gate_state.GENERATOR_REQUIRES:
            raise RuntimeError("gated document registry drift")
        posture = gate_state.pending_approvals(kind, gate_root, engagement=engagement)
        if posture["status"] in ("ungated", "clear"):
            # ungated occurs only for a caller that declared no engagement; AssessHub declares one.
            # clear here means the owner verified every required upstream approval.
            return None
        # Everything else is disclosed — `pending`, and also the remaining UNEVALUATED statuses
        # (`unreadable`, `bad_root`). Unknown ≠ absent: the CLI fails CLOSED there, we cannot, so we
        # say so loudly rather than let an operator read silence as "gates are fine". On those,
        # `missing == []` means UNKNOWN, not "nothing missing", which is why `status` leads the
        # header. `detail` mirrors `summary` for the pre-existing key shape.
        return {**posture, "detail": posture.get("summary", "")}
    except Exception:                         # noqa: BLE001 - a broken ledger never blocks a download
        # Do not interpolate the request-derived kind or an exception whose message can contain
        # ledger content.  Either can contain control characters and forge adjacent log records.
        logger.warning("[GATE UNREADABLE] gated document disclosure evaluation failed")
        return {
            "generator": kind,
            "status": "unreadable",
            "verified": False,
            "missing": [],
            "revoked": [],
            "summary": "gate state evaluation failed; approvals are unconfirmed",
            "detail": "gate state evaluation failed; approvals are unconfirmed",
        }


# Compatibility name retained for callers that used the old web-local dataclass. Membership,
# labels, media, dependencies and producers now all come from docmeta.ARTIFACT_SPECS.
Spec = ArtifactSpec
SPECS: Dict[str, ArtifactSpec] = {spec.key: spec for spec in ASSESSHUB_DOWNLOAD_SPECS}
_LIB_MODULE = {spec.needs: spec.dependency_module for spec in ASSESSHUB_DOWNLOAD_SPECS
               if spec.needs and spec.dependency_module}


def _have(lib: str) -> bool:
    return importlib.util.find_spec(_LIB_MODULE[lib]) is not None


def availability() -> Dict[str, bool]:
    """Which deliverables can be generated on this server (depends on the optional libs)."""
    by_library = {lib: _have(lib) for lib in _LIB_MODULE}
    return {key: bool(spec.needs and by_library.get(spec.needs, False))
            for key, spec in SPECS.items()}


def have_docx() -> bool:
    """Whether python-docx is installed — for docx producers outside SPECS (e.g. the PIR)."""
    return _have("python-docx")


def resolve_writer(spec: ArtifactSpec):
    """Resolve a registry-owned writer through the backend package that is actually running.

    AssessHub supports both ``backend.*`` (the documented Uvicorn/dev entry) and
    ``webapp.backend.*`` (the installed/production entry). Importing the registry's absolute
    ``webapp.backend.*`` owner while ``backend.*`` is live creates a second module object, so
    patches and module state attach to a producer the route never calls. Keep the registry's
    production owner for packaging, but bind web-local writers to this module's package identity.
    """
    module_name = spec.writer_module or ""
    web_prefix = "webapp.backend."
    if module_name.startswith(web_prefix):
        module_name = __package__ + "." + module_name.removeprefix(web_prefix)
    module = importlib.import_module(module_name)
    return getattr(module, spec.writer_name or "")


def catalogue() -> list:
    """Canonical snapshot-download metadata plus live renderer availability, for the UI."""
    avail = availability()
    return [{"key": s.key, "label": s.label, "ext": s.ext, "available": avail[s.key],
             "producer": s.producer, "engine_cli_member": s.cli_produced, "stage": s.stage}
            for s in SPECS.values()]


def generate(kind: str, snap: dict, label: str, *, gates: dict | None = None,
             gate_root: str = ".", protocol_assurance_bundle: dict | None = None,
             document_gate: object = _GATE_NOT_SUPPLIED,
             gate_engagement: str | None = None) -> str:
    """Write the deliverable to a temp file; return its path. Caller streams it and deletes it.
    `gates` is the campaign's recorded gate sign-offs — consumed only by the engagement plan of
    record (its §4.3 as-signed trail); every other writer is a pure snapshot read.

    **DECISION (2026-07-22): the PPDIOO document gates DISCLOSE here; they do not block.**

    `design` and `mop` are gated deliverables (P0-3/DEC-003), and this path is not a re-render of
    something already approved — it calls the same writers the CLI does, against a snapshot whose
    engine version and detector set may have moved since anything was peer-reviewed. So the gates
    are RELEVANT here; the earlier reading that this is a mere "additive-compatibility surface"
    covered only half the truth. The valid half survives: `write_design_doc_docx` /
    `write_mop_docx` themselves stay ungated (they are pure writers, and ~62 direct-call tests in
    tests/test_design.py + tests/test_mop.py depend on that). Gating a CALLER — which is exactly
    what the CLI's `gate_enforce` is — costs those tests nothing.

    Why disclose rather than refuse, unlike the CLI and unlike `ingest.run_redaction_folder`:
    `gate_state` resolves ONE `docs/engagement-state.json` from a directory root, but an AssessHub
    campaign is a DB row with no filesystem home (`storage.py`: campaigns(id, name, description,
    created_at)). A server-root ledger would therefore govern EVERY campaign — engagement A's
    revoked LLD would withhold campaign B's MOP, a refusal that is simply wrong about B. There is
    also no `--override-gate` on this surface, and Atlas makes AssessHub the one door in the
    field, so a wrongly-refused engineer has no recourse at all. A control with a wrong-refusal
    mode and no override gets routed around, which costs more enforcement than it buys. Disclosure
    has no wrong-refusal mode, and it matches this subsystem's own doctrine: `gates.py`
    (`annotate_out_of_order`) already DISCLOSES an out-of-order sign-off rather than blocking it,
    and `_reconcile_gate` above is fail-soft for the same reason.

    Not-silent is the part that changed. "The CLI is the enforcement point" was defensible when
    the browser was a developer convenience; ADR-0004 made AssessHub the primary field interface,
    where that reads as "the enforcement point is the door nobody uses" — an engineer refused at
    the CLI can click Download and invoke the same writer. That bypass is now visible in the
    log and on the response (`X-Gate-Status`, set by the route) instead of being invisible.

    The disclosure now travels inside Design/MOP DOCX bytes as `UNAPPROVED DRAFT` plus bounded
    machine-readable core metadata. This does not establish campaign ownership: a per-campaign
    document-gate axis still needs a schema migration and precedence rule against the filesystem
    ledger. The in-document state therefore never claims approval.
    """
    spec = SPECS[kind]
    _reconcile_gate(snap, kind)   # W3-5: loudly flag a drifting snapshot before emit (fail-soft, never blocks)
    disclosure = (
        gate_disclosure(kind, gate_root, engagement=gate_engagement)
        if document_gate is _GATE_NOT_SUPPLIED
        else document_gate
    )
    if disclosure:
        logger.warning("[GATE %s] %s generated from AssessHub with unsatisfied document gates: %s "
                       "-- disclosed, not blocked (see deliverables.generate docstring)",
                       disclosure["status"].upper(), kind, disclosure)
    # Producer dispatch is registry-owned too. validate_artifact_registry refuses any listed
    # download without a module/function, so catalogue membership cannot become an orphan.
    write = resolve_writer(spec)

    fd, path = tempfile.mkstemp(suffix="." + spec.ext, prefix=f"assesshub_{kind}_")
    os.close(fd)
    try:
        if kind == "engagement":
            # The writer treats None and {} identically (its own gr filter), so no `and gates`
            # second branch — one call shape per kind (V3.23.159 review simplification).
            write(path, snap, label, gate_record=gates)
        elif kind in {"runbook", "mop", "nrfu"}:
            # The writers are pure renderers and cannot establish exact-byte custody from a parsed
            # snapshot. AssessHub supplies the separately built portfolio-owner sidecar; direct/CLI
            # calls leave it absent and the documents state NOT VERIFIED rather than minting one.
            write(
                path,
                snap,
                label,
                protocol_assurance_bundle=protocol_assurance_bundle,
            )
        else:
            write(path, snap, label)
        _stamp_unapproved_draft(path, kind, disclosure)
    except Exception:
        if os.path.exists(path):
            os.unlink(path)
        raise
    return path
